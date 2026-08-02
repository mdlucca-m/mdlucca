# -*- coding: utf-8 -*-
"""Gera o DOCUMENTO COMPLETO em Word (.docx) a partir do documento HTML já
construído (/tmp/documento.html) — mesma fonte de verdade, com TODAS as seções,
tabelas e figuras (imagens embutidas), além de capa e sumário navegável (campo
TOC do Word, atualizável). É a versão Word do "Documento completo".

Uso: python build_documento.py  (gera /tmp/documento.html) e depois
     python build_documento_word.py  → Documento_completo_BRUMS_HIIT.docx
Dependências: beautifulsoup4 python-docx Pillow
"""
import re, base64, io, os
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

HERE=os.path.dirname(os.path.abspath(__file__))
SRC='/tmp/documento.html'
OUT=os.path.join(HERE,'Documento_completo_BRUMS_HIIT.docx')
INK=RGBColor(0x16,0x27,0x3d); TEAL=RGBColor(0x0e,0x8c,0x86); BLUE=RGBColor(0x24,0x5c,0x8b); MUT=RGBColor(0x5b,0x6b,0x82)
CONTENT_W=6.3  # polegadas úteis (A4, margens ~1.2cm/1.2cm → ~6.9; usamos 6.3 por segurança)

soup=BeautifulSoup(open(SRC,encoding='utf-8').read(),'html.parser')
doc=Document()
# margens A4
sec=doc.sections[0]
from docx.shared import Cm
sec.page_height=Cm(29.7); sec.page_width=Cm(21.0)
sec.top_margin=Cm(1.8); sec.bottom_margin=Cm(1.8); sec.left_margin=Cm(1.8); sec.right_margin=Cm(1.8)
# estilo base
st=doc.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(10.5); st.font.color.rgb=INK
for hi,(sz,col) in {'Heading 1':(17,RGBColor(0x12,0x24,0x38)),'Heading 2':(13,BLUE),'Heading 3':(11.5,BLUE)}.items():
    s=doc.styles[hi]; s.font.name='Cambria'; s.font.size=Pt(sz); s.font.bold=True; s.font.color.rgb=col

def add_runs(p, node):
    """Adiciona runs de um nó, preservando <b>/<strong>/<i>/<em>/<br>."""
    for c in node.children:
        if isinstance(c,NavigableString):
            t=str(c)
            if t.strip() or t==' ':
                p.add_run(re.sub(r'\s+',' ',t))
        elif isinstance(c,Tag):
            if c.name=='br': p.add_run().add_break()
            elif c.name in ('b','strong'):
                r=p.add_run(re.sub(r'\s+',' ',c.get_text())); r.bold=True
            elif c.name in ('i','em'):
                r=p.add_run(re.sub(r'\s+',' ',c.get_text())); r.italic=True
            elif c.name=='span':
                add_runs(p,c)
            else:
                p.add_run(re.sub(r'\s+',' ',c.get_text()))

def para(node, style=None):
    p=doc.add_paragraph(style=style); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    add_runs(p,node); return p

def set_cell_bg(cell,hexc):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),hexc); tcPr.append(sh)

def add_table(tbl):
    heads=[th.get_text(strip=True) for th in tbl.select('thead th')] or \
          [th.get_text(strip=True) for th in tbl.find_all('tr')[0].find_all(['th','td'])]
    body_rows=tbl.select('tbody tr') or tbl.find_all('tr')[1:]
    ncol=len(heads)
    t=doc.add_table(rows=1,cols=ncol); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    try: t.style='Light Grid Accent 1'
    except Exception: t.style='Table Grid'
    hcells=t.rows[0].cells
    for i,h in enumerate(heads):
        hcells[i].text=''; r=hcells[i].paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(9)
        hcells[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if i else WD_ALIGN_PARAGRAPH.LEFT
    for tr in body_rows:
        cells=tr.find_all(['td','th'])
        if not cells: continue
        row=t.add_row().cells
        for i in range(ncol):
            txt=cells[i].get_text(' ',strip=True) if i<len(cells) else ''
            row[i].text=''; rp=row[i].paragraphs[0]; rn=rp.add_run(txt); rn.font.size=Pt(9)
            rp.alignment=WD_ALIGN_PARAGRAPH.CENTER if i else WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

def add_figure(fig):
    img=fig.find('img')
    if img and img.get('src','').startswith('data:image'):
        b64=img['src'].split(',',1)[1]
        raw=base64.b64decode(b64); bio=io.BytesIO(raw)
        w,h=Image.open(io.BytesIO(raw)).size
        width_in=min(CONTENT_W, CONTENT_W)  # largura útil
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(bio,width=Inches(width_in))
    cap=fig.find('figcaption')
    if cap:
        cp=doc.add_paragraph(); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        add_runs(cp,cap)
        for r in cp.runs: r.italic=True; r.font.size=Pt(8.5); r.font.color.rgb=MUT

def eyebrow(text):
    p=doc.add_paragraph(); r=p.add_run(text.upper()); r.bold=True; r.font.size=Pt(8); r.font.color.rgb=TEAL
    p.paragraph_format.space_after=Pt(0)

# ---------------- CAPA ----------------
cov=soup.find('div',class_='cover')
if cov:
    eb=cov.find('div',class_='eyebrow');
    if eb: eyebrow(eb.get_text())
    h1=cov.find('h1')
    if h1:
        p=doc.add_paragraph(); r=p.add_run(h1.get_text(strip=True)); r.bold=True; r.font.name='Cambria'; r.font.size=Pt(26); r.font.color.rgb=RGBColor(0x12,0x24,0x38)
    sub=cov.find('p',class_='sub')
    if sub:
        p=doc.add_paragraph(); r=p.add_run(sub.get_text(strip=True)); r.font.size=Pt(12); r.font.color.rgb=MUT
    lead=cov.find('p',class_='lead')
    if lead: para(lead)
    # KPIs
    kpis=cov.select('.kpi')
    if kpis:
        line=' · '.join(f"{k.find('b').get_text(strip=True)} {k.find('span').get_text(strip=True)}" for k in kpis if k.find('b'))
        p=doc.add_paragraph(); r=p.add_run(line); r.bold=True; r.font.size=Pt(10.5); r.font.color.rgb=TEAL

# ---------------- SUMÁRIO (campo TOC do Word) ----------------
doc.add_paragraph()
p=doc.add_paragraph(); r=p.add_run('Sumário'); r.bold=True; r.font.name='Cambria'; r.font.size=Pt(15); r.font.color.rgb=RGBColor(0x12,0x24,0x38)
tp=doc.add_paragraph(); run=tp.add_run()
fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),r'TOC \o "1-2" \h \z \u')
it=OxmlElement('w:r'); itt=OxmlElement('w:t'); itt.text='Atualize o sumário no Word (clique com o botão direito → Atualizar campo).'; it.append(itt); fld.append(it)
tp._p.append(fld)
doc.add_page_break()

# ---------------- SEÇÕES ----------------
BLOCK={'p','table','figure','h2','h3','h4','div','ul','ol'}
def walk(node):
    for ch in node.children:
        if not isinstance(ch,Tag): continue
        cls=ch.get('class') or []
        if ch.name=='div' and ('toolbar' in cls or 'toc' in cls): continue
        if ch.name in ('script','button','style'): continue
        if ch.name=='div' and 'eyebrow' in cls: eyebrow(ch.get_text())
        elif ch.name=='h2': doc.add_heading(ch.get_text(strip=True),level=1)
        elif ch.name=='h3': doc.add_heading(ch.get_text(strip=True),level=2)
        elif ch.name=='h4': doc.add_heading(ch.get_text(strip=True),level=3)
        elif ch.name=='p':
            cl=ch.get('class') or []
            if 'foot' in cl:
                p=para(ch)
                for r in p.runs: r.font.size=Pt(8.5); r.font.color.rgb=MUT
            else: para(ch)
        elif ch.name=='table': add_table(ch)
        elif ch.name=='figure': add_figure(ch)
        elif ch.name in ('ul','ol'):
            for li in ch.find_all('li',recursive=False):
                para(li,style='List Bullet' if ch.name=='ul' else 'List Number')
        elif ch.name=='div':
            # containers (two/grid) — desce recursivamente
            walk(ch)

for section in soup.find_all('section'):
    # cada <section> vira um bloco; eyebrow+h2 já tratados dentro
    walk(section)

doc.save(OUT)
print('escrito', OUT, f'{os.path.getsize(OUT)//1024} KB', '| seções:', len(soup.find_all("section")))
