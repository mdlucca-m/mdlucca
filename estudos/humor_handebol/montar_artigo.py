# -*- coding: utf-8 -*-
import sys, json
import os
RAIZ=os.environ.get("HH_RAIZ") or os.path.dirname(os.path.abspath(__file__))
DADOS=os.path.join(RAIZ,"dados"); SAIDA=os.path.join(RAIZ,"saida")
os.makedirs(SAIDA, exist_ok=True)
S=RAIZ
sys.path.insert(0,os.path.join(RAIZ,"texto"))
import ET as A
from docx import Document
from docx.oxml.ns import qn as _qn_early
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc=Document()
_z=doc.settings.element.find(_qn_early('w:zoom'))
if _z is not None: _z.set(_qn_early('w:percent'),'100')
sec=doc.sections[0]
sec.page_width, sec.page_height = Cm(21), Cm(29.7)
for m in ('left_margin','right_margin'): setattr(sec,m,Cm(3))
sec.top_margin, sec.bottom_margin = Cm(3), Cm(2)

st=doc.styles['Normal']; st.font.name='Times New Roman'; st.font.size=Pt(12)
st._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')
pf=st.paragraph_format; pf.line_spacing=1.5; pf.space_after=Pt(0); pf.first_line_indent=Cm(1.25)
pf.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

def para(txt, indent=True, size=12, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         before=0, after=0, spacing=1.5):
    p=doc.add_paragraph(); r=p.add_run(txt); r.bold=bold; r.italic=italic; r.font.size=Pt(size)
    f=p.paragraph_format; f.alignment=align; f.line_spacing=spacing
    f.space_before=Pt(before); f.space_after=Pt(after)
    f.first_line_indent=Cm(1.25) if indent else Cm(0)
    return p

def head(txt, lvl=1):
    para(txt, indent=False, bold=True, size=12, align=WD_ALIGN_PARAGRAPH.LEFT,
         before=18 if lvl==1 else 12, after=6, spacing=1.5)

def caption(txt, before=12):
    para(txt, indent=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT, before=before, after=3, spacing=1.0)

def src(txt="Fonte: dados da pesquisa (2024).", nota=None):
    para(txt, indent=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT, before=3, after=0, spacing=1.0)
    if nota: para("Nota: "+nota, indent=False, size=9, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=6, spacing=1.0)

def borders(tbl, rows_top_bottom=True):
    tblPr=tbl._tbl.tblPr; bd=OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        el=OxmlElement(f'w:{edge}')
        if edge in ('top','bottom'): el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'8')
        else: el.set(qn('w:val'),'none')
        el.set(qn('w:color'),'000000'); bd.append(el)
    tblPr.insert_element_before(bd,'w:shd','w:tblLayout','w:tblCellMar','w:tblLook',
                                'w:tblCaption','w:tblDescription')

def mktable(headers, rows, widths=None, first_col_left=True, fs=9):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    borders(t)
    hdr=t.rows[0].cells
    for i,h in enumerate(headers):
        p=hdr[i].paragraphs[0]; p.paragraph_format.first_line_indent=Cm(0)
        p.paragraph_format.line_spacing=1.0; p.paragraph_format.space_after=Pt(2)
        p.alignment=WD_ALIGN_PARAGRAPH.LEFT if (i==0 and first_col_left) else WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(h); r.bold=True; r.font.size=Pt(fs)
    # linha sob o cabeçalho
    for c in hdr:
        pr=c._tc.get_or_add_tcPr(); bds=OxmlElement('w:tcBorders')
        b=OxmlElement('w:bottom'); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6'); b.set(qn('w:color'),'000000')
        bds.append(b)
        pr.insert_element_before(bds,'w:shd','w:noWrap','w:tcMar','w:textDirection',
                                 'w:tcFitText','w:vAlign','w:hideMark')
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            p=cells[i].paragraphs[0]; p.paragraph_format.first_line_indent=Cm(0)
            p.paragraph_format.line_spacing=1.0; p.paragraph_format.space_after=Pt(1)
            p.alignment=WD_ALIGN_PARAGRAPH.LEFT if (i==0 and first_col_left) else WD_ALIGN_PARAGRAPH.CENTER
            r=p.add_run(str(v)); r.font.size=Pt(fs)
            if str(v).startswith('__'): r.bold=True; r.text=str(v)[2:]
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Cm(w)
    return t

def figura(path, num, legenda, w=15.5):
    caption(f"Figura {num} – {legenda}")
    doc.add_picture(path, width=Cm(w))
    doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.first_line_indent=Cm(0)
    src("Fonte: elaborada pelos autores (2024).")


# ================= dados =================
import numpy as np
B=json.load(open(f"{DADOS}/U_base.json")); Q=json.load(open(f"{DADOS}/U_perfis.json"))
St=json.load(open(f"{DADOS}/U_stats.json")); S2=json.load(open(f"{DADOS}/U_stats2.json"))
PP=json.load(open(f"{DADOS}/U_prepos.json")); CT=json.load(open(f"{DADOS}/U_cont.json"))
B1=json.load(open(f"{DADOS}/U_brums1.json")); B2=json.load(open(f"{DADOS}/U_brums2.json"))
B3=json.load(open(f"{DADOS}/U_brums3.json")); E=json.load(open(f"{DADOS}/U_estimulo.json"))
NORMA=B['NORMA']
V7=['Tensão','Depressão','Raiva','Vigor','Fadiga','Confusão','TMD']
L7={'TMD':'PTH'}
def L(k): return L7.get(k,k)
PERF=['Iceberg','Superfície','Submerso','Barbatana de tubarão','Iceberg invertido','Everest invertido']

def n(x,d=2):
    if x is None: return "—"
    if isinstance(x,float) and x!=x: return "—"
    return f"{x:.{d}f}".replace('.',',').replace('-','−')
def pf(p,d=3):
    if p is None or (isinstance(p,float) and p!=p): return "—"
    return "< 0,001" if p<0.001 else f"{p:.{d}f}".replace('.',',')
def Tv(k,x):
    m,s=NORMA[k]; return (x-m)/s*10+50

# ================= corpo =================
para(A.TITULO, indent=False, bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, after=6, spacing=1.15)
para(A.SUB, indent=False, italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=18, spacing=1.15)

head("RESUMO")
para(A.RESUMO, indent=False, size=11, spacing=1.0, after=6)
para("Palavras-chave: "+A.PALAVRAS, indent=False, size=11, spacing=1.0, after=12)
head("ABSTRACT")
para(A.ABSTRACT, indent=False, size=11, italic=True, spacing=1.0, after=6)
para("Keywords: "+A.KEYWORDS, indent=False, size=11, italic=True, spacing=1.0, after=6)

head("1 INTRODUÇÃO")
for p in A.INTRO: para(p)

head("2 MÉTODO")
for i,(sub,ps) in enumerate(A.METODO):
    head(f"2.{i+1} {sub}", lvl=2)
    for p in ps: para(p)
    if i==5:
        figura(f"{SAIDA}/E1fig.png", 1,
               "Desenho do microciclo, janelas de coleta e cadeia de processamento das séries. "
               "O painel A apresenta a sequência de estímulos, a carga acumulada e os momentos de "
               "coleta; o painel B, as sete etapas do tratamento analítico e os dois vereditos que "
               "delas decorrem.", w=16.0)

head("3 RESULTADOS")

head("3.1 Descrição das variáveis do BRUMS", lvl=2)
for p in A.R1[:1]: para(p)
caption("Tabela 1 – Descrição robusta das sete variáveis do BRUMS nos 166 pares atleta-dia")
rows=[]
for k in V7:
    d=B1['DESC'][k]
    rows.append([L(k), f"{n(d['m'])} ({n(d['sd'])})", f"{n(d['lo'])}–{n(d['hi'])}",
                 f"{n(d['md'],1)} [{n(d['q1'],1)}–{n(d['q3'],1)}]", n(d['tm']), n(d['mad']),
                 f"{n(d['mn'],1)}–{n(d['mx'],1)}", n(d['cv'],1)])
mktable(["Variável","Média (DP)","IC 95%","Mediana [Q1–Q3]","Média aparada 20%","MAD","Mín–Máx","CV (%)"],
        rows, widths=[2.4,2.2,2.0,2.4,1.9,1.3,1.7,1.3], fs=8)
src(nota="MAD = desvio absoluto mediano multiplicado por 1,4826; CV = coeficiente de variação; "
         "PTH = perturbação total do humor. A média aparada descarta 20% das observações em cada cauda.")
for p in A.R1[1:3]: para(p)
caption("Tabela 2 – Forma da distribuição, efeito de piso, posição normativa e confiabilidade das medidas repetidas")
rows=[]
for k in V7:
    d=B1['DESC'][k]; ic=St['icc'][k]
    g=Q['grupoT'].get(k)
    tt=f"{n(g['m'],1)} ({n(g['ep'],2)})" if g else "—"
    rows.append([L(k), n(d['sk']), n(d['ku']), f"{n(d['W'],3)} ({pf(d['pW'])})", n(d['piso'],1),
                 tt, n(ic[0],3), n(ic[1],3), n(ic[2]), n(ic[3])])
mktable(["Variável","Assim.","Curt.","W de Shapiro-Wilk (p)","Piso (%)","Escore T (EP)",
         "CCI (1,1)","CCI (1,k)","EPM","MVD"],
        rows, widths=[2.1,1.1,1.1,2.7,1.2,1.9,1.3,1.3,1.0,1.0], fs=8)
src(nota="Piso = percentual de respostas no valor mínimo da escala; o critério de Terwee et al. (2007) "
         "considera problemática a concentração superior a 15%. CCI (1,1) refere-se à medida única e "
         "CCI (1,k), à média das medidas da semana, obtida pela fórmula de Spearman-Brown. EPM = "
         "erro-padrão de medida; MVD = mínima variação detectável. O escore T não se aplica à PTH.")
for p in A.R1[3:]: para(p)

head("3.2 Comportamento temporal e tendência ordenada", lvl=2)
for p in A.R2[:1]: para(p)
figura(f"{SAIDA}/E2fig.png", 2,
       "Trajetória das seis subescalas em escore T ao longo do microciclo e resultado do teste de "
       "tendência de Page. O sombreado de fundo identifica o estímulo de cada dia.", w=16.0)
for p in A.R2[1:]: para(p)
caption("Tabela 3 – Comparação global entre os sete dias, tendência ordenada e contraste entre a linha "
        "de base e a véspera da estreia")
rows=[]
for k in V7:
    fr=B1['FR'][k]; pg=B1['PAGE'][k]; re_=B1['RES'][k]
    rows.append([L(k), n(fr['chi']), pf(fr['p']), n(fr['W'],3),
                 n(pg['z']), pf(pg['p']), n(re_['d']), n(re_['r'],3), pf(re_['p']), n(re_['dz'],3)])
mktable(["Variável","χ² Friedman","p","W","z de Page","p","Δ D1→D7","r","p","d"],
        rows, widths=[2.1,1.7,1.2,1.1,1.3,1.2,1.3,1.0,1.2,1.0], fs=8)
src(nota="Friedman com gl = 6, restrito aos 19 atletas com registro completo nos sete dias; W = "
         "coeficiente de concordância de Kendall. Teste L de Page com alternativa ordenada de D1 a D7. "
         "Contraste D1–D7 pelo teste de Wilcoxon em 21 atletas; r = z/√n; d = diferença média dividida "
         "pelo desvio-padrão das diferenças.")

head("3.3 Sinal, ruído e suavização", lvl=2)
for p in A.R3[:1]: para(p)
figura(f"{SAIDA}/E3fig.png", 3,
       "Decomposição sinal–ruído das seis subescalas. Cada painel sobrepõe a série observada com o "
       "erro-padrão diário, a série suavizada pelo filtro binomial e a banda do piso de ruído em torno "
       "do valor basal.", w=16.0)
for p in A.R3[1:]: para(p)
caption("Tabela 4 – Piso de ruído, deslocamento total, veredito de sinal, transições de choque e pontos "
        "de inflexão")
rows=[]
for k in V7:
    d=S2['SER'][k]; dt=d['med'][6]-d['med'][0]
    ch=", ".join(f"D{c}→D{c+1}" for c in d['choque']) or "—"
    inf=", ".join(n(x) for x in d['infl']) or "—"
    rows.append([L(k), n(d['piso']), n(dt), n(abs(dt)/d['piso'],1),
                 "__SINAL" if d['sinal'] else "ruído", ch, inf])
mktable(["Variável","Piso de ruído","Δ D1→D7","|Δ| / piso","Veredito","Transições de choque",
         "Ponto de inflexão"],
        rows, widths=[2.2,1.8,1.5,1.4,1.5,3.6,2.0], fs=8)
src(nota="O piso de ruído é a média dos sete erros-padrão diários. Declara-se sinal quando |Δ| supera o "
         "piso. Transição de choque é aquela cuja primeira derivada, em valor absoluto, supera o piso. O "
         "ponto de inflexão é a abscissa, em fração de dia, na qual a segunda derivada muda de sinal. Os "
         "deslocamentos desta tabela derivam das médias diárias de todos os pares disponíveis e diferem "
         "por pequena margem dos contrastes pareados da Tabela 3, restritos aos 21 atletas com registro "
         "no primeiro e no sétimo dia.")
figura(f"{SAIDA}/E4fig.png", 4,
       "Primeira e segunda derivadas das séries diárias, expressas em unidades do piso de ruído de cada "
       "subescala. A moldura destaca as células cujo valor absoluto excede uma unidade de piso.", w=16.0)
for p in A.R4[1:]: para(p)

head("3.4 Prevalência diária dos seis perfis de humor", lvl=2)
for p in A.R5[:1]: para(p)
figura(f"{SAIDA}/E5fig.png", 5,
       "Prevalência diária dos seis perfis de humor, com série observada, série suavizada, banda do piso "
       "de ruído e transições de choque.", w=16.0)
for p in A.R5[1:]: para(p)
caption("Tabela 5 – Prevalência diária dos seis perfis e das três faixas de humor, veredito de sinal e "
        "teste de estabilidade")
SP=E['SER_PERF']
rows=[]
for nm in PERF+['Favorável','Neutra','De risco']:
    d=SP[nm]
    cq=S2['CQ'].get(nm) or (S2['CQ'].get('Faixa de risco') if nm=='De risco' else None)
    q=f"{n(cq['Q'])}" if cq else "—"; pq=pf(cq['p']) if cq else "—"
    ver=("não avaliável" if max(d['y'])<10 else ("__SINAL" if d['sinal'] else "ruído"))
    rows.append([nm]+[n(v,1) for v in d['y']]+[n(d['dtot'],1), n(d['piso'],1), ver, q, pq])
mktable(["Perfil ou faixa","D1","D2","D3","D4","D5","D6","D7","Δ","Piso","Veredito","Q","p"],
        rows, widths=[3.0,0.9,0.9,0.9,0.9,0.9,0.9,0.9,1.0,0.9,1.2,0.9,1.0], fs=7.5)
src(nota="Valores em percentual dos pares atleta-dia do dia; n por dia = 27, 26, 26, 21, 23, 22 e 21. "
         "Δ e piso em pontos percentuais. Q de Cochran com gl = 6, restrito aos 19 atletas com registro "
         "completo; a faixa favorável coincide com o perfil iceberg. O Everest invertido aparece como não "
         "avaliável porque envolve dois pares no conjunto inteiro e o piso binomial deixa de discriminar "
         "em prevalências próximas de zero.")

head("3.5 Resposta dos perfis e das variáveis aos diferentes estímulos", lvl=2)
for p in A.R6[:1]: para(p)
figura(f"{SAIDA}/E6fig.png", 6,
       "Distribuição dos seis perfis por tipo de estímulo e composição das três faixas de humor. As "
       "barras de erro representam o erro-padrão binomial.", w=16.0)
for p in A.R6[1:2]: para(p)
caption("Tabela 6 – Distribuição dos perfis e das faixas de humor por tipo de estímulo")
EST4=['Basal','HIIT','Amistoso','Técnico']
rows=[]
for nm in PERF:
    rows.append([nm]+[n(E['PREV_EST'][nm][e],1) for e in EST4])
rows.append(["__Faixa de humor","","","",""])
for f_ in ['Favorável','Neutra','Risco']:
    rows.append([f_]+[n(E['FAIXA_EST'][f_][e],1) for e in EST4])
mktable(["Perfil ou faixa"]+[f"{e} (n = {E['NPOR'][e]})" for e in EST4],
        rows, widths=[4.4,2.6,2.6,2.6,2.6], fs=8)
src(nota=f"Valores em percentual dos pares atleta-dia de cada tipo de estímulo. Associação entre "
         f"estímulo e perfil: χ² = {n(E['chi'])}; gl = {E['gl']}; p = {pf(E['p_chi'])}. Associação entre "
         f"estímulo e faixa: χ² = {n(E['chi_f'])}; gl = {E['gl_f']}; p = {pf(E['p_f'])}. O dia basal "
         "corresponde ao primeiro dia, com coleta única noturna, e o dia técnico ocorreu uma única vez.")
for p in A.R6[2:]: para(p)
caption("Tabela 7 – Nível médio diário e resposta aguda intradiária de cada variável por tipo de estímulo")
rows=[["__Nível médio do dia","","","","",""]]
for k in V7:
    d=S2['NIV'][k]
    rows.append([L(k), n(d['H']), n(d['A']), n(d['T']), n(d['r'],3), pf(d['p'])])
rows.append(["__Resposta aguda (pós − pré)","","","","",""])
for k in V7:
    a=S2['AG'][k]
    rows.append([L(k), f"{n(a['HIIT']['d'])} ({pf(a['HIIT']['p'])})",
                 f"{n(a['Amistoso']['d'])} ({pf(a['Amistoso']['p'])})",
                 f"{n(a['Técnico']['d'])} ({pf(a['Técnico']['p'])})",
                 f"{n(a['HIIT']['r'],2)} / {n(a['Amistoso']['r'],2)} / {n(a['Técnico']['r'],2)}", "—"])
mktable(["Variável","HIIT","Amistoso","Técnico/força","r","p"],
        rows, widths=[2.4,2.9,2.9,2.9,2.4,1.3], fs=8)
src(nota="No bloco superior, r e p referem-se ao teste de Friedman entre os três tipos de dia, em 25 "
         "atletas com registro nos três. No bloco inferior, o valor entre parênteses é o p do teste de "
         "Wilcoxon para a diferença entre a noite e a manhã naquele tipo de dia, e a coluna r reúne os "
         "três tamanhos de efeito na ordem HIIT, amistoso e técnico.")

head("3.6 Migração intradiária por tipo de estímulo", lvl=2)
for p in A.R7[:1]: para(p)
figura(f"{SAIDA}/E7fig.png", 7,
       "Prevalência dos seis perfis pela manhã e à noite em cada tipo de estímulo e teste de McNemar "
       "para a migração à faixa de risco.", w=16.0)
for p in A.R7[1:]: para(p)
caption("Tabela 8 – Migração intradiária para a faixa de risco por tipo de estímulo")
rows=[]
g=PP['MCN']; fx=PP['FAIXA']
rows.append(["Todos os dias","120",str(g['c']),str(g['b']),
             f"{n(fx['pre']['sfav'],1)} → {n(fx['pos']['sfav'],1)}",
             f"{n(fx['pre']['sris'],1)} → {n(fx['pos']['sris'],1)}",
             n(g['chi']), pf(g['p']), "—"])
for t in ['HIIT','Amistoso','Técnico']:
    m=E['MCN_EST'][t]; d=E['PREV_PP'][t]
    rows.append([t, str(m['n']), str(m['entra']), str(m['sai']),
                 f"{n(d['fav'][0],1)} → {n(d['fav'][1],1)}",
                 f"{n(d['ris'][0],1)} → {n(d['ris'][1],1)}",
                 n(m['chi']), pf(m['p']), pf(m['ph'])])
mktable(["Estímulo","n","Entram","Saem","Iceberg pré → pós (%)","Risco pré → pós (%)","χ²","p","p Holm"],
        rows, widths=[2.1,0.9,1.2,1.0,3.0,3.0,1.1,1.2,1.3], fs=8)
src(nota="Teste de McNemar com correção de continuidade e gl = 1. Correção de Holm aplicada às três "
         "comparações por estímulo. A linha do conjunto reúne os 120 pares completos de manhã e noite.")

head("3.7 Teste formal de cruzamento entre trajetórias", lvl=2)
for p in A.R8[:1]: para(p)
figura(f"{SAIDA}/E8fig.png", 8,
       "Teste formal de cruzamento. O painel A acompanha as três faixas de humor; o painel B apresenta a "
       "diferença entre a faixa favorável e a faixa de risco contra o limiar combinado; o painel C mostra "
       "a inversão estabelecida entre vigor e fadiga.", w=16.0)
for p in A.R8[1:]: para(p)

head("3.8 Estrutura de associação entre as variáveis", lvl=2)
for p in A.R9[:1]: para(p)
figura(f"{SAIDA}/E9fig.png", 9,
       "Estrutura de associação entre as variáveis do BRUMS: matriz de Spearman, decomposição entre "
       "atletas e dentro do atleta, persistência de defasagem um e precedência temporal.", w=16.0)
for p in A.R9[1:]: para(p)

head("4 DISCUSSÃO")
for i,(sub,ps) in enumerate(A.DISCUSSAO):
    head(f"4.{i+1} {sub}", lvl=2)
    for p in ps: para(p)

head("5 LIMITAÇÕES")
for p in A.LIMITACOES: para(p)

head("6 CONCLUSÃO")
for p in A.CONCLUSAO: para(p)

head("DECLARAÇÕES")
para("Aprovação ética: parecer CAAE [inserir número do CAAE]. Consentimento: todos os participantes "
     "assinaram termo de consentimento livre e esclarecido. Financiamento: [inserir]. Conflito de "
     "interesses: os autores declaram não haver conflito de interesses. Contribuição dos autores: "
     "[inserir]. Disponibilidade de dados: a base anonimizada e os roteiros de análise em Python podem "
     "ser disponibilizados mediante solicitação ao autor correspondente; os arquivos com identificação "
     "nominal permanecem sob guarda restrita e não são compartilhados.", indent=False, size=11,
     spacing=1.15)

head("REFERÊNCIAS")
for r in A.REFS:
    para(r, indent=False, size=11, spacing=1.0, after=6, align=WD_ALIGN_PARAGRAPH.LEFT)

out=f"{SAIDA}/ARTIGO_INOVACAO_PERFIS_HUMOR_HANDEBOL.docx"
doc.save(out); print("salvo:", out)
