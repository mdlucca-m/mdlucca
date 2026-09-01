# -*- coding: utf-8 -*-
import sys, json
import os
RAIZ=os.environ.get("HH_RAIZ") or os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
S=os.path.join(RAIZ,"saida")
DADOS=os.path.join(RAIZ,"dados")
sys.path.insert(0,os.path.join(RAIZ,"texto"))

from docx import Document
from docx.oxml.ns import qn as _qn_early
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc=Document()
_z=doc.settings.element.find(_qn_early('w:zoom'))
if _z is not None: _z.set(_qn_early('w:percent'),'100')
sec=doc.sections[0]
sec.page_width, sec.page_height = Cm(21), Cm(29.7)
for m in ('left_margin','right_margin'): setattr(sec,m,Cm(3))
sec.top_margin, sec.bottom_margin = Cm(3), Cm(2)

st=doc.styles['Normal']; st.font.name='Times New Roman'; st.font.size=Pt(12)
st._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')
pf=st.paragraph_format; pf.line_spacing=1.5; pf.space_after=Pt(0); pf.first_line_indent=Cm(1.25)
pf.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

_CONT={'tab':0,'fig':0,'quadro':0}
def tab():   _CONT['tab']+=1;   return _CONT['tab']
def fig():   _CONT['fig']+=1;   return _CONT['fig']
def quadro():_CONT['quadro']+=1;return _CONT['quadro']

def para(txt, indent=True, size=12, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         before=0, after=0, spacing=1.5):
    p=doc.add_paragraph(); r=p.add_run(txt); r.bold=bold; r.italic=italic; r.font.size=Pt(size)
    f=p.paragraph_format; f.alignment=align; f.line_spacing=spacing
    f.space_before=Pt(before); f.space_after=Pt(after)
    f.first_line_indent=Cm(1.25) if indent else Cm(0)
    return p

def head(txt, lvl=1):
    para(txt, indent=False, bold=True, size=12, align=WD_ALIGN_PARAGRAPH.LEFT,
         before=18 if lvl==1 else 12, after=6, spacing=1.5)

def caption(txt, before=12):
    para(txt, indent=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT, before=before, after=3, spacing=1.0)

def src(txt="Fonte: dados da pesquisa (2024).", nota=None):
    para(txt, indent=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT, before=3, after=0, spacing=1.0)
    if nota: para("Nota: "+nota, indent=False, size=9, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=6, spacing=1.0)

def borders(tbl, rows_top_bottom=True):
    tblPr=tbl._tbl.tblPr; bd=OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        el=OxmlElement(f'w:{edge}')
        if edge in ('top','bottom'): el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'8')
        else: el.set(qn('w:val'),'none')
        el.set(qn('w:color'),'000000'); bd.append(el)
    tblPr.insert_element_before(bd,'w:shd','w:tblLayout','w:tblCellMar','w:tblLook',
                                'w:tblCaption','w:tblDescription')

def mktable(headers, rows, widths=None, first_col_left=True, fs=9):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    borders(t)
    hdr=t.rows[0].cells
    for i,h in enumerate(headers):
        p=hdr[i].paragraphs[0]; p.paragraph_format.first_line_indent=Cm(0)
        p.paragraph_format.line_spacing=1.0; p.paragraph_format.space_after=Pt(2)
        p.alignment=WD_ALIGN_PARAGRAPH.LEFT if (i==0 and first_col_left) else WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(h); r.bold=True; r.font.size=Pt(fs)
    # linha sob o cabeçalho
    for c in hdr:
        pr=c._tc.get_or_add_tcPr(); bds=OxmlElement('w:tcBorders')
        b=OxmlElement('w:bottom'); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6'); b.set(qn('w:color'),'000000')
        bds.append(b)
        pr.insert_element_before(bds,'w:shd','w:noWrap','w:tcMar','w:textDirection',
                                 'w:tcFitText','w:vAlign','w:hideMark')
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            p=cells[i].paragraphs[0]; p.paragraph_format.first_line_indent=Cm(0)
            p.paragraph_format.line_spacing=1.0; p.paragraph_format.space_after=Pt(1)
            p.alignment=WD_ALIGN_PARAGRAPH.LEFT if (i==0 and first_col_left) else WD_ALIGN_PARAGRAPH.CENTER
            r=p.add_run(str(v)); r.font.size=Pt(fs)
            if str(v).startswith('__'): r.bold=True; r.text=str(v)[2:]
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Cm(w)
    return t

def figura(path, num, legenda, w=15.5):
    caption(f"Figura {num} – {legenda}")
    doc.add_picture(path, width=Cm(w))
    doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.first_line_indent=Cm(0)
    src("Fonte: elaborada pelos autores (2024).")

