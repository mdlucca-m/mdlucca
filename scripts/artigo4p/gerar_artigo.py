#!/usr/bin/env python3
"""Gera o artigo curto em padrão ABNT.

    python3 scripts/artigo4p/gerar_artigo.py -o data/ARTIGO_SENSIBILIDADE.docx

A4 retrato, margens de 3 cm à esquerda e no topo e 2 cm à direita e embaixo,
Times New Roman 12, recuo de 1,25 cm na primeira linha, tabelas com título
acima e fonte abaixo, figuras com legenda acima e fonte abaixo. A entrelinha do
corpo é de 1,0 para que o texto, as quatro tabelas e as três figuras caibam em
quatro páginas.
"""
from __future__ import annotations

import argparse
import re
import shutil
import sys
import zipfile
from pathlib import Path

import docx
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent))
import conteudo  # noqa: E402
import figuras  # noqa: E402

FONTE = "Times New Roman"
CORPO, MENOR, MIUDO = 12, 10, 8
ENTRELINHA = 1.0
PRETO = RGBColor(0, 0, 0)


def preparar(doc):
    s = doc.sections[0]
    s.orientation = WD_ORIENT.PORTRAIT
    s.page_width, s.page_height = Cm(21.0), Cm(29.7)
    s.left_margin, s.top_margin = Cm(3.0), Cm(3.0)
    s.right_margin, s.bottom_margin = Cm(2.0), Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = FONTE
    normal.font.size = Pt(CORPO)
    normal.font.color.rgb = PRETO
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONTE)

    for nivel in (1, 2):
        st = doc.styles[f"Heading {nivel}"]
        st.font.name = FONTE
        st.font.size = Pt(CORPO)
        st.font.bold = True
        st.font.all_caps = (nivel == 1)
        st.font.italic = False
        st.font.color.rgb = PRETO
        st.element.rPr.rFonts.set(qn("w:eastAsia"), FONTE)
        pf = st.paragraph_format
        pf.space_before = Pt(10 if nivel == 1 else 8)
        pf.space_after = Pt(4)
        pf.line_spacing = ENTRELINHA
        pf.keep_with_next = True


def paragrafo(doc, texto, *, recuo=True, tamanho=CORPO, espaco=ENTRELINHA,
              centro=False, negrito=False, italico=False, justificar=True,
              depois=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = espaco
    pf.space_after = Pt(depois)
    pf.first_line_indent = Cm(1.25) if recuo else Cm(0)
    p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if centro else
                   WD_ALIGN_PARAGRAPH.JUSTIFY if justificar else
                   WD_ALIGN_PARAGRAPH.LEFT)
    r = p.add_run(texto)
    r.font.name, r.font.size, r.font.color.rgb = FONTE, Pt(tamanho), PRETO
    r.bold, r.italic = negrito, italico
    return p


# O esquema OOXML exige esta ordem dentro de <w:tcBorders>.
ORDEM_BORDAS = ("top", "left", "bottom", "right", "insideH", "insideV")


def definir_bordas(celula, visiveis: set[str], cor="000000", sz=6):
    tcPr = celula._tc.get_or_add_tcPr()
    antigas = tcPr.find(qn("w:tcBorders"))
    if antigas is not None:
        tcPr.remove(antigas)
    marcas = tcPr.makeelement(qn("w:tcBorders"), {})
    for lado in ORDEM_BORDAS:
        if lado in visiveis:
            atributos = {qn("w:val"): "single", qn("w:sz"): str(sz),
                         qn("w:space"): "0", qn("w:color"): cor}
        else:
            atributos = {qn("w:val"): "none", qn("w:sz"): "0",
                         qn("w:space"): "0", qn("w:color"): "auto"}
        marcas.append(marcas.makeelement(qn(f"w:{lado}"), atributos))
    tcPr.append(marcas)


def inserir_tabela(doc, dados):
    """Apresentação tabular: sem bordas laterais, com traço no topo, sob o
    cabeçalho e no rodapé."""
    paragrafo(doc, f"Tabela {dados['numero']} - {dados['titulo']}",
              recuo=False, justificar=False, tamanho=MENOR, espaco=1.0,
              depois=3)

    n = len(dados["cabecalho"])
    corpo = 7.5 if n <= 6 else 6.5
    tab = doc.add_table(rows=1, cols=n)
    tab.alignment = WD_TABLE_ALIGNMENT.CENTER
    tab.autofit = True
    tab.style = "Table Grid"

    for i, texto in enumerate(dados["cabecalho"]):
        cel = tab.rows[0].cells[i]
        cel.text = ""
        p = cel.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(texto)
        r.bold = True
        r.font.name, r.font.size, r.font.color.rgb = FONTE, Pt(corpo), PRETO

    for linha in dados["linhas"]:
        celulas = tab.add_row().cells
        for i, texto in enumerate(linha):
            celulas[i].text = ""
            p = celulas[i].paragraphs[0]
            p.alignment = (WD_ALIGN_PARAGRAPH.LEFT if i == 0
                           else WD_ALIGN_PARAGRAPH.CENTER)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(texto)
            r.font.name, r.font.size, r.font.color.rgb = FONTE, Pt(corpo), PRETO

    ultima = len(tab.rows) - 1
    for i, linha in enumerate(tab.rows):
        if i == 0:
            visiveis = {"top", "bottom"}
        elif i == ultima:
            visiveis = {"bottom"}
        else:
            visiveis = set()
        for cel in linha.cells:
            definir_bordas(cel, visiveis)

    paragrafo(doc, conteudo.FONTE_TABELA, recuo=False, justificar=False,
              tamanho=MIUDO, espaco=1.0, depois=1)
    paragrafo(doc, dados["nota"], recuo=False, tamanho=MIUDO, espaco=1.0,
              depois=6)


def inserir_figura(doc, caminho: Path, legenda: str, largura_cm: float):
    paragrafo(doc, legenda, recuo=False, justificar=False, tamanho=MENOR,
              espaco=1.0, depois=3)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(caminho), width=Cm(largura_cm))
    paragrafo(doc, conteudo.FONTE_FIGURA, recuo=False, justificar=False,
              tamanho=MIUDO, espaco=1.0, depois=6)


def montar(saida: Path, dir_figuras: Path) -> dict:
    figuras.gerar_todas(dir_figuras)

    doc = docx.Document()
    preparar(doc)

    paragrafo(doc, conteudo.TITULO, recuo=False, centro=True, negrito=True,
              espaco=1.0, depois=3)
    paragrafo(doc, conteudo.SUBTITULO, recuo=False, centro=True, italico=True,
              tamanho=MENOR, espaco=1.0, depois=10)

    paragrafo(doc, "RESUMO", recuo=False, justificar=False, negrito=True,
              tamanho=MENOR, espaco=1.0, depois=2)
    paragrafo(doc, conteudo.RESUMO, recuo=False, tamanho=MENOR, espaco=1.0,
              depois=6)
    paragrafo(doc, conteudo.PALAVRAS, recuo=False, tamanho=MENOR, espaco=1.0,
              depois=6)

    n_tab = n_fig = n_par = 0
    for secao in conteudo.SECOES:
        doc.add_heading(secao["titulo"], level=1)
        for bloco in secao["blocos"]:
            tipo = bloco[0]
            if tipo == "p":
                paragrafo(doc, bloco[1], depois=4)
                n_par += 1
            elif tipo == "h2":
                doc.add_heading(bloco[1], level=2)
            elif tipo == "tab":
                inserir_tabela(doc, conteudo.TABELAS[bloco[1]])
                n_tab += 1
            elif tipo == "fig":
                inserir_figura(doc, dir_figuras / bloco[1], bloco[3], bloco[2])
                n_fig += 1
            else:
                raise ValueError(f"bloco desconhecido: {tipo}")

    saida.parent.mkdir(parents=True, exist_ok=True)
    doc.save(saida)
    _limpar_settings(saida)
    return {"seções": len(conteudo.SECOES), "parágrafos": n_par,
            "tabelas": n_tab, "figuras": n_fig}


def _limpar_settings(caminho: Path):
    """Remove o <w:zoom> sem atributo que o template do python-docx grava e
    que viola o esquema OOXML."""
    temp = caminho.with_suffix(".tmp.docx")
    with zipfile.ZipFile(caminho) as orig, \
            zipfile.ZipFile(temp, "w", zipfile.ZIP_DEFLATED) as novo:
        for item in orig.infolist():
            dados = orig.read(item.filename)
            if item.filename == "word/settings.xml":
                txt = dados.decode("utf-8")
                if "w:percent" not in txt:
                    txt = re.sub(r"<w:zoom\b[^>]*/>", "", txt)
                    txt = re.sub(r"<w:zoom\b[^>]*>.*?</w:zoom>", "", txt,
                                 flags=re.S)
                dados = txt.encode("utf-8")
            novo.writestr(item, dados)
    shutil.move(temp, caminho)


def main() -> int:
    ap = argparse.ArgumentParser(
        description=__doc__,
        formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("-o", "--saida", type=Path,
                    default=Path("data/ARTIGO_SENSIBILIDADE_HUMOR.docx"))
    ap.add_argument("--figuras", type=Path, default=Path("data/fig4p"))
    a = ap.parse_args()
    info = montar(a.saida, a.figuras)
    print(f"\ngerado: {a.saida}")
    for k, v in info.items():
        print(f"  {k}: {v}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
