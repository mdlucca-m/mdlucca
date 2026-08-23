# -*- coding: utf-8 -*-
# Carta de apresentação (didática) para o orientador: de onde vieram as análises e como foram feitas.
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEAL = RGBColor(0x0E, 0x8C, 0x7A)
INK  = RGBColor(0x22, 0x2A, 0x33)
GREY = RGBColor(0x55, 0x5F, 0x6B)

doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(2.2); sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.4); sec.right_margin = Cm(2.4)
st = doc.styles['Normal']; st.font.name = 'Calibri'; st.font.size = Pt(11.5); st.font.color.rgb = INK

def shade(p, color):
    pPr = p._p.get_or_add_pPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),color)
    pPr.append(sh)

def eyebrow(t):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(t); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = TEAL
    r.font.name = 'Consolas'
    return p

def title(t):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10)
    r = p.add_run(t); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = INK
    return p

def H(t):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(t); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = TEAL
    return p

def P(t, after=8):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.line_spacing = 1.28
    # suporta **negrito** simples
    parts = t.split('**'); bold = False
    for seg in parts:
        if seg:
            r = p.add_run(seg); r.bold = bold; r.font.size = Pt(11.5)
        bold = not bold
    return p

def bullet(lead, rest):
    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(lead); r.bold = True; r.font.size = Pt(11.5)
    r2 = p.add_run(rest); r2.font.size = Pt(11.5)
    return p

def note(t):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.3); p.paragraph_format.right_indent = Cm(0.3)
    shade(p, 'F0F7F5')
    r = p.add_run(t); r.font.size = Pt(10.5); r.font.color.rgb = GREY; r.italic = True
    return p

# ---------- Cabeçalho ----------
eyebrow('CARTA DE APRESENTAÇÃO  ·  MONITORAMENTO DO HUMOR')
title('De onde vieram as análises e como foram feitas')
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
r = p.add_run('Microciclo pré-competitivo · 27 atletas de handebol · 21 a 27 de abril de 2024')
r.font.size = Pt(11); r.font.color.rgb = GREY

# linha
hr = doc.add_paragraph(); hr.paragraph_format.space_before = Pt(4); hr.paragraph_format.space_after = Pt(10)
pPr = hr._p.get_or_add_pPr(); pb = OxmlElement('w:pBdr'); bot = OxmlElement('w:bottom')
bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'8'); bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),'0E8C7A')
pb.append(bot); pPr.append(pb)

# ---------- Saudação ----------
P('Professor(a), esta carta acompanha o painel interativo (dashboard) e serve como um guia rápido: '
  'explica **de onde vieram os dados**, **o que foi analisado** e **por que cada análise foi escolhida**. '
  'A ideia é que qualquer pessoa consiga entender o caminho completo — da coleta ao resultado — sem precisar '
  'ser especialista em estatística.')

# ---------- 1. Os dados ----------
H('1 · De onde vieram os dados')
P('Todos os números partem de coletas reais feitas com a equipe durante **uma semana de pré-temporada** '
  '(21 a 27/04/2024). Foram **27 atletas** e **456 respostas** de questionário. Cada atleta respondia até '
  'três vezes por dia (antes, no meio e depois do treino), o que permite ver o humor variando dentro do dia '
  'e ao longo da semana.')
bullet('BRUMS — ', 'o humor propriamente dito: vigor, fadiga, tensão, depressão, raiva e confusão (e a soma delas, a Perturbação Total do Humor).')
bullet('Epworth — ', 'sonolência diurna (o quanto o atleta está com sono ao longo do dia).')
bullet('PSS — ', 'estresse percebido (o quanto o atleta se sente sobrecarregado psicologicamente).')
bullet('HIIT (FC e PSE) — ', 'a carga interna das sessões de tiro: frequência cardíaca e esforço percebido, sessão por sessão.')
bullet('T-CAR — ', 'o teste de aptidão aeróbia que definiu a intensidade do HIIT (104% do pico de velocidade).')
note('Todos os dados foram tratados de forma anonimizada (cada atleta é um código A01–A27). Nenhum nome real foi '
     'usado nas análises nem no painel. A conferência dos bancos originais está registrada no documento de auditoria.')

# ---------- 2. Como foi organizado ----------
H('2 · Como os dados foram organizados')
P('Antes de qualquer análise, foi feita uma **auditoria dos bancos originais** para garantir que tudo batia. '
  'Nessa etapa três cuidados foram decisivos:')
bullet('O dia 1 é a linha de base. ', 'Em 21/04 os atletas responderam o questionário, mas fizeram apenas uma sessão técnico-tática à noite; o primeiro HIIT foi em 22/04. Isso foi corrigido para não superestimar os efeitos.')
bullet('A média de cada dia usa todos os momentos. ', 'Cada valor diário é a média do atleta no dia inteiro (não só o "antes do treino"), o que dá um retrato mais justo da semana.')
bullet('A semana tem 7 dias, não 8. ', 'Os rótulos de planilha ("21–28") eram apenas nomes de intervalo; a coleta real vai de 21 a 27/04.')

# ---------- 3. As análises e por quê ----------
H('3 · Quais análises foram feitas — e por que cada uma')
P('As escolhas estatísticas seguem o que a literatura recomenda para **amostras pequenas e medidas repetidas** '
  '(o mesmo atleta medido vários dias). Em linguagem simples:')
bullet('Testes não paramétricos (Friedman, Wilcoxon). ', 'Como são poucos atletas, não se assume que os dados sigam a "curva normal"; esses testes são mais seguros nesse cenário. Eles respondem: "o humor mudou de verdade ao longo da semana?".')
bullet('Tamanho de efeito (dz, W de Kendall). ', 'O valor de p diz se a mudança é real; o tamanho de efeito diz se ela é grande ou pequena — é o que realmente importa na prática esportiva.')
bullet('Correlações (Spearman). ', 'Para ver o que anda junto: por exemplo, sonolência sobe quando a fadiga sobe.')
bullet('Curvas, derivadas e inflexões. ', 'Em vez de olhar só o começo e o fim, as curvas mostram a velocidade da mudança e o dia exato em que uma variável vira (o ponto de inflexão) — daí os cruzamentos entre vigor e fadiga.')
bullet('Análise por tipo de dia. ', 'Separar dias de HIIT, de jogo amistoso e de recuperação para ver se o humor negativo se comporta diferente em cada um.')
bullet('Classificação de perfis de humor. ', 'Cada atleta-dia é enquadrado em um perfil clássico (iceberg, iceberg invertido, etc.) para enxergar quem está bem e quem está em alerta.')
bullet('Modelos de predição / IoT. ', 'Uma prova de conceito: usar o humor de hoje para sinalizar risco amanhã, como um sistema de monitoramento contínuo faria.')

# ---------- 4. O que os resultados dizem ----------
H('4 · A leitura honesta dos resultados')
P('O achado central é consistente: ao longo da semana o **vigor cai e a fadiga sobe** (o "iceberg" do humor se '
  'achata), a **sonolência aumenta** acompanhando a fadiga, mas o **estresse percebido permanece estável** e o '
  'humor negativo é até menor nos dias de jogo. Ou seja: o desgaste é sobretudo **físico-energético**, compatível '
  'com uma **sobrecarga funcional planejada** — e não com um quadro de sofrimento psicológico.')
note('Importante para a discussão: as análises de risco e predição são apresentadas como demonstração de '
     'viabilidade de triagem (rastreio), e não como associação causal comprovada com lesão. Os dados de lesão '
     'disponíveis são de um torneio posterior e não podem ser ligados diretamente a esta semana. Esse limite é '
     'declarado abertamente no painel.')

# ---------- 5. Como navegar ----------
H('5 · Como usar o painel')
P('O dashboard reúne tudo isto de forma visual e navegável. Sugestão de percurso: comece pela **Apresentação '
  'guiada**, passe pela **Trajetória** (onde é possível exibir todas as variáveis cruzando ao mesmo tempo e ver '
  'a tabela com os pontos de análise de cada uma), veja **Dias de HIIT × amistoso**, a **Sonolência e estresse**, '
  'os **Perfis de risco por atleta** e feche em **Resultados & Discussão**, que explica cada achado em texto '
  'didático.')

# fechamento
clо = doc.add_paragraph(); clо.paragraph_format.space_before = Pt(10)
r = clо.add_run('Fico à disposição para detalhar qualquer etapa — dos dados brutos aos testes estatísticos.')
r.font.size = Pt(11.5); r.italic = True; r.font.color.rgb = GREY

doc.save('/home/user/mdlucca/Artigos/Carta_Apresentacao_Professor.docx')
print('[docx: Carta_Apresentacao_Professor.docx]')
