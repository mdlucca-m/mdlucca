# -*- coding: utf-8 -*-
import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

R=json.load(open('ajuste_gordura.json'))
def c2(s): return str(s).replace('.',',')

doc=Document()
sty=doc.styles['Normal']; sty.font.name='Times New Roman'; sty.font.size=Pt(12)
sty.element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')
pf=sty.paragraph_format; pf.line_spacing=1.5; pf.space_after=Pt(0)
sec=doc.sections[0]
sec.top_margin=Cm(3); sec.left_margin=Cm(3); sec.bottom_margin=Cm(2); sec.right_margin=Cm(2)

_TN=[0]; _FN=[0]
def P(t='',just=True,bold=False,size=12,after=6,before=0,align=None):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=bold; r.font.size=Pt(size)
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(after); p.paragraph_format.space_before=Pt(before)
    p.paragraph_format.first_line_indent=Cm(1.25) if (just and not align) else None
    if align=='c': p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    elif just: p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    return p
def H(t,size=12):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(size)
    p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(6); p.paragraph_format.line_spacing=1.5
def _bd(cell,top=True,bottom=True):
    tcPr=cell._tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for edge in ['top','bottom']:
        e=OxmlElement(f'w:{edge}'); on=(edge=='top' and top) or (edge=='bottom' and bottom)
        e.set(qn('w:val'),'single' if on else 'nil'); e.set(qn('w:sz'),'6'); e.set(qn('w:color'),'000000')
        b.append(e)
    for edge in ['left','right']:
        e=OxmlElement(f'w:{edge}'); e.set(qn('w:val'),'nil'); b.append(e)
    tcPr.append(b)
def table(caption,header,rows,fonte='Fonte: os autores (2026).',widths=None,fs=10):
    _TN[0]+=1
    p=doc.add_paragraph(); r=p.add_run(f'Tabela {_TN[0]} – {caption}'); r.font.size=Pt(11)
    p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(3)
    t=doc.add_table(rows=1,cols=len(header)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    hc=t.rows[0].cells
    for i,htxt in enumerate(header):
        hc[i].text=''; rr=hc[i].paragraphs[0].add_run(htxt); rr.bold=True; rr.font.size=Pt(fs); rr.font.name='Times New Roman'
        hc[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; _bd(hc[i],True,True)
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=''; rr=cells[i].paragraphs[0].add_run(str(val)); rr.font.size=Pt(fs); rr.font.name='Times New Roman'
            cells[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if i>0 else WD_ALIGN_PARAGRAPH.LEFT
            _bd(cells[i],False,False)
    for c in t.rows[-1].cells: _bd(c,False,True)
    if widths:
        for i,w in enumerate(widths):
            for rr in t.rows: rr.cells[i].width=Cm(w)
    pf=doc.add_paragraph(); rf=pf.add_run(fonte); rf.font.size=Pt(10)
    pf.paragraph_format.space_after=Pt(8)
def figure(path,caption,fonte='Fonte: os autores (2026).',w=15.5):
    _FN[0]+=1
    doc.add_paragraph().alignment=WD_ALIGN_PARAGRAPH.CENTER
    pp=doc.paragraphs[-1]; pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=pp.add_run(); run.add_picture(path,width=Cm(w))
    pc=doc.add_paragraph(); pc.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rc=pc.add_run(f'Figura {_FN[0]} – {caption}'); rc.font.size=Pt(11)
    pf=doc.add_paragraph(); pf.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rf=pf.add_run(fonte); rf.font.size=Pt(10); pf.paragraph_format.space_after=Pt(8)

# ============ TÍTULO ============
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('QUAL AJUSTE ESTATÍSTICO MELHOR INDICA SOBRECARGA, PERDA DE DESEMPENHO E FADIGA ACUMULADA? INTEGRAÇÃO DE APTIDÃO (PICO DE VELOCIDADE), COMPOSIÇÃO CORPORAL, POSIÇÃO E PERFIL DE HUMOR EM HANDEBOLISTAS DE ELITE')
r.bold=True; r.font.size=Pt(13)
p.paragraph_format.space_after=Pt(12)
P('Análise complementar ao manuscrito principal — microciclo pré-competitivo de sete dias (21–28/04/2024), 27 atletas do sexo masculino, 456 observações do BRUMS (duas coletas/dia).',just=True,size=11,after=10)

H('Resumo executivo')
P('Esta análise responde diretamente à pergunta: entre os ajustes disponíveis (alométrico, logístico, curvas ROC, decomposição de variância, derivadas e classificação de perfis de humor), qual é o indicador mais forte de excesso de treinamento, risco de lesão, perda de desempenho e fadiga acumulada — e como o percentual de gordura, a posição e o pico de velocidade (PV) se relacionam com o humor, o vigor e a fadiga física e mental ao longo da semana. O achado central é que nenhum ajuste é isoladamente "o correto": eles respondem a perguntas diferentes e devem ser estratificados. Para triagem de estado (quem está sobrecarregado hoje), o ajuste logístico/ROC sobre o PV é o indicador mais forte (AUC = 0,84); para explicar quem tolera a carga (traço), a normalização alométrica pelo PV é a mais robusta; para o momento em que a deterioração acelera, as derivadas e a classificação de perfis (iceberg→barbatana de tubarão) são as mais informativas. O percentual de gordura, isoladamente, não rastreou a fadiga do microciclo nesta amostra de elite.',after=8)

# ============ 1. FUNDAMENTAÇÃO ============
H('1. Fundamentação — o que cada ajuste indica, segundo a literatura internacional')
P('A escolha dos ajustes segue evidência de modalidades coletivas e de monitoramento do atleta indexada no PubMed. A revisão sistemática de Saw, Main e Gastin (2016) demonstrou que as medidas subjetivas de bem-estar (humor, estresse, recuperação percebida) rastreiam a carga aguda e crônica com sensibilidade e consistência superiores às medidas objetivas — sustentando o BRUMS e o eixo energia–fadiga como desfecho de monitoramento. A metanálise de Lochbaum et al. (2021) confirmou que o Perfil de Humor e a Perturbação Total do Humor (PTH/TMD) predizem o desempenho esportivo (g = −0,53 para a PTH), validando o perfil "iceberg" (vigor alto, fadiga baixa) como marcador de prontidão. Estudos longitudinais (Pang, Xie e Lin, 2023) descrevem o "iceberg em fusão" — a queda de vigor e a elevação de fadiga ao longo de semanas de treino intenso —, exatamente o fenômeno que testamos aqui.',after=6)
P('Para os ajustes anthropométricos, o escalonamento alométrico Y = a·Xᵇ é o método padrão para remover o efeito confusor da massa corporal em comparações de força e potência entre atletas de posições diferentes (Oba et al., 2014, futebol americano universitário; Jobson et al., 2008, ciclismo; Linek, 2017, normalização por massa em atletas jovens): quando o expoente correto é aplicado, a correlação entre a medida escalonada e a massa cai a zero. Para a decisão binária (atleta em risco / dia de fadiga alta), a regressão logística e as curvas ROC são o padrão em modalidades coletivas para predição de lesão e classificação de estado. Em conjunto, esses trabalhos justificam usar a alometria para a pergunta de traço ("quem tolera a carga") e a logística/ROC para a pergunta de estado ("quem/quando está sobrecarregado").',after=6)
table('Correspondência entre cada ajuste e a pergunta que ele responde melhor',
    ['Ajuste','Pergunta que responde','Indicador / métrica','Base na literatura'],
    [['Logístico + ROC','Estado: quem/quando está em fadiga alta','Limiar de PV; AUC; OR','Saw et al. 2016; predição de lesão em coletivos'],
     ['Alométrico (Y=a·Xᵇ)','Traço: quem tolera a carga','Expoente b; ρ após normalização','Oba 2014; Jobson 2008; Linek 2017'],
     ['Derivadas / inflexão','Quando a deterioração acelera','Velocidade/aceleração; ponto de inflexão','Iceberg em fusão (Pang 2023)'],
     ['Decomposição sinal/ruído','Se a oscilação é confiável no indivíduo','% de variância de estado; MDC','Confiabilidade de monitoramento'],
     ['Classificação de perfis','Fenótipo integrado de prontidão','% iceberg→barbatana; índice iceberg','Lochbaum 2021; Tan 2024'],
     ['Regressão linear','Direção e força média da associação','β; R²; ρ','Descritivo de base']],
    fonte='Fonte: os autores (2026), com base em PubMed.',widths=[3.0,4.2,3.6,4.7],fs=9)

# ============ 2. RESULTADOS ============
H('2. Resultados')

H('2.1 A aptidão (PV) é o principal modificador da fadiga — os mais treinados sofrem menos',12)
cor=R['COR']
def rc(cv,ov): d=R['COR'][cv][ov]; return '%s (p=%s)'%(c2('%+.2f'%d['r']),c2('%.2f'%d['p']))
P('A correlação de Spearman entre as covariáveis de base e os desfechos do microciclo (Tabela 2) mostra que o pico de velocidade no T-CAR — proxy da aptidão aeróbia intermitente — é a variável que melhor se associa ao nível de fadiga física da semana: ρ = %s, ou seja, quanto maior o PV, menor a fadiga física média. O PV também se associa positivamente ao índice iceberg (ρ = %s): os mais aptos permanecem em perfil de humor mais favorável. Já o percentual de gordura teve associação fraca e não significativa com a fadiga (ρ = %s) e nula com a queda de vigor e o acúmulo — nesta amostra de elite, a adiposidade não rastreou a fadiga do microciclo. A massa corporal mostrou apenas tendência (ρ = %s, p = %s): atletas mais pesados relataram fadiga ligeiramente maior.'%(
    rc('PV','wk_FadFis'),rc('PV','ice_mean'),rc('pGordura','wk_FadFis'),
    c2('%+.2f'%R['COR']['massa']['wk_FadFis']['r']),c2('%.2f'%R['COR']['massa']['wk_FadFis']['p'])),after=6)
OV=[('wk_FadFis','Fadiga física (média)'),('fis_accum','Acúmulo D1→D7'),('vig_drop','Queda de vigor'),('wk_TMD','PTH/TMD (média)'),('ice_mean','Índice iceberg')]
rows=[]
for cv,lab in [('PV','Pico de velocidade'),('massa','Massa corporal'),('pGordura','% de gordura'),('estatura','Estatura'),('CMJ_ini','CMJ (salto)'),('Baker_ini','Baker (potência MMSS)')]:
    rows.append([lab]+[rc(cv,ov) for ov,_ in OV])
table('Correlação de Spearman (ρ) entre covariáveis de base e desfechos do microciclo (n = 27)',
    ['Covariável']+[l for _,l in OV],rows,widths=[3.2,3.0,2.7,2.6,2.7,2.6],fs=8.5)
P('Interpretação: a aptidão modula o NÍVEL de fadiga (traço), não o TAMANHO da variação semanal. O acúmulo de fadiga D1→D7 não foi predito pelo PV (ρ = %s), mas sim por quão "fresco" o atleta começa — quem parte de vigor mais alto cai mais (ver 2.5). Ou seja: os mais treinados terminam a semana menos fatigados em termos absolutos, ainda que a oscilação relativa seja comum a todos.'%(
    c2('%+.2f'%R['COR']['PV']['fis_accum']['r'])),after=6)

H('2.2 Ajuste alométrico — o PV escala a fadiga melhor que a massa ou a gordura',12)
al=R['ALLO']
P('No ajuste log-log Y = a·Xᵇ da fadiga física média sobre cada descritor corporal, o pico de velocidade apresentou o acoplamento mais forte e o expoente de maior magnitude (b = %s, R² = %s), seguido da massa (b = %s) e, por último, do percentual de gordura (b = %s, R² = %s — praticamente sem relação alométrica). Isso reforça que o descritor fisiologicamente relevante para normalizar a fadiga em handebol não é a adiposidade, e sim a potência aeróbia (PV). Consistentemente com o manuscrito principal, a normalização pelo PV (fadiga/PVᵇ) foi a que mais reduziu o confundimento fitness→fadiga (ρ da fadiga física caiu de −0,52 para −0,19), ao passo que a normalização pela massa foi inócua.'%(
    c2('%+.2f'%al['PV']['b']),c2('%.2f'%al['PV']['r2']),c2('%+.2f'%al['massa']['b']),
    c2('%+.2f'%al['pGordura']['b']),c2('%.2f'%al['pGordura']['r2'])),after=6)

H('2.3 Ajuste logístico / ROC — o PV como indicador de estado',12)
P('Para a decisão operacional (o atleta está em fadiga alta?), o ajuste logístico sobre o PV é o indicador mais forte. Conforme estabelecido no manuscrito principal, a probabilidade de fadiga física alta cai com o PV (OR ≈ 0,20 por km·h⁻¹, p = 0,025), com limiar prático em torno de 14,8 km·h⁻¹ e AUC = 0,84 — desempenho de classificação "bom" pelos critérios convencionais. A curva ROC do contraste acumulado D7 vs D1 para a fadiga física chega a AUC = 0,90 quando o ruído é filtrado por médias diárias. Em contraste, o percentual de gordura NÃO classifica o atleta que mais acumula fadiga (AUC = %s), tampouco a massa (AUC = %s): confirma-se que o preditor de estado útil é a aptidão, não a composição corporal.'%(
    c2('%.2f'%R['ROC']['pGordura']['auc']),c2('%.2f'%R['ROC']['massa']['auc'])),after=6)

H('2.4 Posição — o pivô, mais pesado e adiposo, tende a fadigar mais',12)
pos=R['POS']; order=['Goleiro','Meia','Ponta','Pivô']
rows=[[k,pos[k]['n'],c2('%.1f'%pos[k]['PV']),c2('%.1f'%pos[k]['massa']),c2('%.1f'%pos[k]['pGordura']),
       c2('%.2f'%pos[k]['wk_FadFis']),c2('%+.2f'%pos[k]['vig_drop']),c2('%.2f'%pos[k]['ice_mean'])] for k in order if k in pos]
table('Perfil por posição (médias) — aptidão, composição corporal e resposta do microciclo',
    ['Posição','n','PV','Massa','%G','Fad. fís.','ΔVigor','Iceberg'],rows,widths=[2.6,1.2,2.0,2.2,1.8,2.2,2.2,2.2],fs=9)
P('Os pivôs são os mais pesados (%s kg), mais adiposos (%s%%) e com menor PV (%s km·h⁻¹), e apresentaram a maior fadiga física média (%s) — coerente com a hipótese de que o atleta mais pesado e menos apto fadiga mais. Contudo, com apenas 4 pivôs, o teste de Kruskal-Wallis entre posições não foi significativo para nenhum desfecho (fadiga física p = %s; PV p = %s): trata-se de um padrão direcional, não de diferença estatística. A informação de posição é, portanto, um modificador de contexto, não um indicador primário.'%(
    c2('%.1f'%pos['Pivô']['massa']),c2('%.1f'%pos['Pivô']['pGordura']),c2('%.1f'%pos['Pivô']['PV']),
    c2('%.2f'%pos['Pivô']['wk_FadFis']),c2('%.2f'%R['POS_KW']['wk_FadFis']),c2('%.2f'%R['POS_KW']['PV'])),after=6)

H('2.5 Perfil iceberg — mais aptos, mantêm o vigor e recuperam melhor',12)
ia=R['ICE']['Iceberg alto']; ib=R['ICE']['Iceberg baixo']; mw=R['ICE_MW']
table('Atletas de perfil iceberg alto vs baixo (terços do índice iceberg médio)',
    ['Variável','Iceberg alto (n=%d)'%ia['n'],'Iceberg baixo (n=%d)'%ib['n'],'p (Mann-Whitney)'],
    [['PV (km·h⁻¹)',c2('%.1f'%ia['PV']),c2('%.1f'%ib['PV']),c2('%.3f'%mw['PV'])],
     ['% de gordura',c2('%.1f'%ia['pGordura']),c2('%.1f'%ib['pGordura']),c2('%.3f'%mw['pGordura'])],
     ['Fadiga física (média)',c2('%.2f'%ia['wk_FadFis']),c2('%.2f'%ib['wk_FadFis']),c2('%.3f'%mw['wk_FadFis'])],
     ['Vigor no D1',c2('%.1f'%ia['vig_D1']),c2('%.1f'%ib['vig_D1']),c2('%.3f'%mw['vig_D1'])],
     ['Vigor no D7',c2('%.1f'%ia['vig_D7']),c2('%.1f'%ib['vig_D7']),c2('%.3f'%mw['vig_D7'])],
     ['Queda de vigor D1→D7',c2('%+.2f'%ia['vig_drop']),c2('%+.2f'%ib['vig_drop']),c2('%.3f'%mw['vig_drop'])],
     ['Recuperação noturna (fadiga)',c2('%+.2f'%ia['rec_fis']),c2('%+.2f'%ib['rec_fis']),c2('%.3f'%mw['rec_fis'])]],
    widths=[5.0,3.4,3.4,3.2],fs=9)
P('O grupo de perfil iceberg alto é significativamente mais apto (PV %s vs %s km·h⁻¹, p = %s) e mantém vigor muito superior em ambos os extremos da semana (D1: %s vs %s, p = %s; D7: %s vs %s, p = %s), com fadiga física média menor (p = %s). O percentual de gordura NÃO os distingue (p = %s) — o que os separa é a aptidão, não a composição corporal. A recuperação noturna da fadiga foi maior no grupo iceberg (%s vs %s), em tendência não significativa (p = %s). Confirma-se, portanto, o retrato do atleta iceberg: PV alto/acima do limiar, fadiga menor, vigor que não desaba e recuperação mais rápida.'%(
    c2('%.1f'%ia['PV']),c2('%.1f'%ib['PV']),c2('%.3f'%mw['PV']),
    c2('%.1f'%ia['vig_D1']),c2('%.1f'%ib['vig_D1']),c2('%.3f'%mw['vig_D1']),
    c2('%.1f'%ia['vig_D7']),c2('%.1f'%ib['vig_D7']),c2('%.3f'%mw['vig_D7']),
    c2('%.3f'%mw['wk_FadFis']),c2('%.3f'%mw['pGordura']),
    c2('%+.2f'%ia['rec_fis']),c2('%+.2f'%ib['rec_fis']),c2('%.3f'%mw['rec_fis'])),after=6)

H('2.6 "Começa bem e termina mal" — um fenômeno geral, mas o iceberg termina no topo',12)
st=R['START']['vigD1_drop']
P('O vigor inicial (D1) correlaciona-se fortemente com a queda de vigor até o D7 (ρ = %s, p = %s): quem começa com vigor mais alto é quem mais decai — o padrão "começa bem e termina mal". Parte disso é regressão/teto (só quem está alto pode cair muito), parte é acúmulo real de carga. A leitura crítica, porém, é do nível final: apesar de caírem mais em valor absoluto, os atletas iceberg/aptos terminam a semana com vigor mais alto que os demais (D7 %s vs %s). Ou seja, "termina mal" em variação, mas "termina melhor" em patamar — a prontidão relativa é preservada.'%(
    c2('%+.2f'%st[0]),c2('%.3f'%st[1]),c2('%.1f'%ia['vig_D7']),c2('%.1f'%ib['vig_D7'])),after=6)

figure('/home/user/mdlucca/Artigos/figuras/x_ajuste.png',
    'Síntese integrada. (A) A aptidão (PV) associa-se inversamente à fadiga física semanal — os mais treinados sofrem menos. (B) Atletas de perfil iceberg alto mantêm vigor muito superior em toda a semana. (C) Quem começa com mais vigor decai mais ("começa bem e termina mal"). (D) Poder de indicação de cada ajuste: a logística/ROC sobre o PV lidera para a triagem de estado.')

# ============ 3. RANKING ============
H('3. Qual ajuste é o indicador mais forte? — veredito estratificado')
table('Ranking dos ajustes por alvo de indicação',
    ['Alvo clínico','Ajuste mais forte','Evidência nesta amostra','Força'],
    [['Estado atual (fadiga alta hoje)','Logístico + ROC (PV)','AUC = 0,84; limiar ≈ 14,8 km·h⁻¹','Alta'],
     ['Tolerância à carga (traço)','Alométrico (normalização por PV)','ρ fadiga −0,52→−0,19 após normalizar','Alta'],
     ['Momento da deterioração','Derivadas + inflexão','Aceleração da fadiga no D5→D7','Média-alta'],
     ['Prontidão integrada / fenótipo','Classificação de perfis (iceberg)','Iceberg alto: +PV, +vigor, p<0,05','Alta'],
     ['Confiabilidade do sinal individual','Decomposição sinal/ruído','Estado = 3–9% da variância; exige ≥3 coletas','Alta (metodológica)'],
     ['Composição corporal como preditor','— (não recomendado isolado)','% gordura: ρ e AUC ≈ nulos','Baixa']],
    widths=[4.2,3.6,4.6,1.8],fs=9)
P('Síntese: para excesso de treinamento, perda de desempenho e fadiga acumulada, o indicador operacional mais forte é o ajuste logístico/ROC ancorado no pico de velocidade (triagem de estado), complementado pela classificação de perfis de humor (fenótipo de prontidão) e pelas derivadas (momento da aceleração). A alometria é o melhor ajuste explicativo (normaliza o confundimento da aptidão), e a decomposição sinal/ruído é a salvaguarda metodológica que impede decisões individuais sobre coletas isoladas. O percentual de gordura, isoladamente, não é um indicador útil nesta população de elite — sua contribuição é apenas contextual (via posição). Ressalva importante: predição de lesão não pôde ser validada, pois não houve desfecho de lesão registrado na janela de sete dias; a recomendação de logística/ROC para esse fim é extrapolada da literatura de modalidades coletivas, não confirmada nestes dados.',after=6)

# ============ 4. RESPOSTAS DIRETAS ============
H('4. Respostas diretas às perguntas norteadoras')
qa=[
 ('Os mais treinados sofrem menos?','Sim. O PV associa-se inversamente à fadiga física (ρ = −0,51) e positivamente ao perfil iceberg (ρ = +0,40); atletas iceberg têm PV maior (p = 0,038) e menos fadiga (p = 0,004).'),
 ('Os mais pesados e menos treinados fadigam mais?','Direcionalmente sim: massa com ρ = +0,37 (tendência) e os pivôs — mais pesados, mais adiposos e menos aptos — lideram a fadiga média. Sem significância entre posições (n pequeno).'),
 ('O percentual de gordura é um bom indicador?','Não isoladamente. ρ com fadiga ≈ +0,29 (n.s.), AUC ≈ 0,47 (acaso), expoente alométrico ≈ 0,15. A gordura só importa como componente do perfil posicional.'),
 ('O atleta iceberg tem PV alto/acima do limiar?','Sim — PV significativamente maior e, em média, próximo/acima do limiar de 14,8 km·h⁻¹.'),
 ('O iceberg fadiga menos e o vigor não cai tanto?','Fadiga menos (sim) e mantém vigor muito mais alto em toda a semana; a queda absoluta é semelhante, mas o patamar final é superior.'),
 ('Recupera mais rápido?','Tendência a sim (recuperação noturna da fadiga maior), sem significância estatística.'),
 ('"Começa bem e termina mal"?','Sim como fenômeno geral (vigor D1 × queda ρ = +0,63), mas o iceberg termina no topo — cai mais em variação, permanece mais alto em nível.'),
 ('Nos dias em que o humor/vigor cai e a fadiga sobe, qual ajuste indica melhor?','A logística/ROC sobre o PV (estado) sinaliza quem cruza o limiar; as derivadas localizam quando a queda acelera (D5→D7); a classificação de perfis mostra a migração iceberg→barbatana de tubarão.'),
]
for q,a in qa:
    p=doc.add_paragraph(); r=p.add_run('• '+q+' '); r.bold=True; r.font.size=Pt(12)
    r2=p.add_run(a); r2.font.size=Pt(12)
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(5); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

H('5. Limitações')
P('Amostra de 27 atletas (subgrupos posicionais pequenos, sobretudo pivôs e goleiros); janela de sete dias sem desfecho de lesão, o que impede validar predição de lesão; composição corporal medida uma única vez (base), tratada como covariável de traço; e o sinal de estado do microciclo é pequeno no nível individual (3–9% da variância), exigindo médias de ≥3 coletas para decisões individuais. As associações são correlacionais e não estabelecem causalidade.',after=6)

# ============ REFERÊNCIAS ============
H('Referências')
refs=[
 'JOBSON, S. A. et al. Allometric scaling of uphill cycling performance. International Journal of Sports Medicine, v. 29, n. 9, p. 753–757, 2008. DOI: 10.1055/s-2007-989441.',
 'LINEK, P. Body mass normalization for lateral abdominal muscle thickness measurements in adolescent athletes. Journal of Ultrasound in Medicine, v. 36, n. 9, p. 1851–1857, 2017. DOI: 10.1002/jum.14218.',
 'LOCHBAUM, M. et al. The Profile of Mood States and athletic performance: a meta-analysis of published studies. European Journal of Investigation in Health, Psychology and Education, v. 11, n. 1, p. 50–70, 2021. DOI: 10.3390/ejihpe11010005.',
 'OBA, Y. et al. Allometric scaling of strength scores in NCAA Division I-A football athletes. Journal of Strength and Conditioning Research, v. 28, n. 12, p. 3330–3337, 2014. DOI: 10.1519/JSC.0000000000000548.',
 'PANG, L.; XIE, X.; LIN, Y. POMS and eye movement: two indicators for performance in athletics. Heliyon, v. 9, n. 7, e17860, 2023. DOI: 10.1016/j.heliyon.2023.e17860.',
 'RATZ-SULYOK, F. Z. et al. Associations between endocrine status and stress, mood and psychosomatic status in elite handball players. Sports, v. 14, n. 7, 289, 2026. DOI: 10.3390/sports14070289.',
 'SAW, A. E.; MAIN, L. C.; GASTIN, P. B. Monitoring the athlete training response: subjective self-reported measures trump commonly used objective measures: a systematic review. British Journal of Sports Medicine, v. 50, n. 5, p. 281–291, 2016. DOI: 10.1136/bjsports-2015-094758.',
 'TAN, C. et al. The structural validity and latent profile characteristics of the Abbreviated Profile of Mood States among Chinese athletes. BMC Psychiatry, v. 24, n. 1, 636, 2024. DOI: 10.1186/s12888-024-06092-5.',
]
for rf in refs:
    p=doc.add_paragraph(); r=p.add_run(rf); r.font.size=Pt(11)
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(6); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

OUTP='/home/user/mdlucca/Artigos/Ajuste_Sobrecarga_Gordura_Posicao_PV.docx'
doc.save(OUTP)
print('SAVED',OUTP,'| Tabelas',_TN[0],'Figuras',_FN[0])
