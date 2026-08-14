import warnings; warnings.filterwarnings('ignore')
import numpy as np, pandas as pd, json, os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
SC='/tmp/claude-0/-home-user-mdlucca/e1dba24c-b1d7-5908-9106-f2f4aaf3f56a/scratchpad'
FIG='/home/user/mdlucca/Artigos/figuras'
amp=json.load(open(f'{SC}/amp_docx.json')); grp=amp['grp']; per=amp['per']
DN=json.load(open(f'{SC}/denoise.json')); PF=json.load(open(f'{SC}/pair_fatigue.json')); FF=PF['grp']['FadFisica']
NORM=json.load(open(f'{SC}/norm_data.json')); APP=json.load(open(f'{SC}/app_data.json')); INF=json.load(open(f'{SC}/inflexao.json'))

doc=Document()
s=doc.sections[0]; s.left_margin=Cm(3); s.top_margin=Cm(3); s.right_margin=Cm(2); s.bottom_margin=Cm(2)
n=doc.styles['Normal']; n.font.name='Times New Roman'; n.font.size=Pt(12)
n.paragraph_format.line_spacing=1.5; n.paragraph_format.space_after=Pt(0); n.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

FN='Times New Roman'
def _set(run,sz=12,b=False,it=False,color=None):
    run.font.name=FN; run.font.size=Pt(sz); run.bold=b; run.italic=it
    if color: run.font.color.rgb=color
def sec(num,txt):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(6)
    p.paragraph_format.line_spacing=1.5; p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    _set(p.add_run(f'{num} {txt}'.upper() if num.count('.')==0 else f'{num} {txt}'),12,b=True); return p
def sub(num,txt):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(4); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    _set(p.add_run(f'{num} {txt}'),12,b=True); return p
def P(t,indent=True,just=True,size=12,it=False,b=False):
    p=doc.add_paragraph(); r=p.add_run(t); _set(r,size,b=b,it=it)
    p.paragraph_format.line_spacing=1.5
    if indent: p.paragraph_format.first_line_indent=Cm(1.25)
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY if just else WD_ALIGN_PARAGRAPH.LEFT
    return p
def fig(name,num,titulo,fonte='Elaborado pelos autores (2026).',w=6.0):
    cap=doc.add_paragraph(); cap.alignment=WD_ALIGN_PARAGRAPH.CENTER; cap.paragraph_format.space_before=Pt(8); cap.paragraph_format.line_spacing=1.0
    _set(cap.add_run(f'Figura {num} – {titulo}'),10)
    path=f'{FIG}/{name}.png'
    if os.path.exists(path):
        doc.add_picture(path,width=Inches(w)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    else: print('MISS FIG',name)
    fp=doc.add_paragraph(); fp.alignment=WD_ALIGN_PARAGRAPH.CENTER; fp.paragraph_format.space_after=Pt(8); fp.paragraph_format.line_spacing=1.0
    _set(fp.add_run(f'Fonte: {fonte}'),10)
def _borders(t):
    tblPr=t._tbl.tblPr; b=OxmlElement('w:tblBorders')
    for e in ('top','bottom','insideH'):
        x=OxmlElement(f'w:{e}'); x.set(qn('w:val'),'single'); x.set(qn('w:sz'),'6'); x.set(qn('w:space'),'0'); x.set(qn('w:color'),'000000'); b.append(x)
    for e in ('left','right','insideV'):
        x=OxmlElement(f'w:{e}'); x.set(qn('w:val'),'none'); b.append(x)
    tblPr.append(b)
def table(num,titulo,headers,rows,fonte='Dados da pesquisa (2026).',fs=10):
    cap=doc.add_paragraph(); cap.alignment=WD_ALIGN_PARAGRAPH.LEFT; cap.paragraph_format.space_before=Pt(8); cap.paragraph_format.line_spacing=1.0
    _set(cap.add_run(f'Tabela {num} – {titulo}'),10)
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=True
    _borders(t)
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.paragraphs[0].line_spacing=1.0
        r=c.paragraphs[0].add_run(str(h)); _set(r,fs,b=True)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].paragraphs[0].line_spacing=1.0
            r=cells[i].paragraphs[0].add_run(str(v)); _set(r,fs)
    fp=doc.add_paragraph(); fp.alignment=WD_ALIGN_PARAGRAPH.LEFT; fp.paragraph_format.space_after=Pt(8); fp.paragraph_format.line_spacing=1.0
    _set(fp.add_run(f'Fonte: {fonte}'),10)

V4=[('Vigor','Vigor'),('Fadiga','Fadiga BRUMS'),('TMD','PTH/TMD'),('FadMental','Fadiga mental')]

# ===== TÍTULO =====
tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
_set(tp.add_run('SINAL, RUÍDO E AMPLITUDE NO MONITORAMENTO PSICOMÉTRICO DA FADIGA E DO HUMOR DURANTE UM MICROCICLO DE CHOQUE DE HIIT EM ATLETAS DE HANDEBOL DE ELITE'),14,b=True)
sub2=doc.add_paragraph(); sub2.alignment=WD_ALIGN_PARAGRAPH.CENTER; _set(sub2.add_run('[Autoria e afiliação a preencher]'),11,it=True)
doc.add_paragraph()

# ===== RESUMO =====
rp=doc.add_paragraph(); _set(rp.add_run('RESUMO'),12,b=True); rp.paragraph_format.line_spacing=1.0
P('O monitoramento psicométrico do humor confunde rotineiramente flutuação real (sinal) e erro de medida (ruído). Este estudo quantificou a amplitude das respostas de humor e fadiga a um microciclo de choque de HIIT e separou, para cada variável, quanto da oscilação é sinal do microciclo, traço estável e ruído, determinando quando a variação é interpretável no grupo e no atleta. Vinte e sete atletas de handebol de elite responderam ao BRUMS-24 e a autoavaliações de fadiga física e mental em três momentos diários (pré, intra-sessão e pós) ao longo de sete dias com três sessões de HIIT 4×4 min a 104% do pico de velocidade do T-CAR. A variância de cada variável foi decomposta (modelo atleta + dia) em sinal, traço e ruído; derivaram-se o erro típico de medida, a mudança mínima detectável (individual e de grupo) e a menor mudança relevante. O processamento e a auditoria interna foram inteiramente roteirizados em Python. O sinal do microciclo é a menor fração da variância (1–12%); a fadiga física concentra o maior sinal. No grupo a oscilação semanal supera o ruído por 4–10 vezes, mas no indivíduo exige médias de ≥3–5 coletas. A trajetória do perfil de humor é não linear, com ponto de inflexão robusto no dia ≈ 4, dividindo a fase de alívio da precipitação do Dia 7. Removido o ruído, as conclusões tornam-se mais nítidas, sem manufaturar sinal onde há efeito de piso.',indent=False)
kw=doc.add_paragraph(); _set(kw.add_run('Palavras-chave: '),12,b=True); _set(kw.add_run('Monitoramento da carga. BRUMS. Mudança mínima detectável. Relação sinal-ruído. HIIT.'),12); kw.paragraph_format.line_spacing=1.5

# ===== 1 INTRODUÇÃO =====
sec('1','Introdução')
P('O monitoramento psicométrico do estado de humor e da fadiga é amplamente empregado para individualizar a carga e prevenir o excesso de treinamento (SAW; MAIN; GASTIN, 2016). Instrumentos como a Brunel Mood Scale (BRUMS) são rápidos e sensíveis, mas partilham um problema raramente enfrentado de forma explícita: parte substancial da variação diária de um atleta não reflete o efeito do treino, e sim erro de medida e flutuação aleatória. Interpretar essa flutuação como resposta ao treino gera decisões espúrias.')
P('Três quantidades coexistem em cada medida: o sinal do microciclo (a mudança sistemática que os dias de treino impõem), o traço (diferenças estáveis entre atletas) e o ruído (erro de medida e variação não sistemática). Separá-las é pré-requisito para saber qual variação merece ação. As ferramentas psicométricas para tal — erro típico de medida (ETM), mudança mínima detectável (MDC95) e menor mudança relevante (SWC) — são conhecidas (HOPKINS, 2000), mas raramente aplicadas de forma integrada a um microciclo real.')
P('Este estudo aplica essa decomposição a um microciclo de choque de HIIT em atletas de handebol de elite, com os objetivos de: (a) descrever a amplitude das respostas de humor e fadiga; (b) decompor a variância em sinal, traço e ruído; (c) estabelecer a detectabilidade no grupo e no indivíduo; (d) demonstrar o efeito da remoção do ruído; e (e) caracterizar o comportamento dos seis perfis de humor (PARSONS-SMITH; TERRY; MACHIN, 2017) e localizar o ponto de inflexão de sua trajetória, com e sem ruído.')

# ===== 2 MÉTODOS =====
sec('2','Materiais e Métodos')
sub('2.1','Amostra e delineamento')
ag=APP['socio']['agg']
P(f"Participaram 27 atletas de handebol de elite do sexo masculino (idade {ag['idade'][0]:.0f} ± {ag['idade'][1]:.0f} anos; experiência {ag['exp'][0]:.1f} ± {ag['exp'][1]:.1f} anos; estatura {ag['estatura'][0]:.1f} ± {ag['estatura'][1]:.1f} cm; massa {ag['massa'][0]:.1f} ± {ag['massa'][1]:.1f} kg; pico de velocidade no T-CAR {ag['PV'][0]:.2f} ± {ag['PV'][1]:.2f} km/h), monitorados durante um microciclo de choque de sete dias (21 a 28 de abril de 2024), totalizando 456 observações. O delineamento é observacional longitudinal, com medidas repetidas aninhadas (observações ⊂ dias ⊂ atletas). Os dados são apresentados de forma anonimizada (A01–A27).")
sub('2.2','Instrumentos e carga')
el=APP['extload']
P(f"O humor foi avaliado pelo BRUMS-24 (subescalas tensão, depressão, raiva, vigor, fadiga e confusão) e pela Perturbação Total do Humor (PTH/TMD); a fadiga física e a fadiga mental por escalas de 0 a 10. As coletas ocorreram em três momentos (pré, intra-sessão e pós). As sessões de HIIT (dias 2, 4 e 7) consistiram em 4×4 min a 104% do pico de velocidade do T-CAR; a carga externa, de prescrição relativa, foi fixa entre sessões (velocidade {el['vel_mean']:.1f} ± {el['vel_sd']:.1f} km/h; distância {el['dist_mean']:.0f} ± {el['dist_sd']:.0f} m/sessão).")
sub('2.3','Processamento em Python e auditoria interna')
P('Todo o tratamento foi roteirizado em Python (pandas, numpy, SciPy, statsmodels): extração das planilhas, anonimização irreversível na fronteira do processamento (códigos A01–A27, com varredura automática que impede vazamento de nomes), construção de bases derivadas e recomputação independente de cada estatística de destaque (auditoria interna por script). A auditoria confirmou a reprodução idêntica dos achados longitudinais centrais (queda do vigor e elevação da fadiga do Dia 1 ao Dia 7, com respectivos tamanhos de efeito e níveis de significância).')
sub('2.4','Decomposição sinal–traço–ruído e limiares')
P('A variância de cada variável foi decomposta por modelo de dois fatores (y ~ atleta + dia; soma de quadrados tipo I) em sinal do microciclo, traço e ruído. Do ruído derivaram-se o ETM = √(QM residual), a MDC95 individual = 1,96·√2·ETM, a MDC95 de grupo = MDC95/√n e a SWC = 0,2·DP-entre-atletas. A relação sinal-ruído (SNR) é a amplitude do sinal semanal dividida pelo ETM.')
sub('2.5','Perfis, ponto de inflexão e análise estatística')
P('Cada observação foi classificada em um dos seis perfis de humor por atribuição ao centroide canônico mais próximo (subescalas em z); o índice-iceberg (vigor − média das negativas, em z) resume o perfil. O ponto de inflexão da trajetória foi estimado como o dia em que a segunda derivada troca de sinal, sobre o sinal bruto (com ruído) e sobre o sinal suavizado por LOWESS (sem ruído), com intervalo de confiança por bootstrap (reamostragem de atletas). A normalidade foi testada por Shapiro–Wilk; dada a não normalidade das variáveis de humor, adotou-se a via não paramétrica (mediana, Spearman, rmcorr, Wilcoxon, Friedman), com a agregação por atleta como unidade dos contrastes. Adotou-se α = 0,05.')

# ===== 3 RESULTADOS =====
sec('3','Resultados')
sub('3.1','Normalidade e estatística descritiva')
dist=[d for d in NORM['dist'] if d['level']=='obs'][:8]
rows=[[d['lab'],d['n'],f"{d['W']:.3f}".replace('.',','),(f"{d['p']:.1e}"),('não' if not d['normal'] else 'sim')] for d in dist]
table('1','Teste de normalidade (Shapiro–Wilk) das distribuições brutas',['Variável','n','W','p','Normal?'],rows)
P('Quatorze das dezoito variáveis são não normais (todas as de humor; p de 3×10⁻⁴ a 3×10⁻³⁵); as subescalas negativas concentram 50–80% das respostas no piso, o que fundamenta a via não paramétrica e antecipa sua baixa responsividade.')

sub('3.2','Amplitude')
rows=[['Fadiga física',f"{FF['sig']:.2f}".replace('.',','),f"{PF['per']['FadFisica']['amp_mean']:.2f} ± {PF['per']['FadFisica']['amp_sd']:.2f}".replace('.',',')]]
for v,lab in V4:
    rows.append([lab,f"{grp[v]['sig']:.2f}".replace('.',','),f"{per[v]['amp_mean']:.2f} ± {per[v]['amp_sd']:.2f}".replace('.',',')])
table('2','Amplitude do sinal do microciclo versus amplitude individual',['Variável','Amplitude do sinal (grupo)','Amplitude intra-atleta (M ± DP)'],rows)
P('Em todas as variáveis a amplitude do sinal (a trajetória média do grupo) é muito menor que a amplitude intra-atleta: o atleta oscila 2 a 3 vezes mais que o microciclo médio, porque sua oscilação carrega o próprio ruído.')
fig('abnt_f1_trajetoria','1','Evolução semanal do grupo (M ± DP; suavização LOWESS pontilhada; linhas tracejadas verticais = dias de HIIT)')

sub('3.3','Decomposição em sinal, traço e ruído')
rows=[['Fadiga física',f"{FF['pday']:.1f}%".replace('.',','),f"{FF['ptrait']:.1f}%".replace('.',','),f"{FF['pnoise']:.1f}%".replace('.',','),f"{FF['snr']:.2f}".replace('.',',')]]
for v,lab in V4:
    g=grp[v]; rows.append([lab,f"{g['pday']:.1f}%".replace('.',','),f"{g['ptrait']:.1f}%".replace('.',','),f"{g['pnoise']:.1f}%".replace('.',','),f"{g['snr']:.2f}".replace('.',',')])
table('3','Decomposição da variância e relação sinal-ruído (SNR)',['Variável','% Sinal','% Traço','% Ruído','SNR'],rows)
P('O sinal do microciclo é sempre a menor fatia da variância (1–12%); o grosso é traço estável (32–73%) e ruído (26–64%). A fadiga física lidera em sinal (12%) e SNR (2,02); a fadiga mental é quase pura estrutura de traço (73%; SNR < 1).')
fig('abnt_f_variancia','2','Partição da variância por variável: sinal do microciclo, traço e ruído',w=5.6)

sub('3.4','Trajetória semanal e curso intra-sessão')
P('A trajetória do eixo energia–fadiga é não linear (componentes linear e cúbica significativas; crescimento cúbico vence por AIC/LRT): subida inicial, alívio no meio da semana e precipitação no Dia 7 — o maior passo diário ocorre de D6 para D7. O momento intra-sessão revela que o curso agudo não é monotônico: a fadiga frequentemente pica no meio da sessão e recua ligeiramente ao final; a resposta aguda é mais nítida no Dia 1 e no fim da semana.')
fig('abnt_f_prepos','3','Curso intra-dia (pré → mid → pós) das quatro variáveis do eixo energia–fadiga (* Wilcoxon pareado pré×pós p < 0,05)')

sub('3.5','Variabilidade individual e detectabilidade')
fig('abnt_f_heatmap','4','Amplitude semanal por atleta (z por variável): heterogeneidade estrutural entre atletas',w=5.6)
rows=[['Fadiga física',f"{FF['sig']:.2f}".replace('.',','),f"{FF['mdc']:.2f}".replace('.',','),'não',f"{FF['mdc_g']:.2f}".replace('.',','),'SIM']]
for v,lab in V4:
    g=grp[v]; rows.append([lab,f"{g['sig']:.2f}".replace('.',','),f"{g['mdc']:.2f}".replace('.',','),('SIM' if g['detect_ind'] else 'não'),f"{g['mdc_g']:.2f}".replace('.',','),('SIM' if g['detect_grp'] else 'não')])
table('4','Detectabilidade em dois níveis: a oscilação semanal supera o ruído?',['Variável','Ampl. sinal','MDC95 indiv.','> indiv.?','MDC95 grupo','> grupo?'],rows)
P('No grupo, a oscilação semanal supera a MDC de grupo por 4 a 10 vezes — sinal robusto no coletivo. No indivíduo, a mesma oscilação fica abaixo da MDC95 de uma coleta: a decisão individual exige médias de ≥3–5 coletas, o que quantifica a recomendação de referenciar cada atleta à sua própria linha de base.')

sub('3.6','Efeito da remoção do ruído')
Robs=pd.DataFrame(DN['Robs']); Rden=pd.DataFrame(DN['Rden'])
rows=[]
for x,y,l in [('Vigor','Fadiga','Vigor × Fadiga'),('Fadiga','TMD','Fadiga × PTH/TMD'),('Vigor','TMD','Vigor × PTH/TMD')]:
    rows.append([l,f"{Robs.loc[x,y]:+.2f}".replace('.',','),f"{Rden.loc[x,y]:+.2f}".replace('.',',')])
table('5','Correlações do eixo energia–fadiga antes e depois de remover o ruído',['Par','r observado (com ruído)','r sem ruído (grupo)'],rows)
P('Removido o ruído, as associações fortalecem-se (Vigor × Fadiga passa de −0,44 para −0,93 no nível-grupo) e os tamanhos de efeito crescem (fadiga física dz +1,74 → +2,36; vigor −0,95 → −1,25): a leitura bruta subestimava o custo do microciclo. Remover o ruído, porém, não cria sinal onde há efeito de piso — as subescalas negativas permanecem planas.')
fig('abnt_f_corr','5','Matriz de correlações do eixo energia–fadiga: observada (com ruído) versus sem ruído',w=5.8)
fig('abnt_f_mdck','6','MDC95 individual em função do número de coletas (k); linha tracejada = amplitude do sinal semanal',w=5.6)

sub('3.7','Contrastes entre variáveis')
P('A fadiga física e a mental dissociam ao longo da semana (a física quase dobra, dz +1,74; a mental fica plana, dz +0,33), embora se acoplem no instante quando removido o erro de medida (0,50 → 0,78 desatenuada): o custo é somático, não cognitivo. O Vigor e a PTH/TMD são espelhos inversos do mesmo eixo (−0,62 → −0,78 sem ruído); como o vigor compõe a PTH, esta não é independente e acrescenta pouca informação além do par vigor + fadiga.')
fig('abnt_f_fisica_mental','7','Fadiga física (acumula) versus fadiga mental (estável) ao longo da semana',w=5.4)
fig('abnt_f_vigor_tmd','8','Vigor e PTH/TMD padronizados (z): acoplamento inverso do eixo energia–humor',w=5.4)

sub('3.8','Perfis de humor e ponto de inflexão')
tab=INF['tab']
rows=[[f"Dia {dd}",f"{tab['Iceberg'][str(dd)]:.0f}",f"{tab['Barbatana tubarão'][str(dd)]:.0f}",f"{tab['Superfície'][str(dd)]:.0f}",f"{tab['Iceberg invertido'][str(dd)]:.0f}",f"{tab['Everest invertido'][str(dd)]:.0f}"] for dd in [1,4,7]]
table('6','Distribuição dos seis perfis de humor (%) em dias-chave',['Dia','Iceberg','Barbatana','Superfície','Iceberg inv.','Everest inv.'],rows)
P('O grupo migra do perfil iceberg (40% no D1 para 17% no D7) para a barbatana de tubarão (2% para 28%) — fadiga em pico com vigor preservado, custo somático e não afetivo. A trajetória do índice-iceberg é cúbica, com alívio no meio da semana e precipitação no D7.')
fig('abnt_f_perfis','9','Migração dos seis perfis de humor ao longo da semana; linha tracejada = ponto de inflexão (≈ dia 4)',w=5.8)
rows=[['Cubic sobre médias diárias brutas (com ruído)',f"dia {INF['infl_raw']:.2f}".replace('.',',')],
      ['Cubic sobre trajetória suavizada (sem ruído)',f"dia {INF['infl_sm']:.2f}".replace('.',',')],
      ['LOWESS — 2ª derivada = 0 (sem ruído)',f"dia {INF['infl_smooth']:.2f}".replace('.',',')],
      ['Bootstrap com ruído (mediana [IC95%])',f"dia {INF['ci_noise'][1]:.2f} [{INF['ci_noise'][0]:.2f}; {INF['ci_noise'][2]:.2f}]".replace('.',',')],
      ['Bootstrap sem ruído (mediana [IC95%])',f"dia {INF['ci_clean'][1]:.2f} [{INF['ci_clean'][0]:.2f}; {INF['ci_clean'][2]:.2f}]".replace('.',',')]]
table('7','Ponto de inflexão da trajetória do perfil por diferentes métodos',['Método de estimativa','Ponto de inflexão'],rows)
P('Os métodos convergem para o dia ≈ 4 (3,8–4,1). O ponto estimado é praticamente idêntico com e sem ruído, o que o torna robusto ao tratamento do ruído; o intervalo mais largo do LOWESS decorre da flexibilidade do suavizador (troca viés por variância), não de maior ruído. O dia 4 é o divisor de águas do microciclo: até ele a deterioração do perfil desacelera, a partir dele acelera rumo ao pico do Dia 7.')
fig('abnt_f_indice','10','Índice-iceberg: observações (ruído), média diária e curva LOWESS (sem ruído); faixa cinza = IC95% do ponto de inflexão',w=5.8)

# ===== 4 DISCUSSÃO =====
sec('4','Discussão')
P('Três mensagens sustentam-se. Primeiro, a deterioração do humor é real e não linear, mas ocupa fração pequena da variância total (1–12%): a maior parte do que se observa no dia a dia de um atleta é traço estável e ruído, e superinterpretar oscilações isoladas conduz a decisões espúrias. Segundo, a detectabilidade depende do nível: a média do grupo remove o ruído por √n e torna o sinal inequívoco, ao passo que a leitura de um único atleta com uma coleta fica abaixo da MDC95; a solução não é abandonar o monitoramento individual, mas agregá-lo (≥3–5 coletas), reconciliando utilidade clínica e confiabilidade psicométrica. Terceiro, remover o ruído esclarece o fenômeno sem alterá-lo — as associações do eixo energia–fadiga aproximam-se de uma dimensão bipolar quase perfeita e os efeitos crescem, mas o efeito de piso das subescalas negativas permanece, confirmando que sua não resposta é real.')
P('A análise dos perfis e do ponto de inflexão acrescenta uma leitura temporal acionável: a migração do iceberg para a barbatana de tubarão e a inflexão no dia ≈ 4 delimitam a janela a partir da qual a carga acumulada passa a cobrar seu preço de forma crescente. Do ponto de vista aplicado, recomenda-se monitorar a fadiga física e o vigor por tendência individual, com limiares de MDC, intervindo na recuperação em torno do dia 4 e evitando sobreinterpretar variações isoladas ou variáveis de piso.')

# ===== 5 CONSIDERAÇÕES =====
sec('5','Considerações Finais')
P('Num microciclo de choque de HIIT, a resposta do humor é real, não linear e concentrada no eixo energia–fadiga, com a fadiga física como marcador dominante e ponto de inflexão robusto no dia ≈ 4. A maior parte da oscilação observada, porém, é traço e ruído; separá-los mostra que o sinal é robusto no grupo mas exige agregação de coletas para a decisão individual. Removido o ruído, as conclusões tornam-se mais nítidas, sem mudar de direção. Recomenda-se centrar o monitoramento na fadiga física e no vigor, com limiares de MDC e médias de ≥3–5 coletas. Como limitações, trata-se de estudo observacional restrito a um microciclo, sem grupo-controle, com carga externa fixa que impede a análise dose-resposta.')

# ===== REFERÊNCIAS =====
sec('REFERÊNCIAS','')
refs=[
 'HOPKINS, W. G. Measures of reliability in sports medicine and science. Sports Medicine, v. 30, n. 1, p. 1-15, 2000.',
 'MORGAN, W. P. Selected psychological factors limiting performance: a mental health model. In: CLARKE, D. H.; ECKERT, H. M. (Ed.). Limits of human performance. Champaign: Human Kinetics, 1985. p. 70-80.',
 'PARSONS-SMITH, R. L.; TERRY, P. C.; MACHIN, M. A. Identification and description of novel mood profile clusters. Frontiers in Psychology, v. 8, art. 1958, 2017.',
 'ROHLFS, I. C. P. M. et al. A Escala de Humor de Brunel (Brums): instrumento para detecção precoce da síndrome do excesso de treinamento. Revista Brasileira de Medicina do Esporte, v. 14, n. 3, p. 176-181, 2008.',
 'SAW, A. E.; MAIN, L. C.; GASTIN, P. B. Monitoring the athlete training response: subjective self-reported measures. British Journal of Sports Medicine, v. 50, n. 5, p. 281-291, 2016.',
 'TERRY, P. C.; LANE, A. M.; FOGARTY, G. J. Construct validity of the Profile of Mood States — Adolescents for use with adults. Psychology of Sport and Exercise, v. 4, n. 2, p. 125-139, 2003.',
]
for r in refs:
    p=doc.add_paragraph(); _set(p.add_run(r),12); p.paragraph_format.line_spacing=1.0; p.paragraph_format.space_after=Pt(6); p.alignment=WD_ALIGN_PARAGRAPH.LEFT

doc.add_paragraph()
P('Nota de reprodutibilidade e governança de dados: todo o processamento e as análises são reproduzíveis pelos scripts em Python depositados no repositório; os atletas são identificados exclusivamente por códigos A01–A27 e nenhum dado bruto com identificação é distribuído.',size=9,it=True,indent=False)

OUTP='/home/user/mdlucca/Artigos/Artigo_Cientifico_ABNT_Sinal_Ruido.docx'
doc.save(OUTP); print('SAVED',OUTP)
d2=Document(OUTP); print('paras',len(d2.paragraphs),'tables',len(d2.tables),'imgs',len(d2.inline_shapes))
