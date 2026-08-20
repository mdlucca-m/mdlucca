# -*- coding: utf-8 -*-
import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
R=json.load(open('all_results.json')); PHY=json.load(open('physio2.json'))
FG='/home/user/mdlucca/Artigos/figuras'
def c2(s): return str(s).replace('.',',')
def num(x,d=1):
    return 'n/d' if x is None else c2(f'{x:.{d}f}')
def pvt(p): return '< 0,001' if p<0.001 else c2(f'{p:.3f}')
def mag(dz):
    a=abs(dz); return 'trivial' if a<0.2 else 'pequeno' if a<0.5 else 'médio' if a<0.8 else 'grande'
doc=Document()
st=doc.styles['Normal']; st.font.name='Times New Roman'; st.font.size=Pt(12)
st.element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman'); st.paragraph_format.line_spacing=1.5; st.paragraph_format.space_after=Pt(0)
sec=doc.sections[0]; sec.top_margin=Cm(2.5); sec.left_margin=Cm(2.5); sec.bottom_margin=Cm(2); sec.right_margin=Cm(2)
_TN=[0]; _FN=[0]
def P(t='',b=False,after=0,ind=True):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=b; r.font.size=Pt(12); r.font.name='Times New Roman'
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(after); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if ind: p.paragraph_format.first_line_indent=Cm(1.25)
    return p
def H(t,before=10):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(12); r.font.name='Times New Roman'
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(4)
def _bd(c):
    tcPr=c._tc.get_or_add_tcPr(); bd=OxmlElement('w:tcBorders')
    for e in ['top','bottom']:
        el=OxmlElement('w:'+e); el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'6'); el.set(qn('w:color'),'000000'); bd.append(el)
    tcPr.append(bd)
def table(cap,header,rows,fs=9,note=None):
    _TN[0]+=1
    p=doc.add_paragraph(); r=p.add_run('Tabela %d – %s'%(_TN[0],cap)); r.font.size=Pt(11); r.font.name='Times New Roman'
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
    t=doc.add_table(rows=1,cols=len(header)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=True
    for i,ht in enumerate(header):
        cc=t.rows[0].cells[i]; cc.text=''; rr=cc.paragraphs[0].add_run(ht); rr.bold=True; rr.font.size=Pt(fs); rr.font.name='Times New Roman'
        cc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if i>0 else WD_ALIGN_PARAGRAPH.LEFT; _bd(cc)
    for row in rows:
        cs=t.add_row().cells
        for i,v in enumerate(row):
            cs[i].text=''; rr=cs[i].paragraphs[0].add_run(str(v)); rr.font.size=Pt(fs); rr.font.name='Times New Roman'
            cs[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if i>0 else WD_ALIGN_PARAGRAPH.LEFT; _bd(cs[i])
    if note:
        pn=doc.add_paragraph(); rn=pn.add_run(note); rn.font.size=Pt(8.5); rn.italic=True; rn.font.name='Times New Roman'; pn.paragraph_format.space_after=Pt(6)
    else: doc.add_paragraph().paragraph_format.space_after=Pt(2)
def figure(path,cap,w=16.0):
    _FN[0]+=1
    pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; pp.add_run().add_picture(path,width=Cm(w)); pp.paragraph_format.space_before=Pt(6)
    pc=doc.add_paragraph(); pc.alignment=WD_ALIGN_PARAGRAPH.CENTER; rc=pc.add_run('Figura %d – %s'%(_FN[0],cap)); rc.font.size=Pt(11); rc.font.name='Times New Roman'
    pc.paragraph_format.space_after=Pt(6)

DIMS=[('Vigor','Vigor'),('Fadiga','Fadiga'),('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),
      ('Confusao','Confusão'),('TMD','PTH'),('FadFisica','Fadiga física'),('FadMental','Fadiga mental')]

# ===== TÍTULO / OBJETIVO =====
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('RESULTADOS: PERFIL E MUDANÇA DO HUMOR DE ATLETAS DE HANDEBOL DE ELITE NA ÚLTIMA SEMANA DE PRÉ-TEMPORADA')
r.bold=True; r.font.size=Pt(13); p.paragraph_format.space_after=Pt(8)
H('Objetivo',before=2)
P('Classificar o perfil e a mudança do humor de atletas de handebol de elite ao longo da última semana de pré-temporada '
 'e relacionar o pico de velocidade no teste de Carminatti (T-car) com a deterioração do humor, com a análise do '
 'comportamento de todas as dimensões do BRUMS ao longo da semana, por dia e entre os dias, por posição e por atleta.')

# ===== 1. DESCRITIVA (semana) =====
H('1 Descrição de todas as variáveis na semana')
ds=R['desc_semana']
table('Estatística descritiva de todas as dimensões do BRUMS no conjunto das observações da semana.',
 ['Dimensão','Média','DP','Mediana','Mín.–Máx.','CV (%)'],
 [[lab,num(ds[k]['m'],2),num(ds[k]['sd'],2),num(ds[k]['med'],1),f"{num(ds[k]['mn'],0)}–{num(ds[k]['mx'],0)}",num(ds[k]['cv'],0)] for k,lab in DIMS],
 note='DP: desvio-padrão; CV: coeficiente de variação. Escores das dimensões de 0 a 16; PTH: perturbação total do humor.')

# ===== 2. POR DIA =====
H('2 Comportamento por dia (D1 a D7)')
P('A Figura 1 apresenta a trajetória média de cada dimensão ao longo dos sete dias, com a mudança entre o primeiro e o '
 'último dia (Δ) e o tamanho de efeito (dz) de cada uma.')
figure(f'{FG}/res_byday.png','Comportamento de todas as dimensões do BRUMS por dia (média diária, D1 a D7). Cada painel traz a mudança D1 para D7 (Δ) e o tamanho de efeito (dz; * p < 0,05).',w=16.5)
bd=R['por_dia']
table('Média diária de todas as dimensões do BRUMS, do primeiro ao sétimo dia.',
 ['Dimensão','D1','D2','D3','D4','D5','D6','D7'],
 [[lab]+[num(bd[k][str(d)],1) for d in range(1,8)] for k,lab in DIMS],fs=9)

# ===== 3. ENTRE DIAS =====
H('3 Comparação entre os dias e mudança D1 para D7')
P('A Tabela 3 reúne o teste de Friedman (variação entre os sete dias, com o W de Kendall) e a comparação entre o '
 'primeiro e o último dia, com o tamanho e a magnitude do efeito, para todas as dimensões.')
fr=R['friedman']; d17=R['d1d7']
def wmag(w): return 'trivial' if w<0.1 else 'pequeno' if w<0.3 else 'moderado' if w<0.5 else 'grande'
table('Comparação das dimensões entre os dias (Friedman e W de Kendall) e entre o primeiro e o último dia (Δ, dz, magnitude).',
 ['Dimensão','χ²','p (Friedman)','W','D1','D7','Δ','dz','p (D1–D7)','Magnitude'],
 [[lab,num(fr[k]['chi'],1) if k in fr else 'n/d',pvt(fr[k]['p']) if k in fr else 'n/d',num(fr[k]['W'],2) if k in fr else 'n/d',
   num(d17[k]['d1'],1),num(d17[k]['d7'],1),c2(f"{d17[k]['delta']:+.1f}"),c2(f"{d17[k]['dz']:+.2f}"),pvt(d17[k]['p']),mag(d17[k]['dz'])] for k,lab in DIMS],
 note='W de Kendall: trivial (<0,1), pequeno (0,1–0,3), moderado (0,3–0,5), grande (>0,5); magnitude de dz: trivial (<0,2), pequeno (0,2–0,5), médio (0,5–0,8), grande (>0,8).',fs=8)

# ===== 4. PRÉ vs PÓS =====
H('4 Resposta aguda pré e pós-treino')
pp=R['prepos']
table('Comparação de todas as dimensões entre os momentos pré e pós-treino (Wilcoxon pareado, com tamanho de efeito).',
 ['Dimensão','Pré','Pós','dz','p','Magnitude'],
 [[lab,num(pp[k]['pre'],2),num(pp[k]['pos'],2),c2(f"{pp[k]['dz']:+.2f}"),pvt(pp[k]['p']),mag(pp[k]['dz'])] for k,lab in DIMS],fs=9)

# ===== 5. PERFIS =====
H('5 Perfis de humor: prevalência e migração')
mign=R['perfil_migracao_n']; est=R['perfil_estavel_n']
P(f'Dos {mign+est} atletas com perfil no primeiro e no último dia, {mign} mudaram de perfil ao longo da semana e {est} '
 f'permaneceram estáveis. A prevalência deslocou-se do predomínio do iceberg para a barbatana de tubarão (Tabela 5).')
PRO=['Iceberg','Superfície','Everest invertido','Iceberg invertido','Submerso','Barbatana tubarão']
pv=R['perfil_prev']
table('Prevalência dos perfis de humor no primeiro e no último dia (número de atletas).',
 ['Perfil','Dia 1','Dia 7'],[[p,pv['D1'].get(p,0),pv['D7'].get(p,0)] for p in PRO],fs=9)
pp2=R['perfil_por_pos']
table('Migração de perfil por posição: número de atletas que mudaram de perfil entre o primeiro e o último dia.',
 ['Posição','n','Mudaram de perfil'],[[pos,pp2[pos]['n'],pp2[pos]['mudaram']] for pos in pp2],fs=9)

# ===== 6. POR POSIÇÃO =====
H('6 Deterioração do humor por posição')
P('A Figura 2 e a Tabela 7 mostram a mudança do humor (D1 para D7) por posição de jogo, com o pico de velocidade médio '
 'de cada grupo.')
figure(f'{FG}/res_bypos.png','Mudança do humor (D1 para D7) do vigor, da fadiga e da PTH por posição de jogo.',w=14.5)
bp=R['por_pos']
table('Média semanal e mudança D1 para D7 por posição, com o pico de velocidade médio (T-car).',
 ['Posição','n','PV (km/h)','ΔVigor','ΔFadiga','ΔPTH','ΔFadiga física'],
 [[pos,bp[pos]['n'],num(bp[pos]['PV'],1),num(bp[pos]['delta_d1d7']['Vigor'],1),num(bp[pos]['delta_d1d7']['Fadiga'],1),
   num(bp[pos]['delta_d1d7']['TMD'],1),num(bp[pos]['delta_d1d7']['FadFisica'],1)] for pos in ['Goleiro','Meia','Pivô','Ponta']],
 note='Δ = mudança do escore entre o primeiro e o último dia. Meia inclui armadores; Ponta inclui alas.',fs=9)

# ===== 7. POR ATLETA =====
H('7 Deterioração individual (por atleta)')
P('A resposta foi heterogênea entre os atletas (Figura 3). Alguns apresentaram forte elevação da perturbação total do '
 'humor entre o primeiro e o último dia, ao passo que outros mantiveram ou reduziram os escores, o que reforça a '
 'recomendação de um monitoramento individualizado referenciado à linha de base de cada atleta.')
figure(f'{FG}/res_athlete.png','Mudança da perturbação total do humor (D7 menos D1) por atleta, ordenada e colorida por posição.',w=13.0)

# ===== 8. APTIDÃO x HUMOR =====
H('8 Aptidão (T-car) e humor: nível e deterioração')
P('Duas perguntas foram separadas. A primeira, se a aptidão se associa ao nível do humor na semana; a segunda, se a '
 'aptidão protege contra a mudança (deterioração) do humor ao longo da semana.')
P('Quanto ao nível, o pico de velocidade associou-se ao vigor e à fadiga física médios da semana (Tabela 8): atletas '
 'mais aptos apresentaram mais vigor e menos fadiga física. Quanto à deterioração, porém, a mudança do humor entre o '
 'primeiro e o último dia não se relacionou de forma significativa com o pico de velocidade em nenhuma dimensão '
 '(Tabela 9; Figura 4). Em outras palavras, os atletas mais treinados não sofreram menor deterioração do humor; a '
 'aptidão define o patamar do estado, e não a magnitude da sua queda ao longo da semana.')
pw=PHY['pv_week']; PVW=[('Vigor','Vigor'),('Fadiga','Fadiga'),('FadFisica','Fadiga física'),('TMD','PTH'),
    ('Depressao','Depressão'),('Tensao','Tensão'),('Raiva','Raiva'),('Confusao','Confusão')]
table('Correlação de Spearman entre o pico de velocidade no T-car e o nível médio do humor na semana.',
 ['Dimensão','ρ','p','n'],[[lab,c2(f"{pw[k]['rho']:+.2f}"),pvt(pw[k]['p']),pw[k]['n']] for k,lab in PVW],fs=9)
pd_=R['pv_vs_deterioracao']
table('Correlação de Spearman entre o pico de velocidade no T-car e a deterioração do humor (mudança D1 para D7).',
 ['Mudança (D7 - D1)','ρ','p','n'],
 [[lab,c2(f"{pd_[k]['rho']:+.2f}"),pvt(pd_[k]['p']),pd_[k]['n']] for k,lab in
  [('Vigor','ΔVigor'),('Fadiga','ΔFadiga'),('TMD','ΔPTH'),('Depressao','ΔDepressão'),('FadFisica','ΔFadiga física'),('global_abs','Deterioração global (|Δ|)')]],
 note='Nenhuma correlação foi significativa (todos p > 0,20): a aptidão não prediz a deterioração do humor na semana.',fs=9)
figure(f'{FG}/res_pv_deter.png','Relação entre o pico de velocidade no T-car e a mudança do vigor e da PTH (D7 menos D1). A ausência de inclinação significativa indica que a aptidão não prediz a deterioração do humor.',w=15.0)

OUTP='/home/user/mdlucca/Artigos/Resultados_Completos_BRUMS.docx'
doc.save(OUTP); print('SAVED',OUTP,'| Tabelas',_TN[0],'Figuras',_FN[0])
