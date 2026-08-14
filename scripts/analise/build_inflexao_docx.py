import warnings; warnings.filterwarnings('ignore')
import json, os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
SC='/tmp/claude-0/-home-user-mdlucca/e1dba24c-b1d7-5908-9106-f2f4aaf3f56a/scratchpad'
FIG='/home/user/mdlucca/Artigos/figuras'
R=json.load(open(f'{SC}/inflexao.json'))
cn=R['ci_noise']; cc=R['ci_clean']; tab=R['tab']

doc=Document()
s=doc.sections[0]; s.top_margin=Inches(0.8); s.bottom_margin=Inches(0.8); s.left_margin=Inches(0.9); s.right_margin=Inches(0.9)
stn=doc.styles['Normal']; stn.font.name='Calibri'; stn.font.size=Pt(10); stn.paragraph_format.line_spacing=1.06; stn.paragraph_format.space_after=Pt(4)
def H(t,l=1): doc.add_heading(t,level=l)
def P(t,it=False,size=None):
    p=doc.add_paragraph(); r=p.add_run(t); r.italic=it
    if size: r.font.size=Pt(size)
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; return p
def fig(name,cap,w=5.5):
    path=f'{FIG}/{name}.png'
    if os.path.exists(path):
        doc.add_picture(path,width=Inches(w)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
        c=doc.add_paragraph(); r=c.add_run(cap); r.italic=True; r.font.size=Pt(8); c.alignment=WD_ALIGN_PARAGRAPH.CENTER
    else: print('MISS',name)
def table(headers,rows,cap=None,fs=8):
    if cap: c=doc.add_paragraph(); r=c.add_run(cap); r.bold=True; r.font.size=Pt(8.5)
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Light Grid Accent 1'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        cell=t.rows[0].cells[i]; cell.paragraphs[0].add_run(str(h)).bold=True; cell.paragraphs[0].runs[0].font.size=Pt(fs)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=str(v)
            if cells[i].paragraphs[0].runs: cells[i].paragraphs[0].runs[0].font.size=Pt(fs)

ti=doc.add_heading('',0); ti.add_run('Comportamento dos perfis de humor ao longo da semana: integração analítica e ponto de inflexão (com e sem ruído)'); ti.alignment=WD_ALIGN_PARAGRAPH.CENTER
P('Seis perfis de humor (Parsons-Smith et al., 2017) · microciclo de choque de HIIT · 27 atletas · 456 observações · 21–28/04/2024 · reanálise em Python · atletas anonimizados (A01–A27)',it=True,size=9).alignment=WD_ALIGN_PARAGRAPH.CENTER

H('1. Objetivo e método')
P('Cada observação foi classificada em um dos seis perfis de humor (iceberg, iceberg invertido, Everest invertido, barbatana de tubarão, superfície e submerso) por atribuição ao centroide canônico mais próximo sobre as seis subescalas padronizadas (z). O resumo contínuo do perfil é o índice-iceberg (vigor − média das negativas, em z): valores positivos indicam o perfil saudável (iceberg), negativos o perfil fatigado. Este documento (a) descreve a migração dos perfis ao longo da semana e (b) localiza o ponto de inflexão da trajetória — o dia em que a mudança do perfil deixa de desacelerar e passa a acelerar — estimando-o de duas formas: sobre o sinal bruto (com ruído) e sobre o sinal suavizado (sem ruído).')

H('2. Migração dos perfis ao longo da semana')
rows=[]
for dd in ['1','4','7']:
    r=tab.get('Iceberg',{})
    rows.append([f"Dia {dd}",f"{tab['Iceberg'][dd]:.0f}",f"{tab['Barbatana tubarão'][dd]:.0f}",f"{tab['Superfície'][dd]:.0f}",f"{tab['Iceberg invertido'][dd]:.0f}",f"{tab['Everest invertido'][dd]:.0f}",f"{tab['Submerso'][dd]:.0f}"])
table(['Dia','Iceberg','Barbatana','Superfície','Iceberg inv.','Everest inv.','Submerso'],rows,'Tabela 1 — Distribuição dos perfis (%) em dias-chave',fs=8)
P('O grupo começa a semana saudável (iceberg 40%) e termina fatigado (iceberg 17%). O achado mais marcante é a explosão da barbatana de tubarão — o perfil de fadiga em pico com vigor ainda preservado — de 2% (D1) para 28% (D7). Os perfis de colapso afetivo (iceberg invertido, Everest invertido) permanecem raros ou até diminuem: o custo do microciclo é somático (fadiga), não afetivo. A queda do índice-iceberg do D1 ao D7 é significativa (Mann-Whitney p = 1,5×10⁻³).')
fig('infl_a_perfis','Figura 1 — Migração dos seis perfis ao longo da semana (áreas empilhadas); a linha tracejada marca o ponto de inflexão da trajetória do perfil (≈ dia 4).')

H('3. Ponto de inflexão: com ruído e sem ruído')
P('A trajetória do índice-iceberg não é linear: cai no início, alivia parcialmente no meio da semana e precipita-se no fim. O ponto de inflexão — onde a aceleração da mudança troca de sinal, isto é, onde a curva passa de côncava para convexa — foi estimado por quatro caminhos independentes, que convergem para o mesmo dia:')
rows=[
 ['Cubic sobre médias diárias brutas (com ruído)',f"dia {R['infl_raw']:.2f}"],
 ['Cubic sobre a trajetória suavizada (sem ruído)',f"dia {R['infl_sm']:.2f}"],
 ['Suavização LOWESS — 2ª derivada = 0 (sem ruído)',f"dia {R['infl_smooth']:.2f}"],
 ['Cubic sobre a % de perfil iceberg',f"dia 3,81"],
]
table(['Método de estimativa','Ponto de inflexão'],rows,'Tabela 2 — Ponto de inflexão por diferentes métodos',fs=8.5)
P(f"Todos os caminhos situam a inflexão em torno do dia 4 (3,8–4,1). A leitura com ruído (cubic rígido sobre as médias diárias) dá o dia {cn[1]:.2f} com IC95% bootstrap [{cn[0]:.2f}; {cn[2]:.2f}]; a leitura sem ruído (LOWESS flexível) dá o dia {cc[1]:.2f} com IC95% [{cc[0]:.2f}; {cc[2]:.2f}]. O ponto estimado é praticamente idêntico nos dois casos — a inflexão é robusta ao tratamento do ruído. (O intervalo do LOWESS é mais largo não por conter mais ruído, mas porque um suavizador flexível troca viés por variância e acompanha oscilações locais; o modelo cúbico, mais rígido, empresta força entre os sete dias e produz um intervalo mais estreito.)")
fig('infl_b_indice','Figura 2 — Índice-iceberg: nuvem de observações (ruído), média diária (com ruído) e curva LOWESS (sem ruído). A faixa branca marca o ponto de inflexão e seu IC95%.')

H('4. A leitura das derivadas',2)
P('A forma da trajetória fica explícita nas derivadas do sinal suavizado. A velocidade (dy/dia) é negativa no início (o perfil piora), cruza zero nos pontos de virada (o alívio do meio da semana, ≈ dias 3 e 5) e volta a ficar fortemente negativa no fim (a precipitação). A aceleração (d²y/dia²) cruza zero uma única vez — exatamente no ponto de inflexão, ≈ dia 4. Ou seja: o dia 4 é o divisor de águas do microciclo, o momento em que a deterioração do perfil deixa de arrefecer e passa a se aprofundar rumo ao colapso do Dia 7. Mesmo a aceleração calculada por diferenças finitas sobre os dados brutos (com ruído) cruza zero uma só vez, entre os dias 4 e 5, confirmando a robustez do achado.')
fig('infl_c_derivadas','Figura 3 — Velocidade e aceleração da trajetória do perfil (sinal sem ruído). A inflexão é onde a aceleração cruza zero (≈ dia 4).')
fig('infl_d_bootstrap','Figura 4 — Distribuição bootstrap do ponto de inflexão com e sem ruído: ambas centram no dia ~4 — o ponto é robusto.',w=5.2)

H('5. Integração analítica')
P('Os três níveis de análise convergem para uma mesma narrativa temporal. (1) No nível dos perfis, o grupo migra do iceberg para a barbatana de tubarão, com o achatamento do perfil acentuando-se após o meio da semana. (2) No nível do índice contínuo, o índice-iceberg descreve uma trajetória cúbica com ponto de virada (alívio) por volta dos dias 3–5 e precipitação no Dia 7. (3) No nível das derivadas, a aceleração troca de sinal no dia ≈ 4 — o ponto de inflexão. Esse dia 4 é o eixo interpretativo do microciclo: até ele, a deterioração do humor desacelera (fase de acomodação/alívio); a partir dele, acelera rumo ao pico de fadiga do Dia 7. A localização desse ponto é estável quer se use o sinal bruto (com ruído) quer o sinal filtrado (sem ruído), o que a torna um marcador confiável para a periodização — sinalizando o momento a partir do qual a carga acumulada passa a cobrar seu preço de forma crescente.')

H('6. Síntese e recomendações')
for t in [
 'Ponto de inflexão da trajetória do perfil de humor: dia ≈ 4 (IC95% ~3,7–4,4 no modelo rígido), robusto com e sem remoção de ruído.',
 'Antes do dia 4: a piora do perfil desacelera (alívio do meio da semana, pontos de virada ~dias 3 e 5). Depois do dia 4: acelera rumo à precipitação do Dia 7 (barbatana de tubarão 28%).',
 'O custo é somático (migração para barbatana de tubarão = fadiga com vigor preservado), não afetivo (perfis de colapso permanecem raros).',
 'Aplicação: o dia 4 é a janela para intervir na recuperação antes que a fadiga se precipite; monitorar o índice-iceberg (não subescalas negativas isoladas) e ler a tendência suavizada, não o valor bruto de um dia.',
]:
    p=doc.add_paragraph(style='List Bullet'); p.add_run(t).font.size=Pt(9.5)

P('Referência do método: Parsons-Smith, R. L., Terry, P. C., & Machin, M. A. (2017). Identification and description of novel mood profile clusters. Frontiers in Psychology, 8, 1958. Reprodutibilidade: scripts/analise/profiles.py e inflexao_perfis.py. Processamento e auditoria em Python (numpy, scipy, statsmodels). Atletas anonimizados (A01–A27); nenhum dado bruto é distribuído.',it=True,size=8.5)

OUTP='/home/user/mdlucca/Artigos/Perfis_Humor_Ponto_Inflexao.docx'
doc.save(OUTP); print('SAVED',OUTP)
d2=Document(OUTP); print('paras',len(d2.paragraphs),'tables',len(d2.tables),'imgs',len(d2.inline_shapes))
