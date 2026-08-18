# -*- coding: utf-8 -*-
import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
R=json.load(open('brums_desc2.json')); STAT=json.load(open('brums_stats3.json')); S4=json.load(open('brums_stats4.json'))
MS=json.load(open('model_stats.json')); MV=json.load(open('manova.json')); PHJ=json.load(open('posthoc.json'))
PV=json.load(open('pv_stats.json')); LIM=json.load(open('tcar_limiar.json')); TCD=json.load(open('tcar_desc.json'))
TCV=json.load(open('tcar_curvas.json'))
LOG=json.load(open('logistica.json'))
BG=json.load(open('bayes_growth.json')); NET=json.load(open('network.json')); MDC=json.load(open('mdc.json')); PSY=json.load(open('psychometric.json'))
L2=json.load(open('limiar2.json'))
MOD=json.load(open('moduladores.json'))
CRS=json.load(open('cross.json')); PK=json.load(open('peaks.json')); SS=json.load(open('sono_stress.json')); AL=json.load(open('alom.json')); DV=json.load(open('deriv.json')); SM=json.load(open('smooth.json')); IND=json.load(open('indiv.json')); AUD=json.load(open('audit.json')); CP=json.load(open('cross_pt.json'))
FG='/home/user/mdlucca/Artigos/figuras'
def c2(s): return str(s).replace('.',',')
doc=Document()
stl=doc.styles['Normal']; stl.font.name='Times New Roman'; stl.font.size=Pt(12)
stl.element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')
stl.paragraph_format.line_spacing=1.5; stl.paragraph_format.space_after=Pt(0)
sec=doc.sections[0]; sec.top_margin=Cm(3); sec.left_margin=Cm(3); sec.bottom_margin=Cm(2); sec.right_margin=Cm(2)
_TN=[0]; _FN=[0]
def P(t='',just=True,size=12,after=6,bold=False,ind=True):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(size); r.bold=bold
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(after)
    if just: p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if just and ind: p.paragraph_format.first_line_indent=Cm(1.25)
    return p
def RUN(pairs,after=6,ind=True):
    p=doc.add_paragraph(); p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(after)
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if ind: p.paragraph_format.first_line_indent=Cm(1.25)
    for txt,bold in pairs:
        r=p.add_run(txt); r.bold=bold; r.font.size=Pt(12)
    return p
def H(t,size=12,before=10):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(size)
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.5
def _bd(c):
    tcPr=c._tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for e_ in ['top','bottom','left','right']:
        el=OxmlElement(f'w:{e_}'); on=e_ in ('top','bottom')
        el.set(qn('w:val'),'single' if on else 'nil'); el.set(qn('w:sz'),'6'); el.set(qn('w:color'),'000000'); b.append(el)
    tcPr.append(b)
def table(cap,header,rows,fonte='Fonte: dados da pesquisa (2026).',fs=9,note=None):
    _TN[0]+=1
    p=doc.add_paragraph(); r=p.add_run('Tabela %d – %s'%(_TN[0],cap)); r.font.size=Pt(11)
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
    t=doc.add_table(rows=1,cols=len(header)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,htx in enumerate(header):
        cc=t.rows[0].cells[i]; cc.text=''; rr=cc.paragraphs[0].add_run(htx); rr.bold=True; rr.font.size=Pt(fs); rr.font.name='Times New Roman'
        cc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; _bd(cc)
    for row in rows:
        cs=t.add_row().cells
        for i,val in enumerate(row):
            cs[i].text=''; rr=cs[i].paragraphs[0].add_run(str(val)); rr.font.size=Pt(fs); rr.font.name='Times New Roman'
            cs[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if i>0 else WD_ALIGN_PARAGRAPH.LEFT; _bd(cs[i])
    if note:
        pn=doc.add_paragraph(); rn=pn.add_run(note); rn.font.size=Pt(8.5); rn.italic=True; pn.paragraph_format.space_after=Pt(0)
    pf=doc.add_paragraph(); rf=pf.add_run(fonte); rf.font.size=Pt(9); pf.paragraph_format.space_after=Pt(6)
    return _TN[0]
def figure(path,cap,w=15.0):
    _FN[0]+=1
    pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; pp.add_run().add_picture(path,width=Cm(w)); pp.paragraph_format.space_before=Pt(6)
    pc=doc.add_paragraph(); pc.alignment=WD_ALIGN_PARAGRAPH.CENTER; rc=pc.add_run('Figura %d – %s'%(_FN[0],cap)); rc.font.size=Pt(11)
    pf=doc.add_paragraph(); pf.alignment=WD_ALIGN_PARAGRAPH.CENTER; rf=pf.add_run('Fonte: elaboração dos autores (2026).'); rf.font.size=Pt(9); pf.paragraph_format.space_after=Pt(6)
    return _FN[0]
def pstr(p): return '< 0,001' if p<0.001 else '= '+c2('%.3f'%p)
mvv=lambda tab,k,f: next(x[f] for x in MV[tab]['rows'] if x['k']==k)
pr=R['prepos']; d17=R['d1d7']; prof=R['profiles']; sm=R['sample']; desc=R['desc']; PREV=MS['prev']
SH=STAT['shapiro']; IC=STAT['icc']; LP=LIM['LIM']['PVini']; FR=S4['friedman']; SENS=S4['sens']
# ordem canônica POMS/BRUMS: Tensão–Depressão–Raiva–Vigor–Fadiga–Confusão–PTH
ORD=[('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Vigor','Vigor'),('Fadiga','Fadiga'),('Confusao','Confusão'),('TMD','PTH')]

# ===== TÍTULO =====
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('DINÂMICA DO HUMOR NUM MICROCICLO PRÉ-COMPETITIVO DE HANDEBOL DE ELITE: EIXO ENERGIA–FADIGA, MIGRAÇÃO DE PERFIS E MODELAGEM DE TRAJETÓRIAS')
r.bold=True; r.font.size=Pt(13); p.paragraph_format.space_after=Pt(10)

# ===== RESUMO =====
H('RESUMO',before=2)
RUN([('Objetivo: ',True),('caracterizar e modelar a dinâmica do humor de atletas de handebol de elite ao longo de um '
 'microciclo pré-competitivo, com a comparação de cada dimensão do BRUMS entre todos os dias, a identificação dos dias '
 'de maior e menor expressão de cada estado, a classificação dos perfis de humor e a modelagem das trajetórias. ',False),
 ('Método: ',True),('%d atletas do sexo masculino responderam ao BRUMS-24 ao longo de sete dias (uma coleta de linha '
 'de base no primeiro dia e duas coletas diárias — pré e pós-treino — nos seis dias de treino), com um total de %d '
 'observações. Empregaram-se estatística descritiva, Shapiro-Wilk, Wilcoxon com tamanho de efeito, Friedman com W de '
 'Kendall, pós-teste (Tukey), ICC, MANOVA em escores T, correlação de Spearman, classificação de perfis com regressão '
 'logística da migração, modelagem polinomial das trajetórias (derivadas e cruzamento) e decomposição sinal–ruído; a '
 'robustez foi testada por reexecução após filtragem de observações atípicas. '%(sm['n'],sm['n_obs']),False),
 ('Resultados: ',True),('a deterioração concentrou-se no eixo energia–fadiga — o vigor caiu (d = %s) e a fadiga subiu '
 '(d = %s), com confirmação por MANOVA (Wilks λ = %s; p %s). O vigor foi máximo no Dia %d e a fadiga no Dia %d. Os seis '
 'perfis de humor descritos por Terry e Parsons-Smith estiveram representados, e a prevalência deslocou-se do iceberg '
 '(%s%% no baseline) para a barbatana de tubarão (%s%% no Dia 7); a regressão logística de tendência confirmou o aumento '
 'da chance do perfil de barbatana de tubarão (OR = %s por dia), sem instalação dos perfis de maior risco à saúde '
 'mental. A modelagem das trajetórias localizou o cruzamento entre vigor e fadiga na metade da semana, e a decomposição '
 'sinal–ruído revelou que o vigor e a fadiga são quase inteiramente sinal, ao passo que as dimensões negativas são '
 'dominadas por ruído. Todos os achados centrais permaneceram estáveis após a auditoria de robustez. '%(
   c2('%+.2f'%mvv('d1d7','Vigor','d')),c2('%+.2f'%mvv('d1d7','Fadiga','d')),c2('%.3f'%MV['d1d7']['wilks']),pstr(MV['d1d7']['p_mv']),
   PK['Vigor']['max_day'],PK['Fadiga']['max_day'],
   c2('%.0f'%(100*PREV['D1']['Iceberg']/PREV['n_d1'])),c2('%.0f'%(100*PREV['D7']['Barbatana tubarão']/PREV['n_d7'])),
   c2('%.2f'%LOG['migracao']['barbatana']['OR_dia'])),False),
 ('Conclusão: ',True),('o humor migrou da prontidão (perfil iceberg) para a fadiga funcional (perfil barbatana de '
 'tubarão), em um padrão compatível com sobre-esforço funcional; o vigor e a fadiga foram as dimensões mais sensíveis e '
 'sistemáticas — seletividade do eixo energia–fadiga que reproduz, num microciclo de handebol, o comportamento já '
 'documentado em outras modalidades sob carga —, o que recomenda centrar nele o monitoramento do humor, integrado a um '
 'acompanhamento multidomínio do estado do atleta.',False)],after=6)
P('Palavras-chave: humor; BRUMS; handebol; monitoramento do atleta; fadiga; perfis de humor.',size=11,after=8,ind=False)

# ===== 1 INTRODUÇÃO =====
H('1 INTRODUÇÃO')
H('1.1 Monitoramento do humor no esporte de rendimento',12,before=6)
P('O acompanhamento do estado psicológico dos atletas consolidou-se como parte essencial da gestão do treinamento no '
 'esporte de rendimento. Instrumentos de autorrelato do humor são práticos, econômicos e sensíveis às variações da carga '
 'de treino, com utilidade preditiva para o bem-estar e o desempenho esportivo (SAW; MAIN; GASTIN, 2016; LOCHBAUM et al., '
 '2021), razão pela qual documentos de consenso recomendam seu uso rotineiro para monitorar a fadiga e orientar decisões '
 'de treino e recuperação (KELLMANN et al., 2018). O interesse por essa via de monitoramento sustenta-se em uma '
 'observação recorrente: o humor capta a dimensão subjetiva da fadiga que os marcadores biológicos nem sempre revelam. '
 'Em programas de treino intensificado, a perturbação do humor pode persistir mesmo depois de indicadores fisiológicos '
 'terem regredido, dissociação que reforça o valor de um monitoramento integrado do estado do atleta '
 '(OSTAPIUK-KAROLCZUK et al., 2025).')
P('Entre esses instrumentos, a Escala de Humor de Brunel (BRUMS), versão abreviada e adaptada do Profile of Mood States '
 '(POMS), destaca-se pela rapidez de aplicação e pela solidez psicométrica, com sucessivas validações transculturais, '
 'entre elas a versão brasileira (BRAMS) (TERRY; LANE; FOGARTY, 2003; ROHLFS et al., 2008, 2023). Sua aplicação tem-se '
 'mostrado particularmente informativa em modalidades esportivas coletivas e no esporte de rendimento, contextos de '
 'elevada demanda física, emocional e interpessoal: a BRUMS tem sido empregada para monitorar o humor de atletas de '
 'elite ao longo de períodos competitivos e sua relação com o sono, o desempenho, os resultados de partida e o risco de '
 'lesão (ANDRADE et al., 2016; BRANDT; BEVILACQUA; ANDRADE, 2017; ANDRADE et al., 2020; DE MIRANDA ROHLFS et al., 2025; '
 'DO NASCIMENTO et al., 2026). Nas modalidades coletivas, em particular, a escala acompanhou o humor de seleções e '
 'equipes de elite ao longo de competições internacionais no basquetebol (BIRD et al., 2025), subsidiou a triagem do '
 'excesso de treinamento e da recuperação em atletas dessa mesma modalidade (BATTAGLINI et al., 2022) e captou a queda '
 'de vigor e a elevação de fadiga sob treino intensificado no futebol (FERREIRA et al., 2026); firma-se, assim, como um '
 'indicador de baixo custo e alta sensibilidade para a triagem do bem-estar psicológico e do risco de saúde mental em '
 'contextos de rendimento.')
P('Na prática, a escala tem sido empregada sobretudo em dois momentos: como um retrato pontual em torno da competição, '
 'associado à probabilidade de vitória e ao estado pré-jogo (BRANDT et al., 2019; DO NASCIMENTO et al., 2026), e em '
 'comparações entre fases de treino separadas por semanas ou meses, muitas vezes em paralelo a marcadores hormonais e '
 'de sono (ROUVEIX et al., 2006; ANDRADE et al., 2019). Permanece menos explorada, contudo, a dinâmica fina do humor no '
 'interior de um único microciclo — aquela que distingue a oscilação aguda de cada sessão do acúmulo ao longo dos dias '
 '—, sobretudo quando acompanhada de forma contínua e articulada à aptidão física em uma modalidade coletiva.')
H('1.2 Perfis de humor e sua aplicação em atletas de elite',12,before=6)
P('Para além dos escores isolados de cada subescala, a leitura do humor evoluiu para a identificação de perfis '
 'prototípicos. Morgan (1985) descreveu o clássico “perfil iceberg” — vigor elevado sobre dimensões negativas baixas — '
 'como assinatura de prontidão e de saúde mental positiva. Mais recentemente, análises de agrupamento (cluster) em '
 'grandes amostras formalizaram seis perfis de humor — iceberg, superfície, submerso, barbatana de tubarão, iceberg '
 'invertido e Everest invertido —, dos quais os três últimos se associam a maior risco à saúde mental e à subperformance '
 '(PARSONS-SMITH; TERRY; MACHIN, 2017; HAN; PARSONS-SMITH; TERRY, 2020). Esses perfis têm sido replicados em diferentes '
 'culturas e aplicados ao rastreamento da prontidão e do bem-estar de atletas de elite, inclusive no contexto brasileiro: '
 'em uma grande amostra de um clube de rendimento do Rio de Janeiro (898 atletas), o perfil iceberg foi o mais prevalente '
 'em avaliação momentânea, ao passo que os perfis de risco foram os menos frequentes (DE MIRANDA ROHLFS et al., 2024). A '
 'abordagem por perfis oferece uma leitura integrada e visual do estado psicológico, sensível às variações de carga e '
 'útil para identificar precocemente atletas em deterioração, com valor documentado para a saúde mental sustentável e '
 'para desfechos como lesão (TERRY et al., 2021; DE MIRANDA ROHLFS et al., 2025).')
P('Dois desses perfis são especialmente informativos para a leitura de um microciclo de carga. O iceberg — vigor '
 'elevado sobrepondo-se a dimensões negativas baixas — é a assinatura da prontidão e do bem-estar e associa-se a melhor '
 'desempenho esportivo; em meta-análise de estudos com atletas competitivos, a perturbação total do humor predisse o '
 'rendimento (efeito médio; LOCHBAUM et al., 2021). A barbatana de tubarão (shark fin), por sua vez, descreve o atleta '
 'ainda energizado, porém com a fadiga já elevada acima do vigor — e, em graus variados, com tensão e raiva acentuadas —, '
 'e configura a assinatura afetiva típica do acúmulo de treino. Como o humor é um estado transitório e sensível à carga, '
 'é de se esperar que, ao longo de uma semana intensa de treinamento, a prevalência se desloque do iceberg para a '
 'barbatana de tubarão: um “derretimento do iceberg” que traduz, no plano dos perfis, a deterioração progressiva do eixo '
 'energia–fadiga (MORGAN, 1985; HAN; PARSONS-SMITH; TERRY, 2020). Descrever esse deslocamento ao longo dos dias — e não '
 'apenas em um instantâneo — é o que permite distinguir a fadiga funcional, esperada e reversível, de uma deterioração '
 'que exigiria intervenção, e constitui um dos focos centrais deste estudo.')
H('1.3 Handebol: modalidade coletiva intermitente e de alta intensidade',12,before=6)
P('O handebol de quadra é uma modalidade coletiva de invasão, de caráter marcadamente intermitente e de alta intensidade. '
 'Ao longo da partida, ações máximas e explosivas — sprints curtos, saltos, arremessos, bloqueios, mudanças de direção e '
 'contatos físicos — alternam-se, de forma imprevisível, com períodos de recuperação incompleta, o que exige '
 'simultaneamente potência anaeróbia, capacidade aeróbia intermitente e elevada tolerância à fadiga (KARCHER; BUCHHEIT, '
 '2014; MICHALSIK; AAGAARD, 2015; WAGNER et al., 2014). É justamente a aptidão aeróbia que sustenta a capacidade de '
 'repetir e manter esforços de alta intensidade ao longo de uma partida: uma maior potência aeróbia acelera a '
 'ressíntese de fosfocreatina e a remoção de metabólitos nas pausas incompletas, atenua a queda de desempenho entre '
 'ações sucessivas e retarda a instalação da fadiga, com a preservação da qualidade das ações decisivas nos minutos finais '
 '(KARCHER; BUCHHEIT, 2014; MICHALSIK; AAGAARD, 2015). Esse padrão de esforço eleva a carga interna nos microciclos de '
 'acúmulo e repercute no estado afetivo, sobretudo na última semana de pré-temporada, quando a carga que antecede a '
 'competição se concentra — contexto no qual o monitoramento do humor se torna particularmente informativo para '
 'distinguir a fadiga esperada de uma deterioração que exija intervenção.')
H('1.4 Objetivos e hipóteses',12,before=6)
P('Apesar da ampla adoção da BRUMS, escasseiam descrições detalhadas do comportamento do humor ao longo de um microciclo '
 'pré-competitivo de handebol de elite — precisamente a fase de acumulação em que se instala o sobre-esforço funcional e '
 'na qual a triagem do estado psicofisiológico se torna mais decisiva. A lacuna é dupla: empírica, pela ausência de '
 'registros em alta resolução dentro de um único microciclo de uma modalidade coletiva de invasão; e metodológica, pela '
 'predominância de uma tradição descritiva — perfis e médias em instantâneos — que raramente modela a trajetória do '
 'humor no tempo. Diante dessa dupla lacuna, e ao unir a coleta duas vezes ao dia à modelagem contínua das trajetórias, '
 'o objetivo geral consistiu '
 'em caracterizar e modelar a dinâmica do humor desses atletas na última semana de pré-temporada, com atenção ao '
 'comportamento de cada dimensão do BRUMS, à comparação entre todos os dias, à identificação dos dias de maior e menor '
 'expressão de cada estado e à evolução dos perfis de humor. De modo específico, o estudo '
 'propôs-se a: (i) caracterizar a amostra; (ii) verificar a normalidade; '
 '(iii) descrever a resposta aguda pré → pós com tamanho de efeito; (iv) comparar as dimensões entre todos os dias; (v) '
 'confirmar as diferenças por análise multivariada (escores T); (vi) descrever o comportamento individual de cada '
 'variável e suas relações; (vii) identificar os dias de maior fadiga, vigor, tensão, raiva e depressão; (viii) '
 'classificar a evolução dos perfis de humor; e (ix) modelar a dinâmica temporal das trajetórias — '
 'taxa de variação, aceleração e cruzamento energia–fadiga —, com separação entre o sinal sistemático e o ruído amostral. '
 'Hipotetizou-se que a deterioração se concentraria no eixo energia–fadiga, com migração do perfil iceberg (prontidão) '
 'para o de barbatana de tubarão (fadiga funcional) e um cruzamento entre as trajetórias de vigor e fadiga ao longo da '
 'semana.')

# ===== 2 MATERIAIS E MÉTODOS =====
H('2 MATERIAIS E MÉTODOS')
H('2.1 Delineamento e amostra',12,before=6)
P('Estudo descritivo-comparativo, observacional e de medidas repetidas, conduzido em condições ecológicas de treinamento '
 'durante o microciclo pré-competitivo (21 a 27 de abril de 2024), com %d atletas de handebol do sexo masculino de nível '
 'competitivo, conforme os princípios éticos da Declaração de Helsinque e com consentimento informado.'%sm['n'])
H('2.2 Instrumentos',12,before=6)
P('O humor foi avaliado pela BRUMS-24 (Escala de Humor de Brunel), com 24 itens em escala de 0 (“nada”) a 4 '
 '(“extremamente”), agrupados em seis subescalas de 0 a 16 pontos (tensão, depressão, raiva, vigor, fadiga e confusão); '
 'a Perturbação Total do Humor (PTH) resume o perfil (soma das dimensões negativas menos o vigor). Para a classificação '
 'dos perfis e a análise multivariada, os escores foram convertidos em escores T (M = 50; DP = 10). O instrumento foi '
 'autoaplicado por formulário eletrônico, uma vez na linha de base e duas vezes ao dia — pré e pós-treino — nos dias de '
 'treino, o que permite separar a resposta afetiva aguda a cada sessão do acúmulo ao longo da semana.')
H('2.3 Procedimentos',12,before=6)
P('O BRUMS foi autoaplicado por formulário eletrônico ao longo de sete dias consecutivos (21 a 27 de abril de 2024). '
 'A data e o horário de cada resposta foram definidos pelo carimbo automático de registro do formulário — e não pela '
 'data informada pelo respondente —, o que garante a alocação correta de cada observação ao dia e ao momento de coleta. '
 'O primeiro dia (21/04) constituiu a avaliação de linha de base (baseline), com uma única coleta por atleta; nos seis '
 'dias de treino subsequentes (22 a 27/04) previram-se duas coletas diárias — uma ao início (pré) e outra ao final '
 '(pós) do treino. Em cada dia de treino, a primeira resposta de cada atleta foi tomada como pré e a última como pós. '
 'Por se tratar de registro em condições ecológicas, houve adesão irregular: nem todos os atletas responderam em todos '
 'os dias ou nas duas janelas, e eventuais respostas repetidas dentro de uma mesma janela foram reduzidas a um único '
 'valor (a primeira, no baseline; a primeira e a última, como pré e pós, nos dias de treino). Após essa padronização e '
 'a exclusão de um registro isolado fora da janela (29/04), o conjunto analítico totalizou %d observações válidas dos '
 '%d atletas (uma no baseline e até duas por dia de treino).'%(sm['n_obs'],sm['n']))
H('2.4 Análise estatística',12,before=6)
P('A análise seguiu uma sequência do descritivo ao inferencial. Inicialmente, empregou-se estatística descritiva (média, '
 'desvio-padrão, mediana e amplitude) para caracterizar as variáveis, e o teste de Shapiro-Wilk para verificar a '
 'normalidade — etapa que orienta a escolha entre testes paramétricos e não paramétricos. Como as distribuições violaram '
 'a normalidade, com forte concentração de escores baixos nas dimensões negativas (efeito de piso), os dados não foram '
 'transformados (NEVILL; LANE, 2007) e privilegiaram-se testes não paramétricos.')
P('Para descrever a resposta aguda ao treino, os momentos pré e pós foram comparados pelo teste de Wilcoxon para amostras '
 'pareadas, acompanhado do tamanho de efeito (d de Cohen), que expressa a magnitude da diferença independentemente do '
 'tamanho da amostra (trivial < 0,2; pequeno < 0,5; médio < 0,8; grande ≥ 0,8); a mesma lógica aplicou-se à comparação '
 'entre o primeiro e o último dia. Para verificar se cada dimensão variou ao longo dos sete dias, utilizou-se o teste de '
 'Friedman — equivalente não paramétrico da ANOVA de medidas repetidas —, com o W de Kendall como tamanho de efeito '
 '(0,1 pequeno; 0,3 moderado; 0,5 grande); quando significativo, aplicou-se o pós-teste das médias marginais de um modelo '
 'misto (correção de Tukey), que identifica entre quais dias, especificamente, houve diferença. A consistência das '
 'medidas repetidas ao longo da semana foi estimada pelo coeficiente de correlação intraclasse (ICC), interpretado como '
 'pobre (< 0,50), moderado (0,50–0,75), bom (0,75–0,90) ou excelente (> 0,90). As propriedades psicométricas das '
 'subescalas foram descritas pela consistência interna (alfa de Cronbach e correlação média entre itens), pela '
 'assimetria da distribuição e pela presença de efeito de piso, considerado relevante quando mais de 15% das '
 'observações se concentravam no menor escore possível (TERWEE et al., 2007). A partir do ICC e do desvio-padrão, '
 'derivaram-se o erro-padrão de medida (SEM = DP × √(1 − ICC)) e a mudança mínima detectável (MDC = 1,96 × √2 × SEM, a '
 '90% e a 95% de confiança), que expressa a menor variação individual distinguível do erro de medida; como referência '
 'de relevância prática, computou-se ainda a menor mudança relevante (SWC = 0,2 × desvio-padrão entre atletas).')
P('Como confirmação robusta, as comparações Dia 1 → Dia 7 e pré → pós foram reanalisadas por análise multivariada de '
 'variância (MANOVA) de medidas repetidas sobre os escores T, que testa as seis dimensões em conjunto e controla o erro '
 'de múltiplas comparações (lambda de Wilks, F e eta-quadrado parcial — η²ₚ; 0,01 pequeno; 0,06 médio; 0,14 grande); nos '
 'testes univariados de acompanhamento da MANOVA, aplicou-se o ajuste de Bonferroni ao nível de significância, com a divisão de '
 '0,05 pelas seis dimensões (α = 0,008), para controlar o erro do tipo I. A '
 'associação entre as dimensões foi quantificada pela correlação de Spearman (ρ). Os perfis de humor foram '
 'classificados a partir dos escores T das seis dimensões, e a migração de perfil ao longo da semana foi testada como '
 'tendência por regressão logística da probabilidade de cada perfil (iceberg e barbatana de tubarão) em função do dia '
 'do microciclo, com os intervalos de confiança das razões de chances obtidos por reamostragem (bootstrap, 1000 '
 'repetições) de atletas, para acomodar a estrutura de medidas repetidas. A dinâmica temporal foi modelada por meio do '
 'ajuste de uma função polinomial às médias diárias de cada dimensão, da qual se obtiveram, por derivação, a taxa de '
 'variação instantânea (derivada primeira), os pontos críticos (máximos e mínimos), o ponto de inflexão (raiz da '
 'derivada segunda) e o cruzamento entre as trajetórias de vigor e fadiga. A mesma componente suave decompôs cada '
 'trajetória em sinal e ruído, com a quantificação da razão sinal-ruído (SNR) e uma visualização suavizada das curvas. '
 'A resposta foi ainda modelada por atleta (trajetórias individuais) e reconstruída em alta resolução com 13 pontos '
 'temporais (linha de base e momentos pré e pós de cada dia de treino). Por fim, para verificar a dependência dos '
 'achados em relação a observações atípicas, todas as análises centrais foram reexecutadas após a remoção de valores '
 'multivariadamente atípicos (distância de Mahalanobis).')
P('Para aprofundar a inferência sobre a dinâmica temporal, acrescentaram-se três abordagens complementares. Primeiro, '
 'ajustou-se um modelo bayesiano multinível de curva de crescimento ao vigor e à fadiga, com o dia do microciclo '
 '(componente crônica) e o momento pré/pós (componente aguda) como efeitos fixos e interceptos e inclinações aleatórios '
 'por atleta; a estimação por amostrador de Gibbs com prioris conjugadas fracamente informativas (20 000 iterações, '
 'descarte inicial de 6 000, semente 2024) forneceu intervalos de credibilidade de 95%, a partição da variância entre '
 'atletas (correlação intraclasse) e a dispersão das inclinações individuais, com as médias a posteriori confrontadas '
 'com um modelo misto de máxima verossimilhança restrita para validação (McELREATH, 2020). Segundo, estimou-se uma rede '
 'psicométrica (modelo gráfico gaussiano) das seis dimensões por correlações parciais, no dia descansado e no dia de '
 'maior carga, com a força global (soma das arestas absolutas) comparada por teste de permutação de 5 000 reamostragens '
 '(EPSKAMP; FRIED, 2018). Terceiro, como verificação de robustez que respeita a natureza ordinal e limitada das '
 'subescalas, ajustou-se um modelo de odds proporcionais (cumulative link) do dia e do momento pré/pós sobre o vigor e a '
 'fadiga categorizados em tercis (LIDDELL; KRUSCHKE, 2018).')
P('Adotou-se nível de significância de 5% '
 '(p < 0,05), com as análises conduzidas em ambiente Python (bibliotecas pandas, NumPy, SciPy e statsmodels).')

# ===== 3 RESULTADOS =====
H('3 RESULTADOS')
H('3.1 Caracterização da amostra',12,before=6)
tp=table('Distribuição da amostra por posição de jogo (n = %d).'%sm['n'],['Posição','n','%'],
    [[k,v,c2('%.1f'%(100*v/sm['n']))] for k,v in sm['pos'].items()]+[['Total',sm['n'],'100,0']],fs=9)
def srow(lab,k):
    v=sm[k]; return [lab,c2('%.1f'%v['mean']),c2('%.1f'%v['sd']),'%s–%s'%(c2('%.1f'%v['mn']),c2('%.1f'%v['mx']))]
ta=table('Caracterização sociodemográfica e antropométrica (n = %d).'%sm['n'],['Variável','Média','DP','Mín–Máx'],
    [srow('Idade (anos)','idade'),srow('Estatura (cm)','estatura'),srow('Massa corporal (kg)','massa'),srow('Percentual de gordura (%)','pG'),srow('Experiência (anos)','exp')])
P('Participaram %d atletas do sexo masculino (idade %s ± %s anos; estatura %s ± %s cm; massa corporal %s ± %s kg), '
 'distribuídos entre as posições de jogo (Tabela %d) e com caracterização sociodemográfica e antropométrica compatível '
 'com a de handebolistas de nível competitivo (Tabela %d).'%(
   sm['n'],c2('%.1f'%sm['idade']['mean']),c2('%.1f'%sm['idade']['sd']),c2('%.1f'%sm['estatura']['mean']),c2('%.1f'%sm['estatura']['sd']),
   c2('%.1f'%sm['massa']['mean']),c2('%.1f'%sm['massa']['sd']),tp,ta))
H('3.2 Normalidade das distribuições',12,before=6)
tn=table('Teste de normalidade (Shapiro-Wilk) das dimensões do BRUMS.',['Dimensão','Estatística (W)','p'],
    [[lab,c2('%.3f'%SH[k]['W']),'< 0,001' if SH[k]['p']<0.001 else c2('%.3f'%SH[k]['p'])] for k,lab in [('Vigor','Vigor'),('Fadiga','Fadiga'),('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Confusao','Confusão')]],fs=9)
P('As seis dimensões não seguem distribuição normal (p < 0,001; Tabela %d), o que justifica os testes não paramétricos.'%tn)
H('3.3 Estatística descritiva e propriedades psicométricas das dimensões',12,before=6)
def drow(k,lab):
    v=desc[k]; return [lab,c2('%.2f'%v['mean']),c2('%.1f'%v['md']),c2('%.2f'%v['sd']),'%s–%s'%(c2('%.0f'%v['mn']),c2('%.0f'%v['mx']))]
td=table('Estatística descritiva das dimensões do BRUMS e da PTH (%d observações).'%sm['n_obs'],
    ['Dimensão','Média','Mediana','DP','Mín–Máx'],[drow(k,l) for k,l in ORD],fs=9)
P('A descritiva geral consta na Tabela %d: o vigor e a fadiga concentram as maiores médias e variabilidade.'%td)
def psyrow(k,lab):
    o=PSY[k]; return [lab,c2('%.2f'%o['alpha']),c2('%.2f'%o['miic']),c2('%.0f'%(100*o['floor0']))+'%',c2('%+.2f'%o['skew']),'sim' if o['floor_effect'] else 'não']
tpsy=table('Propriedades psicométricas das subescalas do BRUMS: consistência interna, efeito de piso e assimetria (%d observações).'%sm['n_obs'],
    ['Dimensão','α de Cronbach','r médio entre itens','% no piso (=0)','Assimetria','Efeito de piso'],
    [psyrow(k,l) for k,l in [('Vigor','Vigor'),('Fadiga','Fadiga'),('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Confusao','Confusão')]],
    note='Efeito de piso considerado presente quando > 15% das observações no menor escore (TERWEE et al., 2007). α = alfa de Cronbach; r médio entre itens = correlação inter-item média.',fs=8.5)
P('As propriedades psicométricas (Tabela %d) esclarecem por que a resposta se restringiu ao eixo energia–fadiga. O '
 'vigor e a fadiga distribuíram-se ao longo de toda a escala, com assimetria próxima de zero (%s e %s) e sem efeito de '
 'piso (%s%% e %s%% no menor escore), e reuniram consistência interna adequada a boa (α = %s e %s). As quatro dimensões '
 'negativas, ao contrário, exibiram forte efeito de piso — de %s%% (tensão) a %s%% (confusão) das observações no escore '
 'zero — e acentuada assimetria positiva, o que comprime a variância disponível. Essa compressão repercutiu na '
 'consistência interna de modo heterogêneo: a tensão e a confusão tiveram alfa reduzido (α = %s e %s), pois a baixa '
 'variância enfraquece a covariância entre itens, ao passo que a depressão e a raiva mantiveram alfa elevado '
 '(α = %s e %s) — quando de fato se elevam, os seus itens variam de forma coerente, ainda que isso ocorra em uma '
 'minoria das observações. Em suma, não é a fidedignidade da escala, mas o efeito de piso próprio de uma amostra de '
 'elite mentalmente saudável, que limita a sensibilidade das dimensões negativas à carga.'%(
   tpsy,c2('%+.2f'%PSY['Vigor']['skew']),c2('%+.2f'%PSY['Fadiga']['skew']),
   c2('%.0f'%(100*PSY['Vigor']['floor0'])),c2('%.0f'%(100*PSY['Fadiga']['floor0'])),
   c2('%.2f'%PSY['Vigor']['alpha']),c2('%.2f'%PSY['Fadiga']['alpha']),
   c2('%.0f'%(100*PSY['Tensao']['floor0'])),c2('%.0f'%(100*PSY['Confusao']['floor0'])),
   c2('%.2f'%PSY['Tensao']['alpha']),c2('%.2f'%PSY['Confusao']['alpha']),
   c2('%.2f'%PSY['Depressao']['alpha']),c2('%.2f'%PSY['Raiva']['alpha'])))
H('3.4 Diferenças entre pré e pós-treino (com tamanho de efeito)',12,before=6)
def ppr(k,lab):
    v=pr[k]; return [lab,c2('%.2f'%v['pre']),c2('%.2f'%v['pos']),c2('%+.0f'%v['pct'])+'%',pstr(v['p']),c2('%+.2f'%v['dz']),v['mag']]
tpp=table('Diferenças entre pré e pós-treino (agregado da semana; Wilcoxon e d de Cohen).',
    ['Dimensão','Pré (M)','Pós (M)','Variação (%)','p','d de Cohen','Magnitude'],[ppr(k,l) for k,l in ORD],fs=9)
P('A resposta aguda foi consistente (Tabela %d): o vigor caiu (%s%%; d = %s) e a fadiga e a PTH subiram do pré para o '
 'pós, com efeito médio.'%(tpp,c2('%.0f'%abs(pr['Vigor']['pct'])),c2('%.2f'%pr['Vigor']['dz'])))
H('3.5 Comportamento diário e comparação entre todos os dias',12,before=6)
f_traj=figure(f'{FG}/xb2_traj.png','Comportamento das dimensões do humor ao longo da semana (médias diárias; áreas sombreadas = IC95%).')
f_box=figure(f'{FG}/xb3_box.png','Diagramas de caixa (box plot) das seis dimensões do BRUMS por dia da semana.')
def ddrow(k,lab):
    dm=CRS['dec'][k]['dm']; f=FR[k]; return [lab]+[c2('%.1f'%dm[str(d)]) for d in range(1,8)]+[c2('%.1f'%f['chi']),pstr(f['p']),c2('%.2f'%f['W'])]
tdd=table('Médias diárias das dimensões do BRUMS e teste de Friedman com W de Kendall (variação entre os sete dias).',
    ['Dimensão','D1','D2','D3','D4','D5','D6','D7','χ²','p','W de Kendall'],[ddrow(k,l) for k,l in ORD],
    note='W de Kendall = tamanho de efeito do teste de Friedman (0,1 pequeno; 0,3 moderado; 0,5 grande).',fs=8)
P('As médias diárias, o teste de Friedman e o W de Kendall constam na Tabela %d (Figuras %d e %d): houve diferença '
 'significativa entre os dias para o vigor (p %s; W = %s), a fadiga (p %s; W = %s), a tensão (p %s) e a confusão '
 '(p %s), sempre com magnitude pequena a moderada, coerente com um microciclo de acúmulo dentro da faixa funcional.'%(
   tdd,f_traj,f_box,pstr(FR['Vigor']['p']),c2('%.2f'%FR['Vigor']['W']),pstr(FR['Fadiga']['p']),c2('%.2f'%FR['Fadiga']['W']),pstr(FR['Tensao']['p']),pstr(FR['Confusao']['p'])))
def icrow(k,lab):
    i=IC[k]; return [lab,c2('%.2f'%i['icc1']),c2('%.2f'%i['icck']),i['cls']]
tic=table('Consistência das medidas repetidas ao longo da semana (coeficiente de correlação intraclasse, ICC).',
    ['Dimensão','ICC(2,1)','ICC(2,k)','Consistência'],
    [icrow(k,l) for k,l in [('Vigor','Vigor'),('Fadiga','Fadiga'),('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Confusao','Confusão')]],fs=9)
P('A consistência das medidas repetidas (Tabela %d) foi moderada a boa, com os valores mais baixos na raiva e na confusão — '
 'dimensões mais reativas dia a dia.'%tic)
tmd=table('Erro de medida e limiares de mudança do vigor e da fadiga (escala 0–16): erro-padrão de medida (SEM), mudança mínima detectável (MDC) e menor mudança relevante (SWC).',
    ['Dimensão','SEM','MDC90','MDC95','SWC'],
    [[l,c2('%.1f'%MDC[k]['sem']),c2('%.1f'%MDC[k]['mdc90']),c2('%.1f'%MDC[k]['mdc95']),c2('%.1f'%MDC[k]['swc'])]
     for k,l in [('Vigor','Vigor'),('Fadiga','Fadiga')]],
    note='SEM = DP × √(1 − ICC); MDC = 1,65 (90%) ou 1,96 (95%) × √2 × SEM; SWC = 0,2 × DP entre atletas. Uma mudança individual só é interpretada como real quando excede a MDC.',fs=9)
P('Para orientar a leitura individual, derivaram-se do ICC os limiares de mudança (Tabela %d). No vigor, o erro-padrão '
 'de medida foi de %s ponto e a mudança mínima detectável, de %s (90%%) a %s (95%%) pontos; na fadiga, %s e %s a %s '
 'pontos, respectivamente. Esses valores contrastam com a menor mudança relevante (SWC = %s no vigor; %s na fadiga): a '
 'menor variação que teria significado prático é bem inferior ao erro de medida de uma leitura isolada. A implicação é '
 'metodológica e direta — no plano individual, apenas oscilações de cerca de %s a %s pontos podem ser tomadas como '
 'mudança real, ao passo que variações menores, embora possivelmente relevantes, confundem-se com o ruído de medida; '
 'por isso a inferência deste estudo apoia-se nas tendências de grupo e nos limiares, e não em leituras isoladas por '
 'atleta.'%(
   tmd,c2('%.1f'%MDC['Vigor']['sem']),c2('%.1f'%MDC['Vigor']['mdc90']),c2('%.1f'%MDC['Vigor']['mdc95']),
   c2('%.1f'%MDC['Fadiga']['sem']),c2('%.1f'%MDC['Fadiga']['mdc90']),c2('%.1f'%MDC['Fadiga']['mdc95']),
   c2('%.1f'%MDC['Vigor']['swc']),c2('%.1f'%MDC['Fadiga']['swc']),
   c2('%.1f'%MDC['Vigor']['mdc90']),c2('%.1f'%MDC['Fadiga']['mdc90'])))
def emmc(v,d): return c2('%.2f'%PHJ[v]['emm'][str(d)])
def sig1(v,d): return '' if d==1 or PHJ[v]['pairs']['1_%d'%d]['ptukey']>=0.05 else '*'
te=table('Pós-teste (médias marginais do modelo misto) por dia, com comparação de cada dia ao Dia 1.',
    ['Dia','Vigor','Fadiga','Fadiga física'],
    [['Dia %d'%d,emmc('Vigor',d)+sig1('Vigor',d),emmc('Fadiga',d)+sig1('Fadiga',d),emmc('FadFisica',d)+sig1('FadFisica',d)] for d in range(1,8)],
    note='* diferença significativa em relação ao Dia 1 (Tukey, p < 0,05).',fs=9)
f_ph=figure(f'{FG}/ph_emm.png','Trajetória diária (médias marginais) com comparação de todos os dias ao Dia 1 (* p < 0,05).')
def npairs(v): return sum(1 for kk,pp in PHJ[v]['pairs'].items() if pp['ptukey']<0.05)
P('Na comparação de todos os dias entre si (pós-teste de Tukey; Tabela %d; Figura %d), o vigor diferiu significativamente em '
 '%d dos 21 pares de dias e a fadiga em %d — sempre no sentido de piora em relação aos primeiros dias —, com a confirmação da '
 'deterioração progressiva do eixo energia–fadiga.'%(te,f_ph,npairs('Vigor'),npairs('Fadiga')))
H('3.6 Confirmação por análise multivariada (MANOVA em escores T)',12,before=6)
f_prof=figure(f'{FG}/xb5_profile_d1d7.png','Perfil de humor em escores T (M = 50; DP = 10) no Dia 1 e no Dia 7.',w=13.5)
def mvrow(x):
    return [x['lab'],c2('%.1f'%x['m1']),c2('%.1f'%x['s1']),c2('%.1f'%x['m2']),c2('%.1f'%x['s2']),c2('%.2f'%x['F']),pstr(x['p'])+('*' if x['p']<0.008 else ''),c2('%+.2f'%x['d']),c2('%.3f'%x['eta'])]
mv=MV['d1d7']
tmv=table('Comparação Dia 1 → Dia 7 das seis dimensões em escores T (MANOVA de medidas repetidas; n = %d).'%mv['n'],
    ['Dimensão','D1 M','D1 DP','D7 M','D7 DP','F','p','d','η²ₚ'],[mvrow(x) for x in mv['rows']],
    note='Wilks λ = %s; F(%d,%d) = %s; p %s; η²ₚ = %s. * p < 0,008 (α ajustado por Bonferroni para as seis dimensões).'%(c2('%.3f'%mv['wilks']),mv['df1'],mv['df2'],c2('%.2f'%mv['Fmv']),pstr(mv['p_mv']),c2('%.2f'%mv['eta_mv'])),fs=8.5)
P('A análise multivariada confirmou a diferença entre o Dia 1 e o Dia 7 (Wilks λ = %s; F(%d,%d) = %s; p %s; η²ₚ = %s). '
 'Nos testes univariados com o critério corrigido de Bonferroni (α = 0,008), apenas o vigor (d = %s) e a fadiga (d = %s) '
 'permaneceram significativos — o que evidencia a concentração do efeito no eixo energia–fadiga —, enquanto a tensão '
 '(p %s) e a confusão (p %s), significativas apenas sob α = 0,05, não resistiram à correção (Tabela %d; Figura %d). A '
 'resposta aguda pré → pós também foi multivariadamente significativa (Wilks λ = %s; p %s), o que reforça o achado.'%(
   c2('%.3f'%mv['wilks']),mv['df1'],mv['df2'],c2('%.2f'%mv['Fmv']),pstr(mv['p_mv']),c2('%.2f'%mv['eta_mv']),
   c2('%+.2f'%mvv('d1d7','Vigor','d')),c2('%+.2f'%mvv('d1d7','Fadiga','d')),pstr(mvv('d1d7','Tensao','p')),pstr(mvv('d1d7','Confusao','p')),tmv,f_prof,c2('%.3f'%MV['prepos']['wilks']),pstr(MV['prepos']['p_mv'])))
H('3.7 Dias de maior expressão de cada estado de humor',12,before=6)
def pkrow(k,lab):
    p=PK[k]; return [lab,'Dia %d'%p['max_day'],c2('%.2f'%p['max_val']),'Dia %d'%p['min_day'],c2('%.2f'%p['min_val'])]
tpk=table('Dia de maior e de menor expressão de cada dimensão do humor (médias diárias).',
    ['Dimensão','Dia de maior valor','Valor','Dia de menor valor','Valor'],[pkrow(k,l) for k,l in ORD],fs=9)
P('A Tabela %d sintetiza os dias de pico: o vigor foi máximo no Dia %d e mínimo no Dia %d; a fadiga foi máxima no Dia %d; '
 'a tensão, máxima no Dia %d; a raiva, máxima no Dia %d; e a depressão, máxima no Dia %d. Em síntese, o início da semana '
 'reúne maior prontidão (vigor, tensão e confusão mais altos) e o final concentra a fadiga e a raiva.'%(
   tpk,PK['Vigor']['max_day'],PK['Vigor']['min_day'],PK['Fadiga']['max_day'],PK['Tensao']['max_day'],PK['Raiva']['max_day'],PK['Depressao']['max_day']))
H('3.8 Comportamento individual de cada variável e suas relações',12,before=6)
P('As Figuras %d a %d apresentam, para cada dimensão, o comportamento ao longo da semana com a média diária, a banda de '
 'confiança de 95%% (área sombreada), os diagramas de caixa por dia e o efeito do microciclo, o que permite visualizar '
 'individualmente como cada estado de humor evolui. Nessas figuras, o efeito Dia 1 → Dia 7 é expresso pelo dz (mudança '
 'padronizada intraindividual), numericamente idêntico ao d relatado na análise multivariada (Tabela 9), e o valor de p '
 'é o do teste de Friedman apresentado na Tabela 6.'%(_FN[0]+1,_FN[0]+6))
for k,fn,lab in [('Tensao','xb4_v_Tensao.png','Tensão'),('Depressao','xb4_v_Depressao.png','Depressão'),('Raiva','xb4_v_Raiva.png','Raiva'),('Vigor','xb4_v_Vigor.png','Vigor'),('Fadiga','xb4_v_Fadiga.png','Fadiga'),('Confusao','xb4_v_Confusao.png','Confusão')]:
    figure(f'{FG}/{fn}','%s ao longo da semana: média diária (banda = IC95%%), diagramas de caixa por dia e efeito Dia 1 → Dia 7.'%lab,w=12.5)
PA=STAT['pairs']; FP=STAT['focusp']; sig=[x for x in PA if x['p']<0.05]
tcorr=table('Correlação de Spearman entre as dimensões do BRUMS com associação significativa (n = %d atletas).'%sm['n'],
    ['Par de dimensões','ρ','p'],[['%s × %s'%(x['a'],x['b']),c2('%+.2f'%x['rho']),pstr(x['p'])] for x in sig],fs=9)
P('Quanto às relações entre as dimensões (Tabela %d), as dimensões negativas associam-se entre si e a fadiga relaciona-se '
 'com a depressão e a raiva; o vigor mantém-se relativamente independente. Esse acoplamento depende da carga: no dia de '
 'maior vigor a depressão é independente da fadiga (ρ = %s), mas no dia de maior fadiga torna-se forte (ρ = %s).'%(
   tcorr,c2('%+.2f'%FP['D1']['Depressao']['rho_fad']),c2('%+.2f'%FP['D7']['Depressao']['rho_fad'])))
H('3.9 Classificação dos perfis de humor',12,before=6)
P('A classificação de cada observação em um dos seis perfis de humor oferece uma leitura integrada e visual do estado '
 'psicológico da equipe. Ao longo do microciclo, os seis perfis estiveram representados na amostra, com reorganização '
 'de prevalência entre o baseline e o último dia de treino.')
f_prev=figure(f'{FG}/xb5_prev.png','Distribuição (%) dos seis perfis de humor no baseline (Dia 1) e no Dia 7.',w=13.5)
PROFR=[('Iceberg','Iceberg'),('Everest invertido','Everest invertido'),('Iceberg invertido','Iceberg invertido'),('Submerso','Submerso'),('Barbatana tubarão','Barbatana de tubarão'),('Superfície','Superfície')]
def prow(p,lab):
    d1=PREV['D1'][p]; d7=PREV['D7'][p]; return [lab,'%d (%s)'%(d1,c2('%.1f'%(100*d1/PREV['n_d1']))),'%d (%s)'%(d7,c2('%.1f'%(100*d7/PREV['n_d7'])))]
tprev=table('Distribuição dos perfis de humor no baseline e no último dia de treino: n (%).',
    ['Perfil','Baseline n (%)','Dia 7 n (%)'],[prow(p,lab) for p,lab in PROFR],fs=9)
P('A prevalência deslocou-se do iceberg (%s%% no baseline) para a barbatana de tubarão (%s%% no Dia 7), com o iceberg '
 'em recuo para %s%% (Figura %d; Tabela %d). Essa reorganização categórica, contudo, não alcançou significância '
 'estatística (χ² = %s; p %s) — resultado esperado dado o pequeno número de observações distribuído por seis perfis, '
 'que reduz a contagem esperada por célula e a potência do teste qui-quadrado; a mudança é, portanto, uma tendência '
 'descritiva, coerente com a queda multivariada do vigor e a elevação da fadiga.'%(
   c2('%.0f'%(100*PREV['D1']['Iceberg']/PREV['n_d1'])),c2('%.0f'%(100*PREV['D7']['Barbatana tubarão']/PREV['n_d7'])),c2('%.0f'%(100*PREV['D7']['Iceberg']/PREV['n_d7'])),f_prev,tprev,
   c2('%.2f'%PREV['chi']),pstr(PREV['p'])))
f_clu=figure(f'{FG}/xb6_clusters.png','Perfis de humor (clusters) identificados na amostra, representados em escores T (M = 50; DP = 10) nas seis dimensões; percentuais entre parênteses.',w=14.5)
P('A forma de cada cluster identificado na amostra (Figura %d) reproduz a taxonomia consagrada dos seis perfis de humor '
 '(PARSONS-SMITH; TERRY; MACHIN, 2017): o iceberg, com vigor elevado sobre dimensões negativas baixas; a barbatana de '
 'tubarão, com pico isolado de fadiga; o submerso, com todas as dimensões abaixo da média; e o Everest invertido, com '
 'todas as negativas elevadas. Comparada à distribuição de referência de uma grande amostra do mesmo contexto brasileiro '
 '(898 atletas de elite e de base; DE MIRANDA ROHLFS et al., 2024) — na qual, em avaliação momentânea, o iceberg foi o '
 'perfil mais prevalente (34,3%%) e os perfis de risco os menos prevalentes (barbatana de tubarão 11,6%%; iceberg '
 'invertido 6,2%%; Everest invertido 2,7%%) —, a nossa amostra parte de um predomínio de iceberg semelhante no Dia 1 '
 '(%s%%), mas, ao final do microciclo, eleva marcadamente a barbatana de tubarão (%s%% no Dia 7), muito acima do patamar '
 'de referência. Isso evidencia o efeito específico do acúmulo de carga da última semana de pré-temporada sobre o perfil '
 'de fadiga, sem instalação relevante dos perfis de maior risco à saúde mental (Everest invertido e submerso).'%(
   f_clu,c2('%.0f'%(100*PREV['D1']['Iceberg']/PREV['n_d1'])),c2('%.0f'%(100*PREV['D7']['Barbatana tubarão']/PREV['n_d7']))))
# --- regressão logística: migração de perfil ~ dia (tendência) ---
MIG=LOG['migracao']
def migrow(key,lab):
    f=MIG[key]; return [lab,'%s [%s–%s]'%(c2('%.2f'%f['OR_dia']),c2('%.2f'%f['OR_lo']),c2('%.2f'%f['OR_hi'])),
        c2('%.2f'%f['p_d1']),c2('%.2f'%f['p_d7']),c2('%.2f'%f['auc']),c2('%.3f'%f['mcfadden'])]
tmig=table('Regressão logística da probabilidade de perfil de humor em função do dia do microciclo (tendência): razão de chances por dia, probabilidade prevista no Dia 1 e no Dia 7, discriminação e ajuste.',
    ['Perfil (desfecho)','OR por dia [IC95%]','P(Dia 1)','P(Dia 7)','AUC','Pseudo-R²'],
    [migrow('iceberg','Iceberg'),migrow('barbatana','Barbatana de tubarão')],
    note='OR = razão de chances (odds ratio) por incremento de um dia; IC95% por reamostragem (bootstrap) de atletas; Pseudo-R² de McFadden. OR > 1 indica aumento da chance do perfil ao longo da semana.',fs=8.5)
mb=MIG['barbatana']; mi=MIG['iceberg']
P('Para testar a migração de perfil como uma tendência ao longo da semana — e não pela via categórica de baixa potência '
 'do qui-quadrado —, ajustou-se uma regressão logística da probabilidade de cada perfil em função do dia do microciclo '
 '(Tabela %d). A chance do perfil de barbatana de tubarão cresceu de forma consistente a cada dia (OR = %s por dia; '
 'IC95%% %s–%s), com elevação da probabilidade prevista de %s no Dia 1 para %s no Dia 7 — um aumento estatisticamente '
 'sustentado (o IC95%% do OR exclui 1), que dá suporte inferencial à migração que o teste categórico não captara. Em '
 'espelho, a chance do perfil iceberg declinou ao longo da semana, embora sem significância (OR = %s por dia; IC95%% '
 '%s–%s). O contraste confirma, no plano dos perfis, o mesmo eixo energia–fadiga: à medida que a semana avança, a '
 'prontidão cede lugar à assinatura de fadiga.'%(
   tmig,c2('%.2f'%mb['OR_dia']),c2('%.2f'%mb['OR_lo']),c2('%.2f'%mb['OR_hi']),
   c2('%.2f'%mb['p_d1']),c2('%.2f'%mb['p_d7']),
   c2('%.2f'%mi['OR_dia']),c2('%.2f'%mi['OR_lo']),c2('%.2f'%mi['OR_hi'])))
# ----- 3.10 Modelagem polinomial: derivadas e taxa de variação -----
H('3.10 Modelagem polinomial das trajetórias: derivadas e taxa de variação',12,before=6)
def dvrow(k,l):
    v=DV['vars'][k]
    infl=', '.join('%d'%round(x) for x in v['infl']) or '—'
    return [l,c2('%.2f'%v['r2']),c2('%+.2f'%v['taxa_media']),c2('%+.2f'%v['dP1']),c2('%+.2f'%v['dP7']),
            c2('%+.2f'%v['d2P1']),c2('%+.2f'%v['d2P7']),infl]
DVORD=[('Vigor','Vigor'),('Fadiga','Fadiga'),('TMD','PTH')]   # eixo energia–fadiga (dimensões sem efeito de piso)
tdv=table('Modelagem polinomial (grau %d) das médias diárias do eixo energia–fadiga: qualidade do ajuste (R²), taxa de variação média, taxa instantânea (P′) e aceleração (P″) nas bordas, e ponto de inflexão.'%DV['grau'],
    ['Dimensão','R²','Taxa média/dia','P′(D1)','P′(D7)','P″(D1)','P″(D7)','Inflexão (dia)'],
    [dvrow(k,l) for k,l in DVORD],
    note='P′ = derivada primeira (taxa de variação); P″ = derivada segunda (aceleração/concavidade); inflexão = raiz de P″ = 0. As dimensões negativas, com efeito de piso e ajuste fraco, foram omitidas.',fs=8.5)
f_dv=figure(f'{FG}/deriv_poly.png','Ajuste polinomial P(t) das médias diárias de vigor e fadiga, com a derivada P′(t) (taxa de variação) e o ponto de inflexão.',w=15.0)
VG=DV['vars']['Vigor']; FD=DV['vars']['Fadiga']
P('Para caracterizar formalmente a dinâmica temporal, ajustou-se uma função polinomial de grau %d às médias diárias de '
 'cada dimensão e derivaram-se as respectivas taxas de variação (Tabela %d; Figura %d). A modelagem restringiu-se ao '
 'eixo energia–fadiga (vigor, fadiga e PTH), onde os ajustes foram excelentes (vigor R² = %s; fadiga R² = %s); as '
 'dimensões negativas, próximas do piso e com variância reduzida, ajustaram-se mal e foram omitidas da modelagem. O '
 'vigor caiu a uma taxa '
 'média de %s ponto/dia, com a maior velocidade de queda no início da semana (P′ no Dia 1 = %s ponto/dia) e '
 'desaceleração progressiva, ao passo que a fadiga subiu em espelho (taxa média = %s ponto/dia; P′ no Dia 1 = %s). A '
 'derivada segunda localizou um ponto de inflexão em torno do dia %s para ambas — o momento em que a taxa de '
 'deterioração deixa de acelerar —, o que fornece uma leitura quantitativa do ritmo do desgaste e sinaliza a metade do '
 'microciclo como o marco a partir do qual a resposta afetiva muda de regime.'%(
   DV['grau'],tdv,f_dv,c2('%.2f'%VG['r2']),c2('%.2f'%FD['r2']),
   c2('%.2f'%VG['taxa_media']),c2('%+.2f'%VG['dP1']),c2('%+.2f'%FD['taxa_media']),c2('%+.2f'%FD['dP1']),
   ('%d'%round(VG['infl'][0]))))
cVF=DV['couple'].get('Vigor|Fadiga'); cFT=DV['couple'].get('Fadiga|TMD'); cVT=DV['couple'].get('Vigor|TMD')
P('A derivada segunda (P″) descreve a aceleração da mudança e a concavidade da trajetória. No vigor, P″ é positiva no '
 'início (%s) — a queda, embora acentuada, desacelera rumo ao meio da semana — e torna-se negativa ao final (%s), quando '
 'a curva volta a se inclinar para o mínimo do Dia 7; a fadiga apresenta o padrão espelhado (P″ = %s no Dia 1; %s no Dia '
 '7). A troca de sinal de P″ no dia ≈ 4 confirma, formalmente, a inflexão sincronizada das duas dimensões. Quanto ao '
 'comportamento conjunto, as curvas de taxa de variação (P′) do vigor e da fadiga são praticamente imagens especulares: '
 'a correlação entre elas é de %s, o que indica que a perda de energia e o ganho de fadiga aceleram e desaceleram em '
 'uníssono ao longo do microciclo. Ambas acoplam-se fortemente à Perturbação Total do Humor (fadiga–PTH = %s; '
 'vigor–PTH = %s), o que confirma, no plano das derivadas, que energia e fadiga operam como um eixo único e integrado — '
 'o mesmo eixo que domina toda a resposta afetiva à carga.'%(
   c2('%+.2f'%VG['d2P1']),c2('%+.2f'%VG['d2P7']),c2('%+.2f'%FD['d2P1']),c2('%+.2f'%FD['d2P7']),
   c2('%+.2f'%cVF),c2('%+.2f'%cFT),c2('%+.2f'%cVT)))
dayf=lambda x: '%d'%round(x)
def exrow(k,l):
    v=DV['vars'][k]
    return [l,'%s (Dia %s)'%(c2('%.1f'%v['fmax_val']),dayf(v['fmax_day'])),
            '%s (Dia %s)'%(c2('%.1f'%v['fmin_val']),dayf(v['fmin_day'])),
            '%s (Dia %s)'%(c2('%+.2f'%v['dmax_val']),dayf(v['dmax_day'])),
            '%s (Dia %s)'%(c2('%+.2f'%v['dmin_val']),dayf(v['dmin_day']))]
def exrowr(k,l):
    return exrow(k,l)+[c2('%.2f'%DV['vars'][k]['r2'])]
tex=table('Extremos da função ajustada e limites da derivada, por dimensão: valor máximo e mínimo modelados (com o dia) e os limites da taxa de variação P′ — subida máxima e queda máxima.',
    ['Dimensão','Valor máx. (dia)','Valor mín. (dia)','P′ máx./subida (dia)','P′ mín./queda (dia)','R²'],
    [exrowr(k,l) for k,l in ORD],
    note='Valores e taxas extraídos da função polinomial P(t) no intervalo [Dia 1; Dia 7]. Para as dimensões negativas (baixo R²; efeito de piso), os extremos modelados são apenas indicativos. Os máximos/mínimos observados constam na tabela de dias de pico (seção 3.7).',fs=8.5)
P('Os extremos das funções ajustadas e os limites de suas derivadas (Tabela %d) delimitam quantitativamente a resposta. '
 'O vigor varia entre um máximo de %s pontos no Dia %s e um mínimo de %s no Dia %s, com a taxa de variação confinada '
 'entre %s ponto/dia (queda máxima, no Dia %s) e %s ponto/dia (leve recuperação em torno do meio da semana); a fadiga '
 'percorre o intervalo inverso (mínimo de %s no Dia %s; máximo de %s no Dia %s), com taxa entre %s e %s ponto/dia. Esses '
 'limites confirmam que a maior velocidade de mudança ocorre no início do microciclo e que o sistema afetivo opera '
 'dentro de uma faixa de variação bem definida.'%(
   tex,
   c2('%.1f'%VG['fmax_val']),dayf(VG['fmax_day']),c2('%.1f'%VG['fmin_val']),dayf(VG['fmin_day']),
   c2('%+.2f'%VG['dmin_val']),dayf(VG['dmin_day']),c2('%+.2f'%VG['dmax_val']),
   c2('%.1f'%FD['fmin_val']),dayf(FD['fmin_day']),c2('%.1f'%FD['fmax_val']),dayf(FD['fmax_day']),
   c2('%+.2f'%FD['dmin_val']),c2('%+.2f'%FD['dmax_val'])))

f_cx=figure(f'{FG}/crossover.png','Cruzamento energia–fadiga: curvas suavizadas de vigor e fadiga com os pontos exatos de interseção (× ) e a área entre elas. O vigor parte muito acima e a fadiga o ultrapassa ao longo da semana.',w=15.5)
cds=CP['cross_days']; cd1=('%d'%round(cds[0])) if cds else '—'; cdl=('%d'%round(cds[-1])) if cds else '—'
P('A sobreposição das curvas suavizadas de vigor e fadiga permite localizar o momento exato de sua inversão (Figura %d). '
 'No primeiro dia, o vigor supera a fadiga por ampla margem (%s pontos); essa vantagem colapsa já no Dia 2, e as duas '
 'dimensões passam a se cruzar repetidamente na região central da semana — os pontos de interseção das funções ocorrem '
 'nos dias %s — reflexo de um período em que energia e fadiga permanecem estatisticamente entrelaçadas. A partir do último '
 'cruzamento (Dia %s), a fadiga assume e se distancia do vigor, e encerra a semana %s pontos acima. Esse cruzamento '
 'energia–fadiga é a assinatura gráfica da migração do perfil de prontidão para o perfil de fadiga.'%(
   f_cx,c2('%.1f'%CP['gap_d1']),', '.join('%d'%round(x) for x in cds),cdl,c2('%.1f'%abs(CP['gap_d7']))))

# ----- 3.11 Decomposição sinal-ruído e suavização das trajetórias -----
H('3.11 Decomposição sinal–ruído e suavização das trajetórias',12,before=6)
def smrow(k,l):
    v=SM['vars'][k]; snr='%.1f'%v['snr'] if v['snr']<900 else '—'
    return [l,c2('%.0f'%v['signal_pct'])+'%',c2('%.0f'%v['noise_pct'])+'%',c2(snr),c2('%+.1f'%v['snr_db'])]
tsm=table('Decomposição sinal–ruído das trajetórias diárias: proporção de sinal (variância explicada pela componente suave) e de ruído (resíduo), razão sinal-ruído (SNR) e SNR em decibéis.',
    ['Dimensão','Sinal','Ruído','SNR','SNR (dB)'],
    [smrow(k,l) for k,l in ORD],
    note='Sinal = componente suave (polinômio ajustado); ruído = resíduo em torno do sinal. SNR = variância do sinal / variância do ruído.',fs=9)
f_sm=figure(f'{FG}/smooth_signal.png','Trajetórias diárias suavizadas: sinal (curva) sobreposto às médias diárias observadas (pontos), com separação entre a tendência e oruído. A) eixo energia–fadiga; B) subescalas negativas.',w=15.0)
sV=SM['vars']['Vigor']; sF=SM['vars']['Fadiga']; sD=SM['vars']['Depressao']
P('Para separar a tendência sistemática (sinal) do ruído de amostragem e obter uma leitura visual mais limpa, cada '
 'trajetória diária foi decomposta em uma componente suave — a função polinomial ajustada — e um resíduo (Tabela %d; '
 'Figura %d). A suavização é altamente informativa no eixo energia–fadiga: o vigor e a fadiga são quase inteiramente '
 'sinal (%s%% e %s%% da variância; razão sinal-ruído de %s e %s, ou %s e %s dB), de modo que suas curvas suavizadas '
 'revelam com nitidez o cruzamento energia–fadiga em torno da metade da semana. Nas dimensões negativas, próximas do '
 'piso, a fração de ruído é muito maior — na depressão, o ruído (%s%%) supera o sinal (%s%%; SNR = %s), o que explica '
 'por que suas oscilações diárias não devem ser sobreinterpretadas. Em termos aplicados, a decomposição indica que o '
 'monitoramento deve priorizar o vigor e a fadiga, cujo sinal é robusto, e tratar as pequenas variações das dimensões '
 'negativas com a devida cautela.'%(
   tsm,f_sm,c2('%.0f'%sV['signal_pct']),c2('%.0f'%sF['signal_pct']),c2('%.1f'%sV['snr']),c2('%.1f'%sF['snr']),
   c2('%+.1f'%sV['snr_db']),c2('%+.1f'%sF['snr_db']),
   c2('%.0f'%sD['noise_pct']),c2('%.0f'%sD['signal_pct']),c2('%.1f'%sD['snr'])))

# ----- 3.12 Modelagem individual (por atleta) -----
H('3.12 Modelagem individual: trajetórias por atleta',12,before=6)
paV=IND['perath']['Vigor']; paF=IND['perath']['Fadiga']; cov=IND['cover']
f_ind=figure(f'{FG}/indiv_spaghetti.png','Trajetórias individuais de vigor e fadiga por atleta (linhas finas) e a média do grupo (linha grossa), que expõem aheterogeneidade da resposta.',w=15.0)
P('Além da tendência do grupo, a resposta foi modelada por atleta, com o ajuste, para cada participante, de sua própria '
 'trajetória ao longo da semana (Figura %d). A cobertura permitiu a modelagem individual: %d dos %d atletas '
 'responderam nos sete dias e %d têm dados suficientes (≥ 4 dias) para um ajuste polinomial individual; para todos foi '
 'possível estimar a taxa de variação individual. Em média, o vigor caiu %s ponto/dia (DP %s) e a fadiga subiu %s '
 'ponto/dia (DP %s), mas com nítida heterogeneidade entre atletas: %s%% apresentaram a queda esperada de vigor e %s%% '
 'o aumento esperado de fadiga, enquanto os demais mantiveram-se estáveis ou responderam na direção oposta. Essa '
 'dispersão das inclinações individuais mostra que a média do grupo convive com respondedores e não respondedores, '
 'o que reforça o valor do monitoramento individualizado — a mesma carga produz trajetórias afetivas distintas.'%(
   f_ind,cov['n7'],cov['ntot'],paV['n_cub'],
   c2('%+.2f'%paV['mean']),c2('%.2f'%paV['sd']),c2('%+.2f'%paF['mean']),c2('%.2f'%paF['sd']),
   c2('%.0f'%paV['resp_pct']),c2('%.0f'%paF['resp_pct'])))

# ----- 3.13 Trajetória em alta resolução (pré/pós) -----
H('3.13 Trajetória em alta resolução: 13 pontos pré/pós',12,before=6)
f_p13=figure(f'{FG}/pts13_curve.png','Trajetória do grupo em 13 pontos temporais (baseline + pré e pós-treino de cada dia), com a dinâmica intradiária além da média diária.',w=15.5)
P('Com base nas duas coletas diárias dos dias de treino, a trajetória do grupo foi reconstruída em alta resolução, '
 'com 13 pontos temporais — a linha de base e os momentos pré e pós de cada um dos seis dias de treino (Figura %d). '
 'Essa resolução revela a dinâmica intradiária que a média diária mascara: em cada dia de treino, o vigor tende a cair '
 'e a fadiga (e a PTH) a subir do pré para o pós, com recuperação parcial na manhã seguinte — um padrão de dente de '
 'serra sobreposto à tendência semanal de queda de energia e acúmulo de fadiga. A leitura de alta resolução, portanto, '
 'confirma que a deterioração semanal resulta da soma de perturbações agudas incompletamente revertidas entre as '
 'sessões.'%f_p13)

# ----- 3.14 Auditoria de robustez (filtragem de ruído) -----
H('3.14 Auditoria de robustez: filtragem de sinal e ruído',12,before=6)
def audrow(m,lab,fmt='%+.2f'):
    r=next((x for x in AUD['rows'] if x['metric']==m),None)
    if not r: return [lab,'—','—']
    return [lab,c2(fmt%r['raw']),c2(fmt%r['filt'])]
taud=table('Auditoria de robustez: principais resultados antes e depois da remoção de %d observações atípicas (%.1f%% do total) por distância de Mahalanobis.'%(AUD['n_out'],100*AUD['n_out']/AUD['n_total']),
    ['Métrica','Dados brutos','Dados filtrados'],
    [audrow('dz_Vigor','Vigor D1→D7 (dz)'),audrow('dz_Fadiga','Fadiga D1→D7 (dz)'),
     audrow('manova_wilks','MANOVA (Wilks λ)','%.3f'),
     audrow('friedman_Vigor','Vigor Friedman (p)','%.4f'),audrow('friedman_Fadiga','Fadiga Friedman (p)','%.4f'),
     audrow('rho_Fadiga_FadFisica','ρ fadiga × fadiga física','%+.2f')],
    note='Filtro: observações com distância de Mahalanobis nas seis subescalas acima do limiar de p < 0,001.',fs=9)
P('Por fim, para verificar se as conclusões dependem de ruído amostral, todas as análises centrais foram reexecutadas '
 'após filtragem: identificaram-se e removeram-se %d observações multivariadamente atípicas (%s%% do total) e os '
 'principais resultados foram recomputados (Tabela %d). A filtragem praticamente não alterou nada — o tamanho de efeito '
 'da queda de vigor (dz de %s para %s) e do aumento de fadiga, a MANOVA (Wilks λ de %s para %s), os testes de Friedman '
 'e as correlações-chave mantiveram magnitude, direção e significância. Essa estabilidade confirma que os achados do '
 'estudo são robustos e refletem sinal sistemático, e não artefatos de observações ruidosas.'%(
   AUD['n_out'],c2('%.1f'%(100*AUD['n_out']/AUD['n_total'])),taud,
   c2('%+.2f'%next(x for x in AUD['rows'] if x['metric']=='dz_Vigor')['raw']),
   c2('%+.2f'%next(x for x in AUD['rows'] if x['metric']=='dz_Vigor')['filt']),
   c2('%.3f'%next(x for x in AUD['rows'] if x['metric']=='manova_wilks')['raw']),
   c2('%.3f'%next(x for x in AUD['rows'] if x['metric']=='manova_wilks')['filt'])))

H('3.15 Modelagem bayesiana multinível: componentes aguda e crônica e heterogeneidade individual',12,before=6)
def bg(v,k,f='%+.2f'): b=BG[v][k]; return c2(f%b['mean']),c2(f%b['lo']),c2(f%b['hi'])
tbg=table('Modelo bayesiano multinível de curva de crescimento para o vigor e a fadiga: estimativas a posteriori (média e intervalo de credibilidade de 95%).',
    ['Parâmetro','Vigor','Fadiga'],
    [['Efeito agudo — pré→pós','%s [%s; %s]'%bg('Vigor','b_pos'),'%s [%s; %s]'%bg('Fadiga','b_pos')],
     ['Efeito crônico — por dia','%s [%s; %s]'%bg('Vigor','b_dia'),'%s [%s; %s]'%bg('Fadiga','b_dia')],
     ['Correlação intraclasse (ICC)','%s [%s; %s]'%bg('Vigor','icc','%.2f'),'%s [%s; %s]'%bg('Fadiga','icc','%.2f')],
     ['DP das inclinações individuais','%s [%s; %s]'%bg('Vigor','sd_slope','%.2f'),'%s [%s; %s]'%bg('Fadiga','sd_slope','%.2f')]],
    note='Estimação por amostrador de Gibbs (semente 2024); intervalos de credibilidade que excluem zero indicam efeito consistente. ICC = proporção da variância atribuível a diferenças entre atletas.',fs=9)
P('Um modelo bayesiano multinível de curva de crescimento decompôs formalmente a resposta em uma componente aguda '
 '(pré→pós, intrassessão) e uma componente crônica (deriva entre dias), com trajetórias específicas de cada atleta '
 '(Tabela %d). Ambas as componentes foram consistentes: para o vigor, o choque agudo (%s pontos por sessão; IC95%% %s a '
 '%s) superou em cerca de quatro vezes a deriva crônica (%s por dia; IC95%% %s a %s), e um padrão simétrico repetiu-se '
 'na fadiga (agudo %s; IC95%% %s a %s | crônico %s; IC95%% %s a %s) — todos os intervalos de credibilidade excluíram o '
 'zero. As médias a posteriori praticamente coincidiram com as do modelo misto de máxima verossimilhança restrita, o '
 'que confirma a estabilidade da estimação. A partição da variância revelou naturezas distintas: o vigor comportou-se '
 'como traço, com a maior parte da variância entre atletas (ICC = %s), ao passo que a fadiga foi mais um estado, com '
 'variância predominantemente intraindividual (ICC = %s). Sobretudo, a dispersão das inclinações individuais foi '
 'apreciável (DP = %s no vigor; %s na fadiga), da ordem da própria deriva média: embora o grupo derive na mesma '
 'direção, a velocidade da deterioração difere entre atletas — heterogeneidade que o modelo quantifica e que respalda '
 'a leitura individualizada do monitoramento.'%(
   tbg,bg('Vigor','b_pos')[0],bg('Vigor','b_pos')[1],bg('Vigor','b_pos')[2],
   bg('Vigor','b_dia')[0],bg('Vigor','b_dia')[1],bg('Vigor','b_dia')[2],
   bg('Fadiga','b_pos')[0],bg('Fadiga','b_pos')[1],bg('Fadiga','b_pos')[2],
   bg('Fadiga','b_dia')[0],bg('Fadiga','b_dia')[1],bg('Fadiga','b_dia')[2],
   bg('Vigor','icc','%.2f')[0],bg('Fadiga','icc','%.2f')[0],
   bg('Vigor','sd_slope','%.2f')[0],bg('Fadiga','sd_slope','%.2f')[0]))
f_rede=figure(f'{FG}/rede_humor.png','Redes de correlações parciais (modelo gráfico gaussiano) das seis dimensões do BRUMS no dia descansado (Dia 1) e no dia de maior carga (Dia 7). Arestas vermelhas: associação parcial positiva; azuis: negativa; a espessura é proporcional à magnitude. Vigor e fadiga destacados em laranja.',w=15.5)
P('A rede psicométrica das seis dimensões (Figura %d) tornou explícita a reorganização do humor sob carga. A mudança '
 'foi seletiva, e não global: a força total da rede não aumentou do dia descansado ao dia de maior carga (%s → %s; '
 'teste de permutação p = %s), mas a aresta entre fadiga e raiva, quase nula no início da semana, tornou-se a mais '
 'forte do painel ao final (correlação parcial %s → %s). Em outras palavras, sob carga acumulada, é especificamente a '
 'irritabilidade que se acopla à exaustão — refinamento que a leitura por correlações simples não isola e que traduz, '
 'no plano da rede, o “fechamento” do perfil de fadiga.'%(
   f_rede,c2('%.2f'%NET['strength_d1']),c2('%.2f'%NET['strength_d7']),pstr(NET['p_strength']),
   c2('%+.2f'%NET['edges']['Raiva']['dia1']),c2('%+.2f'%NET['edges']['Raiva']['dia7'])))
ov=NET['ordinal']['Vigor']; ofa=NET['ordinal']['Fadiga']
P('Como verificação de robustez que respeita a natureza ordinal e limitada das subescalas, um modelo de odds '
 'proporcionais reproduziu o padrão central sem pressupor normalidade. O efeito agudo pré→pós manteve-se nítido — a '
 'chance de estar num tercil superior de fadiga mais que dobrou do pré ao pós (OR = %s; IC95%% %s–%s), enquanto a de '
 'vigor caiu na mesma transição (OR = %s; IC95%% %s–%s) —, e o efeito crônico do dia permaneceu significativo para o '
 'vigor (OR = %s por dia; IC95%% %s–%s) e no limiar da significância para a fadiga (OR = %s; p = %s). A convergência '
 'entre o modelo contínuo, o bayesiano e o ordinal reforça que a deterioração do eixo energia–fadiga não é artefato da '
 'escala de medida.'%(
   c2('%.2f'%ofa['pos']['OR']),c2('%.2f'%ofa['pos']['OR_lo']),c2('%.2f'%ofa['pos']['OR_hi']),
   c2('%.2f'%ov['pos']['OR']),c2('%.2f'%ov['pos']['OR_lo']),c2('%.2f'%ov['pos']['OR_hi']),
   c2('%.2f'%ov['dia']['OR']),c2('%.2f'%ov['dia']['OR_lo']),c2('%.2f'%ov['dia']['OR_hi']),
   c2('%.2f'%ofa['dia']['OR']),pstr(ofa['dia']['p'])))

H('3.16 Classificação individual por mudança confiável (MDC)',12,before=6)
MC=json.load(open('mdc_class.json'))
tmc=table('Classificação dos atletas pela mudança confiável (Δ Dia 1 → Dia 7 acima da MDC90) no vigor e na fadiga.',
    ['Desfecho','Melhora','Estável','Piora','Com mudança confiável'],
    [['Vigor',str(MC['Vigor']['melhora']),str(MC['Vigor']['estável']),str(MC['Vigor']['piora']),'%d/%d'%(MC['n_change_vigor'],MC['n'])],
     ['Fadiga',str(MC['Fadiga']['melhora']),str(MC['Fadiga']['estável']),str(MC['Fadiga']['piora']),'%d/%d'%(MC['n_change_fadiga'],MC['n'])]],
    note='Mudança confiável = |Δ| > MDC90 (vigor %s; fadiga %s pontos). Piora = queda de vigor ou elevação de fadiga acima da MDC. n = %d atletas com medidas completas no Dia 1 e no Dia 7.'%(
        c2('%.1f'%MC['mdc90']['Vigor']),c2('%.1f'%MC['mdc90']['Fadiga']),MC['n']),fs=9)
f_mc=figure(f'{FG}/mdc_individual.png','Mudança individual do vigor e da fadiga (Δ Dia 1 → Dia 7) por atleta, ordenada. A faixa cinza central corresponde à mudança mínima detectável (MDC90); barras fora dela indicam mudança confiável — verde: melhora; vermelho: piora; cinza: dentro do erro de medida.',w=15.5)
P('A aplicação do critério de mudança confiável ao nível individual (Tabela %d; Figura %d) matiza o achado de grupo. '
 'Embora a equipe tenha derivado no eixo energia–fadiga, a maioria dos atletas manteve-se, individualmente, dentro do '
 'erro de medida: apenas %d dos %d apresentaram queda de vigor superior à MDC e %d, elevação de fadiga acima dela. Um '
 'subconjunto restrito — %d atletas — reuniu simultaneamente queda confiável de vigor e elevação confiável de fadiga, '
 'que representa a assinatura individual mais nítida de sobre-esforço e o grupo que mais claramente demandaria atenção da '
 'comissão técnica. Esse contraste entre a resposta média do grupo e a heterogeneidade individual — coerente com a '
 'dispersão das inclinações estimada pelo modelo bayesiano — ilustra por que o monitoramento aplicado deve combinar a '
 'tendência coletiva com a leitura, atleta a atleta, das mudanças que excedem o erro de medida.'%(
   tmc,f_mc,MC['Vigor']['piora'],MC['n'],MC['Fadiga']['piora'],MC['n_piora_ambos']))

# ===== 4 DISCUSSÃO =====
H('4 DISCUSSÃO')
P('O presente estudo caracterizou o comportamento do humor de handebolistas de elite ao longo da última semana de '
 'pré-temporada, e o achado central é inequívoco: a deterioração concentrou-se no eixo energia–fadiga. O vigor caiu e a '
 'fadiga subiu de forma consistente, tanto na resposta aguda a cada treino (pré → pós: vigor −%s%%, d = %s; fadiga +%s%%, '
 'd = %s) quanto na comparação do primeiro ao último dia (vigor d = %s; fadiga d = %s). A robustez desse achado é '
 'sustentada por três abordagens convergentes — o pós-teste do modelo misto na comparação de todos os dias, o teste de Friedman '
 'e a análise multivariada em escores T (Wilks λ = %s; F(%d,%d) = %s; p %s; η²ₚ = %s). Sob o critério conservador de '
 'Bonferroni, apenas o vigor (η²ₚ = %s) e a fadiga (η²ₚ = %s) permaneceram significativos, enquanto as dimensões negativas '
 'de valência não fadiga, próximas do piso, mantiveram-se estáveis. Confirma-se, assim, que o vigor e a fadiga são as '
 'dimensões subjetivas mais sensíveis à carga de treino, em consonância com a literatura de monitoramento (SAW; MAIN; '
 'GASTIN, 2016; THORPE et al., 2017; KELLMANN et al., 2018). Essa seletividade não é acidental: ao acompanhar nadadores '
 'ao longo de uma temporada de 24 semanas, Pierce (2002) verificou que o vigor e a fadiga acompanharam o volume de treino '
 '(r = −0,54 e r = +0,53, respectivamente), ao passo que tensão, depressão, confusão e a perturbação global permaneceram '
 'insensíveis — resultado que replica, em outra modalidade e em '
 'outra escala temporal, o mesmo eixo de resposta aqui identificado e reforça a leitura de que o par vigor–fadiga '
 'constitui o núcleo informativo do painel de humor sob carga.'%(
   c2('%.0f'%abs(pr['Vigor']['pct'])),c2('%+.2f'%pr['Vigor']['dz']),c2('%.0f'%pr['Fadiga']['pct']),c2('%+.2f'%pr['Fadiga']['dz']),
   c2('%+.2f'%mvv('d1d7','Vigor','d')),c2('%+.2f'%mvv('d1d7','Fadiga','d')),
   c2('%.3f'%MV['d1d7']['wilks']),MV['d1d7']['df1'],MV['d1d7']['df2'],c2('%.2f'%MV['d1d7']['Fmv']),pstr(MV['d1d7']['p_mv']),c2('%.2f'%MV['d1d7']['eta_mv']),
   c2('%.2f'%mvv('d1d7','Vigor','eta')),c2('%.2f'%mvv('d1d7','Fadiga','eta'))))
P('Para além de confirmar a deterioração, os resultados detalham a sua dinâmica temporal — o “quando”, tão relevante para '
 'a prática. O humor mais positivo concentrou-se no início da semana (vigor máximo no Dia %d) e a fadiga acumulou-se '
 'progressivamente até o pico no último dia (Dia %d), acompanhada da raiva e da Perturbação Total do Humor. Sobrepõem-se '
 'dois processos: um choque agudo intra-sessão, em que o vigor cai e a fadiga sobe do pré ao pós a cada treino, e um '
 'acúmulo ao longo da semana, que desloca o eixo energia–fadiga rumo ao final do microciclo — como evidenciado pelo '
 'pós-teste, em que os dias finais diferem significativamente do Dia 1. Essa leitura temporal é diretamente acionável: '
 'sinaliza que o reforço da recuperação deve concentrar-se na transição do início da semana (choque de carga) e no '
 'encerramento do microciclo (fadiga acumulada), pontos em que o estado psicofisiológico é mais desfavorável.'%(
   PK['Vigor']['max_day'],PK['Fadiga']['max_day']))
P('Um aspecto metodológico distingue este estudo da tradição descritiva do perfilamento do humor. A literatura clássica '
 'caracteriza o humor de atletas por instantâneos — o perfil de um momento — ou, quando muito, por comparações de médias '
 'entre condições (LOCHBAUM et al., 2021). Aqui, ao ajustar uma função polinomial às trajetórias diárias e dela extrair '
 'as derivadas, passou-se de uma descrição estática para uma leitura dinâmica e quantitativa do humor, na linha do que '
 'Cockerill, Nevill e Lyons (1991) propuseram ao modelar estados de humor com finalidade prescritiva, e não apenas '
 'descritiva. A derivada primeira quantifica a velocidade da deterioração e a derivada segunda localiza o momento em que '
 'essa velocidade deixa de acelerar (inflexão em torno do Dia %s) e oferece à comissão técnica não apenas o quanto, mas '
 'o ritmo e a aceleração do desgaste — informação alinhada ao papel do monitoramento do humor como indicador precoce de '
 'sobrecarga e de excesso de treinamento (FEIJEN et al., 2020; NEDERHOF et al., 2007) e à dinâmica dia-a-dia entre carga '
 'de treino e humor descrita em outras modalidades (HAMLIN et al., 2019; ALFONSO; CAPDEVILA, 2022). A decomposição '
 'sinal–ruído, por sua vez, agrega rigor à interpretação: ao mostrar que o vigor e a fadiga são quase inteiramente sinal, '
 'ao passo que as dimensões negativas próximas do piso são dominadas por ruído, ela fundamenta quantitativamente a '
 'decisão de centrar o monitoramento no eixo energia–fadiga e de não sobreinterpretar as pequenas oscilações das demais '
 'subescalas — uma resposta direta às conhecidas limitações de fidedignidade e de efeito de piso dessas dimensões em '
 'amostras de elite (TERRY; LANE; FOGARTY, 2003). A análise psicométrica desta amostra corrobora e explica esse padrão: '
 'as quatro dimensões negativas excederam o critério de efeito de piso (de 46%% a 79%% das observações no escore zero), '
 'ao passo que o vigor e a fadiga, distribuídos por toda a escala, não o fizeram (TERWEE et al., 2007). Trata-se de uma '
 'expressão esperada do modelo de saúde mental do atleta de elite — o perfil iceberg, com afeto negativo naturalmente '
 'baixo (MORGAN, 1985) —, e não de uma falha do instrumento: a compressão da variância junto ao piso reduz a variação '
 'possível e, com ela, a sensibilidade estatística à carga, o que legitima interpretar as dimensões negativas com '
 'parcimônia e sobretudo como sinalizadoras quando de fato se elevam. Por fim, o cruzamento energia–fadiga, obtido das curvas suavizadas, '
 'condensa em um único marcador visual e objetivo o instante em que a fadiga supera o vigor — evento que o '
 'acompanhamento diário isolado dificilmente tornaria tão explícito.'%(
   '%d'%round(DV['vars']['Vigor']['infl'][0])))
P('A esse arcabouço somam-se três recursos estatísticos que raramente comparecem à literatura de monitoramento do humor '
 'e que reforçam a solidez das conclusões. O modelo bayesiano multinível é o mais consequente: ao dissociar formalmente '
 'a componente aguda da crônica com trajetórias por atleta, ele confere intervalos de credibilidade estáveis mesmo com '
 'amostra reduzida — onde a máxima verossimilhança encontra dificuldade de convergência — e devolve duas leituras que a '
 'análise agregada oculta. A primeira é a partição da variância: o vigor mostrou-se mais um traço (elevada variância '
 'entre atletas) e a fadiga, mais um estado (variância sobretudo intraindividual), distinção com consequências diretas '
 'para o que cada dimensão sinaliza no monitoramento. A segunda é a heterogeneidade das inclinações: embora o grupo '
 'derive na mesma direção, a velocidade da deterioração varia de atleta para atleta, o que fundamenta '
 'quantitativamente a vigilância individualizada e dialoga com a recomendação de referências personalizadas no '
 'acompanhamento do atleta (McELREATH, 2020; SAW; MAIN; GASTIN, 2016). A rede psicométrica, por sua vez, precisa a '
 'natureza da reorganização afetiva: não houve adensamento global da rede, mas um acoplamento seletivo entre fadiga e '
 'irritabilidade ao fim da semana, leitura que as correlações par a par não isolam e que a análise de redes tornou '
 'central no estudo do afeto (EPSKAMP; FRIED, 2018). Por fim, a convergência com o modelo ordinal de odds '
 'proporcionais — que não pressupõe normalidade e respeita o efeito-piso das subescalas — afasta a hipótese de que a '
 'deterioração observada seja artefato da escala de medida (LIDDELL; KRUSCHKE, 2018). Em conjunto, essas abordagens '
 'elevam o rigor inferencial de um desenho de amostra pequena e ilustram um repertório analítico transferível ao '
 'monitoramento intramicrociclo em outros contextos de rendimento.')
P('O caráter do handebol ajuda a contextualizar a magnitude dessa resposta. Trata-se de uma modalidade coletiva de '
 'invasão, marcadamente intermitente, na qual ações de alta intensidade — sprints, saltos, arremessos, bloqueios, '
 'mudanças de direção e contatos — alternam-se com períodos de recuperação incompleta e reclamam, a um só tempo, potência '
 'anaeróbia e capacidade aeróbia intermitente para sustentar o esforço repetido e retardar a instalação da fadiga '
 '(KARCHER; BUCHHEIT, 2014; MICHALSIK; AAGAARD, 2015). Evidências atuais confirmam que o treino e o jogo de handebol '
 'constituem exercício de alta intensidade, com elevadas demandas aeróbias e anaeróbias (PEREIRA et al., 2024), e que os '
 'esportes de invasão se caracterizam por atividade intermitente de alta intensidade intercalada por recuperação de '
 'menor intensidade (VACCARO-BENET et al., 2024). Nesse contexto, o acúmulo de carga da última semana de pré-temporada '
 'explica tanto a deterioração do eixo energia–fadiga quanto a migração do perfil de humor observadas, e esclarece por que '
 'o vigor e a fadiga foram as dimensões mais sensíveis do painel. A pertinência do humor como via de monitoramento no '
 'handebol de elite encontra respaldo adicional em investigação de larga escala com centenas de jogadores de alto nível, '
 'na qual o estado de humor, combinado à percepção de estresse e a marcadores endócrinos, compôs um arcabouço '
 'multidomínio de acompanhamento (RATZ-SULYOK et al., 2026) — moldura em que os presentes achados se inserem ao detalhar, '
 'em alta resolução temporal, a face afetiva dessa resposta ao longo de um microciclo.')
P('Um achado de particular interesse foi a natureza dependente da carga na relação entre o afeto negativo e a fadiga. No '
 'dia de maior vigor (Dia 1, grupo descansado), as dimensões negativas mostraram-se praticamente independentes da fadiga '
 '(depressão × fadiga ρ = %s; raiva × fadiga ρ = %s); no dia de maior fadiga (Dia 7), esse acoplamento tornou-se forte '
 '(depressão × fadiga ρ = %s; raiva × fadiga ρ = %s). Em outras palavras, quando a equipe está recuperada, o atleta mais '
 'irritado ou abatido não é necessariamente o mais cansado; sob carga acumulada, porém, a irritabilidade e o abatimento '
 'organizam-se em torno da exaustão e o perfil de humor “fecha”. Esse padrão indica que a interpretação das dimensões '
 'negativas deve considerar o estado de fadiga do grupo, e que a elevação conjunta de afeto negativo e fadiga ao fim de '
 'um microciclo intenso é, em atletas de elite, uma resposta esperada à carga, e não necessariamente sinal de '
 'desajuste.'%(
   c2('%+.2f'%FP['D1']['Depressao']['rho_fad']),c2('%+.2f'%FP['D1']['Raiva']['rho_fad']),
   c2('%+.2f'%FP['D7']['Depressao']['rho_fad']),c2('%+.2f'%FP['D7']['Raiva']['rho_fad'])))
P('No plano dos perfis de humor, o deslocamento de prevalência do iceberg (prontidão) para a barbatana de tubarão '
 '(fadiga funcional) ao longo da semana — que a regressão logística de tendência confirmou para o perfil de barbatana '
 'de tubarão, cuja chance cresceu cerca de %s%% por dia (OR = %s por dia; IC95%% exclui 1), ainda que a via categórica '
 'do qui-quadrado, de menor potência, não a tenha detectado — reproduz, em um microciclo de handebol, o “derretimento '
 'do iceberg” descrito na literatura de sobrecarga (MORGAN, 1985; HAN; PARSONS-SMITH; TERRY, 2020).'%(
   c2('%.0f'%(100*(LOG['migracao']['barbatana']['OR_dia']-1))),c2('%.2f'%LOG['migracao']['barbatana']['OR_dia'])))
P('É relevante que, mesmo sob o acúmulo de carga da '
 'pré-temporada, os perfis associados a maior risco à saúde mental (Everest invertido, submerso e iceberg invertido) não '
 'se tornaram prevalentes: a deterioração restringiu-se ao perfil de fadiga, cuja prevalência no último dia (%s%%) '
 'superou nitidamente o patamar de referência de atletas do mesmo contexto brasileiro (11,6%% em avaliação momentânea; '
 'DE MIRANDA ROHLFS et al., 2024). Esse contraste sugere que, em handebolistas de elite bem condicionados, a resposta ao '
 'acúmulo de carga é funcional e transitória — uma fadiga esperada — e não um sinal de risco psicológico. Tal leitura '
 'respalda o emprego do perfilamento do humor como triagem de prontidão e de bem-estar em modalidades coletivas de '
 'rendimento (TERRY et al., 2021), no qual a BRUMS já demonstrou sensibilidade a fatores como sono, desempenho e '
 'resultados de partida (ANDRADE et al., 2016; BRANDT; BEVILACQUA; ANDRADE, 2017; DO NASCIMENTO et al., 2026) e '
 'associação com desfechos como a lesão (DE MIRANDA ROHLFS et al., 2025).'%(
   c2('%.0f'%(100*PREV['D7']['Barbatana tubarão']/PREV['n_d7'])),))
P('O caráter transitório dessa resposta deve ser lido à luz da posição do microciclo no planejamento. A última semana '
 'de pré-temporada corresponde a uma fase de acumulação de carga, que antecede o afinamento (tapering) e a competição. '
 'Nesse enquadramento, a queda de vigor e a elevação de fadiga aqui documentadas configuram a assinatura afetiva '
 'esperada do acúmulo, e não um estado estável: o modelo de equilíbrio entre estresse e recuperação prevê que, com a '
 'redução da carga e o reforço da recuperação nos dias subsequentes, o vigor tende a ser restaurado e o perfil de '
 'prontidão (iceberg) a se reinstalar (KELLMANN et al., 2018). Interpretar a fadiga do fim do microciclo como um vale '
 'programado — e não como deterioração — é justamente o que habilita a comissão técnica a distinguir a resposta '
 'funcional daquela que exigiria intervenção, e reforça o valor de manter o monitoramento do humor na transição para o '
 'afinamento, quando a recuperação do estado afetivo é esperada e verificável.')
P('Sob a ótica psicofisiológica, o eixo energia–fadiga funciona como um leitor do equilíbrio entre a carga de treino e '
 'a capacidade de recuperação. Em uma semana de acumulação, esse equilíbrio inclina-se transitoriamente para o polo da '
 'sobrecarga, e a queda do vigor com a ascensão da fadiga exprime, no plano afetivo, a resposta integrada do organismo '
 'ao esforço acumulado — uma cascata de ajustes metabólicos, hormonais e neurais que visam restaurar a homeostase e '
 'poupar energia (WOODS et al., 2018; VRIJKOTTE et al., 2019). A concentração da resposta no eixo energético — e não '
 'nas dimensões de valência emocional — respalda essa leitura: a perturbação nasceu da carga física e da restauração '
 'incompleta entre as sessões. O acoplamento crescente entre o afeto '
 'negativo e a fadiga, do primeiro ao último dia, acompanha essa transição de um estado diferenciado para um estado '
 'integrado de estresse, no qual as respostas afetivas convergem sob um substrato central comum. Convém sublinhar que '
 'essa via afetiva não é redundante com a leitura fisiológica: evidência recente mostra que, sob treino intensificado, '
 'a perturbação do humor pode persistir mesmo quando marcadores biológicos já retornaram à linha de base, dissociação '
 'que expõe assincronia entre a recuperação psicológica e a fisiológica (OSTAPIUK-KAROLCZUK et al., 2025). Tal '
 'descompasso legitima o humor como sentinela independente da fadiga — capaz de sinalizar o custo do esforço quando os '
 'índices periféricos já se normalizaram — e justifica ancorar o monitoramento no eixo energia–fadiga.')
P('Esse conjunto de sinais delineia a assinatura de um sobre-esforço funcional (functional overreaching) — a fadiga '
 'planejada e reversível que precede a supercompensação — e não a de um estado disfuncional. A literatura recente '
 'ampara tal distinção: períodos de treino intensificado agravam a perturbação do humor e comprometem o desempenho, '
 'com plena recuperação após dias de afinamento, tanto em esportes coletivos (CAMPBELL et al., 2020; FERREIRA et al., '
 '2026) quanto em modalidades de resistência (PIACENTINI et al., 2016; WOODS et al., 2018), padrão que revisões '
 'sistemáticas reconhecem '
 'como um marcador do sobre-esforço funcional (ROETE et al., 2021). Cabe, todavia, a cautela metodológica que essas '
 'mesmas revisões assinalam: os instrumentos de humor sinalizam o sobre-esforço, porém nem sempre separam a fadiga '
 'aguda do acúmulo funcional (ROETE et al., 2021) — limitação que o presente desenho amortece ao dissociar a resposta '
 'aguda (pré → pós de cada sessão) do acúmulo semanal e ao modelar a trajetória temporal, o que torna visíveis tanto o '
 'choque intrassessão quanto a deriva ao longo dos dias. A preservação das dimensões de valência não fadiga e a '
 'ausência dos perfis de risco à saúde mental completam o quadro de uma adaptação esperada, e não de um processo '
 'patológico.')
P('Situados no conjunto da literatura, esses achados ocupam uma lacuna específica. A BRUMS tem sido aplicada, '
 'predominantemente, como um retrato pontual em torno da competição (BRANDT et al., 2019; DO NASCIMENTO et al., '
 '2026) ou em comparações entre fases de treino distantes semanas ou meses entre si (ROUVEIX et al., 2006), quase '
 'sempre associada ao sono, ao desempenho ou a marcadores hormonais (ANDRADE et al., 2019; ROUVEIX et al., 2006). São '
 'escassas, contudo, as descrições que acompanham o humor em alta resolução dentro de um único microciclo — com '
 'medidas duas vezes ao dia que separam a resposta aguda do acúmulo crônico —, que modelam a trajetória de forma '
 'contínua em uma modalidade coletiva de invasão de elevada demanda física e emocional. É precisamente nesse ponto '
 'que residem as contribuições originais deste estudo, cada qual voltada a uma face da lacuna. Primeiro, a resolução '
 'temporal: a coleta duas vezes ao dia dissocia o choque agudo intrassessão do acúmulo semanal, distinção que o instantâneo '
 'pré-competitivo e as comparações entre fases distantes não alcançam. Segundo, a passagem da descrição à modelagem: '
 'ao ajustar uma função contínua às trajetórias e dela extrair as derivadas, o estudo quantifica a velocidade e a '
 'aceleração da deterioração e localiza o cruzamento vigor–fadiga — indo além do perfilamento estático que domina a '
 'literatura. Terceiro, a decomposição sinal–ruído, que fundamenta de forma objetiva a decisão de centrar o '
 'monitoramento no eixo energia–fadiga e de não sobreinterpretar as demais subescalas. Quarto, a formalização da '
 'migração de perfis por regressão logística de tendência, que traduz em uma medida de risco o “derretimento do '
 'iceberg” até aqui descrito apenas qualitativamente. Em conjunto, essas quatro frentes convertem o monitoramento do '
 'humor de um retrato pontual em uma leitura dinâmica, quantitativa e acionável, e sinalizam um caminho para o '
 'acompanhamento intramicrociclo em esportes de rendimento.')
P('Algumas limitações delimitam o alcance dos achados. A primeira, e mais importante, é o tamanho amostral (%d atletas), '
 'associado a uma única equipe e a um único microciclo: o desenho observacional de fase única não autoriza inferência '
 'causal sobre a carga, e a generalização a outras equipes, categorias, sexos e momentos da temporada permanece a '
 'confirmar. A segunda diz respeito às dimensões negativas do humor, próximas do piso, com variância e fidedignidade '
 'reduzidas — o que limita a leitura das suas pequenas oscilações e restringe a interpretação da parte da rede que as '
 'envolve. A terceira é que a conversão dos escores T foi referenciada à própria amostra, e não a tabelas normativas '
 'nacionais de atletas, o que recomenda cautela na comparação absoluta de prevalências entre estudos e na classificação '
 'estrita dos perfis. Some-se que a inferência no nível individual esbarra no erro de medida: como a mudança mínima '
 'detectável é ampla, apenas variações individuais expressivas podem ser afirmadas com segurança. Por fim, parte das '
 'análises — a rede psicométrica e o modelo ordinal — tem caráter exploratório e valor sobretudo gerador de hipóteses, '
 'dado o número de observações.'%sm['n'])
P('Essas limitações, contudo, são atenuadas por características do desenho e da análise, o que preserva a validade dos '
 'achados centrais. A medição duas vezes ao dia gerou %d observações e permitiu dissociar a resposta aguda do acúmulo '
 'crônico; a auditoria de robustez (remoção de atípicos por distância de Mahalanobis) não alterou as conclusões; o '
 'modelo bayesiano forneceu intervalos de credibilidade estáveis justamente onde a máxima verossimilhança vacila em '
 'amostras pequenas; e a convergência entre o modelo contínuo, o bayesiano e o ordinal indica que a deterioração do '
 'eixo energia–fadiga não é artefato de uma escolha analítica particular. Em conjunto, o achado central sustenta-se '
 'apesar do tamanho amostral, ainda que a sua extensão a outros contextos dependa de replicação. Como direções '
 'futuras, recomenda-se ampliar a amostra e acompanhar múltiplos microciclos, integrar o humor a marcadores de carga '
 'externa e interna mensurados continuamente, adotar tabelas normativas e agrupamento semeado (seeded k-means) para a '
 'classificação de perfis e testar o valor preditivo dos perfis negativos para desfechos como lesão e sobrecarga '
 '(DE MIRANDA ROHLFS et al., 2024, 2025).'%sm['n_obs'])

# ===== 5 CONCLUSÕES =====
H('5 CONCLUSÕES')
P('Na última semana de pré-temporada, o humor dos handebolistas transitou da prontidão para a fadiga: o vigor recuou e a '
 'fadiga avançou, do primeiro ao último dia e no interior de cada sessão, com confirmação pelas comparações entre todos '
 'os dias e pela análise multivariada. A modelagem das trajetórias situou o cruzamento entre vigor e fadiga na metade da '
 'semana e restringiu o sinal robusto ao eixo energia–fadiga, ao passo que as demais dimensões, próximas do piso, pouco '
 'informaram. O conjunto — deterioração limitada ao eixo energético, sem instalação dos perfis de risco à saúde mental e '
 'sem instalação dos perfis de maior risco à saúde mental — corresponde a um sobre-esforço funcional, ou seja, a uma '
 'fadiga esperada e reversível, própria de uma fase de acumulação. A migração do perfil iceberg para o de barbatana de '
 'tubarão, confirmada por regressão logística de tendência, e o cruzamento entre vigor e fadiga sintetizam essa '
 'dinâmica. Recomenda-se, portanto, o acompanhamento contínuo do humor na pré-temporada, com atenção '
 'primordial ao eixo energia–fadiga e à sua restauração na transição para o afinamento.')

# ===== DECLARAÇÕES (back-matter) =====
H('Contribuições dos autores',12,before=10)
P('[Autor 1] e [Autor 2] conceberam o estudo e o delineamento. [Autor 1] conduziu a coleta de dados. [Autor 1] e '
 '[Autor 3] realizaram as análises estatísticas. [Autor 1] redigiu a primeira versão do manuscrito. Todos os autores '
 'revisaram criticamente o texto e aprovaram a versão final.',ind=False)
H('Financiamento',12,before=6)
P('[Esta pesquisa não recebeu financiamento externo. / Esta pesquisa foi financiada por [agência], processo nº '
 '[placeholder].]',ind=False)
H('Aprovação ética',12,before=6)
P('O estudo foi conduzido de acordo com a Declaração de Helsinque e aprovado pelo Comitê de Ética em Pesquisa de '
 '[instituição] (parecer nº / CAAE [placeholder]).',ind=False)
H('Consentimento informado',12,before=6)
P('Todos os participantes assinaram o termo de consentimento livre e esclarecido antes da coleta de dados.',ind=False)
H('Disponibilidade dos dados',12,before=6)
P('Os dados que sustentam os achados deste estudo, em forma anonimizada, e os scripts de análise estão disponíveis '
 'mediante solicitação razoável ao autor correspondente.',ind=False)
H('Conflitos de interesse',12,before=6)
P('Os autores declaram não haver conflitos de interesse.',ind=False)
H('Artigo companheiro',12,before=6)
P('Este manuscrito integra um par de estudos-companheiros que compartilham a mesma amostra e o mesmo microciclo, com '
 'desfechos primários distintos. O artigo companheiro — sobre os correlatos fisiológicos e de bem-estar da resposta de '
 'humor — foi submetido a outro periódico; ambos se citam mutuamente.',ind=False)

# ===== REFERÊNCIAS =====
H('REFERÊNCIAS')
refs=[
 'ALFONSO, C.; CAPDEVILA, L. Heart rate variability, mood and performance: a pilot study on the interrelation of these variables in amateur road cyclists. PeerJ, v. 10, e13094, 2022. DOI: 10.7717/peerj.13094.',
 'ANDRADE, A. et al. Sleep quality, mood and performance: a study of elite Brazilian volleyball athletes. Journal of Sports Science and Medicine, v. 15, n. 4, p. 601–605, 2016.',
 'ANDRADE, A. et al. Sleep quality associated with mood in elite athletes. The Physician and Sportsmedicine, v. 47, n. 3, p. 312–317, 2019. DOI: 10.1080/00913847.2018.1553467.',
 'ANDRADE, A. et al. Effect of practice exergames on the mood states and self-esteem of elementary school boys and girls during physical education classes: a cluster-randomized controlled trial. PLoS ONE, v. 15, n. 6, e0232392, 2020. DOI: 10.1371/journal.pone.0232392.',
 'BATTAGLINI, M. P. et al. Analysis of progressive muscle relaxation on psychophysiological variables in basketball athletes. International Journal of Environmental Research and Public Health, v. 19, n. 24, 17065, 2022. DOI: 10.3390/ijerph192417065.',
 'BIRD, S. P. et al. Wellness, mood, sleep, and performance in a women’s national basketball team during international competition. Journal of Human Kinetics, v. 96, p. 163–175, 2025. DOI: 10.5114/jhk/200117.',
 'BRANDT, R.; BEVILACQUA, G. G.; ANDRADE, A. Perceived sleep quality, mood states, and their relationship with performance among Brazilian elite athletes during a competitive period. Journal of Strength and Conditioning Research, v. 31, n. 4, p. 1033–1039, 2017.',
 'BRANDT, R. et al. Comparisons of mood states associated with outcomes achieved by female and male athletes in high-level judo and Brazilian jiu-jitsu championships: psychological factors associated with the probability of success. Journal of Strength and Conditioning Research, v. 35, n. 9, p. 2518–2524, 2019. DOI: 10.1519/JSC.0000000000003218.',
 'CAMPBELL, P. G. et al. The effect of overreaching on neuromuscular performance and wellness responses in Australian rules football athletes. Journal of Strength and Conditioning Research, v. 34, n. 6, p. 1530–1538, 2020. DOI: 10.1519/JSC.0000000000003603.',
 'COCKERILL, I. M.; NEVILL, A. M.; LYONS, N. Modelling mood states in athletic performance. Journal of Sports Sciences, v. 9, n. 2, p. 205–212, 1991. DOI: 10.1080/02640419108729881.',
 'DE MIRANDA ROHLFS, I. C. P. et al. Prevalence of specific mood profile clusters among elite and youth athletes at a Brazilian sports club. Sports, v. 12, n. 7, 195, 2024. DOI: 10.3390/sports12070195.',
 'DE MIRANDA ROHLFS, I. C. P. et al. Mood states, injury status, and countermovement jump performance in Brazilian high-level sports. Sports, v. 13, n. 9, 303, 2025. DOI: 10.3390/sports13090303.',
 'DO NASCIMENTO, M. H. et al. Acute psychological responses to official match outcomes in male youth volleyball: an observational repeated-measures study within a single national-level team. Frontiers in Psychology, v. 17, 1826372, 2026. DOI: 10.3389/fpsyg.2026.1826372.',
 'EPSKAMP, S.; FRIED, E. I. A tutorial on regularized partial correlation networks. Psychological Methods, v. 23, n. 4, p. 617–634, 2018. DOI: 10.1037/met0000167.',
 'FEIJEN, S. et al. Monitoring the swimmer’s training load: a narrative review of monitoring strategies applied in research. Scandinavian Journal of Medicine & Science in Sports, v. 30, n. 11, p. 2037–2043, 2020. DOI: 10.1111/sms.13798.',
 'FERREIRA, A. B. M. et al. Impact of sleep restriction and intensified training on mucosal immunity and psychological responses in young soccer players. Journal of Strength and Conditioning Research, v. 40, n. 7, p. e703–e713, 2026. DOI: 10.1519/JSC.0000000000005416.',
 'HAMLIN, M. J. et al. Monitoring training loads and perceived stress in young elite university athletes. Frontiers in Physiology, v. 10, 34, 2019. DOI: 10.3389/fphys.2019.00034.',
 'HAN, C.; PARSONS-SMITH, R. L.; TERRY, P. C. Mood profiling in Singapore: cross-cultural validation and potential applications of mood profile clusters. Frontiers in Psychology, v. 11, 665, 2020. DOI: 10.3389/fpsyg.2020.00665.',
 'KARCHER, C.; BUCHHEIT, M. On-court demands of elite handball, with special reference to playing positions. Sports Medicine, v. 44, n. 6, p. 797–814, 2014. DOI: 10.1007/s40279-014-0164-z.',
 'KELLMANN, M. et al. Recovery and performance in sport: consensus statement. International Journal of Sports Physiology and Performance, v. 13, n. 2, p. 240–245, 2018. DOI: 10.1123/ijspp.2017-0759.',
 'LIDDELL, T. M.; KRUSCHKE, J. K. Analyzing ordinal data with metric models: what could possibly go wrong? Journal of Experimental Social Psychology, v. 79, p. 328–348, 2018. DOI: 10.1016/j.jesp.2018.08.009.',
 'LOCHBAUM, M. et al. The Profile of Mood States and athletic performance: a meta-analysis of published studies. European Journal of Investigation in Health, Psychology and Education, v. 11, n. 1, p. 50–70, 2021. DOI: 10.3390/ejihpe11010005.',
 'McELREATH, R. Statistical rethinking: a Bayesian course with examples in R and Stan. 2. ed. Boca Raton: CRC Press, 2020.',
 'MICHALSIK, L. B.; AAGAARD, P. Physical demands in elite team handball: comparisons between male and female players. Journal of Sports Medicine and Physical Fitness, v. 55, n. 9, p. 878–891, 2015.',
 'MORGAN, W. P. Selected psychological factors limiting performance: a mental health model. In: CLARKE, D. H.; ECKERT, H. M. (Ed.). Limits of human performance. Champaign: Human Kinetics, 1985. p. 70–80.',
 'NEDERHOF, E. et al. Different diagnostic tools in nonfunctional overreaching. International Journal of Sports Medicine, v. 29, n. 7, p. 590–597, 2007. DOI: 10.1055/s-2007-989264.',
 'NEVILL, A. M.; LANE, A. M. Why self-report “Likert” scale data should not be log-transformed. Journal of Sports Sciences, v. 25, n. 1, p. 1–2, 2007. DOI: 10.1080/02640410601111183.',
 'OSTAPIUK-KAROLCZUK, J. et al. Biochemical and psychological markers of fatigue and recovery in mixed martial arts athletes during strength and conditioning training. Scientific Reports, v. 15, n. 1, 24234, 2025. DOI: 10.1038/s41598-025-09719-z.',
 'PARSONS-SMITH, R. L.; TERRY, P. C.; MACHIN, M. A. Identification and description of novel mood profile clusters. Frontiers in Psychology, v. 8, 1958, 2017. DOI: 10.3389/fpsyg.2017.01958.',
 'PEREIRA, R. et al. Exercise intensity and reliability during recreational team handball training for 50–77-year-old unexperienced women. Biology of Sport, v. 41, n. 4, p. 253–261, 2024. DOI: 10.5114/biolsport.2024.132995.',
 'PIACENTINI, M. F. et al. Effect of intensive training on mood with no effect on brain-derived neurotrophic factor. International Journal of Sports Physiology and Performance, v. 11, n. 6, p. 824–830, 2016. DOI: 10.1123/ijspp.2015-0279.',
 'PIERCE, E. F. Relationship between training volume and mood states in competitive swimmers during a 24-week season. Perceptual and Motor Skills, v. 94, n. 3, p. 1009–1012, 2002. DOI: 10.2466/pms.2002.94.3.1009.',
 'RATZ-SULYOK, F. Z. et al. Associations between endocrine status and stress, mood and psychosomatic status in elite handball players. Sports, v. 14, n. 7, 289, 2026. DOI: 10.3390/sports14070289.',
 'ROETE, A. J. et al. A systematic review on markers of functional overreaching in endurance athletes. International Journal of Sports Physiology and Performance, v. 16, n. 8, p. 1065–1073, 2021. DOI: 10.1123/ijspp.2021-0024.',
 'ROHLFS, I. C. P. M. et al. A Escala de Humor de Brunel (Brums): instrumento para detecção precoce da síndrome do excesso de treinamento. Revista Brasileira de Medicina do Esporte, v. 14, n. 3, p. 176–181, 2008.',
 'ROHLFS, I. C. P. M. et al. Psychometric characteristics of the Brazil Mood Scale among youth and elite athletes using two response time frames. Sports, v. 11, n. 12, 244, 2023. DOI: 10.3390/sports11120244.',
 'ROUVEIX, M. et al. The 24 h urinary cortisol/cortisone ratio and epinephrine/norepinephrine ratio for monitoring training in young female tennis players. International Journal of Sports Medicine, v. 27, n. 11, p. 856–863, 2006. DOI: 10.1055/s-2006-923778.',
 'SAW, A. E.; MAIN, L. C.; GASTIN, P. B. Monitoring the athlete training response: subjective self-reported measures trump commonly used objective measures: a systematic review. British Journal of Sports Medicine, v. 50, n. 5, p. 281–291, 2016. DOI: 10.1136/bjsports-2015-094758.',
 'TERRY, P. C.; LANE, A. M.; FOGARTY, G. J. Construct validity of the Profile of Mood States — Adolescents for use with adults. Psychology of Sport and Exercise, v. 4, n. 2, p. 125–139, 2003. DOI: 10.1016/S1469-0292(02)00035-8.',
 'TERRY, P. C. et al. Mood profiling for sustainable mental health among athletes. Sustainability, v. 13, n. 11, 6116, 2021. DOI: 10.3390/su13116116.',
 'TERWEE, C. B. et al. Quality criteria were proposed for measurement properties of health status questionnaires. Journal of Clinical Epidemiology, v. 60, n. 1, p. 34–42, 2007. DOI: 10.1016/j.jclinepi.2006.03.012.',
 'THORPE, R. T. et al. Monitoring fatigue status in elite team-sport athletes: implications for practice. International Journal of Sports Physiology and Performance, v. 12, n. S2, p. S227–S234, 2017. DOI: 10.1123/ijspp.2016-0434.',
 'VACCARO-BENET, P. et al. Internal and external load profile during beach invasion sports match-play by electronic performance and tracking systems: a systematic review. Sensors, v. 24, n. 12, 3738, 2024. DOI: 10.3390/s24123738.',
 'VRIJKOTTE, S. et al. The overtraining syndrome in soldiers: insights from the sports domain. Military Medicine, v. 184, n. 5-6, p. e192–e200, 2019. DOI: 10.1093/milmed/usy274.',
 'WAGNER, H. et al. Individual and team performance in team-handball: a review. Journal of Sports Science and Medicine, v. 13, n. 4, p. 808–816, 2014.',
 'WOODS, A. L. et al. The effects of intensified training on resting metabolic rate, body composition and performance in trained cyclists. PLoS ONE, v. 13, n. 2, e0191644, 2018. DOI: 10.1371/journal.pone.0191644.']
for rf in refs:
    p=doc.add_paragraph(); r=p.add_run(rf); r.font.size=Pt(11); p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(6); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

OUTP='/home/user/mdlucca/Artigos/Paper1_Humor_dinamica.docx'
doc.save(OUTP); print('SAVED',OUTP,'| Tabelas',_TN[0],'Figuras',_FN[0])
