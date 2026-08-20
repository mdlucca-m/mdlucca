# -*- coding: utf-8 -*-
import json, numpy as np, pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
FG='/home/user/mdlucca/Artigos/figuras'
AR=json.load(open('all_results.json')); T=json.load(open('temporal.json'))
h=pd.read_csv('hum_prof.csv')
DIMS=[('Vigor','Vigor'),('Fadiga','Fadiga'),('Tensao','Tensão'),('Depressao','Depressão'),
      ('Raiva','Raiva'),('Confusao','Confusão'),('TMD','Perturbação Total do Humor (PTH)'),
      ('FadFisica','Fadiga física'),('FadMental','Fadiga mental')]
def c2(s): return str(s).replace('.',',')
def num(x,d=1): return 'n/d' if x is None else c2(f'{x:.{d}f}')
def sg(x,d=2): return c2(f'{x:+.{d}f}')
def pvt(p): return 'p < 0,001' if p<0.001 else 'p = '+c2(f'{p:.3f}')

doc=Document()
stl=doc.styles['Normal']; stl.font.name='Times New Roman'; stl.font.size=Pt(12)
stl.element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')
stl.paragraph_format.line_spacing=1.5; stl.paragraph_format.space_after=Pt(0)
sec=doc.sections[0]; sec.top_margin=Cm(2.3); sec.left_margin=Cm(2.2); sec.bottom_margin=Cm(2); sec.right_margin=Cm(2.2)
_FN=[0]
def P(t='',ind=True,after=0):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(12); r.font.name='Times New Roman'
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(after); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if ind: p.paragraph_format.first_line_indent=Cm(1.25)
def H(t,before=12,size=12,center=False):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(size); r.font.name='Times New Roman'
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(4)
    if center: p.alignment=WD_ALIGN_PARAGRAPH.CENTER
def figure(path,cap,w=15.5):
    _FN[0]+=1
    pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; pp.add_run().add_picture(path,width=Cm(w)); pp.paragraph_format.space_before=Pt(6)
    pc=doc.add_paragraph(); pc.alignment=WD_ALIGN_PARAGRAPH.CENTER; rc=pc.add_run('Figura %d – %s'%(_FN[0],cap)); rc.font.size=Pt(11); rc.font.name='Times New Roman'
    pc.paragraph_format.space_after=Pt(6)

def floor(k): return 100*np.mean(pd.to_numeric(h[k],errors='coerce').dropna()<=0)
def sig_trans(k):
    out=[]
    for t in T['trans']:
        v=t['vars'][k]
        if v['sig']: out.append((t['lab'],t['tipo'],v['dz'],v['p']))
    return out
def byday(k):
    d=AR['por_dia'][k]; return {int(kk):vv for kk,vv in d.items()}

# ---- Título ----
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('COMPORTAMENTO DAS VARIÁVEIS DO HUMOR AO LONGO DA SEMANA DE PRÉ-TEMPORADA: '
            'ANÁLISE DESCRITIVA E DINÂMICA POR DIMENSÃO'); r.bold=True; r.font.size=Pt(13)
p.paragraph_format.space_after=Pt(8)
H('Nota metodológica',before=2)
P('Cada dimensão do BRUMS é apresentada em um painel individual com seis quadros: (A) trajetória diária suavizada '
  'por spline com os limites de mínimo e máximo, marcação dos dias de treino intervalado de alta intensidade e do '
  'ponto de inflexão; (B) velocidade instantânea de mudança, dada pela primeira derivada; (C) aceleração, dada pela '
  'segunda derivada; (D) tamanho de efeito de cada transição sequencial, com separação entre o efeito agudo intradia e o efeito '
  'de recuperação entre dias; (E) efeito agudo do pré ao pós-treino em cada dia; e (F) síntese descritiva. Os escores '
  'variam de zero a dezesseis por dimensão. O tamanho de efeito padronizado para dados pareados é o dz, e o asterisco '
  'indica significância no teste de Wilcoxon a cinco por cento. A confiabilidade é expressa pelo coeficiente de '
  'correlação intraclasse (ICC) e a mudança mínima detectável a noventa e cinco por cento (MDC95) define o limiar de '
  'mudança confiável para cada variável.')

for k,lab in DIMS:
    ds=AR['desc_semana'][k]; fr=AR['friedman'][k]; d17=AR['d1d7'][k]; pp=AR['prepos'][k]
    bd=byday(k); dv=T['derivadas'][k]; ac=T['agudo_cronico'][k]; lim=T['limiares'].get(k,{})
    pv=AR['pv_vs_deterioracao'].get(k,{}); fl=floor(k); st=sig_trans(k)
    inf=('dia %s'%num(dv['inflexao'][0])) if dv['inflexao'] else 'sem inflexão detectada'
    d1v=bd.get(1); d7v=bd.get(7)
    doc.add_page_break()
    H(lab.upper(),before=2,size=13,center=True)
    figure(f'{FG}/vd_{k}.png','Painel dinâmico da dimensão %s: trajetória, limites, derivadas, transições e síntese descritiva.'%lab)

    # Parágrafo 1 — tendência central, dispersão, distribuição e confiabilidade
    P('A dimensão %s apresentou média semanal de %s ± %s pontos, mediana de %s e amplitude observada de %s a %s pontos, '
      'com coeficiente de variação de %s por cento, o que indica %s dispersão relativa entre atletas e momentos. '
      'A proporção de escores nulos foi de %s por cento, que caracteriza %s efeito de piso; essa condição sustenta o uso '
      'de estatística não paramétrica e de tamanhos de efeito pareados. A confiabilidade entre dias foi %s (ICC = %s), '
      'e a mudança mínima detectável foi de %s pontos (MDC95), valor a partir do qual uma variação individual pode ser '
      'considerada real e não atribuível ao erro de medida.'%(
        lab,num(ds['m'],2),num(ds['sd'],2),num(ds['med']),num(ds['mn'],0),num(ds['mx'],0),num(ds['cv'],0),
        'ampla' if (ds['cv'] or 0)>=70 else 'moderada' if (ds['cv'] or 0)>=40 else 'baixa',
        num(fl,0),'acentuado' if fl>=60 else 'moderado' if fl>=25 else 'discreto',
        'elevada' if lim.get('icc',0)>=0.75 else 'moderada' if lim.get('icc',0)>=0.5 else 'baixa',
        num(lim.get('icc',0),2),num(lim.get('mdc95',0),1)))

    # Parágrafo 2 — trajetória semanal, Friedman, D1->D7
    P('Ao longo do microciclo, %s variou de forma %s (teste de Friedman: qui-quadrado = %s, %s; W de Kendall = %s). '
      'Entre o primeiro e o sétimo dia, o escore passou de %s para %s, uma variação de %s ponto (dz = %s, %s), '
      'considerada %s. A leitura por dia mostra os valores médios de %s (D1), %s (D2), %s (D3), %s (D4), %s (D5), '
      '%s (D6) e %s (D7), que evidenciam %s ao longo da semana.'%(
        lab,'significativa' if fr['p']<0.05 else 'não significativa',num(fr['chi'],2),pvt(fr['p']),num(fr['W'],2),
        num(d1v,1) if d1v is not None else 'n/d',num(d7v,1) if d7v is not None else 'n/d',
        sg(d17['delta'],1),sg(d17['dz']),pvt(d17['p']),
        'grande' if abs(d17['dz'])>=0.8 else 'moderada' if abs(d17['dz'])>=0.5 else 'pequena' if abs(d17['dz'])>=0.2 else 'trivial',
        *[num(bd.get(d),1) if bd.get(d) is not None else 'n/d' for d in range(1,8)],
        'deterioração progressiva' if d17['delta']<0 and k=='Vigor' else 'acúmulo progressivo' if d17['delta']>0 else 'estabilidade relativa'))

    # Parágrafo 3 — derivadas, forma da curva, agudo x crônico
    P('A dinâmica da curva confirma esse padrão: a velocidade de mudança atingiu pico de %s ponto por dia por volta do %s, '
      'e a curva apresentou %s, ponto em que a tendência inverte o sinal da aceleração. O contraste entre efeitos mostra '
      'um efeito agudo médio intradia de %s ponto e um efeito crônico acumulado (D1 a D7) de %s ponto, o que indica que a '
      'mudança da variável foi predominantemente %s. No conjunto pré versus pós-treino da semana, a variação foi de %s '
      '(dz = %s, %s).'%(
        sg(dv['vel_pico'],1),'dia %s'%num(dv['vel_pico_dia']),inf,
        sg(ac['agudo_medio']),sg(ac['cronico_delta']),
        'de natureza crônica, por acúmulo entre dias' if abs(ac['cronico_delta'])>abs(ac['agudo_medio'])*2 else 'resultante tanto do estímulo agudo quanto do acúmulo entre dias',
        sg(pp['pos']-pp['pre'],2),num(pp['dz'],2),pvt(pp['p'])))

    # Parágrafo 4 — transições significativas + aptidão
    if st:
        trs='; '.join('%s (%s, dz = %s, %s)'%(l.replace(' → ','→'),tp.lower(),sg(dz),pvt(pp2)) for l,tp,dz,pp2 in st)
        s4='As transições estatisticamente significativas concentraram-se em: %s. '%trs
    else:
        s4='Nenhuma transição sequencial isolada alcançou significância estatística, o que é coerente com o efeito de piso e com uma variação distribuída ao longo da semana. '
    s4+=('Quanto à aptidão, a correlação entre o pico de velocidade no T-car e a deterioração desta variável entre D1 e D7 '
         'foi de rho = %s (%s, n = %s), resultado que mostra que o nível de aptidão aeróbia %s a magnitude da mudança desta dimensão '
         'ao longo da semana.'%(
        sg(pv.get('rho',0)),pvt(pv.get('p',1)),pv.get('n','n/d'),
        'não explicou' if pv.get('p',1)>=0.05 else 'associou-se a'))
    P(s4)

doc.save('/home/user/mdlucca/Artigos/Comportamento_Variaveis_Humor.docx')
print('OK Comportamento_Variaveis_Humor.docx')
