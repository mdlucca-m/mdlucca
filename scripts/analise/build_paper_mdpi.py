# -*- coding: utf-8 -*-
# Versao do artigo enxuto no LAYOUT MDPI (periodico Sports), em portugues,
# com citacoes numeradas [n] e referencias em estilo MDPI. Autoria/etica em placeholders.
import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
R=json.load(open('brums_desc2.json')); MS=json.load(open('model_stats.json'))
PSY=json.load(open('psychometric.json')); PK=json.load(open('peaks.json'))
PCA=json.load(open('pca.json')); LOG=json.load(open('logistica.json')); MV=json.load(open('manova.json'))
STAT=json.load(open('brums_stats3.json')); S4=json.load(open('brums_stats4.json')); CRS=json.load(open('cross.json'))
MDC=json.load(open('mdc.json')); PHJ=json.load(open('posthoc.json')); SMO=json.load(open('smooth_deriv.json'))
POLY=json.load(open('poly_fit.json')); CRX=json.load(open('cross_pts.json'))
FG='/home/user/mdlucca/Artigos/figuras'
BODY='Palatino Linotype'
def c2(s): return str(s).replace('.',',')
doc=Document()
stl=doc.styles['Normal']; stl.font.name=BODY; stl.font.size=Pt(10)
stl.element.rPr.rFonts.set(qn('w:eastAsia'),BODY)
stl.paragraph_format.line_spacing=1.15; stl.paragraph_format.space_after=Pt(0)
sec=doc.sections[0]; sec.top_margin=Cm(2.2); sec.left_margin=Cm(1.9); sec.bottom_margin=Cm(2.0); sec.right_margin=Cm(1.9)
_TN=[0]; _FN=[0]; _SN=[0]; SUPP=[]
def P(t='',just=True,size=10,after=0,bold=False,ind=True,italic=False,color=None):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(size); r.bold=bold; r.italic=italic; r.font.name=BODY
    if color: r.font.color.rgb=color
    p.paragraph_format.line_spacing=1.15; p.paragraph_format.space_after=Pt(after)
    if just: p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if just and ind: p.paragraph_format.first_line_indent=Cm(0.5)
    return p
def RUN(pairs,after=0,ind=True,size=10):
    p=doc.add_paragraph(); p.paragraph_format.line_spacing=1.15; p.paragraph_format.space_after=Pt(after)
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if ind: p.paragraph_format.first_line_indent=Cm(0.5)
    for txt,bold in pairs:
        r=p.add_run(txt); r.bold=bold; r.font.size=Pt(size); r.font.name=BODY
    return p
def H(t,size=10,before=8,after=3):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(size); r.font.name=BODY
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after); p.paragraph_format.line_spacing=1.15
    return p
def _bd(c):
    tcPr=c._tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for e_ in ['top','bottom','left','right']:
        el=OxmlElement(f'w:{e_}'); on=e_ in ('top','bottom')
        el.set(qn('w:val'),'single' if on else 'nil'); el.set(qn('w:sz'),'6'); el.set(qn('w:color'),'000000'); b.append(el)
    tcPr.append(b)
def table(cap,header,rows,fs=8.5,note=None):
    _TN[0]+=1
    p=doc.add_paragraph(); r=p.add_run('Tabela %d. '%_TN[0]); r.bold=True; r.font.size=Pt(9); r.font.name=BODY
    r2=p.add_run(cap); r2.font.size=Pt(9); r2.font.name=BODY
    p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(2); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    t=doc.add_table(rows=1,cols=len(header)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    t.autofit=False; tblPr=t._tbl.tblPr
    tl=OxmlElement('w:tblLayout'); tl.set(qn('w:type'),'fixed'); tblPr.append(tl)
    cm=OxmlElement('w:tblCellMar')
    for side in ('left','right'):
        e=OxmlElement('w:'+side); e.set(qn('w:w'),'40'); e.set(qn('w:type'),'dxa'); cm.append(e)
    tblPr.append(cm)
    nC=len(header); TOT=17.0
    if nC>2:
        c0=min(3.1,TOT/nC*1.5); cL=2.8 if len(str(header[-1]))>6 else (TOT-c0)/(nC-1)
        mid=(TOT-c0-cL)/(nC-2); _cw=[c0]+[mid]*(nC-2)+[cL]
    else:
        _cw=[TOT/nC]*nC
    grid=t._tbl.find(qn('w:tblGrid'))
    if grid is not None:
        for i,gc in enumerate(grid.findall(qn('w:gridCol'))):
            gc.set(qn('w:w'),str(int(_cw[i]*567)))
    for i,htx in enumerate(header):
        cc=t.rows[0].cells[i]; cc.text=''; cc.width=Cm(_cw[i]); rr=cc.paragraphs[0].add_run(htx); rr.bold=True; rr.font.size=Pt(fs); rr.font.name=BODY
        cc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; _bd(cc)
    for row in rows:
        cs=t.add_row().cells
        for i,val in enumerate(row):
            cs[i].text=''; cs[i].width=Cm(_cw[i]); rr=cs[i].paragraphs[0].add_run(str(val)); rr.font.size=Pt(fs); rr.font.name=BODY
            cs[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if i>0 else WD_ALIGN_PARAGRAPH.LEFT; _bd(cs[i])
    if note:
        pn=doc.add_paragraph(); rn=pn.add_run(note); rn.font.size=Pt(8); rn.italic=True; rn.font.name=BODY; pn.paragraph_format.space_after=Pt(4)
    else:
        doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return _TN[0]
def figure(path,cap,w=14.0):
    _FN[0]+=1
    pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; pp.add_run().add_picture(path,width=Cm(w)); pp.paragraph_format.space_before=Pt(6)
    pc=doc.add_paragraph(); pc.alignment=WD_ALIGN_PARAGRAPH.CENTER; pc.paragraph_format.space_after=Pt(6)
    rc=pc.add_run('Figura %d. '%_FN[0]); rc.bold=True; rc.font.size=Pt(9); rc.font.name=BODY
    rc2=pc.add_run(cap); rc2.font.size=Pt(9); rc2.font.name=BODY
    return _FN[0]
def sfig(path,cap,w=14.0):
    # figura suplementar: registra para renderizacao diferida (bloco Material Suplementar)
    _SN[0]+=1; SUPP.append((path,cap,w,_SN[0])); return _SN[0]
def render_supp():
    if not SUPP: return
    H('Material Suplementar',before=12)
    P('As figuras a seguir complementam os resultados do corpo do artigo e detalham a distribuição diária, a resposta '
      'aguda pré e pós-treino, as comparações par a par, a análise por derivadas e o ajuste polinomial das trajetórias.',after=4)
    for path,cap,w,n in SUPP:
        pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; pp.add_run().add_picture(path,width=Cm(w)); pp.paragraph_format.space_before=Pt(6)
        pc=doc.add_paragraph(); pc.alignment=WD_ALIGN_PARAGRAPH.CENTER; pc.paragraph_format.space_after=Pt(6)
        rc=pc.add_run('Figura S%d. '%n); rc.bold=True; rc.font.size=Pt(9); rc.font.name=BODY
        rc2=pc.add_run(cap); rc2.font.size=Pt(9); rc2.font.name=BODY

sm=R['sample']; desc=R['desc']; pr=R['prepos']; d17=R['d1d7']; PREV=MS['prev']
def num(x,d=1): return c2(f'{x:.{d}f}')
def pv(p): return '< 0,001' if p<0.001 else '= '+c2(f'{p:.3f}')
def pvt(p): return '< 0,001' if p<0.001 else c2(f'{p:.3f}')
fr_vig_p=pv(S4['friedman']['Vigor']['p']); fr_vig_w=num(S4['friedman']['Vigor']['W'],2)
fr_fad_p=pv(S4['friedman']['Fadiga']['p']); fr_fad_w=num(S4['friedman']['Fadiga']['W'],2)
vig_dz=num(pr['Vigor']['dz'],2); fad_dz=num(pr['Fadiga']['dz'],2)
vig_pct=num(abs(pr['Vigor']['pct']),0); fad_pct=num(pr['Fadiga']['pct'],0)
d1d7_vig_dz=num(d17['Vigor']['dz'],2); d1d7_fad_dz=num(d17['Fadiga']['dz'],2)
d1d7_vig_pct=num(abs(d17['Vigor']['pct']),0); d1d7_fad_pct=num(d17['Fadiga']['pct'],0)
wilks=num(MV['d1d7']['wilks'],3)
vmaxd=PK['Vigor']['max_day']; vmind=PK['Vigor']['min_day']
fmaxd=PK['Fadiga']['max_day']; fmind=PK['Fadiga']['min_day']
ice_d1=num(100*PREV['D1']['Iceberg']/PREV['n_d1'],0); ice_d7=num(100*PREV['D7']['Iceberg']/PREV['n_d7'],0)
bar_d1=num(100*PREV['D1']['Barbatana tubarão']/PREV['n_d1'],0); bar_d7=num(100*PREV['D7']['Barbatana tubarão']/PREV['n_d7'],0)
sub_d1=num(100*PREV['D1']['Submerso']/PREV['n_d1'],0); sub_d7=num(100*PREV['D7']['Submerso']/PREV['n_d7'],0)
or_bar=num(LOG['migracao']['barbatana']['OR_dia'],2)
pth_rv=num(abs(PSY['PTH']['rho_vigor']),2); pth_rf=num(PSY['PTH']['rho_fadiga'],2); pth_r2=num(100*PSY['PTH']['r2_axis'],0)
a_vig=num(PSY['Vigor']['alpha'],2); a_fad=num(PSY['Fadiga']['alpha'],2)
idade=num(sm['idade']['mean'],1); idade_sd=num(sm['idade']['sd'],1)
exp=num(sm['exp']['mean'],1); exp_sd=num(sm['exp']['sd'],1)
ORD=[('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Vigor','Vigor'),('Fadiga','Fadiga'),('Confusao','Confusão'),('TMD','PTH')]
GREY=RGBColor(0x66,0x66,0x66)

# ===== GERENCIADOR DE CITAÇÕES (numeracao MDPI por ordem de aparicao) =====
REFS={
 'saw2016':'Saw, A.E.; Main, L.C.; Gastin, P.B. Monitoring the athlete training response: Subjective self-reported measures trump commonly used objective measures: A systematic review. Br. J. Sports Med. 2016, 50, 281–291. https://doi.org/10.1136/bjsports-2015-094758.',
 'lochbaum2021':'Lochbaum, M.; Zanatta, T.; Kirschling, D.; May, E. The Profile of Mood States and athletic performance: A meta-analysis of published studies. Eur. J. Investig. Health Psychol. Educ. 2021, 11, 50–70. https://doi.org/10.3390/ejihpe11010005.',
 'helwig2023':'Helwig, J.; Diels, J.; Röll, M.; Mahler, H.; Gollhofer, A.; Roecker, K.; Willwacher, S. Relationships between external, wearable sensor-based, and internal parameters: A systematic review. Sensors 2023, 23, 827. https://doi.org/10.3390/s23020827.',
 'kellmann2018':'Kellmann, M.; Bertollo, M.; Bosquet, L.; Brink, M.; Coutts, A.J.; Duffield, R.; Erlacher, D.; Halson, S.L.; Hecksteden, A.; Heidari, J.; et al. Recovery and performance in sport: Consensus statement. Int. J. Sports Physiol. Perform. 2018, 13, 240–245. https://doi.org/10.1123/ijspp.2017-0759.',
 'terry2003':'Terry, P.C.; Lane, A.M.; Fogarty, G.J. Construct validity of the Profile of Mood States, Adolescents for use with adults. Psychol. Sport Exerc. 2003, 4, 125–139. https://doi.org/10.1016/S1469-0292(02)00035-8.',
 'rohlfs2008':'Rohlfs, I.C.P.M.; Rotta, T.M.; Luft, C.D.B.; Andrade, A.; Krebs, R.J.; Carvalho, T. A Escala de Humor de Brunel (BRUMS): Instrumento para detecção precoce da síndrome do excesso de treinamento. Rev. Bras. Med. Esporte 2008, 14, 176–181.',
 'terryltu2022':'Terry, P.C.; Skurvydas, A.; Lisinskiene, A.; Majauskiene, D.; Valanciene, D.; Cooper, S.; Lochbaum, M. Validation of a Lithuanian-language version of the Brunel Mood Scale: The BRUMS-LTU. Int. J. Environ. Res. Public Health 2022, 19, 4867. https://doi.org/10.3390/ijerph19084867.',
 'morgan1985':'Morgan, W.P. Selected psychological factors limiting performance: A mental health model. In Limits of Human Performance; Clarke, D.H., Eckert, H.M., Eds.; Human Kinetics: Champaign, IL, USA, 1985; pp. 70–80.',
 'parsons2017':'Parsons-Smith, R.L.; Terry, P.C.; Machin, M.A. Identification and description of novel mood profile clusters. Front. Psychol. 2017, 8, 1958. https://doi.org/10.3389/fpsyg.2017.01958.',
 'luojumaki2026':'Luojumäki, R.J.; Ruiz, M.C.; Adie, J.M.; Terry, P.C. Exploring mood profile clusters across physical activity level, gender and age in a Finnish population. Eur. J. Sport Sci. 2026, 26, e70131. https://doi.org/10.1002/ejsc.70131.',
 'terry2021':'Terry, P.C.; Parsons-Smith, R.L.; Terry, V.R. Mood profiling for sustainable mental health among athletes. Sustainability 2021, 13, 6116. https://doi.org/10.3390/su13116116.',
 'han2020':'Han, C.; Parsons-Smith, R.L.; Terry, P.C. Mood profiling in Singapore: Cross-cultural validation and potential applications of mood profile clusters. Front. Psychol. 2020, 11, 665. https://doi.org/10.3389/fpsyg.2020.00665.',
 'lew2023':'Lew, P.C.F.; Parsons-Smith, R.L.; Lamont-Mills, A.; Terry, P.C. Cross-cultural validation of the Malaysian Mood Scale and tests of between-group mood differences. Int. J. Environ. Res. Public Health 2023, 20, 3348. https://doi.org/10.3390/ijerph20043348.',
 'karcher2014':'Karcher, C.; Buchheit, M. On-court demands of elite handball, with special reference to playing positions. Sports Med. 2014, 44, 797–814. https://doi.org/10.1007/s40279-014-0164-z.',
 'garcia2023':'García-Sánchez, C.; Navarro, R.M.; Karcher, C.; de la Rubia, A. Physical demands during official competitions in elite handball: A systematic review. Int. J. Environ. Res. Public Health 2023, 20, 3353. https://doi.org/10.3390/ijerph20043353.',
 'carton2023':'Carton-Llorente, A.; Lozano, D.; Gilart Iglesias, V.; Jorquera, D.M.; Manchado, C. Worst-case scenario analysis of physical demands in elite men handball players by playing position through big data analytics. Biol. Sport 2023, 40, 1219–1227. https://doi.org/10.5114/biolsport.2023.126665.',
 'thorpe2017':'Thorpe, R.T.; Atkinson, G.; Drust, B.; Gregson, W. Monitoring fatigue status in elite team-sport athletes: Implications for practice. Int. J. Sports Physiol. Perform. 2017, 12, S227–S234. https://doi.org/10.1123/ijspp.2016-0434.',
 'roete2021':'Roete, A.J.; Elferink-Gemser, M.T.; Otter, R.T.A.; Stoter, I.K.; Lamberts, R.P. A systematic review on markers of functional overreaching in endurance athletes. Int. J. Sports Physiol. Perform. 2021, 16, 1065–1073. https://doi.org/10.1123/ijspp.2021-0024.',
 'manescu2026':'Mănescu, D.C.; Voinea, A.; Plastoi, C.D.; Iacobini, A.R.; Vulpe, A.A.; Pîrvan, A.; Dinciu, C.C.; Vulpe, B.I.; Băltărețu, C.; Iacobini, A. Molecular biomarkers of training responses: A systems framework for exercise adaptation and athlete monitoring. Int. J. Mol. Sci. 2026, 27, 3601. https://doi.org/10.3390/ijms27083601.',
 'latorre2023':'la Torre, M.E.; Monda, A.; Messina, A.; de Stefano, M.I.; Monda, V.; Moscatelli, F.; Tafuri, F.; Saraiello, E.; Latino, F.; Monda, M.; et al. The potential role of nutrition in overtraining syndrome: A narrative review. Nutrients 2023, 15, 4916. https://doi.org/10.3390/nu15234916.',
 'rohlfs2024':'de Miranda Rohlfs, I.C.P.; et al. Prevalence of specific mood profile clusters among elite and youth athletes at a Brazilian sports club. Sports 2024, 12, 195. https://doi.org/10.3390/sports12070195.',
 'pierce2002':'Pierce, E.F. Relationship between training volume and mood states in competitive swimmers during a 24-week season. Percept. Mot. Skills 2002, 94, 1009–1012. https://doi.org/10.2466/pms.2002.94.3.1009.',
 'ferreira2026':'Ferreira, A.B.M.; Halson, S.; Galvão-Coelho, N.L.; Almeida, R.N.; Nakamura, F.Y.; Mortatti, A.L. Impact of sleep restriction and intensified training on mucosal immunity and psychological responses in young soccer players. J. Strength Cond. Res. 2026, 40, e703–e713. https://doi.org/10.1519/JSC.0000000000005416.',
 'rohlfs2023':'Rohlfs, I.C.P.M.; et al. Psychometric characteristics of the Brazil Mood Scale among youth and elite athletes using two response time frames. Sports 2023, 11, 244. https://doi.org/10.3390/sports11120244.',
 'struzik2026':'Struzik, A.; Nadobnik, J.; Stępień-Słodkowska, M. TRIMP and session-RPE monitoring in elite women’s handball: A full-season descriptive analysis. Sci. Rep. 2026, 16, 53134. https://doi.org/10.1038/s41598-026-53134-x.',
 'staiano2025':'Staiano, W.; Sternberg, L.M.; Ferri-Caruana, A.; Salazar-Bonet, L.R.; Romagnoli, M.; Henriksen, K.; Kirk, U. Overcoming mental fatigue through mindfulness: Improving physical and cognitive performance in elite handball players. J. Sci. Med. Sport 2025, 29, 91–99. https://doi.org/10.1016/j.jsams.2025.08.004.',
 'bird2025':'Bird, S.P.; Hawley, C.; Cavallerio, F.; Hobbs, M.; Sáez de Villarreal, E. Wellness, mood, sleep, and performance in a women’s national basketball team during international competition. J. Hum. Kinet. 2025, 96, 163–175. https://doi.org/10.5114/jhk/200117.',
 'ratz2026':'Ratz-Sulyok, F.Z.; et al. Associations between endocrine status and stress, mood and psychosomatic status in elite handball players. Sports 2026, 14, 289. https://doi.org/10.3390/sports14070289.',
 'nascimento2026':'do Nascimento, M.H.; et al. Acute psychological responses to official match outcomes in male youth volleyball: An observational repeated-measures study within a single national-level team. Front. Psychol. 2026, 17, 1826372. https://doi.org/10.3389/fpsyg.2026.1826372.',
 'ostapiuk2025':'Ostapiuk-Karolczuk, J.; et al. Biochemical and psychological markers of fatigue and recovery in mixed martial arts athletes during strength and conditioning training. Sci. Rep. 2025, 15, 24234. https://doi.org/10.1038/s41598-025-09719-z.',
}
_ORD=[]
def _n(slug):
    if slug not in REFS: raise KeyError('ref desconhecida: '+slug)
    if slug not in _ORD: _ORD.append(slug)
    return _ORD.index(slug)+1
def cite(*slugs):
    nums=sorted(_n(s) for s in slugs)
    # agrupa sequencias consecutivas: 3+ vira faixa a–b; 1-2 vira lista
    toks=[]; i=0
    while i<len(nums):
        j=i
        while j+1<len(nums) and nums[j+1]==nums[j]+1: j+=1
        if j-i>=2: toks.append('%d–%d'%(nums[i],nums[j]))
        else: toks.extend(str(nums[k]) for k in range(i,j+1))
        i=j+1
    return '['+','.join(toks)+']'

# ===== BARRA/FRONT MATTER (placeholders de produção MDPI) =====
P('Citation: [Autores]. [Título]. Sports 2026, [vol], [artigo]. Received: [data]; Accepted: [data]; '
 'Published: [data]. © 2026 pelos autores. Licenciado sob CC BY 4.0.',size=8,ind=False,italic=True,color=GREY,after=8)
p=doc.add_paragraph(); r=p.add_run('Sports'); r.bold=True; r.italic=True; r.font.size=Pt(15); r.font.name=BODY
r2=p.add_run('    Article'); r2.bold=True; r2.font.size=Pt(11); r2.font.name=BODY; p.paragraph_format.space_after=Pt(4)

# TÍTULO
p=doc.add_paragraph(); r=p.add_run('Dinâmica do Humor em um Microciclo Pré-Competitivo de Handebol de Elite: '
 'Perfis de Humor e o Eixo Energia–Fadiga')
r.bold=True; r.font.size=Pt(16); r.font.name=BODY; p.paragraph_format.space_after=Pt(6)
# AUTORES (placeholders)
p=doc.add_paragraph();
for t,sup in [('[Autor 1] ',False),('1,*',True),(', [Autor 2] ',False),('2',True),(', [Autor 3] ',False),('3',True)]:
    r=p.add_run(t); r.font.size=Pt(11); r.font.name=BODY
    if sup: r.font.superscript=True
p.paragraph_format.space_after=Pt(4)
# AFILIAÇÕES
for n,txt in [('1','[Afiliação 1: departamento, instituição, cidade, país; e-mail]'),
              ('2','[Afiliação 2: departamento, instituição, cidade, país; e-mail]'),
              ('3','[Afiliação 3: departamento, instituição, cidade, país; e-mail]')]:
    p=doc.add_paragraph(); rs=p.add_run(n); rs.font.superscript=True; rs.font.size=Pt(8); rs.font.name=BODY
    r=p.add_run(' '+txt); r.font.size=Pt(8); r.font.name=BODY; p.paragraph_format.space_after=Pt(0)
p=doc.add_paragraph(); rs=p.add_run('*'); rs.font.size=Pt(8); rs.font.name=BODY
r=p.add_run(' Correspondência: [e-mail do autor correspondente]'); r.font.size=Pt(8); r.font.name=BODY; p.paragraph_format.space_after=Pt(8)

# RESUMO
RUN([('Resumo: ',True),(f'Objetivo: descrever e caracterizar a dinâmica do humor de atletas de handebol de elite ao '
 f'longo de um microciclo pré-competitivo, com ênfase no eixo energia–fadiga, no comportamento dos seis perfis de humor, '
 f'na resposta aguda pré e pós-treino e na forma temporal das trajetórias. Método: {sm["n"]} atletas do sexo masculino '
 f'responderam à Escala de Humor de Brunel (BRUMS-24) durante sete dias, com uma coleta de linha de base e duas coletas '
 f'diárias (pré e pós-treino) nos seis dias de treino, e um total de {sm["n_obs"]} observações. Além da estatística '
 f'descritiva, da consistência interna e da classificação dos perfis, as comparações entre dias e entre pré e pós-treino '
 f'incluíram o tamanho e a magnitude do efeito, e as trajetórias foram analisadas por suavização, segundas derivadas, '
 f'ajuste polinomial e localização dos cruzamentos entre as dimensões. Resultados: a deterioração concentrou-se no eixo '
 f'energia–fadiga, com o vigor em queda e a fadiga em elevação do primeiro para o último dia, com efeito grande '
 f'(d = {d1d7_vig_dz} e d = {d1d7_fad_dz}) e confirmação multivariada (Wilks λ = {wilks}). As trajetórias suavizadas '
 f'apresentaram um ponto de inflexão na metade da semana, e a fadiga ultrapassou o vigor de forma definitiva no fim do '
 f'microciclo, com a PTH a superar ambas as dimensões. A prevalência dos perfis deslocou-se do iceberg ({ice_d1}%% no '
 f'primeiro dia) para a barbatana de tubarão ({bar_d7}%% no último dia), com aumento da chance desse perfil a cada dia '
 f'(OR = {or_bar}). Conclusão: o humor migrou da prontidão para a fadiga funcional, em um padrão compatível com '
 f'sobre-esforço funcional, o que recomenda centrar o monitoramento no par vigor–fadiga e no cruzamento entre as suas '
 f'curvas.'.replace('%%','%'),False)],after=4)
RUN([('Palavras-chave: ',True),('humor; BRUMS; handebol; perfis de humor; fadiga; monitoramento do atleta',False)],after=8,ind=False)

# ===== 1. INTRODUÇÃO =====
H('1. Introdução')
P(f'O monitoramento do estado do atleta consolidou-se como parte da gestão do treinamento no esporte de rendimento, '
 f'e integra medidas objetivas de carga externa e interna a instrumentos subjetivos de autorrelato {cite("helwig2023","kellmann2018")}. '
 f'Entre esses instrumentos, os questionários de humor destacam-se pela praticidade, pelo baixo custo e pela '
 f'sensibilidade às variações da carga, e mostram utilidade para o acompanhamento do bem-estar e do desempenho '
 f'{cite("saw2016","lochbaum2021")}. Revisões sistemáticas indicam, inclusive, que as medidas subjetivas costumam '
 f'responder às alterações de carga com sensibilidade igual ou superior à de muitos marcadores objetivos, o que '
 f'sustenta o seu uso rotineiro {cite("saw2016")}. Por esse motivo, documentos de consenso recomendam a vigilância '
 f'regular da fadiga e da recuperação como base para as decisões de treino {cite("kellmann2018")}.')
P(f'A Escala de Humor de Brunel (BRUMS) operacionaliza essa avaliação por meio de seis dimensões, a saber, tensão, '
 f'depressão, raiva, vigor, fadiga e confusão, com propriedades psicométricas replicadas em diferentes idiomas e '
 f'culturas {cite("terry2003","terryltu2022")}. A versão em português dispõe de validação para o contexto brasileiro e '
 f'foi concebida, desde a origem, como ferramenta de detecção precoce de sinais associados ao excesso de treinamento '
 f'{cite("rohlfs2008")}. A partir das seis dimensões calcula-se a perturbação total do humor (PTH), um índice-resumo '
 f'do desequilíbrio afetivo que sintetiza o estado do atleta em um único valor e cuja associação com o desempenho tem '
 f'sido documentada em meta-análise {cite("lochbaum2021")}.')
P(f'Para além dos escores isolados, o modelo dos perfis de humor organiza as seis dimensões em padrões reconhecíveis. '
 f'O clássico perfil iceberg, com o vigor acima da média e as dimensões negativas rebaixadas, foi descrito no modelo de '
 f'saúde mental como marca do atleta em prontidão {cite("morgan1985")}. Estudos posteriores, com grandes amostras e '
 f'análise de agrupamento, identificaram e replicaram seis perfis distintos, entre os quais o iceberg, a barbatana de '
 f'tubarão, que sinaliza fadiga com vigor ainda preservado, e perfis de maior risco, como o iceberg invertido e o '
 f'submerso {cite("parsons2017","luojumaki2026")}. A leitura por perfis aproxima o dado psicométrico da linguagem do '
 f'treinador, facilita a comunicação com a comissão técnica e tem sido proposta como recurso de rastreio da saúde '
 f'mental no esporte {cite("terry2021","han2020","lew2023")}.')
P(f'Entre as seis dimensões, o vigor e a fadiga formam o eixo mais responsivo à carga e concentram boa parte do valor '
 f'prático do monitoramento. Sob intensificação do treino, o vigor tende a cair e a fadiga a subir, um padrão observado '
 f'em atletas de diferentes modalidades e sensível também à privação de sono {cite("ferreira2026","pierce2002")}. Esse '
 f'eixo responde ainda de forma aguda a cada sessão ou competição, com oscilações mensuráveis entre os momentos pré e '
 f'pós-esforço, o que recomenda coletas repetidas dentro do dia, e não apenas entre dias {cite("nascimento2026")}.')
P(f'A leitura conjunta dessas variações remete ao continuum do sobre-esforço. A distinção entre o sobre-esforço '
 f'funcional, o sobre-esforço não funcional e a síndrome de overtraining é hoje descrita como um processo gradual, de '
 f'difícil diagnóstico por um único marcador, no qual a fadiga central e a piora do humor figuram entre os sinais mais '
 f'precoces {cite("roete2021","latorre2023","manescu2026")}. Nesse quadro, o acompanhamento do humor oferece um '
 f'marcador sensível, de baixo custo e não invasivo, complementar aos marcadores fisiológicos, para vigiar a janela em '
 f'que o sobre-esforço funcional é buscado de forma planejada {cite("thorpe2017")}.')
P(f'O handebol oferece um cenário exigente para essa vigilância. A modalidade é coletiva, intermitente e de alta '
 f'intensidade, com sprints curtos, mudanças de direção, saltos, arremessos e contato físico, o que impõe elevada '
 f'demanda neuromuscular e psicofisiológica, variável por posição de jogo {cite("karcher2014","garcia2023","carton2023")}. '
 f'O monitoramento longitudinal da carga interna ao longo da temporada e a atenção à fadiga, inclusive a mental, já '
 f'foram descritos como relevantes para o rendimento na modalidade {cite("struzik2026","staiano2025")}. A acumulação '
 f'de carga na semana tende, assim, a corroer o vigor e a elevar a fadiga, um padrão que, quando controlado, '
 f'caracteriza o sobre-esforço funcional e antecede a recuperação planejada.')
P(f'Apesar do interesse crescente, poucos estudos descrevem, dentro de um único microciclo de handebol de elite, a '
 f'migração dos perfis de humor e a forma temporal exata das trajetórias do eixo energia–fadiga, com coletas pré e '
 f'pós-treino que capturam também a variação dentro do dia {cite("rohlfs2024","bird2025","ratz2026")}. O objetivo deste '
 f'estudo foi descrever e caracterizar a dinâmica do humor de atletas de handebol de elite ao longo de um microciclo '
 f'pré-competitivo, com quatro focos articulados: (i) o comportamento dos seis perfis de humor e a sua migração entre o '
 f'primeiro e o último dia; (ii) a magnitude dos efeitos no eixo energia–fadiga, entre dias e entre pré e pós-treino; '
 f'(iii) a forma temporal das trajetórias de vigor, fadiga e PTH, por meio de suavização, análise das segundas '
 f'derivadas, ajuste polinomial e localização dos cruzamentos exatos entre as dimensões; e (iv) a leitura desse conjunto '
 f'como um quadro de sobre-esforço funcional útil ao monitoramento aplicado. A Figura {_FN[0]+1} resume o quadro '
 f'conceitual que orienta essas análises.')
figure(f'{FG}/framework.png','Quadro conceitual do monitoramento do humor no eixo energia–fadiga: a carga do microciclo altera as dimensões do BRUMS, com queda do vigor e elevação da fadiga e da PTH, o que desloca os perfis de humor em direção à barbatana de tubarão, dentro da janela do sobre-esforço funcional, e retroalimenta as decisões de treino e recuperação.',w=15.5)

# ===== 2. MATERIAIS E MÉTODOS =====
H('2. Materiais e Métodos')
H('2.1. Participantes',after=2)
P(f'Participaram {sm["n"]} atletas de handebol do sexo masculino, de nível de elite (idade de {idade} ± {idade_sd} anos; '
 f'{exp} ± {exp_sd} anos de prática), das posições de armador, ala, pivô e goleiro. Todos treinavam regularmente no '
 f'clube e estavam aptos no período do estudo.')
H('2.2. Instrumento',after=2)
P(f'O humor foi avaliado pela BRUMS-24, que reúne seis dimensões (tensão, depressão, raiva, vigor, fadiga e confusão) '
 f'com escores de 0 a 16, a partir das quais se calculou também a perturbação total do humor (PTH). O instrumento dispõe '
 f'de validação e de propriedades psicométricas descritas para atletas brasileiros {cite("rohlfs2008","rohlfs2023")}.')
H('2.3. Procedimento',after=2)
P(f'O delineamento cobriu sete dias de um microciclo pré-competitivo, com uma coleta de linha de base no primeiro dia e '
 f'duas coletas diárias, uma antes e outra depois do treino, nos seis dias subsequentes, o que totalizou {sm["n_obs"]} '
 f'observações. As coletas ocorreram em ambiente padronizado, antes das refeições e do início da sessão, no caso do '
 f'momento pré-treino. A Figura {_FN[0]+1} sintetiza o delineamento e o plano de análise.')
figure(f'{FG}/organograma.png','Organograma do delineamento do estudo: da amostra e do microciclo às coletas de linha de base e pré/pós-treino, às 286 observações do BRUMS-24 e aos cinco blocos do plano de análise.',w=15.5)
H('2.4. Análise dos Dados',after=2)
P('Empregaram-se estatística descritiva das seis dimensões e da PTH e avaliação da consistência interna por alfa de '
 'Cronbach. Como o teste de Shapiro-Wilk apontou desvios da normalidade em parte das dimensões, adotaram-se testes não '
 'paramétricos: o teste de Wilcoxon para a comparação entre pré e pós-treino, o teste de Friedman com o W de Kendall '
 'para a comparação entre os sete dias e a correlação de Spearman para as relações entre as dimensões, sempre com o '
 'cálculo do tamanho do efeito, classificado por magnitude. A confirmação multivariada da diferença entre o primeiro e '
 'o último dia recorreu à MANOVA em escores T. Cada observação foi classificada em um dos seis perfis de humor a partir '
 'dos escores padronizados, e a distribuição dos perfis foi avaliada pelo teste do qui-quadrado, com a tendência de '
 'migração resumida por uma regressão logística. A consistência das medidas repetidas foi estimada pelo coeficiente de '
 'correlação intraclasse (ICC), do qual se derivaram os limiares de mudança, a saber, o erro-padrão de medida, a '
 'mudança mínima detectável e a menor mudança relevante.')
P('A forma temporal das trajetórias foi analisada por três abordagens complementares. Primeiro, as séries de vigor, '
 'fadiga e PTH foram suavizadas por spline, e a segunda derivada de cada curva localizou o ponto de inflexão, definido '
 'como a raiz da segunda derivada, no qual a concavidade muda e a taxa de variação atinge o seu limite. Segundo, cada '
 'trajetória foi modelada por um ajuste polinomial de grau três, em duas resoluções, sobre as sete médias diárias e '
 'sobre os quatorze pontos pré e pós-treino, do qual se extraíram a equação, o coeficiente de determinação (R²) e a '
 'inflexão analítica (x = -b/3a). Terceiro, sobre as curvas ajustadas foram localizados os cruzamentos exatos entre as '
 'dimensões, ou seja, os dias em que uma curva iguala e ultrapassa a outra. As análises foram conduzidas em Python '
 '(pacotes NumPy, SciPy, pandas, statsmodels e Plotly), com nível de significância de 5%.')

# ===== 3. RESULTADOS =====
H('3. Resultados')
H('3.1. Descrição e Consistência das Dimensões',after=2)
P('A Tabela 1 resume as seis dimensões do BRUMS e a PTH no conjunto das observações. O vigor apresentou a maior média '
 'entre as dimensões, ao passo que as dimensões negativas concentraram-se em valores baixos, com médias próximas do '
 'limite inferior da escala.')
rows1=[[lab,num(desc[k]['mean'],1),num(desc[k]['sd'],1),f"{num(desc[k]['mn'],0)}–{num(desc[k]['mx'],0)}"] for k,lab in ORD]
table('Estatística descritiva das dimensões do BRUMS e da perturbação total do humor (286 observações).',
 ['Dimensão','Média','DP','Mín.–Máx.'],rows1,
 note='DP: desvio-padrão. PTH: perturbação total do humor. Escores das dimensões variam de 0 a 16.')
P(f'A Figura {_FN[0]+1} apresenta os diagramas de caixa das seis dimensões, em painéis separados; como as dimensões '
 f'diferem muito em amplitude, cada painel usa a escala da própria variável. O vigor e a fadiga percorrem a maior parte '
 f'da escala e apresentam média e mediana próximas, ao passo que as dimensões negativas concentram a caixa junto ao zero '
 f'e exibem a média acima da mediana, o que revela assimetria positiva e efeito de piso.')
figure(f'{FG}/box_humor.png','Diagramas de caixa das seis dimensões do BRUMS, em painéis separados; cada painel usa a escala da própria variável, com a mediana (linha sólida) e a média (linha tracejada e losango).',w=14.0)
P(f'A consistência interna foi adequada nas duas dimensões do eixo energia–fadiga (alfa de {a_vig} para o vigor e {a_fad} '
 f'para a fadiga). A PTH associou-se de forma forte ao vigor e à fadiga (correlações de {pth_rv} e {pth_rf}), e essas '
 f'duas dimensões explicaram {pth_r2}% da sua variância, o que confirma o eixo energia–fadiga como o núcleo do sinal.')
IC=STAT['icc']
table('Consistência das medidas repetidas ao longo da semana (coeficiente de correlação intraclasse, ICC).',
 ['Dimensão','ICC(2,1)','ICC(2,k)','Consistência'],
 [[l,num(IC[k]['icc1'],2),num(IC[k]['icck'],2),IC[k]['cls']] for k,l in [('Vigor','Vigor'),('Fadiga','Fadiga'),('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Confusao','Confusão')]],
 note='ICC(2,1) = medida isolada; ICC(2,k) = média das medidas. Consistência: < 0,50 pobre; 0,50–0,75 moderada; > 0,75 boa.')
P(f'A partir do ICC derivaram-se os limiares de mudança do eixo energia–fadiga (Tabela {_TN[0]+1}). No vigor, o '
 f'erro-padrão de medida foi de {num(MDC["Vigor"]["sem"],1)} ponto e a mudança mínima detectável, de '
 f'{num(MDC["Vigor"]["mdc90"],1)} a {num(MDC["Vigor"]["mdc95"],1)} pontos; na fadiga, de {num(MDC["Fadiga"]["sem"],1)} e '
 f'de {num(MDC["Fadiga"]["mdc90"],1)} a {num(MDC["Fadiga"]["mdc95"],1)} pontos. Esses valores superam a menor mudança '
 f'relevante (SWC), o que indica que, no plano individual, apenas oscilações de cerca de {num(MDC["Vigor"]["mdc90"],1)} a '
 f'{num(MDC["Fadiga"]["mdc90"],1)} pontos podem ser lidas como mudança real; por isso a inferência se apoia nas '
 f'tendências de grupo.')
table('Erro de medida e limiares de mudança do vigor e da fadiga (escala 0–16): erro-padrão de medida (SEM), mudança mínima detectável (MDC) e menor mudança relevante (SWC).',
 ['Dimensão','SEM','MDC90','MDC95','SWC'],
 [[l,num(MDC[k]['sem'],1),num(MDC[k]['mdc90'],1),num(MDC[k]['mdc95'],1),num(MDC[k]['swc'],1)] for k,l in [('Vigor','Vigor'),('Fadiga','Fadiga')]],
 note='SEM = DP × √(1 - ICC); MDC = 1,65 (90%) ou 1,96 (95%) × √2 × SEM; SWC = 0,2 × DP entre atletas.')

H('3.2. Perfis de Humor e sua Migração',after=2)
P(f'Os seis perfis de humor estiveram representados na amostra. A Figura {_FN[0]+1} apresenta cada perfil identificado no '
 f'estudo em escores T, com a respectiva prevalência. O perfil iceberg exprime prontidão, com o vigor acima da média e '
 f'as dimensões negativas abaixo, ao passo que a barbatana de tubarão sinaliza pico isolado de fadiga [7].')
figure(f'{FG}/xb6_clusters.png','Perfis de humor identificados na amostra, em escores T (M = 50; DP = 10); prevalência de cada perfil entre parênteses.',w=13.0)
P(f'A comparação entre o primeiro e o último dia revelou a reconfiguração do perfil médio do grupo (Figura {_FN[0]+1}). '
 f'No primeiro dia, o perfil assumiu o formato iceberg; no último, inverteu-se para a barbatana de tubarão, com a fadiga '
 f'no topo e o vigor rebaixado.')
figure(f'{FG}/xb5_profile_d1d7.png','Perfil de humor em escores T no primeiro e no último dia do microciclo.',w=11.5)
P(f'Essa mudança correspondeu a um deslocamento da prevalência dos perfis (Tabela {_TN[0]+1}). O perfil iceberg caiu de '
 f'{ice_d1}% para {ice_d7}%, a barbatana de tubarão subiu de {bar_d1}% para {bar_d7}% e o submerso passou de {sub_d1}% '
 f'para {sub_d7}%. A reorganização categórica não alcançou significância no qui-quadrado (χ² = {num(PREV["chi"],2)}; '
 f'p {pv(PREV["p"])}), resultado esperado pela baixa contagem por célula, mas a regressão logística confirmou o aumento '
 f'da chance da barbatana de tubarão a cada dia (OR = {or_bar}).')
PROFR=[('Iceberg','Iceberg'),('Everest invertido','Everest invertido'),('Iceberg invertido','Iceberg invertido'),('Submerso','Submerso'),('Barbatana tubarão','Barbatana de tubarão'),('Superfície','Superfície')]
table('Distribuição dos seis perfis de humor no primeiro e no último dia do microciclo: n (%).',
 ['Perfil','Dia 1, n (%)','Dia 7, n (%)'],
 [[l,f"{PREV['D1'][p]} ({num(100*PREV['D1'][p]/PREV['n_d1'],1)}%)",f"{PREV['D7'][p]} ({num(100*PREV['D7'][p]/PREV['n_d7'],1)}%)"] for p,l in PROFR],
 note='χ² = %s; p = %s. A migração é descritiva e converge com a queda do vigor e a elevação da fadiga.'%(num(PREV['chi'],2),num(PREV['p'],3)))

H('3.3. Diferença entre o Primeiro e o Último Dia',after=2)
P(f'A comparação entre o primeiro e o último dia quantificou a magnitude da mudança em cada dimensão (Tabela {_TN[0]+1}). '
 f'O efeito foi grande no vigor, que caiu {d1d7_vig_pct}% (dz = {d1d7_vig_dz}), e na fadiga, que subiu {d1d7_fad_pct}% '
 f'(dz = {d1d7_fad_dz}). A análise multivariada confirmou a diferença global (Wilks λ = {wilks}; p {pv(MV["d1d7"]["p_mv"])}), '
 f'e, sob a correção de Bonferroni, apenas o vigor e a fadiga permaneceram significativos.')
D7ORD=[('Vigor','Vigor'),('Fadiga','Fadiga'),('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Confusao','Confusão'),('TMD','PTH')]
table('Diferença das dimensões do BRUMS e da PTH entre o primeiro e o último dia (com tamanho e magnitude do efeito).',
 ['Dimensão','Dia 1 (M)','Dia 7 (M)','Variação (%)','p','dz','Magnitude'],
 [[l,num(d17[k]['d1'],2),num(d17[k]['d7'],2),c2(f"{d17[k]['pct']:+.0f}")+'%',pvt(d17[k]['p']),c2(f"{d17[k]['dz']:+.2f}"),d17[k]['mag']] for k,l in D7ORD],
 note='dz = tamanho de efeito intraindividual; magnitude: trivial (< 0,2); pequeno (0,2–0,5); médio (0,5–0,8); grande (> 0,8).')

H('3.4. Comparação das Dimensões entre os Dias',after=2)
P(f'O teste de Friedman (Tabela {_TN[0]+1}) apontou variação significativa no vigor (p {fr_vig_p}; W = {fr_vig_w}) e na '
 f'fadiga (p {fr_fad_p}; W = {fr_fad_w}), com magnitude pequena a moderada, além de diferenças na tensão e na confusão.')
FR=S4['friedman']
def wmag(w): return 'trivial' if w<0.1 else 'pequeno' if w<0.3 else 'moderado' if w<0.5 else 'grande'
table('Médias diárias das dimensões do BRUMS e teste de Friedman com W de Kendall e magnitude do efeito.',
 ['Dimensão','D1','D2','D3','D4','D5','D6','D7','χ²','p','W','Magnitude'],
 [[lab]+[num(CRS['dec'][k]['dm'][str(d)],1) for d in range(1,8)]+[num(FR[k]['chi'],1),pvt(FR[k]['p']),num(FR[k]['W'],2),wmag(FR[k]['W'])] for k,lab in ORD],
 note='W de Kendall: trivial (< 0,1); pequeno (0,1–0,3); moderado (0,3–0,5); grande (> 0,5). PTH: perturbação total do humor.',fs=8)
P(f'A Figura {_FN[0]+1} ilustra a trajetória do vigor, da fadiga e da PTH, com a queda progressiva do vigor e a elevação '
 f'da fadiga em direção ao fim do microciclo, e a Figura S1 detalha, por dia, a distribuição de cada dimensão.')
figure(f'{FG}/abnt_f1_trajetoria.png','Trajetória de vigor, fadiga e perturbação total do humor ao longo dos sete dias (curvas suavizadas; ponto de inflexão e extremos sinalizados).',w=12.5)
sfig(f'{FG}/xb3_box.png','Diagramas de caixa das seis dimensões do BRUMS por dia do microciclo (áreas sombreadas: início e acúmulo da semana).',w=15.0)
P(f'A comparação direta da PTH entre três momentos do microciclo (Figura S2) ilustra a mesma tendência: a '
 f'perturbação total do humor aumentou do primeiro para o último dia, e a diferença foi significativa entre o Dia 1 e o '
 f'Dia 7 (p = 0,008), embora as comparações entre dias adjacentes não tenham alcançado significância.')
sfig(f'{FG}/box_signif.png','Perturbação total do humor (PTH) no Dia 1, no Dia 4 e no Dia 7, com as comparações par a par (teste de Mann-Whitney; * p < 0,05; ** p < 0,01; ns = não significativo).',w=13.0)
P(f'A distribuição completa da PTH em cada dia (Figura S3) e a sua distribuição acumulada (Figura S4) confirmam o deslocamento ao longo da semana. Os diagramas de violino mostram a densidade da PTH a subir do Dia 1 ao Dia 7, e as curvas de distribuição acumulada revelam um deslocamento sistemático para a direita, com o Dia 7 situado acima do Dia 1 em praticamente toda a faixa de escores.')
sfig(f'{FG}/tec_violin.png','Distribuição da PTH por dia (diagramas de violino com caixa e média; a escala foi ajustada à faixa central da distribuição para facilitar a leitura).',w=13.5)
sfig(f'{FG}/tec_ecdf.png','Distribuição acumulada (ECDF) da PTH no Dia 1, no Dia 4 e no Dia 7; o deslocamento das curvas para a direita indica o aumento da perturbação ao longo do microciclo.',w=13.0)

H('3.5. Suavização, Derivadas e Cruzamento das Trajetórias',after=2)
P(f'Para separar o sinal do ruído, as trajetórias do vigor, da fadiga e da PTH foram suavizadas sobre os doze pontos '
 f'pré e pós-treino por meio de uma spline, e a segunda derivada de cada curva localizou o seu ponto de inflexão, ou '
 f'seja, o momento em que a concavidade muda e a taxa de variação atinge o seu limite (Figura {_FN[0]+1}; '
 f'Tabela {_TN[0]+1}). Depois de removido o ruído, cada dimensão exibiu um único ponto de inflexão, situado na metade da '
 f'semana (em torno do Dia {num(SMO["V"]["infl"][0],1) if SMO["V"]["infl"] else "n/d"}), o que marca a transição entre a '
 f'fase inicial de prontidão e a fase de acúmulo de carga. A curva suavizada do vigor subiu até um máximo em torno do '
 f'Dia {num(SMO["V"]["ymax_x"],1)} e caiu para o seu mínimo ao fim da semana, enquanto a fadiga e a PTH percorreram o '
 f'caminho inverso e atingiram os seus máximos no último dia.')
figure(f'{FG}/smooth_deriv.png','Trajetórias suavizadas do vigor, da fadiga e da PTH sobre os doze pontos pré e pós-treino. Marcadores translúcidos: sinal bruto; linha grossa: sinal suavizado; linha pontilhada: ponto de inflexão (segunda derivada nula); triângulos: máximo e mínimo. Áreas sombreadas: início e acúmulo da semana.',w=10.5)
def smrow(k,lab):
    r=SMO[k]; inf=('Dia '+num(r['infl'][0],1)) if r['infl'] else 'n/d'
    return [lab,inf,f"Dia {num(r['ymax_x'],1)} ({num(r['ymax'],1)})",f"Dia {num(r['ymin_x'],1)} ({num(r['ymin'],1)})"]
table('Ponto de inflexão (segunda derivada nula) e extremos das trajetórias suavizadas do eixo energia–fadiga.',
 ['Dimensão','Inflexão (dia)','Máximo: dia (escore)','Mínimo: dia (escore)'],
 [smrow('V','Vigor'),smrow('F','Fadiga'),smrow('P','PTH')],
 note='Curvas suavizadas por spline sobre os 12 pontos pré/pós; inflexão = raiz da segunda derivada. PTH: perturbação total do humor.')
P(f'A mesma inflexão foi examinada por dois caminhos complementares, detalhados no material suplementar. O painel da '
 f'segunda derivada localiza o cruzamento por zero de cada curva e realça os extremos com o respectivo valor (Figura S7). '
 f'Em paralelo, um ajuste polinomial de grau três resume cada trajetória por uma equação e reproduz a inflexão de forma '
 f'analítica, na raiz da segunda derivada (x = -b/3a), em duas resoluções (Figura S8). Sobre as sete médias diárias, o '
 f'ajuste foi forte, com R² de {num(POLY["V"]["d7"]["r2"],2)} para o vigor, {num(POLY["F"]["d7"]["r2"],2)} para a fadiga e '
 f'{num(POLY["P"]["d7"]["r2"],2)} para a PTH, e inflexões nos dias {num(POLY["V"]["d7"]["infl"],1)}, '
 f'{num(POLY["F"]["d7"]["infl"],1)} e {num(POLY["P"]["d7"]["infl"],1)}. Sobre os quatorze pontos pré e pós-treino, '
 f'incluída a linha de base do primeiro dia, o R² foi de {num(POLY["V"]["d14"]["r2"],2)}, {num(POLY["F"]["d14"]["r2"],2)} '
 f'e {num(POLY["P"]["d14"]["r2"],2)}, com inflexões praticamente idênticas. As duas resoluções convergem para a transição '
 f'na metade da semana e reforçam, por via paramétrica, o achado da spline.')
vf=CRX['cross']['VF']; vp=CRX['cross']['VP']; fp=CRX['cross']['FP']
P(f'Como o vigor e a fadiga compartilham a mesma escala, as suas trajetórias podem ser comparadas de forma direta e os '
 f'seus cruzamentos exatos podem ser localizados (Figura {_FN[0]+1}). No primeiro dia, o vigor superou a fadiga por ampla '
 f'margem (média de {num(d17["Vigor"]["d1"],1)} contra {num(d17["Fadiga"]["d1"],1)} pontos). O vigor caiu e a fadiga '
 f'subiu até que as duas curvas se igualaram pela primeira vez em torno do dia {num(vf[0][0],1)} (escore '
 f'{num(vf[0][1],1)}). A partir daí, o vigor e a fadiga percorreram uma faixa estreita e próxima e voltaram a cruzar-se '
 f'nos dias {num(vf[1][0],1)} e {num(vf[-1][0],1)}, sinal de um equilíbrio instável entre energia e fadiga durante a '
 f'maior parte da semana. Após o último cruzamento, no dia {num(vf[-1][0],1)} (escore {num(vf[-1][1],1)}), a fadiga '
 f'afastou-se de forma definitiva acima do vigor. No mesmo trecho final, a perturbação total do humor, que vinha abaixo '
 f'das duas dimensões, subiu e ultrapassou primeiro o vigor (dia {num(vp[0][0],1)}; escore {num(vp[0][1],1)}) e depois a '
 f'fadiga (dia {num(fp[0][0],1)}; escore {num(fp[0][1],1)}), o que marca a deterioração conjunta do estado de humor no '
 f'fim do microciclo.')
figure(f'{FG}/cross_traj.png','Cruzamentos exatos das trajetórias de vigor, fadiga e PTH ao longo do microciclo (curvas do ajuste polinomial sobre as médias diárias; losangos: pontos de cruzamento, com o dia e o escore). Áreas sombreadas: início e acúmulo da semana.',w=15.0)

H('3.6. Dias de Pico e Comparação de Cada Dia ao Primeiro',after=2)
P(f'A localização dos picos sintetiza a dinâmica semanal (Tabela {_TN[0]+1}): o vigor foi máximo no Dia {vmaxd} e mínimo '
 f'no Dia {vmind}, ao passo que a fadiga foi máxima no Dia {fmaxd}. O início da semana reúne maior prontidão e o final '
 f'concentra a fadiga e a raiva.')
table('Dia de maior e de menor expressão de cada dimensão do humor (médias diárias).',
 ['Dimensão','Dia de maior valor','Valor','Dia de menor valor','Valor'],
 [[l,f'Dia {PK[k]["max_day"]}',num(PK[k]['max_val'],2),f'Dia {PK[k]["min_day"]}',num(PK[k]['min_val'],2)] for k,l in ORD[:6]])
def emm(v,d): return num(PHJ[v]['emm'][str(d)],2)
def sig1(v,d): return '' if d==1 or PHJ[v]['pairs'].get('1_%d'%d,{}).get('ptukey',1)>=0.05 else '*'
np_vig=sum(1 for _,pp in PHJ['Vigor']['pairs'].items() if pp['ptukey']<0.05)
np_fad=sum(1 for _,pp in PHJ['Fadiga']['pairs'].items() if pp['ptukey']<0.05)
P(f'No pós-teste que compara todos os dias entre si (Tabela {_TN[0]+1}; Figura S5), o vigor diferiu de forma '
 f'significativa em {np_vig} dos 21 pares de dias e a fadiga em {np_fad}, sempre no sentido de piora em relação aos '
 f'primeiros dias, o que confirma a deterioração progressiva do eixo energia–fadiga.')
table('Médias marginais estimadas por dia do vigor e da fadiga, com comparação de cada dia ao Dia 1 (pós-teste de Tukey).',
 ['Dia','Vigor','Fadiga'],
 [[f'Dia {d}',emm('Vigor',d)+sig1('Vigor',d),emm('Fadiga',d)+sig1('Fadiga',d)] for d in range(1,8)],
 note='* diferença significativa em relação ao Dia 1 (Tukey, p < 0,05).')
sfig(f'{FG}/ph_emm.png','Médias marginais diárias do vigor, da fadiga e da fadiga física, com comparação de todos os dias ao Dia 1 (* p < 0,05).',w=14.5)

H('3.7. Variação entre Pré e Pós-Treino',after=2)
P(f'O teste de Wilcoxon (Tabela {_TN[0]+1}) evidenciou uma resposta aguda coerente com o esforço: o vigor caiu '
 f'(d = {vig_dz}) e a fadiga e a PTH subiram do momento pré para o pós, com efeito de magnitude pequena a moderada. A '
 f'Figura S6 apresenta esses escores por dia e mostra que a diferença entre pré e pós-treino se mantém ao longo da semana.')
table('Comparação entre pré e pós-treino das dimensões do BRUMS e da PTH (teste de Wilcoxon e d de Cohen).',
 ['Dimensão','Pré (M)','Pós (M)','Variação (%)','p','d','Magnitude'],
 [[l,num(pr[k]['pre'],2),num(pr[k]['pos'],2),c2(f"{pr[k]['pct']:+.0f}")+'%',pvt(pr[k]['p']),c2(f"{pr[k]['dz']:+.2f}"),pr[k]['mag']] for k,l in ORD],
 note='p do teste de Wilcoxon; d = tamanho de efeito de Cohen. PTH: perturbação total do humor.')
sfig(f'{FG}/abnt_f_prepos.png','Escores de vigor, fadiga e perturbação total do humor no pré e no pós-treino, por dia.',w=12.5)

H('3.8. Relações entre as Dimensões',after=2)
P(f'As correlações de Spearman (Tabela {_TN[0]+1}) mostraram que as dimensões negativas se associam entre si, com '
 f'destaque para os pares depressão–raiva e tensão–confusão, e que a fadiga se relaciona com a depressão e a raiva. O '
 f'vigor manteve-se relativamente independente, o que reforça a sua leitura como um polo próprio do eixo energia–fadiga.')
sigpairs=[x for x in STAT['pairs'] if x['p']<0.05]
table('Correlações de Spearman significativas entre as dimensões do BRUMS (n = %d atletas).'%sm['n'],
 ['Par de dimensões','ρ','p'],
 [[f"{x['a']} × {x['b']}",c2(f"{x['rho']:+.2f}"),pvt(x['p'])] for x in sigpairs],
 note='ρ = coeficiente de correlação de Spearman. Apresentam-se apenas os pares com p < 0,05.')
P(f'A relação entre o vigor e a fadiga, além de negativa, tornou-se mais estreita à medida que a carga se acumulou '
 f'(Figura {_FN[0]+1}): a regressão por fase do microciclo mostrou um acoplamento fraco na fase inicial (r = -0,35) e '
 f'mais forte na fase de acúmulo (r = -0,51), o que sugere que os dois polos do eixo energia–fadiga passam a variar de '
 f'forma mais solidária sob fadiga acumulada.')
figure(f'{FG}/tec_regr.png','Relação entre vigor e fadiga por fase do microciclo, com a reta de regressão de cada grupo (início: dias 1 a 4; acúmulo: dias 5 a 7).',w=12.5)

H('3.9. Estrutura Dimensional (Análise de Componentes Principais)',after=2)
P(f'Uma análise exploratória de componentes principais resumiu a estrutura das seis dimensões. Os dois primeiros '
 f'componentes explicaram {num(100*(PCA["var_ratio"][0]+PCA["var_ratio"][1]),0)}% da variância '
 f'({num(100*PCA["var_ratio"][0],0)}% no primeiro e {num(100*PCA["var_ratio"][1],0)}% no segundo). O círculo de '
 f'correlação (Figura {_FN[0]+1}) mostra que as dimensões negativas (depressão, raiva, confusão e fadiga) se projetam '
 f'juntas no lado positivo do primeiro componente, que funciona como um eixo geral de perturbação, ao passo que o vigor '
 f'aponta em sentido oposto. Essa oposição entre o vigor e as dimensões de perturbação reforça, por via independente, a '
 f'centralidade do eixo energia–fadiga.')
figure(f'{FG}/pca_circulo.png','Círculo de correlação da análise de componentes principais das seis dimensões do BRUMS. A espessura de cada seta é proporcional à contribuição da variável ao plano dos dois primeiros componentes.',w=11.5)

# ===== 4. DISCUSSÃO =====
H('4. Discussão')
H('4.1. Migração do Humor e o Eixo Energia–Fadiga',after=2)
P(f'Os resultados descrevem, em um microciclo pré-competitivo de handebol de elite, a migração do humor da prontidão '
 f'para a fadiga funcional. A queda do vigor e a elevação da fadiga entre o primeiro e o último dia alcançaram efeito '
 f'grande e foram confirmadas pela análise multivariada, o que situa o eixo energia–fadiga como o principal responsável '
 f'pela mudança do estado do atleta. Esse padrão reproduz, dentro de um microciclo de handebol, o comportamento já '
 f'documentado em outras modalidades sob acúmulo de carga e sob treino intensificado {cite("pierce2002","ferreira2026","thorpe2017")}.')
P(f'À tendência semanal somou-se uma resposta aguda coerente com o esforço. Entre o momento pré e o pós-treino, o vigor '
 f'caiu e a fadiga e a PTH subiram, com magnitude de pequena a moderada, o que acrescenta uma camada de variação dentro '
 f'do dia à deterioração observada entre dias. Esse comportamento converge com o de outros esportes coletivos, nos quais '
 f'o humor responde de forma imediata ao evento esportivo, e reforça o valor de coletas repetidas dentro do dia, e não '
 f'apenas entre dias {cite("nascimento2026")}.')
P(f'A perturbação total do humor comportou-se como um integrador desse eixo. O vigor e a fadiga associaram-se de forma '
 f'forte à PTH e explicaram a maior parte da sua variância, ao passo que as dimensões negativas, concentradas em valores '
 f'baixos, contribuíram pouco para a variação. A consistência interna adequada do vigor e da fadiga e a sua correlação '
 f'intraclasse de moderada a boa convergem para a mesma conclusão, a saber, que o núcleo do sinal do monitoramento '
 f'reside no par energia–fadiga {cite("lochbaum2021")}.')
H('4.2. Perfis de Humor e Sobre-esforço Funcional',after=2)
P(f'A leitura por perfis acrescentou clareza a essa descrição. O deslocamento do perfil iceberg para a barbatana de '
 f'tubarão traduz a acumulação da carga sem os sinais de comprometimento da saúde mental, uma vez que os perfis de maior '
 f'risco, como o iceberg invertido e o submerso, não se instalaram. Quando comparados à distribuição de referência de '
 f'uma grande amostra brasileira e aos padrões descritos em análises de agrupamento, os nossos dados partem de um '
 f'predomínio semelhante de iceberg e elevam a barbatana de tubarão apenas ao fim da semana {cite("rohlfs2024","parsons2017","luojumaki2026")}.')
P(f'Essa reconfiguração tem valor comunicativo direto. O perfil condensa as seis dimensões em uma imagem única, o que '
 f'aproxima o dado psicométrico da linguagem da comissão técnica e favorece a decisão de treino e de recuperação. A '
 f'mesma leitura por perfis tem sido proposta como recurso de rastreio da saúde mental no esporte, o que amplia o '
 f'alcance do monitoramento para além do desempenho {cite("terry2021","han2020","lew2023")}.')
H('4.3. Contribuição Metodológica: Derivadas, Ajuste Polinomial e Cruzamentos',after=2)
P(f'Além de comparar médias, este estudo mapeou a forma temporal exata das trajetórias. A suavização das séries e o '
 f'cálculo da segunda derivada localizaram, em cada dimensão do eixo energia–fadiga, um único ponto de inflexão situado '
 f'na metade da semana, no qual a concavidade muda e a taxa de variação atinge o seu limite. Esse marcador objetivo '
 f'demarca a transição entre a fase inicial de prontidão e a fase de acúmulo de carga e oferece uma leitura que a '
 f'simples comparação entre o primeiro e o último dia não revela.')
P(f'O ajuste polinomial de grau três acrescentou uma síntese paramétrica a essa análise. O modelo reproduziu a inflexão '
 f'de forma analítica, na raiz da segunda derivada, e resumiu cada trajetória por uma equação, com ajuste forte sobre as '
 f'médias diárias (R² de {num(POLY["V"]["d7"]["r2"],2)} a {num(POLY["F"]["d7"]["r2"],2)} no eixo energia–fadiga). As '
 f'duas resoluções, sobre sete médias diárias e sobre quatorze pontos pré e pós-treino, apontaram inflexões quase '
 f'idênticas, o que indica que o achado da transição na metade da semana não depende do método nem da granularidade da '
 f'medida. A convergência entre a spline não paramétrica e o polinômio paramétrico reforça a robustez do resultado.')
P(f'A localização dos cruzamentos exatos representa a contribuição mais original desta abordagem. O vigor superou a '
 f'fadiga por ampla margem no início e foi alcançado por ela ainda na primeira metade da semana, mas apenas ao fim do '
 f'microciclo a fadiga afastou-se de forma definitiva acima do vigor, seguida pela PTH, que ultrapassou primeiro o vigor '
 f'e depois a fadiga. A datação desses cruzamentos converte a ideia qualitativa de inversão do eixo energia–fadiga em um '
 f'evento com dia e escore definidos, o que fornece um candidato a marcador temporal do início do sobre-esforço '
 f'funcional, coerente com o continuum descrito entre sobre-esforço e overtraining {cite("roete2021","manescu2026","latorre2023")}.')
H('4.4. Aplicação Prática e Perspectivas',after=2)
P(f'As características do handebol ajudam a explicar esse comportamento e orientam a aplicação. A modalidade impõe '
 f'esforço intermitente de alta intensidade, com sprints curtos, mudanças de direção, saltos e contato, o que gera '
 f'elevada demanda neuromuscular e psicofisiológica ao longo da semana e sustenta tanto a tendência de acúmulo quanto a '
 f'resposta aguda entre o pré e o pós-treino {cite("karcher2014","garcia2023","carton2023")}. Na prática, o par '
 f'vigor–fadiga concentra o sinal útil, e o cruzamento entre as suas curvas oferece um alerta simples e visual, capaz de '
 f'complementar o monitoramento de carga interna e de fadiga já descrito na modalidade {cite("struzik2026","staiano2025")}.')
P(f'O acompanhamento subjetivo do humor deve ser lido como parte de um monitoramento multidomínio, e não como medida '
 f'isolada. A sua sensibilidade, o baixo custo e o caráter não invasivo tornam-no complementar aos marcadores '
 f'fisiológicos, endócrinos e de bem-estar praticados em esportes coletivos {cite("ratz2026","bird2025","helwig2023")}. '
 f'Como agenda futura, recomenda-se integrar medidas objetivas de carga a este delineamento, ampliar a amostra e testar '
 f'se a datação do cruzamento do eixo energia–fadiga antecipa desfechos de fadiga e de desempenho, de modo a validar o '
 f'marcador temporal aqui proposto {cite("kellmann2018","saw2016")}.')

H('4.5. Limitações',after=2)
P('O estudo tem limitações, entre as quais a amostra de um único clube, o recorte de um microciclo e a ausência de '
 'medidas objetivas de carga neste recorte, o que restringe a generalização e impede inferências de causa. O pequeno '
 'número de observações distribuído por seis perfis reduz a potência do teste categórico, e o erro de medida de uma '
 'leitura isolada recomenda cautela na interpretação individual, de modo que a inferência se apoia nas tendências de '
 'grupo.')
H('4.6. Conclusões',after=2)
P(f'Em um microciclo pré-competitivo de handebol de elite, o humor migrou da prontidão para a fadiga funcional, com o '
 f'eixo energia–fadiga a concentrar os maiores efeitos e a sustentar a reconfiguração dos perfis do iceberg para a '
 f'barbatana de tubarão. A análise das trajetórias localizou um ponto de inflexão na metade da semana e datou o '
 f'cruzamento em que a fadiga ultrapassa o vigor no fim do microciclo, o que converte a inversão do eixo energia–fadiga '
 f'em um evento mensurável. O padrão é compatível com sobre-esforço funcional e não indicou instalação dos perfis de '
 f'maior risco. Do ponto de vista prático, recomenda-se centrar o monitoramento rotineiro no par vigor–fadiga ao longo '
 f'da semana, com atenção ao cruzamento entre as suas curvas, integrado a um acompanhamento multidomínio do estado do '
 f'atleta {cite("kellmann2018","ostapiuk2025")}.')

# ===== BACK MATTER (placeholders) =====
def bm(lbl,txt):
    p=doc.add_paragraph(); r=p.add_run(lbl+' '); r.bold=True; r.italic=True; r.font.size=Pt(9); r.font.name=BODY
    r2=p.add_run(txt); r2.font.size=Pt(9); r2.font.name=BODY; p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(0); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
doc.add_paragraph().paragraph_format.space_after=Pt(4)
bm('Materiais Suplementares:','As figuras suplementares (Figuras S1 a S8) estão disponíveis ao fim deste documento e detalham a distribuição diária das dimensões, as comparações par a par da PTH, os diagramas de violino e a distribuição acumulada, as médias marginais diárias, a resposta pré e pós-treino, a análise por derivadas e o ajuste polinomial das trajetórias.')
bm('Contribuições dos Autores:','[conceituação; metodologia; análise formal; investigação; redação do rascunho original; redação, revisão e edição]. Todos os autores leram e concordaram com a versão publicada do manuscrito.')
bm('Financiamento:','[Esta pesquisa não recebeu financiamento externo / nome da agência e número do processo].')
bm('Declaração do Comitê de Ética:','O estudo foi conduzido de acordo com a Declaração de Helsinque e aprovado pelo Comitê de Ética em Pesquisa de [Instituição] (CAAE [número]; parecer [número], [data]).')
bm('Declaração de Consentimento Informado:','O termo de consentimento livre e esclarecido foi obtido de todos os participantes envolvidos no estudo.')
bm('Declaração de Disponibilidade dos Dados:','[Os dados que sustentam os achados deste estudo estão disponíveis com o autor correspondente mediante solicitação razoável / restrições aplicam-se conforme acordo com a instituição].')
bm('Agradecimentos:','[Agradecimentos à comissão técnica e aos atletas participantes].')
bm('Conflitos de Interesse:','Os autores declaram não haver conflito de interesse.')

# ===== REFERÊNCIAS (estilo MDPI, ordem de citação, geradas pelo gerenciador) =====
H('Referências',before=10)
faltantes=[s for s in REFS if s not in _ORD]
if faltantes: print('AVISO refs nao citadas:',faltantes)
for i,slug in enumerate(_ORD,1):
    p=doc.add_paragraph(); rn=p.add_run('%d. '%i); rn.font.size=Pt(8.5); rn.font.name=BODY
    r=p.add_run(REFS[slug]); r.font.size=Pt(8.5); r.font.name=BODY
    p.paragraph_format.line_spacing=1.1; p.paragraph_format.space_after=Pt(2); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent=Cm(0.6); p.paragraph_format.first_line_indent=Cm(-0.6)

# ===== MATERIAL SUPLEMENTAR (figuras diferidas) =====
sfig(f'{FG}/deriv_impact.png','Análise por derivadas do eixo energia–fadiga. Em cada variável, o painel superior traz a curva suavizada com o máximo e o mínimo realçados em cápsulas de valor, e o painel inferior traz a segunda derivada (concavidade) com o cruzamento por zero destacado, que marca o ponto de inflexão da trajetória.',w=15.0)
sfig(f'{FG}/poly_fit.png','Ajuste polinomial de grau três das trajetórias do vigor, da fadiga e da PTH em duas resoluções: sobre as sete médias diárias (marcadores quadrados e linha cheia) e sobre os quatorze pontos pré e pós-treino, incluída a linha de base do primeiro dia (pontos translúcidos e linha tracejada). Cada painel traz a equação do ajuste diário, o R² de cada resolução e a inflexão analítica (raiz da segunda derivada, x = -b/3a).',w=14.5)
render_supp()

OUTP='/home/user/mdlucca/Artigos/Paper1_Humor_MDPI.docx'
doc.save(OUTP); print('SAVED',OUTP,'| Tabelas',_TN[0],'Figuras corpo',_FN[0],'Figuras supl.',_SN[0])
