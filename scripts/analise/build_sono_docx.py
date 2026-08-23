# -*- coding: utf-8 -*-
# Documento autônomo (pronto para o artigo): Resultados e Discussão — Sonolência (Epworth) e Estresse (PSS).
import warnings; warnings.filterwarnings('ignore')
import pandas as pd, numpy as np
from scipy import stats
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

d=pd.read_csv('humor_epworth_pss_anon.csv')
ad=d.groupby(['ID','dia'])[['Epworth','PSS','TMD','Vigor','Fadiga']].mean().reset_index()
def pv(p): return '< 0,001' if p<0.001 else ('%.3f'%p).replace('.',',')
def num(x,n=1): return ('%.*f'%(n,x)).replace('.',',')

doc=Document()
st=doc.styles['Normal']; st.font.name='Times New Roman'; st.font.size=Pt(12)
def H(t,sz=13):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(sz); r.font.name='Times New Roman'
    p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(4); return p
def P(t):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.space_after=Pt(8)
    p.paragraph_format.line_spacing=1.5; r=p.add_run(t); r.font.size=Pt(12); r.font.name='Times New Roman'; return p
def cap(t):
    p=doc.add_paragraph(); r=p.add_run(t); r.italic=True; r.font.size=Pt(10); r.font.name='Times New Roman'
    p.paragraph_format.space_after=Pt(10); return p
def setcell(c,txt,bold=False,align='center',size=10):
    c.text=''; pr=c.paragraphs[0]; pr.alignment={'center':WD_ALIGN_PARAGRAPH.CENTER,'left':WD_ALIGN_PARAGRAPH.LEFT}[align]
    r=pr.add_run(str(txt)); r.font.size=Pt(size); r.font.name='Times New Roman'; r.bold=bold

t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run('Sonolência diurna e estresse percebido ao longo de um microciclo pré-competitivo'); r.bold=True; r.font.size=Pt(14)

H('Métodos (resumo)')
P('Em cada envio do questionário diário (27 atletas; 456 observações; 21–27/04/2024), além do BRUMS foram '
  'registrados a sonolência diurna pela Escala de Sonolência de Epworth (0–24) e o estresse percebido pela '
  'Perceived Stress Scale (PSS). Ambos os instrumentos tiveram 100% de preenchimento. As variáveis foram '
  'resumidas por atleta e por dia; a mudança do primeiro ao sétimo dia foi testada por Wilcoxon pareado (com '
  'tamanho de efeito dz) e as associações com o humor por correlação de Spearman no nível atleta-dia.')

H('Resultados')
# tabela D1->D7
cap('Tabela. Sonolência (Epworth) e estresse (PSS): primeiro versus último dia do microciclo.')
hdr=['Variável','Dia 1 (M)','Dia 7 (M)','dz','p (D1→D7)']
tb=doc.add_table(rows=1,cols=5); tb.style='Table Grid'; tb.alignment=WD_TABLE_ALIGNMENT.CENTER
for j,hh in enumerate(hdr): setcell(tb.rows[0].cells[j],hh,bold=True)
for k,lab in [('Epworth','Sonolência (Epworth)'),('PSS','Estresse (PSS)')]:
    w=ad.pivot_table(index='ID',columns='dia',values=k); j=pd.concat([w[1],w[7]],axis=1).dropna()
    dz=(j[7]-j[1]).mean()/(j[7]-j[1]).std(ddof=1); p=stats.wilcoxon(j[1],j[7]).pvalue
    c=tb.add_row().cells
    for jj,v in enumerate([lab,num(w[1].mean()),num(w[7].mean()),('%+.2f'%dz).replace('.',','),pv(p)]):
        setcell(c[jj],v,align='left' if jj==0 else 'center')
cap('M: média. dz: tamanho de efeito intraindividual (Wilcoxon pareado).')

P('A sonolência diurna aumentou de forma significativa ao longo da semana (Epworth: 8,8 no primeiro dia para '
  '11,5 no sétimo; dz = +0,58; p = 0,019), com valor máximo no dia da última sessão de HIIT. O estresse '
  'percebido, em contraste, permaneceu estável (PSS: 22,7 para 21,6; dz = −0,19; p = 0,414), sendo inclusive '
  'mais baixo nos dias de jogo amistoso.')
P('No nível atleta-dia, a sonolência associou-se positivamente à fadiga (ρ = +0,34; p < 0,001) e à perturbação '
  'total do humor (ρ = +0,37; p < 0,001) e negativamente ao vigor (ρ = −0,23; p = 0,002). O estresse percebido '
  'apresentou associação fraca com a perturbação total do humor (ρ = +0,24; p = 0,002) e não se relacionou com a '
  'fadiga (ρ = −0,04). A Figura ilustra as duas trajetórias e as correlações.')
cap('Figura. (A) Sonolência ao longo da semana; (B) estresse ao longo da semana; (C) correlações com o humor. '
    'Faixas alaranjadas = dias de HIIT; azuis = jogos amistosos.')

H('Discussão')
P('A sonolência diurna comportou-se como um marcador de recuperação: acompanhou o eixo da fadiga e cresceu à '
  'medida que a carga se acumulou no microciclo, atingindo o pico no dia da sessão de HIIT mais exigente. Esse '
  'padrão é coerente com a maior necessidade de sono associada ao estresse fisiológico do treinamento e reforça, '
  'de forma independente, a interpretação de que o desgaste observado é sobretudo energético.')
P('O estresse percebido, por outro lado, manteve-se estável e desacoplado da fadiga, tendo sido menor justamente '
  'nos dias de jogo. Em conjunto, os dois instrumentos indicam que os atletas ficaram progressivamente mais '
  'sonolentos e fadigados, mas não mais estressados — um perfil compatível com sobrecarga funcional, e não com '
  'um quadro de distresse psicológico. A dissociação entre um eixo energético (fadiga, sonolência, vigor), que '
  'responde à quantidade de carga, e um eixo de estresse/afeto negativo, que permanece controlado, sugere que o '
  'monitoramento do bem-estar do atleta se beneficia de acompanhar as duas dimensões separadamente.')

doc.save('/home/user/mdlucca/Artigos/Sono_Estresse_Resultados_Discussao.docx')
print('[docx: Sono_Estresse_Resultados_Discussao.docx]')
