# -*- coding: utf-8 -*-
import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
R=json.load(open('brums_desc2.json')); STAT=json.load(open('brums_stats3.json')); S4=json.load(open('brums_stats4.json'))
MS=json.load(open('model_stats.json')); MV=json.load(open('manova.json')); PHJ=json.load(open('posthoc.json'))
PV=json.load(open('pv_stats.json')); LIM=json.load(open('tcar_limiar.json')); TCD=json.load(open('tcar_desc.json'))
CRS=json.load(open('cross.json')); PK=json.load(open('peaks.json'))
FG='/home/user/mdlucca/Artigos/figuras'
def c2(s): return str(s).replace('.',',')
doc=Document()
stl=doc.styles['Normal']; stl.font.name='Times New Roman'; stl.font.size=Pt(12)
stl.element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')
stl.paragraph_format.line_spacing=1.5; stl.paragraph_format.space_after=Pt(0)
sec=doc.sections[0]; sec.top_margin=Cm(3); sec.left_margin=Cm(3); sec.bottom_margin=Cm(2); sec.right_margin=Cm(2)
_TN=[0]; _FN=[0]
def P(t='',just=True,size=12,after=6,bold=False,ind=True):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(size); r.bold=bold
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(after)
    if just: p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if just and ind: p.paragraph_format.first_line_indent=Cm(1.25)
    return p
def RUN(pairs,after=6,ind=True):
    p=doc.add_paragraph(); p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(after)
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if ind: p.paragraph_format.first_line_indent=Cm(1.25)
    for txt,bold in pairs:
        r=p.add_run(txt); r.bold=bold; r.font.size=Pt(12)
    return p
def H(t,size=12,before=10):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(size)
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.5
def _bd(c):
    tcPr=c._tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for e_ in ['top','bottom','left','right']:
        el=OxmlElement(f'w:{e_}'); on=e_ in ('top','bottom')
        el.set(qn('w:val'),'single' if on else 'nil'); el.set(qn('w:sz'),'6'); el.set(qn('w:color'),'000000'); b.append(el)
    tcPr.append(b)
def table(cap,header,rows,fonte='Fonte: dados da pesquisa (2026).',fs=9,note=None):
    _TN[0]+=1
    p=doc.add_paragraph(); r=p.add_run('Tabela %d – %s'%(_TN[0],cap)); r.font.size=Pt(11)
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
    t=doc.add_table(rows=1,cols=len(header)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,htx in enumerate(header):
        cc=t.rows[0].cells[i]; cc.text=''; rr=cc.paragraphs[0].add_run(htx); rr.bold=True; rr.font.size=Pt(fs); rr.font.name='Times New Roman'
        cc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; _bd(cc)
    for row in rows:
        cs=t.add_row().cells
        for i,val in enumerate(row):
            cs[i].text=''; rr=cs[i].paragraphs[0].add_run(str(val)); rr.font.size=Pt(fs); rr.font.name='Times New Roman'
            cs[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if i>0 else WD_ALIGN_PARAGRAPH.LEFT; _bd(cs[i])
    if note:
        pn=doc.add_paragraph(); rn=pn.add_run(note); rn.font.size=Pt(8.5); rn.italic=True; pn.paragraph_format.space_after=Pt(0)
    pf=doc.add_paragraph(); rf=pf.add_run(fonte); rf.font.size=Pt(9); pf.paragraph_format.space_after=Pt(6)
    return _TN[0]
def figure(path,cap,w=15.0):
    _FN[0]+=1
    pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; pp.add_run().add_picture(path,width=Cm(w)); pp.paragraph_format.space_before=Pt(6)
    pc=doc.add_paragraph(); pc.alignment=WD_ALIGN_PARAGRAPH.CENTER; rc=pc.add_run('Figura %d – %s'%(_FN[0],cap)); rc.font.size=Pt(11)
    pf=doc.add_paragraph(); pf.alignment=WD_ALIGN_PARAGRAPH.CENTER; rf=pf.add_run('Fonte: elaboração dos autores (2026).'); rf.font.size=Pt(9); pf.paragraph_format.space_after=Pt(6)
    return _FN[0]
def pstr(p): return '< 0,001' if p<0.001 else c2('%.3f'%p)
mvv=lambda tab,k,f: next(x[f] for x in MV[tab]['rows'] if x['k']==k)
pr=R['prepos']; d17=R['d1d7']; prof=R['profiles']; sm=R['sample']; desc=R['desc']; PREV=MS['prev']
SH=STAT['shapiro']; IC=STAT['icc']; LP=LIM['LIM']['PVini']; FR=S4['friedman']; SENS=S4['sens']
ORD=[('Vigor','Vigor'),('Fadiga','Fadiga'),('TMD','PTH'),('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Confusao','Confusão')]

# ===== TÍTULO =====
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('PERFIL DE HUMOR DE ATLETAS DE HANDEBOL DE ELITE NA ÚLTIMA SEMANA DE PRÉ-TEMPORADA')
r.bold=True; r.font.size=Pt(13); p.paragraph_format.space_after=Pt(10)

# ===== RESUMO =====
H('RESUMO',before=2)
RUN([('Objetivo: ',True),('caracterizar e analisar o perfil de humor de atletas de handebol de elite ao longo da última '
 'semana de pré-temporada, comparando cada dimensão do BRUMS entre todos os dias e identificando os dias de maior e menor '
 'expressão de cada estado de humor, com a aptidão aeróbia intermitente como parâmetro fisiológico. ',False),
 ('Método: ',True),('%d atletas do sexo masculino responderam ao BRUMS-24 duas vezes ao dia por sete dias (%d '
 'observações). Empregaram-se estatística descritiva, Shapiro-Wilk, Wilcoxon com tamanho de efeito, Friedman com W de '
 'Kendall, pós-teste (Tukey), ICC, MANOVA em escores T e regressão do pico de velocidade do T-CAR, com limiar por índice '
 'de Youden. '%(sm['n'],sm['n_obs']),False),
 ('Resultados: ',True),('a deterioração concentrou-se no eixo energia–fadiga — o vigor caiu (d = %s) e a fadiga subiu '
 '(d = %s), confirmado por MANOVA (Wilks λ = %s; p = %s). O vigor foi máximo no Dia %d e a fadiga no Dia %d; o perfil '
 'iceberg caiu e o de fadiga subiu (χ² = %s; p = %s). Maior aptidão associou-se a menos fadiga (ρ = %s), com limiar de '
 '%s km/h. '%(
   c2('%+.2f'%mvv('d1d7','Vigor','d')),c2('%+.2f'%mvv('d1d7','Fadiga','d')),c2('%.3f'%MV['d1d7']['wilks']),pstr(MV['d1d7']['p_mv']),
   PK['Vigor']['max_day'],PK['Fadiga']['max_day'],c2('%.2f'%PREV['chi']),pstr(PREV['p']),
   c2('%+.2f'%PV['pv']['wk_FadFisica']['TCAR1']['rho']),c2('%.1f'%LP['thr'])),False),
 ('Conclusão: ',True),('o humor migrou da prontidão para a fadiga funcional; o vigor e a fadiga foram as dimensões mais '
 'sensíveis, e a aptidão intermitente ajudou a explicar a resposta, fundamentando o monitoramento individualizado.',False)],after=6)
P('Palavras-chave: humor; BRUMS; handebol; monitoramento do atleta; fadiga; aptidão intermitente.',size=11,after=8,ind=False)

# ===== 1 INTRODUÇÃO =====
H('1 INTRODUÇÃO')
H('1.1 Monitoramento do humor no esporte de rendimento',12,before=6)
P('O acompanhamento do estado psicológico dos atletas consolidou-se como parte essencial da gestão do treinamento no '
 'esporte de rendimento. Instrumentos de autorrelato do humor são práticos, econômicos e sensíveis às variações da carga '
 'de treino, com utilidade preditiva para o bem-estar e o desempenho esportivo (SAW; MAIN; GASTIN, 2016; LOCHBAUM et al., '
 '2021), razão pela qual documentos de consenso recomendam seu uso rotineiro para monitorar a fadiga e orientar decisões '
 'de treino e recuperação (KELLMANN et al., 2018).')
P('Entre esses instrumentos, a Escala de Humor de Brunel (BRUMS), versão abreviada e adaptada do Profile of Mood States '
 '(POMS), destaca-se pela rapidez de aplicação e pela solidez psicométrica, com sucessivas validações transculturais, '
 'incluindo a versão brasileira (BRAMS) (TERRY; LANE; FOGARTY, 2003; ROHLFS et al., 2008, 2023). Sua aplicação tem-se '
 'mostrado particularmente informativa em modalidades esportivas coletivas e no esporte de rendimento, contextos de '
 'elevada demanda física, emocional e interpessoal: a BRUMS vem sendo empregada para monitorar o humor de atletas de '
 'elite ao longo de períodos competitivos e sua relação com o sono, o desempenho, os resultados de partida e o risco de '
 'lesão (ANDRADE et al., 2016; BRANDT; BEVILACQUA; ANDRADE, 2017; ANDRADE et al., 2020; DE MIRANDA ROHLFS et al., 2025; '
 'DO NASCIMENTO et al., 2026). Além dos escores isolados, a literatura recente formalizou seis perfis prototípicos de '
 'humor — iceberg, Everest invertido, iceberg invertido, submerso, barbatana de tubarão e superfície —, replicados em '
 'grandes amostras internacionais e utilizados no rastreamento da prontidão de atletas de elite (MORGAN, 1985; '
 'PARSONS-SMITH; TERRY; MACHIN, 2017; HAN; PARSONS-SMITH; TERRY, 2020; TERRY et al., 2021).')
H('1.2 Handebol: modalidade coletiva intermitente e de alta intensidade',12,before=6)
P('O handebol de quadra é uma modalidade coletiva de invasão, de caráter marcadamente intermitente e de alta intensidade. '
 'Ao longo da partida, ações máximas e explosivas — sprints curtos, saltos, arremessos, bloqueios, mudanças de direção e '
 'contatos físicos — alternam-se, de forma imprevisível, com períodos de recuperação incompleta, exigindo '
 'simultaneamente potência anaeróbia, capacidade aeróbia intermitente e elevada tolerância à fadiga (KARCHER; BUCHHEIT, '
 '2014; MICHALSIK; AAGAARD, 2015; WAGNER et al., 2014). Esse padrão de esforço eleva a carga interna nos microciclos de '
 'acúmulo e repercute no estado afetivo, sobretudo na última semana de pré-temporada, quando a carga que antecede a '
 'competição se concentra. A tolerância a essa carga depende da aptidão aeróbia intermitente, cujo marcador de campo — o '
 'pico de velocidade do Teste de Carminatti (T-CAR) — associa-se ao desempenho físico em modalidades intermitentes '
 '(FERNANDES-DA-SILVA et al., 2016).')
H('1.3 Objetivos e hipóteses',12,before=6)
P('Diante da ampla adoção da BRUMS, mas da escassez de descrições detalhadas do comportamento do humor em um microciclo '
 'pré-competitivo de handebol de elite, o objetivo geral foi caracterizar e analisar o perfil de humor desses atletas ao '
 'longo da última semana de pré-temporada, descrevendo o comportamento de cada dimensão do BRUMS, comparando todos os '
 'dias entre si, identificando os dias de maior e menor expressão de cada estado de humor e relacionando a resposta à '
 'aptidão aeróbia intermitente. Especificamente, buscou-se: (i) caracterizar a amostra; (ii) verificar a normalidade; '
 '(iii) descrever a resposta aguda pré → pós com tamanho de efeito; (iv) comparar as dimensões entre todos os dias; (v) '
 'confirmar as diferenças por análise multivariada (escores T); (vi) descrever o comportamento individual de cada '
 'variável e suas relações; (vii) identificar os dias de maior fadiga, vigor, tensão, raiva e depressão; (viii) '
 'classificar a evolução dos perfis de humor; e (ix) analisar o T-CAR e o limiar de pico de velocidade. Hipotetizou-se '
 'que a deterioração se concentraria no eixo energia–fadiga, com migração do perfil iceberg para o de fadiga, e que a '
 'aptidão intermitente modularia a resposta.')

# ===== 2 MATERIAIS E MÉTODOS =====
H('2 MATERIAIS E MÉTODOS')
H('2.1 Delineamento e amostra',12,before=6)
P('Estudo descritivo-comparativo, observacional e de medidas repetidas, conduzido em condições ecológicas de treinamento '
 'durante o microciclo pré-competitivo (21 a 27 de abril de 2024), com %d atletas de handebol do sexo masculino de nível '
 'competitivo, conforme os princípios éticos da Declaração de Helsinque e com consentimento informado.'%sm['n'])
H('2.2 Instrumentos',12,before=6)
P('O humor foi avaliado pela BRUMS-24, com 24 itens em escala de 0 (“nada”) a 4 (“extremamente”), agrupados em seis '
 'subescalas de 0 a 16 pontos (tensão, depressão, raiva, vigor, fadiga e confusão); a Perturbação Total do Humor (PTH) '
 'resume o perfil (soma das negativas menos o vigor). Para a classificação de perfis e a análise multivariada, os escores '
 'foram convertidos em escores T (M = 50; DP = 10). A aptidão aeróbia intermitente foi avaliada pelo Teste de Carminatti '
 '(T-CAR), teste de campo progressivo e intermitente (repetições de 12 s de corrida em vaivém intercaladas por 6 s de '
 'recuperação, até a exaustão), tendo o pico de velocidade (PV) como desfecho (FERNANDES-DA-SILVA et al., 2016). O T-CAR '
 'foi aplicado em 15 de abril de 2024, quatro dias de treino antes do início do microciclo, de modo a servir como '
 'parâmetro fisiológico de linha de base das análises.')
H('2.3 Procedimentos',12,before=6)
P('O BRUMS foi autoaplicado por formulário eletrônico com registro de data e hora, duas vezes por dia de treino — a '
 'primeira resposta tomada como pré e a última como pós —, ao longo de sete dias, totalizando %d observações válidas.'%sm['n_obs'])
H('2.4 Análise estatística',12,before=6)
P('A análise seguiu uma sequência do descritivo ao inferencial. Inicialmente, empregou-se estatística descritiva (média, '
 'desvio-padrão, mediana e amplitude) para caracterizar as variáveis, e o teste de Shapiro-Wilk para verificar a '
 'normalidade — etapa que orienta a escolha entre testes paramétricos e não paramétricos. Como as distribuições violaram '
 'a normalidade, com forte concentração de escores baixos nas dimensões negativas (efeito de piso), os dados não foram '
 'transformados (NEVILL; LANE, 2007) e privilegiaram-se testes não paramétricos.')
P('Para descrever a resposta aguda ao treino, os momentos pré e pós foram comparados pelo teste de Wilcoxon para amostras '
 'pareadas, acompanhado do tamanho de efeito (d de Cohen), que expressa a magnitude da diferença independentemente do '
 'tamanho da amostra (trivial < 0,2; pequeno < 0,5; médio < 0,8; grande ≥ 0,8); a mesma lógica aplicou-se à comparação '
 'entre o primeiro e o último dia. Para verificar se cada dimensão variou ao longo dos sete dias, utilizou-se o teste de '
 'Friedman — equivalente não paramétrico da ANOVA de medidas repetidas —, com o W de Kendall como tamanho de efeito '
 '(0,1 pequeno; 0,3 moderado; 0,5 grande); quando significativo, aplicou-se o pós-teste das médias marginais de um modelo '
 'misto (correção de Tukey), que identifica entre quais dias, especificamente, houve diferença. A consistência das '
 'medidas repetidas ao longo da semana foi estimada pelo coeficiente de correlação intraclasse (ICC), interpretado como '
 'pobre (< 0,50), moderado (0,50–0,75), bom (0,75–0,90) ou excelente (> 0,90).')
P('Como confirmação robusta, as comparações Dia 1 → Dia 7 e pré → pós foram reanalisadas por análise multivariada de '
 'variância (MANOVA) de medidas repetidas sobre os escores T, que testa as seis dimensões em conjunto e controla o erro '
 'de múltiplas comparações (lambda de Wilks, F e eta-quadrado parcial — η²ₚ; 0,01 pequeno; 0,06 médio; 0,14 grande). A '
 'associação entre as dimensões foi quantificada pela correlação de Spearman (ρ). Por fim, a relação entre o pico de '
 'velocidade do T-CAR e a fadiga foi analisada por regressão, e um limiar de pico de velocidade foi estabelecido pelo '
 'índice de Youden — que maximiza a soma de sensibilidade e especificidade —, com a qualidade de discriminação avaliada '
 'pela área sob a curva ROC. Adotou-se nível de significância de 5% (p < 0,05), com as análises conduzidas em ambiente '
 'Python (bibliotecas pandas, SciPy e statsmodels).')

# ===== 3 RESULTADOS =====
H('3 RESULTADOS')
H('3.1 Caracterização da amostra',12,before=6)
tp=table('Distribuição da amostra por posição de jogo (n = %d).'%sm['n'],['Posição','n','%'],
    [[k,v,c2('%.1f'%(100*v/sm['n']))] for k,v in sm['pos'].items()]+[['Total',sm['n'],'100,0']],fs=9)
def srow(lab,k):
    v=sm[k]; return [lab,c2('%.1f'%v['mean']),c2('%.1f'%v['sd']),'%s–%s'%(c2('%.1f'%v['mn']),c2('%.1f'%v['mx']))]
ta=table('Caracterização sociodemográfica e antropométrica (n = %d).'%sm['n'],['Variável','Média','DP','Mín–Máx'],
    [srow('Idade (anos)','idade'),srow('Estatura (cm)','estatura'),srow('Massa corporal (kg)','massa'),srow('Percentual de gordura (%)','pG'),srow('Experiência (anos)','exp')])
P('A amostra é descrita nas Tabelas %d e %d.'%(tp,ta))
H('3.2 Normalidade das distribuições',12,before=6)
tn=table('Teste de normalidade (Shapiro-Wilk) das dimensões do BRUMS.',['Dimensão','Estatística (W)','p'],
    [[lab,c2('%.3f'%SH[k]['W']),'< 0,001' if SH[k]['p']<0.001 else c2('%.3f'%SH[k]['p'])] for k,lab in [('Vigor','Vigor'),('Fadiga','Fadiga'),('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Confusao','Confusão')]],fs=9)
P('As seis dimensões não seguem distribuição normal (p < 0,001; Tabela %d), justificando os testes não paramétricos.'%tn)
H('3.3 Estatística descritiva das dimensões',12,before=6)
def drow(k,lab):
    v=desc[k]; return [lab,c2('%.2f'%v['mean']),c2('%.1f'%v['md']),c2('%.2f'%v['sd']),'%s–%s'%(c2('%.0f'%v['mn']),c2('%.0f'%v['mx']))]
td=table('Estatística descritiva das dimensões do BRUMS e da PTH (%d observações).'%sm['n_obs'],
    ['Dimensão','Média','Mediana','DP','Mín–Máx'],[drow(k,l) for k,l in ORD],fs=9)
P('A descritiva geral consta na Tabela %d: o vigor e a fadiga concentram as maiores médias e variabilidade.'%td)
H('3.4 Diferenças entre pré e pós-treino (com tamanho de efeito)',12,before=6)
def ppr(k,lab):
    v=pr[k]; return [lab,c2('%.2f'%v['pre']),c2('%.2f'%v['pos']),c2('%+.0f'%v['pct'])+'%',pstr(v['p']),c2('%+.2f'%v['dz']),v['mag']]
tpp=table('Diferenças entre pré e pós-treino (agregado da semana; Wilcoxon e d de Cohen).',
    ['Dimensão','Pré (M)','Pós (M)','Variação (%)','p','d de Cohen','Magnitude'],[ppr(k,l) for k,l in ORD],fs=9)
P('A resposta aguda foi consistente (Tabela %d): o vigor caiu (%s%%; d = %s) e a fadiga e a PTH subiram do pré para o '
 'pós, com efeito médio.'%(tpp,c2('%.0f'%abs(pr['Vigor']['pct'])),c2('%.2f'%pr['Vigor']['dz'])))
H('3.5 Comportamento diário e comparação entre todos os dias',12,before=6)
f_traj=figure(f'{FG}/xb2_traj.png','Comportamento das dimensões do humor ao longo da semana (médias diárias; áreas sombreadas = IC95%).')
f_box=figure(f'{FG}/xb3_box.png','Diagramas de caixa (box plot) das seis dimensões do BRUMS por dia da semana.')
def ddrow(k,lab):
    dm=CRS['dec'][k]['dm']; f=FR[k]; return [lab]+[c2('%.1f'%dm[str(d)]) for d in range(1,8)]+[c2('%.1f'%f['chi']),pstr(f['p']),c2('%.2f'%f['W'])]
tdd=table('Médias diárias das dimensões do BRUMS e teste de Friedman com W de Kendall (variação entre os sete dias).',
    ['Dimensão','D1','D2','D3','D4','D5','D6','D7','χ²','p','W de Kendall'],[ddrow(k,l) for k,l in ORD],
    note='W de Kendall = tamanho de efeito do teste de Friedman (0,1 pequeno; 0,3 moderado; 0,5 grande).',fs=8)
P('As médias diárias, o teste de Friedman e o W de Kendall constam na Tabela %d (Figuras %d e %d): houve diferença '
 'significativa entre os dias para o vigor (p = %s; W = %s), a fadiga (p = %s; W = %s), a tensão (p = %s) e a confusão '
 '(p = %s), sempre com magnitude pequena a moderada, coerente com um microciclo de acúmulo dentro da faixa funcional.'%(
   tdd,f_traj,f_box,pstr(FR['Vigor']['p']),c2('%.2f'%FR['Vigor']['W']),pstr(FR['Fadiga']['p']),c2('%.2f'%FR['Fadiga']['W']),pstr(FR['Tensao']['p']),pstr(FR['Confusao']['p'])))
def icrow(k,lab):
    i=IC[k]; return [lab,c2('%.2f'%i['icc1']),c2('%.2f'%i['icck']),i['cls']]
tic=table('Consistência das medidas repetidas ao longo da semana (coeficiente de correlação intraclasse, ICC).',
    ['Dimensão','ICC(2,1)','ICC(2,k)','Consistência'],
    [icrow(k,l) for k,l in [('Vigor','Vigor'),('Fadiga','Fadiga'),('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Confusao','Confusão')]],fs=9)
P('A consistência das medidas repetidas (Tabela %d) foi moderada a boa, sendo mais baixa para a raiva e a confusão — '
 'dimensões mais reativas dia a dia.'%tic)
def emmc(v,d): return c2('%.2f'%PHJ[v]['emm'][str(d)])
def sig1(v,d): return '' if d==1 or PHJ[v]['pairs']['1_%d'%d]['ptukey']>=0.05 else '*'
te=table('Pós-teste (médias marginais do modelo misto) por dia, com comparação de cada dia ao Dia 1.',
    ['Dia','Vigor','Fadiga','Fadiga física'],
    [['Dia %d'%d,emmc('Vigor',d)+sig1('Vigor',d),emmc('Fadiga',d)+sig1('Fadiga',d),emmc('FadFisica',d)+sig1('FadFisica',d)] for d in range(1,8)],
    note='* diferença significativa em relação ao Dia 1 (Tukey, p < 0,05).',fs=9)
f_ph=figure(f'{FG}/ph_emm.png','Trajetória diária (médias marginais) com comparação de todos os dias ao Dia 1 (* p < 0,05).')
def npairs(v): return sum(1 for kk,pp in PHJ[v]['pairs'].items() if pp['ptukey']<0.05)
P('Comparando todos os dias entre si (pós-teste de Tukey; Tabela %d; Figura %d), o vigor diferiu significativamente em '
 '%d dos 21 pares de dias e a fadiga em %d — sempre no sentido de piora em relação aos primeiros dias —, confirmando a '
 'deterioração progressiva do eixo energia–fadiga.'%(te,f_ph,npairs('Vigor'),npairs('Fadiga')))
H('3.6 Confirmação por análise multivariada (MANOVA em escores T)',12,before=6)
f_prof=figure(f'{FG}/xb5_profile_d1d7.png','Perfil de humor em escores T (M = 50; DP = 10) no Dia 1 e no Dia 7.',w=13.5)
def mvrow(x):
    return [x['lab'],c2('%.1f'%x['m1']),c2('%.1f'%x['s1']),c2('%.1f'%x['m2']),c2('%.1f'%x['s2']),c2('%.2f'%x['F']),pstr(x['p'])+('*' if x['p']<0.05 else ''),c2('%+.2f'%x['d']),c2('%.3f'%x['eta'])]
mv=MV['d1d7']
tmv=table('Comparação Dia 1 → Dia 7 das seis dimensões em escores T (MANOVA de medidas repetidas; n = %d).'%mv['n'],
    ['Dimensão','D1 M','D1 DP','D7 M','D7 DP','F','p','d','η²ₚ'],[mvrow(x) for x in mv['rows']],
    note='Wilks λ = %s; F(%d,%d) = %s; p = %s; η²ₚ = %s. * p < 0,05.'%(c2('%.3f'%mv['wilks']),mv['df1'],mv['df2'],c2('%.2f'%mv['Fmv']),pstr(mv['p_mv']),c2('%.2f'%mv['eta_mv'])),fs=8.5)
P('A análise multivariada confirmou a diferença entre o Dia 1 e o Dia 7 (Wilks λ = %s; F(%d,%d) = %s; p = %s; η²ₚ = %s), '
 'com efeito concentrado no vigor (d = %s) e na fadiga (d = %s) (Tabela %d; Figura %d). A resposta aguda pré → pós também '
 'foi multivariadamente significativa (Wilks λ = %s; p = %s), reforçando o achado do eixo energia–fadiga.'%(
   c2('%.3f'%mv['wilks']),mv['df1'],mv['df2'],c2('%.2f'%mv['Fmv']),pstr(mv['p_mv']),c2('%.2f'%mv['eta_mv']),
   c2('%+.2f'%mvv('d1d7','Vigor','d')),c2('%+.2f'%mvv('d1d7','Fadiga','d')),tmv,f_prof,c2('%.3f'%MV['prepos']['wilks']),pstr(MV['prepos']['p_mv'])))
H('3.7 Dias de maior expressão de cada estado de humor',12,before=6)
def pkrow(k,lab):
    p=PK[k]; return [lab,'Dia %d'%p['max_day'],c2('%.2f'%p['max_val']),'Dia %d'%p['min_day'],c2('%.2f'%p['min_val'])]
tpk=table('Dia de maior e de menor expressão de cada dimensão do humor (médias diárias).',
    ['Dimensão','Dia de maior valor','Valor','Dia de menor valor','Valor'],[pkrow(k,l) for k,l in ORD],fs=9)
P('A Tabela %d sintetiza os dias de pico: o vigor foi máximo no Dia %d e mínimo no Dia %d; a fadiga foi máxima no Dia %d; '
 'a tensão, máxima no Dia %d; a raiva, máxima no Dia %d; e a depressão, máxima no Dia %d. Em síntese, o início da semana '
 'reúne maior prontidão (vigor, tensão e confusão mais altos) e o final concentra a fadiga e a raiva.'%(
   tpk,PK['Vigor']['max_day'],PK['Vigor']['min_day'],PK['Fadiga']['max_day'],PK['Tensao']['max_day'],PK['Raiva']['max_day'],PK['Depressao']['max_day']))
H('3.8 Comportamento individual de cada variável e suas relações',12,before=6)
P('As Figuras %d a %d apresentam, para cada dimensão, o comportamento ao longo da semana com a média diária, a banda de '
 'confiança de 95%% (área sombreada), os diagramas de caixa por dia e o efeito do microciclo, permitindo visualizar '
 'individualmente como cada estado de humor evolui.'%(_FN[0]+1,_FN[0]+6))
for k,fn,lab in [('Vigor','xb4_v_Vigor.png','Vigor'),('Fadiga','xb4_v_Fadiga.png','Fadiga'),('Tensao','xb4_v_Tensao.png','Tensão'),('Depressao','xb4_v_Depressao.png','Depressão'),('Raiva','xb4_v_Raiva.png','Raiva'),('Confusao','xb4_v_Confusao.png','Confusão')]:
    figure(f'{FG}/{fn}','%s ao longo da semana: média diária (banda = IC95%%), diagramas de caixa por dia e efeito Dia 1 → Dia 7.'%lab,w=12.5)
PA=STAT['pairs']; FP=STAT['focusp']; sig=[x for x in PA if x['p']<0.05]
tcorr=table('Correlação de Spearman entre as dimensões do BRUMS com associação significativa (n = %d atletas).'%sm['n'],
    ['Par de dimensões','ρ','p'],[['%s × %s'%(x['a'],x['b']),c2('%+.2f'%x['rho']),pstr(x['p'])] for x in sig],fs=9)
P('Quanto às relações entre as dimensões (Tabela %d), as dimensões negativas associam-se entre si e a fadiga relaciona-se '
 'com a depressão e a raiva; o vigor mantém-se relativamente independente. Esse acoplamento depende da carga: no dia de '
 'maior vigor a depressão é independente da fadiga (ρ = %s), mas no dia de maior fadiga torna-se forte (ρ = %s).'%(
   tcorr,c2('%+.2f'%FP['D1']['Depressao']['rho_fad']),c2('%+.2f'%FP['D7']['Depressao']['rho_fad'])))
H('3.9 Classificação dos perfis de humor',12,before=6)
f_prev=figure(f'{FG}/xb5_prev.png','Distribuição (%) dos seis perfis de humor no Dia 1 e no Dia 7.',w=13.5)
PROFR=[('Iceberg','Iceberg'),('Everest invertido','Everest invertido'),('Iceberg invertido','Iceberg invertido'),('Submerso','Submerso'),('Barbatana tubarão','Barbatana de tubarão'),('Superfície','Superfície')]
def prow(p,lab):
    d1=PREV['D1'][p]; d7=PREV['D7'][p]; return [lab,'%d (%s)'%(d1,c2('%.1f'%(100*d1/PREV['n_d1']))),'%d (%s)'%(d7,c2('%.1f'%(100*d7/PREV['n_d7'])))]
tprev=table('Distribuição dos perfis de humor no primeiro e no último dia da semana: n (%).',
    ['Perfil','Dia 1 n (%)','Dia 7 n (%)'],[prow(p,lab) for p,lab in PROFR],fs=9)
P('O perfil migrou do iceberg (%s%% no Dia 1) para a barbatana de tubarão (%s%% no Dia 7), com o iceberg caindo para '
 '%s%% (Figura %d; Tabela %d).'%(
   c2('%.0f'%(100*PREV['D1']['Iceberg']/PREV['n_d1'])),c2('%.0f'%(100*PREV['D7']['Barbatana tubarão']/PREV['n_d7'])),c2('%.0f'%(100*PREV['D7']['Iceberg']/PREV['n_d7'])),f_prev,tprev))
H('3.10 Aptidão intermitente (T-CAR) e limiar de pico de velocidade',12,before=6)
def tcrow(o):
    lab='Pico de velocidade — T-CAR (km/h)' if o['k']=='PVini' else o['lab']
    return [lab,o['n'],c2('%.1f'%o['m']),c2('%.1f'%o['sd']),'%s–%s'%(c2('%.1f'%o['mn']),c2('%.1f'%o['mx']))]
TCD2=[o for o in TCD if o['k'] not in ('PV','dPV')]
ttc=table('Estatística descritiva do desempenho no Teste de Carminatti (T-CAR).',
    ['Parâmetro','n','Média','DP','Mín–Máx'],[tcrow(o) for o in TCD2],
    note='PV = pico de velocidade; FC = frequência cardíaca; TRIMP = training impulse.',fs=8.5)
f_pv=figure(f'{FG}/pv7_scatter.png','Dispersão e reta de regressão (com banda de IC95%) entre o pico de velocidade do T-CAR e as médias semanais de vigor, fadiga, PTH e fadiga física.',w=15.0)
P('O desempenho no T-CAR consta na Tabela %d. A regressão mostrou que atletas com maior pico de velocidade reportaram '
 'menos fadiga física (ρ = %s; p = %s) e mais vigor (ρ = %s; p = %s) ao longo da semana (Figura %d). A partir de um '
 'modelo que estima a probabilidade de um dia de fadiga elevada em função do pico de velocidade, identificou-se um limiar '
 'de aproximadamente %s km/h (área sob a curva = %s; sensibilidade = %s; especificidade = %s): abaixo desse valor, os '
 'atletas apresentaram maior probabilidade de dias de fadiga elevada, o que fornece uma referência objetiva para '
 'individualizar a carga.'%(
   ttc,c2('%+.2f'%PV['pv']['wk_FadFisica']['TCAR1']['rho']),pstr(PV['pv']['wk_FadFisica']['TCAR1']['rho_p']),
   c2('%+.2f'%PV['pv']['wk_Vigor']['TCAR1']['rho']),pstr(PV['pv']['wk_Vigor']['TCAR1']['rho_p']),f_pv,
   c2('%.1f'%LP['thr']),c2('%.2f'%LP['auc']),c2('%.2f'%LP['sens']),c2('%.2f'%LP['spec'])))

# ===== 4 DISCUSSÃO =====
H('4 DISCUSSÃO')
P('A resposta do humor concentrou-se no eixo energia–fadiga, confirmada tanto pela comparação entre todos os dias '
 '(pós-teste) quanto pela análise multivariada em escores T: o vigor caiu com efeito grande e a fadiga subiu com efeito '
 'médio, enquanto as dimensões negativas de valência não fadiga permaneceram estáveis. Os dias de pico esclarecem a '
 'dinâmica — maior prontidão no início (vigor máximo no Dia %d) e maior fadiga e raiva no final (Dia %d) —, padrão '
 'coerente com o acúmulo de carga da pré-temporada e com a literatura de monitoramento (SAW; MAIN; GASTIN, 2016; '
 'KELLMANN et al., 2018).'%(PK['Vigor']['max_day'],PK['Fadiga']['max_day']))
P('O caráter do handebol ajuda a contextualizar a magnitude dessa resposta. Trata-se de uma modalidade coletiva de '
 'invasão, marcadamente intermitente, na qual ações de alta intensidade — sprints, saltos, arremessos, bloqueios, '
 'mudanças de direção e contatos — alternam-se com períodos de recuperação incompleta, exigindo simultaneamente potência '
 'anaeróbia e capacidade aeróbia intermitente para sustentar o esforço repetido e retardar a instalação da fadiga '
 '(KARCHER; BUCHHEIT, 2014; MICHALSIK; AAGAARD, 2015). Evidências atuais confirmam que o treino e o jogo de handebol '
 'constituem exercício de alta intensidade, com elevadas demandas aeróbias e anaeróbias (PEREIRA et al., 2024), e que os '
 'esportes de invasão se caracterizam por atividade intermitente de alta intensidade intercalada por recuperação de '
 'menor intensidade (VACCARO-BENET et al., 2024). Nesse contexto, o acúmulo de carga da última semana de pré-temporada '
 'explica tanto a deterioração do eixo energia–fadiga quanto a migração do perfil de humor observadas, reforçando por que '
 'o vigor e a fadiga foram as dimensões mais sensíveis do painel.')
P('Os achados também reforçam o valor da BRUMS como ferramenta de monitoramento no esporte coletivo de rendimento. Em '
 'modalidades de equipe, o humor responde de forma sensível a fatores como carga, sono e resultados competitivos: em '
 'atletas brasileiros de elite, a BRUMS já discriminou respostas ligadas ao sono e ao desempenho (ANDRADE et al., 2016; '
 'BRANDT; BEVILACQUA; ANDRADE, 2017) e, mais recentemente, respostas psicológicas agudas a resultados de partida em '
 'atletas jovens de modalidades coletivas (DO NASCIMENTO et al., 2026), além de associar-se à ocorrência de lesão em '
 'coortes de alto rendimento (DE MIRANDA ROHLFS et al., 2025). A rapidez de aplicação e a leitura por perfis tornam a '
 'BRUMS especialmente adequada ao acompanhamento longitudinal de equipes, permitindo à comissão técnica identificar '
 'precocemente os atletas em deterioração e individualizar as estratégias de recuperação.')
P('A migração do perfil iceberg para o de fadiga reproduz o “derretimento do iceberg” descrito na literatura (MORGAN, '
 '1985; HAN; PARSONS-SMITH; TERRY, 2020), e a relação entre a aptidão aeróbia intermitente e a fadiga — sintetizada no '
 'limiar de pico de velocidade — reforça a utilidade de integrar um parâmetro fisiológico simples ao monitoramento do '
 'humor, individualizando a interpretação da resposta à carga (FERNANDES-DA-SILVA et al., 2016). Como limitações, '
 'destacam-se o tamanho amostral e o desenho observacional de fase única, que não permite inferência causal.')

# ===== 5 CONCLUSÕES =====
H('5 CONCLUSÕES')
P('Na última semana de pré-temporada, o humor dos handebolistas migrou da prontidão para a fadiga, com o vigor em queda '
 'e a fadiga em ascensão, do início ao fim da semana e dentro de cada treino, confirmado por comparações entre todos os '
 'dias e por análise multivariada. Os dias críticos foram identificados — maior vigor no início e maior fadiga no final '
 '—, e a aptidão intermitente mostrou-se associada à fadiga, com um limiar de pico de velocidade útil para individualizar '
 'a carga. Recomenda-se o acompanhamento contínuo do humor, com atenção especial ao eixo energia–fadiga.')

# ===== REFERÊNCIAS =====
H('REFERÊNCIAS')
refs=[
 'ANDRADE, A. et al. Sleep quality, mood and performance: a study of elite Brazilian volleyball athletes. Journal of Sports Science and Medicine, v. 15, n. 4, p. 601–605, 2016.',
 'ANDRADE, A. et al. Effect of practice exergames on the mood states and self-esteem of elementary school boys and girls during physical education classes: a cluster-randomized controlled trial. PLoS ONE, v. 15, n. 6, e0232392, 2020. DOI: 10.1371/journal.pone.0232392.',
 'BRANDT, R.; BEVILACQUA, G. G.; ANDRADE, A. Perceived sleep quality, mood states, and their relationship with performance among Brazilian elite athletes during a competitive period. Journal of Strength and Conditioning Research, v. 31, n. 4, p. 1033–1039, 2017.',
 'DE MIRANDA ROHLFS, I. C. P. et al. Mood states, injury status, and countermovement jump performance in Brazilian high-level sports. Sports, v. 13, n. 9, 303, 2025. DOI: 10.3390/sports13090303.',
 'DO NASCIMENTO, M. H. et al. Acute psychological responses to official match outcomes in male youth volleyball: an observational repeated-measures study within a single national-level team. Frontiers in Psychology, v. 17, 1826372, 2026. DOI: 10.3389/fpsyg.2026.1826372.',
 'FERNANDES-DA-SILVA, J. et al. The peak velocity derived from the Carminatti Test is related to physical match performance in young soccer players. Journal of Sports Sciences, v. 34, n. 24, p. 2238–2245, 2016. DOI: 10.1080/02640414.2015.1093646.',
 'HAN, C.; PARSONS-SMITH, R. L.; TERRY, P. C. Mood profiling in Singapore: cross-cultural validation and potential applications of mood profile clusters. Frontiers in Psychology, v. 11, 665, 2020. DOI: 10.3389/fpsyg.2020.00665.',
 'KARCHER, C.; BUCHHEIT, M. On-court demands of elite handball, with special reference to playing positions. Sports Medicine, v. 44, n. 6, p. 797–814, 2014. DOI: 10.1007/s40279-014-0164-z.',
 'KELLMANN, M. et al. Recovery and performance in sport: consensus statement. International Journal of Sports Physiology and Performance, v. 13, n. 2, p. 240–245, 2018. DOI: 10.1123/ijspp.2017-0759.',
 'LOCHBAUM, M. et al. The Profile of Mood States and athletic performance: a meta-analysis of published studies. European Journal of Investigation in Health, Psychology and Education, v. 11, n. 1, p. 50–70, 2021. DOI: 10.3390/ejihpe11010005.',
 'MICHALSIK, L. B.; AAGAARD, P. Physical demands in elite team handball: comparisons between male and female players. Journal of Sports Medicine and Physical Fitness, v. 55, n. 9, p. 878–891, 2015.',
 'MORGAN, W. P. Selected psychological factors limiting performance: a mental health model. In: CLARKE, D. H.; ECKERT, H. M. (Ed.). Limits of human performance. Champaign: Human Kinetics, 1985. p. 70–80.',
 'NEVILL, A. M.; LANE, A. M. Why self-report “Likert” scale data should not be log-transformed. Journal of Sports Sciences, v. 25, n. 1, p. 1–2, 2007. DOI: 10.1080/02640410601111183.',
 'PARSONS-SMITH, R. L.; TERRY, P. C.; MACHIN, M. A. Identification and description of novel mood profile clusters. Frontiers in Psychology, v. 8, 1958, 2017. DOI: 10.3389/fpsyg.2017.01958.',
 'PEREIRA, R. et al. Exercise intensity and reliability during recreational team handball training for 50–77-year-old unexperienced women. Biology of Sport, v. 41, n. 4, p. 253–261, 2024. DOI: 10.5114/biolsport.2024.132995.',
 'ROHLFS, I. C. P. M. et al. A Escala de Humor de Brunel (Brums): instrumento para detecção precoce da síndrome do excesso de treinamento. Revista Brasileira de Medicina do Esporte, v. 14, n. 3, p. 176–181, 2008.',
 'ROHLFS, I. C. P. M. et al. Psychometric characteristics of the Brazil Mood Scale among youth and elite athletes using two response time frames. Sports, v. 11, n. 12, 244, 2023. DOI: 10.3390/sports11120244.',
 'SAW, A. E.; MAIN, L. C.; GASTIN, P. B. Monitoring the athlete training response: subjective self-reported measures trump commonly used objective measures: a systematic review. British Journal of Sports Medicine, v. 50, n. 5, p. 281–291, 2016. DOI: 10.1136/bjsports-2015-094758.',
 'TERRY, P. C.; LANE, A. M.; FOGARTY, G. J. Construct validity of the Profile of Mood States — Adolescents for use with adults. Psychology of Sport and Exercise, v. 4, n. 2, p. 125–139, 2003. DOI: 10.1016/S1469-0292(02)00035-8.',
 'TERRY, P. C. et al. Mood profiling for sustainable mental health among athletes. Sustainability, v. 13, n. 11, 6116, 2021. DOI: 10.3390/su13116116.',
 'VACCARO-BENET, P. et al. Internal and external load profile during beach invasion sports match-play by electronic performance and tracking systems: a systematic review. Sensors, v. 24, n. 12, 3738, 2024. DOI: 10.3390/s24123738.',
 'WAGNER, H. et al. Individual and team performance in team-handball: a review. Journal of Sports Science and Medicine, v. 13, n. 4, p. 808–816, 2014.']
for rf in refs:
    p=doc.add_paragraph(); r=p.add_run(rf); r.font.size=Pt(11); p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(6); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

OUTP='/home/user/mdlucca/Artigos/Relatorio_Avancado_Perfil_Humor.docx'
doc.save(OUTP); print('SAVED',OUTP,'| Tabelas',_TN[0],'Figuras',_FN[0])
