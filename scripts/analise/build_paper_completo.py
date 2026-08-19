# -*- coding: utf-8 -*-
import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
R=json.load(open('brums_desc2.json')); MS=json.load(open('model_stats.json'))
PSY=json.load(open('psychometric.json')); PK=json.load(open('peaks.json'))
PCA=json.load(open('pca.json')); LOG=json.load(open('logistica.json')); MV=json.load(open('manova.json'))
STAT=json.load(open('brums_stats3.json')); S4=json.load(open('brums_stats4.json')); CRS=json.load(open('cross.json'))
MDC=json.load(open('mdc.json')); PHJ=json.load(open('posthoc.json')); SMO=json.load(open('smooth_deriv.json'))
POLY=json.load(open('poly_fit.json')); CRX=json.load(open('cross_pts.json')); ADV=json.load(open('adv.json'))
FG='/home/user/mdlucca/Artigos/figuras'
def c2(s): return str(s).replace('.',',')
doc=Document()
stl=doc.styles['Normal']; stl.font.name='Times New Roman'; stl.font.size=Pt(12)
stl.element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')
stl.paragraph_format.line_spacing=1.5; stl.paragraph_format.space_after=Pt(0)
sec=doc.sections[0]; sec.top_margin=Cm(3); sec.left_margin=Cm(3); sec.bottom_margin=Cm(2); sec.right_margin=Cm(2)
_TN=[0]; _FN=[0]; _SN=[0]; SUPP=[]
def P(t='',just=True,size=12,after=0,bold=False,ind=True):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(size); r.bold=bold
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(after)
    if just: p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if just and ind: p.paragraph_format.first_line_indent=Cm(1.25)
    return p
def RUN(pairs,after=0,ind=True):
    p=doc.add_paragraph(); p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(after)
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if ind: p.paragraph_format.first_line_indent=Cm(1.25)
    for txt,bold in pairs:
        r=p.add_run(txt); r.bold=bold; r.font.size=Pt(12)
    return p
def H(t,size=12,before=10):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(size)
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.5
def _bd(c):
    tcPr=c._tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for e_ in ['top','bottom','left','right']:
        el=OxmlElement(f'w:{e_}'); on=e_ in ('top','bottom')
        el.set(qn('w:val'),'single' if on else 'nil'); el.set(qn('w:sz'),'6'); el.set(qn('w:color'),'000000'); b.append(el)
    tcPr.append(b)
def table(cap,header,rows,fonte='Fonte: dados da pesquisa (2026).',fs=9,note=None):
    _TN[0]+=1
    p=doc.add_paragraph(); r=p.add_run('Tabela %d – %s'%(_TN[0],cap)); r.font.size=Pt(11)
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
    t=doc.add_table(rows=1,cols=len(header)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    t.autofit=False
    tblPr=t._tbl.tblPr
    tl=OxmlElement('w:tblLayout'); tl.set(qn('w:type'),'fixed'); tblPr.append(tl)
    cm=OxmlElement('w:tblCellMar')
    for side in ('left','right'):
        e=OxmlElement('w:'+side); e.set(qn('w:w'),'40'); e.set(qn('w:type'),'dxa'); cm.append(e)
    tblPr.append(cm)
    nC=len(header); TOT=16.0
    if nC>2:
        c0=min(3.0,TOT/nC*1.5); cL=2.7 if len(str(header[-1]))>6 else (TOT-c0)/(nC-1)
        mid=(TOT-c0-cL)/(nC-2); _cw=[c0]+[mid]*(nC-2)+[cL]
    else:
        _cw=[TOT/nC]*nC
    grid=t._tbl.find(qn('w:tblGrid'))
    if grid is not None:
        for i,gc in enumerate(grid.findall(qn('w:gridCol'))):
            gc.set(qn('w:w'),str(int(_cw[i]*567)))
    for i,htx in enumerate(header):
        cc=t.rows[0].cells[i]; cc.text=''; cc.width=Cm(_cw[i]); rr=cc.paragraphs[0].add_run(htx); rr.bold=True; rr.font.size=Pt(fs); rr.font.name='Times New Roman'
        cc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; _bd(cc)
    for row in rows:
        cs=t.add_row().cells
        for i,val in enumerate(row):
            cs[i].text=''; cs[i].width=Cm(_cw[i]); rr=cs[i].paragraphs[0].add_run(str(val)); rr.font.size=Pt(fs); rr.font.name='Times New Roman'
            cs[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if i>0 else WD_ALIGN_PARAGRAPH.LEFT; _bd(cs[i])
    if note:
        pn=doc.add_paragraph(); rn=pn.add_run(note); rn.font.size=Pt(8.5); rn.italic=True; pn.paragraph_format.space_after=Pt(0)
    pf=doc.add_paragraph(); rf=pf.add_run(fonte); rf.font.size=Pt(9); pf.paragraph_format.space_after=Pt(6)
    return _TN[0]
def figure(path,cap,w=15.0):
    _FN[0]+=1
    pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; pp.add_run().add_picture(path,width=Cm(w)); pp.paragraph_format.space_before=Pt(6)
    pc=doc.add_paragraph(); pc.alignment=WD_ALIGN_PARAGRAPH.CENTER; rc=pc.add_run('Figura %d – %s'%(_FN[0],cap)); rc.font.size=Pt(11)
    pf=doc.add_paragraph(); pf.alignment=WD_ALIGN_PARAGRAPH.CENTER; rf=pf.add_run('Fonte: elaboração dos autores (2026).'); rf.font.size=Pt(9); pf.paragraph_format.space_after=Pt(6)
    return _FN[0]
def sfig(path,cap,w=15.0):
    # versao completa: todas as figuras entram no corpo, em ordem (sem material suplementar)
    return figure(path,cap,w)
def render_supp():
    pass

sm=R['sample']; desc=R['desc']; pr=R['prepos']; d17=R['d1d7']; PREV=MS['prev']
# ---- valores pré-calculados (f-strings, sem armadilha de %) ----
def num(x,d=1): return c2(f'{x:.{d}f}')
def pv(p): return '< 0,001' if p<0.001 else '= '+c2(f'{p:.3f}')
def pvt(p): return '< 0,001' if p<0.001 else c2(f'{p:.3f}')
fr_vig_p=pv(S4['friedman']['Vigor']['p']); fr_vig_w=num(S4['friedman']['Vigor']['W'],2)
fr_fad_p=pv(S4['friedman']['Fadiga']['p']); fr_fad_w=num(S4['friedman']['Fadiga']['W'],2)
vig_dz=num(pr['Vigor']['dz'],2); fad_dz=num(pr['Fadiga']['dz'],2)
vig_pct=num(abs(pr['Vigor']['pct']),0); fad_pct=num(pr['Fadiga']['pct'],0)
d1d7_vig_dz=num(d17['Vigor']['dz'],2); d1d7_fad_dz=num(d17['Fadiga']['dz'],2)
d1d7_vig_pct=num(abs(d17['Vigor']['pct']),0); d1d7_fad_pct=num(d17['Fadiga']['pct'],0)
wilks=num(MV['d1d7']['wilks'],3); p_mv=pv(MV['d1d7']['p_mv'])
vmaxd=PK['Vigor']['max_day']; vmind=PK['Vigor']['min_day']
fmaxd=PK['Fadiga']['max_day']; fmind=PK['Fadiga']['min_day']
ice_d1=num(100*PREV['D1']['Iceberg']/PREV['n_d1'],0)
ice_d7=num(100*PREV['D7']['Iceberg']/PREV['n_d7'],0)
bar_d1=num(100*PREV['D1']['Barbatana tubarão']/PREV['n_d1'],0)
bar_d7=num(100*PREV['D7']['Barbatana tubarão']/PREV['n_d7'],0)
sub_d1=num(100*PREV['D1']['Submerso']/PREV['n_d1'],0)
sub_d7=num(100*PREV['D7']['Submerso']/PREV['n_d7'],0)
or_bar=num(LOG['migracao']['barbatana']['OR_dia'],2)
pc1=num(100*PCA['var_ratio'][0],0); pc2=num(100*PCA['var_ratio'][1],0)
pc12=num(100*(PCA['var_ratio'][0]+PCA['var_ratio'][1]),0); nkaiser=PCA['n_kaiser']
pth_rv=num(abs(PSY['PTH']['rho_vigor']),2); pth_rf=num(PSY['PTH']['rho_fadiga'],2)
pth_r2=num(100*PSY['PTH']['r2_axis'],0)
a_vig=num(PSY['Vigor']['alpha'],2); a_fad=num(PSY['Fadiga']['alpha'],2)
idade=num(sm['idade']['mean'],1); idade_sd=num(sm['idade']['sd'],1)
exp=num(sm['exp']['mean'],1); exp_sd=num(sm['exp']['sd'],1)

# ===== TÍTULO =====
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('DINÂMICA DO HUMOR EM UM MICROCICLO PRÉ-COMPETITIVO DE HANDEBOL DE ELITE: PERFIS DE HUMOR E O EIXO ENERGIA–FADIGA')
r.bold=True; r.font.size=Pt(13); p.paragraph_format.space_after=Pt(4)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('(versão completa, com todas as análises no corpo do texto)'); r.italic=True; r.font.size=Pt(11); p.paragraph_format.space_after=Pt(10)

# ===== RESUMO =====
H('RESUMO',before=2)
RUN([('Objetivo: ',True),(f'descrever e caracterizar a dinâmica do humor de atletas de handebol de elite ao longo de um '
 f'microciclo pré-competitivo, com ênfase no eixo energia–fadiga, no comportamento dos seis perfis de humor, na resposta '
 f'aguda pré e pós-treino e na forma temporal das trajetórias. ',False),
 ('Método: ',True),(f'{sm["n"]} atletas do sexo masculino responderam ao BRUMS-24 durante sete dias, com uma coleta de '
 f'linha de base e duas coletas diárias (pré e pós-treino) nos seis dias de treino, e um total de {sm["n_obs"]} '
 f'observações. Além da estatística descritiva, da consistência interna e da classificação dos perfis, as comparações '
 f'entre dias e entre pré e pós-treino incluíram o tamanho e a magnitude do efeito, e as trajetórias foram analisadas '
 f'por suavização, segundas derivadas, ajuste polinomial, localização dos cruzamentos entre as dimensões, separação '
 f'entre sinal e ruído e comparação sequencial do efeito agudo (pré→pós) e de recuperação (pós→pré do dia seguinte) por '
 f'teste de Wilcoxon pareado. ',False),
 ('Resultados: ',True),(f'a deterioração concentrou-se no eixo energia–fadiga: o vigor caiu e a fadiga subiu do primeiro '
 f'para o último dia, com efeito grande (d = {d1d7_vig_dz} e d = {d1d7_fad_dz}) e confirmação multivariada (Wilks λ = {wilks}; p {p_mv}). '
 f'As trajetórias suavizadas apresentaram um ponto de inflexão na metade da semana, e a fadiga ultrapassou o vigor de '
 f'forma definitiva no fim do microciclo, com a PTH a superar ambas as dimensões. O vigor e a fadiga concentraram a '
 f'maior razão sinal/ruído, ao passo que as dimensões negativas mostraram forte efeito de piso. A análise sequencial '
 f'revelou um padrão em dente de serra, com piora aguda no treino, sobretudo no Dia 6 (PTH com dz = {num(ADV["trans"][9]["vars"]["TMD"]["dz"],2)}), '
 f'e recuperação apenas parcial entre as sessões. A prevalência dos perfis deslocou-se do iceberg ({ice_d1}% no primeiro '
 f'dia) para a barbatana de tubarão ({bar_d7}% no último dia), com aumento da chance desse perfil a cada dia (OR = {or_bar}). ',False),
 ('Conclusão: ',True),(f'o humor migrou da prontidão para a fadiga funcional, em um padrão compatível com sobre-esforço '
 f'funcional, o que recomenda centrar o monitoramento no par vigor–fadiga e no cruzamento entre as suas curvas.',False)],after=6)
P('Palavras-chave: humor; BRUMS; handebol; perfis de humor; fadiga; monitoramento do atleta.',size=11,after=8,ind=False)

# ===== 1 INTRODUÇÃO =====
H('1 INTRODUÇÃO')
P('O monitoramento do estado do atleta consolidou-se como parte da gestão do treinamento no esporte de rendimento, e '
 'integra medidas objetivas de carga externa e interna a instrumentos subjetivos de autorrelato (HELWIG et al., 2023; '
 'KELLMANN et al., 2018). Entre esses instrumentos, os questionários de humor destacam-se pela praticidade, pelo baixo '
 'custo e pela sensibilidade às variações da carga, e mostram utilidade para o acompanhamento do bem-estar e do '
 'desempenho (SAW; MAIN; GASTIN, 2016; LOCHBAUM et al., 2021). Revisões sistemáticas indicam, inclusive, que as medidas '
 'subjetivas costumam responder às alterações de carga com sensibilidade igual ou superior à de muitos marcadores '
 'objetivos, o que sustenta o seu uso rotineiro (SAW; MAIN; GASTIN, 2016). Por esse motivo, documentos de consenso '
 'recomendam a vigilância regular da fadiga e da recuperação como base para as decisões de treino (KELLMANN et al., 2018).')
P('A Escala de Humor de Brunel (BRUMS) operacionaliza essa avaliação por meio de seis dimensões, a saber, tensão, '
 'depressão, raiva, vigor, fadiga e confusão, com propriedades psicométricas replicadas em diferentes idiomas e culturas '
 '(TERRY; LANE; FOGARTY, 2003; TERRY et al., 2022). A versão em português dispõe de validação para o contexto brasileiro '
 'e foi concebida, desde a origem, como ferramenta de detecção precoce de sinais associados ao excesso de treinamento '
 '(ROHLFS et al., 2008). A partir das seis dimensões calcula-se a perturbação total do humor (PTH), um índice-resumo do '
 'desequilíbrio afetivo que sintetiza o estado do atleta em um único valor e cuja associação com o desempenho tem sido '
 'documentada em meta-análise (LOCHBAUM et al., 2021).')
P('Para além dos escores isolados, o modelo dos perfis de humor organiza as seis dimensões em padrões reconhecíveis. O '
 'clássico perfil iceberg, com o vigor acima da média e as dimensões negativas rebaixadas, foi descrito no modelo de '
 'saúde mental como marca do atleta em prontidão (MORGAN, 1985). Estudos posteriores, com grandes amostras e análise de '
 'agrupamento, identificaram e replicaram seis perfis distintos, entre os quais o iceberg, a barbatana de tubarão, que '
 'sinaliza fadiga com vigor ainda preservado, e perfis de maior risco, como o iceberg invertido e o submerso '
 '(PARSONS-SMITH; TERRY; MACHIN, 2017; LUOJUMÄKI et al., 2026). A leitura por perfis aproxima o dado psicométrico da '
 'linguagem do treinador, facilita a comunicação com a comissão técnica e tem sido proposta como recurso de rastreio da '
 'saúde mental no esporte (TERRY et al., 2021; HAN; PARSONS-SMITH; TERRY, 2020; LEW et al., 2023).')
P('Entre as seis dimensões, o vigor e a fadiga formam o eixo mais responsivo à carga e concentram boa parte do valor '
 'prático do monitoramento. Sob intensificação do treino, o vigor tende a cair e a fadiga a subir, um padrão observado '
 'em atletas de diferentes modalidades e sensível também à privação de sono (FERREIRA et al., 2026; PIERCE, 2002). Esse '
 'eixo responde ainda de forma aguda a cada sessão ou competição, com oscilações mensuráveis entre os momentos pré e '
 'pós-esforço, o que recomenda coletas repetidas dentro do dia, e não apenas entre dias (DO NASCIMENTO et al., 2026).')
P('A leitura conjunta dessas variações remete ao continuum do sobre-esforço. A distinção entre o sobre-esforço '
 'funcional, o sobre-esforço não funcional e a síndrome de overtraining é hoje descrita como um processo gradual, de '
 'difícil diagnóstico por um único marcador, no qual a fadiga central e a piora do humor figuram entre os sinais mais '
 'precoces (ROETE et al., 2021; LA TORRE et al., 2023; MĂNESCU et al., 2026). Nesse quadro, o acompanhamento do humor '
 'oferece um marcador sensível, de baixo custo e não invasivo, complementar aos marcadores fisiológicos, para vigiar a '
 'janela em que o sobre-esforço funcional é buscado de forma planejada (THORPE et al., 2017).')
P('O handebol oferece um cenário exigente para essa vigilância. A modalidade é coletiva, intermitente e de alta '
 'intensidade, com sprints curtos, mudanças de direção, saltos, arremessos e contato físico, o que impõe elevada demanda '
 'neuromuscular e psicofisiológica, variável por posição de jogo (KARCHER; BUCHHEIT, 2014; GARCÍA-SÁNCHEZ et al., 2023; '
 'CARTON-LLORENTE et al., 2023). O monitoramento longitudinal da carga interna ao longo da temporada e a atenção à '
 'fadiga, inclusive a mental, já foram descritos como relevantes para o rendimento na modalidade (STRUZIK; NADOBNIK; '
 'STĘPIEŃ-SŁODKOWSKA, 2026; STAIANO et al., 2025). A acumulação de carga na semana tende, assim, a corroer o vigor e a '
 'elevar a fadiga, um padrão que, quando controlado, caracteriza o sobre-esforço funcional e antecede a recuperação '
 'planejada.')
P(f'Apesar do interesse crescente, poucos estudos descrevem, dentro de um único microciclo de handebol de elite, a '
 f'migração dos perfis de humor e a forma temporal exata das trajetórias do eixo energia–fadiga, com coletas pré e '
 f'pós-treino que capturam também a variação dentro do dia (DE MIRANDA ROHLFS et al., 2024; BIRD et al., 2025; '
 f'RATZ-SULYOK et al., 2026). A Figura {_FN[0]+1} resume o quadro conceitual que orienta este trabalho.')
figure(f'{FG}/framework.png','Quadro conceitual do monitoramento do humor no eixo energia–fadiga: a carga do microciclo altera as dimensões do BRUMS, com queda do vigor e elevação da fadiga e da PTH, o que desloca os perfis de humor em direção à barbatana de tubarão, dentro da janela do sobre-esforço funcional, e retroalimenta as decisões de treino e recuperação.',w=15.5)

# ===== 2 JUSTIFICATIVA =====
H('2 JUSTIFICATIVA')
P('O monitoramento do humor reúne três atributos que o tornam atraente para a rotina do esporte de rendimento: é '
 'sensível às variações da carga, tem baixo custo e não é invasivo, o que favorece a coleta frequente sem acréscimo '
 'relevante de logística (SAW; MAIN; GASTIN, 2016; HELWIG et al., 2023). Em um calendário competitivo denso, no qual as '
 'medidas fisiológicas e endócrinas nem sempre são viáveis no dia a dia, um marcador subjetivo bem escolhido oferece à '
 'comissão técnica uma leitura rápida do estado do atleta e um apoio concreto às decisões de treino e recuperação '
 '(KELLMANN et al., 2018).')
P('A literatura, contudo, ainda descreve de forma incompleta a dinâmica do humor dentro de um único microciclo de '
 'handebol de elite. A maior parte dos estudos compara momentos isolados ou temporadas inteiras, sem capturar, ao mesmo '
 'tempo, a variação entre dias e a variação aguda entre pré e pós-treino, e sem descrever a forma temporal das '
 'trajetórias do eixo energia–fadiga (DE MIRANDA ROHLFS et al., 2024; DO NASCIMENTO et al., 2026). Essa lacuna é '
 'relevante porque é justamente no interior do microciclo que se instala o sobre-esforço funcional buscado de forma '
 'planejada, e cuja passagem para o sobre-esforço não funcional deve ser vigiada (ROETE et al., 2021; MĂNESCU et al., '
 '2026).')
P('Este estudo justifica-se, portanto, por reunir, em um mesmo delineamento, a descrição completa das seis dimensões e '
 'dos perfis de humor, a quantificação do tamanho do efeito e uma análise da forma temporal das trajetórias que vai além '
 'da comparação de médias. O uso de suavização, de segundas derivadas, de ajuste polinomial e da localização dos '
 'cruzamentos exatos entre vigor, fadiga e PTH converte a ideia qualitativa de inversão do eixo energia–fadiga em '
 'eventos datados, o que oferece à prática do handebol de elite um conjunto de marcadores objetivos e visuais, '
 'complementares ao monitoramento de carga e de fadiga já descrito na modalidade (STRUZIK; NADOBNIK; '
 'STĘPIEŃ-SŁODKOWSKA, 2026; STAIANO et al., 2025).')

# ===== 3 OBJETIVOS =====
H('3 OBJETIVOS')
H('3.1 Objetivo geral',12,before=6)
P('Descrever e caracterizar a dinâmica do humor de atletas de handebol de elite ao longo de um microciclo '
 'pré-competitivo, com ênfase no eixo energia–fadiga, no comportamento dos seis perfis de humor, na resposta aguda pré '
 'e pós-treino e na forma temporal das trajetórias.')
H('3.2 Objetivos específicos',12,before=6)
P('a) Descrever a estatística descritiva, a consistência interna e os limiares de mudança das seis dimensões do BRUMS e '
 'da perturbação total do humor.',ind=False)
P('b) Caracterizar os seis perfis de humor representados na amostra e quantificar a sua migração entre o primeiro e o '
 'último dia do microciclo.',ind=False)
P('c) Estimar a magnitude do efeito da mudança do humor entre o primeiro e o último dia, entre os sete dias e entre os '
 'momentos pré e pós-treino.',ind=False)
P('d) Modelar a forma temporal das trajetórias do eixo energia–fadiga por meio de suavização e da análise das segundas '
 'derivadas, com a localização dos pontos de inflexão.',ind=False)
P('e) Ajustar modelos polinomiais às trajetórias, em duas resoluções, e localizar os cruzamentos exatos entre o vigor, a '
 'fadiga e a perturbação total do humor.',ind=False)
P('f) Examinar a estrutura dimensional e as correlações entre as dimensões do humor e discutir o conjunto como um quadro '
 'de sobre-esforço funcional útil ao monitoramento aplicado.',ind=False)

# ===== 3 MÉTODO =====
H('4 MÉTODO')
P(f'Participaram {sm["n"]} atletas de handebol do sexo masculino, de nível de elite (idade de {idade} ± {idade_sd} anos; '
 f'{exp} ± {exp_sd} anos de prática), das posições de armador, ala, pivô e goleiro. O humor foi avaliado pela BRUMS-24, '
 f'que reúne seis dimensões (tensão, depressão, raiva, vigor, fadiga e confusão) com escores de 0 a 16, a partir das '
 f'quais se calculou também a perturbação total do humor (PTH). O delineamento cobriu sete dias de um microciclo '
 f'pré-competitivo, com uma coleta de linha de base no primeiro dia e duas coletas diárias, uma antes e outra depois do '
 f'treino, nos seis dias subsequentes, o que totalizou {sm["n_obs"]} observações. A Figura {_FN[0]+1} sintetiza o '
 f'delineamento e o plano de análise.')
figure(f'{FG}/organograma.png','Organograma do delineamento do estudo: da amostra e do microciclo às coletas de linha de base e pré/pós-treino, às 286 observações do BRUMS-24 e aos cinco blocos do plano de análise.',w=15.5)
P('Na análise, empregaram-se estatística descritiva das seis dimensões e da PTH e avaliação da consistência interna por '
 'alfa de Cronbach. Como o teste de Shapiro-Wilk apontou desvios da normalidade em parte das dimensões, adotaram-se '
 'testes não paramétricos: o teste de Wilcoxon para a comparação entre pré e pós-treino, o teste de Friedman com o W de '
 'Kendall para a comparação entre os sete dias e a correlação de Spearman para as relações entre as dimensões, sempre com '
 'o cálculo do tamanho do efeito, classificado por magnitude (trivial, pequeno, médio ou grande no d de Cohen e no dz; '
 'trivial, pequeno, moderado ou grande no W de Kendall). A confirmação multivariada da diferença entre o primeiro e o '
 'último dia recorreu à MANOVA em escores T. Cada observação foi classificada em um dos seis perfis de humor a partir '
 'dos escores padronizados, e a distribuição dos perfis entre o primeiro e o último dia foi avaliada pelo teste do '
 'qui-quadrado, com a tendência de migração resumida por uma regressão logística ao longo dos dias. A consistência das '
 'medidas repetidas foi estimada pelo coeficiente de correlação intraclasse (ICC), do qual se derivaram os limiares de '
 'mudança.')
P('A forma temporal das trajetórias foi analisada por três abordagens complementares. Primeiro, as séries de vigor, '
 'fadiga e PTH foram suavizadas por spline, e a segunda derivada de cada curva localizou o ponto de inflexão, definido '
 'como a raiz da segunda derivada. Segundo, cada trajetória foi modelada por um ajuste polinomial de grau três, em duas '
 'resoluções, sobre as sete médias diárias e sobre os quatorze pontos pré e pós-treino, do qual se extraíram a equação, '
 'o coeficiente de determinação (R²) e a inflexão analítica (x = -b/3a). Terceiro, sobre as curvas ajustadas foram '
 'localizados os cruzamentos exatos entre as dimensões, ou seja, os dias em que uma curva iguala e ultrapassa a outra. '
 'As análises foram conduzidas em Python (pacotes NumPy, SciPy, pandas, statsmodels e Plotly), com nível de '
 'significância de 5%.')

# ===== 4 RESULTADOS =====
H('5 RESULTADOS')
P('Os resultados são apresentados em três níveis de profundidade, em ordem lógica. O primeiro nível, descritivo, resume '
 'as seis dimensões e a PTH, avalia a consistência interna e os limiares de mudança, e caracteriza os perfis de humor e '
 'a sua migração. O segundo nível, exploratório, compara os dias e os momentos pré e pós-treino com o cálculo do tamanho '
 'e da magnitude do efeito, e examina as relações entre as dimensões e a estrutura de componentes principais. O terceiro '
 'nível, avançado, modela a forma temporal das trajetórias por suavização, segundas derivadas, ajuste polinomial e '
 'localização dos cruzamentos, separa o sinal do ruído em cada dimensão e decompõe a resposta em efeito agudo e de '
 'recuperação ao longo das treze coletas. Cada análise é descrita e, em seguida, interpretada quanto ao seu significado '
 'para o monitoramento do atleta.')

H('5.1 Descrição e consistência das dimensões',12,before=6)
P('A Tabela 1 resume as seis dimensões do BRUMS e a PTH no conjunto das observações. O vigor apresentou a maior média '
 'entre as dimensões, ao passo que as dimensões negativas concentraram-se em valores baixos, com médias próximas do '
 'limite inferior da escala. Esse padrão indica um grupo, em geral, bem ajustado, no qual a fadiga e o vigor ocupam a '
 'maior parte da faixa de resposta e sustentam a leitura do estado do atleta.')
ORD=[('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Vigor','Vigor'),('Fadiga','Fadiga'),('Confusao','Confusão'),('TMD','PTH')]
rows1=[]
for k,lab in ORD:
    dd=desc[k]
    rows1.append([lab,num(dd['mean'],1),num(dd['sd'],1),f"{num(dd['mn'],0)}–{num(dd['mx'],0)}"])
table('Estatística descritiva das dimensões do BRUMS e da perturbação total do humor (286 observações).',
 ['Dimensão','Média','DP','Mín.–Máx.'],rows1,
 note='DP: desvio-padrão. PTH: perturbação total do humor. Escores das dimensões variam de 0 a 16.')
P(f'A Figura {_FN[0]+1} apresenta os diagramas de caixa das seis dimensões, em ordem canônica e em painéis separados, '
 f'com a mediana e a média de cada uma. Como as dimensões diferem muito em amplitude, cada painel usa a escala da própria '
 f'variável, o que torna a distribuição de cada uma mais legível. O vigor e a fadiga percorrem a maior parte da escala e '
 f'apresentam média e mediana próximas, sinal de distribuições razoavelmente simétricas, ao passo que as dimensões '
 f'negativas concentram a caixa junto ao zero e exibem a média acima da mediana, o que revela assimetria positiva e um '
 f'efeito de piso.')
figure(f'{FG}/box_humor.png','Diagramas de caixa das seis dimensões do BRUMS, em painéis separados e ordem canônica; cada painel usa a escala da própria variável, com a mediana (linha sólida) e a média (linha tracejada e losango).',w=13.5)
P(f'A consistência interna foi adequada nas duas dimensões do eixo energia–fadiga (alfa de {a_vig} para o vigor e {a_fad} '
 f'para a fadiga), o que reforça a confiança na sua medida. A PTH, por reunir as seis dimensões em um único índice, '
 f'associou-se de forma forte ao vigor e à fadiga (correlações de {pth_rv} e {pth_rf}, respectivamente), e essas duas '
 f'dimensões explicaram {pth_r2}% da sua variância, o que confirma o eixo energia–fadiga como o núcleo do sinal.')
IC=STAT['icc']
def icrow(k,lab):
    i=IC[k]; return [lab,num(i['icc1'],2),num(i['icck'],2),i['cls']]
table('Consistência das medidas repetidas ao longo da semana (coeficiente de correlação intraclasse, ICC).',
 ['Dimensão','ICC(2,1)','ICC(2,k)','Consistência'],
 [icrow(k,l) for k,l in [('Vigor','Vigor'),('Fadiga','Fadiga'),('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Confusao','Confusão')]],
 note='ICC(2,1) = medida isolada; ICC(2,k) = média das medidas. Consistência: < 0,50 pobre; 0,50–0,75 moderada; > 0,75 boa.',fs=9)
P(f'A consistência das medidas repetidas ao longo da semana (Tabela {_TN[0]}) foi moderada a boa, com os valores mais '
 f'baixos na raiva e na confusão, dimensões mais reativas de um dia para o outro, o que também recomenda cautela na '
 f'leitura isolada dessas dimensões.')
P(f'A partir do ICC derivaram-se os limiares de mudança das duas dimensões do eixo energia–fadiga (Tabela {_TN[0]+1}), '
 f'úteis para separar a variação real do ruído de medida. No vigor, o erro-padrão de medida foi de '
 f'{num(MDC["Vigor"]["sem"],1)} ponto e a mudança mínima detectável, de {num(MDC["Vigor"]["mdc90"],1)} (90%) a '
 f'{num(MDC["Vigor"]["mdc95"],1)} (95%) pontos; na fadiga, de {num(MDC["Fadiga"]["sem"],1)} e de '
 f'{num(MDC["Fadiga"]["mdc90"],1)} a {num(MDC["Fadiga"]["mdc95"],1)} pontos, respectivamente. Esses valores superam a '
 f'menor mudança relevante (SWC = {num(MDC["Vigor"]["swc"],1)} no vigor e {num(MDC["Fadiga"]["swc"],1)} na fadiga), o que '
 f'indica que, no plano individual, apenas oscilações de cerca de {num(MDC["Vigor"]["mdc90"],1)} a '
 f'{num(MDC["Fadiga"]["mdc90"],1)} pontos podem ser lidas como mudança real, ao passo que variações menores se confundem '
 f'com o ruído. Por esse motivo, a interpretação deste estudo apoia-se nas tendências de grupo e não em leituras '
 f'isoladas por atleta.')
table('Erro de medida e limiares de mudança do vigor e da fadiga (escala 0–16): erro-padrão de medida (SEM), mudança mínima detectável (MDC) e menor mudança relevante (SWC).',
 ['Dimensão','SEM','MDC90','MDC95','SWC'],
 [[l,num(MDC[k]['sem'],1),num(MDC[k]['mdc90'],1),num(MDC[k]['mdc95'],1),num(MDC[k]['swc'],1)] for k,l in [('Vigor','Vigor'),('Fadiga','Fadiga')]],
 note='SEM = DP × √(1 - ICC); MDC = 1,65 (90%) ou 1,96 (95%) × √2 × SEM; SWC = 0,2 × DP entre atletas. Uma mudança individual só é tomada como real quando excede a MDC.',fs=9)

H('5.2 Perfis de humor e sua migração',12,before=6)
P(f'Os seis perfis de humor descritos na literatura estiveram representados na amostra. A Figura {_FN[0]+1} apresenta '
 f'cada perfil identificado no estudo em escores T, com a respectiva prevalência, o que permite reconhecer a sua forma '
 f'característica. O perfil iceberg exprime prontidão, com o vigor acima da média e as dimensões negativas abaixo, ao '
 f'passo que a barbatana de tubarão sinaliza pico isolado de fadiga, o submerso reúne todas as dimensões abaixo da média '
 f'e o Everest invertido eleva todas as dimensões negativas (PARSONS-SMITH; TERRY; MACHIN, 2017).')
figure(f'{FG}/xb6_clusters.png','Perfis de humor identificados na amostra, em escores T (M = 50; DP = 10) nas seis dimensões; prevalência de cada perfil entre parênteses.',w=12.5)
P(f'A comparação entre o primeiro e o último dia do microciclo revelou a reconfiguração do perfil médio do grupo. A '
 f'Figura {_FN[0]+1} mostra que, no primeiro dia, o perfil assumiu o formato iceberg, com o vigor no topo e as dimensões '
 f'negativas abaixo da média populacional. No último dia, o perfil inverteu-se para a forma de barbatana de tubarão, com '
 f'a fadiga no topo e o vigor rebaixado, o que traduz a acumulação da carga ao longo da semana.')
figure(f'{FG}/xb5_profile_d1d7.png','Perfil de humor em escores T no primeiro e no último dia do microciclo.',w=13.0)
P(f'Essa mudança de forma correspondeu a um deslocamento da prevalência dos perfis (Tabela {_TN[0]+1}). O perfil iceberg '
 f'caiu de {ice_d1}% no primeiro dia para {ice_d7}% no último, enquanto a barbatana de tubarão subiu de {bar_d1}% para '
 f'{bar_d7}% e o perfil submerso passou de {sub_d1}% para {sub_d7}%. A reorganização categórica não alcançou '
 f'significância no teste do qui-quadrado (χ² = {num(PREV["chi"],2)}; p {pv(PREV["p"])}), resultado esperado pela baixa '
 f'contagem por célula quando poucas observações se distribuem por seis perfis. A regressão logística de tendência, '
 f'porém, confirmou o aumento da chance do perfil de barbatana de tubarão a cada dia (OR = {or_bar}), sem que os perfis '
 f'de maior risco à saúde mental se instalassem.')
PROFR=[('Iceberg','Iceberg'),('Everest invertido','Everest invertido'),('Iceberg invertido','Iceberg invertido'),('Submerso','Submerso'),('Barbatana tubarão','Barbatana de tubarão'),('Superfície','Superfície')]
def prow(p,lab):
    d1=PREV['D1'][p]; d7=PREV['D7'][p]
    return [lab,f"{d1} ({num(100*d1/PREV['n_d1'],1)}%)",f"{d7} ({num(100*d7/PREV['n_d7'],1)}%)"]
table('Distribuição dos seis perfis de humor no primeiro e no último dia do microciclo: n (%).',
 ['Perfil','Dia 1, n (%)','Dia 7, n (%)'],[prow(p,l) for p,l in PROFR],
 note='χ² = %s; p = %s. A migração é descritiva e converge com a queda do vigor e a elevação da fadiga.'%(num(PREV['chi'],2),num(PREV['p'],3)),fs=9)

H('5.3 Diferença entre o primeiro e o último dia (com tamanho de efeito)',12,before=6)
P(f'A comparação entre o primeiro e o último dia do microciclo quantificou a magnitude da mudança em cada dimensão '
 f'(Tabela {_TN[0]+1}). O efeito foi grande no vigor, que caiu {d1d7_vig_pct}% (dz = {d1d7_vig_dz}), e na fadiga, que '
 f'subiu {d1d7_fad_pct}% (dz = {d1d7_fad_dz}), ao passo que as demais dimensões apresentaram efeitos de menor magnitude. '
 f'A análise multivariada confirmou a diferença global entre os dois dias (Wilks λ = {wilks}; p {pv(MV["d1d7"]["p_mv"])}), '
 f'e, sob a correção de Bonferroni, apenas o vigor e a fadiga permaneceram significativos, o que concentra o efeito no '
 f'eixo energia–fadiga.')
def d7row(k,lab):
    v=d17[k]; return [lab,num(v['d1'],2),num(v['d7'],2),c2(f"{v['pct']:+.0f}")+'%',pvt(v['p']),c2(f"{v['dz']:+.2f}"),v['mag']]
D7ORD=[('Vigor','Vigor'),('Fadiga','Fadiga'),('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Confusao','Confusão'),('TMD','PTH')]
table('Diferença das dimensões do BRUMS e da PTH entre o primeiro e o último dia do microciclo (com tamanho e magnitude do efeito).',
 ['Dimensão','Dia 1 (M)','Dia 7 (M)','Variação (%)','p','dz','Magnitude'],[d7row(k,l) for k,l in D7ORD],
 note='dz = tamanho de efeito intraindividual (mudança padronizada); magnitude: trivial (< 0,2); pequeno (0,2–0,5); médio (0,5–0,8); grande (> 0,8). PTH: perturbação total do humor.',fs=9)

H('5.4 Comparação das dimensões entre os dias (Friedman)',12,before=6)
P(f'A comparação das dimensões entre os sete dias pelo teste de Friedman (Tabela {_TN[0]+1}) apontou variação '
 f'significativa no vigor (p {fr_vig_p}; W = {fr_vig_w}) e na fadiga (p {fr_fad_p}; W = {fr_fad_w}), com magnitude '
 f'pequena a moderada, além de diferenças na tensão e na confusão. O vigor foi máximo no Dia {vmaxd} e mínimo no Dia '
 f'{vmind}, enquanto a fadiga percorreu o caminho inverso, com pico no Dia {fmaxd}, um padrão coerente com o acúmulo de '
 f'carga dentro da faixa funcional.')
FR=S4['friedman']
def wmag(w):
    return 'trivial' if w<0.1 else 'pequeno' if w<0.3 else 'moderado' if w<0.5 else 'grande'
def ddrow(k,lab):
    dm=CRS['dec'][k]['dm']; f=FR[k]
    return [lab]+[num(dm[str(d)],1) for d in range(1,8)]+[num(f['chi'],1),pvt(f['p']),num(f['W'],2),wmag(f['W'])]
table('Médias diárias das dimensões do BRUMS e teste de Friedman com W de Kendall e magnitude do efeito (comparação entre os sete dias).',
 ['Dimensão','D1','D2','D3','D4','D5','D6','D7','χ²','p','W','Magnitude'],[ddrow(k,l) for k,l in ORD],
 note='W de Kendall = tamanho de efeito do teste de Friedman; magnitude: trivial (< 0,1); pequeno (0,1–0,3); moderado (0,3–0,5); grande (> 0,5). PTH: perturbação total do humor.',fs=8)
P(f'A Figura {_FN[0]+1} ilustra a trajetória do vigor, da fadiga e da PTH ao longo da semana, com a queda progressiva do '
 f'vigor e a elevação da fadiga e da PTH em direção ao fim do microciclo.')
figure(f'{FG}/abnt_f1_trajetoria.png','Trajetória de vigor, fadiga e perturbação total do humor ao longo dos sete dias (curvas suavizadas; ponto de inflexão e extremos sinalizados).',w=12.5)
P(f'A Figura {_FN[0]+1} detalha, por meio de diagramas de caixa por dia, a distribuição de cada dimensão ao longo da '
 f'semana. O deslocamento das caixas do vigor para baixo e das caixas da fadiga para cima, sobretudo na segunda metade do '
 f'microciclo, torna visível a mesma deterioração do eixo energia–fadiga, ao passo que as dimensões negativas de menor '
 f'expressão mantêm caixas comprimidas junto ao zero.')
sfig(f'{FG}/xb3_box.png','Diagramas de caixa das seis dimensões do BRUMS por dia do microciclo (áreas sombreadas: início e acúmulo da semana).',w=15.0)
P(f'A comparação direta da PTH entre três momentos do microciclo (Figura {_FN[0]+1}) ilustra a mesma tendência: a '
 f'perturbação total do humor aumentou do primeiro para o último dia, e a diferença foi significativa entre o Dia 1 e o '
 f'Dia 7 (p = 0,008), ao passo que as comparações entre dias adjacentes não alcançaram significância.')
sfig(f'{FG}/box_signif.png','Perturbação total do humor (PTH) no Dia 1, no Dia 4 e no Dia 7, com as comparações par a par (teste de Mann-Whitney; * p < 0,05; ** p < 0,01; ns = não significativo).',w=11.5)
P(f'A distribuição completa da PTH em cada dia (Figura {_FN[0]+1}) e a sua distribuição acumulada (Figura {_FN[0]+2}) confirmam o deslocamento ao longo da semana. Os diagramas de violino mostram a densidade da PTH a subir do Dia 1 ao Dia 7, e as curvas de distribuição acumulada revelam um deslocamento sistemático para a direita, com o Dia 7 situado acima do Dia 1 em praticamente toda a faixa de escores.')
sfig(f'{FG}/tec_violin.png','Distribuição da PTH por dia (diagramas de violino com caixa e média; a escala foi ajustada à faixa central da distribuição para facilitar a leitura).',w=13.0)
sfig(f'{FG}/tec_ecdf.png','Distribuição acumulada (ECDF) da PTH no Dia 1, no Dia 4 e no Dia 7; o deslocamento das curvas para a direita indica o aumento da perturbação ao longo do microciclo.',w=12.5)

H('5.5 Suavização, derivadas e cruzamento das trajetórias',12,before=6)
P(f'Para separar o sinal do ruído, as trajetórias do vigor, da fadiga e da PTH foram suavizadas sobre os doze pontos pré '
 f'e pós-treino por meio de uma spline, e a segunda derivada de cada curva localizou o seu ponto de inflexão, ou seja, o '
 f'momento em que a concavidade muda e a taxa de variação atinge o seu limite (Figura {_FN[0]+1}; Tabela {_TN[0]+1}). '
 f'Depois de removido o ruído, cada dimensão exibiu um único ponto de inflexão, situado na metade da semana (em torno do '
 f'Dia {num(SMO["V"]["infl"][0],1)}), o que marca a transição entre a fase inicial de prontidão e a fase de acúmulo de '
 f'carga. A curva suavizada do vigor subiu até um máximo em torno do Dia {num(SMO["V"]["ymax_x"],1)} e caiu ao seu '
 f'mínimo no fim da semana, enquanto a fadiga e a PTH percorreram o caminho inverso e atingiram os seus máximos no '
 f'último dia.')
figure(f'{FG}/smooth_deriv.png','Trajetórias suavizadas do vigor, da fadiga e da PTH sobre os doze pontos pré e pós-treino. Marcadores translúcidos: sinal bruto; linha grossa: sinal suavizado; linha pontilhada: ponto de inflexão (segunda derivada nula); triângulos: máximo e mínimo. Áreas sombreadas: início e acúmulo da semana.',w=12.5)
def smrow(k,lab):
    r=SMO[k]; inf=('Dia '+num(r['infl'][0],1)) if r['infl'] else 'n/d'
    return [lab,inf,f"Dia {num(r['ymax_x'],1)} ({num(r['ymax'],1)})",f"Dia {num(r['ymin_x'],1)} ({num(r['ymin'],1)})"]
table('Ponto de inflexão (segunda derivada nula) e extremos das trajetórias suavizadas do eixo energia–fadiga.',
 ['Dimensão','Inflexão (dia)','Máximo: dia (escore)','Mínimo: dia (escore)'],
 [smrow('V','Vigor'),smrow('F','Fadiga'),smrow('P','PTH')],
 note='Curvas suavizadas por spline sobre os 12 pontos pré/pós; inflexão = raiz da segunda derivada. PTH: perturbação total do humor.')
P(f'A mesma inflexão foi examinada por dois caminhos complementares. O painel da segunda derivada localiza o cruzamento '
 f'por zero de cada curva e realça os extremos com o respectivo valor (Figura {_FN[0]+1}). Em paralelo, um ajuste '
 f'polinomial de grau três resume cada trajetória por uma equação e reproduz a inflexão de forma analítica, na raiz da '
 f'segunda derivada (x = -b/3a), em duas resoluções (Figura {_FN[0]+2}). Sobre as sete médias diárias, o ajuste foi '
 f'forte, com R² de {num(POLY["V"]["d7"]["r2"],2)} para o vigor, {num(POLY["F"]["d7"]["r2"],2)} para a fadiga e '
 f'{num(POLY["P"]["d7"]["r2"],2)} para a PTH, e inflexões nos dias {num(POLY["V"]["d7"]["infl"],1)}, '
 f'{num(POLY["F"]["d7"]["infl"],1)} e {num(POLY["P"]["d7"]["infl"],1)}. Sobre os quatorze pontos pré e pós-treino, '
 f'incluída a linha de base do primeiro dia, o R² foi de {num(POLY["V"]["d14"]["r2"],2)}, {num(POLY["F"]["d14"]["r2"],2)} '
 f'e {num(POLY["P"]["d14"]["r2"],2)}, com inflexões praticamente idênticas. As duas resoluções convergem para a transição '
 f'na metade da semana e reforçam, por via paramétrica, o achado da spline.')
figure(f'{FG}/deriv_impact.png','Análise por derivadas do eixo energia–fadiga. Em cada variável, o painel superior traz a curva suavizada com o máximo e o mínimo realçados em cápsulas de valor, e o painel inferior traz a segunda derivada (concavidade) com o cruzamento por zero destacado, que marca o ponto de inflexão da trajetória.',w=15.0)
figure(f'{FG}/poly_fit.png','Ajuste polinomial de grau três das trajetórias do vigor, da fadiga e da PTH em duas resoluções: sobre as sete médias diárias (marcadores quadrados e linha cheia) e sobre os quatorze pontos pré e pós-treino, incluída a linha de base do primeiro dia (pontos translúcidos e linha tracejada). Cada painel traz a equação do ajuste diário, o R² de cada resolução e a inflexão analítica (raiz da segunda derivada, x = -b/3a).',w=14.5)
vf=CRX['cross']['VF']; vp=CRX['cross']['VP']; fp=CRX['cross']['FP']
P(f'Como o vigor e a fadiga compartilham a mesma escala, as suas trajetórias podem ser comparadas de forma direta e os '
 f'seus cruzamentos exatos podem ser localizados (Figura {_FN[0]+1}). No primeiro dia, o vigor superou a fadiga por ampla '
 f'margem (média de {num(d17["Vigor"]["d1"],1)} contra {num(d17["Fadiga"]["d1"],1)} pontos). O vigor caiu e a fadiga '
 f'subiu até que as duas curvas se igualaram pela primeira vez em torno do dia {num(vf[0][0],1)} (escore '
 f'{num(vf[0][1],1)}). A partir daí, o vigor e a fadiga percorreram uma faixa estreita e próxima e voltaram a cruzar-se '
 f'nos dias {num(vf[1][0],1)} e {num(vf[-1][0],1)}, sinal de um equilíbrio instável entre energia e fadiga durante a '
 f'maior parte da semana. Após o último cruzamento, no dia {num(vf[-1][0],1)} (escore {num(vf[-1][1],1)}), a fadiga '
 f'afastou-se de forma definitiva acima do vigor. No mesmo trecho final, a perturbação total do humor, que vinha abaixo '
 f'das duas dimensões, subiu e ultrapassou primeiro o vigor (dia {num(vp[0][0],1)}; escore {num(vp[0][1],1)}) e depois a '
 f'fadiga (dia {num(fp[0][0],1)}; escore {num(fp[0][1],1)}), o que marca a deterioração conjunta do estado de humor no '
 f'fim do microciclo.')
figure(f'{FG}/cross_traj.png','Cruzamentos exatos das trajetórias de vigor, fadiga e PTH ao longo do microciclo (curvas do ajuste polinomial sobre as médias diárias; losangos: pontos de cruzamento, com o dia e o escore). Áreas sombreadas: início e acúmulo da semana.',w=15.0)
vfc=ADV['vf_cross']; vfl=ADV['vf_last']
P(f'Para examinar o cruzamento decisivo do eixo energia–fadiga com maior detalhe, a Figura {_FN[0]+1} amplia a região '
 f'do último cruzamento e representa a diferença suavizada entre o vigor e a fadiga com a área sombreada de acordo com o '
 f'sinal: a área é positiva enquanto o vigor domina e negativa quando a fadiga passa a dominar. Sobre as treze coletas, '
 f'o vigor e a fadiga cruzaram-se três vezes (dias {num(vfc[0]["x"],1)}, {num(vfc[1]["x"],1)} e {num(vfl["x"],1)}), o '
 f'que confirma um longo trecho de quase equilíbrio no miolo da semana. O cruzamento do dia {num(vfl["x"],1)} é o '
 f'definitivo: a partir dele a diferença torna-se cada vez mais negativa e alcança cerca de três pontos de vantagem da '
 f'fadiga ao fim do microciclo, sem retorno.')
figure(f'{FG}/deriv_zoom.png','Cruzamento vigor × fadiga. Painel superior: curvas suavizadas com a área sombreada entre elas e os três cruzamentos assinalados. Painel inferior: recorte ampliado da diferença suavizada (vigor - fadiga), com a área sombreada segundo o sinal (verde: vigor domina; laranja: fadiga domina) e o cruzamento definitivo destacado.',w=14.5)
P(f'A análise por suavização e derivadas foi então estendida às seis dimensões e à PTH (Figura {_FN[0]+1}; '
 f'Tabela {_TN[0]+1}). A separação entre sinal e ruído variou muito entre as dimensões. O vigor e a fadiga '
 f'apresentaram a maior razão sinal/ruído (S/R de {num(ADV["snr"]["Vigor"]["snr_amp"],1)} e '
 f'{num(ADV["snr"]["Fadiga"]["snr_amp"],1)}), com trajetórias limpas e um único ponto de inflexão na metade da semana. '
 f'Já as dimensões negativas exibiram sinal fraco e forte ruído (S/R entre {num(min(ADV["snr"][k]["snr_amp"] for k in ["Tensao","Depressao","Raiva","Confusao"]),1)} '
 f'e {num(max(ADV["snr"][k]["snr_amp"] for k in ["Tensao","Depressao","Raiva","Confusao"]),1)}), com trajetórias '
 f'irregulares e de pequena amplitude, marcadas por um forte efeito de piso.')
figure(f'{FG}/all_deriv.png','Suavização e derivadas das seis dimensões do BRUMS e da PTH, sobre as treze coletas. Marcadores: sinal bruto (ruído); linha grossa: sinal suavizado; linha pontilhada: inflexão (segunda derivada nula); triângulos: máximo e mínimo. S/R: razão sinal/ruído (amplitude); piso: percentual de escores nulos.',w=15.5)
SNRV=[('Vigor','Vigor'),('Fadiga','Fadiga'),('TMD','PTH'),('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Confusao','Confusão')]
def snrow(k,lab):
    s=ADV['snr'][k]; inf=('Dia '+num(s['infl'][0],1)) if s['infl'] else 'n/d'
    return [lab,num(s['signal_amp'],1),num(s['noise_sd'],2),num(s['snr_amp'],1),c2(f"{s['floor0']:.0f}")+'%',inf]
table('Separação entre sinal e ruído, efeito de piso e ponto de inflexão de cada dimensão (curvas suavizadas sobre as 13 coletas).',
 ['Dimensão','Amplitude do sinal','Ruído (DP)','S/R','Piso (% de zeros)','Inflexão (1º dia)'],
 [snrow(k,l) for k,l in SNRV],
 note='Amplitude do sinal = máximo - mínimo da curva suavizada; ruído = desvio-padrão dos resíduos; S/R = razão sinal/ruído; piso = percentual de observações com escore zero.',fs=8.5)

H('5.6 Dias de pico e comparação de cada dia ao primeiro',12,before=6)
P(f'A localização dos picos sintetiza a dinâmica semanal (Tabela {_TN[0]+1}). O vigor foi máximo no Dia {vmaxd} e mínimo '
 f'no Dia {vmind}; a fadiga foi máxima no Dia {fmaxd}; a tensão, no Dia {PK["Tensao"]["max_day"]}; a raiva, no Dia '
 f'{PK["Raiva"]["max_day"]}; e a depressão, no Dia {PK["Depressao"]["max_day"]}. Em síntese, o início da semana reúne '
 f'maior prontidão, com vigor e tensão mais altos, ao passo que o final concentra a fadiga e a raiva.')
def pkrow(k,lab):
    p=PK[k]; return [lab,f'Dia {p["max_day"]}',num(p['max_val'],2),f'Dia {p["min_day"]}',num(p['min_val'],2)]
table('Dia de maior e de menor expressão de cada dimensão do humor (médias diárias).',
 ['Dimensão','Dia de maior valor','Valor','Dia de menor valor','Valor'],
 [pkrow(k,l) for k,l in ORD[:6]],fs=9)
def emm(v,d): return num(PHJ[v]['emm'][str(d)],2)
def sig1(v,d): return '' if d==1 or PHJ[v]['pairs'].get('1_%d'%d,{}).get('ptukey',1)>=0.05 else '*'
np_vig=sum(1 for k,pp in PHJ['Vigor']['pairs'].items() if pp['ptukey']<0.05)
np_fad=sum(1 for k,pp in PHJ['Fadiga']['pairs'].items() if pp['ptukey']<0.05)
P(f'No pós-teste que compara todos os dias entre si (Tabela {_TN[0]+1}; Figura {_FN[0]+1}), o vigor diferiu de forma '
 f'significativa em {np_vig} dos 21 pares de dias e a fadiga em {np_fad}, sempre no sentido de piora em relação aos '
 f'primeiros dias, o que confirma a deterioração progressiva do eixo energia–fadiga ao longo do microciclo.')
table('Médias marginais estimadas por dia do vigor e da fadiga, com comparação de cada dia ao Dia 1 (pós-teste de Tukey).',
 ['Dia','Vigor','Fadiga'],
 [[f'Dia {d}',emm('Vigor',d)+sig1('Vigor',d),emm('Fadiga',d)+sig1('Fadiga',d)] for d in range(1,8)],
 note='* diferença significativa em relação ao Dia 1 (Tukey, p < 0,05).',fs=9)
sfig(f'{FG}/ph_emm.png','Médias marginais diárias do vigor, da fadiga e da fadiga física, com comparação de todos os dias ao Dia 1 (* p < 0,05).',w=14.5)

H('5.7 Variação entre Pré e Pós-treino e Efeito Agudo/Recuperação',12,before=6)
P(f'A comparação entre pré e pós-treino pelo teste de Wilcoxon (Tabela {_TN[0]+1}) evidenciou uma resposta aguda '
 f'coerente com o esforço: o vigor caiu (d = {vig_dz}) e a fadiga e a PTH subiram do momento pré para o pós, com efeito '
 f'de magnitude pequena a moderada. Essa oscilação dentro do dia somou-se à tendência semanal e ajudou a compor o quadro '
 f'de fadiga funcional observado no fim do microciclo. A Figura {_FN[0]+1} apresenta esses escores por dia e mostra que a '
 f'diferença entre pré e pós-treino se mantém ao longo da semana.')
def wrow(k,lab):
    v=pr[k]; return [lab,num(v['pre'],2),num(v['pos'],2),c2(f"{v['pct']:+.0f}")+'%',pvt(v['p']),c2(f"{v['dz']:+.2f}"),v['mag']]
table('Comparação entre pré e pós-treino das dimensões do BRUMS e da PTH (teste de Wilcoxon e d de Cohen).',
 ['Dimensão','Pré (M)','Pós (M)','Variação (%)','p','d','Magnitude'],[wrow(k,l) for k,l in ORD],
 note='p do teste de Wilcoxon; d = tamanho de efeito de Cohen. PTH: perturbação total do humor.',fs=9)
sfig(f'{FG}/abnt_f_prepos.png','Escores de vigor, fadiga e perturbação total do humor no pré e no pós-treino, por dia.',w=12.5)
P(f'Para descrever a resposta aguda com maior rigor, cada uma das treze coletas foi comparada à coleta imediatamente '
 f'anterior por meio do teste de Wilcoxon pareado, o que separa o efeito agudo do treino (transição pré→pós, dentro do '
 f'dia) do efeito de recuperação entre sessões (transição pós→pré do dia seguinte). A Figura {_FN[0]+1} representa esse '
 f'percurso em forma de dente de serra e a Tabela {_TN[0]+1} reúne o tamanho de efeito de cada transição no eixo '
 f'energia–fadiga.')
figure(f'{FG}/sequential.png','Efeito agudo (pré→pós, segmentos laranja) e de recuperação (pós→pré do dia seguinte, segmentos azuis) da PTH ao longo das treze coletas, com o vigor e a fadiga como contexto. Os rótulos indicam o tamanho de efeito das transições significativas da PTH (* p < 0,05, Wilcoxon pareado).',w=15.5)
def tz(v):
    return (c2(f"{v['dz']:+.2f}")+('*' if v['sig'] else ''))
TR=ADV['trans']
table('Transições sequenciais entre coletas: efeito agudo (pré→pós) e de recuperação (pós→pré) no eixo energia–fadiga (tamanho de efeito dz; teste de Wilcoxon pareado).',
 ['Transição','Tipo','Vigor (dz)','Fadiga (dz)','PTH (dz)'],
 [[t['lab'],'Agudo' if t['tipo'].startswith('Agudo') else 'Recuperação',tz(t['vars']['Vigor']),tz(t['vars']['Fadiga']),tz(t['vars']['TMD'])] for t in TR],
 note='dz = tamanho de efeito intraindividual; * transição significativa (p < 0,05). Agudo = pré→pós no mesmo dia; Recuperação = pós→pré do dia seguinte (base = coleta única do Dia 1). PTH: perturbação total do humor.',fs=8)
n_ag=sum(1 for t in TR if t['tipo'].startswith('Agudo') and t['vars']['TMD']['sig'])
n_rec=sum(1 for t in TR if t['tipo'].startswith('Recup') and t['vars']['TMD']['sig'])
P(f'O padrão foi consistente com o de um sistema que oscila sob carga e se recupera de forma parcial. As transições '
 f'agudas tenderam a elevar a fadiga e a PTH e a reduzir o vigor, com destaque para a sessão do Dia 6, na qual a PTH '
 f'subiu com efeito grande (dz = {num(TR[9]["vars"]["TMD"]["dz"],2)}) e o vigor caiu (dz = {num(TR[9]["vars"]["Vigor"]["dz"],2)}). '
 f'As transições de recuperação, entre o pós de um dia e o pré do seguinte, moveram-se no sentido inverso e devolveram '
 f'parte do vigor e reduziram a fadiga e a PTH, com o episódio mais nítido entre o Dia 2 e o Dia 3 (PTH com '
 f'dz = {num(TR[2]["vars"]["TMD"]["dz"],2)}). Ainda assim, a recuperação noturna não anulou por completo o efeito agudo '
 f'acumulado, o que explica a deriva descendente do vigor e a subida líquida da fadiga ao longo da semana.')

H('5.8 Relações entre as dimensões (Spearman)',12,before=6)
P(f'As correlações de Spearman entre as dimensões (Tabela {_TN[0]+1}) mostraram que as dimensões negativas se associam '
 f'entre si, com destaque para os pares depressão–raiva e tensão–confusão, e que a fadiga se relaciona com a depressão e '
 f'a raiva. O vigor manteve-se relativamente independente das demais dimensões, o que reforça a sua leitura como um polo '
 f'próprio do eixo energia–fadiga.')
sigpairs=[x for x in STAT['pairs'] if x['p']<0.05]
def srow(x): return [f"{x['a']} × {x['b']}",c2(f"{x['rho']:+.2f}"),pvt(x['p'])]
table('Correlações de Spearman significativas entre as dimensões do BRUMS (n = %d atletas).'%sm['n'],
 ['Par de dimensões','ρ','p'],[srow(x) for x in sigpairs],
 note='ρ = coeficiente de correlação de Spearman. Apresentam-se apenas os pares com p < 0,05.',fs=9)
P(f'A relação entre o vigor e a fadiga, além de negativa, tornou-se mais estreita à medida que a carga se acumulou '
 f'(Figura {_FN[0]+1}): a regressão por fase do microciclo mostrou um acoplamento fraco na fase inicial (r = -0,35) e '
 f'mais forte na fase de acúmulo (r = -0,51), o que sugere que os dois polos do eixo energia–fadiga passam a variar de '
 f'forma mais solidária sob fadiga acumulada.')
figure(f'{FG}/tec_regr.png','Relação entre vigor e fadiga por fase do microciclo, com a reta de regressão de cada grupo (início: dias 1 a 4; acúmulo: dias 5 a 7).',w=12.5)

H('5.9 Estrutura dimensional (análise de componentes principais)',12,before=6)
P(f'Uma análise exploratória de componentes principais resumiu a estrutura das seis dimensões. Os dois primeiros '
 f'componentes explicaram {pc12}% da variância ({pc1}% no primeiro e {pc2}% no segundo). O círculo de correlação '
 f'(Figura {_FN[0]+1}) mostra que as dimensões negativas (depressão, raiva, confusão e fadiga) se projetam juntas no '
 f'lado positivo do primeiro componente, que funciona como um eixo geral de perturbação, ao passo que o vigor aponta em '
 f'sentido oposto. Essa oposição reforça, por via independente, a centralidade do eixo energia–fadiga.')
figure(f'{FG}/pca_circulo.png','Círculo de correlação da análise de componentes principais das seis dimensões do BRUMS. A espessura de cada seta é proporcional à contribuição da variável ao plano dos dois primeiros componentes.',w=12.0)

# ===== 5 DISCUSSÃO =====
H('6 DISCUSSÃO')
P('O conjunto dos resultados sustenta uma tese central: em um microciclo pré-competitivo de handebol de elite, o estado '
 'de humor comporta-se como um sistema dinâmico de um único eixo dominante, o eixo energia–fadiga, que se deteriora de '
 'modo ordenado sob a carga e cujo curso pode ser lido, datado e quantificado. Cada camada de análise, da descrição '
 'simples à modelagem das trajetórias, converge para esse mesmo núcleo por caminhos independentes, o que confere '
 'robustez à interpretação e afasta a leitura de achados isolados. A discussão a seguir organiza essa convergência em '
 'quatro eixos: o significado do padrão observado, a sua tradução por perfis de humor, a contribuição metodológica das '
 'análises de trajetória e a sua aplicação prática, com atenção reflexiva aos limites psicométricos do instrumento e à '
 'natureza dos dados.')
H('6.1 Migração do humor e o eixo energia–fadiga',12,before=6)
P('Os resultados descrevem, em um microciclo pré-competitivo de handebol de elite, a migração do humor da prontidão para '
 'a fadiga funcional. A queda do vigor e a elevação da fadiga entre o primeiro e o último dia alcançaram efeito grande e '
 'foram confirmadas pela análise multivariada, o que situa o eixo energia–fadiga como o principal responsável pela '
 'mudança do estado do atleta. Esse padrão reproduz, dentro de um microciclo de handebol, o comportamento já documentado '
 'em outras modalidades sob acúmulo de carga e sob treino intensificado (PIERCE, 2002; FERREIRA et al., 2026; THORPE et '
 'al., 2017).')
P('À tendência semanal somou-se uma resposta aguda coerente com o esforço. Entre o momento pré e o pós-treino, o vigor '
 'caiu e a fadiga e a PTH subiram, com magnitude de pequena a moderada, o que acrescenta uma camada de variação dentro '
 'do dia à deterioração observada entre dias. Esse comportamento converge com o de outros esportes coletivos, nos quais '
 'o humor responde de forma imediata ao evento esportivo, e reforça o valor de coletas repetidas dentro do dia, e não '
 'apenas entre dias (DO NASCIMENTO et al., 2026).')
P('A perturbação total do humor comportou-se como um integrador desse eixo. O vigor e a fadiga associaram-se de forma '
 'forte à PTH e explicaram a maior parte da sua variância, ao passo que as dimensões negativas, concentradas em valores '
 'baixos, contribuíram pouco para a variação. A consistência interna adequada do vigor e da fadiga e a sua correlação '
 'intraclasse de moderada a boa convergem para a mesma conclusão, a saber, que o núcleo do sinal do monitoramento reside '
 'no par energia–fadiga (LOCHBAUM et al., 2021; ROHLFS et al., 2023).')
P(f'Do ponto de vista psicométrico, a concentração das dimensões negativas junto ao limite inferior da escala é '
 f'informativa, e não um mero artefato. O forte efeito de piso, com percentuais de escores nulos que chegaram a '
 f'{num(ADV["snr"]["Confusao"]["floor0"],0)}% na confusão e a {num(ADV["snr"]["Depressao"]["floor0"],0)}% na depressão, '
 f'reduz a variância disponível dessas dimensões, comprime a sua distribuição e rebaixa a sua razão sinal/ruído, o que '
 f'limita a capacidade de detectar mudança e ajuda a explicar a sua menor confiabilidade entre medidas repetidas. Em '
 f'atletas saudáveis e bem ajustados, esse piso é esperado e, longe de invalidar o instrumento, delimita o seu uso: as '
 f'dimensões negativas funcionam como sentinelas de estados clínicos raros, ao passo que o vigor e a fadiga, livres do '
 f'piso e com sinal limpo, carregam a informação útil ao monitoramento da carga (TERRY et al., 2022; ROHLFS et al., '
 f'2023). Esse raciocínio recomenda cautela ao interpretar médias e testes das dimensões de piso e favorece, para elas, '
 f'a leitura por prevalência e por perfis em vez da leitura por escore contínuo.')
H('6.2 Perfis de humor e sobre-esforço funcional',12,before=6)
P('A leitura por perfis acrescentou clareza a essa descrição. O deslocamento do perfil iceberg para a barbatana de '
 'tubarão traduz, em uma única imagem, a acumulação da carga sem os sinais de comprometimento da saúde mental, uma vez '
 'que os perfis de maior risco, como o iceberg invertido e o submerso, não se instalaram. Quando comparados à '
 'distribuição de referência de uma grande amostra brasileira e aos padrões descritos em análises de agrupamento, os '
 'nossos dados partem de um predomínio semelhante de iceberg e elevam a barbatana de tubarão apenas ao fim da semana '
 '(PARSONS-SMITH; TERRY; MACHIN, 2017; LUOJUMÄKI et al., 2026; DE MIRANDA ROHLFS et al., 2024).')
P('Essa reconfiguração tem valor comunicativo direto. O perfil condensa as seis dimensões em uma imagem única, o que '
 'aproxima o dado psicométrico da linguagem da comissão técnica e favorece a decisão de treino e de recuperação. A mesma '
 'leitura por perfis tem sido proposta como recurso de rastreio da saúde mental no esporte, o que amplia o alcance do '
 'monitoramento para além do desempenho (TERRY et al., 2021; HAN; PARSONS-SMITH; TERRY, 2020; LEW et al., 2023).')
H('6.3 Contribuição metodológica: derivadas, ajuste polinomial e cruzamentos',12,before=6)
P('Além de comparar médias, este estudo mapeou a forma temporal exata das trajetórias. A suavização das séries e o '
 'cálculo da segunda derivada localizaram, em cada dimensão do eixo energia–fadiga, um único ponto de inflexão situado '
 'na metade da semana, no qual a concavidade muda e a taxa de variação atinge o seu limite. Esse marcador objetivo '
 'demarca a transição entre a fase inicial de prontidão e a fase de acúmulo de carga e oferece uma leitura que a simples '
 'comparação entre o primeiro e o último dia não revela.')
P(f'O ajuste polinomial de grau três acrescentou uma síntese paramétrica a essa análise. O modelo reproduziu a inflexão '
 f'de forma analítica, na raiz da segunda derivada, e resumiu cada trajetória por uma equação, com ajuste forte sobre as '
 f'médias diárias (R² de {num(POLY["V"]["d7"]["r2"],2)} a {num(POLY["F"]["d7"]["r2"],2)} no eixo energia–fadiga). As '
 f'duas resoluções, sobre sete médias diárias e sobre quatorze pontos pré e pós-treino, apontaram inflexões quase '
 f'idênticas, o que indica que o achado da transição na metade da semana não depende do método nem da granularidade da '
 f'medida. A convergência entre a spline não paramétrica e o polinômio paramétrico reforça a robustez do resultado.')
P('A localização dos cruzamentos exatos representa a contribuição mais original desta abordagem. O vigor superou a '
 'fadiga por ampla margem no início e foi alcançado por ela ainda na primeira metade da semana, mas apenas ao fim do '
 'microciclo a fadiga afastou-se de forma definitiva acima do vigor, seguida pela PTH, que ultrapassou primeiro o vigor '
 'e depois a fadiga. A datação desses cruzamentos converte a ideia qualitativa de inversão do eixo energia–fadiga em um '
 'evento com dia e escore definidos, o que fornece um candidato a marcador temporal do início do sobre-esforço '
 'funcional, coerente com o continuum descrito entre sobre-esforço e overtraining (ROETE et al., 2021; MĂNESCU et al., '
 '2026; LA TORRE et al., 2023).')
P('O rigor deste conjunto de análises apoia-se na convergência de métodos independentes, e não em um único '
 'procedimento. A escolha de testes não paramétricos, adotada após a verificação da normalidade, protege as comparações '
 'contra os desvios de distribuição observados nas dimensões negativas, e o relato sistemático do tamanho e da magnitude '
 'do efeito, ao lado dos valores de p, evita a leitura de significância como se fosse relevância. Na análise de '
 'trajetória, a spline não paramétrica e o ajuste polinomial paramétrico localizaram a mesma inflexão, e as duas '
 'resoluções do polinômio, sobre médias diárias e sobre pontos pré e pós-treino, apontaram valores quase idênticos, o '
 'que caracteriza uma triangulação metodológica. Os limiares de mudança derivados do coeficiente de correlação '
 'intraclasse delimitam ainda o que pode ser lido no plano individual e o que só se sustenta como tendência de grupo, o '
 'que mantém a inferência dentro dos limites do delineamento.')
P(f'A decomposição entre sinal e ruído acrescenta uma leitura reflexiva sobre a qualidade da informação de cada '
 f'dimensão. Ao separar a tendência lenta (o sinal) das flutuações de alta frequência (o ruído de medida e a '
 f'variabilidade biológica do dia a dia), a suavização mostrou que o vigor e a fadiga concentram a maior razão '
 f'sinal/ruído (cerca de {num(ADV["snr"]["Vigor"]["snr_amp"],1)} e {num(ADV["snr"]["Fadiga"]["snr_amp"],1)}), enquanto '
 f'as dimensões negativas, comprimidas pelo piso, mal ultrapassam o ruído. Esse resultado converge com a análise '
 f'psicométrica e com a estrutura de componentes principais, e fundamenta, por um terceiro caminho, a decisão de centrar '
 f'o monitoramento no eixo energia–fadiga, no qual o sinal é forte o bastante para sustentar inferências sobre a carga '
 f'(SAW; MAIN; GASTIN, 2016; HELWIG et al., 2023).')
P(f'A análise sequencial das treze coletas revelou, por fim, a microestrutura da fadiga funcional. A decomposição de '
 f'cada dia em um efeito agudo (pré→pós) e um efeito de recuperação (pós→pré do dia seguinte) mostrou um sistema que '
 f'oscila em dente de serra: a sessão eleva a perturbação e reduz o vigor, e o intervalo entre sessões devolve parte do '
 f'estado, com destaque para a carga elevada do Dia 6 (PTH com dz = {num(ADV["trans"][9]["vars"]["TMD"]["dz"],2)}) e '
 f'para a recuperação entre o Dia 2 e o Dia 3 (PTH com dz = {num(ADV["trans"][2]["vars"]["TMD"]["dz"],2)}). O ponto '
 f'central é que a recuperação foi apenas parcial: como o retorno noturno não anulou o efeito agudo, o saldo acumulou-se '
 f'e produziu a deriva descendente do vigor ao longo da semana. Essa leitura dá conteúdo operacional ao conceito de '
 f'sobre-esforço funcional, no qual a fadiga é induzida de forma planejada e monitorada justamente pela relação entre a '
 f'carga aguda e a recuperação entre sessões, e sugere que o desequilíbrio persistente entre esses dois termos seria o '
 f'sinal de alerta para a transição ao sobre-esforço não funcional (KELLMANN et al., 2018; ROETE et al., 2021; '
 f'MĂNESCU et al., 2026).')
H('6.4 Aplicação prática e perspectivas',12,before=6)
P('As características do handebol ajudam a explicar esse comportamento e orientam a aplicação. A modalidade impõe esforço '
 'intermitente de alta intensidade, com sprints curtos, mudanças de direção, saltos e contato, o que gera elevada '
 'demanda neuromuscular e psicofisiológica ao longo da semana e sustenta tanto a tendência de acúmulo quanto a resposta '
 'aguda entre o pré e o pós-treino (KARCHER; BUCHHEIT, 2014; GARCÍA-SÁNCHEZ et al., 2023; CARTON-LLORENTE et al., 2023). '
 'Na prática, o par vigor–fadiga concentra o sinal útil, e o cruzamento entre as suas curvas oferece um alerta simples e '
 'visual, capaz de complementar o monitoramento de carga interna e de fadiga já descrito na modalidade (STRUZIK; '
 'NADOBNIK; STĘPIEŃ-SŁODKOWSKA, 2026; STAIANO et al., 2025).')
P('O acompanhamento subjetivo do humor deve ser lido como parte de um monitoramento multidomínio, e não como medida '
 'isolada. A sua sensibilidade, o baixo custo e o caráter não invasivo tornam-no complementar aos marcadores '
 'fisiológicos, endócrinos e de bem-estar praticados em esportes coletivos (RATZ-SULYOK et al., 2026; BIRD et al., 2025; '
 'HELWIG et al., 2023). Como agenda futura, recomenda-se integrar medidas objetivas de carga a este delineamento, '
 'ampliar a amostra e testar se a datação do cruzamento do eixo energia–fadiga antecipa desfechos de fadiga e de '
 'desempenho, de modo a validar o marcador temporal aqui proposto (KELLMANN et al., 2018; SAW; MAIN; GASTIN, 2016).')
H('6.5 Limitações',12,before=6)
P('O estudo tem limitações, entre as quais a amostra de um único clube, o recorte de um microciclo e a ausência de '
 'medidas objetivas de carga neste recorte, o que restringe a generalização e impede inferências de causa. O pequeno '
 'número de observações distribuído por seis perfis reduz a potência do teste categórico, e o erro de medida de uma '
 'leitura isolada recomenda cautela na interpretação individual, de modo que a inferência se apoia nas tendências de '
 'grupo. Ainda assim, a descrição oferece um retrato direto da dinâmica do humor no handebol de elite e sustenta a '
 'recomendação prática de acompanhar o eixo energia–fadiga ao longo da semana, integrado a um monitoramento multidomínio '
 'do estado do atleta (OSTAPIUK-KAROLCZUK et al., 2025).')

# ===== REFERÊNCIAS =====
H('REFERÊNCIAS',before=6)
refs=[
 'BIRD, S. P. et al. Wellness, mood, sleep, and performance in a women’s national basketball team during international competition. Journal of Human Kinetics, v. 96, p. 163–175, 2025. DOI: 10.5114/jhk/200117.',
 'CARTON-LLORENTE, A. et al. Worst-case scenario analysis of physical demands in elite men handball players by playing position through big data analytics. Biology of Sport, v. 40, n. 4, p. 1219–1227, 2023. DOI: 10.5114/biolsport.2023.126665.',
 'DE MIRANDA ROHLFS, I. C. P. et al. Prevalence of specific mood profile clusters among elite and youth athletes at a Brazilian sports club. Sports, v. 12, n. 7, 195, 2024. DOI: 10.3390/sports12070195.',
 'DO NASCIMENTO, M. H. et al. Acute psychological responses to official match outcomes in male youth volleyball: an observational repeated-measures study within a single national-level team. Frontiers in Psychology, v. 17, 1826372, 2026. DOI: 10.3389/fpsyg.2026.1826372.',
 'FERREIRA, A. B. M. et al. Impact of sleep restriction and intensified training on mucosal immunity and psychological responses in young soccer players. Journal of Strength and Conditioning Research, v. 40, n. 7, p. e703–e713, 2026. DOI: 10.1519/JSC.0000000000005416.',
 'GARCÍA-SÁNCHEZ, C. et al. Physical demands during official competitions in elite handball: a systematic review. International Journal of Environmental Research and Public Health, v. 20, n. 4, 3353, 2023. DOI: 10.3390/ijerph20043353.',
 'HAN, C.; PARSONS-SMITH, R. L.; TERRY, P. C. Mood profiling in Singapore: cross-cultural validation and potential applications of mood profile clusters. Frontiers in Psychology, v. 11, 665, 2020. DOI: 10.3389/fpsyg.2020.00665.',
 'HELWIG, J. et al. Relationships between external, wearable sensor-based, and internal parameters: a systematic review. Sensors, v. 23, n. 2, 827, 2023. DOI: 10.3390/s23020827.',
 'KARCHER, C.; BUCHHEIT, M. On-court demands of elite handball, with special reference to playing positions. Sports Medicine, v. 44, n. 6, p. 797–814, 2014. DOI: 10.1007/s40279-014-0164-z.',
 'KELLMANN, M. et al. Recovery and performance in sport: consensus statement. International Journal of Sports Physiology and Performance, v. 13, n. 2, p. 240–245, 2018. DOI: 10.1123/ijspp.2017-0759.',
 'LA TORRE, M. E. et al. The potential role of nutrition in overtraining syndrome: a narrative review. Nutrients, v. 15, n. 23, 4916, 2023. DOI: 10.3390/nu15234916.',
 'LEW, P. C. F. et al. Cross-cultural validation of the Malaysian Mood Scale and tests of between-group mood differences. International Journal of Environmental Research and Public Health, v. 20, n. 4, 3348, 2023. DOI: 10.3390/ijerph20043348.',
 'LOCHBAUM, M. et al. The Profile of Mood States and athletic performance: a meta-analysis of published studies. European Journal of Investigation in Health, Psychology and Education, v. 11, n. 1, p. 50–70, 2021. DOI: 10.3390/ejihpe11010005.',
 'LUOJUMÄKI, R. J. et al. Exploring mood profile clusters across physical activity level, gender and age in a Finnish population. European Journal of Sport Science, v. 26, n. 2, e70131, 2026. DOI: 10.1002/ejsc.70131.',
 'MĂNESCU, D. C. et al. Molecular biomarkers of training responses: a systems framework for exercise adaptation and athlete monitoring. International Journal of Molecular Sciences, v. 27, n. 8, 3601, 2026. DOI: 10.3390/ijms27083601.',
 'MORGAN, W. P. Selected psychological factors limiting performance: a mental health model. In: CLARKE, D. H.; ECKERT, H. M. (Ed.). Limits of human performance. Champaign: Human Kinetics, 1985. p. 70–80.',
 'OSTAPIUK-KAROLCZUK, J. et al. Biochemical and psychological markers of fatigue and recovery in mixed martial arts athletes during strength and conditioning training. Scientific Reports, v. 15, n. 1, 24234, 2025. DOI: 10.1038/s41598-025-09719-z.',
 'PARSONS-SMITH, R. L.; TERRY, P. C.; MACHIN, M. A. Identification and description of novel mood profile clusters. Frontiers in Psychology, v. 8, 1958, 2017. DOI: 10.3389/fpsyg.2017.01958.',
 'PIERCE, E. F. Relationship between training volume and mood states in competitive swimmers during a 24-week season. Perceptual and Motor Skills, v. 94, n. 3, p. 1009–1012, 2002. DOI: 10.2466/pms.2002.94.3.1009.',
 'RATZ-SULYOK, F. Z. et al. Associations between endocrine status and stress, mood and psychosomatic status in elite handball players. Sports, v. 14, n. 7, 289, 2026. DOI: 10.3390/sports14070289.',
 'ROETE, A. J. et al. A systematic review on markers of functional overreaching in endurance athletes. International Journal of Sports Physiology and Performance, v. 16, n. 8, p. 1065–1073, 2021. DOI: 10.1123/ijspp.2021-0024.',
 'ROHLFS, I. C. P. M. et al. A Escala de Humor de Brunel (Brums): instrumento para detecção precoce da síndrome do excesso de treinamento. Revista Brasileira de Medicina do Esporte, v. 14, n. 3, p. 176–181, 2008.',
 'ROHLFS, I. C. P. M. et al. Psychometric characteristics of the Brazil Mood Scale among youth and elite athletes using two response time frames. Sports, v. 11, n. 12, 244, 2023. DOI: 10.3390/sports11120244.',
 'SAW, A. E.; MAIN, L. C.; GASTIN, P. B. Monitoring the athlete training response: subjective self-reported measures trump commonly used objective measures: a systematic review. British Journal of Sports Medicine, v. 50, n. 5, p. 281–291, 2016. DOI: 10.1136/bjsports-2015-094758.',
 'STAIANO, W. et al. Overcoming mental fatigue through mindfulness: improving physical and cognitive performance in elite handball players. Journal of Science and Medicine in Sport, v. 29, n. 1, p. 91–99, 2025. DOI: 10.1016/j.jsams.2025.08.004.',
 'STRUZIK, A.; NADOBNIK, J.; STĘPIEŃ-SŁODKOWSKA, M. TRIMP and session-RPE monitoring in elite women’s handball: a full-season descriptive analysis. Scientific Reports, v. 16, n. 1, 2026. DOI: 10.1038/s41598-026-53134-x.',
 'TERRY, P. C.; LANE, A. M.; FOGARTY, G. J. Construct validity of the Profile of Mood States, Adolescents for use with adults. Psychology of Sport and Exercise, v. 4, n. 2, p. 125–139, 2003. DOI: 10.1016/S1469-0292(02)00035-8.',
 'TERRY, P. C. et al. Mood profiling for sustainable mental health among athletes. Sustainability, v. 13, n. 11, 6116, 2021. DOI: 10.3390/su13116116.',
 'TERRY, P. C. et al. Validation of a Lithuanian-language version of the Brunel Mood Scale: the BRUMS-LTU. International Journal of Environmental Research and Public Health, v. 19, n. 8, 4867, 2022. DOI: 10.3390/ijerph19084867.',
 'THORPE, R. T. et al. Monitoring fatigue status in elite team-sport athletes: implications for practice. International Journal of Sports Physiology and Performance, v. 12, n. S2, p. S227–S234, 2017. DOI: 10.1123/ijspp.2016-0434.']
for rf in refs:
    p=doc.add_paragraph(); r=p.add_run(rf); r.font.size=Pt(11); p.paragraph_format.line_spacing=1.3; p.paragraph_format.space_after=Pt(4); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

OUTP='/home/user/mdlucca/Artigos/Paper1_Humor_completo.docx'
doc.save(OUTP); print('SAVED',OUTP,'| Tabelas',_TN[0],'Figuras corpo',_FN[0])
