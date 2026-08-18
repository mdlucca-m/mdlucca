# -*- coding: utf-8 -*-
import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
R=json.load(open('brums_desc2.json')); STAT=json.load(open('brums_stats3.json')); S4=json.load(open('brums_stats4.json'))
MS=json.load(open('model_stats.json')); MV=json.load(open('manova.json')); PHJ=json.load(open('posthoc.json'))
PV=json.load(open('pv_stats.json')); LIM=json.load(open('tcar_limiar.json')); TCD=json.load(open('tcar_desc.json'))
TCV=json.load(open('tcar_curvas.json'))
LOG=json.load(open('logistica.json'))
L2=json.load(open('limiar2.json'))
MOD=json.load(open('moduladores.json'))
RSA=json.load(open('rsa_stats.json'))
COM=json.load(open('commonality.json')); SPC=json.load(open('speccurve.json')); MDC=json.load(open('mdc.json')); PSY=json.load(open('psychometric.json')); BP=json.load(open('bayes_phys.json'))
CRS=json.load(open('cross.json')); PK=json.load(open('peaks.json')); SS=json.load(open('sono_stress.json')); AL=json.load(open('alom.json')); DV=json.load(open('deriv.json')); SM=json.load(open('smooth.json')); IND=json.load(open('indiv.json')); AUD=json.load(open('audit.json')); CP=json.load(open('cross_pt.json'))
FG='/home/user/mdlucca/Artigos/figuras'
def c2(s): return str(s).replace('.',',')
doc=Document()
stl=doc.styles['Normal']; stl.font.name='Times New Roman'; stl.font.size=Pt(12)
stl.element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')
stl.paragraph_format.line_spacing=1.5; stl.paragraph_format.space_after=Pt(0)
sec=doc.sections[0]; sec.top_margin=Cm(3); sec.left_margin=Cm(3); sec.bottom_margin=Cm(2); sec.right_margin=Cm(2)
_TN=[0]; _FN=[0]
def P(t='',just=True,size=12,after=6,bold=False,ind=True):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(size); r.bold=bold
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(after)
    if just: p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if just and ind: p.paragraph_format.first_line_indent=Cm(1.25)
    return p
def RUN(pairs,after=6,ind=True):
    p=doc.add_paragraph(); p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(after)
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if ind: p.paragraph_format.first_line_indent=Cm(1.25)
    for txt,bold in pairs:
        r=p.add_run(txt); r.bold=bold; r.font.size=Pt(12)
    return p
def H(t,size=12,before=10):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(size)
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.5
def _bd(c):
    tcPr=c._tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for e_ in ['top','bottom','left','right']:
        el=OxmlElement(f'w:{e_}'); on=e_ in ('top','bottom')
        el.set(qn('w:val'),'single' if on else 'nil'); el.set(qn('w:sz'),'6'); el.set(qn('w:color'),'000000'); b.append(el)
    tcPr.append(b)
def table(cap,header,rows,fonte='Fonte: dados da pesquisa (2026).',fs=9,note=None):
    _TN[0]+=1
    p=doc.add_paragraph(); r=p.add_run('Tabela %d – %s'%(_TN[0],cap)); r.font.size=Pt(11)
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
    t=doc.add_table(rows=1,cols=len(header)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,htx in enumerate(header):
        cc=t.rows[0].cells[i]; cc.text=''; rr=cc.paragraphs[0].add_run(htx); rr.bold=True; rr.font.size=Pt(fs); rr.font.name='Times New Roman'
        cc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; _bd(cc)
    for row in rows:
        cs=t.add_row().cells
        for i,val in enumerate(row):
            cs[i].text=''; rr=cs[i].paragraphs[0].add_run(str(val)); rr.font.size=Pt(fs); rr.font.name='Times New Roman'
            cs[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if i>0 else WD_ALIGN_PARAGRAPH.LEFT; _bd(cs[i])
    if note:
        pn=doc.add_paragraph(); rn=pn.add_run(note); rn.font.size=Pt(8.5); rn.italic=True; pn.paragraph_format.space_after=Pt(0)
    pf=doc.add_paragraph(); rf=pf.add_run(fonte); rf.font.size=Pt(9); pf.paragraph_format.space_after=Pt(6)
    return _TN[0]
def figure(path,cap,w=15.0):
    _FN[0]+=1
    pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; pp.add_run().add_picture(path,width=Cm(w)); pp.paragraph_format.space_before=Pt(6)
    pc=doc.add_paragraph(); pc.alignment=WD_ALIGN_PARAGRAPH.CENTER; rc=pc.add_run('Figura %d – %s'%(_FN[0],cap)); rc.font.size=Pt(11)
    pf=doc.add_paragraph(); pf.alignment=WD_ALIGN_PARAGRAPH.CENTER; rf=pf.add_run('Fonte: elaboração dos autores (2026).'); rf.font.size=Pt(9); pf.paragraph_format.space_after=Pt(6)
    return _FN[0]
def pstr(p): return '< 0,001' if p<0.001 else '= '+c2('%.3f'%p)
mvv=lambda tab,k,f: next(x[f] for x in MV[tab]['rows'] if x['k']==k)
pr=R['prepos']; d17=R['d1d7']; prof=R['profiles']; sm=R['sample']; desc=R['desc']; PREV=MS['prev']
SH=STAT['shapiro']; IC=STAT['icc']; LP=LIM['LIM']['PVini']; FR=S4['friedman']; SENS=S4['sens']
# ordem canônica POMS/BRUMS: Tensão–Depressão–Raiva–Vigor–Fadiga–Confusão–PTH
ORD=[('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Vigor','Vigor'),('Fadiga','Fadiga'),('Confusao','Confusão'),('TMD','PTH')]

# ===== TÍTULO =====
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('CORRELATOS FISIOLÓGICOS E DE BEM-ESTAR DA RESPOSTA DE HUMOR NUM MICROCICLO PRÉ-COMPETITIVO DE HANDEBOL DE ELITE: APTIDÃO AERÓBIA, CAPACIDADE ANAERÓBIA, SONO E ESTRESSE')
r.bold=True; r.font.size=Pt(13); p.paragraph_format.space_after=Pt(10)

# ===== RESUMO =====
H('RESUMO',before=2)
RUN([('Objetivo: ',True),('examinar os correlatos fisiológicos e de bem-estar da resposta de humor de handebolistas de '
 'elite num microciclo pré-competitivo — aptidão aeróbia intermitente (pico de velocidade do Teste de Carminatti), '
 'capacidade anaeróbia (sprints repetidos), composição corporal, sonolência e estresse percebido. ',False),
 ('Método: ',True),('%d atletas do sexo masculino responderam ao BRUMS-24 ao longo de sete dias (uma coleta de linha '
 'de base no primeiro dia e duas coletas diárias — pré e pós-treino — nos seis dias de treino; %d observações) e foram '
 'avaliados por antropometria, salto vertical (CMJ), teste de sprints repetidos (Baker), Teste de Carminatti (T-CAR), '
 'Escala de Sonolência de Epworth e Escala de Estresse Percebido (PSS-14). Empregaram-se correlação de Spearman, '
 'regressão, limiares (índice de Youden e tercis), ajuste alométrico, correlação parcial e análise de moduladores da '
 'magnitude dos efeitos (coeficiente de determinação R²). '%(sm['n'],sm['n_obs']),False),
 ('Resultados: ',True),('a deterioração do humor concentrou-se no eixo energia–fadiga (vigor d = %s; fadiga d = %s; '
 'MANOVA Wilks λ = %s; p %s). A aptidão aeróbia intermitente foi o correlato fisiológico mais forte: maior pico de '
 'velocidade associou-se a menos fadiga física (ρ = %s) e a mais vigor (ρ = %s), com limiar de aproximadamente %s km/h '
 'e faixas de risco definidas por dois cortes. A capacidade anaeróbia (sprints repetidos) não se associou de forma '
 'significativa à fadiga (melhor tempo × fadiga física ρ = %s), e o vínculo aparente desapareceu ao controlar a massa '
 'corporal (ρ parcial = %s), o que evidencia especificidade aeróbia. A sonolência elevou-se ao longo da semana e '
 'acompanhou a fadiga e a perturbação total do humor, ao passo que o estresse percebido se manteve estável. A magnitude '
 'dos efeitos agudo e crônico foi amplamente independente do perfil individual (coeficientes de determinação baixos). '%(
   c2('%+.2f'%mvv('d1d7','Vigor','d')),c2('%+.2f'%mvv('d1d7','Fadiga','d')),c2('%.3f'%MV['d1d7']['wilks']),pstr(MV['d1d7']['p_mv']),
   c2('%+.2f'%PV['pv']['wk_FadFisica']['TCAR1']['rho']),c2('%+.2f'%PV['pv']['wk_Vigor']['TCAR1']['rho']),c2('%.1f'%LP['thr']),
   c2('%+.2f'%RSA['assoc']['BkMel']['FadFisica']['raw_r']),c2('%+.2f'%RSA['assoc']['BkMel']['FadFisica']['par_r'])),False),
 ('Conclusão: ',True),('entre os parâmetros examinados, a aptidão aeróbia intermitente destacou-se como o correlato '
 'fisiológico genuíno da fadiga do humor, ao passo que a capacidade anaeróbia e o porte corporal não a explicaram; a '
 'sonolência acrescentou uma leitura de recuperação e o estresse percebido delimitou a especificidade da resposta à '
 'carga de treino. Por captar uma dimensão subjetiva que os marcadores periféricos nem sempre revelam, o humor atua '
 'como sentinela complementar, e não redundante, à leitura fisiológica; recomenda-se, assim, um monitoramento '
 'multidomínio centrado no eixo energia–fadiga, no qual a aptidão aeróbia oferece a referência para individualizar a '
 'carga.',False)],after=6)
P('Palavras-chave: aptidão aeróbia; Teste de Carminatti; humor; BRUMS; handebol; monitoramento do atleta.',size=11,after=8,ind=False)

# ===== 1 INTRODUÇÃO =====
H('1 INTRODUÇÃO')
H('1.1 Monitoramento do humor no esporte de rendimento',12,before=6)
P('O acompanhamento do estado psicológico dos atletas consolidou-se como parte essencial da gestão do treinamento no '
 'esporte de rendimento. Instrumentos de autorrelato do humor são práticos, econômicos e sensíveis às variações da carga '
 'de treino, com utilidade preditiva para o bem-estar e o desempenho esportivo (SAW; MAIN; GASTIN, 2016; LOCHBAUM et al., '
 '2021), razão pela qual documentos de consenso recomendam seu uso rotineiro para monitorar a fadiga e orientar decisões '
 'de treino e recuperação (KELLMANN et al., 2018). O valor desse autorrelato não se esgota na conveniência: sob treino '
 'intensificado, a perturbação do humor pode persistir mesmo depois de marcadores bioquímicos de fadiga terem regredido, '
 'assincronia entre a recuperação psicológica e a fisiológica que fundamenta a integração das duas leituras num '
 'monitoramento multidomínio (OSTAPIUK-KAROLCZUK et al., 2025).')
P('Entre esses instrumentos, a Escala de Humor de Brunel (BRUMS), versão abreviada e adaptada do Profile of Mood States '
 '(POMS), destaca-se pela rapidez de aplicação e pela solidez psicométrica, com sucessivas validações transculturais, '
 'entre elas a versão brasileira (BRAMS) (TERRY; LANE; FOGARTY, 2003; ROHLFS et al., 2008, 2023). Sua aplicação tem-se '
 'mostrado particularmente informativa em modalidades esportivas coletivas e no esporte de rendimento, contextos de '
 'elevada demanda física, emocional e interpessoal: a BRUMS tem sido empregada para monitorar o humor de atletas de '
 'elite ao longo de períodos competitivos e sua relação com o sono, o desempenho, os resultados de partida e o risco de '
 'lesão (ANDRADE et al., 2016; BRANDT; BEVILACQUA; ANDRADE, 2017; ANDRADE et al., 2020; DE MIRANDA ROHLFS et al., 2025; '
 'DO NASCIMENTO et al., 2026). Nas modalidades coletivas, em particular, a escala acompanhou o humor de seleções e '
 'equipes de elite ao longo de competições internacionais no basquetebol (BIRD et al., 2025), subsidiou a triagem do '
 'excesso de treinamento e da recuperação em atletas dessa mesma modalidade (BATTAGLINI et al., 2022) e captou a queda '
 'de vigor e a elevação de fadiga sob treino intensificado no futebol (FERREIRA et al., 2026); firma-se, assim, como um '
 'indicador de baixo custo e alta sensibilidade para a triagem do bem-estar psicológico e do risco de saúde mental em '
 'contextos de rendimento.')
P('Na prática, a escala tem sido empregada sobretudo em dois momentos: como um retrato pontual em torno da competição, '
 'associado à probabilidade de vitória e ao estado pré-jogo (BRANDT et al., 2019; DO NASCIMENTO et al., 2026), e em '
 'comparações entre fases de treino separadas por semanas ou meses, muitas vezes em paralelo a marcadores hormonais e '
 'de sono (ROUVEIX et al., 2006; ANDRADE et al., 2019). Permanece menos explorada, contudo, a dinâmica fina do humor no '
 'interior de um único microciclo — aquela que distingue a oscilação aguda de cada sessão do acúmulo ao longo dos dias '
 '—, sobretudo quando acompanhada de forma contínua e articulada à aptidão física em uma modalidade coletiva.')
H('1.2 Perfis de humor e sua aplicação em atletas de elite',12,before=6)
P('Para além dos escores isolados de cada subescala, a leitura do humor evoluiu para a identificação de perfis '
 'prototípicos. Morgan (1985) descreveu o clássico “perfil iceberg” — vigor elevado sobre dimensões negativas baixas — '
 'como assinatura de prontidão e de saúde mental positiva. Mais recentemente, análises de agrupamento (cluster) em '
 'grandes amostras formalizaram seis perfis de humor — iceberg, superfície, submerso, barbatana de tubarão, iceberg '
 'invertido e Everest invertido —, dos quais os três últimos se associam a maior risco à saúde mental e à subperformance '
 '(PARSONS-SMITH; TERRY; MACHIN, 2017; HAN; PARSONS-SMITH; TERRY, 2020). Esses perfis têm sido replicados em diferentes '
 'culturas e aplicados ao rastreamento da prontidão e do bem-estar de atletas de elite, inclusive no contexto brasileiro: '
 'em uma grande amostra de um clube de rendimento do Rio de Janeiro (898 atletas), o perfil iceberg foi o mais prevalente '
 'em avaliação momentânea, ao passo que os perfis de risco foram os menos frequentes (DE MIRANDA ROHLFS et al., 2024). A '
 'abordagem por perfis oferece uma leitura integrada e visual do estado psicológico, sensível às variações de carga e '
 'útil para identificar precocemente atletas em deterioração, com valor documentado para a saúde mental sustentável e '
 'para desfechos como lesão (TERRY et al., 2021; DE MIRANDA ROHLFS et al., 2025).')
P('Dois desses perfis são especialmente informativos para a leitura de um microciclo de carga. O iceberg — vigor '
 'elevado sobrepondo-se a dimensões negativas baixas — é a assinatura da prontidão e do bem-estar e associa-se a melhor '
 'desempenho esportivo; em meta-análise de estudos com atletas competitivos, a perturbação total do humor predisse o '
 'rendimento (efeito médio; LOCHBAUM et al., 2021). A barbatana de tubarão (shark fin), por sua vez, descreve o atleta '
 'ainda energizado, porém com a fadiga já elevada acima do vigor — e, em graus variados, com tensão e raiva acentuadas —, '
 'e configura a assinatura afetiva típica do acúmulo de treino. Como o humor é um estado transitório e sensível à carga, '
 'é de se esperar que, ao longo de uma semana intensa de treinamento, a prevalência se desloque do iceberg para a '
 'barbatana de tubarão: um “derretimento do iceberg” que traduz, no plano dos perfis, a deterioração progressiva do eixo '
 'energia–fadiga (MORGAN, 1985; HAN; PARSONS-SMITH; TERRY, 2020). Descrever esse deslocamento ao longo dos dias — e não '
 'apenas em um instantâneo — é o que permite distinguir a fadiga funcional, esperada e reversível, de uma deterioração '
 'que exigiria intervenção, e constitui um dos focos centrais deste estudo.')
H('1.3 Handebol: modalidade coletiva intermitente e de alta intensidade',12,before=6)
P('O handebol de quadra é uma modalidade coletiva de invasão, de caráter marcadamente intermitente e de alta intensidade. '
 'Ao longo da partida, ações máximas e explosivas — sprints curtos, saltos, arremessos, bloqueios, mudanças de direção e '
 'contatos físicos — alternam-se, de forma imprevisível, com períodos de recuperação incompleta, o que exige '
 'simultaneamente potência anaeróbia, capacidade aeróbia intermitente e elevada tolerância à fadiga (KARCHER; BUCHHEIT, '
 '2014; MICHALSIK; AAGAARD, 2015; WAGNER et al., 2014). É justamente a aptidão aeróbia que sustenta a capacidade de '
 'repetir e manter esforços de alta intensidade ao longo de uma partida: uma maior potência aeróbia acelera a '
 'ressíntese de fosfocreatina e a remoção de metabólitos nas pausas incompletas, atenua a queda de desempenho entre '
 'ações sucessivas e retarda a instalação da fadiga, com a preservação da qualidade das ações decisivas nos minutos finais '
 '(KARCHER; BUCHHEIT, 2014; MICHALSIK; AAGAARD, 2015). Por atuar sobre a recuperação entre esforços, e não apenas sobre '
 'a intensidade máxima isolada, a aptidão aeróbia intermitente é determinante da capacidade de sustentar altas '
 'intensidades no jogo intermitente. Avaliá-la, contudo, exige um teste que reproduza esse mesmo padrão de esforço e '
 'pausa: é o que faz o Teste de Carminatti (T-CAR), teste de campo progressivo e intermitente cujo pico de velocidade '
 '(PV) reflete a capacidade aeróbia intermitente e se associa ao desempenho físico e à aptidão em modalidades dessa '
 'natureza (FERNANDES-DA-SILVA et al., 2016). Esse padrão de esforço eleva a carga interna nos microciclos de '
 'acúmulo e repercute no estado afetivo, sobretudo na última semana de pré-temporada, quando a carga que antecede a '
 'competição se concentra. Como a tolerância a essa carga depende da aptidão aeróbia intermitente, é plausível que o PV '
 'do T-CAR module a magnitude da resposta de humor ao acúmulo de treino — hipótese que este estudo examina.')
H('1.4 Objetivos e hipóteses',12,before=6)
P('Embora a resposta de humor à carga de treino esteja bem documentada, permanece pouco esclarecido quais parâmetros '
 'fisiológicos e de bem-estar rastreiam a sua magnitude num microciclo pré-competitivo de handebol de elite — questão '
 'central para um monitoramento multidomínio, que combina marcadores neuromusculares, cardioautonômicos e de '
 'autorrelato (NAUGHTON et al., 2023). A lacuna vai além da simples ausência de dados: quando a aptidão física é '
 'considerada, raramente se testa a especificidade do correlato — isto é, se é a capacidade aeróbia ou a anaeróbia que '
 'de fato acompanha a fadiga afetiva —, e o porte corporal, um confundidor conhecido, poucas vezes é controlado. '
 'Diante dessa lacuna, o objetivo geral consistiu em examinar os correlatos '
 'fisiológicos e de bem-estar da resposta de humor desses atletas na última semana de pré-temporada. De modo '
 'específico, o estudo propôs-se a: (i) caracterizar a amostra e confirmar a deterioração do humor no eixo '
 'energia–fadiga; (ii) analisar a relação entre a aptidão aeróbia intermitente (pico de velocidade do T-CAR) e a '
 'resposta de humor, com limiares de decisão; (iii) contrastar essa relação com a da capacidade anaeróbia (sprints '
 'repetidos) e com a da composição corporal, com controle do porte do atleta; (iv) caracterizar a sonolência e o estresse '
 'percebido e a sua relação com o humor; e (v) quantificar em que medida as características do atleta modulam a magnitude '
 'dos efeitos agudo e crônico. Hipotetizou-se que a aptidão aeróbia intermitente seria o correlato fisiológico mais '
 'consistente da fadiga do humor, ao passo que a capacidade anaeróbia e o porte corporal teriam papel secundário.')

# ===== 2 MATERIAIS E MÉTODOS =====
H('2 MATERIAIS E MÉTODOS')
H('2.1 Delineamento e amostra',12,before=6)
P('Estudo descritivo-comparativo, observacional e de medidas repetidas, conduzido em condições ecológicas de treinamento '
 'durante o microciclo pré-competitivo (21 a 27 de abril de 2024), com %d atletas de handebol do sexo masculino de nível '
 'competitivo, conforme os princípios éticos da Declaração de Helsinque e com consentimento informado.'%sm['n'])
H('2.2 Instrumentos',12,before=6)
P('O humor foi avaliado pela BRUMS-24, com 24 itens em escala de 0 (“nada”) a 4 (“extremamente”), agrupados em seis '
 'subescalas de 0 a 16 pontos (tensão, depressão, raiva, vigor, fadiga e confusão); a Perturbação Total do Humor (PTH) '
 'resume o perfil (soma das negativas menos o vigor). Para a classificação de perfis e a análise multivariada, os escores '
 'foram convertidos em escores T (M = 50; DP = 10). A aptidão aeróbia intermitente foi avaliada pelo Teste de Carminatti '
 '(T-CAR), teste de campo progressivo e intermitente (repetições de 12 s de corrida em vaivém intercaladas por 6 s de '
 'recuperação, até a exaustão), que adota o pico de velocidade (PV) como desfecho (FERNANDES-DA-SILVA et al., 2016). A '
 'estrutura de esforço e pausa do T-CAR reproduz o padrão intermitente das ações do handebol, de modo que o PV expressa '
 'a capacidade aeróbia de sustentar e repetir esforços de alta intensidade — e não apenas a intensidade máxima isolada '
 '—, o que o qualifica como marcador fisiológico específico para esta amostra. O T-CAR '
 'foi aplicado em 15 de abril de 2024, quatro dias de treino antes do início do microciclo, de modo a servir como '
 'parâmetro fisiológico de linha de base das análises. No mesmo diário eletrônico, registraram-se ainda a sonolência, '
 'pela Escala de Sonolência de Epworth em versão de 6 itens (probabilidade de cochilar em seis situações; 0–18), e o '
 'estresse percebido, pela Escala de Estresse Percebido de 14 itens (PSS-14; 0–56, com sete itens de pontuação '
 'invertida), tomados como marcadores de recuperação e de carga psicossocial. A composição corporal foi avaliada por '
 'antropometria (massa e estatura) e por dobras cutâneas (percentual de gordura), de utilidade para a caracterização da amostra '
 'e ao ajuste alométrico do desempenho no T-CAR.')
tinv=table('Inventário dos instrumentos: domínio avaliado, estrutura, amplitude e momento de coleta.',
    ['Instrumento','Domínio avaliado','Estrutura e amplitude','Coleta'],
    [['BRUMS-24 (Escala de Humor de Brunel)','Humor: tensão, depressão, raiva, vigor, fadiga, confusão (+ PTH)','24 itens (4 por subescala), 0–4 por item; subescalas 0–16','Baseline (1×) e pré/pós de cada treino'],
     ['Fadiga física e mental','Fadiga percebida (física e mental)','2 itens únicos, 0–10 cada','Junto ao diário eletrônico'],
     ['Epworth (versão de 6 itens)','Sonolência diurna','6 itens, 0–18','Diária'],
     ['PSS-14 (Escala de Estresse Percebido)','Estresse percebido','14 itens (7 de pontuação invertida), 0–56','Diária'],
     ['Teste de Carminatti (T-CAR)','Aptidão aeróbia intermitente','Teste de campo progressivo e intermitente; pico de velocidade (km/h)','Linha de base e final'],
     ['Salto vertical (CMJ)','Potência de membros inferiores','Altura do salto (cm)','Linha de base e final'],
     ['Sprints repetidos (Baker)','Capacidade de sprints repetidos (componente anaeróbio)','8 sprints; melhor tempo (s) e índice de fadiga %F; menor = melhor','Linha de base e final'],
     ['Antropometria e dobras cutâneas','Composição corporal','Massa (kg), estatura (cm), gordura (%)','Linha de base'],
     ['Monitoramento de carga','Carga interna','Frequência cardíaca, TRIMP e percepção de esforço (PSE)','Sessões de treino']],
    note='Autorrelato: BRUMS-24, fadiga física/mental, Epworth e PSS-14. Avaliações físicas: T-CAR, CMJ, sprints repetidos (Baker) e antropometria. PTH = Perturbação Total do Humor; PSE = percepção subjetiva de esforço; TRIMP = training impulse.',fs=8.5)
P('A Tabela %d sintetiza os instrumentos empregados, o domínio de cada um, a sua estrutura e o momento de coleta.'%tinv)
H('2.3 Procedimentos',12,before=6)
P('O BRUMS foi autoaplicado por formulário eletrônico ao longo de sete dias consecutivos (21 a 27 de abril de 2024). '
 'A data e o horário de cada resposta foram definidos pelo carimbo automático de registro do formulário — e não pela '
 'data informada pelo respondente —, o que garante a alocação correta de cada observação ao dia e ao momento de coleta. '
 'O primeiro dia (21/04) constituiu a avaliação de linha de base (baseline), com uma única coleta por atleta; nos seis '
 'dias de treino subsequentes (22 a 27/04) previram-se duas coletas diárias — uma ao início (pré) e outra ao final '
 '(pós) do treino. Em cada dia de treino, a primeira resposta de cada atleta foi tomada como pré e a última como pós. '
 'Por se tratar de registro em condições ecológicas, houve adesão irregular: nem todos os atletas responderam em todos '
 'os dias ou nas duas janelas, e eventuais respostas repetidas dentro de uma mesma janela foram reduzidas a um único '
 'valor (a primeira, no baseline; a primeira e a última, como pré e pós, nos dias de treino). Após essa padronização e '
 'a exclusão de um registro isolado fora da janela (29/04), o conjunto analítico totalizou %d observações válidas dos '
 '%d atletas (uma no baseline e até duas por dia de treino).'%(sm['n_obs'],sm['n']))
H('2.4 Análise estatística',12,before=6)
P('A análise seguiu uma sequência do descritivo ao inferencial. Inicialmente, empregou-se estatística descritiva (média, '
 'desvio-padrão, mediana e amplitude) para caracterizar as variáveis, e o teste de Shapiro-Wilk para verificar a '
 'normalidade — etapa que orienta a escolha entre testes paramétricos e não paramétricos. Como as distribuições violaram '
 'a normalidade, com forte concentração de escores baixos nas dimensões negativas (efeito de piso), os dados não foram '
 'transformados (NEVILL; LANE, 2007) e privilegiaram-se testes não paramétricos.')
P('Para descrever a resposta aguda ao treino, os momentos pré e pós foram comparados pelo teste de Wilcoxon para amostras '
 'pareadas, acompanhado do tamanho de efeito (d de Cohen), que expressa a magnitude da diferença independentemente do '
 'tamanho da amostra (trivial < 0,2; pequeno < 0,5; médio < 0,8; grande ≥ 0,8); a mesma lógica aplicou-se à comparação '
 'entre o primeiro e o último dia. Para verificar se cada dimensão variou ao longo dos sete dias, utilizou-se o teste de '
 'Friedman — equivalente não paramétrico da ANOVA de medidas repetidas —, com o W de Kendall como tamanho de efeito '
 '(0,1 pequeno; 0,3 moderado; 0,5 grande); quando significativo, aplicou-se o pós-teste das médias marginais de um modelo '
 'misto (correção de Tukey), que identifica entre quais dias, especificamente, houve diferença. A consistência das '
 'medidas repetidas ao longo da semana foi estimada pelo coeficiente de correlação intraclasse (ICC), interpretado como '
 'pobre (< 0,50), moderado (0,50–0,75), bom (0,75–0,90) ou excelente (> 0,90). As propriedades psicométricas das '
 'subescalas foram descritas pela consistência interna (alfa de Cronbach), pela assimetria e pelo efeito de piso '
 '(presente quando mais de 15% das observações no menor escore; TERWEE et al., 2007). A partir do ICC e do desvio-padrão, '
 'derivaram-se o erro-padrão de medida (SEM), a mudança mínima detectável (MDC = 1,96 × √2 × SEM) e a menor mudança '
 'relevante (SWC = 0,2 × desvio-padrão entre atletas), a fim de balizar a interpretação de mudanças individuais no '
 'humor e o uso prático do limiar de aptidão.')
P('Como confirmação robusta, as comparações Dia 1 → Dia 7 e pré → pós foram reanalisadas por análise multivariada de '
 'variância (MANOVA) de medidas repetidas sobre os escores T, que testa as seis dimensões em conjunto e controla o erro '
 'de múltiplas comparações (lambda de Wilks, F e eta-quadrado parcial — η²ₚ; 0,01 pequeno; 0,06 médio; 0,14 grande); nos '
 'testes univariados de acompanhamento da MANOVA, aplicou-se o ajuste de Bonferroni ao nível de significância, com a divisão de '
 '0,05 pelas seis dimensões (α = 0,008), para controlar o erro do tipo I. A '
 'associação entre as dimensões foi quantificada pela correlação de Spearman (ρ). Por fim, a relação entre o pico de '
 'velocidade do T-CAR e a fadiga foi analisada por regressão, e um limiar de pico de velocidade foi estabelecido por '
 'regressão logística (probabilidade de um dia de fadiga elevada em função do pico de velocidade), com o ponto de corte '
 'definido pelo índice de Youden — que maximiza a soma de sensibilidade e especificidade —, a discriminação avaliada '
 'pela área sob a curva ROC e a qualidade de ajuste do modelo logístico reportada pela log-verossimilhança, pelo '
 'pseudo-R² de McFadden e pelos critérios de informação de Akaike (AIC) e bayesiano (BIC). Como verificação a '
 'posteriori da forma funcional dessa relação, a associação entre o pico de velocidade e a resposta semanal de humor '
 '(fadiga física e vigor) foi ajustada por três modelos concorrentes — linear, logarítmico (ln do pico de velocidade) e '
 'polinomial de 2º grau —, com a escolha daquele de melhor compromisso entre ajuste e parcimônia pelo menor AIC/BIC, além do '
 'R² ajustado e do RMSE; para o modelo logístico, comparou-se ainda o preditor na escala bruta contra a logarítmica. '
 'Modelou-se ainda a probabilidade de um dia de baixo vigor (tercil inferior) em função do pico de velocidade, com os '
 'intervalos de confiança das razões de chances e das áreas sob a curva obtidos por reamostragem (bootstrap, 1000 '
 'repetições) de atletas, para acomodar a estrutura de medidas repetidas. Além do ponto de corte único, o pico de velocidade foi '
 'estratificado por dois limiares — os tercis de sua distribuição — em três faixas de aptidão, com a comparação da '
 'prevalência de dias críticos entre as faixas ordenadas pelo teste de tendência linear de Cochran-Armitage. '
 'Para separar o efeito da aptidão do efeito do tamanho corporal, o pico de velocidade foi '
 'ainda submetido a ajuste alométrico: estimou-se o expoente da relação alométrica pela regressão log(PV)–log(massa) e '
 'normalizou-se o PV pela massa elevada a esse expoente, com o contraste entre as associações do PV bruto e do PV normalizado '
 'com a fadiga e o vigor. A capacidade anaeróbia foi avaliada pelo teste de sprints repetidos (Baker), do qual se '
 'tomaram o melhor tempo, a soma dos tempos e o índice de fadiga (percentual de decremento); a sua associação com a '
 'fadiga e o vigor foi examinada por correlação de Spearman, por correlação parcial (com controle da massa corporal) e '
 'por ajuste alométrico, de modo a separar o efeito da capacidade do efeito do porte corporal. '
 'Para identificar os moduladores da magnitude dos efeitos, definiu-se o efeito agudo como a '
 'variação intrassessão (pós − pré) por atleta-dia e o efeito crônico como a inclinação semanal (taxa por dia) da média '
 'diária de cada dimensão; a associação de cada efeito com as características do atleta — aptidão aeróbia intermitente '
 '(PV), composição corporal, potência de membros inferiores (salto CMJ), sprints repetidos (Baker), sonolência, estresse e carga interna '
 '(TRIMP) — foi quantificada pela correlação de Pearson e pelo coeficiente de determinação (R²), além de modelos de '
 'regressão múltipla com coeficientes padronizados (β*); dado o número de testes, aplicou-se a correção de Holm para '
 'comparações múltiplas.')
P('Para quantificar de forma rigorosa a especificidade do correlato físico, empregou-se a análise de comunalidade e '
 'dominância: decompôs-se o R² da fadiga e do vigor entre a aptidão aeróbia, a capacidade anaeróbia e a massa corporal, '
 'com estimativa, para cada preditor, da variância única (ganho de R² sobre os demais) e da dominância geral (contribuição '
 'média sobre todos os subconjuntos de preditores), bem como a variância compartilhada entre aptidão e massa '
 '(NIMON; OSWALD, 2013; AZEN; BUDESCU, 2003). A robustez do limiar de pico de velocidade foi ainda testada por análise '
 'de curva de especificação (specification-curve): enumeraram-se todas as combinações razoáveis de valor de limiar, '
 'definição de dia crítico e forma do preditor (bruto ou ajustado pela massa) e examinou-se a distribuição do tamanho '
 'de efeito sobre o conjunto das especificações, de modo a aferir se o achado depende de escolhas analíticas '
 'particulares (SIMONSOHN; SIMMONS; NELSON, 2020). Por fim, para integrar a aptidão à trajetória do humor num único '
 'modelo, ajustou-se uma regressão bayesiana multinível de curva de crescimento com o pico de velocidade como preditor '
 'de nível 2, com a separação entre o seu efeito sobre o nível do humor (intercepto) e o seu efeito sobre a taxa de '
 'mudança (interação com o dia), com interceptos e inclinações aleatórios por atleta e estimação por amostrador de Gibbs (prioris conjugadas '
 'fracamente informativas; semente 2024), validada contra um modelo misto de máxima verossimilhança restrita '
 '(McELREATH, 2020). Adotou-se nível de significância de 5% '
 '(p < 0,05), com as análises conduzidas em ambiente Python (bibliotecas pandas, NumPy, SciPy e statsmodels).')

# ===== 3 RESULTADOS =====
H('3 RESULTADOS')
H('3.1 Caracterização da amostra',12,before=6)
tp=table('Distribuição da amostra por posição de jogo (n = %d).'%sm['n'],['Posição','n','%'],
    [[k,v,c2('%.1f'%(100*v/sm['n']))] for k,v in sm['pos'].items()]+[['Total',sm['n'],'100,0']],fs=9)
def srow(lab,k):
    v=sm[k]; return [lab,c2('%.1f'%v['mean']),c2('%.1f'%v['sd']),'%s–%s'%(c2('%.1f'%v['mn']),c2('%.1f'%v['mx']))]
ta=table('Caracterização sociodemográfica e antropométrica (n = %d).'%sm['n'],['Variável','Média','DP','Mín–Máx'],
    [srow('Idade (anos)','idade'),srow('Estatura (cm)','estatura'),srow('Massa corporal (kg)','massa'),srow('Percentual de gordura (%)','pG'),srow('Experiência (anos)','exp')])
P('Participaram %d atletas do sexo masculino (idade %s ± %s anos; estatura %s ± %s cm; massa corporal %s ± %s kg), '
 'distribuídos entre as posições de jogo (Tabela %d) e com caracterização sociodemográfica e antropométrica compatível '
 'com a de handebolistas de nível competitivo (Tabela %d).'%(
   sm['n'],c2('%.1f'%sm['idade']['mean']),c2('%.1f'%sm['idade']['sd']),c2('%.1f'%sm['estatura']['mean']),c2('%.1f'%sm['estatura']['sd']),
   c2('%.1f'%sm['massa']['mean']),c2('%.1f'%sm['massa']['sd']),tp,ta))
H('3.2 Normalidade das distribuições',12,before=6)
tn=table('Teste de normalidade (Shapiro-Wilk) das dimensões do BRUMS.',['Dimensão','Estatística (W)','p'],
    [[lab,c2('%.3f'%SH[k]['W']),'< 0,001' if SH[k]['p']<0.001 else c2('%.3f'%SH[k]['p'])] for k,lab in [('Vigor','Vigor'),('Fadiga','Fadiga'),('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Confusao','Confusão')]],fs=9)
P('As seis dimensões não seguem distribuição normal (p < 0,001; Tabela %d), o que justifica os testes não paramétricos.'%tn)
H('3.3 Estatística descritiva e propriedades psicométricas das dimensões',12,before=6)
def drow(k,lab):
    v=desc[k]; return [lab,c2('%.2f'%v['mean']),c2('%.1f'%v['md']),c2('%.2f'%v['sd']),'%s–%s'%(c2('%.0f'%v['mn']),c2('%.0f'%v['mx']))]
td=table('Estatística descritiva das dimensões do BRUMS e da PTH (%d observações).'%sm['n_obs'],
    ['Dimensão','Média','Mediana','DP','Mín–Máx'],[drow(k,l) for k,l in ORD],fs=9)
P('A descritiva geral consta na Tabela %d: o vigor e a fadiga concentram as maiores médias e variabilidade.'%td)
def psyrow(k,lab):
    o=PSY[k]; return [lab,c2('%.2f'%o['alpha']),c2('%.0f'%(100*o['floor0']))+'%',c2('%+.2f'%o['skew']),'sim' if o['floor_effect'] else 'não']
tpsy=table('Propriedades psicométricas das subescalas do BRUMS: consistência interna, efeito de piso e assimetria (%d observações).'%sm['n_obs'],
    ['Dimensão','α de Cronbach','% no piso (=0)','Assimetria','Efeito de piso'],
    [psyrow(k,l) for k,l in [('Vigor','Vigor'),('Fadiga','Fadiga'),('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Confusao','Confusão')]],
    note='Efeito de piso presente quando > 15% das observações no menor escore (TERWEE et al., 2007).',fs=8.5)
P('As propriedades psicométricas (Tabela %d) contextualizam a insensibilidade das dimensões negativas. O vigor e a '
 'fadiga distribuíram-se por toda a escala, sem efeito de piso (%s%% e %s%% no menor escore) e com consistência interna '
 'adequada a boa (α = %s e %s). As quatro dimensões negativas, ao contrário, apresentaram forte efeito de piso — de '
 '%s%% (tensão) a %s%% (confusão) das observações no escore zero — e acentuada assimetria positiva, o que comprime a '
 'variância e limita, por construção, a sua responsividade à carga; trata-se de uma expressão do perfil mentalmente '
 'saudável (iceberg) do atleta de elite, e não de falha do instrumento (MORGAN, 1985; TERWEE et al., 2007). A '
 'Perturbação Total do Humor, por ser composta, não sofreu efeito de piso e foi governada pelo eixo energia–fadiga '
 '(o vigor e a fadiga explicam %s%% da sua variância), o que a mantém sensível à carga.'%(
   tpsy,c2('%.0f'%(100*PSY['Vigor']['floor0'])),c2('%.0f'%(100*PSY['Fadiga']['floor0'])),
   c2('%.2f'%PSY['Vigor']['alpha']),c2('%.2f'%PSY['Fadiga']['alpha']),
   c2('%.0f'%(100*PSY['Tensao']['floor0'])),c2('%.0f'%(100*PSY['Confusao']['floor0'])),
   c2('%.0f'%(100*PSY['PTH']['r2_axis']))))
H('3.4 Diferenças entre pré e pós-treino (com tamanho de efeito)',12,before=6)
def ppr(k,lab):
    v=pr[k]; return [lab,c2('%.2f'%v['pre']),c2('%.2f'%v['pos']),c2('%+.0f'%v['pct'])+'%',pstr(v['p']),c2('%+.2f'%v['dz']),v['mag']]
tpp=table('Diferenças entre pré e pós-treino (agregado da semana; Wilcoxon e d de Cohen).',
    ['Dimensão','Pré (M)','Pós (M)','Variação (%)','p','d de Cohen','Magnitude'],[ppr(k,l) for k,l in ORD],fs=9)
P('A resposta aguda foi consistente (Tabela %d): o vigor caiu (%s%%; d = %s) e a fadiga e a PTH subiram do pré para o '
 'pós, com efeito médio.'%(tpp,c2('%.0f'%abs(pr['Vigor']['pct'])),c2('%.2f'%pr['Vigor']['dz'])))
H('3.5 Comportamento diário e comparação entre todos os dias',12,before=6)
f_traj=figure(f'{FG}/xb2_traj.png','Comportamento das dimensões do humor ao longo da semana (médias diárias; áreas sombreadas = IC95%).')
f_box=figure(f'{FG}/xb3_box.png','Diagramas de caixa (box plot) das seis dimensões do BRUMS por dia da semana.')
def ddrow(k,lab):
    dm=CRS['dec'][k]['dm']; f=FR[k]; return [lab]+[c2('%.1f'%dm[str(d)]) for d in range(1,8)]+[c2('%.1f'%f['chi']),pstr(f['p']),c2('%.2f'%f['W'])]
tdd=table('Médias diárias das dimensões do BRUMS e teste de Friedman com W de Kendall (variação entre os sete dias).',
    ['Dimensão','D1','D2','D3','D4','D5','D6','D7','χ²','p','W de Kendall'],[ddrow(k,l) for k,l in ORD],
    note='W de Kendall = tamanho de efeito do teste de Friedman (0,1 pequeno; 0,3 moderado; 0,5 grande).',fs=8)
P('As médias diárias, o teste de Friedman e o W de Kendall constam na Tabela %d (Figuras %d e %d): houve diferença '
 'significativa entre os dias para o vigor (p %s; W = %s), a fadiga (p %s; W = %s), a tensão (p %s) e a confusão '
 '(p %s), sempre com magnitude pequena a moderada, coerente com um microciclo de acúmulo dentro da faixa funcional.'%(
   tdd,f_traj,f_box,pstr(FR['Vigor']['p']),c2('%.2f'%FR['Vigor']['W']),pstr(FR['Fadiga']['p']),c2('%.2f'%FR['Fadiga']['W']),pstr(FR['Tensao']['p']),pstr(FR['Confusao']['p'])))
def icrow(k,lab):
    i=IC[k]; return [lab,c2('%.2f'%i['icc1']),c2('%.2f'%i['icck']),i['cls']]
tic=table('Consistência das medidas repetidas ao longo da semana (coeficiente de correlação intraclasse, ICC).',
    ['Dimensão','ICC(2,1)','ICC(2,k)','Consistência'],
    [icrow(k,l) for k,l in [('Vigor','Vigor'),('Fadiga','Fadiga'),('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Confusao','Confusão')]],fs=9)
P('A consistência das medidas repetidas (Tabela %d) foi moderada a boa, com os valores mais baixos na raiva e na confusão — '
 'dimensões mais reativas dia a dia.'%tic)
P('Os limiares de mudança do humor delimitam o alcance da leitura individual. O erro-padrão de medida foi de %s ponto '
 'no vigor e %s na fadiga (escala 0–16), de modo que a mudança mínima detectável — a menor variação individual '
 'distinguível do erro de medida — situou-se em torno de %s a %s pontos no vigor e %s a %s na fadiga (90%% e 95%% de '
 'confiança). Como esses valores superam largamente a menor mudança relevante (SWC = %s e %s), uma oscilação isolada de '
 'um único atleta só deve ser tomada como real quando ultrapassa a mudança mínima detectável; do contrário, a decisão '
 'ganha robustez ao apoiar-se na média semanal e na posição do atleta em relação ao limiar de pico de velocidade, e '
 'não em leituras pontuais.'%(
   c2('%.1f'%MDC['Vigor']['sem']),c2('%.1f'%MDC['Fadiga']['sem']),
   c2('%.1f'%MDC['Vigor']['mdc90']),c2('%.1f'%MDC['Vigor']['mdc95']),
   c2('%.1f'%MDC['Fadiga']['mdc90']),c2('%.1f'%MDC['Fadiga']['mdc95']),
   c2('%.1f'%MDC['Vigor']['swc']),c2('%.1f'%MDC['Fadiga']['swc'])))
def emmc(v,d): return c2('%.2f'%PHJ[v]['emm'][str(d)])
def sig1(v,d): return '' if d==1 or PHJ[v]['pairs']['1_%d'%d]['ptukey']>=0.05 else '*'
te=table('Pós-teste (médias marginais do modelo misto) por dia, com comparação de cada dia ao Dia 1.',
    ['Dia','Vigor','Fadiga','Fadiga física'],
    [['Dia %d'%d,emmc('Vigor',d)+sig1('Vigor',d),emmc('Fadiga',d)+sig1('Fadiga',d),emmc('FadFisica',d)+sig1('FadFisica',d)] for d in range(1,8)],
    note='* diferença significativa em relação ao Dia 1 (Tukey, p < 0,05).',fs=9)
f_ph=figure(f'{FG}/ph_emm.png','Trajetória diária (médias marginais) com comparação de todos os dias ao Dia 1 (* p < 0,05).')
def npairs(v): return sum(1 for kk,pp in PHJ[v]['pairs'].items() if pp['ptukey']<0.05)
P('Na comparação de todos os dias entre si (pós-teste de Tukey; Tabela %d; Figura %d), o vigor diferiu significativamente em '
 '%d dos 21 pares de dias e a fadiga em %d — sempre no sentido de piora em relação aos primeiros dias —, com a confirmação da '
 'deterioração progressiva do eixo energia–fadiga.'%(te,f_ph,npairs('Vigor'),npairs('Fadiga')))
H('3.6 Confirmação por análise multivariada (MANOVA em escores T)',12,before=6)
f_prof=figure(f'{FG}/xb5_profile_d1d7.png','Perfil de humor em escores T (M = 50; DP = 10) no Dia 1 e no Dia 7.',w=13.5)
def mvrow(x):
    return [x['lab'],c2('%.1f'%x['m1']),c2('%.1f'%x['s1']),c2('%.1f'%x['m2']),c2('%.1f'%x['s2']),c2('%.2f'%x['F']),pstr(x['p'])+('*' if x['p']<0.008 else ''),c2('%+.2f'%x['d']),c2('%.3f'%x['eta'])]
mv=MV['d1d7']
tmv=table('Comparação Dia 1 → Dia 7 das seis dimensões em escores T (MANOVA de medidas repetidas; n = %d).'%mv['n'],
    ['Dimensão','D1 M','D1 DP','D7 M','D7 DP','F','p','d','η²ₚ'],[mvrow(x) for x in mv['rows']],
    note='Wilks λ = %s; F(%d,%d) = %s; p %s; η²ₚ = %s. * p < 0,008 (α ajustado por Bonferroni para as seis dimensões).'%(c2('%.3f'%mv['wilks']),mv['df1'],mv['df2'],c2('%.2f'%mv['Fmv']),pstr(mv['p_mv']),c2('%.2f'%mv['eta_mv'])),fs=8.5)
P('A análise multivariada confirmou a diferença entre o Dia 1 e o Dia 7 (Wilks λ = %s; F(%d,%d) = %s; p %s; η²ₚ = %s). '
 'Nos testes univariados com o critério corrigido de Bonferroni (α = 0,008), apenas o vigor (d = %s) e a fadiga (d = %s) '
 'permaneceram significativos — o que evidencia a concentração do efeito no eixo energia–fadiga —, enquanto a tensão '
 '(p %s) e a confusão (p %s), significativas apenas sob α = 0,05, não resistiram à correção (Tabela %d; Figura %d). A '
 'resposta aguda pré → pós também foi multivariadamente significativa (Wilks λ = %s; p %s), o que reforça o achado.'%(
   c2('%.3f'%mv['wilks']),mv['df1'],mv['df2'],c2('%.2f'%mv['Fmv']),pstr(mv['p_mv']),c2('%.2f'%mv['eta_mv']),
   c2('%+.2f'%mvv('d1d7','Vigor','d')),c2('%+.2f'%mvv('d1d7','Fadiga','d')),pstr(mvv('d1d7','Tensao','p')),pstr(mvv('d1d7','Confusao','p')),tmv,f_prof,c2('%.3f'%MV['prepos']['wilks']),pstr(MV['prepos']['p_mv'])))
H('3.7 Aptidão intermitente (T-CAR): regressão, comparação de modelos e limiar de pico de velocidade',12,before=6)
def tcrow(o):
    lab='Pico de velocidade — T-CAR (km/h)' if o['k']=='PVini' else o['lab']
    return [lab,o['n'],c2('%.1f'%o['m']),c2('%.1f'%o['sd']),'%s–%s'%(c2('%.1f'%o['mn']),c2('%.1f'%o['mx']))]
TCD2=[o for o in TCD if o['k'] not in ('PV','dPV')]
ttc=table('Estatística descritiva do desempenho no Teste de Carminatti (T-CAR).',
    ['Parâmetro','n','Média','DP','Mín–Máx'],[tcrow(o) for o in TCD2],
    note='PV = pico de velocidade; FC = frequência cardíaca; TRIMP = training impulse.',fs=8.5)
f_pv=figure(f'{FG}/pv7_scatter.png','Dispersão e reta de regressão (com banda de IC95%) entre o pico de velocidade do T-CAR e as médias semanais de vigor, fadiga, PTH e fadiga física.',w=15.0)
P('O desempenho no T-CAR consta na Tabela %d. A regressão mostrou que atletas com maior pico de velocidade reportaram '
 'menos fadiga física (ρ = %s; p %s) e mais vigor (ρ = %s; p %s) ao longo da semana (Figura %d). A partir de um '
 'modelo que estima a probabilidade de um dia de fadiga elevada em função do pico de velocidade, identificou-se um limiar '
 'de aproximadamente %s km/h (área sob a curva = %s; sensibilidade = %s; especificidade = %s): abaixo desse valor, os '
 'atletas apresentaram maior probabilidade de dias de fadiga elevada, o que fornece uma referência objetiva para '
 'individualizar a carga.'%(
   ttc,c2('%+.2f'%PV['pv']['wk_FadFisica']['TCAR1']['rho']),pstr(PV['pv']['wk_FadFisica']['TCAR1']['rho_p']),
   c2('%+.2f'%PV['pv']['wk_Vigor']['TCAR1']['rho']),pstr(PV['pv']['wk_Vigor']['TCAR1']['rho_p']),f_pv,
   c2('%.1f'%LP['thr']),c2('%.2f'%LP['auc']),c2('%.2f'%LP['sens']),c2('%.2f'%LP['spec'])))
# --- comparação a posteriori de modelos de curva (linear x logarítmico x polinomial) ---
NMC={'linear':'Linear','log':'Logarítmica','quad':'Polinomial (2º grau)'}
def cvrow(resp):
    rows=[]; C=TCV['cont'][resp]; best=C['best']
    for i,mdl in enumerate(('linear','log','quad')):
        v=C[mdl]; name=NMC[mdl]+(' †' if mdl==best else '')
        rel=(C['lab']+' × T-CAR') if i==0 else ''
        rows.append([rel,name,c2('%.3f'%v['r2']),c2('%.3f'%v['r2a']),c2('%.1f'%v['aic']),c2('%.1f'%v['bic']),c2('%.3f'%v['rmse'])])
    return rows
tcv=table('Comparação a posteriori de modelos de curva (linear, logarítmico e polinomial) para a relação entre o pico de velocidade do T-CAR e a resposta semanal de humor, por qualidade de ajuste.',
    ['Relação','Modelo','R²','R² ajustado','AIC','BIC','RMSE'],
    cvrow('FadFisica')+cvrow('Vigor'),
    note='AIC = critério de informação de Akaike; BIC = critério bayesiano; RMSE = raiz do erro quadrático médio. † modelo de melhor ajuste (menor AIC). Valores menores de AIC/BIC indicam melhor compromisso entre ajuste e parcimônia.',fs=8.5)
f_cv=figure(f'{FG}/tcar_curvas.png','Dispersão do pico de velocidade do T-CAR contra as médias semanais de fadiga física e de vigor, com os três modelos ajustados sobrepostos (o de menor AIC em linha cheia mais grossa).',w=15.0)
LGa=TCV['logistic']['linear_pred']; LGb=TCV['logistic']['log_pred']
dAIC=abs(TCV['cont']['FadFisica']['log']['aic']-TCV['cont']['FadFisica']['linear']['aic'])
P('Para verificar se a relação entre a aptidão e a resposta de humor seria mais bem descrita por uma forma não linear, '
 'ajustaram-se a posteriori três modelos de curva à relação entre o pico de velocidade do T-CAR e as médias semanais de '
 'fadiga física e de vigor — linear, logarítmico (ln do pico de velocidade) e polinomial de 2º grau —, comparados por R² '
 'ajustado, AIC e BIC (Tabela %d; Figura %d). Em ambas as relações, o modelo linear ofereceu o melhor compromisso entre '
 'ajuste e parcimônia (fadiga física: AIC = %s; vigor: AIC = %s); o ajuste logarítmico foi praticamente equivalente '
 '(ΔAIC ≈ %s) e o termo quadrático não compensou o parâmetro adicional (AIC maior, R² ajustado menor). Dentro da faixa '
 'de aptidão observada, portanto, a associação aptidão–humor é essencialmente linear, o que respalda descrevê-la pela '
 'reta de regressão. Para o desfecho binário — a probabilidade de um dia de fadiga elevada —, a regressão logística '
 'apresentou poder discriminativo moderado (pseudo-R² de McFadden = %s; AUC = %s), e a comparação do preditor na escala '
 'bruta contra a logarítmica favoreceu marginalmente a escala bruta (AIC %s vs. %s), mantida, assim, na definição do '
 'limiar de pico de velocidade.'%(
   tcv,f_cv,c2('%.1f'%TCV['cont']['FadFisica']['linear']['aic']),c2('%.1f'%TCV['cont']['Vigor']['linear']['aic']),
   c2('%.1f'%dAIC),c2('%.3f'%LGa['mcfadden']),c2('%.2f'%LGa['auc']),c2('%.1f'%LGa['aic']),c2('%.1f'%LGb['aic'])))
# --- logísticas complementares: baixo vigor ~ PV e iceberg ~ PV ---
BV=LOG['baixo_vigor']; BPV=LOG['aptidao']['iceberg_pv']
P('Duas regressões logísticas complementares fecham a leitura da aptidão. Simetricamente ao limiar de fadiga, modelou-se '
 'a probabilidade de um dia de baixo vigor (vigor no tercil inferior) em função do pico de velocidade: cada km/h a mais '
 'reduziu a chance de um dia de baixo vigor (OR = %s por km/h) e o índice de Youden localizou um limiar de '
 'aproximadamente %s km/h (sensibilidade = %s; especificidade = %s; área sob a curva = %s [IC95%% %s–%s]), próximo ao '
 'limiar de fadiga e confirma o pico de velocidade como referência objetiva de prontidão. Por fim, a probabilidade de '
 'o atleta exibir o perfil iceberg tendeu a crescer com a aptidão (OR = %s por km/h; área sob a curva = %s), porém sem '
 'significância (IC95%% %s–%s) — direção coerente com o papel protetor da capacidade aeróbia, que esta amostra, contudo, '
 'não teve potência para confirmar.'%(
   c2('%.2f'%BV['OR_kmh']),c2('%.1f'%BV['thr']),c2('%.2f'%BV['sens']),c2('%.2f'%BV['spec']),
   c2('%.2f'%BV['auc']),c2('%.2f'%BV['auc_lo']),c2('%.2f'%BV['auc_hi']),
   c2('%.2f'%BPV['OR_pv']),c2('%.2f'%BPV['auc']),c2('%.2f'%BPV['OR_lo']),c2('%.2f'%BPV['OR_hi'])))
# --- dois limiares de PV: três faixas de aptidão ---
BND=L2['fadiga']['bands']
tl2=table('Estratificação do pico de velocidade do T-CAR em três faixas de aptidão por dois limiares (tercis da distribuição): prevalência de dias de fadiga elevada e de baixo vigor em cada faixa.',
    ['Faixa de aptidão','PV (km/h)','n (atleta-dia)','Fadiga elevada','Baixo vigor'],
    [[BND[k]['lab'],c2(BND[k]['rng']),str(BND[k]['n']),c2('%.0f'%L2['fadiga']['bands'][k]['prev'])+'%',c2('%.0f'%L2['vigor']['bands'][k]['prev'])+'%'] for k in range(3)],
    note='Dois limiares (t₁ = %s km/h; t₂ = %s km/h) definem três faixas. Dia de fadiga elevada = fadiga física no tercil superior; dia de baixo vigor = vigor no tercil inferior. Tendência linear da prevalência (Cochran-Armitage): fadiga p = %s; baixo vigor p = %s.'%(
        c2('%.1f'%L2['t1']),c2('%.1f'%L2['t2']),c2('%.3f'%L2['fadiga']['trend_p']),c2('%.3f'%L2['vigor']['trend_p'])),fs=8.5)
f_l2=figure(f'{FG}/limiar2_faixas.png','Dois limiares do pico de velocidade do T-CAR (%s e %s km/h) que delimitam três faixas de aptidão (esquerda) e a prevalênciade dias de fadiga elevada e de baixo vigor em cada faixa (direita).'%(c2('%.1f'%L2['t1']),c2('%.1f'%L2['t2'])),w=15.5)
P('Em vez de um único ponto de corte, o pico de velocidade foi ainda estratificado por dois limiares — os tercis de sua '
 'distribuição (t₁ = %s km/h; t₂ = %s km/h) —, que definem três faixas de aptidão com valor de decisão direto (Tabela %d; '
 'Figura %d). A prevalência de dias críticos caiu de forma consistente da faixa de baixa para a de alta aptidão, tanto '
 'para a fadiga elevada (%s%% → %s%% → %s%%; tendência p = %s) quanto para o baixo vigor (%s%% → %s%% → %s%%; tendência '
 'p = %s). O ganho da estratificação dupla é sobretudo prático: enquanto o limiar único de Youden (≈ %s km/h) traça uma '
 'só linha, os dois limiares isolam uma faixa superior de aptidão (PV > %s km/h) em que o risco de dias críticos cai '
 'para cerca de um quinto das observações, e uma faixa inferior (PV < %s km/h) que concentra o maior risco — e oferece '
 'à comissão técnica uma referência graduada, e não binária, para individualizar a carga e priorizar a recuperação dos '
 'atletas menos aptos.'%(
   c2('%.1f'%L2['t1']),c2('%.1f'%L2['t2']),tl2,f_l2,
   c2('%.0f'%L2['fadiga']['bands'][0]['prev']),
   c2('%.0f'%L2['fadiga']['bands'][1]['prev']),c2('%.0f'%L2['fadiga']['bands'][2]['prev']),c2('%.3f'%L2['fadiga']['trend_p']),
   c2('%.0f'%L2['vigor']['bands'][0]['prev']),c2('%.0f'%L2['vigor']['bands'][1]['prev']),c2('%.0f'%L2['vigor']['bands'][2]['prev']),c2('%.3f'%L2['vigor']['trend_p']),
   c2('%.1f'%L2['fadiga']['youden']),c2('%.1f'%L2['t2']),c2('%.1f'%L2['t1'])))
f_lc=figure(f'{FG}/logistic_curvas.png','Curvas logísticas ajustadas: à esquerda, a probabilidade de um dia de fadiga elevada decresce com o pico de velocidade do T-CAR (pontos = proporção observada por faixa, tamanho proporcional ao n); à direita, a probabilidade do perfil barbatana de tubarão cresce ao longo do microciclo.',w=15.5)
P('As curvas logísticas ajustadas sintetizam, em forma gráfica, os dois modelos de probabilidade discutidos: a queda '
 'monotônica da chance de um dia de fadiga elevada à medida que aumenta o pico de velocidade e a elevação da chance do '
 'perfil de barbatana de tubarão ao longo dos dias (Figura %d).'%f_lc)
f_spc=figure(f'{FG}/spec_curve.png','Análise de curva de especificação (specification-curve) dos limiares de pico de velocidade: cada ponto é o risco relativo de um dia crítico (PV baixo vs alto) sob uma especificação analítica; o painel inferior indica a forma do PV (bruto ou ajustado pela massa). A quase totalidade das especificações situa-se acima de 1 (linha tracejada).',w=15.5)
P('Como um ponto de corte é, em si, uma decisão analítica arbitrária, submeteu-se o achado a uma análise de curva de '
 'especificação (specification-curve): varreram-se %d especificações que combinam o valor do limiar, a definição de dia '
 'crítico (fadiga elevada, baixo vigor ou qualquer um) e a forma do pico de velocidade (bruto ou ajustado pela massa), '
 'e registrou-se, em cada uma, o risco relativo de dia crítico entre os atletas abaixo e acima do limiar (Figura %d). O '
 'efeito mostrou-se notavelmente robusto: em %s%% das especificações o risco relativo excedeu 1 — ou seja, os atletas '
 'menos aptos concentraram mais dias críticos —, com risco relativo mediano de %s (faixa de %s a %s). A robustez '
 'sobreviveu, ademais, ao uso do pico de velocidade ajustado pela massa (RR mediano %s), o que reitera que o valor '
 'preditivo da aptidão não se reduz ao porte corporal. A conclusão prática é que a existência de um limiar útil de '
 'pico de velocidade não depende de uma escolha particular de corte, mas se sustenta ao longo de todo o espectro de '
 'decisões razoáveis.'%(
   SPC['n_specs'],f_spc,c2('%.0f'%(100*SPC['frac_rr_gt1'])),c2('%.2f'%SPC['median_rr']),
   c2('%.2f'%SPC['rr_range'][0]),c2('%.2f'%SPC['rr_range'][1]),c2('%.2f'%SPC['by_form']['ajustado_massa'])))

# ----- 3.11 Sonolência e estresse percebido -----
H('3.8 Sonolência (Epworth) e estresse percebido (PSS-14)',12,before=6)
def ssrow(v,lab):
    dd=SS['desc'][v]; pp=SS['perath'][v]
    return [lab,'%s ± %s'%(c2('%.1f'%dd['M']),c2('%.1f'%dd['SD'])),'%s (%s–%s)'%(c2('%.0f'%dd['Md']),c2('%.0f'%dd['mn']),c2('%.0f'%dd['mx'])),
            '%s ± %s'%(c2('%.1f'%pp['M']),c2('%.1f'%pp['SD'])),pstr(SS['traj'][v]['p']).replace('= ','')]
tss=table('Sonolência (Epworth, 6 itens, 0–18) e estresse percebido (PSS-14, 0–56) no microciclo: descritivos por observação e por atleta, e variação entre os dias.',
    ['Variável','M ± DP (obs.)','Mediana (mín–máx)','M ± DP (atleta)','Friedman p'],
    [ssrow('Epworth','Sonolência (Epworth)'),ssrow('PSS','Estresse (PSS-14)')],
    note='Friedman: teste da variação entre os sete dias. Epworth em versão de 6 itens (sem ponto de corte clínico padrão).',fs=9)
f_ss=figure(f'{FG}/sono_traj.png','Trajetória diária (média ± IC95%) da sonolência (Epworth) e do estresse percebido (PSS-14) ao longo do microciclo.',w=15.0)
ec=SS['corr']['Epworth']; pc=SS['corr']['PSS']
P('A sonolência e o estresse percebido, coletados no mesmo diário, tiveram comportamentos distintos ao longo do '
 'microciclo (Tabela %d; Figura %d). A sonolência (Epworth) elevou-se progressivamente do primeiro ao último dia '
 '(Friedman p %s; W = %s), padrão compatível com o acúmulo de fadiga e de débito de sono ao fim da semana de carga. No '
 'plano interindividual (média semanal por atleta), atletas mais sonolentos reportaram mais fadiga (ρ = %s; p %s) e pior '
 'humor global (PTH; ρ = %s; p %s), o que posiciona a sonolência como um marcador de recuperação que acompanha a '
 'resposta afetiva à carga. O estresse percebido (PSS-14), ao contrário, manteve-se estável entre os dias (Friedman p '
 '%s) e em patamar moderado, sem associação significativa com o vigor (ρ = %s; p %s) ou com a fadiga (ρ = %s; p %s); '
 'nenhuma das duas variáveis diferiu entre pré e pós-treino (Epworth p %s; PSS p %s), coerente com a natureza mais '
 'estável desses instrumentos. A estabilidade do estresse indica que a deterioração do humor no eixo energia–fadiga foi '
 'específica da carga de treino, e não reflexo de um aumento concomitante do estresse percebido.'%(
   tss,f_ss,pstr(SS['traj']['Epworth']['p']),c2('%.2f'%SS['traj']['Epworth']['W']),
   c2('%+.2f'%ec['Fadiga']['rho']),pstr(ec['Fadiga']['p']),c2('%+.2f'%ec['TMD']['rho']),pstr(ec['TMD']['p']),
   pstr(SS['traj']['PSS']['p']),c2('%+.2f'%pc['Vigor']['rho']),pstr(pc['Vigor']['p']),c2('%+.2f'%pc['Fadiga']['rho']),pstr(pc['Fadiga']['p']),
   pstr(SS['prepos']['Epworth']['p']),pstr(SS['prepos']['PSS']['p'])))

# ----- 3.9 Composição corporal e ajuste alométrico -----
H('3.9 Composição corporal e ajuste alométrico do pico de velocidade',12,before=6)
def alr(c,lab,un):
    dd=AL['desc'][c]; return [lab,'%s ± %s'%(c2('%.1f'%dd['M']),c2('%.1f'%dd['SD'])),'%s–%s %s'%(c2('%.1f'%dd['mn']),c2('%.1f'%dd['mx']),un)]
talom=table('Composição corporal da amostra (n = %d) e associação com a fadiga.'%AL['n'],
    ['Variável','M ± DP','Amplitude'],
    [alr('massa','Massa corporal','kg'),alr('estatura','Estatura','cm'),alr('pGordura','Gordura corporal','%')],
    note='Massa e estatura por antropometria; gordura corporal por dobras cutâneas.',fs=9)
AB=AL['pv_adj']
P('A caracterização da composição corporal (Tabela %d) mostrou massa de %s kg e gordura corporal de %s%%. Nem a massa '
 '(ρ = %s com a fadiga física; p %s) nem o percentual de gordura (ρ = %s; p %s) associaram-se significativamente à '
 'fadiga, embora com tendência positiva. Como o pico de velocidade do T-CAR depende do porte corporal, examinou-se sua '
 'escala alométrica: o expoente da relação log(PV)–log(massa) foi de %s (r = %s; p %s), o que confirma que os atletas mais '
 'pesados atingem menor PV. Ao normalizar o PV pela massa elevada a esse expoente (PV·massa^%s), a associação com o '
 'vigor manteve-se praticamente inalterada (ρ = %s; p %s), ao passo que a associação com a fadiga física atenuou-se e '
 'perdeu significância (de ρ = %s para ρ = %s; p %s). Ou seja, o vínculo entre aptidão e vigor é um efeito genuíno de '
 'capacidade aeróbia, independente do tamanho corporal, enquanto parte da relação entre aptidão e fadiga física é '
 'atribuível ao porte do atleta.'%(
   talom,c2('%.1f'%AL['desc']['massa']['M']),c2('%.1f'%AL['desc']['pGordura']['M']),
   c2('%+.2f'%AL['corr']['massa']['FadFisica']['rho']),pstr(AL['corr']['massa']['FadFisica']['p']),
   c2('%+.2f'%AL['corr']['pGordura']['Fadiga']['rho']),pstr(AL['corr']['pGordura']['Fadiga']['p']),
   c2('%.2f'%AL['alom']['b']),c2('%.2f'%AL['alom']['loglog_r']),pstr(AL['alom']['loglog_p']),
   c2('%.2f'%(-AL['alom']['b'])),
   c2('%+.2f'%AB['alom']['vigor_rho']),pstr(AB['alom']['vigor_p']),
   c2('%+.2f'%AB['bruto']['fadfis_rho']),c2('%+.2f'%AB['alom']['fadfis_rho']),pstr(AB['alom']['fadfis_p'])))

# ----- 3.10 Capacidade anaeróbia (RSA) -----
H('3.10 Capacidade anaeróbia (RSA) e fadiga: especificidade e ajuste pela massa',12,before=6)
def rsarow(mk):
    a=RSA['assoc'][mk]
    return [RSA['labels'][mk],
        '%s / %s'%(c2('%+.2f'%a['FadFisica']['raw_r']),c2('%+.2f'%a['FadFisica']['par_r'])),
        '%s / %s'%(c2('%+.2f'%a['Fadiga']['raw_r']),c2('%+.2f'%a['Fadiga']['par_r'])),
        '%s / %s'%(c2('%+.2f'%a['Vigor']['raw_r']),c2('%+.2f'%a['Vigor']['par_r']))]
trsa=table('Associação entre a capacidade anaeróbia (teste de sprints repetidos, Baker) e a fadiga e o vigor semanais: correlação bruta e parcial (com controle da massa corporal).',
    ['Medida do RSA','Fadiga física (bruto / parcial)','Fadiga BRUMS (bruto / parcial)','Vigor (bruto / parcial)'],
    [rsarow('BkMel'),rsarow('BkSoma'),rsarow('BkF')],
    note='Correlação de Spearman (ρ); parcial = com controle da massa corporal. Melhor tempo, soma e índice de fadiga: menor = melhor. n = %d atletas. Como contraste, o pico de velocidade aeróbio associou-se à fadiga física (ρ = %s; p = %s).'%(RSA['n'],c2('%+.2f'%RSA['pv_fad'][0]),c2('%.3f'%RSA['pv_fad'][1])),fs=8.5)
f_rsa=figure(f'{FG}/rsa_especificidade.png','Correlatos da fadiga física semanal: à esquerda, a capacidade anaeróbia (melhor tempo em sprints repetidos) apresenta associação fraca e não significativa; à direita, a aptidão aeróbia (pico de velocidade do T-CAR) apresenta associação moderada e significativa.',w=15.5)
P('A capacidade anaeróbia foi avaliada por um teste de sprints repetidos (oito sprints; Baker), do qual se tomaram o '
 'melhor tempo, a soma dos tempos e o índice de fadiga (percentual de decremento). Nenhuma das três medidas se associou '
 'de forma significativa à fadiga física, à fadiga do BRUMS ou ao vigor (Tabela %d; Figura %d); as correlações foram '
 'fracas, no limite da significância apenas para o índice de fadiga com o vigor (ρ = %s). A capacidade anaeróbia '
 'mostrou-se, ademais, independente da aptidão aeróbia (RSA × PV: ρ = %s), o que indica que os dois testes captam '
 'capacidades distintas.'%(trsa,f_rsa,c2('%+.2f'%RSA['assoc']['BkF']['Vigor']['raw_r']),c2('%+.2f'%RSA['rsa_pv'][0])))
P('A normalização pela massa corporal esclarece a natureza dessas associações. Os atletas mais pesados foram mais '
 'lentos nos sprints (massa × melhor tempo: ρ = %s) e tenderam a relatar mais fadiga (massa × fadiga física: ρ = %s); '
 'ao controlar a massa, a fraca associação entre o melhor tempo e a fadiga física praticamente desaparece (bruto ρ = %s '
 '→ parcial ρ = %s), e o ajuste alométrico conduz à mesma conclusão (ρ ajustado = %s). Em contraste, a associação entre '
 'o pico de velocidade aeróbio e a fadiga física — significativa (ρ = %s) — resiste ao ajuste. O padrão revela '
 'especificidade fisiológica: no acúmulo de pré-temporada, a fadiga do humor é rastreada pela aptidão aeróbia, e não '
 'pela capacidade anaeróbia de sprints repetidos nem pelo porte corporal.'%(
   c2('%+.2f'%RSA['mass_bkmel'][0]),c2('%+.2f'%RSA['mass_fad'][0]),
   c2('%+.2f'%RSA['assoc']['BkMel']['FadFisica']['raw_r']),c2('%+.2f'%RSA['assoc']['BkMel']['FadFisica']['par_r']),
   c2('%+.2f'%RSA['allo']['FadFisica']['adj']),c2('%+.2f'%RSA['pv_fad'][0])))
def comrow(v,p):
    o=COM[v]; return [COM['labels'][p],c2('%.3f'%o['dominance'][p]),c2('%.3f'%o['unique'][p])]
tcom=table('Análise de comunalidade e dominância: decomposição do R² da fadiga física e do vigor semanais entre a aptidão aeróbia (PV do T-CAR), a capacidade anaeróbia (sprints repetidos) e a massa corporal.',
    ['Preditor','Fadiga: dominância / única','Vigor: dominância / única'],
    [[COM['labels'][p],
      '%s / %s'%(c2('%.3f'%COM['FadFisica']['dominance'][p]),c2('%.3f'%COM['FadFisica']['unique'][p])),
      '%s / %s'%(c2('%.3f'%COM['Vigor']['dominance'][p]),c2('%.3f'%COM['Vigor']['unique'][p]))] for p in COM['labels']],
    note='Dominância geral = contribuição média ao R² sobre todos os subconjuntos de preditores; variância única = ganho de R² do preditor sobre os demais (semiparcial²). R² total: fadiga %s; vigor %s. n = %d atletas.'%(
        c2('%.3f'%COM['FadFisica']['full']),c2('%.3f'%COM['Vigor']['full']),COM['n']),fs=8.5)
P('Para além do contraste bivariado, uma análise de comunalidade e dominância decompôs a variância explicada da fadiga '
 'e do vigor entre os três candidatos físicos — aptidão aeróbia, capacidade anaeróbia e massa corporal — e quantificou '
 'o que cada um contribui de forma única e o que partilham entre si (Tabela %d). O resultado formaliza a '
 'especificidade. Na fadiga física, a aptidão aeróbia foi o preditor dominante (dominância geral %s) e reteve a maior '
 'variância única (%s); a massa, embora se associe à fadiga isoladamente (R² = %s), praticamente não acrescenta '
 'variância única (%s), pois a quase totalidade da sua contribuição é partilhada com a aptidão aeróbia (variância comum '
 'aeróbia∩massa = %s) — isto é, o aparente vínculo do porte com a fadiga é, em grande parte, aptidão aeróbia disfarçada '
 'de massa. A capacidade anaeróbia contribui de forma única desprezível (%s). No vigor, o quadro é ainda mais nítido: a '
 'aptidão aeróbia domina (variância única %s) e a massa opera como supressora (comunalidade negativa, %s), o que '
 'explica por que a relação aptidão–vigor se fortalece — e não se atenua — ao controlar o porte, em consonância com o '
 'ajuste alométrico. A capacidade anaeróbia, novamente, é irrelevante (%s). A decomposição, portanto, sustenta '
 'quantitativamente que é a aptidão aeróbia — e não a anaeróbia nem a massa — o correlato específico do humor sob '
 'carga.'%(
   tcom,c2('%.3f'%COM['FadFisica']['dominance']['PVini']),c2('%.3f'%COM['FadFisica']['unique']['PVini']),
   c2('%.3f'%COM['FadFisica']['r_mass']),c2('%.3f'%COM['FadFisica']['unique']['massa']),
   c2('%.3f'%COM['FadFisica']['common_pv_mass']),c2('%.3f'%COM['FadFisica']['unique']['BkMel']),
   c2('%.3f'%COM['Vigor']['unique']['PVini']),c2('%.3f'%COM['Vigor']['common_pv_mass']),
   c2('%.3f'%COM['Vigor']['unique']['BkMel'])))

# ----- 3.11 Moduladores da magnitude dos efeitos -----
H('3.11 Moduladores da magnitude dos efeitos agudo e crônico',12,before=6)
MODm=MOD['chronic']['matrix']
def modcell(s): return ('%s (%s)%s'%(c2('%+.2f'%s['r']),c2('%.2f'%s['r2']),'*' if s['p']<0.05 else '')) if s else '—'
def modrow(mk): return [MOD['modlab'][mk],modcell(MODm['Vigor'][mk]),modcell(MODm['Fadiga'][mk])]
tmod=table('Moduladores da magnitude do efeito crônico (inclinação semanal): correlação (r) e coeficiente de determinação (R²) entre cada característica do atleta e a taxa de variação do vigor e da fadiga.',
    ['Moderador','Vigor — r (R²)','Fadiga — r (R²)'],[modrow(m) for m in MOD['mods']],
    note='r = correlação de Pearson; R² = coeficiente de determinação (variância explicada); * p < 0,05 sem correção. Nenhuma associação sobrevive à correção de Holm para %d testes. CMJ = salto vertical; Baker = melhor tempo em sprints repetidos (menor = melhor); TRIMP = carga interna.'%MOD['n_tests'],fs=8.5)
f_mod=figure(f'{FG}/moduladores.png','Modulação da taxa semanal pelo perfil do atleta: à esquerda, o coeficiente de determinação (R²) de cada característica sobre a taxa do vigor (verde) e da fadiga (laranja) — a fadiga é praticamente não modulada; à direita, a taxa de queda do vigor decresce com a potência de membros inferiores (salto CMJ).',w=15.5)
_acp=min(MOD['acute']['day_trend'][v]['p'] for v in MOD['vars'])
_acR=max(s['r2'] for v in MOD['vars'] for s in (MOD['acute_matrix'][v][m] for m in MOD['mods']) if s)
_fdR=max(s['r2'] for s in (MODm['Fadiga'][m] for m in MOD['mods']) if s)
mlv=MOD['mlr']['sl_Vigor~PVini+CMJ_mai']
P('Uma questão aplicada relevante é saber quais características do atleta modulam a intensidade da resposta — e em que '
 'medida. A magnitude do efeito agudo (variação pré → pós) revelou-se estável e independente do perfil individual: a '
 'queda média do vigor (%s ponto por sessão), a elevação da fadiga (+%s) e a da PTH (+%s) não se acentuaram ao longo da '
 'semana (tendência por dia: p ≥ %s) nem se vincularam de modo consistente à aptidão, à composição corporal, à potência, '
 'ao sono ou à carga interna — todos os coeficientes de determinação foram baixos (R² < %s). O choque afetivo '
 'intrassessão comporta-se, portanto, como uma resposta estereotipada, compartilhada pelo grupo.'%(
   c2('%.2f'%MOD['acute']['mean']['Vigor']),c2('%.2f'%MOD['acute']['mean']['Fadiga']),c2('%.2f'%MOD['acute']['mean']['TMD']),
   c2('%.2f'%_acp),c2('%.2f'%(_acR+0.005))))
P('O efeito crônico (inclinação semanal) exibiu modulação apenas parcial, restrita ao vigor (Tabela %d; Figura %d). A '
 'taxa de acúmulo de fadiga foi praticamente idêntica entre os atletas — nenhum moderador explicou fração relevante de '
 'sua variância (R² ≤ %s) —, ao passo que a taxa de queda do vigor associou-se à potência de membros inferiores (salto '
 'CMJ: r = %s; R² = %s) e, em menor grau, à sonolência: os atletas mais explosivos perderam vigor mais depressa. Um '
 'modelo de regressão múltipla com aptidão e potência explicou cerca de %s%% da variância da taxa do vigor (R² = %s; R² '
 'ajustado = %s), com o maior peso padronizado atribuído à potência (β* = %s). Impõe-se, porém, cautela: entre os %d '
 'testes de moderação, nenhuma associação resistiu à correção de Holm para comparações múltiplas, de sorte que esses '
 'vínculos permanecem exploratórios. Predomina, assim, uma resposta de humor amplamente independente das características '
 'individuais mensuradas — sobretudo no eixo da fadiga —, o que reforça o caráter geral e esperado da reação à carga da '
 'pré-temporada.'%(
   tmod,f_mod,c2('%.2f'%_fdR),
   c2('%+.2f'%MODm['Vigor']['CMJ_mai']['r']),c2('%.2f'%MODm['Vigor']['CMJ_mai']['r2']),
   c2('%.0f'%(100*mlv['r2'])),c2('%.2f'%mlv['r2']),c2('%.2f'%mlv['adj']),c2('%.2f'%mlv['beta']['CMJ_mai']),MOD['n_tests']))

H('3.12 Modelo bayesiano multinível: a aptidão modula o nível, não a taxa',12,before=6)
def bpc(v,k): b=BP[v][k]; return '%s [%s; %s]'%(c2('%+.2f'%b['mean']),c2('%+.2f'%b['lo']),c2('%+.2f'%b['hi']))
tbp=table('Modelo bayesiano multinível do humor com a aptidão aeróbia (pico de velocidade) como preditor de nível 2: efeito sobre o nível (intercepto) e sobre a taxa (interação com o dia). Estimativas a posteriori (média e IC de credibilidade de 95%).',
    ['Desfecho','Aptidão → nível','Aptidão → taxa (× dia)'],
    [['Fadiga física',bpc('FadFisica','b_pv'),bpc('FadFisica','b_pvxdia')],
     ['Vigor',bpc('Vigor','b_pv'),bpc('Vigor','b_pvxdia')],
     ['PTH',bpc('TMD','b_pv'),bpc('TMD','b_pvxdia')]],
    note='Modelo y = b0 + b1·dia + b2·(pré/pós) + b3·PV + b4·(PV×dia) + interceptos e inclinações aleatórios por atleta; amostrador de Gibbs (semente 2024). n = %d observações, %d atletas. IC que exclui zero indica efeito consistente.'%(BP['FadFisica']['n_obs'],BP['FadFisica']['n_ath']),fs=8.5)
P('Para integrar a aptidão fisiológica à trajetória do humor num único arcabouço, ajustou-se um modelo bayesiano '
 'multinível com o pico de velocidade como preditor de nível 2, com a separação entre o seu efeito sobre o nível do humor '
 '(intercepto) e o seu efeito sobre a taxa de mudança (interação com o dia) (Tabela %d). O padrão foi consistente entre '
 'os desfechos: a aptidão deslocou o nível na direção esperada — mais apto, menos fadiga física (%s) e mais vigor '
 '(%s) —, embora os intervalos de credibilidade tocassem o zero, o que reflete a imprecisão de estimar um efeito de nível '
 '2 com %d atletas. Já a interação aptidão × dia foi claramente nula em todos os desfechos (fadiga %s; vigor %s; PTH '
 '%s): a aptidão não modulou a taxa de acúmulo do humor ao longo do microciclo. A leitura que emerge, coerente com a '
 'análise de moduladores, é que a aptidão aeróbia atua como um fator de nível — fixa a linha de base do estado afetivo '
 '—, e não como um fator de ritmo: a trajetória de deterioração desenrola-se em velocidade comum ao grupo, '
 'independentemente da aptidão (McELREATH, 2020).'%(
   tbp,bpc('FadFisica','b_pv'),bpc('Vigor','b_pv'),BP['FadFisica']['n_ath'],
   bpc('FadFisica','b_pvxdia'),bpc('Vigor','b_pvxdia'),bpc('TMD','b_pvxdia')))

# ===== 4 DISCUSSÃO =====
H('4 DISCUSSÃO')
P('O presente estudo examinou os correlatos fisiológicos e de bem-estar da resposta de humor de handebolistas de elite '
 'ao longo da última semana de pré-temporada. Como pano de fundo, a deterioração do humor concentrou-se no eixo '
 'energia–fadiga: o vigor caiu e a '
 'fadiga subiu de forma consistente, tanto na resposta aguda a cada treino (pré → pós: vigor −%s%%, d = %s; fadiga +%s%%, '
 'd = %s) quanto na comparação do primeiro ao último dia (vigor d = %s; fadiga d = %s). A robustez desse achado é '
 'sustentada por três abordagens convergentes — o pós-teste do modelo misto na comparação de todos os dias, o teste de Friedman '
 'e a análise multivariada em escores T (Wilks λ = %s; F(%d,%d) = %s; p %s; η²ₚ = %s). Sob o critério conservador de '
 'Bonferroni, apenas o vigor (η²ₚ = %s) e a fadiga (η²ₚ = %s) permaneceram significativos, enquanto as dimensões negativas '
 'de valência não fadiga, próximas do piso, mantiveram-se estáveis. Confirma-se, assim, que o vigor e a fadiga são as '
 'dimensões subjetivas mais sensíveis à carga de treino, em consonância com a literatura de monitoramento (SAW; MAIN; '
 'GASTIN, 2016; THORPE et al., 2017; KELLMANN et al., 2018). Tal seletividade replica um padrão já descrito noutras '
 'modalidades: em nadadores acompanhados por 24 semanas, o vigor e a fadiga acompanharam o volume de treino '
 '(r = −0,54 e r = +0,53), ao passo que tensão, depressão e confusão permaneceram inertes (PIERCE, 2002) — o que sustenta '
 'centrar a leitura fisiológica subsequente no par vigor–fadiga.'%(
   c2('%.0f'%abs(pr['Vigor']['pct'])),c2('%+.2f'%pr['Vigor']['dz']),c2('%.0f'%pr['Fadiga']['pct']),c2('%+.2f'%pr['Fadiga']['dz']),
   c2('%+.2f'%mvv('d1d7','Vigor','d')),c2('%+.2f'%mvv('d1d7','Fadiga','d')),
   c2('%.3f'%MV['d1d7']['wilks']),MV['d1d7']['df1'],MV['d1d7']['df2'],c2('%.2f'%MV['d1d7']['Fmv']),pstr(MV['d1d7']['p_mv']),c2('%.2f'%MV['d1d7']['eta_mv']),
   c2('%.2f'%mvv('d1d7','Vigor','eta')),c2('%.2f'%mvv('d1d7','Fadiga','eta'))))
P('O caráter do handebol ajuda a contextualizar a magnitude dessa resposta. Trata-se de uma modalidade coletiva de '
 'invasão, marcadamente intermitente, na qual ações de alta intensidade — sprints, saltos, arremessos, bloqueios, '
 'mudanças de direção e contatos — alternam-se com períodos de recuperação incompleta e reclamam, a um só tempo, potência '
 'anaeróbia e capacidade aeróbia intermitente para sustentar o esforço repetido e retardar a instalação da fadiga '
 '(KARCHER; BUCHHEIT, 2014; MICHALSIK; AAGAARD, 2015). Evidências atuais confirmam que o treino e o jogo de handebol '
 'constituem exercício de alta intensidade, com elevadas demandas aeróbias e anaeróbias (PEREIRA et al., 2024), e que os '
 'esportes de invasão se caracterizam por atividade intermitente de alta intensidade intercalada por recuperação de '
 'menor intensidade (VACCARO-BENET et al., 2024). Nesse contexto, o acúmulo de carga da última semana de pré-temporada '
 'explica a deterioração do eixo energia–fadiga observada e esclarece por que '
 'o vigor e a fadiga foram as dimensões mais sensíveis do painel.')
P('É relevante que, mesmo sob o acúmulo de carga da '
 'pré-temporada, os perfis associados a maior risco à saúde mental (Everest invertido, submerso e iceberg invertido) não '
 'se tornaram prevalentes: a deterioração restringiu-se ao perfil de fadiga, cuja prevalência no último dia (%s%%) '
 'superou nitidamente o patamar de referência de atletas do mesmo contexto brasileiro (11,6%% em avaliação momentânea; '
 'DE MIRANDA ROHLFS et al., 2024). Esse contraste sugere que, em handebolistas de elite bem condicionados, a resposta ao '
 'acúmulo de carga é funcional e transitória — uma fadiga esperada — e não um sinal de risco psicológico. Tal leitura '
 'respalda o emprego do perfilamento do humor como triagem de prontidão e de bem-estar em modalidades coletivas de '
 'rendimento (TERRY et al., 2021), no qual a BRUMS já demonstrou sensibilidade a fatores como sono, desempenho e '
 'resultados de partida (ANDRADE et al., 2016; BRANDT; BEVILACQUA; ANDRADE, 2017; DO NASCIMENTO et al., 2026) e '
 'associação com desfechos como a lesão (DE MIRANDA ROHLFS et al., 2025).'%(
   c2('%.0f'%(100*PREV['D7']['Barbatana tubarão']/PREV['n_d7'])),))
P('O caráter transitório dessa resposta deve ser lido à luz da posição do microciclo no planejamento. A última semana '
 'de pré-temporada corresponde a uma fase de acumulação de carga, que antecede o afinamento (tapering) e a competição. '
 'Nesse enquadramento, a queda de vigor e a elevação de fadiga aqui documentadas configuram a assinatura afetiva '
 'esperada do acúmulo, e não um estado estável: o modelo de equilíbrio entre estresse e recuperação prevê que, com a '
 'redução da carga e o reforço da recuperação nos dias subsequentes, o vigor tende a ser restaurado e o perfil de '
 'prontidão (iceberg) a se reinstalar (KELLMANN et al., 2018). Interpretar a fadiga do fim do microciclo como um vale '
 'programado — e não como deterioração — é justamente o que habilita a comissão técnica a distinguir a resposta '
 'funcional daquela que exigiria intervenção, e reforça o valor de manter o monitoramento do humor na transição para o '
 'afinamento, quando a recuperação do estado afetivo é esperada e verificável.')
P('Sob a ótica psicofisiológica, o eixo energia–fadiga funciona como um leitor do equilíbrio entre a carga de treino e '
 'a capacidade de recuperação. Em uma semana de acumulação, esse equilíbrio inclina-se transitoriamente para o polo da '
 'sobrecarga, e a queda do vigor com a ascensão da fadiga exprime, no plano afetivo, a resposta integrada do organismo '
 'ao esforço acumulado — uma cascata de ajustes metabólicos, hormonais e neurais que visam restaurar a homeostase e '
 'poupar energia (WOODS et al., 2018; VRIJKOTTE et al., 2019). No próprio handebol, blocos de treino que elevam '
 'marcadores de dano muscular e de estresse oxidativo fazem-se acompanhar de perturbação do humor, o que ancora a '
 'resposta afetiva a um substrato fisiológico de sobrecarga (ASSUNÇÃO CARVALHO et al., 2018). A elevação da sonolência '
 'ao término da semana, sem '
 'alteração do estresse percebido, respalda essa leitura: a perturbação nasceu da carga física e da restauração '
 'incompleta entre as sessões, não de um estressor psicossocial concomitante. O acoplamento crescente entre o afeto '
 'negativo e a fadiga, do primeiro ao último dia, acompanha essa transição de um estado diferenciado para um estado '
 'integrado de estresse, no qual as respostas afetivas convergem sob um substrato central comum — leitura que o próprio '
 'delineamento sustenta, uma vez que os marcadores medidos (humor, sono e estresse) apontam para a mesma origem na '
 'carga de treino. Essa via afetiva, contudo, guarda autonomia em relação aos índices periféricos: sob treino '
 'intensificado, a perturbação do humor pode subsistir mesmo quando marcadores bioquímicos de fadiga já regrediram '
 '(OSTAPIUK-KAROLCZUK et al., 2025), assincronia que credencia o eixo energia–fadiga como sentinela complementar — e '
 'não redundante — em relação à leitura fisiológica.')
P('Esse conjunto de sinais delineia a assinatura de um sobre-esforço funcional (functional overreaching) — a fadiga '
 'planejada e reversível que precede a supercompensação — e não a de um estado disfuncional. A literatura recente '
 'ampara tal distinção: períodos de treino intensificado agravam a perturbação do humor e comprometem o desempenho, '
 'com plena recuperação após dias de afinamento, tanto em esportes coletivos (CAMPBELL et al., 2020) quanto em '
 'modalidades de resistência (PIACENTINI et al., 2016; WOODS et al., 2018), padrão que revisões sistemáticas reconhecem '
 'como um marcador do sobre-esforço funcional (ROETE et al., 2021). Cabe, todavia, a cautela metodológica que essas '
 'mesmas revisões assinalam: os instrumentos de humor sinalizam o sobre-esforço, porém nem sempre separam a fadiga '
 'aguda do acúmulo funcional (ROETE et al., 2021) — limitação que o presente desenho amortece ao dissociar a resposta '
 'aguda (pré → pós de cada sessão) do acúmulo semanal e ao modelar a trajetória temporal, o que torna visíveis tanto o '
 'choque intrassessão quanto a deriva ao longo dos dias. A preservação das dimensões de valência não fadiga e a '
 'ausência dos perfis de risco à saúde mental completam o quadro de uma adaptação esperada, e não de um processo '
 'patológico.')
P('A análise dos moduladores refina esse entendimento e traz uma mensagem de utilidade prática. A magnitude tanto do '
 'choque agudo quanto da deriva crônica mostrou-se, em larga medida, independente do perfil do atleta: nenhum dos '
 'preditores — aptidão, composição corporal, potência, sono, estresse ou carga interna — explicou parcela expressiva da '
 'variância, e nenhuma associação resistiu à correção para comparações múltiplas. A taxa de acúmulo de fadiga, em '
 'especial, revelou-se homogênea entre os atletas, o que sugere uma resposta obrigatória e comum à carga da '
 'pré-temporada — coerente com a leitura de um sobre-esforço funcional que atinge o grupo de maneira relativamente '
 'uniforme. A única modulação digna de nota, ainda que exploratória, recaiu sobre a taxa de queda do vigor, associada à '
 'potência de membros inferiores: os atletas mais explosivos perderam energia mais depressa, padrão compatível com a '
 'hipótese de que perfis de maior demanda neuromuscular acumulam fadiga central e periférica com maior rapidez sob um '
 'bloco de acumulação aeróbia. Do ponto de vista aplicado, a baixa modulação reforça que a vigilância do humor deve ser '
 'universal na equipe — e não reservada a subgrupos supostamente mais vulneráveis —, embora o eixo do vigor mereça '
 'atenção adicional nos atletas de perfil mais potente.')
P('Situados no conjunto da literatura, esses achados ocupam uma lacuna específica. A BRUMS tem sido aplicada, '
 'predominantemente, como um retrato pontual em torno da competição (BRANDT et al., 2019; DO NASCIMENTO et al., '
 '2026) ou em comparações entre fases de treino distantes semanas ou meses entre si (ROUVEIX et al., 2006), quase '
 'sempre associada ao sono, ao desempenho ou a marcadores hormonais (ANDRADE et al., 2019; ROUVEIX et al., 2006). São '
 'escassas, contudo, as descrições que articulam essa resposta afetiva às capacidades físicas do atleta e que '
 'discriminam qual componente da aptidão efetivamente a rastreia em uma modalidade coletiva de invasão. É nesse ponto '
 'que se concentram as contribuições originais deste estudo, cada qual voltada a uma face da lacuna. Primeiro, a '
 'passagem do descritivo ao correlacional: em vez de apenas documentar a deterioração do humor, o estudo a vincula à '
 'aptidão física e mostra que a magnitude da resposta afetiva tem leitura interindividual. Segundo, e de forma central, '
 'a demonstração de especificidade: entre as capacidades físicas, é a aptidão aeróbia intermitente — e não a anaeróbia '
 '— que rastreia a fadiga do humor, contraste raramente testado e que o presente desenho isola ao confrontar os dois '
 'sistemas na mesma amostra. Terceiro, o controle do porte corporal: por correlação parcial e ajuste alométrico, '
 'evidencia-se que o vínculo aparente da capacidade anaeróbia com a fadiga decorre da massa, e não da qualidade física '
 'em si — uma depuração que a leitura bruta não permitiria. Quarto, a tradução do achado em decisão prática, com '
 'limiares de pico de velocidade (único e em faixas) que sinalizam atletas de maior risco. Em conjunto, essas frentes '
 'convertem a associação humor–aptidão de uma constatação genérica em um critério fisiológico específico e acionável, e '
 'sinalizam um caminho para o monitoramento multidomínio intramicrociclo em esportes de rendimento.')
P('A inclusão do pico de velocidade do T-CAR como parâmetro fisiológico acrescentou uma leitura interindividual à '
 'resposta afetiva. Atletas com maior aptidão aeróbia intermitente reportaram mais vigor (ρ = %s; p %s) e menos fadiga '
 'física (ρ = %s; p %s) ao longo da semana, e um limiar de pico de velocidade de aproximadamente %s km/h discriminou os '
 'dias de maior fadiga (área sob a curva = %s). Esse resultado é coerente com o papel da capacidade aeróbia intermitente '
 'em sustentar e repetir esforços de alta intensidade — pois acelera a recuperação entre as ações e retarda a fadiga — e, '
 'portanto, em modular a tolerância à carga do handebol, e com a validade do pico de velocidade do T-CAR como marcador '
 'de desempenho físico em modalidades intermitentes (FERNANDES-DA-SILVA et al., 2016). Do ponto de vista aplicado, normalizar a '
 'resposta de humor pela aptidão física ajuda a distinguir a fadiga esperada — de atletas menos aptos sob a mesma carga '
 '— daquela que possa sinalizar sobrecarga, e o limiar identificado oferece à comissão técnica uma referência objetiva '
 'para individualizar a prescrição da carga e o reforço da recuperação. A estratificação por dois limiares refina essa '
 'referência ao substituir a linha única por três faixas de aptidão: a prevalência de dias críticos caiu de modo '
 'ordenado da faixa inferior para a superior (tendência significativa para fadiga e para baixo vigor), e a faixa de alta '
 'aptidão isolou atletas com risco de dias críticos reduzido a cerca de um quinto — uma leitura graduada, e não binária, '
 'mais aderente à decisão prática de dosar carga e recuperação atleta a atleta. O ajuste alométrico refinou essa '
 'leitura: como '
 'o pico de velocidade escala negativamente com a massa (expoente %s), parte da relação entre aptidão e fadiga física '
 'reflete o porte corporal e não apenas a capacidade aeróbia — ao normalizar o PV pela massa, essa associação se '
 'atenua. Já a relação entre aptidão e vigor resistiu ao ajuste, o que evidencia um efeito genuíno de capacidade aeróbia '
 'sobre o estado de energia, independentemente do tamanho do atleta.'%(
   c2('%+.2f'%PV['pv']['wk_Vigor']['TCAR1']['rho']),pstr(PV['pv']['wk_Vigor']['TCAR1']['rho_p']),
   c2('%+.2f'%PV['pv']['wk_FadFisica']['TCAR1']['rho']),pstr(PV['pv']['wk_FadFisica']['TCAR1']['rho_p']),
   c2('%.1f'%LP['thr']),c2('%.2f'%LP['auc']),c2('%.2f'%AL['alom']['b'])))
P('A contraposição com a capacidade anaeróbia delimita a especificidade dessa relação. O teste de sprints repetidos '
 '(Baker) — um marcador da capacidade de repetir esforços máximos de curta duração — não se associou de forma '
 'significativa à fadiga do humor em nenhuma de suas medidas, e a fraca associação bruta entre o melhor tempo e a '
 'fadiga física desapareceu ao controlar a massa corporal, o que a atribui ao porte do atleta, e não à capacidade '
 'anaeróbia em si. Como a capacidade anaeróbia se mostrou independente da aptidão aeróbia, os dois testes captam '
 'qualidades físicas distintas; e é a aptidão aeróbia — não a anaeróbia — que rastreia a fadiga afetiva do acúmulo de '
 'pré-temporada. Esse contraste é coerente com a natureza predominantemente aeróbia da carga de uma semana de '
 'acumulação e com a recomendação de sistemas de monitoramento multidomínio, que combinam marcadores neuromusculares, '
 'cardioautonômicos e de autorrelato de humor, sono e estresse (NAUGHTON et al., 2023; TAVARES et al., 2017). A '
 'pertinência dessa moldura no próprio handebol de elite encontra respaldo em investigação de larga escala com centenas '
 'de jogadores de alto nível, na qual o estado de humor, a percepção de estresse e marcadores endócrinos foram '
 'articulados num arcabouço multidomínio de acompanhamento (RATZ-SULYOK et al., 2026) — quadro em que o presente '
 'contraste entre aptidão aeróbia e anaeróbia especifica qual componente físico efetivamente rastreia a resposta '
 'afetiva. Convém, '
 'ainda, a cautela metodológica de que o índice de fadiga do teste de sprints repetidos, por derivar de uma pequena '
 'queda de desempenho, tem confiabilidade limitada e deve ser lido com reservas (OLIVER, 2007; GLAISTER et al., 2008).')
P('Duas escolhas analíticas, incomuns na literatura de monitoramento, sustentam a força dessa conclusão. A primeira é a '
 'análise de comunalidade e dominância, que traduz a especificidade em números: a aptidão aeróbia não apenas dominou a '
 'variância explicada da fadiga como reteve a quase totalidade da contribuição única, ao passo que a variância da massa '
 'se revelou, em grande parte, compartilhada com a própria aptidão — evidência direta de que o aparente vínculo do '
 'porte com a fadiga é aptidão aeróbia sob outro nome — e a capacidade anaeróbia nada acrescentou de único. No vigor, a '
 'massa comportou-se como supressora, o que explica por que a associação aeróbia se robustece ao controlá-la. Esse tipo '
 'de partição, que decompõe o R² em frações única e compartilhada, evita as armadilhas de interpretar coeficientes de '
 'regressão isolados sob colinearidade (NIMON; OSWALD, 2013; AZEN; BUDESCU, 2003). A segunda é a curva de especificação: '
 'em vez de defender um único ponto de corte — sempre discutível —, mostrou-se que a utilidade do limiar de pico de '
 'velocidade persiste ao longo de praticamente todo o espectro de decisões analíticas plausíveis, inclusive quando o '
 'preditor é ajustado pela massa, o que blinda o achado contra a crítica de seletividade de especificação e o torna '
 'mais crível como ferramenta de triagem (SIMONSOHN; SIMMONS; NELSON, 2020). Em conjunto, esses recursos elevam o rigor '
 'da inferência sobre a especificidade aeróbia para além do que a correlação parcial isolada permitiria.')
P('A comparação a posteriori de modelos de curva reforça a leitura dessa relação e agrega parcimônia à sua '
 'interpretação. Testadas as formas linear, logarítmica e polinomial, o modelo linear mostrou o melhor compromisso entre '
 'ajuste e parcimônia para a fadiga física e para o vigor (menores AIC e BIC), com o ajuste logarítmico praticamente '
 'equivalente e sem ganho do termo quadrático — o que revela que, na faixa de aptidão desta amostra, cada incremento de '
 'pico de velocidade se associa a uma variação aproximadamente constante da resposta de humor, sem evidência de '
 'saturação ou de limiar de rendimentos decrescentes. Esse resultado justifica descrever a associação pela reta de '
 'regressão e, ao mesmo tempo, delimita o alcance do modelo: uma relação linear estimada em uma faixa estreita de '
 'aptidão não deve ser extrapolada para além dela. Para o desfecho binário, a regressão logística — cuja qualidade de '
 'ajuste foi documentada pelo pseudo-R² de McFadden e pelos critérios de informação — sustentou o limiar de pico de '
 'velocidade como ferramenta prática de triagem, com preferência pelo preditor na escala bruta, de leitura mais direta para '
 'a comissão técnica.')
P('A sonolência e o estresse percebido acrescentaram duas leituras convergentes com o padrão central. A sonolência '
 '(Epworth) acompanhou a semana de carga: elevou-se rumo ao último dia e associou-se, entre os atletas, a mais fadiga '
 'e a um pior humor global — um marcador de recuperação/sono que corrobora, no plano comportamental, a deterioração do eixo '
 'energia–fadiga e reforça a recomendação de vigiar o sono na fase de acumulação (KELLMANN et al., 2018). Esse vínculo '
 'entre sono e humor sob carga encontra paralelo direto no futebol de base: em jovens jogadores submetidos a treino '
 'intensificado, a restrição de sono precipitou a queda do vigor e a elevação da fadiga, ao passo que a preservação do '
 'sono, sob a mesma carga, manteve o humor estável (FERREIRA et al., 2026) — o que credencia o sono como um moderador '
 'da resposta afetiva ao acúmulo, e não mero acompanhante dela. Já o estresse '
 'percebido (PSS-14) permaneceu estável e moderado, sem se relacionar à oscilação do humor; essa ausência de variação '
 'é informativa, pois indica que a resposta afetiva ao microciclo teve origem na carga de treino, e não em um aumento '
 'concomitante do estresse psicossocial percebido — um argumento a favor da especificidade do achado.')
P('Algumas limitações delimitam o alcance dos achados. A primeira é o tamanho amostral (%d atletas), circunscrito a uma '
 'única equipe e a um único microciclo: o desenho observacional de fase única não autoriza inferência causal, e as '
 'associações entre aptidão e humor são, por natureza, transversais. A segunda é que o limiar de pico de velocidade foi '
 'derivado na própria amostra; embora a curva de especificação ateste a sua robustez a escolhas analíticas, o seu '
 'desempenho preditivo ainda requer validação externa antes do uso clínico. A terceira, e relevante para a '
 'interpretação da dose-resposta, é que a carga interna foi registrada apenas no contexto do teste de campo, e não '
 'continuamente ao longo dos sete dias; por isso não foi possível modelar o gradiente dose-resposta intramicrociclo '
 'entre a carga diária e o humor, e a leitura da modulação fisiológica restringe-se ao gradiente de exposição pela '
 'aptidão (pico de velocidade). A essas somam-se limitações de mensuração já assinaladas: o efeito de piso das '
 'dimensões negativas, a referência dos escores T à própria amostra (e não a normas nacionais), a versão de 6 itens do '
 'Epworth (sem ponto de corte clínico validado) e a PSS-14, cujo horizonte excede o microciclo — razão pela qual '
 'sonolência e estresse foram lidos por média semanal, e não por flutuação diária. Convém, ainda, a cautela já '
 'apontada quanto ao índice de fadiga do teste de sprints repetidos, de confiabilidade limitada.'%sm['n'])
P('Essas limitações, porém, são atenuadas pela convergência das evidências, o que preserva a validade da conclusão '
 'central sobre a especificidade aeróbia. Três abordagens independentes — correlação parcial com controle da massa, '
 'ajuste alométrico e análise de comunalidade — apontam para o mesmo resultado, e a curva de especificação mostra que '
 'ele não depende de uma escolha particular de corte; a própria ausência de gradiente entre a carga interna do teste e '
 'o humor foi reportada de forma transparente, sem sobreinterpretação. Assim, embora a magnitude dos coeficientes deva '
 'ser lida com a cautela que o tamanho amostral impõe, a direção e a especificidade do achado sustentam-se sob '
 'múltiplas lentes analíticas. Como direções '
 'futuras, recomenda-se ampliar a amostra e acompanhar múltiplos microciclos, integrar o humor a marcadores de carga '
 'externa e interna mensurados continuamente — o que habilitaria o gradiente dose-resposta aqui inviável —, adotar '
 'tabelas normativas e agrupamento semeado (seeded k-means) para a classificação de perfis, validar externamente os '
 'limiares de pico de velocidade e testar o valor preditivo dos perfis negativos para desfechos como lesão e sobrecarga '
 '(DE MIRANDA ROHLFS et al., 2024, 2025).')

# ===== 5 CONCLUSÕES =====
H('5 CONCLUSÕES')
P('Na última semana de pré-temporada, a deterioração do humor concentrou-se no eixo energia–fadiga, e a sua magnitude '
 'associou-se de forma seletiva às capacidades físicas. Entre os parâmetros examinados, a aptidão aeróbia intermitente '
 '(pico de velocidade do T-CAR) foi o correlato fisiológico genuíno: os atletas mais aptos relataram mais vigor e menos '
 'fadiga física, com um limiar de aproximadamente %s km/h e dois cortes que delimitaram faixas de risco úteis à '
 'individualização da carga. A capacidade anaeróbia de sprints repetidos, ao contrário, não explicou a fadiga do humor '
 '— e o seu vínculo aparente decorria do porte corporal —, o que revela especificidade aeróbia. A sonolência '
 'acrescentou uma leitura de recuperação, ao acompanhar a fadiga e a perturbação do humor, enquanto o estresse '
 'percebido, estável, delimitou a origem da resposta na carga de treino. A magnitude dos efeitos foi, porém, pouco '
 'modulada pelo perfil individual, o que sugere uma resposta amplamente comum ao grupo. Recomenda-se, portanto, um '
 'monitoramento multidomínio — que combine o humor, a aptidão aeróbia e marcadores de recuperação — centrado no eixo '
 'energia–fadiga.'%c2('%.1f'%LP['thr']))

# ===== DECLARAÇÕES (back-matter) =====
H('Contribuições dos autores',12,before=10)
P('[Autor 1] e [Autor 2] conceberam o estudo e o delineamento. [Autor 1] conduziu a coleta de dados. [Autor 1] e '
 '[Autor 3] realizaram as análises estatísticas. [Autor 1] redigiu a primeira versão do manuscrito. Todos os autores '
 'revisaram criticamente o texto e aprovaram a versão final.',ind=False)
H('Financiamento',12,before=6)
P('[Esta pesquisa não recebeu financiamento externo. / Esta pesquisa foi financiada por [agência], processo nº '
 '[placeholder].]',ind=False)
H('Aprovação ética',12,before=6)
P('O estudo foi conduzido de acordo com a Declaração de Helsinque e aprovado pelo Comitê de Ética em Pesquisa de '
 '[instituição] (parecer nº / CAAE [placeholder]).',ind=False)
H('Consentimento informado',12,before=6)
P('Todos os participantes assinaram o termo de consentimento livre e esclarecido antes da coleta de dados.',ind=False)
H('Disponibilidade dos dados',12,before=6)
P('Os dados que sustentam os achados deste estudo, em forma anonimizada, e os scripts de análise estão disponíveis '
 'mediante solicitação razoável ao autor correspondente.',ind=False)
H('Conflitos de interesse',12,before=6)
P('Os autores declaram não haver conflitos de interesse.',ind=False)
H('Artigo companheiro',12,before=6)
P('Este manuscrito integra um par de estudos-companheiros que compartilham a mesma amostra e o mesmo microciclo, com '
 'desfechos primários distintos. O artigo companheiro — sobre a dinâmica do humor e a modelagem das trajetórias — foi '
 'submetido a outro periódico; ambos se citam mutuamente.',ind=False)

# ===== REFERÊNCIAS =====
H('REFERÊNCIAS')
refs=[
 'ANDRADE, A. et al. Sleep quality, mood and performance: a study of elite Brazilian volleyball athletes. Journal of Sports Science and Medicine, v. 15, n. 4, p. 601–605, 2016.',
 'ANDRADE, A. et al. Sleep quality associated with mood in elite athletes. The Physician and Sportsmedicine, v. 47, n. 3, p. 312–317, 2019. DOI: 10.1080/00913847.2018.1553467.',
 'ANDRADE, A. et al. Effect of practice exergames on the mood states and self-esteem of elementary school boys and girls during physical education classes: a cluster-randomized controlled trial. PLoS ONE, v. 15, n. 6, e0232392, 2020. DOI: 10.1371/journal.pone.0232392.',
 'ASSUNÇÃO CARVALHO, L. C. S. et al. Nectar supplementation reduced biomarkers of oxidative stress, muscle damage, and improved psychological response in highly trained young handball players. Frontiers in Physiology, v. 9, 1508, 2018. DOI: 10.3389/fphys.2018.01508.',
 'AZEN, R.; BUDESCU, D. V. The dominance analysis approach for comparing predictors in multiple regression. Psychological Methods, v. 8, n. 2, p. 129–148, 2003. DOI: 10.1037/1082-989X.8.2.129.',
 'BATTAGLINI, M. P. et al. Analysis of progressive muscle relaxation on psychophysiological variables in basketball athletes. International Journal of Environmental Research and Public Health, v. 19, n. 24, 17065, 2022. DOI: 10.3390/ijerph192417065.',
 'BIRD, S. P. et al. Wellness, mood, sleep, and performance in a women’s national basketball team during international competition. Journal of Human Kinetics, v. 96, p. 163–175, 2025. DOI: 10.5114/jhk/200117.',
 'BRANDT, R.; BEVILACQUA, G. G.; ANDRADE, A. Perceived sleep quality, mood states, and their relationship with performance among Brazilian elite athletes during a competitive period. Journal of Strength and Conditioning Research, v. 31, n. 4, p. 1033–1039, 2017.',
 'BRANDT, R. et al. Comparisons of mood states associated with outcomes achieved by female and male athletes in high-level judo and Brazilian jiu-jitsu championships: psychological factors associated with the probability of success. Journal of Strength and Conditioning Research, v. 35, n. 9, p. 2518–2524, 2019. DOI: 10.1519/JSC.0000000000003218.',
 'CAMPBELL, P. G. et al. The effect of overreaching on neuromuscular performance and wellness responses in Australian rules football athletes. Journal of Strength and Conditioning Research, v. 34, n. 6, p. 1530–1538, 2020. DOI: 10.1519/JSC.0000000000003603.',
 'DE MIRANDA ROHLFS, I. C. P. et al. Prevalence of specific mood profile clusters among elite and youth athletes at a Brazilian sports club. Sports, v. 12, n. 7, 195, 2024. DOI: 10.3390/sports12070195.',
 'DE MIRANDA ROHLFS, I. C. P. et al. Mood states, injury status, and countermovement jump performance in Brazilian high-level sports. Sports, v. 13, n. 9, 303, 2025. DOI: 10.3390/sports13090303.',
 'DO NASCIMENTO, M. H. et al. Acute psychological responses to official match outcomes in male youth volleyball: an observational repeated-measures study within a single national-level team. Frontiers in Psychology, v. 17, 1826372, 2026. DOI: 10.3389/fpsyg.2026.1826372.',
 'FERNANDES-DA-SILVA, J. et al. The peak velocity derived from the Carminatti Test is related to physical match performance in young soccer players. Journal of Sports Sciences, v. 34, n. 24, p. 2238–2245, 2016. DOI: 10.1080/02640414.2015.1093646.',
 'FERREIRA, A. B. M. et al. Impact of sleep restriction and intensified training on mucosal immunity and psychological responses in young soccer players. Journal of Strength and Conditioning Research, v. 40, n. 7, p. e703–e713, 2026. DOI: 10.1519/JSC.0000000000005416.',
 'GLAISTER, M. et al. The reliability and validity of fatigue measures during multiple-sprint work: an issue revisited. Journal of Strength and Conditioning Research, v. 22, n. 5, p. 1597–1601, 2008. DOI: 10.1519/JSC.0b013e318181ab80.',
 'HAN, C.; PARSONS-SMITH, R. L.; TERRY, P. C. Mood profiling in Singapore: cross-cultural validation and potential applications of mood profile clusters. Frontiers in Psychology, v. 11, 665, 2020. DOI: 10.3389/fpsyg.2020.00665.',
 'KARCHER, C.; BUCHHEIT, M. On-court demands of elite handball, with special reference to playing positions. Sports Medicine, v. 44, n. 6, p. 797–814, 2014. DOI: 10.1007/s40279-014-0164-z.',
 'KELLMANN, M. et al. Recovery and performance in sport: consensus statement. International Journal of Sports Physiology and Performance, v. 13, n. 2, p. 240–245, 2018. DOI: 10.1123/ijspp.2017-0759.',
 'LOCHBAUM, M. et al. The Profile of Mood States and athletic performance: a meta-analysis of published studies. European Journal of Investigation in Health, Psychology and Education, v. 11, n. 1, p. 50–70, 2021. DOI: 10.3390/ejihpe11010005.',
 'McELREATH, R. Statistical rethinking: a Bayesian course with examples in R and Stan. 2. ed. Boca Raton: CRC Press, 2020.',
 'MICHALSIK, L. B.; AAGAARD, P. Physical demands in elite team handball: comparisons between male and female players. Journal of Sports Medicine and Physical Fitness, v. 55, n. 9, p. 878–891, 2015.',
 'MORGAN, W. P. Selected psychological factors limiting performance: a mental health model. In: CLARKE, D. H.; ECKERT, H. M. (Ed.). Limits of human performance. Champaign: Human Kinetics, 1985. p. 70–80.',
 'NAUGHTON, M. et al. Defining and quantifying fatigue in the rugby codes. PLoS ONE, v. 18, n. 3, e0282390, 2023. DOI: 10.1371/journal.pone.0282390.',
 'NEVILL, A. M.; LANE, A. M. Why self-report “Likert” scale data should not be log-transformed. Journal of Sports Sciences, v. 25, n. 1, p. 1–2, 2007. DOI: 10.1080/02640410601111183.',
 'NIMON, K. F.; OSWALD, F. L. Understanding the results of multiple linear regression: beyond standardized regression coefficients. Organizational Research Methods, v. 16, n. 4, p. 650–674, 2013. DOI: 10.1177/1094428113493929.',
 'OLIVER, J. L. Is a fatigue index a worthwhile measure of repeated sprint ability? Journal of Science and Medicine in Sport, v. 12, n. 1, p. 20–23, 2007. DOI: 10.1016/j.jsams.2007.10.010.',
 'OSTAPIUK-KAROLCZUK, J. et al. Biochemical and psychological markers of fatigue and recovery in mixed martial arts athletes during strength and conditioning training. Scientific Reports, v. 15, n. 1, 24234, 2025. DOI: 10.1038/s41598-025-09719-z.',
 'PARSONS-SMITH, R. L.; TERRY, P. C.; MACHIN, M. A. Identification and description of novel mood profile clusters. Frontiers in Psychology, v. 8, 1958, 2017. DOI: 10.3389/fpsyg.2017.01958.',
 'PEREIRA, R. et al. Exercise intensity and reliability during recreational team handball training for 50–77-year-old unexperienced women. Biology of Sport, v. 41, n. 4, p. 253–261, 2024. DOI: 10.5114/biolsport.2024.132995.',
 'PIACENTINI, M. F. et al. Effect of intensive training on mood with no effect on brain-derived neurotrophic factor. International Journal of Sports Physiology and Performance, v. 11, n. 6, p. 824–830, 2016. DOI: 10.1123/ijspp.2015-0279.',
 'PIERCE, E. F. Relationship between training volume and mood states in competitive swimmers during a 24-week season. Perceptual and Motor Skills, v. 94, n. 3, p. 1009–1012, 2002. DOI: 10.2466/pms.2002.94.3.1009.',
 'RATZ-SULYOK, F. Z. et al. Associations between endocrine status and stress, mood and psychosomatic status in elite handball players. Sports, v. 14, n. 7, 289, 2026. DOI: 10.3390/sports14070289.',
 'ROETE, A. J. et al. A systematic review on markers of functional overreaching in endurance athletes. International Journal of Sports Physiology and Performance, v. 16, n. 8, p. 1065–1073, 2021. DOI: 10.1123/ijspp.2021-0024.',
 'ROHLFS, I. C. P. M. et al. A Escala de Humor de Brunel (Brums): instrumento para detecção precoce da síndrome do excesso de treinamento. Revista Brasileira de Medicina do Esporte, v. 14, n. 3, p. 176–181, 2008.',
 'ROHLFS, I. C. P. M. et al. Psychometric characteristics of the Brazil Mood Scale among youth and elite athletes using two response time frames. Sports, v. 11, n. 12, 244, 2023. DOI: 10.3390/sports11120244.',
 'ROUVEIX, M. et al. The 24 h urinary cortisol/cortisone ratio and epinephrine/norepinephrine ratio for monitoring training in young female tennis players. International Journal of Sports Medicine, v. 27, n. 11, p. 856–863, 2006. DOI: 10.1055/s-2006-923778.',
 'SAW, A. E.; MAIN, L. C.; GASTIN, P. B. Monitoring the athlete training response: subjective self-reported measures trump commonly used objective measures: a systematic review. British Journal of Sports Medicine, v. 50, n. 5, p. 281–291, 2016. DOI: 10.1136/bjsports-2015-094758.',
 'SIMONSOHN, U.; SIMMONS, J. P.; NELSON, L. D. Specification curve analysis. Nature Human Behaviour, v. 4, n. 11, p. 1208–1214, 2020. DOI: 10.1038/s41562-020-0912-z.',
 'TAVARES, F. et al. Wellness, muscle soreness and neuromuscular performance during a training week in volleyball athletes. Journal of Sports Medicine and Physical Fitness, v. 58, n. 12, p. 1852–1858, 2017. DOI: 10.23736/S0022-4707.17.07818-5.',
 'TERRY, P. C.; LANE, A. M.; FOGARTY, G. J. Construct validity of the Profile of Mood States — Adolescents for use with adults. Psychology of Sport and Exercise, v. 4, n. 2, p. 125–139, 2003. DOI: 10.1016/S1469-0292(02)00035-8.',
 'TERRY, P. C. et al. Mood profiling for sustainable mental health among athletes. Sustainability, v. 13, n. 11, 6116, 2021. DOI: 10.3390/su13116116.',
 'TERWEE, C. B. et al. Quality criteria were proposed for measurement properties of health status questionnaires. Journal of Clinical Epidemiology, v. 60, n. 1, p. 34–42, 2007. DOI: 10.1016/j.jclinepi.2006.03.012.',
 'THORPE, R. T. et al. Monitoring fatigue status in elite team-sport athletes: implications for practice. International Journal of Sports Physiology and Performance, v. 12, n. S2, p. S227–S234, 2017. DOI: 10.1123/ijspp.2016-0434.',
 'VACCARO-BENET, P. et al. Internal and external load profile during beach invasion sports match-play by electronic performance and tracking systems: a systematic review. Sensors, v. 24, n. 12, 3738, 2024. DOI: 10.3390/s24123738.',
 'VRIJKOTTE, S. et al. The overtraining syndrome in soldiers: insights from the sports domain. Military Medicine, v. 184, n. 5-6, p. e192–e200, 2019. DOI: 10.1093/milmed/usy274.',
 'WAGNER, H. et al. Individual and team performance in team-handball: a review. Journal of Sports Science and Medicine, v. 13, n. 4, p. 808–816, 2014.',
 'WOODS, A. L. et al. The effects of intensified training on resting metabolic rate, body composition and performance in trained cyclists. PLoS ONE, v. 13, n. 2, e0191644, 2018. DOI: 10.1371/journal.pone.0191644.']
for rf in refs:
    p=doc.add_paragraph(); r=p.add_run(rf); r.font.size=Pt(11); p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(6); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

OUTP='/home/user/mdlucca/Artigos/Paper2_Fisiologia_bemestar.docx'
doc.save(OUTP); print('SAVED',OUTP,'| Tabelas',_TN[0],'Figuras',_FN[0])
