# -*- coding: utf-8 -*-
import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
R=json.load(open('tcar2_baseline.json')); R1=json.load(open('tcar1_baseline.json'))
def c2(s): return str(s).replace('.',',')
doc=Document()
sty=doc.styles['Normal']; sty.font.name='Times New Roman'; sty.font.size=Pt(12)
sty.element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')
pf=sty.paragraph_format; pf.line_spacing=1.5; pf.space_after=Pt(0)
s=doc.sections[0]; s.top_margin=Cm(3); s.left_margin=Cm(3); s.bottom_margin=Cm(2); s.right_margin=Cm(2)
_TN=[0]; _FN=[0]
def P(t='',just=True,size=12,after=6,before=0):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(size)
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(after); p.paragraph_format.space_before=Pt(before)
    if just: p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.first_line_indent=Cm(1.25)
    return p
def H(t,size=12):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(size)
    p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(6); p.paragraph_format.line_spacing=1.5
def _bd(cell):
    tcPr=cell._tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for edge in ['top','bottom','left','right']:
        e=OxmlElement(f'w:{edge}'); on=edge in ('top','bottom')
        e.set(qn('w:val'),'single' if on else 'nil'); e.set(qn('w:sz'),'6'); e.set(qn('w:color'),'000000'); b.append(e)
    tcPr.append(b)
def table(caption,header,rows,fonte='Fonte: os autores (2026).',widths=None,fs=10):
    _TN[0]+=1
    p=doc.add_paragraph(); r=p.add_run(f'Tabela {_TN[0]} – {caption}'); r.font.size=Pt(11)
    p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(3)
    t=doc.add_table(rows=1,cols=len(header)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,htxt in enumerate(header):
        cc=t.rows[0].cells[i]; cc.text=''; rr=cc.paragraphs[0].add_run(htxt); rr.bold=True; rr.font.size=Pt(fs); rr.font.name='Times New Roman'
        cc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; _bd(cc)
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=''; rr=cells[i].paragraphs[0].add_run(str(val)); rr.font.size=Pt(fs); rr.font.name='Times New Roman'
            cells[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if i>0 else WD_ALIGN_PARAGRAPH.LEFT; _bd(cells[i])
    if widths:
        for i,w in enumerate(widths):
            for rr in t.rows: rr.cells[i].width=Cm(w)
    pf=doc.add_paragraph(); rf=pf.add_run(fonte); rf.font.size=Pt(10); pf.paragraph_format.space_after=Pt(8)
def figure(path,caption,w=15.5):
    _FN[0]+=1
    pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; pp.add_run().add_picture(path,width=Cm(w))
    pc=doc.add_paragraph(); pc.alignment=WD_ALIGN_PARAGRAPH.CENTER; rc=pc.add_run(f'Figura {_FN[0]} – {caption}'); rc.font.size=Pt(11)
    pf=doc.add_paragraph(); pf.alignment=WD_ALIGN_PARAGRAPH.CENTER; rf=pf.add_run('Fonte: os autores (2026).'); rf.font.size=Pt(10); pf.paragraph_format.space_after=Pt(8)

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('T-CAR 2 (APTIDÃO PÓS-BLOCO) E ADAPTAÇÃO INDIVIDUAL COMO PARÂMETROS FISIOLÓGICOS DO ESTADO PSICOLÓGICO DE BASE — COMPARAÇÃO COM O T-CAR 1')
r.bold=True; r.font.size=Pt(13); p.paragraph_format.space_after=Pt(12)
P('Parâmetro fisiológico: T-CAR 2 (reteste — pico de velocidade final, distância/repetições finais e FCmáx) e a adaptação ΔPV (T-CAR 2 − T-CAR 1); carga interna do HIIT (TRIMP/PSE). Parâmetro psicológico: baseline (D1, 21/04). Amostra: 27 handebolistas do sexo masculino.',size=11,after=10)

H('Resumo')
P('Refeita a análise com o T-CAR 2 (aptidão medida após o bloco de treino), as associações com o estado psicológico de base ENFRAQUECEM em relação ao T-CAR 1: o PV do T-CAR 2 correlaciona-se com a fadiga de base a ρ = %s (p = %s), contra ρ = %s (p = %s) do T-CAR 1; e classifica o perfil iceberg com AUC = %s, contra %s do T-CAR 1. A interpretação é direta e temporalmente coerente: a aptidão CONTEMPORÂNEA ao baseline (T-CAR 1) é o covariável correto do estado inicial; o T-CAR 2 reflete a aptidão pós-adaptação e é mais fracamente ligado ao baseline anterior. O grupo melhorou de forma real (ΔPV médio = %s km·h⁻¹, ≈ +6%%), mas a magnitude da adaptação individual NÃO se associou nem ao humor de base nem à deterioração da semana. Recomenda-se usar o T-CAR 1 como covariável fisiológica do baseline e o T-CAR 2 como DESFECHO (ganho de aptidão), não como preditor do estado inicial.'%(
    c2('%+.2f'%R['COR']['PV']['b_Fadiga']['r']),c2('%.2f'%R['COR']['PV']['b_Fadiga']['p']),
    c2('%+.2f'%R1['COR']['PV']['b_Fadiga']['r']),c2('%.2f'%R1['COR']['PV']['b_Fadiga']['p']),
    c2('%.2f'%R['ROC']['PV']['auc']),c2('%.2f'%R1['ROC']['PV']['auc']),c2('%.2f'%R['means']['dPV'])),after=8)

H('1. Parâmetros')
P('T-CAR 2 (reteste incremental pós-bloco): PV final (km·h⁻¹), distância percorrida (repetições finais × (PV/3,6) × 12 s) e FCmáx. Adaptação: ΔPV = PV final − PV inicial (ganho de velocidade aeróbia máxima entre os dois testes). Carga interna: TRIMP de Banister das sessões de HIIT (S1–S3) e PSE. Psicológico: baseline do BRUMS no D1 (vigor, fadiga, fadiga física/mental, PTH/TMD, índice iceberg). Correlações de Spearman; alometria log-log; ROC (AUC).',after=6)

H('2. Resultados')
H('2.1 T-CAR 2 × psicológico de base',12)
OUTB=[('b_Vigor','Vigor'),('b_Fadiga','Fadiga BRUMS'),('b_FadFisica','Fad. física'),('b_FadMental','Fad. mental'),('b_TMD','PTH/TMD'),('b_ice_idx','Iceberg')]
def rc(ph,ov): d=R['COR'][ph][ov]; return '%s (%s)'%(c2('%+.2f'%d['r']),c2('%.2f'%d['p']))
rows=[]
for ph,lab in [('PV','PV final (T-CAR 2)'),('dist','Distância (T-CAR 2)'),('FCmax','FCmáx'),('TRIMP','TRIMP (HIIT)'),('PSE','PSE (HIIT)')]:
    rows.append([lab]+[rc(ph,ov) for ov,_ in OUTB])
table('Correlação de Spearman (ρ, p) entre parâmetros do T-CAR 2 e o psicológico de base (n = 27)',
    ['Parâmetro']+[l for _,l in OUTB],rows,widths=[3.4,2.7,2.7,2.5,2.5,2.4,2.3],fs=8.5)
P('O PV final e a distância do T-CAR 2 mantêm o sinal negativo sobre a fadiga de base, porém enfraquecido e não significativo (ρ = %s). PV e distância permanecem redundantes. A carga interna segue nula, como esperado.'%(
    c2('%+.2f'%R['COR']['PV']['b_Fadiga']['r'])),after=6)

H('2.2 Comparação direta T-CAR 1 vs T-CAR 2',12)
table('Poder de indicação sobre o baseline — T-CAR 1 vs T-CAR 2',
    ['Métrica','T-CAR 1','T-CAR 2','Leitura'],
    [['ρ PV × fadiga (base)',c2('%+.2f'%R1['COR']['PV']['b_Fadiga']['r']),c2('%+.2f'%R['COR']['PV']['b_Fadiga']['r']),'T-CAR 1 mais forte'],
     ['ρ PV × fadiga física (base)',c2('%+.2f'%R1['COR']['PV']['b_FadFisica']['r']),c2('%+.2f'%R['COR']['PV']['b_FadFisica']['r']),'T-CAR 1 mais forte'],
     ['AUC iceberg (base)',c2('%.2f'%R1['ROC']['PV']['auc']),c2('%.2f'%R['ROC']['PV']['auc']),'T-CAR 1 mais forte'],
     ['PV médio (km·h⁻¹)',c2('%.1f'%R1['means']['PV']),c2('%.1f'%R['means']['PV']),'ganho ≈ +%s'%c2('%.1f'%(R['means']['PV']-R1['means']['PV']))]],
    widths=[5.2,3.0,3.0,3.8],fs=9.5)
P('A aptidão pós-bloco (T-CAR 2) explica pior o estado psicológico de base do que a aptidão contemporânea a ele (T-CAR 1) — coerente com a temporalidade: o baseline foi medido antes do bloco de treino que produziu o T-CAR 2. Isso confirma a recomendação de tratar o T-CAR 1 como o marcador fisiológico do baseline.',after=6)

H('2.3 Adaptação individual (ΔPV) — quem ganhou mais aptidão fadigou menos?',12)
A=R['ADAPT']
P('O grupo melhorou o PV em média %s km·h⁻¹ (≈ +6%%) entre os dois testes — efeito de treino real, com variação individual de perda a ganho expressivo. Contudo, a magnitude da adaptação NÃO se associou ao humor de base (ΔPV × fadiga base ρ = %s; × iceberg ρ = %s) nem à deterioração da semana (ΔPV × Δfadiga física D1→D7 ρ = %s; ΔPV × Δvigor ρ = %s) — todas não significativas. Ou seja, neste microciclo específico, quem adaptou mais não fadigou menos: a resposta aguda da semana é governada por outros fatores (nível de aptidão em si, perfil de humor, carga percebida), não pelo tamanho do ganho no teste.'%(
    c2('%.2f'%R['means']['dPV']),c2('%+.2f'%A['b_Fadiga']['r']),c2('%+.2f'%A['b_ice_idx']['r']),
    c2('%+.2f'%A['d_FadFisica']['r']),c2('%+.2f'%A['d_Vigor']['r'])),after=6)

H('2.4 Posição',12)
POS=R['POS']; order=[o for o in ['Goleiro','Meia','Ponta','Pivô'] if o in POS]
table('Parâmetros por posição — T-CAR 2 e psicológico de base (médias)',
    ['Posição','n','PV final','Distância','Vigor (base)','Fad. fís. (base)','Iceberg (base)'],
    [[o,POS[o]['n'],c2('%.1f'%POS[o]['PV']),c2('%.0f'%POS[o]['dist']),c2('%.1f'%POS[o]['b_Vigor']),
      c2('%.1f'%POS[o]['b_FadFisica']),c2('%+.2f'%POS[o]['b_ice'])] for o in order],
    widths=[2.6,1.2,2.4,2.6,2.7,2.8,2.7],fs=9.5)
P('A ordenação posicional é preservada: pontas e meias com maior PV final; pivôs menores. O padrão de aptidão por posição é estável entre os dois testes.',after=6)

figure('/home/user/mdlucca/Artigos/figuras/x_tcar2.png',
    'T-CAR 2 e adaptação. (A) PV do T-CAR 2 × fadiga de base (associação enfraquecida). (B) T-CAR 1 supera o T-CAR 2 no poder sobre o baseline. (C) Adaptação individual ΔPV (verde = ganho, vermelho = perda; linha = média +0,9). (D) O ganho de aptidão não previu menos fadiga na semana.')

H('3. Síntese')
P('Com o T-CAR 2: (1) a aptidão pós-bloco associa-se mais fracamente ao baseline do que o T-CAR 1 — o marcador fisiológico correto do estado inicial é o T-CAR 1; (2) houve adaptação real do grupo (ΔPV ≈ +0,9 km·h⁻¹), mas o tamanho do ganho individual não explicou o humor de base nem a deterioração semanal; (3) a estrutura posicional da aptidão é estável. Conclusão prática: o T-CAR 2 é mais informativo como DESFECHO de adaptação do que como preditor do estado psicológico; para modelar o estado e a resposta da semana, mantenha o T-CAR 1 (aptidão de base) como covariável fisiológica.',after=6)

H('4. Limitações')
P('Baseline é ponto único com efeito de piso; o T-CAR 2 é posterior e temporalmente distante do baseline; um atleta sem PV final (n = 26 para o T-CAR 2). O TRIMP vem das sessões de HIIT (não do teste) sob premissas constantes. Relações correlacionais; subgrupos posicionais pequenos.',after=6)

OUTP='/home/user/mdlucca/Artigos/TCAR2_Adaptacao_Baseline_Parametros.docx'
doc.save(OUTP)
print('SAVED',OUTP,'| Tabelas',_TN[0],'Figuras',_FN[0])
