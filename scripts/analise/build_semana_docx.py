import warnings; warnings.filterwarnings('ignore')
import numpy as np, pandas as pd, json, os
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
SC='/tmp/claude-0/-home-user-mdlucca/e1dba24c-b1d7-5908-9106-f2f4aaf3f56a/scratchpad'
FIG='/home/user/mdlucca/Artigos/figuras'
R=json.load(open(f'{SC}/semana.json')); APP=json.load(open(f'{SC}/app_data.json'))
desc=R['desc']; cron=R['cron']; agu=R['agu']; agu_day=R['agu_day']; dec=R['dec']; der=R['der']; rank=R['rank']
def br(x): return str(x).replace('.',',')
CORE=[('Vigor','Vigor'),('Fadiga','Fadiga (BRUMS)'),('FadFisica','Fadiga física'),('FadMental','Fadiga mental')]

doc=Document(); s=doc.sections[0]; s.left_margin=Cm(3); s.top_margin=Cm(3); s.right_margin=Cm(2); s.bottom_margin=Cm(2)
n=doc.styles['Normal']; n.font.name='Times New Roman'; n.font.size=Pt(12); n.paragraph_format.line_spacing=1.5; n.paragraph_format.space_after=Pt(0)
FN='Times New Roman'
def _set(r,sz=12,b=False,it=False): r.font.name=FN; r.font.size=Pt(sz); r.bold=b; r.italic=it
def sec(num,txt):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(6); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    _set(p.add_run((f'{num} {txt}').upper() if '.' not in num else f'{num} {txt}'),12,b=True)
def sub(num,txt):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(4); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    _set(p.add_run(f'{num} {txt}'),12,b=True)
def P(t,indent=True,size=12,it=False,just=True):
    p=doc.add_paragraph(); _set(p.add_run(t),size,it=it); p.paragraph_format.line_spacing=1.5
    if indent: p.paragraph_format.first_line_indent=Cm(1.25)
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY if just else WD_ALIGN_PARAGRAPH.LEFT
def fig(name,num,titulo,w=6.0,fonte='Elaborado pelos autores (2026).'):
    cap=doc.add_paragraph(); cap.alignment=WD_ALIGN_PARAGRAPH.CENTER; cap.paragraph_format.space_before=Pt(8); cap.paragraph_format.line_spacing=1.0
    _set(cap.add_run(f'Figura {num} – {titulo}'),10)
    p=f'{FIG}/{name}.png'
    if os.path.exists(p): doc.add_picture(p,width=Inches(w)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    else: print('MISS',name)
    fp=doc.add_paragraph(); fp.alignment=WD_ALIGN_PARAGRAPH.CENTER; fp.paragraph_format.space_after=Pt(8); fp.paragraph_format.line_spacing=1.0
    _set(fp.add_run(f'Fonte: {fonte}'),10)
def _bd(t):
    tblPr=t._tbl.tblPr; b=OxmlElement('w:tblBorders')
    for e in ('top','bottom','insideH'):
        x=OxmlElement(f'w:{e}'); x.set(qn('w:val'),'single'); x.set(qn('w:sz'),'6'); x.set(qn('w:space'),'0'); x.set(qn('w:color'),'000000'); b.append(x)
    for e in ('left','right','insideV'):
        x=OxmlElement(f'w:{e}'); x.set(qn('w:val'),'none'); b.append(x)
    tblPr.append(b)
def table(num,titulo,headers,rows,fonte='Dados da pesquisa (2026).',fs=10):
    cap=doc.add_paragraph(); cap.alignment=WD_ALIGN_PARAGRAPH.LEFT; cap.paragraph_format.space_before=Pt(8); cap.paragraph_format.line_spacing=1.0
    _set(cap.add_run(f'Tabela {num} – {titulo}'),10)
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; _bd(t)
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.paragraphs[0].line_spacing=1.0; _set(c.paragraphs[0].add_run(str(h)),fs,b=True)
    for row in rows:
        cs=t.add_row().cells
        for i,v in enumerate(row):
            cs[i].paragraphs[0].line_spacing=1.0; _set(cs[i].paragraphs[0].add_run(str(v)),fs)
    fp=doc.add_paragraph(); fp.alignment=WD_ALIGN_PARAGRAPH.LEFT; fp.paragraph_format.space_after=Pt(8); fp.paragraph_format.line_spacing=1.0
    _set(fp.add_run(f'Fonte: {fonte}'),10)

# TÍTULO
tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
_set(tp.add_run('PERFIL DE HUMOR E COMPORTAMENTO DAS VARIÁVEIS AFETIVAS E DE FADIGA NA ÚLTIMA SEMANA DA PRÉ-TEMPORADA COMPETITIVA DE ATLETAS DE HANDEBOL'),14,b=True)
a2=doc.add_paragraph(); a2.alignment=WD_ALIGN_PARAGRAPH.CENTER; _set(a2.add_run('[Autoria e afiliação a preencher]'),11,it=True)
doc.add_paragraph()

# RESUMO
rp=doc.add_paragraph(); _set(rp.add_run('RESUMO'),12,b=True); rp.paragraph_format.line_spacing=1.0
P('Este estudo descreve e classifica o perfil de humor e o comportamento das variáveis afetivas e de fadiga de atletas de handebol ao longo da última semana da pré-temporada competitiva. Vinte e sete atletas responderam ao BRUMS-24 e a autoavaliações de fadiga física e mental duas vezes ao dia — a primeira coleta (pré) e a última (pós) — durante sete dias, totalizando 456 observações. Analisaram-se a caracterização sociodemográfica, a estatística descritiva com limites (amplitude), a comparação do Dia 1 (linha de base) ao Dia 7, a resposta pré→pós no grupo e por atleta, os percentuais e a variação das alterações, a sensibilidade das variáveis, além de derivadas, decomposição da variância com e sem ruído, suavização das curvas e pontos de inflexão. O humor deteriorou-se de forma concentrada no eixo energia–fadiga: o vigor caiu (dz = −1,09) e a fadiga física (dz = +1,69) e a fadiga do BRUMS (dz = +0,68) aumentaram do Dia 1 ao Dia 7, enquanto a fadiga mental variou pouco (dz = +0,33). A fadiga física foi a variável mais responsiva, tanto no efeito agudo (pré→pós) quanto no crônico (semana). As trajetórias apresentaram ponto de inflexão robusto em torno do Dia 4. Conclui-se que o custo do período é somático e acumulativo, sendo a fadiga física e o vigor os marcadores mais sensíveis para o monitoramento.',indent=False)
kw=doc.add_paragraph(); _set(kw.add_run('Palavras-chave: '),12,b=True); _set(kw.add_run('Estado de humor. BRUMS. Fadiga. Pré-temporada. Handebol.'),12); kw.paragraph_format.line_spacing=1.5

# 1 INTRODUÇÃO
sec('1','Introdução')
P('O monitoramento dos estados de humor é uma ferramenta prática, de baixo custo e não invasiva para acompanhar a resposta psicobiológica de atletas à carga de treinamento, com potencial de sinalizar desadaptações antes que se manifestem em queda de desempenho (TERRY; LANE; FOGARTY, 2003; ROHLFS et al., 2008). A Brunel Mood Scale (BRUMS) operacionaliza o construto em seis dimensões — tensão, depressão, raiva, vigor, fadiga e confusão — e é amplamente utilizada no esporte por conjugar brevidade e sensibilidade a variações agudas e crônicas. Atletas saudáveis tendem a exibir o perfil iceberg, com vigor elevado destacando-se acima das dimensões negativas (MORGAN, 1985); o achatamento ou a inversão desse perfil indica fadiga acumulada.')
P('A caracterização do comportamento do humor em diferentes escalas temporais — no conjunto da semana, na variação do início ao fim e na resposta aguda dentro de cada dia — e, simultaneamente, nos níveis coletivo e individual, fornece um panorama interpretável para calibrar a recuperação. Este estudo tem por objetivo avaliar e classificar o perfil de humor de atletas de handebol durante a última semana da pré-temporada competitiva, descrevendo o comportamento geral do grupo por variável, a variação do Dia 1 ao Dia 7, a resposta pré→pós no grupo e por atleta, e a sensibilidade relativa das variáveis, com aprofundamento por derivadas, decomposição com e sem ruído e pontos de inflexão.')

# 2 JUSTIFICATIVA
sec('2','Justificativa')
P('A última semana pré-competitiva é um período crítico do planejamento, no qual a comissão técnica busca dissipar a fadiga acumulada preservando as adaptações, de modo a maximizar a prontidão para competir. O manejo inadequado da carga nesse intervalo pode deteriorar o humor e sinalizar recuperação incompleta, com repercussões diretas sobre o desempenho; ainda assim, o comportamento do humor nessa janela raramente é caracterizado de forma estruturada nas equipes. Descrever esse comportamento de forma quantitativa — distinguindo o que é resposta real do que é ruído de medida, e identificando quais variáveis são mais sensíveis e em que momento a trajetória se inflete — oferece à comissão técnica subsídios objetivos para individualizar a recuperação e intervir no momento oportuno.')

# 3 MÉTODO
sec('3','Método')
sub('3.1','Delineamento e amostra')
ag=APP['socio']['agg']
P('Estudo descritivo de medidas repetidas, conduzido em condições ecológicas de treinamento, com 27 atletas de handebol do sexo masculino acompanhados ao longo de sete dias da última semana da pré-temporada. Cada atleta funcionou como sua própria referência. Os dados são apresentados de forma anonimizada (A01–A27).')
sub('3.2','Instrumentos')
P('O humor foi avaliado pela Brunel Mood Scale (BRUMS-24), composta por 24 itens em escala de cinco pontos, agrupados em seis subescalas (tensão, depressão, raiva, vigor, fadiga e confusão); a Perturbação Total do Humor (PTH) resume o perfil (PTH = negativas − vigor). A fadiga física e a fadiga mental foram avaliadas por escalas de autoavaliação de 0 a 10.')
sub('3.3','Procedimentos de coleta')
P('O instrumento foi autoaplicado por formulário eletrônico com carimbo de data/hora, duas vezes por dia: a primeira resposta do dia foi tomada como “pré” e a última como “pós”. No microciclo registraram-se 456 respostas válidas dos 27 atletas.')
sub('3.4','Análise estatística')
P('Computaram-se estatísticas descritivas (média, desvio-padrão, mediana, intervalo interquartílico e amplitude/limites). A variação do início ao fim da semana foi avaliada pelo contraste Dia 1 (linha de base) versus Dia 7, e a resposta aguda pela comparação pré versus pós, ambos por teste de Wilcoxon pareado com tamanho de efeito pareado (dz), no grupo e por atleta. A sensibilidade das variáveis foi ordenada pela magnitude do efeito. A forma da trajetória foi caracterizada por ajuste cúbico das médias diárias, do qual se extraíram velocidade, aceleração e ponto de inflexão; a variância foi decomposta em sinal da semana, traço (entre atletas) e ruído (resíduo), derivando-se o erro típico de medida (ETM) e a mudança mínima detectável (MDC95). As curvas foram suavizadas por regressão local (LOWESS) para separar sinal de ruído. Adotou-se α = 0,05; processamento em Python (pandas, SciPy, statsmodels).')

# 4 RESULTADOS
sec('4','Resultados')

sub('4.1','Caracterização sociodemográfica e antropométrica')
pc=APP['socio']['pos_counts']
table('1','Caracterização sociodemográfica e antropométrica da amostra (n = 27)',['Variável','Média ± DP (amplitude)'],[
 ['Idade (anos)',f"{ag['idade'][0]:.0f} ± {ag['idade'][1]:.0f} ({ag['idade'][2]}–{ag['idade'][3]})"],
 ['Experiência (anos)',f"{ag['exp'][0]:.1f} ± {ag['exp'][1]:.1f}".replace('.',',')],
 ['Estatura (cm)',f"{ag['estatura'][0]:.1f} ± {ag['estatura'][1]:.1f}".replace('.',',')],
 ['Massa corporal (kg)',f"{ag['massa'][0]:.1f} ± {ag['massa'][1]:.1f}".replace('.',',')],
 ['Gordura corporal (%)',f"{ag['pG'][0]:.1f} ± {ag['pG'][1]:.1f}".replace('.',',')],
 ['Pico de velocidade T-CAR (km/h)',f"{ag['PV'][0]:.2f} ± {ag['PV'][1]:.2f}".replace('.',',')],
])
P(f"A amostra distribuiu-se por posição em {pc.get('Armador',0)} armadores, {pc.get('Ala',0)} alas, {pc.get('Pivô',0)} pivôs e {pc.get('Goleiro',0)} goleiros, caracterizando um grupo homogêneo de atletas adultos de elite.")

sub('4.2','Perfil de humor geral do grupo e limites')
rows=[]
for v,lab in [('Vigor','Vigor'),('Fadiga','Fadiga (BRUMS)'),('FadFisica','Fadiga física'),('FadMental','Fadiga mental'),('TMD','PTH/TMD'),('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Confusao','Confusão')]:
    d=desc[v]; rows.append([lab,f"{d['mean']:.2f} ± {d['sd']:.2f}".replace('.',','),f"{d['med']:.1f} [{d['iqr']:.1f}]".replace('.',','),f"{d['mn']:.0f}–{d['mx']:.0f}",f"{d['amp']:.0f}".replace('.',',')])
table('2','Estatística descritiva e limites (amplitude) do humor e da fadiga na semana',['Variável','M ± DP','Mediana [IIQ]','Mín.–Máx.','Amplitude'],rows)
P('No conjunto da semana, o grupo preservou, em média, a configuração do perfil iceberg, com o vigor posicionando-se acima das subescalas negativas, que se mantiveram próximas ao piso (tensão, depressão, raiva e confusão com medianas baixas). A fadiga destacou-se como a dimensão negativa mais expressiva, e a elevada amplitude da PTH revela marcada heterogeneidade interindividual. Os limites (amplitude) indicam que as variáveis do eixo energia–fadiga percorrem quase toda a escala, ao passo que as negativas concentram-se no piso.')
fig('sem_f_radar','1','Forma do perfil de humor (escores z) no Dia 1 e no Dia 7',w=5.2)

sub('4.3','Comportamento das variáveis ao longo da semana')
P('A leitura conjunta das trajetórias (Figura 2) evidencia que o sinal temporal se concentra no eixo energia–fadiga: o vigor declina e a fadiga — física e do BRUMS — eleva-se de modo progressivo rumo ao Dia 7, ao passo que a fadiga mental e as subescalas negativas permanecem estáveis. As bandas de desvio-padrão amplas em todos os dias confirmam a heterogeneidade entre atletas.')
fig('sem_f_trajetoria','2','Comportamento semanal das quatro variáveis (M ± DP; curva suavizada por LOWESS pontilhada)')

sub('4.4','Do início ao fim da semana: Dia 1 (linha de base) versus Dia 7')
rows=[]
for v,lab in CORE:
    c=cron[v]; rows.append([lab,f"{c['d1']:.2f}".replace('.',','),f"{c['d7']:.2f}".replace('.',','),f"{c['delta']:+.2f}".replace('.',','),f"{c['pct']:+.0f}%",f"{c['dz']:+.2f}".replace('.',','),(f"{c['p']:.3f}".replace('.',',') if c['p']>=0.001 else '<0,001'),f"{c['worse']:.0f}%"])
table('3','Humor no Dia 1 (linha de base) e no Dia 7; comparação pareada (Wilcoxon)',['Variável','Dia 1','Dia 7','Δ','% mudança','dz','p','% que piora'],rows)
P('Da linha de base ao último dia, o vigor caiu acentuadamente (−44%; dz = −1,09; p < 0,001) e a fadiga física (+80%; dz = +1,69) e a fadiga do BRUMS (+73%; dz = +0,68) aumentaram de forma expressiva; a fadiga mental variou pouco e sem significância (+14%; dz = +0,33; p = 0,111). A proporção de atletas que pioraram foi máxima na fadiga física (95%) e no vigor (86%), confirmando que a deterioração é generalizada no grupo e restrita ao eixo energia–fadiga.')

sub('4.5','Resposta pré versus pós: intra-grupo e por dia')
rows=[]
for v,lab in CORE:
    a=agu[v]; rows.append([lab,f"{a['pre']:.2f}".replace('.',','),f"{a['pos']:.2f}".replace('.',','),f"{a['delta']:+.2f}".replace('.',','),f"{a['dz']:+.2f}".replace('.',','),(f"{a['p']:.4f}".replace('.',',') if a['p']>=0.001 else '<0,001')])
table('4','Resposta aguda pré→pós agregada na semana (intra-grupo, por atleta)',['Variável','Pré','Pós','Δ','dz','p'],rows)
P('Dentro dos dias, do primeiro ao último registro, o vigor caiu e a fadiga (física, do BRUMS e mental) aumentou, todos de forma significativa — a fadiga física com o maior efeito agudo (dz = +1,08). Esse padrão, de mesma direção que o observado na escala semanal, indica que a resposta aguda a cada dia é um análogo, em menor amplitude, do processo de acúmulo. A Figura 3 detalha a resposta pré→pós por dia, evidenciando que o contraste se mantém ao longo de toda a semana.')
fig('sem_f_prepos','3','Média pré e pós por dia para as quatro variáveis')

sub('4.6','Percentuais, variação e sensibilidade das variáveis')
rows=[]
for r in sorted(rank,key=lambda x:-x['dz_cr']):
    rows.append([r['lab'],f"{r['dz_ag']:.2f}".replace('.',','),f"{r['dz_cr']:.2f}".replace('.',','),f"{r['snr']:.2f}".replace('.',','),f"{r['pday']:.0f}%"])
table('5','Sensibilidade/responsividade das variáveis (ordenado pelo efeito crônico)',['Variável','|dz| agudo','|dz| crônico','SNR','% sinal'],rows)
P('A ordenação por magnitude de efeito revela uma hierarquia clara de responsividade: a fadiga física é a variável mais sensível, tanto no efeito agudo (|dz| = 1,08) quanto no crônico (|dz| = 1,69), seguida do vigor; a fadiga do BRUMS ocupa posição intermediária e a fadiga mental é a menos responsiva (|dz| crônico = 0,33; SNR < 1). A fadiga física reúne também a maior fração de sinal da semana (13%) e a melhor relação sinal-ruído (1,82), qualificando-a como o marcador de eleição para o monitoramento neste período.')
fig('sem_f_agudo_cronico','4','Tamanho de efeito (|dz|) agudo (pré→pós) versus crônico (Dia 1→Dia 7)',w=5.4)

sub('4.7','Limites, derivadas e pontos de inflexão')
rows=[]
for v,lab in CORE+[('TMD','PTH/TMD')]:
    e=der[v]; d=desc[v]; infl=(f"dia {e['infl']:.1f}".replace('.',',') if e['infl'] else 'n/d')
    turn=', '.join(f"{x:.1f}".replace('.',',') for x in e['turn']) if e['turn'] else '—'
    rows.append([lab,f"{d['mn']:.0f}–{d['mx']:.0f}",f"{d['amp']:.0f}",f"{e['vel_d1']:+.2f}".replace('.',','),f"{e['vel_d7']:+.2f}".replace('.',','),infl,turn])
table('6','Limites, derivadas (velocidade no D1 e D7) e ponto de inflexão por variável',['Variável','Mín.–Máx.','Amplitude','Vel. D1','Vel. D7','Inflexão','Viradas (vel = 0)'],rows)
P('As derivadas descrevem a forma da mudança. A velocidade da fadiga física e da fadiga do BRUMS é máxima nas extremidades (início e Dia 7), com um mínimo de deterioração no meio da semana; a aceleração troca de sinal uma única vez — o ponto de inflexão —, situado em torno do Dia 4 para o vigor, a fadiga física e a fadiga do BRUMS. A fadiga mental, por ser praticamente estável, não apresenta inflexão definida. O Dia 4 marca, assim, a transição da fase de acomodação para a fase de aprofundamento da fadiga rumo ao Dia 7 (Figura 5).')
fig('sem_f_derivadas','5','Velocidade (linha cheia) e aceleração (tracejada) da trajetória de cada variável; linha vertical = ponto de inflexão')

sub('4.8','Decomposição com e sem ruído; filtros e suavização')
rows=[]
for v,lab in CORE:
    d=dec[v]; rows.append([lab,f"{d['pday']:.0f}%",f"{d['ptrait']:.0f}%",f"{d['pnoise']:.0f}%",f"{d['etm']:.2f}".replace('.',','),f"{d['mdc']:.2f}".replace('.',','),f"{d['snr']:.2f}".replace('.',',')])
table('7','Decomposição da variância (sinal da semana, traço e ruído) e limiares de medida',['Variável','% Sinal','% Traço','% Ruído','ETM','MDC95','SNR'],rows)
P('A decomposição da variância mostra que o sinal da semana é a menor fração (1–13%); o grosso é traço estável (entre atletas) e ruído de medida. Removendo o ruído por suavização (LOWESS), a tendência subjacente torna-se nítida: a Figura 6 apresenta a partição da variância e a Figura 7 contrasta o índice-iceberg bruto (com ruído) com a curva suavizada (sem ruído), sobre a qual o ponto de inflexão (≈ Dia 4) é inequívoco. Em termos práticos, a leitura do grupo é robusta, mas a decisão individual exige médias de várias coletas, pois a oscilação de um único registro é dominada pelo ruído.')
fig('sem_f_decomp','6','Decomposição da variância por variável: sinal da semana, traço e ruído',w=5.4)
fig('sem_f_ruido','7','Índice-iceberg com ruído (observações e média diária) e sem ruído (LOWESS); linha vertical = inflexão',w=5.8)

sub('4.9','Perfis de humor ao longo da semana')
tab=R['perfil_tab']
rows=[[f"Dia {dd}",f"{tab['Iceberg'][str(dd)]:.0f}",f"{tab['Barbatana tubarão'][str(dd)]:.0f}",f"{tab['Superfície'][str(dd)]:.0f}",f"{tab['Iceberg invertido'][str(dd)]:.0f}",f"{tab['Everest invertido'][str(dd)]:.0f}"] for dd in [1,4,7]]
table('8','Distribuição dos seis perfis de humor (%) em dias-chave',['Dia','Iceberg','Barbatana','Superfície','Iceberg inv.','Everest inv.'],rows)
P('A classificação nos seis perfis mostra a migração do perfil saudável para o de fadiga: o iceberg cai do Dia 1 ao Dia 7, enquanto a barbatana de tubarão — fadiga em pico com vigor ainda preservado — cresce expressivamente. Os perfis de colapso afetivo permanecem raros, reforçando que o custo do período é somático, e não afetivo (Figura 8).')
fig('sem_f_perfis','8','Distribuição dos seis perfis de humor ao longo da semana',w=5.8)

sub('4.10','Efeito agudo e efeito crônico: comportamento individual')
P('A integração das escalas temporais distingue dois efeitos convergentes. O efeito agudo (pré→pós) reflete a resposta imediata ao esforço diário; o efeito crônico (Dia 1→Dia 7) reflete o acúmulo ao longo da semana. Para a fadiga física e o vigor, o efeito crônico supera o agudo (dz crônico 1,69 e 1,09 versus agudo 1,08 e 0,74), indicando acúmulo progressivo; para a fadiga mental, o efeito agudo (0,53) supera o crônico (0,33), sugerindo uma resposta pontual ao dia que não se consolida ao longo da semana. As trajetórias individuais (Figura 9) revelam ampla dispersão em torno da tendência coletiva: embora a direção seja compartilhada, a magnitude e o ponto de partida variam entre atletas, o que desaconselha pontos de corte normativos únicos e sustenta o referenciamento à linha de base individual.')
fig('sem_f_individual','9','Trajetórias individuais do vigor e da fadiga física ao longo da semana (linha em destaque = média do grupo)')

# 5 CONSIDERAÇÕES
sec('5','Considerações Finais')
P('O perfil de humor deteriorou-se de forma coerente e concentrada no eixo energia–fadiga, em duas escalas temporais convergentes: ao longo da semana (vigor ↓, fadiga ↑) e, em menor amplitude, dentro de cada dia. A fadiga física foi a variável mais sensível e responsiva, com maior fração de sinal e melhor relação sinal-ruído, seguida do vigor; a fadiga mental e as subescalas negativas foram pouco responsivas. As trajetórias inflectiram-se em torno do Dia 4, marcando a transição para o aprofundamento da fadiga. Recomenda-se acompanhar a fadiga física e o vigor por tendência individual e por médias de coletas, com atenção reforçada a partir do meio da semana. Como limitações, trata-se de estudo descritivo, sem grupo de comparação e restrito a um único período de uma equipe masculina.')

# REFERÊNCIAS
sec('REFERÊNCIAS','')
for r in [
 'MORGAN, W. P. Selected psychological factors limiting performance: a mental health model. In: CLARKE, D. H.; ECKERT, H. M. (Ed.). Limits of human performance. Champaign: Human Kinetics, 1985. p. 70-80.',
 'ROHLFS, I. C. P. M. et al. A Escala de Humor de Brunel (Brums): instrumento para detecção precoce da síndrome do excesso de treinamento. Revista Brasileira de Medicina do Esporte, v. 14, n. 3, p. 176-181, 2008.',
 'TERRY, P. C.; LANE, A. M.; FOGARTY, G. J. Construct validity of the Profile of Mood States — Adolescents for use with adults. Psychology of Sport and Exercise, v. 4, n. 2, p. 125-139, 2003.',
]:
    p=doc.add_paragraph(); _set(p.add_run(r),12); p.paragraph_format.line_spacing=1.0; p.paragraph_format.space_after=Pt(6); p.alignment=WD_ALIGN_PARAGRAPH.LEFT

doc.add_paragraph()
P('Nota de reprodutibilidade: análises processadas e auditadas em Python (pandas, SciPy, statsmodels); atletas anonimizados (A01–A27); nenhum dado bruto com identificação é distribuído.',size=9,it=True,indent=False)

OUTP='/home/user/mdlucca/Artigos/Perfil_Humor_Semana_PreTemporada_ABNT.docx'
doc.save(OUTP); print('SAVED',OUTP)
d2=Document(OUTP); print('paras',len(d2.paragraphs),'tables',len(d2.tables),'imgs',len(d2.inline_shapes))
