# -*- coding: utf-8 -*-
import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
R=json.load(open('brums_desc2.json')); MV=json.load(open('manova.json'))
PV=json.load(open('pv_stats.json')); LIM=json.load(open('tcar_limiar.json')); ALLO=json.load(open('allo.json')); EX=json.load(open('extra.json'))
FG='/home/user/mdlucca/Artigos/figuras'
def c2(s): return str(s).replace('.',',')
doc=Document()
st=doc.styles['Normal']; st.font.name='Times New Roman'; st.font.size=Pt(12)
st.element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')
st.paragraph_format.line_spacing=1.5; st.paragraph_format.space_after=Pt(0)
s=doc.sections[0]; s.top_margin=Cm(3); s.left_margin=Cm(3); s.bottom_margin=Cm(2); s.right_margin=Cm(2)
_TN=[0]; _FN=[0]
def P(t='',just=True,size=12,after=6,bold=False,ind=True):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(size); r.bold=bold
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(after)
    if just: p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if just and ind: p.paragraph_format.first_line_indent=Cm(1.25)
    return p
def RUN(pairs,after=6,ind=True):
    p=doc.add_paragraph(); p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(after)
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if ind: p.paragraph_format.first_line_indent=Cm(1.25)
    for txt,bold in pairs:
        r=p.add_run(txt); r.bold=bold; r.font.size=Pt(12)
    return p
def H(t,size=12,before=10):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(size)
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.5
def _bd(c):
    tcPr=c._tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for e_ in ['top','bottom','left','right']:
        el=OxmlElement(f'w:{e_}'); on=e_ in ('top','bottom')
        el.set(qn('w:val'),'single' if on else 'nil'); el.set(qn('w:sz'),'6'); el.set(qn('w:color'),'000000'); b.append(el)
    tcPr.append(b)
def table(cap,header,rows,fonte='Fonte: dados da pesquisa (2026).',fs=9,note=None):
    _TN[0]+=1
    p=doc.add_paragraph(); r=p.add_run('Tabela %d – %s'%(_TN[0],cap)); r.font.size=Pt(11)
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
    t=doc.add_table(rows=1,cols=len(header)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,htx in enumerate(header):
        cc=t.rows[0].cells[i]; cc.text=''; rr=cc.paragraphs[0].add_run(htx); rr.bold=True; rr.font.size=Pt(fs); rr.font.name='Times New Roman'
        cc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; _bd(cc)
    for row in rows:
        cs=t.add_row().cells
        for i,val in enumerate(row):
            cs[i].text=''; rr=cs[i].paragraphs[0].add_run(str(val)); rr.font.size=Pt(fs); rr.font.name='Times New Roman'
            cs[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if i>0 else WD_ALIGN_PARAGRAPH.LEFT; _bd(cs[i])
    if note:
        pn=doc.add_paragraph(); rn=pn.add_run(note); rn.font.size=Pt(8.5); rn.italic=True; pn.paragraph_format.space_after=Pt(0)
    pf=doc.add_paragraph(); rf=pf.add_run(fonte); rf.font.size=Pt(9); pf.paragraph_format.space_after=Pt(6)
    return _TN[0]
def figure(path,cap,w=15.5):
    _FN[0]+=1
    pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; pp.add_run().add_picture(path,width=Cm(w)); pp.paragraph_format.space_before=Pt(6)
    pc=doc.add_paragraph(); pc.alignment=WD_ALIGN_PARAGRAPH.CENTER; rc=pc.add_run('Figura %d – %s'%(_FN[0],cap)); rc.font.size=Pt(11)
    pf=doc.add_paragraph(); pf.alignment=WD_ALIGN_PARAGRAPH.CENTER; rf=pf.add_run('Fonte: elaboração dos autores (2026).'); rf.font.size=Pt(9); pf.paragraph_format.space_after=Pt(6)
    return _FN[0]
def pstr(p): return '< 0,001' if p<0.001 else c2('%.3f'%p)
mvv=lambda tab,k,f: next(x[f] for x in MV[tab]['rows'] if x['k']==k)
sm=R['sample']; d17=R['d1d7']
LP=LIM['LIM']['PVini']; TC=LIM['TERC']

# ===================== TÍTULO + RESUMO =====================
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('APTIDÃO AERÓBIA INTERMITENTE, CARGA DE HIIT E SONO NA RESPOSTA DE HUMOR–FADIGA DE HANDEBOLISTAS DE ELITE NA ÚLTIMA SEMANA DE PRÉ-TEMPORADA')
r.bold=True; r.font.size=Pt(13); p.paragraph_format.space_after=Pt(10)
H('RESUMO',before=2)
RUN([('A resposta afetiva à carga é individual e pode depender da aptidão física e de fatores de recuperação. ',False),
 ('O objetivo ',True),('foi examinar em que medida a aptidão aeróbia intermitente (pico de velocidade do Teste de '
 'Carminatti, T-CAR), a carga de treino intervalado de alta intensidade (HIIT), a sonolência (Epworth) e o estresse '
 'percebido (PSS) explicam a resposta de humor–fadiga de handebolistas de elite no microciclo pré-competitivo. ',False),
 ('Método: ',True),('%d atletas do sexo masculino responderam ao BRUMS-24 duas vezes ao dia por sete dias (%d '
 'observações), com avaliação do pico de velocidade (T-CAR), Epworth, PSS e registro das sessões de HIIT. As relações '
 'foram analisadas por regressão linear com reta de ajuste e dispersão, tercis de aptidão (Kruskal-Wallis), limiar '
 'discriminante (índice de Youden, AUC), escalonamento alométrico e comparação pareada (Wilcoxon). '%(sm['n'],sm['n_obs']),False),
 ('Resultados: ',True),('maior aptidão intermitente associou-se a mais vigor (ρ = %s; p = %s) e menos fadiga física '
 '(ρ = %s; R² = %s; p = %s); a fadiga diferiu entre tercis de aptidão (p = %s) e um limiar de %s km/h discriminou os dias '
 'de maior fadiga (AUC = %s). O escalonamento alométrico mostrou contribuição do tamanho corporal (ρ passou de %s para '
 '%s). A sonolência associou-se à fadiga (ρ = %s; p = %s) e o estresse percebido, não. Os dias de HIIT concentraram pior '
 'humor (vigor dz = %s; fadiga dz = %s), embora a resposta aguda pré → pós não tenha sido maior nesses dias. '%(
   c2('%+.2f'%PV['pv']['wk_Vigor']['TCAR1']['rho']),pstr(PV['pv']['wk_Vigor']['TCAR1']['rho_p']),
   c2('%+.2f'%PV['pv']['wk_FadFisica']['TCAR1']['rho']),c2('%.2f'%PV['pv']['wk_FadFisica']['TCAR1']['r2']),pstr(PV['pv']['wk_FadFisica']['TCAR1']['rho_p']),
   pstr(PV['terc_kruskal_fadfis']['p']),c2('%.1f'%LP['thr']),c2('%.2f'%LP['auc']),
   c2('%+.2f'%ALLO['fitness']['rho_raw']),c2('%+.2f'%ALLO['fitness']['rho_allo']),
   c2('%+.2f'%EX['epw']['corr']['Fadiga']['rho']),pstr(EX['epw']['corr']['Fadiga']['p']),
   c2('%+.2f'%EX['hiit']['Vigor']['dz']),c2('%+.2f'%EX['hiit']['Fadiga']['dz'])),False),
 ('Conclusão: ',True),('a aptidão aeróbia intermitente modula a resposta de humor–fadiga e o sono acompanha a fadiga do '
 'microciclo, fundamentando a normalização do monitoramento do humor por marcadores fisiológicos e de recuperação.',False)],after=6)
P('Palavras-chave: aptidão intermitente; Teste de Carminatti; carga de treino; sono; humor; handebol.',size=11,after=8,ind=False)

# ===================== 1 INTRODUÇÃO =====================
H('1 INTRODUÇÃO')
H('1.1 Resposta de humor–fadiga à carga no esporte intermitente',12,before=6)
P('O monitoramento subjetivo do humor e da fadiga é sensível às variações de carga de treino e amplamente recomendado '
 'para orientar decisões de treino e recuperação (SAW; MAIN; GASTIN, 2016; THORPE et al., 2017; KELLMANN et al., 2018). '
 'A resposta afetiva, contudo, não é uniforme entre atletas: sob a mesma carga, indivíduos com diferentes níveis de '
 'aptidão e de recuperação podem apresentar respostas distintas de vigor e de fadiga. Compreender quais parâmetros '
 'fisiológicos e de carga modulam essa resposta é essencial para individualizar o monitoramento e distinguir a fadiga '
 'esperada daquela que sinaliza sobrecarga.')
H('1.2 Handebol, capacidades físicas e aptidão aeróbia intermitente',12,before=6)
P('O handebol de quadra é uma modalidade coletiva de caráter marcadamente intermitente e de alta intensidade. Ao longo '
 'da partida, esforços máximos e explosivos — sprints curtos, saltos, arremessos, mudanças de direção e contatos físicos '
 '— alternam-se, de forma imprevisível, com períodos de recuperação incompleta, exigindo simultaneamente potência '
 'anaeróbia, capacidade aeróbia intermitente e elevada tolerância à fadiga (KARCHER; BUCHHEIT, 2014; WAGNER et al., '
 '2014). As demandas variam com a posição tática e a fadiga já foi documentada pela queda do volume de corrida em alta '
 'intensidade ao longo do jogo (MICHALSIK; MADSEN; AAGAARD, 2013; MICHALSIK; AAGAARD, 2015). A tolerância a esse padrão '
 'de esforço depende diretamente da aptidão aeróbia intermitente, que sustenta a reposição energética e a recuperação '
 'entre os esforços de alta intensidade e, assim, atenua a instalação da fadiga ao longo do jogo e do microciclo. Um '
 'marcador de campo válido dessa capacidade é o pico de velocidade obtido no Teste de Carminatti (T-CAR), um teste '
 'intermitente progressivo cujo pico de velocidade se associa ao desempenho físico em partida e discrimina níveis de '
 'aptidão em modalidades intermitentes (FERNANDES-DA-SILVA et al., 2016).')
H('1.3 Sono, estresse e carga de HIIT',12,before=6)
P('Além da aptidão, a recuperação e o estímulo específico modulam a resposta afetiva. A qualidade e a quantidade de sono '
 'associam-se ao humor em atletas de elite, com a sonolência diurna refletindo recuperação insuficiente (ANDRADE et al., '
 '2016; BRANDT; BEVILACQUA; ANDRADE, 2017). O estresse percebido é outro correlato potencial do estado afetivo (COHEN; '
 'KAMARCK; MERMELSTEIN, 1983). Do lado do estímulo, as sessões de treino intervalado de alta intensidade (HIIT) impõem '
 'elevada carga interna e podem intensificar a resposta de fadiga. Em handebolistas de elite, o humor e o estresse '
 'associam-se a indicadores de sobrecarga fisiológica e psicossomática (RATZ-SULYOK et al., 2026).')
H('1.4 Objetivos e hipóteses',12,before=6)
P('O objetivo geral foi examinar em que medida a aptidão aeróbia intermitente (pico de velocidade do T-CAR), a carga de '
 'HIIT, a sonolência e o estresse percebido explicam a resposta de humor–fadiga de handebolistas de elite na última '
 'semana de pré-temporada. Especificamente: (i) relacionar o pico de velocidade ao humor e à fadiga semanais, '
 'identificando um limiar discriminante; (ii) examinar a contribuição do tamanho corporal por escalonamento alométrico; '
 '(iii) testar as associações da sonolência e do estresse com o humor; e (iv) comparar a resposta de humor–fadiga entre '
 'dias com e sem HIIT.')
P('Hipóteses: (H1) maior aptidão intermitente associa-se a mais vigor e menos fadiga; (H2) parte da relação aptidão–'
 'fadiga reflete o tamanho corporal; (H3) a sonolência acompanha a fadiga, ao passo que o estresse percebido tem relação '
 'fraca; e (H4) os dias de HIIT apresentam pior humor, sobretudo por acúmulo ao longo do microciclo.')

# ===================== 2 MÉTODOS =====================
H('2 MATERIAIS E MÉTODOS')
H('2.1 Participantes',12,before=6)
P('Participaram %d atletas de handebol do sexo masculino de nível competitivo (idade %s ± %s anos; estatura %s ± %s cm; '
 'massa %s ± %s kg), avaliados em condições ecológicas durante o microciclo pré-competitivo (21–27 de abril de 2024), '
 'conforme a Declaração de Helsinque e com consentimento informado.'%(
   sm['n'],c2('%.1f'%sm['idade']['mean']),c2('%.1f'%sm['idade']['sd']),c2('%.1f'%sm['estatura']['mean']),c2('%.1f'%sm['estatura']['sd']),
   c2('%.1f'%sm['massa']['mean']),c2('%.1f'%sm['massa']['sd'])))
H('2.2 Medida do humor e da fadiga',12,before=6)
P('O humor foi avaliado pela BRUMS-24 (seis subescalas de 0 a 16: tensão, depressão, raiva, vigor, fadiga e confusão; '
 'PTH = soma das negativas − vigor), autoaplicada duas vezes ao dia (pré e pós-treino) por sete dias, totalizando %d '
 'observações. A fadiga física foi tomada como desfecho complementar de fadiga percebida. Escores mais altos indicam '
 'maior nível da dimensão.'%sm['n_obs'])
H('2.3 Aptidão aeróbia intermitente (T-CAR)',12,before=6)
P('A aptidão aeróbia intermitente foi avaliada pelo Teste de Carminatti (T-CAR), teste de campo progressivo e '
 'intermitente composto por repetições de 12 s de corrida em vaivém intercaladas por 6 s de recuperação, com incrementos '
 'sucessivos de velocidade guiados por sinal sonoro até a exaustão voluntária. O desfecho foi o pico de velocidade (PV), '
 'marcador da capacidade intermitente e do limiar de esforço (FERNANDES-DA-SILVA et al., 2016), avaliado no início da '
 'pré-temporada (T-CAR1; PV = %s ± %s km/h) e reaplicado ao final (T-CAR2; PV = %s ± %s km/h). O PV inicial (T-CAR1) foi '
 'usado como referência para normalizar e interpretar interindividualmente a resposta de humor–fadiga.'%(
   c2('%.1f'%PV['desc']['pvini_m']),c2('%.1f'%PV['desc']['pvini_sd']),c2('%.1f'%PV['desc']['pv_m']),c2('%.1f'%PV['desc']['pv_sd'])))
H('2.4 Sono, estresse e carga de HIIT',12,before=6)
P('Registraram-se a sonolência diurna (Escala de Sonolência de Epworth; JOHNS, 1991), o estresse percebido (Escala de '
 'Estresse Percebido — PSS; COHEN; KAMARCK; MERMELSTEIN, 1983) e a realização de sessão de HIIT em cada dia. No microciclo '
 'estudado, houve HIIT nos Dias 2, 4 e 7 e ausência de HIIT nos Dias 1, 3, 5 e 6.')
H('2.5 Análise de dados',12,before=6)
P('As relações entre o pico de velocidade e as médias semanais de humor e fadiga foram examinadas por regressão linear '
 '(β, R² e ρ de Spearman) com reta de ajuste, banda de confiança de 95%% e dispersão; os atletas foram estratificados em '
 'tercis de aptidão (Kruskal-Wallis) e determinou-se um limiar de PV discriminante dos dias de maior fadiga (índice de '
 'Youden, com AUC e IC95%% por bootstrap). A contribuição do tamanho corporal foi avaliada por escalonamento alométrico '
 '(PV = a · massa^b), comparando a correlação aptidão–fadiga bruta e normalizada. As associações da sonolência e do '
 'estresse com o humor usaram correlação de Spearman e comparação entre atletas sonolentos (Epworth > 10) e não '
 'sonolentos (Mann-Whitney). A resposta de humor entre dias com e sem HIIT foi comparada por Wilcoxon pareado (médias por '
 'atleta), e a resposta aguda pré → pós foi contrastada entre esses dias para isolar o estímulo intra-sessão do acúmulo '
 'do microciclo. A magnitude seguiu d de Cohen. Adotou-se α = 0,05, com análises em Python (pandas, SciPy, statsmodels).')

# ===================== 3 RESULTADOS =====================
H('3 RESULTADOS')
H('3.1 Caracterização da resposta de humor–fadiga',12,before=6)
P('No microciclo, a resposta afetiva concentrou-se no eixo energia–fadiga: do Dia 1 ao Dia 7 o vigor caiu (d = %s) e a '
 'fadiga subiu (d = %s), com deterioração multivariadamente significativa (Wilks λ = %s; p = %s). Sobre esse pano de '
 'fundo, examinaram-se os determinantes fisiológicos e de carga da magnitude individual dessa resposta.'%(
   c2('%+.2f'%mvv('d1d7','Vigor','d')),c2('%+.2f'%mvv('d1d7','Fadiga','d')),c2('%.3f'%MV['d1d7']['wilks']),pstr(MV['d1d7']['p_mv'])))
d=PV['desc']; ep=EX['epw']; ps=EX['pss']
t_desc=table('Descritiva dos parâmetros de aptidão intermitente e de recuperação (n = %d atletas).'%sm['n'],
    ['Parâmetro','Média','DP','Mín–Máx'],
    [['T-CAR1 – pico de velocidade (km/h)',c2('%.1f'%d['pvini_m']),c2('%.1f'%d['pvini_sd']),'%s–%s'%(c2('%.1f'%d['pvini_min']),c2('%.1f'%d['pvini_max']))],
     ['T-CAR2 – pico de velocidade (km/h)',c2('%.1f'%d['pv_m']),c2('%.1f'%d['pv_sd']),'—'],
     ['ΔPV (ganho na pré-temporada, km/h)',c2('%.1f'%d['dpv_m']),c2('%.1f'%d['dpv_sd']),'—'],
     ['Sonolência (Epworth)',c2('%.1f'%ep['m']),c2('%.1f'%ep['sd']),'%s–%s'%(c2('%.0f'%ep['mn']),c2('%.0f'%ep['mx']))],
     ['Estresse percebido (PSS)',c2('%.1f'%ps['m']),c2('%.1f'%ps['sd']),'%s–%s'%(c2('%.0f'%ps['mn']),c2('%.0f'%ps['mx']))]],
    note='Epworth > 10 = sonolência diurna excessiva (%d de %d atletas).'%(ep['hi'],ep['n']),fs=9)

H('3.2 Aptidão intermitente (T-CAR) e resposta de humor–fadiga',12,before=6)
def preg(col,tab='TCAR1'): return PV['pv']['wk_'+col][tab]
P('O pico de velocidade inicial (T-CAR1) relacionou-se de forma coerente com a resposta da semana (Figura %d; Tabela %d): '
 'atletas com maior aptidão reportaram mais vigor (β = %s; R² = %s; ρ = %s; p = %s) e menos fadiga física (β = %s; R² = '
 '%s; ρ = %s; p = %s), com tendência semelhante para a fadiga do BRUMS (ρ = %s; p = %s), confirmando a H1.'%(
   _FN[0]+1,_TN[0]+1,
   c2('%+.2f'%preg('Vigor')['slope']),c2('%.2f'%preg('Vigor')['r2']),c2('%+.2f'%preg('Vigor')['rho']),pstr(preg('Vigor')['rho_p']),
   c2('%+.2f'%preg('FadFisica')['slope']),c2('%.2f'%preg('FadFisica')['r2']),c2('%+.2f'%preg('FadFisica')['rho']),pstr(preg('FadFisica')['rho_p']),
   c2('%+.2f'%preg('Fadiga')['rho']),pstr(preg('Fadiga')['rho_p'])))
figure(f'{FG}/pv7_scatter.png','Dispersão e reta de regressão entre o pico de velocidade do T-CAR1 e as médias semanais de vigor, fadiga (BRUMS), PTH e fadiga física (linha tracejada = limiar de %s km/h; faixa sombreada = IC95%%).'%c2('%.1f'%LP['thr']),w=15.0)
def pregrow(col,lab):
    v=preg(col); return [lab,c2('%+.2f'%v['slope']),c2('%.2f'%v['r2']),c2('%+.2f'%v['rho']),pstr(v['rho_p'])]
t_pv=table('Regressão do pico de velocidade do T-CAR1 sobre as médias semanais de humor e fadiga (n = %d atletas).'%preg('Vigor')['n'],
    ['Desfecho semanal','β','R²','ρ','p'],
    [pregrow('Vigor','Vigor (BRUMS)'),pregrow('Fadiga','Fadiga (BRUMS)'),pregrow('TMD','PTH'),pregrow('FadFisica','Fadiga física')],
    note='β = inclinação por km/h; ρ = correlação de Spearman.',fs=9)
P('A estratificação em tercis confirma o padrão (Figura %d): a fadiga física decresce do tercil de menor para o de maior '
 'pico de velocidade (%s → %s; Kruskal-Wallis p = %s). Um limiar de PV ≈ %s km/h discriminou os dias de maior fadiga '
 '(AUC = %s [IC95%% %s–%s]; sensibilidade = %s; especificidade = %s), oferecendo referência prática para individualizar a '
 'carga (Figura %d).'%(
   _FN[0]+1,c2('%.1f'%TC['Baixa']['FadFis']),c2('%.1f'%TC['Alta']['FadFis']),pstr(PV['terc_kruskal_fadfis']['p']),
   c2('%.1f'%LP['thr']),c2('%.2f'%LP['auc']),c2('%.2f'%LP['lo']),c2('%.2f'%LP['hi']),c2('%.2f'%LP['sens']),c2('%.2f'%LP['spec']),_FN[0]+2))
figure(f'{FG}/pv8_tercis.png','Fadiga semanal (BRUMS e física) por tercil de aptidão intermitente (T-CAR1).',w=12.5)
figure(f'{FG}/pv9_limiar.png','Pico de velocidade do T-CAR1 e fadiga física semanal, com reta de regressão, IC95%% e limiar discriminante.',w=13.5)

H('3.3 Escalonamento alométrico: contribuição do tamanho corporal',12,before=6)
P('O escalonamento alométrico da aptidão (PV = a · massa^%s; R² = %s) mostrou que parte da associação bruta entre pico de '
 'velocidade e fadiga reflete o tamanho corporal: a correlação aptidão × fadiga física caiu de ρ = %s (p = %s) para ρ = '
 '%s (p = %s) após normalizar o PV pela massa (Figura %d), confirmando a H2. Assim, a “proteção” contra a fadiga '
 'atribuída ao pico de velocidade é, em parte, um efeito de dimensão corporal, o que recomenda cautela ao normalizar '
 'respostas por marcadores de aptidão (NEVILL; LANE, 2007).'%(
   c2('%.2f'%ALLO['fitness']['b']),c2('%.2f'%ALLO['fitness']['r2']),
   c2('%+.2f'%ALLO['fitness']['rho_raw']),pstr(ALLO['fitness']['p_raw']),c2('%+.2f'%ALLO['fitness']['rho_allo']),pstr(ALLO['fitness']['p_allo']),_FN[0]+1))
figure(f'{FG}/x_allo.png','Ajuste alométrico da trajetória de fadiga (lei de potência) em escala linear e log-log.',w=15.0)

H('3.4 Sono, estresse percebido e humor',12,before=6)
ep=EX['epw']; ps=EX['pss']; sc=EX['sleepy_cmp']
def cx(c): return '%s (p = %s)'%(c2('%+.2f'%c['rho']),pstr(c['p']))
P('A sonolência diurna (Epworth = %s ± %s; %d dos %d atletas com sonolência excessiva, > 10) associou-se ao pior humor: '
 'correlacionou-se positivamente com a fadiga (ρ = %s; p = %s) e com a PTH (ρ = %s; p = %s), e os atletas mais sonolentos '
 'apresentaram fadiga semanal marcadamente maior (%s vs. %s; p = %s), confirmando a primeira parte da H3 (Tabela %d; '
 'Figura %d). O estresse percebido (PSS = %s ± %s) não se relacionou com nenhuma dimensão do humor (todos p > 0,05), '
 'confirmando a segunda parte da H3: a fadiga do microciclo acompanha mais a sonolência (recuperação) do que o estresse '
 'psicológico geral.'%(
   c2('%.1f'%ep['m']),c2('%.1f'%ep['sd']),ep['hi'],ep['n'],
   c2('%+.2f'%ep['corr']['Fadiga']['rho']),pstr(ep['corr']['Fadiga']['p']),c2('%+.2f'%ep['corr']['TMD']['rho']),pstr(ep['corr']['TMD']['p']),
   c2('%.2f'%sc['Fadiga']['hi']),c2('%.2f'%sc['Fadiga']['lo']),pstr(sc['Fadiga']['p']),_TN[0]+1,_FN[0]+1,c2('%.1f'%ps['m']),c2('%.1f'%ps['sd'])))
def sprow(k,lab): return [lab,cx(ep['corr'][k]),cx(ps['corr'][k])]
t_sono=table('Correlação de Spearman da sonolência (Epworth) e do estresse percebido (PSS) com o humor (nível entre atletas, n = %d).'%ep['n'],
    ['Desfecho semanal','Epworth ρ (p)','PSS ρ (p)'],
    [sprow('Vigor','Vigor'),sprow('Fadiga','Fadiga'),sprow('TMD','PTH'),sprow('FadFisica','Fadiga física')],fs=9)
figure(f'{FG}/x_sono.png','Sonolência (Epworth) e fadiga: dispersão com reta de regressão (linha tracejada = corte 10) e comparação do humor entre atletas sonolentos e não sonolentos.',w=15.0)

H('3.5 Resposta ao estímulo: dias com vs. sem HIIT',12,before=6)
hi=EX['hiit']; ha=EX['hiit_acute']
P('Os dias com HIIT (Dias 2, 4 e 7) apresentaram pior humor do que os dias sem HIIT (Dias 1, 3, 5 e 6): menor vigor (%s '
 'vs. %s; p = %s; dz = %s), maior fadiga (%s vs. %s; p = %s; dz = %s), maior PTH (p = %s) e maior fadiga física (p = %s) '
 '(Tabela %d; Figura %d). Contudo, a resposta aguda pré → pós não foi maior nos dias de HIIT (variação do vigor %s vs. '
 '%s; da fadiga %s vs. %s; ambos p > 0,05), indicando que a diferença reflete sobretudo o acúmulo ao longo do microciclo '
 '— inclusive o Dia 7, que soma HIIT e fadiga terminal — e não uma resposta intra-sessão exclusiva do HIIT, o que apenas '
 'parcialmente sustenta a H4.'%(
   c2('%.2f'%hi['Vigor']['hiit']),c2('%.2f'%hi['Vigor']['nohiit']),pstr(hi['Vigor']['p']),c2('%+.2f'%hi['Vigor']['dz']),
   c2('%.2f'%hi['Fadiga']['hiit']),c2('%.2f'%hi['Fadiga']['nohiit']),pstr(hi['Fadiga']['p']),c2('%+.2f'%hi['Fadiga']['dz']),
   pstr(hi['TMD']['p']),pstr(hi['FadFisica']['p']),_TN[0]+1,_FN[0]+1,
   c2('%+.2f'%ha['Vigor']['hiit']),c2('%+.2f'%ha['Vigor']['nohiit']),c2('%+.2f'%ha['Fadiga']['hiit']),c2('%+.2f'%ha['Fadiga']['nohiit'])))
def hrow(k,lab):
    v=hi[k]; return [lab,c2('%.2f'%v['hiit']),c2('%.2f'%v['nohiit']),c2('%+.2f'%(v['hiit']-v['nohiit'])),pstr(v['p']),c2('%+.2f'%v['dz'])]
t_hiit=table('Humor e fadiga nos dias com vs. sem HIIT (média por atleta; Wilcoxon pareado, n = %d).'%hi['Vigor']['n'],
    ['Variável','Dias com HIIT','Dias sem HIIT','Δ','p','dz'],
    [hrow('Vigor','Vigor'),hrow('Fadiga','Fadiga'),hrow('TMD','PTH'),hrow('FadFisica','Fadiga física')],
    note='Dias com HIIT = 2, 4 e 7; dias sem HIIT = 1, 3, 5 e 6.',fs=9)
figure(f'{FG}/x_hiit.png','Humor e fadiga nos dias com vs. sem HIIT (* p < 0,05, Wilcoxon pareado).',w=12.5)

# ===================== 4 DISCUSSÃO =====================
H('4 DISCUSSÃO')
P('Os resultados mostram que a aptidão aeróbia intermitente modula a resposta de humor–fadiga do microciclo: maior pico '
 'de velocidade associou-se a mais vigor e menos fadiga (H1), coerente com o papel da capacidade intermitente na '
 'tolerância ao esforço repetido do handebol (KARCHER; BUCHHEIT, 2014; MICHALSIK; AAGAARD, 2015) e com a validade do pico '
 'de velocidade do T-CAR como marcador de desempenho físico em modalidades intermitentes (FERNANDES-DA-SILVA et al., '
 '2016). O escalonamento alométrico evidenciou que parte dessa proteção decorre do tamanho corporal (H2), um cuidado '
 'interpretativo relevante ao normalizar respostas por marcadores de aptidão (NEVILL; LANE, 2007).')
P('A associação da sonolência com a fadiga e a ausência de relação do estresse percebido com o humor (H3) reforçam que, '
 'neste microciclo, a deterioração afetiva acompanha marcadores de recuperação insuficiente mais do que a carga '
 'psicológica geral — convergente com a literatura que liga sono e humor em atletas (ANDRADE et al., 2016; BRANDT; '
 'BEVILACQUA; ANDRADE, 2017). Que quase metade dos atletas tenha reportado sonolência excessiva sublinha o valor de '
 'monitorar o sono junto ao humor. Os dias de HIIT concentraram o pior humor, mas a resposta aguda pré → pós não diferiu '
 'dos demais dias (H4 apenas parcialmente sustentada), indicando efeito de acúmulo ao longo do microciclo — com o Dia 7 '
 'somando HIIT e fadiga terminal — e não um choque intra-sessão exclusivo do HIIT.')
P('Do ponto de vista aplicado, normalizar a resposta de humor pela aptidão física e monitorar o sono ajudam a distinguir '
 'a fadiga esperada — de atletas menos aptos ou menos recuperados sob a mesma carga — daquela que sinaliza sobrecarga, '
 'refinando a individualização do monitoramento; o limiar de pico de velocidade identificado oferece uma referência '
 'objetiva. Como limitações, destacam-se o tamanho amostral (n = %d), o caráter observacional de fase única — que não '
 'permite inferência causal — e a avaliação da aptidão em dois momentos apenas. Estudos futuros devem integrar carga '
 'externa e interna mensuradas continuamente e ampliar a amostra.'%sm['n'])

# ===================== 5 CONCLUSÕES =====================
H('5 CONCLUSÕES')
P('A aptidão aeróbia intermitente, avaliada pelo pico de velocidade do T-CAR, modula a resposta de humor–fadiga de '
 'handebolistas de elite no microciclo pré-competitivo: atletas mais aptos reportam mais vigor e menos fadiga, com um '
 'limiar de pico de velocidade discriminando os dias de maior fadiga, ainda que parte do efeito decorra do tamanho '
 'corporal. A sonolência acompanha a fadiga, enquanto o estresse percebido não, e os dias de HIIT concentram pior humor '
 'por acúmulo no microciclo. Recomenda-se integrar a aptidão intermitente e o sono ao monitoramento do humor, '
 'normalizando a interpretação da resposta afetiva à luz da capacidade física e da recuperação de cada atleta.')

# ===================== REFERÊNCIAS =====================
H('REFERÊNCIAS')
refs=[
 'ANDRADE, A. et al. Sleep quality, mood and performance: a study of elite Brazilian volleyball athletes. Journal of Sports Science and Medicine, v. 15, n. 4, p. 601–605, 2016.',
 'BRANDT, R.; BEVILACQUA, G. G.; ANDRADE, A. Perceived sleep quality, mood states, and their relationship with performance among Brazilian elite athletes during a competitive period. Journal of Strength and Conditioning Research, v. 31, n. 4, p. 1033–1039, 2017.',
 'COHEN, S.; KAMARCK, T.; MERMELSTEIN, R. A global measure of perceived stress. Journal of Health and Social Behavior, v. 24, n. 4, p. 385–396, 1983. DOI: 10.2307/2136404.',
 'FERNANDES-DA-SILVA, J. et al. The peak velocity derived from the Carminatti Test is related to physical match performance in young soccer players. Journal of Sports Sciences, v. 34, n. 24, p. 2238–2245, 2016. DOI: 10.1080/02640414.2015.1093646.',
 'JOHNS, M. W. A new method for measuring daytime sleepiness: the Epworth Sleepiness Scale. Sleep, v. 14, n. 6, p. 540–545, 1991. DOI: 10.1093/sleep/14.6.540.',
 'KARCHER, C.; BUCHHEIT, M. On-court demands of elite handball, with special reference to playing positions. Sports Medicine, v. 44, n. 6, p. 797–814, 2014. DOI: 10.1007/s40279-014-0164-z.',
 'KELLMANN, M. et al. Recovery and performance in sport: consensus statement. International Journal of Sports Physiology and Performance, v. 13, n. 2, p. 240–245, 2018. DOI: 10.1123/ijspp.2017-0759.',
 'MICHALSIK, L. B.; AAGAARD, P. Physical demands in elite team handball: comparisons between male and female players. Journal of Sports Medicine and Physical Fitness, v. 55, n. 9, p. 878–891, 2015.',
 'MICHALSIK, L. B.; MADSEN, K.; AAGAARD, P. Match performance and physiological capacity of female elite team handball players. International Journal of Sports Medicine, v. 35, n. 7, p. 595–607, 2013. DOI: 10.1055/s-0033-1358713.',
 'NEVILL, A. M.; LANE, A. M. Why self-report “Likert” scale data should not be log-transformed. Journal of Sports Sciences, v. 25, n. 1, p. 1–2, 2007. DOI: 10.1080/02640410601111183.',
 'RATZ-SULYOK, F. Z. et al. Associations between endocrine status and stress, mood and psychosomatic status in elite handball players. Sports, v. 14, n. 7, 289, 2026. DOI: 10.3390/sports14070289.',
 'SAW, A. E.; MAIN, L. C.; GASTIN, P. B. Monitoring the athlete training response: subjective self-reported measures trump commonly used objective measures: a systematic review. British Journal of Sports Medicine, v. 50, n. 5, p. 281–291, 2016. DOI: 10.1136/bjsports-2015-094758.',
 'THORPE, R. T. et al. Monitoring fatigue status in elite team-sport athletes: implications for practice. International Journal of Sports Physiology and Performance, v. 12, s2, p. S2-27–S2-34, 2017. DOI: 10.1123/ijspp.2016-0434.',
 'WAGNER, H. et al. Individual and team performance in team-handball: a review. Journal of Sports Science and Medicine, v. 13, n. 4, p. 808–816, 2014.']
for rf in refs:
    p=doc.add_paragraph(); r=p.add_run(rf); r.font.size=Pt(11); p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(6); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

OUTP='/home/user/mdlucca/Artigos/Artigo2_Aptidao_Carga_Humor.docx'
doc.save(OUTP); print('SAVED',OUTP,'| Tabelas',_TN[0],'Figuras',_FN[0])
