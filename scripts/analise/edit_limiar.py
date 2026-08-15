# -*- coding: utf-8 -*-
import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
DOC='/home/user/mdlucca/Artigos/PERFIL_HUMOR_COMPLETO.docx'
FIG='/home/user/mdlucca/Artigos/figuras/x_limiar.png'
R=json.load(open('tcar_limiar.json'))
def c2(s): return str(s).replace('.',',')
d=Document(DOC)

def _fmt(p,size=12,just=True,bold=False,center=False):
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(6)
    if center: p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    elif just: p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    for r in p.runs: r.font.name='Times New Roman'; r.font.size=Pt(size); r.bold=bold
def after(el,text='',**kw):
    p=d.add_paragraph(text); el.addnext(p._p); _fmt(p,**kw); return p._p
def head(el,text): return after(el,text,bold=True,just=False)
def _bd(cell):
    tcPr=cell._tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for edge in ['top','bottom','left','right']:
        e=OxmlElement(f'w:{edge}'); on=edge in ('top','bottom')
        e.set(qn('w:val'),'single' if on else 'nil'); e.set(qn('w:sz'),'6'); e.set(qn('w:color'),'000000'); b.append(e)
    tcPr.append(b)
def table_after(el,caption,header,rows,fonte='Fonte: dados da pesquisa (2026).',fs=9.5):
    cap=d.add_paragraph(caption); el.addnext(cap._p); _fmt(cap,size=11,just=False); prev=cap._p
    t=d.add_table(rows=1,cols=len(header)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    prev.addnext(t._tbl)
    for i,htxt in enumerate(header):
        cc=t.rows[0].cells[i]; cc.text=''; rr=cc.paragraphs[0].add_run(htxt); rr.bold=True; rr.font.size=Pt(fs); rr.font.name='Times New Roman'
        cc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; _bd(cc)
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=''; rr=cells[i].paragraphs[0].add_run(str(val)); rr.font.size=Pt(fs); rr.font.name='Times New Roman'
            cells[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if i>0 else WD_ALIGN_PARAGRAPH.LEFT; _bd(cells[i])
    fon=d.add_paragraph(fonte); t._tbl.addnext(fon._p); _fmt(fon,size=10,just=False); return fon._p
def fig_after(el,path,caption,w=16.0):
    pimg=d.add_paragraph(); el.addnext(pimg._p); pimg.alignment=WD_ALIGN_PARAGRAPH.CENTER; pimg.add_run().add_picture(path,width=Cm(w))
    cap=d.add_paragraph(caption); pimg._p.addnext(cap._p); _fmt(cap,size=11,just=False,center=True)
    fon=d.add_paragraph('Fonte: elaboração dos autores (2026).'); cap._p.addnext(fon._p); _fmt(fon,size=10,just=False,center=True); return fon._p

def find(prefix):
    for p in d.paragraphs:
        if p.text.strip().startswith(prefix): return p
    raise KeyError(prefix)

WKi=R['WK']['PVini']; WKf=R['WK']['PVfin']; L1=R['LIM']['PVini']; L2=R['LIM']['PVfin']; TE=R['TERC']
p6=find('6 SÍNTESE DOS PRINCIPAIS ACHADOS'); anchor=p6._p.getprevious()

anchor=head(anchor,'5.12 Ajuste do limiar do pico de velocidade e regressão aptidão–fadiga (semana e por dia)')
anchor=after(anchor,
 'Para responder diretamente se os atletas de maior aptidão e potência aeróbica são os menos fadigados, ajustou-se a '
 'relação entre o pico de velocidade (PV) do T-CAR e a fadiga física — no acumulado da semana e dia a dia — e estimou-se '
 'um limiar de PV por regressão logística. Apresentam-se os dois testes: o PV inicial (T-CAR 1, aptidão de entrada do '
 'microciclo) e o PV final (T-CAR 2, reavaliação pós-bloco).')
anchor=after(anchor,
 'No acumulado semanal, a fadiga física decresce linearmente com o PV de entrada: fadiga = %s %s·PV (ρ = %s; p = %s; '
 'R² = %s; Figura 20A). Cada km·h⁻¹ adicional de velocidade de pico associa-se a cerca de %s ponto a menos de fadiga '
 'física média na semana — ou seja, os atletas de maior potência aeróbica são, de fato, os menos fadigados. A relação é '
 'consistente dia a dia: a correlação PV–fadiga é negativa nos sete dias e significativa nos dias 2, 3 e 5 (ρ entre −0,48 '
 'e −0,56; Figura 20B; Tabela 13). Com o PV do T-CAR 2, as mesmas associações são fracas e não significativas (semana '
 'ρ = %s), coerente com o nulo da modelagem preditiva (§5.10): o parâmetro operante para a fadiga desta semana é a '
 'aptidão de entrada, não a reavaliação posterior.'%(
   c2('%.1f'%WKi['intercept']),c2('%+.2f'%WKi['slope']),c2('%+.2f'%WKi['rho']),c2('%.3f'%WKi['p']),
   c2('%.2f'%WKi['r2']),c2('%.1f'%abs(WKi['slope'])),c2('%+.2f'%WKf['rho'])))
anchor=after(anchor,
 'O ajuste logístico estima um limiar de PV que separa os dias de fadiga alta: em torno de %s km·h⁻¹ para o T-CAR 1 '
 '(AUC = %s; IC95%% %s–%s; sensibilidade %s; especificidade %s; OR = %s por km·h⁻¹), abaixo do qual a probabilidade de '
 'fadiga alta cresce acentuadamente (Figura 20C; Tabela 14). O T-CAR 2 não produz limiar utilizável (AUC = %s; '
 'especificidade %s). Estratificando por tercis de aptidão (Figura 20D), o terço menos apto (PV %s) acumula fadiga média '
 'de %s, contra %s do terço mais apto (PV %s) — um gradiente de %s ponto.'%(
   c2('%.1f'%L1['thr']),c2('%.2f'%L1['auc']),c2('%.2f'%L1['lo']),c2('%.2f'%L1['hi']),c2('%.2f'%L1['sens']),
   c2('%.2f'%L1['spec']),c2('%.2f'%L1['OR']),c2('%.2f'%L2['auc']),c2('%.2f'%L2['spec']),
   c2('%.1f'%TE['Baixa']['PV']),c2('%.1f'%TE['Baixa']['FadFis']),c2('%.1f'%TE['Alta']['FadFis']),
   c2('%.1f'%TE['Alta']['PV']),c2('%.1f'%(TE['Baixa']['FadFis']-TE['Alta']['FadFis']))))
anchor=after(anchor,
 'Interpretação à luz da literatura. Os resultados alinham-se ao entendimento de que a capacidade aeróbia intermitente '
 'sustenta a recuperação entre esforços de alta intensidade e, portanto, a resistência à fadiga nas modalidades '
 'coletivas (BANGSBO; IAIA; KRUSTRUP, 2008; MICHALSIK; MADSEN; AAGAARD, 2013). O PV do teste de Carminatti, por integrar '
 'potência aeróbia, economia de corrida e capacidade de repetir esforços, é um índice ecológico dessa aptidão (CHAMARI '
 'et al., 2005). Transformar esse índice contínuo em uma regra de decisão por meio de regressão logística e de um ponto '
 'de corte (limiar de Youden) é procedimento consolidado no rastreamento esportivo, desde que ancorado na linha de base '
 'apropriada (SAW; MAIN; GASTIN, 2016). A leitura aplicada é clara: atletas com PV de entrada abaixo de ~%s km·h⁻¹ '
 'concentram os dias de fadiga alta e merecem individualização de carga e recuperação.'%(c2('%.1f'%L1['thr'])))

# Tabela 13 – regressão por dia
rows13=[['Semana (acumulado)',c2('%+.2f'%WKi['slope']),c2('%+.2f'%WKi['rho']),c2('%.3f'%WKi['p']),c2('%+.2f'%WKf['rho']),c2('%.3f'%WKf['rho_p'])]]
for dd in range(1,8):
    a=R['DAY']['PVini'][str(dd)]; b=R['DAY']['PVfin'][str(dd)]
    rows13.append(['Dia %d'%dd,c2('%+.2f'%a['slope']),c2('%+.2f'%a['rho']),c2('%.3f'%a['rho_p']),c2('%+.2f'%b['rho']),c2('%.3f'%b['rho_p'])])
anchor=table_after(anchor,'Tabela 13 – Regressão aptidão–fadiga por dia e no acumulado semanal: inclinação e correlação (T-CAR 1 e T-CAR 2).',
    ['Recorte','β (T-CAR 1)','ρ (T-CAR 1)','p (T-CAR 1)','ρ (T-CAR 2)','p (T-CAR 2)'],rows13)
# Tabela 14 – limiar
anchor=table_after(anchor,'Tabela 14 – Ajuste do limiar logístico do PV para o dia de fadiga física alta.',
    ['Parâmetro','T-CAR 1 (PV inicial)','T-CAR 2 (PV final)'],
    [['AUC [IC95%]','%s [%s–%s]'%(c2('%.2f'%L1['auc']),c2('%.2f'%L1['lo']),c2('%.2f'%L1['hi'])),
                    '%s [%s–%s]'%(c2('%.2f'%L2['auc']),c2('%.2f'%L2['lo']),c2('%.2f'%L2['hi']))],
     ['Limiar de Youden (km·h⁻¹)',c2('%.1f'%L1['thr']),c2('%.1f'%L2['thr'])],
     ['Sensibilidade / Especificidade','%s / %s'%(c2('%.2f'%L1['sens']),c2('%.2f'%L1['spec'])),'%s / %s'%(c2('%.2f'%L2['sens']),c2('%.2f'%L2['spec']))],
     ['OR por km·h⁻¹',c2('%.2f'%L1['OR']),c2('%.2f'%L2['OR'])]])
anchor=fig_after(anchor,FIG,
 'Figura 20 – Ajuste do limiar do PV e regressão aptidão–fadiga. (A) Regressão da fadiga física semanal sobre o PV de '
 'entrada (T-CAR 1): quanto maior a potência aeróbica, menor a fadiga. (B) A correlação PV–fadiga é negativa em todos os '
 'dias (T-CAR 1), significativa nos dias 2, 3 e 5 (*); o T-CAR 2 é fraco. (C) Ajuste logístico: o T-CAR 1 define um limiar '
 'utilizável (~14,9 km·h⁻¹); o T-CAR 2, não. (D) O terço menos apto acumula mais fadiga que o mais apto.')

# Referência nova (alfabética: BANGSBO antes de BRANDT)
brandt=find('BRANDT, R.')
newp=d.add_paragraph('BANGSBO, J.; IAIA, F. M.; KRUSTRUP, P. The Yo-Yo intermittent recovery test: a useful tool for evaluation of physical performance in intermittent sports. Sports Medicine, v. 38, n. 1, p. 37–51, 2008. DOI: 10.2165/00007256-200838010-00004.')
brandt._p.addprevious(newp._p); _fmt(newp,size=12,just=False)

d.save(DOC)
print('SAVED',DOC)
