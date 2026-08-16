# -*- coding: utf-8 -*-
import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
R=json.load(open('brums_desc2.json')); RC=json.load(open('rci6.json')); STAT=json.load(open('brums_stats3.json'))
S4=json.load(open('brums_stats4.json')); MV=json.load(open('manova.json')); PHJ=json.load(open('posthoc.json'))
ALLO=json.load(open('allo.json')); ROC=json.load(open('roc.json')); DEN=json.load(open('denoise.json'))
CRS=json.load(open('cross.json')); EX=json.load(open('extra.json'))
FG='/home/user/mdlucca/Artigos/figuras'
def c2(s): return str(s).replace('.',',')
doc=Document()
stl=doc.styles['Normal']; stl.font.name='Times New Roman'; stl.font.size=Pt(12)
stl.element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')
stl.paragraph_format.line_spacing=1.5; stl.paragraph_format.space_after=Pt(0)
sec=doc.sections[0]; sec.top_margin=Cm(3); sec.left_margin=Cm(3); sec.bottom_margin=Cm(2); sec.right_margin=Cm(2)
_TN=[0]; _FN=[0]
def P(t='',just=True,size=12,after=6,ind=True):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(size)
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(after)
    if just: p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if just and ind: p.paragraph_format.first_line_indent=Cm(1.25)
    return p
def H(t,size=12,before=12):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(size)
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.5
def _bd(c):
    tcPr=c._tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for e_ in ['top','bottom','left','right']:
        el=OxmlElement(f'w:{e_}'); on=e_ in ('top','bottom')
        el.set(qn('w:val'),'single' if on else 'nil'); el.set(qn('w:sz'),'6'); el.set(qn('w:color'),'000000'); b.append(el)
    tcPr.append(b)
def table(cap,header,rows,fs=9,note=None):
    _TN[0]+=1
    p=doc.add_paragraph(); r=p.add_run('Tabela S%d – %s'%(_TN[0],cap)); r.font.size=Pt(11)
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
    t=doc.add_table(rows=1,cols=len(header)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,htx in enumerate(header):
        cc=t.rows[0].cells[i]; cc.text=''; rr=cc.paragraphs[0].add_run(htx); rr.bold=True; rr.font.size=Pt(fs); rr.font.name='Times New Roman'
        cc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; _bd(cc)
    for row in rows:
        cs=t.add_row().cells
        for i,val in enumerate(row):
            cs[i].text=''; rr=cs[i].paragraphs[0].add_run(str(val)); rr.font.size=Pt(fs); rr.font.name='Times New Roman'
            cs[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if i>0 else WD_ALIGN_PARAGRAPH.LEFT; _bd(cs[i])
    if note:
        pn=doc.add_paragraph(); rn=pn.add_run(note); rn.font.size=Pt(8.5); rn.italic=True; pn.paragraph_format.space_after=Pt(4)
    return _TN[0]
def figure(path,cap,w=15.5):
    _FN[0]+=1
    pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; pp.add_run().add_picture(path,width=Cm(w)); pp.paragraph_format.space_before=Pt(6)
    pc=doc.add_paragraph(); pc.alignment=WD_ALIGN_PARAGRAPH.CENTER; rc=pc.add_run('Figura S%d – %s'%(_FN[0],cap)); rc.font.size=Pt(11)
    pc.paragraph_format.space_after=Pt(8)
    return _FN[0]
def pstr(p): return '< 0,001' if p<0.001 else c2('%.3f'%p)

sm=R['sample']; SENS=S4['sens']; FR=S4['friedman']; CV=S4['cv']; PT=S4['prof_trans']; PCNT=S4['prof_counts']

# ===== CAPA =====
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('MATERIAL SUPLEMENTAR'); r.bold=True; r.font.size=Pt(14); p.paragraph_format.space_after=Pt(2)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Perfil de humor de atletas de handebol de elite na última semana de pré-temporada'); r.italic=True; r.font.size=Pt(12)
p.paragraph_format.space_after=Pt(10)
P('Este material reúne as análises complementares e as figuras detalhadas referidas no artigo principal: análise '
 'exploratória das distribuições, comparações complementares, análises individuais, o pós-teste temporal, os ajustes '
 'alométricos, as derivadas da trajetória e a decomposição sinal–ruído.',ind=False)

# ===== A. EXPLORATÓRIA =====
H('A. Análise exploratória das distribuições')
figure(f'{FG}/xb3_hist.png','Distribuição de frequências das seis dimensões do BRUMS (linha tracejada = mediana).')
figure(f'{FG}/xb3_box.png','Diagramas de caixa das seis dimensões por dia.')
figure(f'{FG}/xb4_splom.png','Matriz de dispersão entre as seis dimensões (médias semanais por atleta).',w=14.0)
figure(f'{FG}/xb2_traj.png','Trajetória das dimensões ao longo da semana (médias diárias; áreas sombreadas = IC95%).')
SH=STAT['shapiro']; IC=STAT['icc']
def nrow(k,lab):
    s=SH[k]; i=IC[k]; return [lab,c2('%.3f'%s['W']),'< 0,001' if s['p']<0.001 else c2('%.3f'%s['p']),c2('%+.2f'%s['skew']),c2('%.2f'%i['icc1']),c2('%.2f'%i['icck']),i['cls']]
table('Normalidade (Shapiro-Wilk), assimetria e consistência (ICC) por dimensão.',
    ['Dimensão','W','p','Assimetria','ICC(2,1)','ICC(2,k)','Consistência'],
    [nrow(k,l) for k,l in [('Vigor','Vigor'),('Fadiga','Fadiga'),('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Confusao','Confusão')]],fs=9)
def cvrow(k,lab):
    v=CV[k]; return [lab,c2('%.0f'%v['inter']),c2('%.0f'%v['intra'])]
table('Coeficiente de variação entre atletas e intraindividual, por dimensão.',
    ['Dimensão','CV entre atletas (%)','CV intraindividual (%)'],
    [cvrow(k,l) for k,l in [('Vigor','Vigor'),('Fadiga','Fadiga'),('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Confusao','Confusão')]],
    note='PTH excluída por assumir valores ≤ 0.',fs=9)

# ===== B. COMPARAÇÕES COMPLEMENTARES =====
H('B. Comparações complementares e sensibilidade')
def mvrow(x):
    sig='*' if x['p']<0.05 else ''
    return [x['lab'],c2('%.1f'%x['m1']),c2('%.1f'%x['s1']),c2('%.1f'%x['m2']),c2('%.1f'%x['s2']),c2('%.2f'%x['F']),pstr(x['p'])+sig,c2('%+.2f'%x['d']),c2('%.3f'%x['eta'])]
mv=MV['prepos']
table('Resposta aguda pré → pós das seis dimensões em escores T (MANOVA de medidas repetidas; n = %d).'%mv['n'],
    ['Dimensão','Pré M','Pré DP','Pós M','Pós DP','F','p','d','η²ₚ'],[mvrow(x) for x in mv['rows']],
    note='Wilks λ = %s; F(%d,%d) = %s; p = %s; η²ₚ = %s. * p < 0,05.'%(c2('%.3f'%mv['wilks']),mv['df1'],mv['df2'],c2('%.2f'%mv['Fmv']),pstr(mv['p_mv']),c2('%.2f'%mv['eta_mv'])),fs=8.5)
figure(f'{FG}/xb5_profile_prepos.png','Perfil de humor em escores T nos momentos pré e pós-treino.',w=14.0)
figure(f'{FG}/xb4_dumbbell.png','Comparação Dia 1 versus Dia 7 dos escores de todas as variáveis.',w=14.0)
snames={'Vigor':'Vigor','Fadiga':'Fadiga','TMD':'PTH','Tensao':'Tensão','Depressao':'Depressão','Raiva':'Raiva','Confusao':'Confusão'}
sord=sorted(snames,key=lambda k:-SENS[k]['absdz'])
def frow2(k,rank):
    fr=FR[k]; return [rank,snames[k],c2('%+.2f'%SENS[k]['dz']),c2('%.1f'%fr['chi']),pstr(fr['p']),c2('%.2f'%fr['W']),'sim' if fr['p']<0.05 else 'não']
table('Sensibilidade à variação semanal e teste de Friedman (7 dias), ordenadas por |dz| (Dia 1 → Dia 7).',
    ['Ordem','Variável','dz (D1→D7)','Friedman χ²','p','W de Kendall','Difere entre dias'],
    [frow2(k,i+1) for i,k in enumerate(sord)],note='Casos completos (n = %d).'%FR['Vigor']['n'],fs=8.5)
PA=STAT['pairs']
table('Correlação de Spearman entre as dimensões do BRUMS (nível entre atletas, n = %d).'%sm['n'],
    ['Par de dimensões','ρ','p'],[['%s × %s'%(pr['a'],pr['b']),c2('%+.2f'%pr['rho']),pstr(pr['p'])] for pr in PA],fs=9)
for k,fn,lab in [('Vigor','xb4_v_Vigor.png','Vigor'),('Fadiga','xb4_v_Fadiga.png','Fadiga'),('Tensao','xb4_v_Tensao.png','Tensão'),('Depressao','xb4_v_Depressao.png','Depressão'),('Raiva','xb4_v_Raiva.png','Raiva'),('Confusao','xb4_v_Confusao.png','Confusão')]:
    figure(f'{FG}/{fn}','%s ao longo da semana: médias diárias (banda = IC95%%), diagramas de caixa por dia e efeito Dia 1 → Dia 7.'%lab,w=12.0)

# ===== C. NÍVEL INDIVIDUAL =====
H('C. Análise no nível do atleta')
figure(f'{FG}/xb4_spaghetti.png','Trajetórias individuais de vigor e fadiga (linhas cinzas = atletas; linha colorida = média do grupo).')
ABBR={'Iceberg':'Iceberg','Everest invertido':'Everest inv.','Iceberg invertido':'Iceberg inv.','Submerso':'Submerso','Barbatana tubarão':'Barbatana','Superfície':'Superfície'}
def tag(d1,d7):
    if d7=='Iceberg' and d1!='Iceberg': return 'melhora'
    if d7=='Barbatana tubarão' and d1!='Barbatana tubarão': return 'fadiga'
    if d1==d7: return 'estável'
    return 'transição'
rows_tr=[[aid,ABBR.get(PT[aid]['D1'],PT[aid]['D1']),ABBR.get(PT[aid]['D7'],PT[aid]['D7']),tag(PT[aid]['D1'],PT[aid]['D7'])] for aid in sorted(PT)]
table('Perfil de humor de cada atleta no Dia 1 e no Dia 7 e classificação da transição (n = %d).'%PCNT['n'],
    ['Atleta','Perfil no Dia 1','Perfil no Dia 7','Transição'],rows_tr,fs=8.5)

# ===== D. DINÂMICA TEMPORAL (pós-teste, cruzamento, derivadas) =====
H('D. Dinâmica temporal: pós-teste, cruzamento e derivadas')
def emmc(v,d): return c2('%.2f'%PHJ[v]['emm'][str(d)])
def sig1(v,d): return '' if d==1 or PHJ[v]['pairs']['1_%d'%d]['ptukey']>=0.05 else '*'
def phrow(d): return ['Dia %d'%d,emmc('Vigor',d)+sig1('Vigor',d),emmc('Fadiga',d)+sig1('Fadiga',d),emmc('FadFisica',d)+sig1('FadFisica',d)]
table('Médias marginais estimadas (modelo misto) por dia e pós-teste vs. Dia 1.',
    ['Dia','Vigor','Fadiga (BRUMS)','Fadiga física'],[phrow(d) for d in range(1,8)],
    note='* diferença significativa vs. Dia 1 (Tukey, p < 0,05).',fs=9)
figure(f'{FG}/ph_emm.png','Trajetória diária (médias marginais do modelo misto) com pós-teste vs. Dia 1 e destaque dos dias de maior variação.')
figure(f'{FG}/x_cross.png','Cruzamento vigor × fadiga por dia e momento (pré → pós dentro de cada dia).',w=15.0)
figure(f'{FG}/x_deriv_exp.png','Velocidade (barras) e aceleração (linha) diárias da trajetória de cada variável.',w=15.0)
figure(f'{FG}/x_vardecomp.png','Decomposição da variância de cada variável em componentes entre atletas, dia e resíduo.',w=13.0)
def ldrow(k,lab):
    d=CRS['dec'][k]; return [lab,c2('%.2f'%d['amp_sig']),c2('%.2f'%d['amp_intra']),'D%d'%d['pkvel_day'],c2('%+.2f'%d['pkvel_val']),c2('%.0f'%d['vc']['ath']),c2('%.0f'%d['vc']['day']),c2('%.0f'%d['vc']['res'])]
table('Limites (amplitude), derivada máxima e decomposição da variância por variável do humor.',
    ['Variável','Amplitude do sinal','Amplitude intra-atleta','Dia de maior taxa','Taxa (pontos/dia)','Var. entre atletas (%)','Var. dia (%)','Var. resíduo (%)'],
    [ldrow(k,l) for k,l in [('Vigor','Vigor'),('Fadiga','Fadiga'),('TMD','PTH'),('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Confusao','Confusão')]],fs=8)

# ===== E. AJUSTE ALOMÉTRICO E SINAL-RUÍDO =====
H('E. Ajuste alométrico e decomposição sinal–ruído')
figure(f'{FG}/x_allo.png','Ajuste alométrico (lei de potência) da trajetória de fadiga em escala linear e log-log.',w=15.0)
def allorow(k,lab):
    v=ALLO[k]; return [lab,c2('%.2f'%v['a']),c2('%.2f'%v['b']),c2('%.2f'%v['r2']),pstr(v['p'])]
table('Ajuste alométrico da trajetória diária (Y = a · dia^b) das variáveis de fadiga.',
    ['Variável','a','b','R²','p'],[allorow('FadFisica','Fadiga física'),allorow('Fadiga','Fadiga (BRUMS)'),allorow('TMD','PTH')],
    note='b < 1 = processo saturante. Escalonamento da aptidão: PV = a·massa^%s; ρ aptidão–fadiga bruta = %s, alométrica = %s.'%(
        c2('%.2f'%ALLO['fitness']['b']),c2('%+.2f'%ALLO['fitness']['rho_raw']),c2('%+.2f'%ALLO['fitness']['rho_allo'])),fs=9)
figure(f'{FG}/x_roc.png','Curvas ROC (Dia 7 vs. Dia 1) com e sem ruído (medidas brutas vs. médias diárias).',w=13.5)
def snrow(k,lab):
    return [lab,c2('%.2f'%DEN['rel'][k]),c2('%.2f'%DEN['etm'][k]),c2('%.2f'%DEN['snr_raw'][k]),c2('%.2f'%ROC[k]['raw'][0]),c2('%.2f'%ROC[k]['filt'][0]),c2('%.1f'%DEN['kmdc'][k]['1']),c2('%.1f'%DEN['kmdc'][k]['2'])]
table('Decomposição sinal–ruído: confiabilidade, erro típico, razão sinal–ruído, discriminação (ROC) e mudança detectável.',
    ['Variável','r (1 medida)','ETM','RSR','AUC c/ ruído','AUC s/ ruído','MDC95 (1)','MDC95 (2)'],
    [snrow('Vigor','Vigor'),snrow('Fadiga','Fadiga (BRUMS)'),snrow('TMD','PTH'),snrow('FadMental','Fadiga mental')],
    note='ETM = erro típico de medida; RSR = razão sinal–ruído; MDC95 (k) = mínima mudança detectável com k coletas.',fs=8.5)

# ===== F. SONO, ESTRESSE E HIIT =====
H('F. Sono, estresse e carga de HIIT')
ep=EX['epw']; ps=EX['pss']; hi=EX['hiit']; ha=EX['hiit_acute']
def cx(c): return '%s (%s)'%(c2('%+.2f'%c['rho']),pstr(c['p']))
def sprow(k,lab): return [lab,cx(ep['corr'][k]),cx(ps['corr'][k])]
table('Correlação de Spearman da sonolência (Epworth) e do estresse percebido (PSS) com o humor (n = %d).'%ep['n'],
    ['Desfecho semanal','Epworth ρ (p)','PSS ρ (p)'],
    [sprow('Vigor','Vigor'),sprow('Fadiga','Fadiga'),sprow('TMD','PTH'),sprow('FadFisica','Fadiga física')],
    note='Epworth: M = %s ± %s; %d de %d atletas com sonolência excessiva (> 10). PSS: M = %s ± %s.'%(
        c2('%.1f'%ep['m']),c2('%.1f'%ep['sd']),ep['hi'],ep['n'],c2('%.1f'%ps['m']),c2('%.1f'%ps['sd'])),fs=9)
figure(f'{FG}/x_sono.png','Sonolência (Epworth) e fadiga: dispersão com reta de regressão e comparação sonolentos vs. não sonolentos.',w=15.0)
def hrow(k,lab):
    v=hi[k]; return [lab,c2('%.2f'%v['hiit']),c2('%.2f'%v['nohiit']),c2('%+.2f'%(v['hiit']-v['nohiit'])),pstr(v['p']),c2('%+.2f'%v['dz'])]
table('Humor e fadiga nos dias com vs. sem HIIT (média por atleta; Wilcoxon pareado, n = %d).'%hi['Vigor']['n'],
    ['Variável','Dias com HIIT','Dias sem HIIT','Δ','p','dz'],
    [hrow('Vigor','Vigor'),hrow('Fadiga','Fadiga'),hrow('TMD','PTH'),hrow('FadFisica','Fadiga física')],
    note='Dias com HIIT = 2, 4 e 7; sem HIIT = 1, 3, 5 e 6. Resposta aguda pré→pós não diferiu entre HIIT e não-HIIT (p > 0,05).',fs=9)
figure(f'{FG}/x_hiit.png','Humor e fadiga nos dias com vs. sem HIIT (* p < 0,05).',w=12.0)
figure(f'{FG}/xb2_intraday.png','Escores pré e pós-treino por dia (vigor, fadiga, PTH) e variação intra-dia das dimensões negativas.')

OUTP='/home/user/mdlucca/Artigos/Material_Suplementar_Perfil_Humor.docx'
doc.save(OUTP); print('SAVED',OUTP,'| Tabelas S',_TN[0],'Figuras S',_FN[0])
