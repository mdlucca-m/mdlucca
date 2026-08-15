# -*- coding: utf-8 -*-
import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
R=json.load(open('brums_desc2.json'))
FG='/home/user/mdlucca/Artigos/figuras'
def c2(s): return str(s).replace('.',',')
doc=Document()
st=doc.styles['Normal']; st.font.name='Times New Roman'; st.font.size=Pt(12)
st.element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')
st.paragraph_format.line_spacing=1.5; st.paragraph_format.space_after=Pt(0)
s=doc.sections[0]; s.top_margin=Cm(3); s.left_margin=Cm(3); s.bottom_margin=Cm(2); s.right_margin=Cm(2)
_TN=[0]; _FN=[0]
def P(t='',just=True,size=12,after=6,bold=False,ind=True):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(size); r.bold=bold
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(after)
    if just: p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if just and ind: p.paragraph_format.first_line_indent=Cm(1.25)
    return p
def H(t,size=12):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(size)
    p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.5
def _bd(c):
    tcPr=c._tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for e_ in ['top','bottom','left','right']:
        el=OxmlElement(f'w:{e_}'); on=e_ in ('top','bottom')
        el.set(qn('w:val'),'single' if on else 'nil'); el.set(qn('w:sz'),'6'); el.set(qn('w:color'),'000000'); b.append(el)
    tcPr.append(b)
def table(cap,header,rows,fonte='Fonte: dados da pesquisa (2026).',fs=9):
    _TN[0]+=1
    p=doc.add_paragraph(); r=p.add_run('Tabela %d – %s'%(_TN[0],cap)); r.font.size=Pt(11)
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
    t=doc.add_table(rows=1,cols=len(header)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,htx in enumerate(header):
        cc=t.rows[0].cells[i]; cc.text=''; rr=cc.paragraphs[0].add_run(htx); rr.bold=True; rr.font.size=Pt(fs); rr.font.name='Times New Roman'
        cc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; _bd(cc)
    for row in rows:
        cs=t.add_row().cells
        for i,val in enumerate(row):
            cs[i].text=''; rr=cs[i].paragraphs[0].add_run(str(val)); rr.font.size=Pt(fs); rr.font.name='Times New Roman'
            cs[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if i>0 else WD_ALIGN_PARAGRAPH.LEFT; _bd(cs[i])
    pf=doc.add_paragraph(); rf=pf.add_run(fonte); rf.font.size=Pt(9); pf.paragraph_format.space_after=Pt(6)
def figure(path,cap,w=15.5):
    _FN[0]+=1
    pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; pp.add_run().add_picture(path,width=Cm(w)); pp.paragraph_format.space_before=Pt(6)
    pc=doc.add_paragraph(); pc.alignment=WD_ALIGN_PARAGRAPH.CENTER; rc=pc.add_run('Figura %d – %s'%(_FN[0],cap)); rc.font.size=Pt(11)
    pf=doc.add_paragraph(); pf.alignment=WD_ALIGN_PARAGRAPH.CENTER; rf=pf.add_run('Fonte: elaboração dos autores (2026).'); rf.font.size=Pt(9); pf.paragraph_format.space_after=Pt(6)

SUBS=[('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Vigor','Vigor'),('Fadiga','Fadiga'),('Confusao','Confusão'),('TMD','PTH/TMD')]
pr=R['prepos']; d17=R['d1d7']; prof=R['profiles']; sm=R['sample']; foc=R['focus']; npk=R['neg_peak']; byd=R['byday']

# ===== TÍTULO =====
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('PERFIL DE HUMOR DE ATLETAS DE HANDEBOL DE ELITE NA ÚLTIMA SEMANA DE PRÉ-TEMPORADA: ANÁLISE DESCRITIVA, COMPARATIVA E DE CORRELAÇÃO DAS VARIÁVEIS DO BRUMS')
r.bold=True; r.font.size=Pt(13); p.paragraph_format.space_after=Pt(10)

H('RESUMO')
P('Objetivo: avaliar o perfil de humor de atletas de handebol de elite na última semana de pré-temporada, descrevendo e '
 'comparando o comportamento das variáveis do BRUMS no grupo e no sujeito. Método: estudo descritivo de medidas repetidas '
 'com %d atletas, %d observações do BRUMS-24 em sete dias, duas coletas/dia (pré e pós-treino). Analisou-se o '
 'comportamento geral, a resposta pré→pós, o padrão diário e intra-dia, a estrutura de correlação entre as subescalas e a '
 'classificação nos seis perfis de humor, com percentuais, intervalos de confiança e tamanhos de efeito. Resultados: a '
 'deterioração concentra-se no eixo energia–fadiga — do pré ao pós o vigor cai %s%% (dz = %s) e a fadiga sobe %s%% '
 '(dz = %s); do Dia 1 ao Dia 7 o vigor recua com efeito grande (dz = %s). As subescalas negativas comportam-se de modo '
 'próprio: a raiva atinge o máximo no dia mais fadigado (D7), a depressão tem pico intermediário e a confusão e a tensão '
 'decrescem. Notavelmente, o acoplamento das negativas com a fadiga é fraco no dia de maior vigor (D1) e forte no dia de '
 'maior fadiga (D7: depressão ρ = %s; raiva ρ = %s). O perfil iceberg cai de %d%% para %d%% e o de fadiga (barbatana de '
 'tubarão) sobe de %d%% para %d%%. Conclusão: o humor migra da prontidão para a fadiga funcional, e o afeto negativo '
 'consolida-se com a fadiga apenas quando a carga se acumula.'%(
   sm['n'],sm['n_obs'],c2('%.0f'%abs(pr['Vigor']['pct'])),c2('%.2f'%pr['Vigor']['dz']),
   c2('%.0f'%pr['Fadiga']['pct']),c2('%.2f'%pr['Fadiga']['dz']),c2('%.2f'%d17['Vigor']['dz']),
   c2('%+.2f'%foc['D7']['Depressao']['rho_fad']),c2('%+.2f'%foc['D7']['Raiva']['rho_fad']),
   int(prof['D1']['Iceberg']),int(prof['D7']['Iceberg']),int(prof['D1']['Barbatana tubarão']),int(prof['D7']['Barbatana tubarão'])),after=8)

# ===== 1 INTRODUÇÃO =====
H('1 INTRODUÇÃO')
P('O monitoramento dos estados de humor consolidou-se como uma das estratégias mais práticas, econômicas e não invasivas '
 'para acompanhar a resposta de atletas à carga de treino e para sinalizar precocemente a fadiga acumulada e o risco de '
 'excesso de treino. Em revisão sistemática, medidas subjetivas de bem-estar rastrearam a carga aguda e crônica com '
 'sensibilidade e consistência superiores às de marcadores objetivos, o que sustenta o uso rotineiro de escalas de humor '
 'no esporte de rendimento (SAW; MAIN; GASTIN, 2016).')
P('A Brunel Mood Scale (BRUMS), derivada do Profile of Mood States e desenvolvida para uso rápido em contextos '
 'esportivos, mede seis dimensões — tensão, depressão, raiva, vigor, fadiga e confusão — e é um dos instrumentos de humor '
 'mais utilizados na ciência do esporte, com propriedades psicométricas sólidas e validações transculturais sucessivas '
 '(TERRY et al., 1999; TERRY; LANE; FOGARTY, 2003; ROHLFS et al., 2008; LEW et al., 2023). Mais do que escores isolados, '
 'o humor pode ser lido como um perfil integrado: Morgan (1985) descreveu o “perfil iceberg” — vigor elevado sobre '
 'dimensões negativas baixas — como marca de prontidão e saúde mental, e trabalhos recentes formalizaram seis perfis '
 'prototípicos (iceberg, Everest invertido, iceberg invertido, submerso, barbatana de tubarão e superfície), hoje '
 'amplamente empregados no rastreamento do estado psicológico e da prontidão de atletas (PARSONS-SMITH; TERRY; MACHIN, '
 '2017; TERRY et al., 2021; TAN et al., 2024). Metanálises confirmam a relação entre a perturbação do humor e o '
 'desempenho esportivo, reforçando o valor aplicado do constructo (LOCHBAUM et al., 2021).')
P('O handebol de quadra é uma modalidade coletiva intermitente de alta intensidade: o jogo alterna, de forma imprevisível, '
 'esforços máximos e explosivos — sprints, saltos, arremessos, mudanças de direção, contatos — com períodos de '
 'recuperação incompleta, exigindo simultaneamente potência anaeróbia, capacidade aeróbia intermitente e tolerância à '
 'fadiga (KARCHER; BUCHHEIT, 2014; WAGNER et al., 2014). As demandas variam com a posição tática, e indícios de fadiga já '
 'foram documentados pela queda do volume de corrida em alta intensidade ao longo da partida (MICHALSIK; MADSEN; AAGAARD, '
 '2013). Esse perfil de esforço torna a carga interna particularmente elevada nos microciclos de acúmulo, com repercussão '
 'direta sobre o estado afetivo dos atletas — em handebolistas de elite, o humor e o estresse associam-se a marcadores de '
 'sobrecarga fisiológica e psicossomática (RATZ-SULYOK et al., 2026).')
P('A resposta afetiva à carga, contudo, não é estática: em fases de redução de carga a fadiga percebida decresce e o '
 'desempenho melhora, ao passo que, em microciclos de acúmulo, espera-se o movimento inverso — vigor em queda e fadiga em '
 'ascensão. A última semana de pré-temporada concentra a carga que antecede a competição e constitui, portanto, uma '
 'janela crítica para observar como cada dimensão do humor se comporta ao longo dos dias e dentro de cada dia de treino, '
 'e como as dimensões se relacionam entre si — em particular, se e quando o afeto negativo (depressão, raiva, confusão) '
 'passa a acompanhar a fadiga. Descrever e comparar esse comportamento, do nível do grupo ao nível do sujeito, fornece à '
 'comissão técnica informação individualizada e diretamente acionável.')

# ===== 2 JUSTIFICATIVA =====
H('2 JUSTIFICATIVA')
P('Descrever e classificar o perfil de humor na semana pré-competitiva, com estatística comparativa robusta (variação '
 'percentual, intervalos de confiança e tamanho de efeito), permite separar o que é sinal de deterioração real do que é '
 'apenas flutuação, distinguir uma fadiga funcional de treino de uma deterioração afetiva mais ampla e identificar quais '
 'dimensões e quais dias concentram a maior sensibilidade. Compreender, ainda, a estrutura de correlação entre as '
 'subescalas — especialmente nos dias de maior vigor e de maior fadiga — esclarece se o afeto negativo é independente da '
 'fadiga ou se com ela se consolida, o que tem implicação direta para a interpretação dos escores no monitoramento diário '
 'e para a individualização da recuperação.')

# ===== 3 OBJETIVO =====
H('3 OBJETIVO')
P('Avaliar o perfil de humor de atletas de handebol de elite na última semana de pré-temporada, descrevendo e comparando '
 'o comportamento das variáveis do BRUMS no grupo (geral, pré→pós, por dia e intra-dia), analisando a correlação entre as '
 'subescalas com ênfase nos dias de maior vigor e de maior fadiga, e classificando os perfis de humor — do nível do grupo '
 'ao nível do sujeito, com estimativas de variação percentual, intervalo de confiança e tamanho de efeito.')

# ===== 4 MÉTODO =====
H('4 MÉTODO')
H('4.1 Delineamento e amostra',12)
P('Estudo descritivo, longitudinal e observacional, de medidas repetidas intraindividuais, em condições ecológicas de '
 'treinamento. Participaram %d atletas de handebol do sexo masculino de nível competitivo (Tabela 1), distribuídos nas '
 'posições da modalidade (%s). O monitoramento ocorreu ao longo de sete dias consecutivos (21–27/04/2024), o microciclo '
 'pré-competitivo que antecede a competição, com duas aplicações do BRUMS por dia de treino — a primeira tomada como pré '
 'e a última como pós —, totalizando %d observações válidas. A Figura 1 detalha o framework das coletas.'%(
   sm['n'],', '.join('%s (%d)'%(k,v) for k,v in sm['pos'].items()),sm['n_obs']))
def srow(lab,k):
    v=sm[k]; return [lab,c2('%.1f'%v['mean']),c2('%.1f'%v['sd']),'%s–%s'%(c2('%.1f'%v['lo']),c2('%.1f'%v['hi'])),'%s–%s'%(c2('%.1f'%v['mn']),c2('%.1f'%v['mx']))]
table('Caracterização sociodemográfica e antropométrica da amostra (n = %d): média, desvio-padrão, IC95%% e amplitude.'%sm['n'],
    ['Variável','Média','DP','IC95%','Mín–Máx'],
    [srow('Idade (anos)','idade'),srow('Estatura (cm)','estatura'),srow('Massa corporal (kg)','massa'),srow('Percentual de gordura','pG'),srow('Experiência na modalidade (anos)','exp')])
figure(f'{FG}/xb2_framework.png','Framework das coletas: da amostra e do microciclo às observações do BRUMS, às subescalas e à classificação de perfis.',w=11.0)
H('4.2 Instrumento',12)
P('O humor foi avaliado pela BRUMS-24, composta por 24 adjetivos respondidos em escala Likert de 0 (“nada”) a 4 '
 '(“extremamente”), organizados em seis subescalas de 0 a 16 pontos (tensão, depressão, raiva, vigor, fadiga e confusão). '
 'A Perturbação Total do Humor (PTH) resume o perfil em um índice único: PTH = tensão + depressão + raiva + fadiga + '
 'confusão − vigor. Os questionários foram autoaplicados por formulário eletrônico com carimbo de data/hora, tomando-se a '
 'primeira resposta do dia como pré e a última como pós.')
H('4.3 Análise estatística',12)
P('Estatística descritiva de cada subescala e da PTH (média, mediana, desvio-padrão, intervalo de confiança de 95%, '
 'coeficiente de variação e amplitude), no geral, por dia e nos momentos pré e pós. As comparações pré→pós e Dia 1→Dia 7 '
 'usaram o teste de Wilcoxon pareado, reportando variação absoluta e percentual, tamanho de efeito pareado (dz de Cohen) '
 'com intervalo de confiança e classificação de magnitude (trivial < 0,2; pequeno < 0,5; médio < 0,8; grande ≥ 0,8). A '
 'estrutura de associação entre as subescalas foi examinada por correlação de Spearman (matriz entre atletas e '
 'correlações diárias), com ênfase nos dias de maior vigor e de maior fadiga. Cada observação foi classificada nos seis '
 'perfis de humor de Terry por proximidade aos protótipos em escores padronizados. As análises foram conduzidas no nível '
 'do grupo e no nível do sujeito.')

# ===== 5 RESULTADOS =====
H('5 RESULTADOS')
P('Os resultados são apresentados primeiro no nível do grupo (5.1 a 5.6) e, em seguida, no nível do sujeito (5.7).',after=6)
H('5.1 Análise descritiva geral do grupo',12)
P('As subescalas negativas situam-se próximas do piso (medianas baixas e coeficientes de variação elevados), enquanto o '
 'vigor e a fadiga concentram a média e a dispersão informativas — é no eixo energia–fadiga que reside a variabilidade '
 'útil do microciclo (Tabela 2). A confusão é a subescala de menor média e a tensão e a raiva, embora baixas, são as '
 'negativas mais expressivas.')
def drow(k,lab):
    v=R['desc'][k]; return [lab,c2('%.2f'%v['mean']),c2('%.1f'%v['md']),c2('%.2f'%v['sd']),
        '%s–%s'%(c2('%.2f'%v['lo']),c2('%.2f'%v['hi'])),c2('%.0f'%v['cv']),'%s–%s'%(c2('%.0f'%v['mn']),c2('%.0f'%v['mx']))]
table('Estatística descritiva geral do grupo por subescala do BRUMS (%d observações).'%sm['n_obs'],
    ['Subescala','Média','Mediana','DP','IC95%','CV (%)','Mín–Máx'],[drow(k,l) for k,l in SUBS])

H('5.2 Comportamento ao longo da semana e do Dia 1 ao Dia 7',12)
P('A Figura 2 mostra, com bandas de confiança de 95%%, as trajetórias das subescalas: no eixo energia–fadiga o vigor '
 'declina e a fadiga e a PTH ascendem de forma sistemática, ao passo que as subescalas negativas de valência não fadiga '
 'oscilam próximas do piso, sem tendência monotônica clara. Na comparação Dia 1 → Dia 7 (Tabela 3), o vigor recua com '
 'efeito grande (Δ = %s; %s%%; dz = %s; p = %s) e a fadiga sobe com efeito médio (Δ = %s; %s%%; dz = %s; p = %s); a tensão '
 'e a confusão, ao contrário, decrescem significativamente — uma acomodação do estado inicial de ativação e apreensão.'%(
   c2('%+.2f'%d17['Vigor']['delta']),c2('%.0f'%d17['Vigor']['pct']),c2('%.2f'%d17['Vigor']['dz']),c2('%.3f'%d17['Vigor']['p']),
   c2('%+.2f'%d17['Fadiga']['delta']),c2('%.0f'%d17['Fadiga']['pct']),c2('%.2f'%d17['Fadiga']['dz']),c2('%.3f'%d17['Fadiga']['p'])))
figure(f'{FG}/xb2_traj.png','Trajetória das subescalas do BRUMS ao longo da semana (médias diárias; áreas sombreadas = IC95%).')
def d17row(k,lab):
    v=d17[k]; return [lab,c2('%.2f'%v['d1']),c2('%.2f'%v['d7']),c2('%+.2f'%v['delta']),c2('%+.0f'%v['pct'])+'%',c2('%+.2f'%v['dz']),c2('%.3f'%v['p']),v['mag']]
table('Comparação Dia 1 → Dia 7 por subescala: médias, variação absoluta e percentual, tamanho de efeito e magnitude.',
    ['Subescala','Dia 1','Dia 7','Δ','Δ%','dz','p','Magnitude'],[d17row(k,l) for k,l in SUBS],fs=8.5)

H('5.3 Resposta aguda pré→pós ao longo da semana',12)
P('Agregando toda a semana, a resposta aguda ao treino (pré → pós) é significativa e de magnitude média no eixo '
 'energia–fadiga (Tabela 4): o vigor cai %s%% (dz = %s; p = %s), a fadiga sobe %s%% (dz = %s; p = %s) e a PTH aumenta %s%% '
 '(dz = %s). As demais subescalas apresentam resposta trivial a pequena, sem significância — a deterioração aguda do '
 'treino concentra-se na energia e na fadiga e poupa o afeto negativo geral.'%(
   c2('%.0f'%abs(pr['Vigor']['pct'])),c2('%.2f'%pr['Vigor']['dz']),c2('%.3f'%pr['Vigor']['p']),
   c2('%.0f'%pr['Fadiga']['pct']),c2('%.2f'%pr['Fadiga']['dz']),c2('%.3f'%pr['Fadiga']['p']),
   c2('%.0f'%pr['TMD']['pct']),c2('%.2f'%pr['TMD']['dz'])))
def prow(k,lab):
    v=pr[k]; return [lab,c2('%.2f'%v['pre']),c2('%.2f'%v['pos']),c2('%+.2f'%v['delta']),c2('%+.0f'%v['pct'])+'%',
        '%s [%s, %s]'%(c2('%+.2f'%v['dz']),c2('%.2f'%v['lo']),c2('%.2f'%v['hi'])),c2('%.3f'%v['p']),v['mag']]
table('Resposta pré → pós agregada na semana: médias, variação percentual, dz (IC95%), p e magnitude.',
    ['Subescala','Pré','Pós','Δ','Δ%','dz [IC95%]','p','Magnitude'],[prow(k,l) for k,l in SUBS],fs=8.5)

H('5.4 Comportamento da depressão, da raiva e da confusão',12)
P('Embora operem próximas do piso, as subescalas negativas apresentam dinâmicas próprias e informativas (Figura 3; Tabela '
 '5). A raiva é a mais reativa e atinge o máximo no dia mais fadigado (pico no Dia %d = %s), acompanhando o acúmulo de '
 'carga. A depressão tem pico intermediário (Dia %d = %s), sugerindo um vale afetivo no meio da semana. A confusão e a '
 'tensão, ao contrário, são mais altas no início (pico no Dia %d), quando a apreensão pré-microciclo é maior, e decrescem '
 'em seguida — a confusão atinge o mínimo (%s) e permanece a subescala de menor expressão. Todas as negativas apresentam '
 'um vale comum por volta do Dia 5, dia de menor carga do microciclo.'%(
   npk['Raiva']['peak_day'],c2('%.2f'%npk['Raiva']['peak_val']),npk['Depressao']['peak_day'],c2('%.2f'%npk['Depressao']['peak_val']),
   npk['Confusao']['peak_day'],c2('%.2f'%npk['Confusao']['min_val'])))
figure(f'{FG}/xb2_negatives.png','Comportamento diário da depressão, da raiva e da confusão (médias com IC95%); faixas sombreadas destacam o dia de maior vigor (D1) e o de maior fadiga (D7).')
def negrow(k,lab):
    row=[lab]+[c2('%.2f'%byd[k][str(d)]['all']) for d in range(1,8)]
    row.append('D%d'%npk[k]['peak_day']); return row
table('Médias diárias das subescalas negativas e dia de pico.',
    ['Subescala']+['D%d'%d for d in range(1,8)]+['Pico'],
    [negrow('Tensao','Tensão'),negrow('Depressao','Depressão'),negrow('Raiva','Raiva'),negrow('Confusao','Confusão')],fs=8.5)

H('5.5 Correlação entre as variáveis do humor e ênfase nos dias de maior vigor e de maior fadiga',12)
CT=R['corr_trait']; LB=R['corr_labels']
def gg(a,b): return c2('%+.2f'%CT[LB.index(a)][LB.index(b)])
P('A matriz de correlação entre atletas (Figura 4) revela um agrupamento coeso de afeto negativo: a depressão '
 'correlaciona-se com a raiva (ρ = %s), a confusão (ρ = %s) e a tensão (ρ = %s), e, de forma relevante, com a fadiga '
 '(ρ = %s); a raiva também acompanha a fadiga (ρ = %s). O eixo energia–fadiga aparece como dimensão invertida '
 '(vigor × fadiga ρ = %s), e o vigor é praticamente independente das negativas — coerente com a leitura de que a '
 'deterioração de treino é, sobretudo, um fenômeno de energia e fadiga.'%(
   gg('Depressão','Raiva'),gg('Depressão','Confusão'),gg('Depressão','Tensão'),gg('Depressão','Fadiga'),gg('Raiva','Fadiga'),gg('Vigor','Fadiga')))
figure(f'{FG}/xb2_heatmap.png','Matriz de correlação de Spearman entre as subescalas do BRUMS (nível entre atletas).',w=11.0)
P('O achado mais expressivo surge ao contrastar os dias extremos (Figura 5; Tabela 6). No dia de maior vigor (Dia 1, '
 'grupo descansado), o afeto negativo é praticamente independente da fadiga: depressão ρ = %s, raiva ρ = %s e confusão '
 'ρ = %s com a fadiga. No dia de maior fadiga (Dia 7), esse acoplamento torna-se forte: a depressão passa a ρ = %s, a '
 'raiva a ρ = %s e a confusão a ρ = %s com a fadiga, e ambas se associam negativamente ao vigor. Ou seja, quando o grupo '
 'está fresco, quem está mais irritado ou abatido não é necessariamente o mais cansado; quando a carga se acumula, as '
 'dimensões negativas consolidam-se com a fadiga — o perfil de humor “fecha” sob exaustão. O padrão cresce de modo '
 'progressivo ao longo dos dias, com um recuo transitório no Dia 5 (menor carga).'%(
   c2('%+.2f'%foc['D1']['Depressao']['rho_fad']),c2('%+.2f'%foc['D1']['Raiva']['rho_fad']),c2('%+.2f'%foc['D1']['Confusao']['rho_fad']),
   c2('%+.2f'%foc['D7']['Depressao']['rho_fad']),c2('%+.2f'%foc['D7']['Raiva']['rho_fad']),c2('%+.2f'%foc['D7']['Confusao']['rho_fad'])))
figure(f'{FG}/xb2_corrfocus.png','Acoplamento das subescalas negativas com a fadiga: (A) no dia de maior vigor (D1) vs. no dia de maior fadiga (D7); (B) evolução dia a dia.')
def frow(neg,lab):
    return [lab,c2('%+.2f'%foc['D1'][neg]['rho_fad']),c2('%+.2f'%foc['D7'][neg]['rho_fad']),
            c2('%+.2f'%foc['D1'][neg]['rho_vig']),c2('%+.2f'%foc['D7'][neg]['rho_vig'])]
table('Correlação (ρ de Spearman) das subescalas negativas com a fadiga e o vigor no dia de maior vigor (D1) e no de maior fadiga (D7).',
    ['Subescala','ρ×Fadiga (D1)','ρ×Fadiga (D7)','ρ×Vigor (D1)','ρ×Vigor (D7)'],
    [frow('Depressao','Depressão'),frow('Raiva','Raiva'),frow('Confusao','Confusão'),frow('Tensao','Tensão')],fs=9)

H('5.6 Comportamento intra-dia (pré → pós por dia) e perfis de humor',12)
P('O padrão intra-dia repete-se com regularidade: em todos os dias de treino o vigor cai e a fadiga e a PTH sobem do pré '
 'para o pós, sinal de um choque agudo consistente a cada sessão, enquanto as subescalas negativas variam pouco dentro do '
 'dia (Figura 6). Quanto à classificação, na linha de base (Dia 1) predomina o perfil iceberg (%d%%), marca de prontidão; '
 'no último dia (Dia 7) o perfil de fadiga (“barbatana de tubarão”) torna-se o mais frequente (%d%%) e o iceberg cai para '
 '%d%% (Figura 7; Tabela 7). Essa migração — do iceberg para a assinatura de fadiga funcional, sem instalação relevante '
 'de perfis de risco (submerso, iceberg invertido) — é a leitura central da semana.'%(
   int(prof['D1']['Iceberg']),int(prof['D7']['Barbatana tubarão']),int(prof['D7']['Iceberg'])))
figure(f'{FG}/xb2_intraday.png','Escores pré e pós-treino por dia (vigor, fadiga, PTH) e variação intra-dia das subescalas negativas.')
figure(f'{FG}/xb2_profiles.png','Distribuição (%) dos seis perfis de humor por recorte: linha de base (D1), semana, último dia (D7), pré e pós.')
PROFR=[('Iceberg','Iceberg'),('Everest invertido','Everest invertido'),('Iceberg invertido','Iceberg invertido'),('Submerso','Submerso'),('Barbatana tubarão','Barbatana de tubarão'),('Superfície','Superfície')]
recs=[('D1','D1'),('Semana','overall'),('D7','D7'),('Pré','pre'),('Pós','pos')]
table('Distribuição (%) dos perfis de humor de Terry por recorte.',
    ['Perfil']+[l for l,_ in recs],
    [[lab]+[c2('%.0f'%prof[key][p]) for _,key in recs] for p,lab in PROFR],fs=9)

H('5.7 Análise no nível do sujeito',12)
sub=R['subject']
P('No nível individual, a resposta é heterogênea, mas majoritariamente de deterioração no eixo energia–fadiga: do Dia 1 '
 'ao Dia 7, %d%% dos atletas reduzem o vigor e %d%% aumentam a fadiga. As médias semanais individuais variam amplamente '
 '(vigor de %s a %s; fadiga de %s a %s), o que reforça a recomendação de interpretar o humor referenciado à linha de base '
 'de cada atleta — e não apenas pela média do grupo. A conjunção dos achados de grupo e de sujeito indica que, embora o '
 'movimento médio seja claro (vigor ↓, fadiga ↑, perfil iceberg → fadiga), a magnitude e o momento da deterioração são '
 'próprios de cada atleta, o que justifica o monitoramento individualizado por tendência.'%(
   int(sub['Vigor']['pct_worse']),int(sub['Fadiga']['pct_worse']),
   c2('%.1f'%sub['Vigor']['wk_min']),c2('%.1f'%sub['Vigor']['wk_max']),
   c2('%.1f'%sub['Fadiga']['wk_min']),c2('%.1f'%sub['Fadiga']['wk_max'])))

# ===== 6 SÍNTESE =====
H('6 SÍNTESE DOS PRINCIPAIS ACHADOS')
for b in [
 'A deterioração do humor concentra-se no eixo energia–fadiga (vigor ↓, fadiga ↑, PTH ↑); as subescalas negativas de valência não fadiga permanecem próximas do piso.',
 'Resposta aguda pré→pós significativa e de magnitude média: vigor −%s%% (dz = %s), fadiga +%s%% (dz = %s), PTH +%s%% (dz = %s).'%(c2('%.0f'%abs(pr['Vigor']['pct'])),c2('%.2f'%pr['Vigor']['dz']),c2('%.0f'%pr['Fadiga']['pct']),c2('%.2f'%pr['Fadiga']['dz']),c2('%.0f'%pr['TMD']['pct']),c2('%.2f'%pr['TMD']['dz'])),
 'Do Dia 1 ao Dia 7, o vigor recua com efeito grande (dz = %s) e a fadiga sobe com efeito médio (dz = %s); tensão e confusão se acomodam.'%(c2('%.2f'%d17['Vigor']['dz']),c2('%.2f'%d17['Fadiga']['dz'])),
 'As negativas têm dinâmica própria: a raiva pica no dia mais fadigado (D7), a depressão no meio da semana e a confusão/tensão no início.',
 'O afeto negativo consolida-se com a fadiga apenas sob carga: depressão×fadiga passa de ρ = %s (D1) para ρ = %s (D7); raiva×fadiga de %s para %s.'%(c2('%+.2f'%foc['D1']['Depressao']['rho_fad']),c2('%+.2f'%foc['D7']['Depressao']['rho_fad']),c2('%+.2f'%foc['D1']['Raiva']['rho_fad']),c2('%+.2f'%foc['D7']['Raiva']['rho_fad'])),
 'O perfil migra do iceberg (%d%% no D1) para a fadiga funcional (barbatana de tubarão, %d%% no D7), sem instalação de perfis de risco.'%(int(prof['D1']['Iceberg']),int(prof['D7']['Barbatana tubarão'])),
 'No nível do sujeito, a maioria deteriora de forma coerente com o grupo, com ampla variabilidade individual — monitorar por tendência e por linha de base individual.']:
    p=doc.add_paragraph(); r=p.add_run('•  '+b); r.font.size=Pt(12); p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(4); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

H('REFERÊNCIAS')
refs=[
 'BRANDT, R.; BEVILACQUA, G. G.; ANDRADE, A. Perceived sleep quality, mood states, and their relationship with performance among Brazilian elite athletes during a competitive period. Journal of Strength and Conditioning Research, v. 31, n. 4, p. 1033–1039, 2017.',
 'KARCHER, C.; BUCHHEIT, M. On-court demands of elite handball, with special reference to playing positions. Sports Medicine, v. 44, n. 6, p. 797–814, 2014. DOI: 10.1007/s40279-014-0164-z.',
 'LEW, P. C. F. et al. Cross-cultural validation of the Malaysian Mood Scale and tests of between-group mood differences. International Journal of Environmental Research and Public Health, v. 20, n. 4, 3348, 2023. DOI: 10.3390/ijerph20043348.',
 'LOCHBAUM, M. et al. The Profile of Mood States and athletic performance: a meta-analysis of published studies. European Journal of Investigation in Health, Psychology and Education, v. 11, n. 1, p. 50–70, 2021. DOI: 10.3390/ejihpe11010005.',
 'MICHALSIK, L. B.; MADSEN, K.; AAGAARD, P. Match performance and physiological capacity of female elite team handball players. International Journal of Sports Medicine, v. 35, n. 7, p. 595–607, 2013. DOI: 10.1055/s-0033-1358713.',
 'MORGAN, W. P. Selected psychological factors limiting performance: a mental health model. In: CLARKE, D. H.; ECKERT, H. M. (Ed.). Limits of human performance. Champaign: Human Kinetics, 1985. p. 70–80.',
 'PARSONS-SMITH, R. L.; TERRY, P. C.; MACHIN, M. A. Identification and description of novel mood profile clusters. Frontiers in Psychology, v. 8, 1958, 2017. DOI: 10.3389/fpsyg.2017.01958.',
 'RATZ-SULYOK, F. Z. et al. Associations between endocrine status and stress, mood and psychosomatic status in elite handball players. Sports, v. 14, n. 7, 289, 2026. DOI: 10.3390/sports14070289.',
 'ROHLFS, I. C. P. M. et al. A Escala de Humor de Brunel (Brums): instrumento para detecção precoce da síndrome do excesso de treinamento. Revista Brasileira de Medicina do Esporte, v. 14, n. 3, p. 176–181, 2008.',
 'SAW, A. E.; MAIN, L. C.; GASTIN, P. B. Monitoring the athlete training response: subjective self-reported measures trump commonly used objective measures: a systematic review. British Journal of Sports Medicine, v. 50, n. 5, p. 281–291, 2016. DOI: 10.1136/bjsports-2015-094758.',
 'TAN, C. et al. The structural validity and latent profile characteristics of the Abbreviated Profile of Mood States among Chinese athletes. BMC Psychiatry, v. 24, n. 1, 636, 2024. DOI: 10.1186/s12888-024-06092-5.',
 'TERRY, P. C.; LANE, A. M.; LANE, H. J.; KEOHANE, L. Development and validation of a mood measure for adolescents. Journal of Sports Sciences, v. 17, n. 11, p. 861–872, 1999. DOI: 10.1080/026404199365425.',
 'TERRY, P. C.; LANE, A. M.; FOGARTY, G. J. Construct validity of the Profile of Mood States — Adolescents for use with adults. Psychology of Sport and Exercise, v. 4, n. 2, p. 125–139, 2003. DOI: 10.1016/S1469-0292(02)00035-8.',
 'TERRY, P. C. et al. Mood profiling for sustainable mental health among athletes. Sustainability, v. 13, n. 11, 6116, 2021. DOI: 10.3390/su13116116.',
 'WAGNER, H. et al. Individual and team performance in team-handball: a review. Journal of Sports Science and Medicine, v. 13, n. 4, p. 808–816, 2014.']
for rf in refs:
    p=doc.add_paragraph(); r=p.add_run(rf); r.font.size=Pt(11); p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(6); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

OUTP='/home/user/mdlucca/Artigos/Perfil_Humor_BRUMS_Descritivo.docx'
doc.save(OUTP); print('SAVED',OUTP,'| Tabelas',_TN[0],'Figuras',_FN[0])
