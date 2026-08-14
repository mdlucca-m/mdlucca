import warnings; warnings.filterwarnings('ignore')
import numpy as np, pandas as pd, json, os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
SC='/tmp/claude-0/-home-user-mdlucca/e1dba24c-b1d7-5908-9106-f2f4aaf3f56a/scratchpad'
FIG='/home/user/mdlucca/Artigos/figuras'
R=json.load(open(f'{SC}/amp_docx.json'))
PP=json.load(open(f'{SC}/prepost.json'))
grp,btw,per=R['grp'],R['btw'],R['per']
VARS=[('Vigor','Vigor (0–20)'),('Fadiga','Fadiga BRUMS (0–20)'),('TMD','PTH/TMD'),('FadMental','Fadiga mental (0–10)')]
pcsv={v:pd.read_csv(f'{SC}/peratlheta_{v}.csv') for v,_ in VARS}

doc=Document()
st=doc.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(10.5)
def H(t,l=1): doc.add_heading(t,level=l)
def P(t,it=False,b=False):
    p=doc.add_paragraph(); r=p.add_run(t); r.italic=it; r.bold=b; return p
def fig(name,cap,w=6.4):
    path=f'{FIG}/{name}.png'
    if os.path.exists(path):
        doc.add_picture(path,width=Inches(w)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
        c=doc.add_paragraph(); r=c.add_run(cap); r.italic=True; r.font.size=Pt(8.5); c.alignment=WD_ALIGN_PARAGRAPH.CENTER
def table(headers,rows,cap=None):
    if cap: c=doc.add_paragraph(); r=c.add_run(cap); r.bold=True; r.font.size=Pt(9.5)
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Light Grid Accent 1'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        cell=t.rows[0].cells[i]; cell.paragraphs[0].add_run(str(h)).bold=True; cell.paragraphs[0].runs[0].font.size=Pt(8.5)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=str(v)
            if cells[i].paragraphs[0].runs: cells[i].paragraphs[0].runs[0].font.size=Pt(8.5)
    doc.add_paragraph()

# ===== CAPA =====
doc.add_heading('Amplitude e Separação Ruído × Sinal',0)
P('Vigor × Fadiga (BRUMS) × PTH/TMD × Fadiga mental ao longo do microciclo de choque — 21 a 28 de abril de 2024',it=True)
P('Análise por grupo e por atleta · dentro da semana e entre dias · 27 atletas · 456 observações · reanálise independente em Python (numpy, scipy, statsmodels). Atletas anonimizados (A01–A27).')
doc.add_paragraph()

# ===== 1. Objetivo e método =====
H('1. Objetivo e método')
P('Este relatório responde a duas perguntas sobre quatro variáveis do eixo energia–fadiga — Vigor, Fadiga do BRUMS, Perturbação Total do Humor (PTH/TMD) e Fadiga mental: (1) qual a AMPLITUDE (range) de cada uma ao longo da semana e entre dias, no grupo e no atleta; e (2) quanto dessa oscilação é SINAL (o efeito real do microciclo) e quanto é RUÍDO (erro de medida + flutuação aleatória).')
P('Cada observação é decomposta como sinal + traço + ruído por um modelo de dois fatores (y ~ atleta + dia; soma de quadrados tipo I), cujas três parcelas somam 100% da variância:')
table(['Parcela','O que representa','Interpretação'],[
 ['Sinal do microciclo (dia)','trajetória sistemática imposta pelos 7 dias','o fenômeno de interesse (a "análise")'],
 ['Traço (entre atletas)','diferenças individuais estáveis','quem o atleta é (estrutura, não evento)'],
 ['Ruído (resíduo)','erro de medida + flutuação aleatória intra-atleta','o que NÃO se deve interpretar'],
])
P('Do ruído derivam os limiares de decisão: ETM (erro típico de medida = DP do resíduo); MDC95 = 1,96·√2·ETM (mudança mínima detectável para 1 atleta/1 coleta); MDC95 do grupo = MDC95/√n (o ruído da média encolhe por √n); e SWC = 0,2·DP-entre-atletas (mudança mínima relevante). Regra: uma variação só é "análise" (sinal) se supera o MDC95; abaixo disso é ruído.')

# ===== 2. Grupo ao longo da semana =====
H('2. Grupo — ao longo da semana')
rows=[]
for v,lab in VARS:
    g=grp[v]
    rows.append([lab,f"{g['d1']:.2f}",f"{g['d7']:.2f}",f"{g['drift']:+.2f}",f"{g['sig']:.2f}",f"{g['etm']:.2f}",f"{g['snr']:.2f}"])
table(['Variável','Dia 1','Dia 7','Deriva D7−D1','Ampl. do sinal','ETM (ruído)','SNR (sinal/ruído)'],rows,
      'Tabela 1 — Trajetória do grupo e amplitude do sinal semanal')
P('A Fadiga (BRUMS) e o PTH/TMD sobem e o Vigor cai ao longo da semana; a Fadiga mental quase não se move (deriva 0,3; amplitude do sinal 0,92). O PTH/TMD tem a maior amplitude de sinal em unidades brutas (6,2), mas também o maior ETM (6,0) — por isso seu SNR é o menor entre as três variáveis responsivas (1,04). O Vigor e a Fadiga BRUMS têm SNR ≈ 1,4–1,5. A Fadiga mental tem SNR < 1: sua oscilação semanal é menor que o próprio ruído — é a única das quatro sem sinal semanal apreciável.')
fig('adx1_grupo_semana','Fig 1 — Trajetória do grupo (média ± DP; LOWESS pontilhado). Linhas vermelhas marcam os dias de HIIT (2, 4, 7). Vigor cai, Fadiga/TMD sobem com precipitação no D7; Fadiga mental permanece plana.')

# variance partition
rows=[]
for v,lab in VARS:
    g=grp[v]; rows.append([lab,f"{g['pday']:.1f}%",f"{g['ptrait']:.1f}%",f"{g['pnoise']:.1f}%"])
table(['Variável','% Sinal (dia)','% Traço','% Ruído'],rows,'Tabela 2 — Decomposição da variância: sinal do microciclo vs traço vs ruído')
P('Em todas as quatro variáveis o sinal do microciclo é a MENOR fatia da variância (1–5%): o humor é dominantemente traço (54–73%) mais ruído (26–41%). A Fadiga BRUMS tem a maior fatia de sinal (4,2%) e a Fadiga mental a menor (1,3%), sendo esta quase toda traço (73%) — coerente com um construto estável, não reativo ao microciclo.')

# ===== 3. Entre dias =====
H('3. Grupo — entre dias')
P('A variação entre dias tem duas facetas: a dispersão ENTRE atletas em cada dia (o quanto discordam entre si) e a mudança dia-a-dia da média do grupo (o passo da trajetória).')
# dispersion table
rows=[]
for v,lab in VARS:
    sdb=btw[v]['sd_between']
    rows.append([lab]+[f"{sdb.get(str(d),sdb.get(d,'—')):.2f}" if str(d) in sdb or d in sdb else '—' for d in range(1,8)]+[f"D{btw[v]['day_maxsd']}"])
table(['Variável','D1','D2','D3','D4','D5','D6','D7','Máx. disp.'],rows,'Tabela 3 — Dispersão ENTRE atletas por dia (DP)')
# day-to-day change
rows=[]
for v,lab in VARS:
    dd=btw[v]['dd']
    rows.append([lab]+[f"{dd.get(str(d),dd.get(d,'—')):.2f}" if (str(d) in dd or d in dd) and (dd.get(str(d),dd.get(d)) is not None) else '—' for d in range(2,8)]+[f"D{btw[v]['day_maxdd']}"])
table(['Variável','D1→2','D2→3','D3→4','D4→5','D5→6','D6→7','Maior passo'],rows,'Tabela 4 — |Mudança dia-a-dia| da média do grupo')
P('O maior passo da trajetória concentra-se no fim da semana (D6→D7) para a Fadiga e o PTH/TMD — a "precipitação" do Dia 7 já descrita na análise polinomial. A dispersão entre atletas mantém-se alta em todos os dias (heterogeneidade intrínseca, coerente com o ICC de traço 0,3–0,7), sem um único dia dominante.')
fig('adx2_dispersao_dia','Fig 2 — Dispersão entre atletas por dia (DP). As quatro variáveis mantêm heterogeneidade elevada ao longo de toda a semana — a diferença entre atletas não é um artefato de um dia isolado.')

# ===== 4. Por atleta =====
H('4. Por atleta — amplitude e sinal-ruído individual')
P('No nível individual, a amplitude é o range de cada atleta na semana (max−min). Ela é sempre MAIOR que a amplitude do sinal do grupo, porque some sinal + ruído do próprio atleta:')
rows=[]
for v,lab in VARS:
    p=per[v]
    rows.append([lab,f"{p['amp_mean']:.2f} ± {p['amp_sd']:.2f}",f"{p['amp_min']:.0f}–{p['amp_max']:.0f}",f"{grp[v]['sig']:.2f}",
                 f"{p['psig_mean']:.2f}",f"{p['pnoise_mean']:.2f}",f"{p['pct_snr_gt1']:.0f}%"])
table(['Variável','Ampl. atleta (média±DP)','Faixa','Ampl. sinal grupo','Sinal pessoal (entre dias)','Ruído pessoal','% atletas SNR>1'],rows,
      'Tabela 5 — Amplitude por atleta vs amplitude do sinal do grupo')
P('O atleta oscila 2–3× mais que a trajetória média do grupo (ex.: PTH/TMD amplitude individual 20,1 ± 9,4 vs sinal do grupo 6,2; Vigor 7,1 vs 2,8). Mas essa oscilação individual não é toda ruído: comparando o sinal pessoal de cada atleta (variação entre suas médias diárias) com o ruído pessoal (resíduo em torno da própria média diária), 89–93% dos atletas têm SNR > 1 — ou seja, a maioria tem um movimento semanal discernível acima do próprio ruído nas quatro variáveis. A leitura individual É possível, desde que baseada em médias diárias (que já reduzem o ruído), não em coletas isoladas.')
fig('adx3_heatmap_atleta','Fig 3 — Amplitude semanal por atleta (z por variável; azul = oscila pouco, vermelho = oscila muito). Atletas ordenados do menor ao maior oscilador global. Alguns atletas são consistentemente estáveis; outros amplificam em todas as variáveis.')
fig('adx4_amp_atleta_box','Fig 4 — Amplitude por atleta (caixa + pontos) vs amplitude do sinal do grupo (losango branco). A caixa inteira fica acima do losango: qualquer atleta oscila mais que o microciclo médio.')
fig('adx5_atletas_exemplo','Fig 5 — Trajetórias individuais (z) dos 3 atletas mais estáveis (topo) e 3 mais oscilantes (base). Ilustra a heterogeneidade: para uns a semana é quase plana; para outros, todas as curvas se movem.')

# ===== 5. Pré e pós por dia =====
H('5. Pré → mid → pós por dia — cada variável')
P('Resposta aguda dentro de cada dia, agora com a fase intermediária (mid, medida intra-sessão entre pré e pós): média ± DP de pré, mid e pós, variação Δ (pós−pré), tamanho de efeito pareado dz, Wilcoxon pareado pré×pós e Friedman pré/mid/pós (quando há mid). Dias de HIIT: 2, 4 e 7. O Dia 1 (baseline) não tem coleta mid. * indica p<0,05.')
for v,lab in VARS:
    rows=[]
    for d in range(1,8):
        r=PP[v][str(d)]
        pstr='—' if r['p'] is None else (f"{r['p']:.3f}"+(' *' if r['p']<0.05 else ''))
        fstr='—' if r['pfried'] is None else (f"{r['pfried']:.3f}"+(' *' if r['pfried']<0.05 else ''))
        dzs='—' if r['dz'] is None else f"{r['dz']:+.2f}"
        mids='—' if r['mid_m'] is None else (f"{r['mid_m']:.2f} ± {r['mid_sd']:.2f}" if r['mid_sd'] is not None else f"{r['mid_m']:.2f}")
        rows.append([f"D{d}"+(" (HIIT)" if r['hiit'] else ""),
                     f"{r['pre_m']:.2f} ± {r['pre_sd']:.2f}",mids,f"{r['pos_m']:.2f} ± {r['pos_sd']:.2f}",
                     f"{r['delta']:+.2f}",dzs,pstr,fstr])
    table(['Dia','Pré (m ± DP)','Mid (m ± DP)','Pós (m ± DP)','Δ (pós−pré)','dz','p pré×pós','p Friedman'],rows,
          f'Tabela {6+VARS.index((v,lab))} — {lab}: pré → mid → pós por dia')
P('A fase mid (intra-sessão) mostra que a resposta não é sempre monotônica pré→pós. Em várias variáveis o mid já é mais alto que o pré (a fadiga sobe durante a sessão) e por vezes supera o pós — sinal de um pico intra-sessão com recuperação parcial ao final (p.ex. Fadiga BRUMS no D7: pré 6,8 → mid 8,6 → pós 7,8; TMD no D4: pré 3,1 → mid 7,1 → pós 7,2). O padrão geral persiste: a maior resposta aguda ocorre no Dia 1 (choque inicial) e reaparece no fim da semana (D6–D7); o Vigor cai e Fadiga/TMD/Fadiga mental sobem. Nem todo dia atinge significância — o sinal robusto está no acúmulo semanal (§2), não na sessão isolada (AUC ≈ 0,5 para a sessão vs 0,86 para o acúmulo).')
fig('adx6_prepos_dia','Fig 6 — Pré (azul) → mid (cinza, intra-sessão) → pós por dia, para cada variável. * marca os dias com Wilcoxon pareado pré×pós p<0,05. O mid revela o curso intra-sessão: a fadiga frequentemente pica no meio e recua um pouco ao final.')

# ===== 6. Veredito dois níveis =====
H('6. Veredito em dois níveis — grupo vs indivíduo')
rows=[]
for v,lab in VARS:
    g=grp[v]; di='SIM' if g['detect_ind'] else 'NÃO'; dg='SIM' if g['detect_grp'] else 'NÃO'
    rows.append([lab,f"{g['sig']:.2f}",f"{g['mdc']:.2f}",di,f"{g['mdc_g']:.2f}",dg])
table(['Variável','Ampl. do sinal','MDC95 individual','> MDC ind.?','MDC95 grupo (÷√n)','> MDC grupo?'],rows,
      'Tabela 10 — A oscilação semanal supera o ruído?')
P('No grupo, a oscilação semanal das três variáveis responsivas (Vigor, Fadiga, TMD) supera o MDC95 do grupo por 4–8×: o efeito do microciclo é sinal real e robusto na média. No indivíduo, a mesma oscilação fica abaixo do MDC95 individual de uma única coleta — não é contradição, é o ganho √n da média de 27 atletas. Consequência prática: para "ler" o microciclo no grupo, uma coleta/dia basta; para decidir sobre UM atleta com confiança de MDC individual, use médias de ≥3 coletas (Φ ≥ 0,80). A Fadiga mental não atinge o limiar nem no grupo em unidades de sinal apreciável (SNR < 1): não deve ser monitorada como marcador de carga aguda.')

# ===== 7. Removendo o ruído =====
DN=json.load(open(f'{SC}/denoise.json'))
keys=['Vigor','Fadiga','TMD','FadMental']; labn={'Vigor':'Vigor','Fadiga':'Fadiga BRUMS','TMD':'PTH/TMD','FadMental':'Fadiga mental'}
H('7. E se removermos o ruído? Como ficam as análises')
P('"Remover o ruído" significa reter apenas a parte sistemática (sinal + traço) e descartar o resíduo. Isso pode ser feito de três formas equivalentes: (a) agregar (a média diária do grupo ou de ≥k coletas por atleta reduz o ruído por √n); (b) suavizar a trajetória (LOWESS); (c) corrigir estatísticas pela confiabilidade (desatenuação). O efeito é sempre o mesmo: as relações reais aparecem mais fortes e nítidas — mas o ruído removido NÃO se transforma em sinal onde não havia.')

H('7.1. As associações ficam muito mais fortes',2)
Robs=pd.DataFrame(DN['Robs']).loc[keys,keys]; Rden=pd.DataFrame(DN['Rden']).loc[keys,keys]
rows=[]
for x,y in [('Vigor','Fadiga'),('Fadiga','TMD'),('Vigor','TMD'),('FadMental','Fadiga')]:
    rows.append([f"{labn[x]} × {labn[y]}",f"{Robs.loc[x,y]:+.2f}",f"{Rden.loc[x,y]:+.2f}"])
table(['Par de variáveis','r observado (com ruído)','r denoised (nível-grupo)'],rows,
      'Tabela 11 — Correlações antes e depois de remover o ruído')
P('Removido o ruído, o eixo energia–fadiga aproxima-se de uma dimensão bipolar quase perfeita: Vigor × Fadiga passa de −0,44 (bruto) para −0,93 (denoised); Fadiga × TMD de +0,79 para +0,88. As correlações brutas eram atenuadas pelo ruído de medida — não porque as relações fossem fracas, mas porque cada medida individual é imprecisa. (A correção por desatenuação leva Fadiga × TMD acima de 1,0 — um caso Heywood esperado, pois o TMD contém a própria subescala Fadiga; por isso a estimativa limpa preferida é a de nível-grupo, sempre ≤ 1.)')
fig('den1_corr','Fig 7 — Matriz de correlações do eixo energia–fadiga: observada (com ruído) vs denoised (nível-grupo). Sem o ruído, Vigor e Fadiga tornam-se quase espelhos (r = −0,93).')

H('7.2. Os tamanhos de efeito crescem',2)
dzo={'Vigor':-0.95,'Fadiga':0.72,'TMD':0.41,'FadMental':0.33}; dzc={'Vigor':-1.25,'Fadiga':0.91,'TMD':0.51,'FadMental':0.38}
rows=[[labn[k],f"{dzo[k]:+.2f}",f"{dzc[k]:+.2f}",f"{DN['rel'][k]:.2f}"] for k in keys]
table(['Variável','dz observado','dz corrigido (ruído removido)','Confiab. (1 medida)'],rows,
      'Tabela 12 — Efeito D1→D7 corrigido pela atenuação do ruído')
P('Corrigido pela confiabilidade de uma medida (r_xx = 0,59–0,74), o efeito acumulado da semana cresce: o Vigor passa de dz = −0,95 (grande) para −1,25 (muito grande) e a Fadiga BRUMS de +0,72 para +0,91. O ruído estava encolhendo os efeitos observados; a magnitude verdadeira do "pedágio" do microciclo é maior do que a leitura bruta sugere.')
fig('den2_effect','Fig 8 — |dz| D1→D7 observado vs corrigido pela confiabilidade. O ruído atenua todos os efeitos; removê-lo revela a magnitude real.')

H('7.3. A decisão individual torna-se viável — com coletas repetidas',2)
km=DN['kmdc']; sig=DN['sigamp']
rows=[]
for k in keys:
    emg=next((kk for kk in [1,2,3,5,7] if sig[k]>km[k][str(kk)]),None)
    rows.append([labn[k],f"{sig[k]:.2f}",f"{km[k]['1']:.2f}",f"{km[k]['3']:.2f}",f"{km[k]['5']:.2f}",f"{km[k]['7']:.2f}",
                 (f"k≥{emg}" if emg else "não emerge")])
table(['Variável','Ampl. sinal','MDC k=1','MDC k=3','MDC k=5','MDC k=7','Sinal supera o ruído em'],rows,
      'Tabela 13 — MDC95 individual cai com k coletas (ruído/√k)')
P('Para um único atleta, o ruído de uma coleta é grande demais (MDC k=1 > amplitude do sinal). Mas o ruído cai por √k com médias de k coletas: para o Vigor e a Fadiga BRUMS, a partir de ~5 coletas a oscilação semanal supera o MDC individual — a decisão individual passa a ser confiável. O PTH/TMD (muito ruído em unidades brutas) e a Fadiga mental (sem sinal) não emergem nem com 7 coletas: a leitura individual deles deve apoiar-se na Fadiga física/BRUMS, não no TMD isolado.')
fig('den3_mdc_k','Fig 9 — MDC95 individual em função do número de coletas (k). A linha branca é a amplitude do sinal semanal: onde a curva do MDC cruza abaixo dela, a mudança individual torna-se detectável.')

H('7.4. Ressalva: remover ruído não inventa sinal',2)
P('Denoising esclarece o que existe; não cria o que não existe. As subescalas negativas com efeito de piso permanecem planas mesmo sem ruído: Tensão amplitude 0,60 (50% no zero), Depressão 0,84 (67%), Raiva 1,81 (60%), Confusão 0,80 (80%). Não há trajetória escondida sob o ruído nessas variáveis — a ausência de resposta é real (efeito piso), confirmada pelo fator de Bayes favorável ao nulo na Confusão. Portanto: removido o ruído, as conclusões do estudo não mudam de direção — ficam mais nítidas. O eixo energia–fadiga é um sistema bipolar forte e responsivo; as negativas continuam silenciosas.')

# ===== 8. Síntese =====
H('8. Síntese')
for t in [
 'Amplitude — grupo: o sinal semanal move o Vigor 2,8, a Fadiga 3,7 e o PTH/TMD 6,2 pontos; a Fadiga mental quase não se move (0,9).',
 'Amplitude — atleta: cada atleta oscila 2–3× mais que a trajetória do grupo, porque carrega o próprio ruído; a heterogeneidade entre atletas é alta e persistente em toda a semana.',
 'Ruído vs análise: só 1–5% da variância é o sinal do microciclo; 54–73% é traço estável e 26–41% é ruído. A Fadiga BRUMS tem o melhor sinal semanal; a Fadiga mental é quase pura estrutura de traço.',
 'Entre dias: o maior passo da trajetória (Fadiga, TMD) é o D6→D7 — a precipitação do Dia 7; nenhum dia domina a dispersão entre atletas.',
 'Decisão: interprete a média do grupo livremente (sinal 4–8× o ruído da média); para o indivíduo, exija mudança acima do MDC95 individual ou use médias de ≥3 coletas. A Fadiga mental não serve de marcador de carga aguda.',
 'Removendo o ruído: as associações fortalecem-se (Vigor × Fadiga −0,44 → −0,93) e os efeitos crescem (Vigor dz −0,95 → −1,25); a decisão individual torna-se viável a partir de ~5 coletas (Vigor, Fadiga). Mas o ruído removido não vira sinal: as subescalas de piso continuam planas — as conclusões ficam mais nítidas, não diferentes.',
]:
    p=doc.add_paragraph(style='List Bullet'); p.add_run(t)

doc.add_paragraph()
P('Reprodutibilidade: scripts/analise/amplitude.py, amp_docx_compute.py, prepost_compute.py, denoise.py, build_amp_docx.py. Consistente com Analise_Amplitude_Ruido_Sinal.md e Analise_Estatistica_Consolidada.md. Atletas anonimizados; nenhum dado bruto com nomes é distribuído.',it=True)

OUTP='/home/user/mdlucca/Artigos/Amplitude_Ruido_Sinal_Vigor_Fadiga_TMD.docx'
doc.save(OUTP)
print('SAVED',OUTP)
d2=Document(OUTP); print('paras',len(d2.paragraphs),'tables',len(d2.tables),'imgs',sum(1 for s in d2.inline_shapes))
