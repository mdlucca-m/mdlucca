# -*- coding: utf-8 -*-
import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
T=json.load(open('temporal.json')); FG='/home/user/mdlucca/Artigos/figuras'
DIMS=T['dims']; KEYS=[k for k,_ in DIMS]; LABS=[l for _,l in DIMS]
def c2(s): return str(s).replace('.',',')
def num(x,d=1): return 'n/d' if x is None else c2(f'{x:.{d}f}')
def sg(x,d=2): return c2(f'{x:+.{d}f}')
def pvt(p): return '< 0,001' if p<0.001 else c2(f'{p:.3f}')
doc=Document()
st=doc.styles['Normal']; st.font.name='Times New Roman'; st.font.size=Pt(12)
st.element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman'); st.paragraph_format.line_spacing=1.5; st.paragraph_format.space_after=Pt(0)
sec=doc.sections[0]; sec.top_margin=Cm(2.3); sec.left_margin=Cm(2.2); sec.bottom_margin=Cm(2); sec.right_margin=Cm(2.2)
_TN=[0]; _FN=[0]
def P(t='',after=0,ind=True):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(12); r.font.name='Times New Roman'
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(after); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if ind: p.paragraph_format.first_line_indent=Cm(1.25)
def H(t,before=10):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(12); r.font.name='Times New Roman'
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(4)
def _bd(c):
    tcPr=c._tc.get_or_add_tcPr(); bd=OxmlElement('w:tcBorders')
    for e in ['top','bottom']:
        el=OxmlElement('w:'+e); el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'6'); el.set(qn('w:color'),'000000'); bd.append(el)
    tcPr.append(bd)
def table(cap,header,rows,fs=8.5,note=None):
    _TN[0]+=1
    p=doc.add_paragraph(); r=p.add_run('Tabela %d – %s'%(_TN[0],cap)); r.font.size=Pt(11); r.font.name='Times New Roman'
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
    t=doc.add_table(rows=1,cols=len(header)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=True
    for i,ht in enumerate(header):
        cc=t.rows[0].cells[i]; cc.text=''; rr=cc.paragraphs[0].add_run(ht); rr.bold=True; rr.font.size=Pt(fs); rr.font.name='Times New Roman'
        cc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if i>0 else WD_ALIGN_PARAGRAPH.LEFT; _bd(cc)
    for row in rows:
        cs=t.add_row().cells
        for i,v in enumerate(row):
            cs[i].text=''; rr=cs[i].paragraphs[0].add_run(str(v)); rr.font.size=Pt(fs); rr.font.name='Times New Roman'
            cs[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if i>0 else WD_ALIGN_PARAGRAPH.LEFT; _bd(cs[i])
    if note:
        pn=doc.add_paragraph(); rn=pn.add_run(note); rn.font.size=Pt(8.5); rn.italic=True; rn.font.name='Times New Roman'; pn.paragraph_format.space_after=Pt(6)
    else: doc.add_paragraph().paragraph_format.space_after=Pt(2)
def figure(path,cap,w=17.0):
    _FN[0]+=1
    pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; pp.add_run().add_picture(path,width=Cm(w)); pp.paragraph_format.space_before=Pt(6)
    pc=doc.add_paragraph(); pc.alignment=WD_ALIGN_PARAGRAPH.CENTER; rc=pc.add_run('Figura %d – %s'%(_FN[0],cap)); rc.font.size=Pt(11); rc.font.name='Times New Roman'
    pc.paragraph_format.space_after=Pt(6)

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('ANÁLISE TEMPORAL DO HUMOR: EFEITO AGUDO E CRÔNICO, DERIVADAS, LIMIARES E APTIDÃO'); r.bold=True; r.font.size=Pt(13); p.paragraph_format.space_after=Pt(8)
H('Objetivo',before=2)
P('Mapear, para todas as dimensões do BRUMS, onde e quanto o humor muda ao longo da semana: em cada transição '
 'sequencial (efeito agudo, do pré ao pós-treino, e efeito de recuperação, do pós de um dia ao pré do dia seguinte), '
 'no efeito agudo médio comparado ao efeito crônico, na forma das curvas por derivadas, e com o ajuste pela aptidão '
 '(pico de velocidade no T-car).')

# 1) MAPA + transições significativas
H('1 Mapa temporal: onde e quanto cada variável muda')
P('A Figura 1 traz o mapa completo das doze transições sequenciais para todas as dimensões, com o tamanho de efeito '
 '(dz) de cada mudança e a marcação das significativas. A Tabela 1 lista apenas as transições significativas, com a '
 'variável, o dia, o momento e o tamanho de efeito.')
figure(f'{FG}/temporal_map.png','Mapa temporal das doze transições sequenciais (dz por transição; * p < 0,05; A = agudo pré para pós; R = recuperação pós para o pré do dia seguinte).',w=17.5)
rows=[]
for t in T['trans']:
    for k,l in DIMS:
        v=t['vars'][k]
        if v['sig']: rows.append([t['lab'],t['tipo'],l,sg(v['dm'],2),sg(v['dz']),pvt(v['p'])])
table('Transições com mudança significativa: exatamente onde (dia e momento), em qual variável e quanto (tamanho de efeito).',
 ['Transição','Tipo','Variável','Δ médio','dz','p'],rows,fs=8,
 note='Wilcoxon pareado. Agudo = pré para pós no mesmo dia; Recuperação = pós para o pré do dia seguinte.')

# 2) PRÉ/PÓS por dia
H('2 Efeito agudo por dia (pré e pós-treino, D2 a D7)')
bd=T['prepos_dia']
def cell(v): return sg(v['dz'])+('*' if v['sig'] else '')
table('Efeito agudo (pré para pós) por dia e por dimensão (tamanho de efeito dz; * p < 0,05).',
 ['Dia']+LABS,[['D%d'%d]+[cell(bd[str(d)][k]) for k in KEYS] for d in range(2,8)],fs=8,
 note='dz do pré para o pós no mesmo dia (Wilcoxon pareado). Valores positivos = aumento do pré para o pós.')
sw=T['prepos_semana']
table('Efeito agudo médio da semana (pré para pós, todas as coletas).',
 ['Dimensão','dz','p'],[[l,sg(sw[k]['dz']),pvt(sw[k]['p'])] for k,l in DIMS],fs=9)

# 3) AGUDO vs CRÔNICO
H('3 Efeito agudo comparado ao efeito crônico')
P('A Figura 2 e a Tabela 4 contrastam o efeito agudo médio dentro do dia com o efeito crônico acumulado entre o '
 'primeiro e o último dia. Em todas as dimensões do eixo energia–fadiga, o efeito crônico supera o agudo, sinal de '
 'que a mudança se acumula ao longo da semana e não se resume à resposta imediata de uma sessão.')
figure(f'{FG}/temporal_agudo_cronico.png','Efeito agudo (média intradia pré para pós) e efeito crônico (mudança D1 para D7) por variável (* efeito crônico significativo).',w=15.5)
ac=T['agudo_cronico']
table('Efeito agudo médio (intradia), efeito de recuperação médio (entre dias) e efeito crônico (D1 para D7).',
 ['Dimensão','Agudo (intradia)','Recuperação (entre dias)','Crônico (D1–D7)','dz crônico','p crônico'],
 [[l,sg(ac[k]['agudo_medio']),sg(ac[k]['recuperacao_media']),sg(ac[k]['cronico_delta']),sg(ac[k]['cronico_dz']),pvt(ac[k]['cronico_p'])] for k,l in DIMS],fs=8.5)

# 4) DERIVADAS e LIMITES
H('4 Limites e derivadas por variável')
P('A Tabela 5 resume, para cada dimensão, os limites da sessão (mínimo e máximo da curva suavizada), o ponto de '
 'inflexão (raiz da segunda derivada, onde a taxa de variação atinge o seu limite) e o pico de velocidade da mudança '
 '(máximo da primeira derivada), com o dia em que ocorre.')
dv=T['derivadas']
def inf(k):
    v=dv[k]['inflexao']; return ('Dia %s'%num(v[0],1)) if v else 'n/d'
table('Limites, ponto de inflexão e pico de velocidade da mudança de cada dimensão (curvas suavizadas).',
 ['Dimensão','Mín.','Máx.','Inflexão','Vel. pico','Dia da vel. pico'],
 [[l,num(dv[k]['min'],1),num(dv[k]['max'],1),inf(k),sg(dv[k]['vel_pico'],1),num(dv[k]['vel_pico_dia'],1)] for k,l in DIMS],fs=8.5,
 note='Vel. pico = máximo do módulo da primeira derivada (pontos por dia). A figura das derivadas de todas as variáveis consta do documento de resultados.')

# 5) LIMIARES de mudança confiável
H('5 Limiares de mudança confiável (MDC)')
lim=T['limiares']
table('Erro-padrão de medida (SEM) e mudança mínima detectável (MDC95) por dimensão: o quanto uma variação precisa exceder para ser considerada real.',
 ['Dimensão','ICC','SEM','MDC95','SWC'],
 [[l,num(lim[k]['icc'],2),num(lim[k]['sem'],1),num(lim[k]['mdc95'],1),num(lim[k]['swc'],1)] for k,l in DIMS if k in lim],fs=9,
 note='SEM = DP × √(1 - ICC); MDC95 = 1,96 × √2 × SEM; SWC = 0,2 × DP entre atletas. Uma mudança individual só é real se exceder a MDC95.')

# 6) APTIDÃO
H('6 Ajuste pela aptidão: mais aptos vs menos aptos')
P('Os atletas foram divididos pela mediana do pico de velocidade (%s km/h). A Figura 3 e a Tabela 7 comparam a '
 'deterioração (D1 para D7) entre os mais e os menos aptos. Os menos aptos tenderam a perder mais vigor e a acumular '
 'mais perturbação, mas nenhuma diferença alcançou significância (todos p > 0,13), e a correlação entre o pico de '
 'velocidade e a deterioração foi fraca em todas as dimensões. A aptidão define o patamar do humor na semana, e não a '
 'magnitude da sua deterioração.'%num(T['pv_mediana'],1))
figure(f'{FG}/temporal_aptidao.png','Deterioração D1 para D7 por variável nos atletas mais aptos e menos aptos (mediana do pico de velocidade).',w=15.5)
sT=T['estratificacao']
table('Deterioração D1 para D7 nos mais e menos aptos, com a comparação entre grupos e a correlação com o pico de velocidade.',
 ['Dimensão','Mais aptos (Δ)','Menos aptos (Δ)','p (grupos)','ρ (PV × Δ)','p (PV)'],
 [[l,sg(sT[k]['mais_aptos_delta'],1),sg(sT[k]['menos_aptos_delta'],1),pvt(sT[k]['p_grupos']),sg(sT[k]['rho_PV']),pvt(sT[k]['p_PV'])] for k,l in DIMS],fs=8.5,
 note='Δ = mudança D1 para D7; p (grupos) por Mann-Whitney; ρ por Spearman entre o pico de velocidade e a deterioração.')

OUTP='/home/user/mdlucca/Artigos/Resultados_Temporais_BRUMS.docx'
doc.save(OUTP); print('SAVED',OUTP,'| Tabelas',_TN[0],'Figuras',_FN[0])
