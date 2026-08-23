# -*- coding: utf-8 -*-
import json, numpy as np
import matplotlib; matplotlib.use('Agg'); import matplotlib.pyplot as plt
import pandas as pd
R=json.load(open('curvas_completa.json'))
xf=np.array(R['_xf']); x=np.arange(1,8)
DIMS=[('Vigor','#2f9e44'),('Fadiga','#e8590c'),('Tensao','#4d9de0'),('Depressao','#c56bd6'),
      ('Raiva','#e0525b'),('Confusao','#f0a848')]
h=pd.read_csv('humor_anon.csv')
def sem(v,d): x=h[h.dia==d][v].dropna(); return x.std(ddof=1)/np.sqrt(len(x))
plt.rcParams.update({'font.family':'DejaVu Sans','font.size':10.5,'axes.spines.top':False,'axes.spines.right':False})

# ===== FIG A: 6 paineis — sem filtro (pontos+linha tracejada) vs com filtro (spline) =====
fig,axs=plt.subplots(2,3,figsize=(15,8),dpi=190)
for ax,(v,c) in zip(axs.ravel(),DIMS):
    r=R[v]; raw=np.array(r['raw']); sm=np.array(R['_smoothed'][v])
    er=[sem(v,d) for d in x]
    for d in [2,4,7]: ax.axvspan(d-.15,d+.15,color='#e8590c',alpha=.07)
    ax.errorbar(x,raw,yerr=er,fmt='o',color=c,ms=6,capsize=3,alpha=.9,zorder=3,label='média diária ± EP (sem filtro)')
    ax.plot(x,raw,'--',color=c,lw=1,alpha=.5,zorder=2,label='ligação bruta')
    ax.plot(xf,sm,'-',color=c,lw=2.8,zorder=4,label='curva suavizada (com filtro)')
    if not np.isnan(r['infl']):
        yi=np.interp(r['infl'],xf,sm); ax.plot(r['infl'],yi,'v',color='black',ms=8,zorder=5)
        ax.annotate(f"inflexão\n~dia {r['infl']:.1f}",(r['infl'],yi),textcoords='offset points',xytext=(6,-24),fontsize=8,color='#333')
    sig='***' if r['fp']<0.001 else ('*' if r['fp']<0.05 else 'ns')
    ax.set_title(f"{r['lab']}  ·  Friedman p={str(round(r['fp'],3)).replace('.',',')} {sig}  ·  dz(D1→D7)={r['dz']:+.2f}".replace('.',','),
                 fontweight='bold',fontsize=10.5,loc='left',color=c)
    ax.set_xticks(x); ax.set_xlim(.7,7.3)
axs[0,0].set_ylabel('Escore'); axs[1,0].set_ylabel('Escore')
for ax in axs[1,:]: ax.set_xlabel('Dia do microciclo')
h1,l1=axs[0,0].get_legend_handles_labels()
fig.legend(h1,l1,frameon=False,fontsize=9.5,loc='lower center',ncol=3,bbox_to_anchor=(0.5,-0.02))
fig.suptitle('Comportamento de cada dimensão do BRUMS ao longo da semana — sem filtro (pontos) vs com filtro (spline) · faixas = HIIT · * p<0,05; *** p<0,001 (Friedman)',
             fontweight='bold',fontsize=12.5,y=0.995)
fig.tight_layout(rect=[0,0.03,1,1]); fig.savefig('/home/user/mdlucca/Artigos/figuras/curvas_por_variavel.png',bbox_inches='tight',facecolor='white')
print('fig A ok')

# ===== FIG B: com/sem filtro + sinal/ruido + cruzamentos (vigor x fadiga) =====
fig,ax=plt.subplots(1,2,figsize=(13.5,5),dpi=190)
# painel 1: filtragem (vigor) — observado, sinal, ruido
v='Vigor'; raw=np.array(R[v]['raw']); sm=np.array(R['_smoothed'][v]); smx=np.interp(x,xf,sm)
ax[0].plot(x,raw,'o',color='#2f9e44',ms=7,label='observado (sem filtro)')
ax[0].plot(xf,sm,'-',color='#2f9e44',lw=2.8,label='sinal (com filtro)')
ax[0].vlines(x,smx,raw,color='#adb5bd',lw=1.5,label='ruído (resíduo)')
ax[0].set_title('Filtragem da curva (exemplo: Vigor)\no filtro separa o sinal (tendência) do ruído',fontweight='bold',fontsize=11,loc='left')
ax[0].set_xticks(x); ax[0].set_xlabel('Dia'); ax[0].set_ylabel('Vigor'); ax[0].legend(frameon=False,fontsize=8.5)
# painel 2: cruzamentos vigor x fadiga (curvas filtradas)
for v,c in [('Vigor','#2f9e44'),('Fadiga','#e8590c')]:
    ax[1].plot(xf,np.array(R['_smoothed'][v]),'-',color=c,lw=2.8,label=v)
    ax[1].plot(x,np.array(R[v]['raw']),'o',color=c,ms=4,alpha=.5)
cr=R['_cross']['VigorFadiga']
for d in [2,4,7]: ax[1].axvspan(d-.15,d+.15,color='#e8590c',alpha=.07)
for xc in cr:
    yc=np.interp(xc,xf,np.array(R['_smoothed']['Vigor']))
    ax[1].plot(xc,yc,'X',color='black',ms=11,zorder=5)
    ax[1].annotate(f'dia {xc}'.replace('.',','),(xc,yc),textcoords='offset points',xytext=(4,8),fontsize=8.5,fontweight='bold')
ax[1].set_title(f'Cruzamentos Vigor × Fadiga (curvas filtradas)\npontos de ligação: dias {", ".join(str(c).replace(".",",") for c in cr)}',fontweight='bold',fontsize=11,loc='left')
ax[1].set_xticks(x); ax[1].set_xlabel('Dia'); ax[1].set_ylabel('Escore'); ax[1].legend(frameon=False,fontsize=9)
fig.suptitle('Com filtro vs sem filtro, sinal/ruído e o cruzamento do eixo energia–fadiga',fontweight='bold',fontsize=12.5,y=1.02)
fig.tight_layout(); fig.savefig('/home/user/mdlucca/Artigos/figuras/curvas_filtro_cruzamento.png',bbox_inches='tight',facecolor='white')
print('fig B ok')

# ===== TABELA docx =====
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
def cm(x): return str(x).replace('.',',')
doc=Document(); st=doc.styles['Normal']; st.font.name='Times New Roman'; st.font.size=Pt(11)
p=doc.add_paragraph(); rr=p.add_run('Tabela. Comportamento de cada dimensão do BRUMS ao longo da semana: mudança, significância (Friedman e Wilcoxon), forma temporal (inflexão) e qualidade do sinal.'); rr.bold=True; rr.font.size=Pt(10.5)
hdr=['Dimensão','D1','D7','Δ','dz (D1→D7)','Wilcoxon p','Friedman χ²','W','Friedman p','Inflexão (dia)','Sinal']
t=doc.add_table(rows=1,cols=len(hdr)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
def setc(c,txt,b=False,sz=9,al='center'):
    c.text=''; pp=c.paragraphs[0]; pp.alignment={'center':WD_ALIGN_PARAGRAPH.CENTER,'left':WD_ALIGN_PARAGRAPH.LEFT}[al]
    z=pp.add_run(str(txt)); z.font.size=Pt(sz); z.font.name='Times New Roman'; z.bold=b
for j,hh in enumerate(hdr): setc(t.rows[0].cells[j],hh,b=True,sz=8.5)
def sigmark(p): return '***' if p<0.001 else ('**' if p<0.01 else ('*' if p<0.05 else 'ns'))
snrq=lambda s: 'alto' if s>=8 else ('moderado' if s>=5 else 'baixo')
for v,_ in DIMS:
    r=R[v]; c=t.add_row().cells
    vals=[r['lab'],cm(round(r['d1'],1)),cm(round(r['d7'],1)),cm(round(r['d7']-r['d1'],1)),
          f"{r['dz']:+.2f}".replace('.',','),cm(round(r['wp'],3))+' '+('*' if r['wp']<0.05 else ''),
          cm(round(r['chi'],1)),cm(round(r['W'],2)),cm(round(r['fp'],3))+' '+sigmark(r['fp']),
          cm(round(r['infl'],1)),snrq(r['snr'])]
    for j,x2 in enumerate(vals): setc(c[j],x2,al='left' if j==0 else 'center',sz=8.5)
nt=doc.add_paragraph(); z=nt.add_run('Δ: variação D1→D7; dz: tamanho de efeito pareado (D1→D7); Wilcoxon: contraste pareado D1 vs D7; Friedman (χ², W de Kendall, p): '
 'teste de medidas repetidas ao longo dos sete dias (subconjunto balanceado); * p<0,05, ** p<0,01, *** p<0,001, ns não significativo. Inflexão: dia em que a '
 '2ª derivada da curva suavizada (spline) cruza zero; os valores exatos de inflexão e dos cruzamentos dependem da intensidade do filtro, ao passo que os testes '
 '(Friedman/Wilcoxon), calculados sobre os dados brutos, são independentes do suavizador. Sinal: qualidade da relação sinal-ruído (alto ≥ 8; moderado 5–8; baixo < 5). '
 'Cruzamentos Vigor × Fadiga (curvas filtradas): dias '+', '.join(str(c).replace('.',',') for c in R['_cross']['VigorFadiga'])+'.')
z.italic=True; z.font.size=Pt(9)
doc.save('/home/user/mdlucca/Artigos/Tabela_Comportamento_Curvas.docx')
print('tabela ok')
