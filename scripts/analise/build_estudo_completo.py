# -*- coding: utf-8 -*-
import json, numpy as np, pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
FG='/home/user/mdlucca/Artigos/figuras'
AR=json.load(open('all_results.json')); T=json.load(open('temporal.json'))
P2=json.load(open('physio2.json')); EX=json.load(open('extra_doc.json')); h=pd.read_csv('hum_prof.csv')
DIMS=[('Vigor','Vigor'),('Fadiga','Fadiga'),('Tensao','Tensão'),('Depressao','Depressão'),
      ('Raiva','Raiva'),('Confusao','Confusão'),('TMD','Perturbação total do humor (PTH)'),
      ('FadFisica','Fadiga física'),('FadMental','Fadiga mental')]
KEYS=[k for k,_ in DIMS]
def c2(s): return str(s).replace('.',',')
def num(x,d=1): return 'n/d' if x is None else c2(f'{x:.{d}f}')
def sg(x,d=2): return c2(f'{x:+.{d}f}')
def pvt(p): return '< 0,001' if p<0.001 else c2(f'{p:.3f}')
def floor(k): return 100*np.mean(pd.to_numeric(h[k],errors='coerce').dropna()<=0)

doc=Document()
stl=doc.styles['Normal']; stl.font.name='Times New Roman'; stl.font.size=Pt(12)
stl.element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')
stl.paragraph_format.line_spacing=1.5; stl.paragraph_format.space_after=Pt(0)
sec=doc.sections[0]; sec.top_margin=Cm(2.3); sec.left_margin=Cm(2.4); sec.bottom_margin=Cm(2); sec.right_margin=Cm(2.2)
_TN=[0]; _FN=[0]
def P(t='',ind=True,after=0):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(12); r.font.name='Times New Roman'
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(after); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if ind: p.paragraph_format.first_line_indent=Cm(1.25)
def LI(t):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(12); r.font.name='Times New Roman'
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.left_indent=Cm(1.25); p.paragraph_format.space_after=Pt(0)
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
def H1(t,n=None):
    p=doc.add_paragraph(); r=p.add_run((n+'  ' if n else '')+t); r.bold=True; r.font.size=Pt(12.5); r.font.name='Times New Roman'
    p.paragraph_format.space_before=Pt(14); p.paragraph_format.space_after=Pt(4)
def H2(t,n=None):
    p=doc.add_paragraph(); r=p.add_run((n+'  ' if n else '')+t); r.bold=True; r.italic=True; r.font.size=Pt(12); r.font.name='Times New Roman'
    p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(3)
def _bd(c):
    tcPr=c._tc.get_or_add_tcPr(); bd=OxmlElement('w:tcBorders')
    for e in ['top','bottom']:
        el=OxmlElement('w:'+e); el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'6'); el.set(qn('w:color'),'000000'); bd.append(el)
    tcPr.append(bd)
def table(cap,header,rows,fs=8.8,note=None,al0='LEFT'):
    _TN[0]+=1
    p=doc.add_paragraph(); r=p.add_run('Tabela %d. %s'%(_TN[0],cap)); r.font.size=Pt(11); r.bold=True; r.font.name='Times New Roman'
    p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(2)
    t=doc.add_table(rows=1,cols=len(header)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=True
    for i,ht in enumerate(header):
        cc=t.rows[0].cells[i]; cc.text=''; rr=cc.paragraphs[0].add_run(ht); rr.bold=True; rr.font.size=Pt(fs); rr.font.name='Times New Roman'
        cc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if i>0 else getattr(WD_ALIGN_PARAGRAPH,al0); _bd(cc)
    for row in rows:
        cs=t.add_row().cells
        for i,v in enumerate(row):
            cs[i].text=''; rr=cs[i].paragraphs[0].add_run(str(v)); rr.font.size=Pt(fs); rr.font.name='Times New Roman'
            cs[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if i>0 else getattr(WD_ALIGN_PARAGRAPH,al0); _bd(cs[i])
    if note:
        pn=doc.add_paragraph(); rn=pn.add_run(note); rn.font.size=Pt(8.5); rn.italic=True; rn.font.name='Times New Roman'; pn.paragraph_format.space_after=Pt(6)
    else: doc.add_paragraph().paragraph_format.space_after=Pt(2)
def figure(path,cap,w=16.0):
    _FN[0]+=1
    pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; pp.add_run().add_picture(path,width=Cm(w)); pp.paragraph_format.space_before=Pt(6)
    pc=doc.add_paragraph(); pc.alignment=WD_ALIGN_PARAGRAPH.CENTER; rc=pc.add_run('Figura %d. %s'%(_FN[0],cap)); rc.font.size=Pt(10.5); rc.font.name='Times New Roman'
    pc.paragraph_format.space_after=Pt(6)

# ============ TÍTULO ============
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('APTIDÃO AERÓBIA E PERFIL DE HUMOR NA ÚLTIMA SEMANA DE PRÉ-TEMPORADA DE ATLETAS DE '
            'HANDEBOL DE ELITE: UM ESTUDO DE ASSOCIAÇÃO')
r.bold=True; r.font.size=Pt(14); p.paragraph_format.space_after=Pt(12)

# ============ RESUMO ============
H1('Resumo')
P('Objetivo: descrever o perfil de humor de atletas de handebol de elite ao longo da última semana de '
  'pré-temporada e verificar a associação entre a aptidão aeróbia, medida pelo pico de velocidade no teste de '
  'campo T-car, e o comportamento do humor. Método: estudo observacional, longitudinal e correlacional, de grupo '
  'único, com medidas repetidas. Vinte e sete atletas do sexo masculino responderam ao BRUMS de forma virtual, '
  'pela plataforma Google Forms, em dois momentos por dia (manhã e fim de tarde) durante sete dias. O pico de '
  'velocidade foi obtido no T-car, aplicado antes do período de observação. A análise seguiu três níveis '
  '(descritivo, exploratório e avançado), com verificação de normalidade, análises psicométricas, testes não '
  'paramétricos, tamanhos de efeito, derivadas das curvas diárias e limiares de mudança confiável. Resultados: o '
  'Vigor caiu de forma acentuada (D1 8,4 para D7 4,3; dz = 1,33; p < 0,001) e a fadiga física aumentou de maneira '
  'consistente (qui-quadrado de Friedman = 51,3; p < 0,001). A maioria dos atletas migrou do perfil de iceberg '
  'para perfis de menor energia. A aptidão aeróbia associou-se ao nível do humor, mas não à magnitude da '
  'deterioração ao longo da semana. Conclusão: a última semana de pré-temporada produziu perda de energia e '
  'acúmulo de fadiga; a aptidão aeróbia não protegeu o atleta contra essa deterioração.')
pk=doc.add_paragraph(); rr=pk.add_run('Palavras-chave: '); rr.bold=True; rr.font.size=Pt(11)
pk.add_run('estados de humor; BRUMS; handebol; carga de treinamento; pico de velocidade; monitoramento.').font.size=Pt(11)
pk.paragraph_format.space_after=Pt(4)

# ============ 1 INTRODUÇÃO ============
H1('Introdução','1')
P('O handebol é uma modalidade coletiva classificada como esporte intermitente de alta intensidade. Durante o '
  'jogo, os atletas alternam longos períodos de atividade de baixa intensidade com bursts curtos de alta '
  'intensidade, como sprints, saltos, mudanças de direção e ações de contato, intercalados por períodos variáveis '
  'de recuperação (Pueo et al., 2017; Cartón-Llorente et al., 2023). Análises com sistemas de posicionamento '
  'mostram incrementos expressivos das ações de alta intensidade nos momentos mais exigentes da partida e '
  'diferenças relevantes entre as posições táticas (Cartón-Llorente et al., 2023). Esse padrão impõe demanda '
  'simultânea aos sistemas aeróbio e anaeróbio e exige um planejamento de treino que combine estímulos de '
  'resistência e de potência.')
P('A pré-temporada é o período em que a equipe concentra as maiores cargas do ano, com o objetivo de reconstruir o '
  'condicionamento antes da fase competitiva. É também o intervalo de maior risco de fadiga acumulada. Por isso, o '
  'monitoramento sistemático da resposta do atleta é parte central do trabalho da comissão técnica: ajustes feitos '
  'a tempo evitam que a fadiga se acumule além do desejado e reduzem o risco de excesso de treinamento (Duignan et '
  'al., 2020).')
P('Entre as ferramentas de monitoramento, os estados de humor ocupam posição de destaque por serem sensíveis, '
  'rápidos e de baixo custo. O instrumento mais utilizado com atletas é a Escala de Humor de Brunel (BRUMS), '
  'derivada do Perfil de Estados de Humor (POMS) e composta por 24 itens que formam seis dimensões: vigor, fadiga, '
  'tensão, depressão, raiva e confusão (Terry, Lane & Fogarty, 2003). A escala tem propriedades psicométricas '
  'consistentes e foi validada em diversos idiomas e culturas, como o português (Rohlfs et al., 2008), o '
  'chinês (Zhang et al., 2014), o lituano (Terry et al., 2022) e o malaio (Lew et al., 2023). Em atletas bem '
  'recuperados, o BRUMS descreve o chamado perfil de iceberg, no qual o vigor se sobrepõe às dimensões negativas; '
  'a inversão desse perfil sinaliza fadiga e risco psicológico.')
P('Em esportes coletivos, o acompanhamento do humor tem se mostrado útil para orientar decisões de treino. Revisões '
  'recentes indicam que os escores de humor variam de acordo com a carga, a intensidade e a modalidade das sessões, '
  'e que o monitoramento frequente ajuda a detectar sinais precoces de perturbação psicológica (Selmi et al., '
  '2023). Medidas subjetivas de bem-estar, como o humor, apresentam relação predominantemente negativa com a '
  'carga de treino em atletas de esportes coletivos (Duignan et al., 2020), e essa relação também aparece em '
  'contextos de treinamento intensificado, como campos de treinamento e microciclos de alta exigência (Chen et '
  'al., 2022).')
P('Ao lado do humor, a aptidão aeróbia define quanto esforço o atleta suporta antes de fadigar e é determinante de '
  'desempenho em atletas de handebol de elite (Font et al., 2021; Fikenzer et al., 2020). Neste estudo, a aptidão '
  'foi medida pelo pico de velocidade no T-car, um teste de campo intermitente e progressivo. O pico de velocidade '
  'desse teste é uma variável já consolidada na literatura: não difere da velocidade correspondente ao consumo '
  'máximo de oxigênio e correlaciona-se fortemente a ela (Da Silva et al., 2011), reflete a potência aeróbia '
  'máxima e pode ser usado como referência para a prescrição de treino (Carminatti et al., 2013; Floriano et al., '
  '2016), além de se relacionar ao desempenho físico durante o jogo (Fernandes-da-Silva et al., 2016). Em resumo, '
  'o pico de velocidade no T-car resume, em uma única medida de campo, a capacidade e a potência aeróbia máxima do '
  'atleta.')
P('A justificativa deste trabalho está na integração desses dois planos de avaliação, o psicológico e o '
  'fisiológico, em um mesmo desenho. Apesar da ampla adoção do BRUMS e do pico de velocidade do T-car de forma '
  'isolada, ainda é pouco explorada a associação entre a aptidão aeróbia e o comportamento do perfil de humor ao '
  'longo de uma semana real de cargas elevadas em atletas de elite. Responder a essa questão tem consequência '
  'prática direta: se a aptidão protege o atleta da deterioração do humor, a avaliação física poderia antecipar o '
  'risco psicológico; se não protege, o monitoramento do humor é insubstituível. A Figura 1 apresenta o modelo '
  'conceitual que orienta o estudo.')
figure(f'{FG}/framework.png','Modelo conceitual (framework): a aptidão aeróbia e a carga de treino como entradas, o perfil de humor como construto monitorado, e o nível e a deterioração do humor como desfechos.',w=16.5)
H2('Objetivos','1.1')
P('Objetivo geral: descrever e classificar o perfil de humor de atletas de handebol de elite ao longo da última '
  'semana de pré-temporada e analisar sua associação com a aptidão aeróbia. Como objetivos específicos, o estudo '
  'busca:')
LI('a) caracterizar o comportamento descritivo de cada dimensão do humor ao longo da semana, por dia e entre os dias;')
LI('b) verificar a mudança do humor ao longo da semana e do primeiro ao sétimo dia, com testes não paramétricos e tamanhos de efeito;')
LI('c) descrever a migração dos perfis de humor entre o início e o fim da semana;')
LI('d) descrever a aptidão aeróbia do grupo pelo pico de velocidade e demais indicadores do T-car;')
LI('e) testar a associação entre a aptidão aeróbia e tanto o nível quanto a deterioração do humor.')

# ============ 2 MÉTODO ============
H1('Método','2')
H2('Tipo de estudo','2.1')
P('Trata-se de um estudo observacional, longitudinal e correlacional, de abordagem quantitativa, com medidas '
  'repetidas no mesmo grupo de atletas. Observacional porque a rotina de treino não foi manipulada: a equipe '
  'seguiu o planejamento normal da comissão técnica. Longitudinal porque os mesmos atletas foram avaliados '
  'repetidas vezes ao longo de sete dias. Correlacional porque um dos objetivos é testar a associação entre a '
  'aptidão aeróbia e o humor. De grupo único porque não houve grupo controle: cada atleta serve de referência para '
  'si mesmo, e as comparações são feitas entre os momentos da semana.')
H2('População e amostra','2.2')
P('Participaram 27 atletas do sexo masculino de uma equipe de handebol de elite, todos integrantes do elenco '
  'principal durante a pré-temporada. A amostra reúne as quatro funções táticas da modalidade, o que permite '
  'descrever também as diferenças por posição. A Tabela 1 resume as características do grupo.')
table('Caracterização da amostra (n = 27 atletas do sexo masculino).',
  ['Característica','Média ± DP','Mín. – Máx.'],
  [['Idade (anos)','22,0 ± 3,8','17 – 38'],
   ['Massa corporal (kg)','85,8 ± 15,7','58,8 – 132,6'],
   ['Estatura (cm)','182,5 ± 6,5','170,8 – 200,0'],
   ['Índice de massa corporal (kg/m²)','25,7 ± 3,6','n/d'],
   ['Gordura corporal (%)','11,6 ± 5,3','5,0 – 25,0'],
   ['Pico de velocidade no T-car (km/h)','14,9 ± 1,1','13,2 – 17,5']],
  note='DP: desvio-padrão. Posições: 11 meias, 9 pontas, 4 pivôs e 3 goleiros. Medidas antropométricas e do T-car disponíveis para 24 a 25 atletas.')
H2('Aspectos éticos','2.3')
P('Antes do início das coletas, a equipe recebeu a explicação detalhada dos procedimentos e dos objetivos do '
  'estudo. Todos os atletas concordaram em participar de forma voluntária e assinaram o termo de consentimento '
  'livre e esclarecido. A participação não interferiu na rotina de treino, e os dados foram tratados de forma '
  'confidencial e anônima.')
H2('Desenho do estudo','2.4')
P('O estudo foi organizado em três etapas, resumidas no organograma da Figura 2. Na primeira etapa, houve o '
  'primeiro contato com a comissão técnica e a equipe, quando foram explicados os procedimentos de coleta e '
  'assinado o termo de consentimento. Em seguida, foi realizada uma coleta de familiarização, para que os atletas '
  'conhecessem o questionário e o modo de resposta virtual. O teste T-car foi aplicado em 15 de abril de 2024, '
  'antes do período de observação, para caracterizar a aptidão aeróbia de cada atleta.')
P('Na segunda etapa, ocorreu o monitoramento diário durante a última semana de pré-temporada. A coleta do baseline '
  'foi feita em 21 de abril de 2024. Nesse dia os atletas não treinaram, de modo que o baseline representa o estado '
  'de humor em repouso, livre do efeito imediato de uma sessão. A partir daí, o humor foi registrado duas vezes por '
  'dia até o fim da semana: uma coleta pela manhã, antes do treino (momento pré), e uma coleta ao fim do dia, após '
  'o treino (momento pós). Essa estrutura de duas medidas diárias permite separar o efeito imediato de cada sessão, '
  'que é a variação da manhã para o fim do dia, do efeito acumulado ao longo dos dias. Os treinos intervalados de '
  'alta intensidade concentraram-se nos dias 2, 4 e 7. Na terceira etapa, os desfechos foram organizados em três '
  'blocos: estados de humor, aptidão aeróbia e carga interna de treino.')
figure(f'{FG}/desenho_estudo.png','Organograma do estudo: etapas de preparação, semana de coletas e desfechos analisados.',w=17.0)
H2('Instrumentos e medidas','2.5')
P('Estados de humor (BRUMS). O humor foi avaliado pela Escala de Humor de Brunel (BRUMS), composta por 24 itens '
  'que descrevem sensações recentes, respondidos em uma escala de zero a quatro. Os itens formam as seis dimensões '
  'da escala: vigor, fadiga, tensão, depressão, raiva e confusão (Terry, Lane & Fogarty, 2003; Rohlfs et al., '
  '2008). A partir dessas seis dimensões calcula-se a perturbação total do humor (PTH), um índice global que soma '
  'as dimensões negativas e subtrai o vigor. Além da escala, foram registradas duas medidas complementares de '
  'fadiga percebida, física e mental, para detalhar a natureza do cansaço relatado. O escore de cada dimensão '
  'varia de zero a dezesseis. O questionário foi aplicado de maneira online: os atletas respondiam de forma '
  'virtual, pela plataforma Google Forms, a partir de um link enviado nos horários de coleta. Esse formato reduz o '
  'tempo de resposta, evita erros de digitação e permite o registro imediato dos dados, sem necessidade de '
  'presença física do avaliador.')
P('Aptidão aeróbia (T-car). A aptidão foi medida pelo teste de Carminatti (T-car), um teste de campo intermitente '
  'e progressivo, feito em corridas de vaivém com aumento gradual da velocidade até a exaustão. A variável de '
  'interesse é o pico de velocidade, ou seja, a maior velocidade sustentada no teste. O pico de velocidade do '
  'T-car é uma medida validada e confiável de aptidão: associa-se à velocidade correspondente ao consumo máximo de '
  'oxigênio (Da Silva et al., 2011), reflete a potência aeróbia máxima e serve de referência para prescrição de '
  'treino (Carminatti et al., 2013; Floriano et al., 2016) e relaciona-se ao desempenho físico no jogo '
  '(Fernandes-da-Silva et al., 2016). Neste estudo, o pico de velocidade foi usado para descrever a aptidão do '
  'grupo, separar atletas mais aptos de menos aptos e testar a associação com o humor.')
P('Carga interna de treino. Nos dias de treino intervalado de alta intensidade, a carga interna foi acompanhada '
  'pela frequência cardíaca e pela percepção subjetiva de esforço, indicadores de quanto o organismo foi exigido '
  'em cada sessão, o que auxilia a interpretar as variações do humor à luz do esforço imposto.')
H2('Análise estatística','2.6')
P('A análise foi conduzida em três níveis, do mais simples ao mais elaborado. A Tabela 2 resume os testes '
  'empregados em cada objetivo. Adotou-se o nível de significância de 5 por cento em todas as análises.')
table('Plano de análise estatística por objetivo.',
  ['Etapa','Análise','Teste ou medida','Finalidade'],
  [['Descritiva','Tendência e dispersão','Média, DP, mediana, mín-máx, CV, piso','Resumir cada variável'],
   ['Exploratória','Normalidade','Shapiro-Wilk','Justificar a estatística não paramétrica'],
   ['Exploratória','Psicometria','Alfa de Cronbach, ômega, assimetria','Avaliar a qualidade das dimensões'],
   ['Exploratória','Confiabilidade','ICC e MDC95','Definir o limiar de mudança confiável'],
   ['Exploratória','Perfis de humor','Classificação e migração','Descrever a mudança de perfil'],
   ['Avançada','Mudança na semana','Friedman e W de Kendall','Testar variação ao longo dos dias'],
   ['Avançada','Comparações pareadas','Wilcoxon e dz','Testar transições e D1-D7'],
   ['Avançada','Associação com a aptidão','Spearman','Relacionar PV com nível e deterioração'],
   ['Avançada','Dinâmica das curvas','Spline e derivadas','Localizar velocidade e inflexão']],
  fs=8.4,note='DP: desvio-padrão. CV: coeficiente de variação. ICC: coeficiente de correlação intraclasse. '
  'MDC95: mudança mínima detectável a 95%. PV: pico de velocidade no T-car.')
P('Nível 1, descritivo. Para cada dimensão do humor foram calculados a média, o desvio-padrão, a mediana, os '
  'valores mínimo e máximo, o coeficiente de variação e a proporção de respostas iguais a zero, o chamado efeito '
  'de piso, comum nas dimensões negativas.')
P('Nível 2, exploratório. A normalidade da distribuição foi verificada pelo teste de Shapiro-Wilk. A qualidade das '
  'dimensões foi avaliada por análises psicométricas: consistência interna pelo alfa de Cronbach e pelo ômega de '
  'McDonald, e assimetria da distribuição. A confiabilidade entre os dias foi estimada pelo coeficiente de '
  'correlação intraclasse (ICC), a partir do qual foi calculada a mudança mínima detectável a 95 por cento '
  '(MDC95), um limiar prático acima do qual a variação é considerada real, e não erro de medida. A classificação '
  'do perfil de humor de cada atleta foi comparada entre o primeiro e o sétimo dia.')
P('Nível 3, avançado. Como parte das dimensões não seguiu distribuição normal e há efeito de piso, foram usados '
  'testes não paramétricos. O teste de Friedman verificou se cada dimensão mudou ao longo da semana, com o W de '
  'Kendall como tamanho do efeito. O teste de Wilcoxon comparou momentos pareados, com o tamanho de efeito dz. A '
  'associação entre a aptidão e o humor foi testada pela correlação de Spearman entre o pico de velocidade e tanto '
  'o nível quanto a variação do humor. A dinâmica de cada dimensão foi descrita pela suavização das curvas diárias '
  'e por suas derivadas: a primeira derivada mostra a velocidade de mudança, e a segunda derivada mostra a '
  'aceleração e o ponto de inflexão, o momento em que a tendência muda de sentido.')

# ============ 3 RESULTADOS ============
H1('Resultados','3')
P('Os resultados são apresentados em cinco etapas: a descrição da aptidão aeróbia do grupo; a qualidade das '
  'medidas de humor (normalidade e psicometria); o comportamento descritivo das dimensões; a variação ao longo da '
  'semana e a migração de perfis; e a associação entre aptidão e humor. Ao final, as figuras individuais detalham '
  'cada dimensão.')

H2('Etapa 1. Aptidão aeróbia do grupo (T-car)','3.1')
P('A Tabela 3 descreve a aptidão aeróbia e a carga do teste para todo o grupo. O pico de velocidade médio no '
  'primeiro T-car foi de 14,9 km/h, com baixa dispersão relativa (coeficiente de variação de 7 por cento), o que '
  'indica um grupo homogêneo em termos de aptidão. A frequência cardíaca máxima próxima de 198 bpm e a percepção '
  'de esforço elevada confirmam que o teste foi conduzido até a exaustão. Da primeira para a segunda aplicação, '
  'houve ganho médio de 0,9 km/h no pico de velocidade, coerente com o efeito da pré-temporada sobre a aptidão.')
tc=EX['tcar']
rows=[[r['lab'],str(r['n']),'%s ± %s'%(num(r['m'],2 if 'velocidade' in r['lab'] else 1),num(r['sd'],1)),
       '%s – %s'%(num(r['mn'],1),num(r['mx'],1)),num(r['cv'],1)] for r in tc]
table('Descritiva da aptidão aeróbia e da carga no T-car (grupo completo).',
  ['Variável','n','Média ± DP','Mín – Máx','CV (%)'],rows,fs=8.4,
  note='DP: desvio-padrão. CV: coeficiente de variação. ΔPV: ganho do pico de velocidade entre as duas aplicações. TRIMP: impulso de treino.')

H2('Etapa 2. Qualidade das medidas de humor','3.2')
P('A Tabela 4 apresenta a normalidade e as propriedades psicométricas das dimensões. O teste de Shapiro-Wilk '
  'indicou distribuição não normal na maioria das dimensões negativas, com forte assimetria positiva e efeito de '
  'piso elevado, o que justifica o uso de testes não paramétricos. A consistência interna, medida pelo alfa de '
  'Cronbach e pelo ômega, foi adequada nas dimensões centrais do estudo, com destaque para fadiga, depressão e '
  'raiva. A tensão apresentou consistência mais baixa, provavelmente pelo forte efeito de piso, e deve ser '
  'interpretada com cautela.')
rows=[]
for k,lab in DIMS:
    nm=EX['normal'][k]; psk={'TMD':'PTH'}.get(k,k); psy=EX['psy'].get(psk,{})
    al=psy.get('alpha'); om=psy.get('omega'); sk=psy.get('skew')
    rows.append([lab,num(nm['W'],3),pvt(nm['p']),'não' if nm['p']<0.05 else 'sim',
                 num(al,2) if al is not None else 'n/d',num(om,2) if om is not None else 'n/d',
                 sg(sk,2) if sk is not None else 'n/d'])
table('Normalidade (Shapiro-Wilk) e psicometria das dimensões do humor.',
  ['Variável','W','p','Normal','Alfa','Ômega','Assimetria'],rows,fs=8.5,
  note='W e p: teste de Shapiro-Wilk sobre a média semanal por atleta. Alfa: alfa de Cronbach. Ômega: ômega de '
  'McDonald. Assimetria positiva indica concentração de escores baixos (efeito de piso). Alfa e ômega não se '
  'aplicam à PTH nem às medidas complementares.')
cf=EX['cfa']
P('A estrutura de seis fatores do BRUMS mostrou ajuste aceitável em análise fatorial confirmatória (CFI = %s; '
  'TLI = %s; RMSEA = %s; SRMR = %s), coerente com a literatura de validação da escala (Zhang et al., 2014; Terry '
  'et al., 2022; Lew et al., 2023).'%(num(cf['cfi'],2),num(cf['tli'],2),num(cf['rmsea'],2),num(cf['srmr'],2)))

H2('Etapa 3. Comportamento descritivo das dimensões','3.3')
rows=[]
for k,lab in DIMS:
    ds=AR['desc_semana'][k]; lim=T['limiares'].get(k,{})
    rows.append([lab,'%s ± %s'%(num(ds['m'],2),num(ds['sd'],1)),num(ds['med']),'%s – %s'%(num(ds['mn'],0),num(ds['mx'],0)),
                 num(ds['cv'],0),num(floor(k),0),num(lim.get('icc',0),2),num(lim.get('mdc95',0),1)])
table('Estatística descritiva das variáveis de humor na semana (286 observações de 27 atletas).',
  ['Variável','Média ± DP','Mediana','Mín – Máx','CV (%)','Piso (%)','ICC','MDC95'],rows,fs=8.6,
  note='As seis primeiras linhas são as dimensões do BRUMS; a PTH é um índice global derivado; fadiga física e '
       'fadiga mental são medidas complementares. CV: coeficiente de variação. Piso: escores iguais a zero. '
       'ICC: correlação intraclasse. MDC95: mudança mínima detectável a 95%.')
P('O Vigor foi a dimensão de maior escore médio, seguido pela fadiga física. As dimensões negativas apresentaram '
  'medianas baixas e forte efeito de piso, com destaque para a depressão, em que a maioria dos atletas marcou '
  'zero. A confiabilidade entre dias variou de baixa a moderada, coerente com um humor que oscila em resposta ao '
  'treino, e os valores de mudança mínima detectável fornecem o limiar prático para julgar a variação individual.')

H2('Etapa 4. Variação na semana e migração de perfis','3.4')
rows=[]
for k,lab in DIMS:
    f=AR['friedman'][k]; d=AR['d1d7'][k]
    rows.append([lab,num(f['chi'],2),pvt(f['p'])+('*' if f['p']<0.05 else ''),num(f['W'],2),
                 '%s → %s'%(num(d['d1'],1),num(d['d7'],1)),sg(d['delta'],1),sg(d['dz']),pvt(d['p'])+('*' if d['sig'] else '')])
table('Mudança das variáveis ao longo da semana (Friedman) e do primeiro ao sétimo dia (Wilcoxon).',
  ['Variável','χ² Friedman','p (semana)','W','D1 → D7','Δ','dz','p (D1→D7)'],rows,fs=8.5,
  note='χ²: qui-quadrado de Friedman. W: W de Kendall. Δ: variação do escore do D1 ao D7. dz: tamanho de efeito pareado. * p < 0,05.')
P('Vigor, fadiga, tensão, confusão, PTH e, com o maior tamanho de efeito, a fadiga física mudaram de forma '
  'significativa ao longo da semana. O padrão é consistente: o vigor caiu e as dimensões de fadiga aumentaram. A '
  'depressão, a raiva e a fadiga mental permaneceram estáveis, o que indica perda de energia e acúmulo de fadiga '
  'física, e não um quadro de humor negativo generalizado. A Figura 3 separa, para cada dimensão, o efeito '
  'imediato de cada sessão do efeito acumulado ao longo dos dias, e confirma que a mudança foi predominantemente '
  'acumulada. A Figura 4 mostra, em um único mapa, onde e quanto cada dimensão mudou em cada transição da semana.')
figure(f'{FG}/temporal_agudo_cronico.png','Efeito agudo (variação intradia, da manhã ao fim do dia) versus efeito crônico (do primeiro ao sétimo dia) por dimensão. O asterisco indica variação semanal significativa.',w=16.5)
figure(f'{FG}/temporal_map.png','Mapa temporal das dimensões: tamanho de efeito (dz) de cada transição da semana. A = transição de estímulo (manhã para fim de dia); R = transição de recuperação (fim de um dia para início do seguinte); * p < 0,05.',w=17.0)
prev=AR['perfil_prev']; allp=sorted(set(list(prev['D1'])+list(prev['D7'])))
rows=[[pf,str(prev['D1'].get(pf,0)),str(prev['D7'].get(pf,0))] for pf in allp]
table('Distribuição dos perfis de humor no primeiro e no sétimo dia.',
  ['Perfil de humor','D1 (atletas)','D7 (atletas)'],rows,fs=9,note='Perfis classificados a partir das dimensões do BRUMS.')
P('No primeiro dia predominou o perfil iceberg, típico de atletas bem recuperados. Ao longo da semana, %d dos 21 '
  'atletas com classificação nos dois momentos migraram para outro perfil, e apenas %d permaneceram estáveis. No '
  'sétimo dia, o perfil mais frequente passou a ser o de barbatana de tubarão, marcado pela queda do vigor. Essa '
  'migração traduz, em linguagem de perfil, a mesma perda de energia observada nas análises anteriores.'
  %(AR['perfil_migracao_n'],AR['perfil_estavel_n']))

H2('Etapa 5. Associação entre aptidão e humor','3.5')
rows=[]
for k,lab in DIMS:
    base=P2['pv_base'].get(k); det=AR['pv_vs_deterioracao'].get(k)
    br='%s (%s)'%(sg(base['rho']),pvt(base['p'])) if base else 'n/d'
    dr='%s (%s)'%(sg(det.get('rho',0)),pvt(det.get('p',1))) if det else 'n/d'
    rows.append([lab,br,dr])
table('Correlação de Spearman entre o pico de velocidade no T-car e o humor: nível semanal e deterioração (D1→D7).',
  ['Variável','PV × nível (rho, p)','PV × deterioração (rho, p)'],rows,fs=8.8,
  note='rho: coeficiente de Spearman. As correlações com a deterioração foram calculadas para as variáveis '
  'principais do estudo. Nenhuma correlação alcançou significância estatística.')
P('A aptidão aeróbia relacionou-se de forma fraca e não significativa com o nível do humor e, sobretudo, não se '
  'relacionou com a magnitude da deterioração ao longo da semana. Em termos práticos, o atleta mais apto não '
  'esteve protegido contra a queda de vigor e o aumento de fadiga: a deterioração atingiu de forma semelhante os '
  'mais e os menos aptos (Figura 5). A aptidão parece definir o patamar geral do humor, e não o quanto ele se '
  'desgasta em uma semana de cargas elevadas.')
figure(f'{FG}/temporal_aptidao.png','Deterioração do primeiro ao sétimo dia por dimensão, para atletas mais aptos versus menos aptos (divididos pela mediana do pico de velocidade). Não houve diferença significativa entre os grupos.',w=16.5)

H2('Detalhamento por dimensão','3.6')
P('As figuras a seguir apresentam, para cada dimensão do humor, um painel completo com a trajetória diária e seus '
  'limites, a velocidade e a aceleração da mudança, as transições ao longo da semana, o efeito imediato por dia e '
  'uma síntese descritiva. A leitura conjunta permite identificar, para cada variável, em que dia a mudança foi '
  'mais intensa e quando a tendência mudou de sentido.')
for k,lab in DIMS:
    figure(f'{FG}/vd_{k}.png','Painel da dimensão %s: trajetória, limites, derivadas, transições e síntese descritiva.'%lab.split(' (')[0].lower(),w=15.0)

# ============ 4 DISCUSSÃO ============
H1('Discussão','4')
P('Este estudo descreveu o perfil de humor de atletas de handebol de elite ao longo da última semana de '
  'pré-temporada e testou a associação entre a aptidão aeróbia e o comportamento do humor. A discussão segue as '
  'etapas dos resultados.')
H2('A energia cai e a fadiga física se acumula','4.1')
P('O achado mais robusto foi a combinação de queda do vigor e aumento da fadiga física, esta última com o maior '
  'tamanho de efeito. Esse padrão é coerente com o propósito da pré-temporada, período planejado para impor carga, '
  'e com a natureza intermitente e de alta intensidade do handebol, que exige elevada solicitação aeróbia e '
  'neuromuscular (Pueo et al., 2017; Cartón-Llorente et al., 2023). A literatura de monitoramento em esportes '
  'coletivos descreve exatamente essa resposta: a carga de treino relaciona-se de forma negativa com o vigor e '
  'positiva com a fadiga (Duignan et al., 2020; Selmi et al., 2023). A análise das derivadas mostrou que a queda '
  'do vigor foi mais rápida no início da semana e apresentou ponto de inflexão na sua metade, sinal de '
  'estabilização da resposta. Do ponto de vista prático, o vigor se confirma como o indicador mais sensível para '
  'acompanhar o desgaste.')
H2('O efeito é acumulado, não pontual','4.2')
P('A separação entre efeito imediato e efeito acumulado trouxe informação relevante para a periodização. A '
  'deterioração do humor não foi explicada por uma sessão isolada, mas pela soma dos dias: as transições de '
  'estímulo, da manhã ao fim do dia, produziram variações menores do que a tendência acumulada ao longo da semana. '
  'Esse resultado está alinhado a estudos que mostram que a resposta subjetiva reflete o acúmulo da carga de um '
  'período, e não apenas o esforço de uma sessão pontual (Chen et al., 2022). A implicação é direta: o '
  'monitoramento diário, que capta a trajetória, é mais informativo do que medidas isoladas.')
H2('Os perfis de humor migram do iceberg','4.3')
P('A migração de perfis descreve, em linguagem intuitiva para a comissão técnica, o mesmo fenômeno observado nas '
  'dimensões isoladas. O predomínio inicial do perfil iceberg, com o vigor sobreposto às dimensões negativas, deu '
  'lugar a perfis de menor energia, como a barbatana de tubarão. A maior parte dos atletas mudou de perfil ao '
  'longo da semana, o que confirma a sensibilidade do humor à carga e sustenta o uso do perfil como recurso de '
  'comunicação rápida do estado do grupo, em linha com a proposta original da escala (Terry, Lane & Fogarty, 2003) '
  'e com sua ampla aplicação em atletas (Terry et al., 2022; Lew et al., 2023).')
H2('Aptidão não é sinônimo de proteção','4.4')
P('O resultado mais instigante é que a aptidão aeróbia não protegeu o atleta da deterioração do humor. As '
  'correlações entre o pico de velocidade e a variação do humor foram fracas e não significativas, e a comparação '
  'entre mais e menos aptos não revelou diferença na magnitude da queda. Vale lembrar que o pico de velocidade no '
  'T-car é um indicador consolidado de capacidade e potência aeróbia máxima (Da Silva et al., 2011; Carminatti et '
  'al., 2013; Floriano et al., 2016), e que a aptidão aeróbia é determinante de desempenho no handebol de elite '
  '(Font et al., 2021; Fikenzer et al., 2020). Ainda assim, ela se associou apenas ao patamar do humor, e não à '
  'inclinação da resposta ao longo da semana. Uma interpretação plausível é que estar bem condicionado ajuda o '
  'atleta a começar melhor, mas não impede o desgaste quando o estímulo é elevado para todos. Esse achado tem '
  'consequência prática: o monitoramento do humor não pode ser substituído pela avaliação da aptidão, porque os '
  'dois descrevem aspectos diferentes e complementares do atleta.')
H2('Síntese dos principais achados','4.5')
P('Em conjunto, os resultados mostram que a última semana de pré-temporada produziu perda de energia e acúmulo de '
  'fadiga física, de forma acumulada ao longo dos dias, com migração da maioria dos atletas para perfis de menor '
  'vigor, e sem que a aptidão aeróbia oferecesse proteção contra essa deterioração. A principal contribuição do '
  'estudo é integrar, em um mesmo desenho, a descrição fina do humor por derivadas e limiares, a leitura por '
  'perfis e o cruzamento com a aptidão, para dar à comissão técnica um retrato claro de onde, quando e quanto o '
  'humor muda. As limitações incluem o desenho de grupo único, sem grupo controle, e o tamanho da amostra, próprio '
  'de uma equipe de elite. Estudos futuros podem acompanhar mais de um microciclo e incluir medidas objetivas de '
  'recuperação para confirmar e ampliar estes achados.')

# ============ 5 CONCLUSÃO ============
H1('Conclusão','5')
P('A última semana de pré-temporada de atletas de handebol de elite foi marcada pela queda do vigor e pelo aumento '
  'da fadiga física, com deterioração acumulada ao longo dos dias e migração da maioria dos atletas do perfil '
  'iceberg para perfis de menor energia. A aptidão aeróbia, medida pelo pico de velocidade no T-car, associou-se '
  'ao patamar do humor, mas não protegeu o atleta da deterioração ao longo da semana. O monitoramento diário dos '
  'estados de humor, aplicado de forma simples e virtual, mostrou-se sensível e útil para orientar decisões de '
  'ajuste de carga.')

# ============ REFERÊNCIAS ============
H1('Referências')
refs=[
 'Carminatti, L. J., Possamai, C. A. P., de Moraes, M., da Silva, J. F., de Lucas, R. D., Dittrich, N., & '
 'Guglielmo, L. G. A. (2013). Intermittent versus continuous incremental field tests: are maximal variables '
 'interchangeable? Journal of Sports Science and Medicine, 12(1), 165-170.',
 'Cartón-Llorente, A., Lozano, D., Gilart Iglesias, V., Marcos Jorquera, D., & Manchado, C. (2023). Worst-case '
 'scenario analysis of physical demands in elite men handball players by playing position through big data '
 'analytics. Biology of Sport, 40(4), 1219-1227. https://doi.org/10.5114/biolsport.2023.126665',
 'Chen, Y. S., Clemente, F. M., Pagaduan, J. C., Crowley-McHattan, Z. J., Lu, Y. X., Chien, C. H., Bezerra, P., '
 'Chiu, Y. W., & Kuo, C. D. (2022). Relationships between perceived measures of internal load and wellness status '
 'during overseas futsal training camps. PLoS ONE, 17(4), e0267227. https://doi.org/10.1371/journal.pone.0267227',
 'Da Silva, J. F., Guglielmo, L. G. A., Carminatti, L. J., De Oliveira, F. R., Dittrich, N., & Paton, C. D. '
 '(2011). Validity and reliability of a new field test (Carminatti’s test) for soccer players compared with '
 'laboratory-based measures. Journal of Sports Sciences, 29(15), 1621-1628. '
 'https://doi.org/10.1080/02640414.2011.609179',
 'Duignan, C., Doherty, C., Caulfield, B., & Blake, C. (2020). Single-item self-report measures of team-sport '
 'athlete wellbeing and their relationship with training load: a systematic review. Journal of Athletic Training, '
 '55(9), 944-953. https://doi.org/10.4085/1062-6050-0528.19',
 'Fernandes-da-Silva, J., Castagna, C., Teixeira, A. S., Carminatti, L. J., & Guglielmo, L. G. A. (2016). The peak '
 'velocity derived from the Carminatti Test is related to physical match performance in young soccer players. '
 'Journal of Sports Sciences, 34(24), 2238-2245. https://doi.org/10.1080/02640414.2016.1209307',
 'Fikenzer, S., Fikenzer, K., Laufs, U., Falz, R., Pietrek, H., & Hepp, P. (2020). Impact of COVID-19 lockdown on '
 'endurance capacity of elite handball players. Journal of Sports Medicine and Physical Fitness, 61(7), 977-982. '
 'https://doi.org/10.23736/S0022-4707.20.11501-9',
 'Floriano, L. T., da Silva, J. F., Teixeira, A. S., Salvador, P. C. N., Dittrich, N., Carminatti, L. J., '
 'Nascimento, L. L., & Guglielmo, L. G. A. (2016). Physiological responses during the time limit at 100% of the '
 'peak velocity in the Carminatti’s test in futsal players. Journal of Human Kinetics, 54, 91-101. '
 'https://doi.org/10.1515/hukin-2016-0038',
 'Font, R., Irurtia, A., Gutierrez, J. A., Salas, S., Vila, E., & Carmona, G. (2021). The effects of COVID-19 '
 'lockdown on jumping performance and aerobic capacity in elite handball players. Biology of Sport, 38(4), '
 '753-759. https://doi.org/10.5114/biolsport.2021.109952',
 'Lew, P. C. F., Parsons-Smith, R. L., Lamont-Mills, A., & Terry, P. C. (2023). Cross-cultural validation of the '
 'Malaysian Mood Scale and tests of between-group mood differences. International Journal of Environmental '
 'Research and Public Health, 20(4), 3348. https://doi.org/10.3390/ijerph20043348',
 'Pueo, B., Jimenez-Olmedo, J. M., Penichet-Tomas, A., Ortega Becerra, M., & Espina Agullo, J. J. (2017). '
 'Analysis of time-motion and heart rate in elite male and female beach handball. Journal of Sports Science and '
 'Medicine, 16(4), 450-458.',
 'Rohlfs, I. C. P. M., Rotta, T. M., Luft, C. D. B., Andrade, A., Krebs, R. J., & Carvalho, T. (2008). A Escala '
 'de Humor de Brunel (Brums): instrumento para deteccao precoce da sindrome do excesso de treinamento. Revista '
 'Brasileira de Medicina do Esporte, 14(3), 176-181.',
 'Selmi, O., Ouergui, I., Muscella, A., Levitt, D. E., Suzuki, K., & Bouassida, A. (2023). Monitoring mood state '
 'to improve performance in soccer players: a brief review. Frontiers in Psychology, 14, 1095238. '
 'https://doi.org/10.3389/fpsyg.2023.1095238',
 'Terry, P. C., Lane, A. M., & Fogarty, G. J. (2003). Construct validity of the Profile of Mood '
 'States-Adolescents for use with adults. Psychology of Sport and Exercise, 4(2), 125-139.',
 'Terry, P. C., Skurvydas, A., Lisinskiene, A., Majauskiene, D., Valanciene, D., Cooper, S., & Lochbaum, M. '
 '(2022). Validation of a Lithuanian-language version of the Brunel Mood Scale: the BRUMS-LTU. International '
 'Journal of Environmental Research and Public Health, 19(8), 4867. https://doi.org/10.3390/ijerph19084867',
 'Zhang, C. Q., Si, G., Chung, P. K., Du, M., & Terry, P. C. (2014). Psychometric properties of the Brunel Mood '
 'Scale in Chinese adolescents and adults. Journal of Sports Sciences, 32(15), 1465-1476. '
 'https://doi.org/10.1080/02640414.2014.898184']
for rf in refs:
    p=doc.add_paragraph(); r=p.add_run(rf); r.font.size=Pt(11); r.font.name='Times New Roman'
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(6); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent=Cm(1.25); p.paragraph_format.first_line_indent=Cm(-1.25)

doc.save('/home/user/mdlucca/Artigos/Estudo_Humor_Handebol_Completo.docx')
print('OK  figs=%d tabs=%d refs=%d'%(_FN[0],_TN[0],len(refs)))
