# -*- coding: utf-8 -*-
import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
SRC='/root/.claude/uploads/e1dba24c-b1d7-5908-9106-f2f4aaf3f56a/557fba4a-PERFIL_HUMOR_COMPLETO.docx'
OUT='/home/user/mdlucca/Artigos/PERFIL_HUMOR_COMPLETO.docx'
FG='/home/user/mdlucca/Artigos/figuras'
N=json.load(open('normalizacao.json')); FULL=json.load(open('tcar1_full.json')); MOD=json.load(open('tcar1_model.json'))
def c2(s): return str(s).replace('.',',')
d=Document(SRC)

def _fmt(p,size=12,just=True,bold=False,center=False):
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(6)
    if center: p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    elif just: p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    for r in p.runs: r.font.name='Times New Roman'; r.font.size=Pt(size); r.bold=bold
def after(el,text='',**kw):
    p=d.add_paragraph(text); el.addnext(p._p); _fmt(p,**kw); return p._p
def head(el,text): return after(el,text,bold=True,just=False)
def _bd(cell):
    tcPr=cell._tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for edge in ['top','bottom','left','right']:
        e=OxmlElement(f'w:{edge}'); on=edge in ('top','bottom')
        e.set(qn('w:val'),'single' if on else 'nil'); e.set(qn('w:sz'),'6'); e.set(qn('w:color'),'000000'); b.append(e)
    tcPr.append(b)
def table_after(el,caption,header,rows,fonte='Fonte: dados da pesquisa (2026).',fs=9.5):
    cap=d.add_paragraph(caption); el.addnext(cap._p); _fmt(cap,size=11,just=False)
    t=d.add_table(rows=1,cols=len(header)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; cap._p.addnext(t._tbl)
    for i,htx in enumerate(header):
        cc=t.rows[0].cells[i]; cc.text=''; rr=cc.paragraphs[0].add_run(htx); rr.bold=True; rr.font.size=Pt(fs); rr.font.name='Times New Roman'
        cc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; _bd(cc)
    for row in rows:
        cs=t.add_row().cells
        for i,val in enumerate(row):
            cs[i].text=''; rr=cs[i].paragraphs[0].add_run(str(val)); rr.font.size=Pt(fs); rr.font.name='Times New Roman'
            cs[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if i>0 else WD_ALIGN_PARAGRAPH.LEFT; _bd(cs[i])
    fon=d.add_paragraph(fonte); t._tbl.addnext(fon._p); _fmt(fon,size=10,just=False); return fon._p
def fig_after(el,path,caption,w=16.0):
    pimg=d.add_paragraph(); el.addnext(pimg._p); pimg.alignment=WD_ALIGN_PARAGRAPH.CENTER; pimg.add_run().add_picture(path,width=Cm(w))
    cap=d.add_paragraph(caption); pimg._p.addnext(cap._p); _fmt(cap,size=11,just=False,center=True)
    fon=d.add_paragraph('Fonte: elaboração dos autores (2026).'); cap._p.addnext(fon._p); _fmt(fon,size=10,just=False,center=True); return fon._p
def setpar(p,text):  # replace paragraph text preserving style of first run
    if not p.runs: p.add_run(text); return
    p.runs[0].text=text
    for r in p.runs[1:]: r.text=''
def find(prefix):
    for p in d.paragraphs:
        if p.text.strip().startswith(prefix): return p
    raise KeyError(prefix)

# ============ 0) GLOBAL RELABEL T-CAR 2 -> T-CAR 1 ; reavaliação -> avaliação inicial ============
for p in d.paragraphs:
    for r in p.runs:
        if r.text:
            r.text=r.text.replace('T-CAR 2','T-CAR 1').replace('reavaliação','avaliação inicial').replace('reavaliar','avaliar')
for t in d.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    if r.text: r.text=r.text.replace('T-CAR 2','T-CAR 1')

# ============ 1) FIX truncated (7) in 4.5 ============
p45=find('(1) Estatística descritiva exploratória')
run=p45.runs[-1]; run.text=run.text.rstrip()
if run.text.endswith('(7)'): run.text=run.text[:-3].rstrip()
run.text=run.text+(' (7) Padronização (escore z) das subescalas para a classificação de perfis e ajuste/normalização '
 'alométrica dos parâmetros do T-CAR 1 pela massa corporal (ver 4.5.1). (8) Suavização das médias diárias por regressão '
 'local ponderada (LOWESS) para separar sinal de ruído. (9) Modelagem preditiva de alta fadiga e de alto vigor pelos '
 'parâmetros fisiológicos, com validação cruzada (leave-one-out) e curvas ROC aparente e validada.')

# ============ 2) 4.5.1 normalização (métodos) ============
a=p45._p
a=head(a,'4.5.1 Normalização, ajuste alométrico e suavização — o que são e por que aplicá-las')
a=after(a,'Três operações preparam os dados antes das inferências, cada uma com uma razão metodológica específica e respaldo na literatura.')
a=after(a,'(a) Padronização (escore z) das subescalas. A classificação nos seis perfis de Terry baseia-se na menor distância euclidiana de cada observação aos protótipos. Como as subescalas têm variâncias muito diferentes — a fadiga varia cerca de três vezes mais do que a confusão (Figura 19A) —, usar os escores brutos faria as subescalas de maior amplitude dominarem a distância e distorcerem o perfil atribuído. A padronização (média 0, desvio 1) coloca todas as dimensões em pé de igualdade, condição necessária para uma classificação de perfil válida e comparável entre atletas (PARSONS-SMITH; TERRY; MACHIN, 2017).')
a=after(a,'(b) Normalização alométrica do parâmetro fisiológico. Medidas fisiológicas dependem do tamanho corporal; compará-las entre atletas de massas diferentes por razão simples (por quilograma) introduz viés sistemático, pois a relação entre a variável e a massa raramente é de proporção direta (NEVILL; RAMSBOTTOM; WILLIAMS, 1992; JARIC, 2002). O ajuste alométrico Y = a·massaᵇ estima o expoente b que torna a medida independente do tamanho, de modo que a medida normalizada (Y/massaᵇ) deixa de se correlacionar com a massa. No futebol e em outras modalidades coletivas, esse procedimento é padrão para comparar a aptidão entre posições e indivíduos sem o confundimento do porte físico (CHAMARI et al., 2005; OBA et al., 2014).')
a=after(a,'(c) Suavização das curvas (LOWESS). O sinal do microciclo é de baixa frequência e pequeno em relação ao ruído de medida do dia a dia. A regressão local ponderada (LOWESS) estima a tendência sistemática sem impor forma paramétrica, separando o sinal (a deriva semanal de vigor e fadiga) do ruído (a oscilação aleatória em torno do piso das subescalas negativas) (CLEVELAND, 1979). Essa filtragem sustenta afirmar que a migração do perfil iceberg para o perfil de fadiga é tendência real e não artefato de medida — em linha com a recomendação de monitorar o atleta por tendência referenciada à sua linha de base individual (SAW; MAIN; GASTIN, 2016).')

# ============ 3) REWRITE 5.10 (T-CAR 1) ============
p117=find('Testou-se se os parâmetros do T-CAR 1')
setpar(p117,'Testou-se se os parâmetros do T-CAR 1 (normalizados e ajustados alometricamente) predizem os dias de fadiga alta na semana (Tabela 12; Figuras 17–18). Diferentemente de uma avaliação pós-bloco, a aptidão de entrada preserva poder discriminativo acima do acaso mesmo na validação cruzada: o pico de velocidade e a distância têm AUC aparente de %s e AUC validada (LOOCV) de %s–%s, com limiar ≈ %s km·h⁻¹ e razão de chances menor que 1 por desvio-padrão (p < 0,001). A associação de grupo é, portanto, significativa, embora a discriminação individual fora da amostra seja apenas modesta — a fadiga da semana tem forte componente perceptual e individual, mas é parcialmente antecipada pela aptidão aeróbia basal. O detalhamento da regressão e do limiar, para a fadiga e também para o vigor, consta na seção 5.12.'%(
    c2('%.2f'%MOD['PV']['auc']),c2('%.2f'%MOD['PV']['cv']),c2('%.2f'%MOD['dist']['cv']),c2('%.1f'%MOD['PV']['thr'])))
# rebuild Table 12 (table with header 'Parâmetro (T-CAR 1)')
tab12=None
for t in d.tables:
    if t.rows[0].cells[0].text.strip().startswith('Parâmetro'): tab12=t; break
def setcell(cell,txt,bold=False):
    cell.text=''; r=cell.paragraphs[0].add_run(txt); r.font.name='Times New Roman'; r.font.size=Pt(9.5); r.bold=bold
    cell.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
rows12=[['PV (T-CAR 1)',c2('%.2f'%MOD['PV']['auc']),c2('%.2f'%MOD['PV']['cv']),c2('%.1f'%MOD['PV']['thr']),c2('%.2f'%MOD['PV']['OR']),'<0,001'],
        ['Distância (T-CAR 1)',c2('%.2f'%MOD['dist']['auc']),c2('%.2f'%MOD['dist']['cv']),c2('%.0f'%MOD['dist']['thr']),c2('%.2f'%MOD['dist']['OR']),'<0,001']]
# header cells
for j,htx in enumerate(['Parâmetro','AUC aparente','AUC validada (LOOCV)','Limiar (km·h⁻¹)','OR (por +1 DP)','p']):
    setcell(tab12.rows[0].cells[j],htx,bold=True)
# fill first 2 data rows, delete the rest
data_rows=tab12.rows[1:]
for i,rvals in enumerate(rows12):
    for j,val in enumerate(rvals): setcell(data_rows[i].cells[j],val)
for extra in data_rows[len(rows12):]:
    extra._tr.getparent().remove(extra._tr)
# update Table 12 caption + Fig17/18 captions
cap12=find('Tabela 12'); setpar(cap12,'Tabela 12 – Modelagem dos dias de fadiga alta pelos parâmetros de aptidão do T-CAR 1: ROC aparente e validada (LOOCV), limiar e razão de chances.')
cap18=find('Figura 18'); setpar(cap18,'Figura 18 – AUC aparente vs. validada (LOOCV): a aptidão de entrada (PV e distância) mantém discriminação modesta acima do acaso, diferentemente de parâmetros sem sinal.')

# swap image blobs for Fig17, Fig18 (last two inline images)
imgs=d.inline_shapes
def swap(shp,path,w_cm,h_cm):
    blip=shp._inline.graphic.graphicData.pic.blipFill.blip; rId=blip.get(qn('r:embed'))
    d.part.related_parts[rId]._blob=open(path,'rb').read()
    shp.width=Cm(w_cm); shp.height=Cm(h_cm)
swap(imgs[len(imgs)-2],f'{FG}/x_fig17.png',13.0,8.2)
swap(imgs[len(imgs)-1],f'{FG}/x_fig18.png',14.0,7.75)

# ============ 4) 5.11 normalização (resultados) + Figura 19 ============
p6=find('6 SÍNTESE DOS PRINCIPAIS ACHADOS'); anchor=p6._p.getprevious()
anchor=head(anchor,'5.11 Normalização, ajuste alométrico e suavização das curvas')
anchor=after(anchor,'A Figura 19 consolida as três operações e suas consequências práticas. (A) As subescalas do BRUMS têm desvios-padrão muito desiguais (fadiga = %s; confusão = %s), o que justifica a padronização antes da classificação de perfis: sem ela, fadiga e vigor dominariam a distância aos protótipos.'%(
    c2('%.1f'%N['sds']['Fadiga']),c2('%.1f'%N['sds']['Confusão'])))
anchor=after(anchor,'(B) A normalização alométrica remove o confundimento do tamanho: a correlação da fadiga física com a massa cai de |ρ| = %s para %s, e a correlação com o PV do T-CAR 1 — depois de descontado o efeito do tamanho — de %s para %s, isolando a contribuição da aptidão. (C) A suavização por LOWESS extrai, da nuvem de observações, a tendência de queda do vigor e de subida da fadiga ao longo da semana — o eixo energia–fadiga é sinal de baixa frequência, não ruído.'%(
    c2('%.2f'%N['mass_raw']),c2('%.2f'%N['mass_norm']),c2('%.2f'%N['pv_raw']),c2('%.2f'%N['pv_norm'])))
anchor=after(anchor,'(D) A tendência suavizada do índice iceberg cai de forma sistemática, cruza o zero por volta do Dia 3 e despenca no Dia 7 — a leitura contínua da mesma migração iceberg→fadiga descrita pelas frequências de perfil (§5.2). Em conjunto, a padronização torna a classificação de perfil válida, a normalização alométrica isola a aptidão do tamanho corporal, e a suavização confirma que a deterioração do humor ao longo do microciclo é sinal e não ruído.')
anchor=fig_after(anchor,f'{FG}/x_normalizacao.png','Figura 19 – Normalização e suavização. (A) Desvios-padrão desiguais das subescalas justificam a padronização (escore z). (B) A normalização alométrica remove o confundimento do tamanho (massa: |ρ| 0,33→0,06; PV do T-CAR 1: 0,54→0,16). (C) A suavização (LOWESS) separa o sinal semanal (vigor ↓, fadiga ↑) do ruído. (D) A tendência suavizada do índice iceberg cai ao longo da semana.')

# ============ 5) 5.12 limiar + regressão (fadiga E vigor) ============
Wf=FULL['FadFisica']['week']; Wv=FULL['Vigor']['week']; Tf=FULL['FadFisica']['thr']; Tv=FULL['Vigor']['thr']; TE=FULL['TERC']
vday=FULL['vigor_by_day']; dmax=max(vday,key=lambda k:vday[k])
anchor=head(anchor,'5.12 Ajuste do limiar do pico de velocidade e regressão aptidão–fadiga–vigor (semana e por dia)')
anchor=after(anchor,'Usando exclusivamente o T-CAR 1 (aptidão de entrada), ajustou-se a relação entre o pico de velocidade (PV) e as duas faces do eixo energia–fadiga — a fadiga física e o vigor —, no acumulado da semana e dia a dia, com estimativa de limiares de PV por regressão logística.')
anchor=after(anchor,'No acumulado semanal, a fadiga física decresce com o PV (fadiga = %s %s·PV; ρ = %s; p = %s; R² = %s) e o vigor cresce com o PV (vigor = %s %s·PV; ρ = %s; p = %s; R² = %s). Ou seja, os atletas de maior potência aeróbica são, ao mesmo tempo, os menos fadigados e os mais vigorosos ao longo da semana (Figura 20A; Tabela 13). A relação é consistente dia a dia: a correlação PV–fadiga é negativa nos sete dias (significativa nos dias 2, 3 e 5) e a correlação PV–vigor é positiva nos sete dias (significativa nos dias 2, 5 e 6).'%(
    c2('%.1f'%Wf['intercept']),c2('%+.2f'%Wf['slope']),c2('%+.2f'%Wf['rho']),c2('%.3f'%Wf['p']),c2('%.2f'%Wf['r2']),
    c2('%.1f'%Wv['intercept']),c2('%+.2f'%Wv['slope']),c2('%+.2f'%Wv['rho']),c2('%.3f'%Wv['p']),c2('%.2f'%Wv['r2'])))
anchor=after(anchor,'O ajuste logístico define limiares de PV: em torno de %s km·h⁻¹ para o dia de fadiga alta (AUC = %s; validação cruzada %s; sensibilidade %s/especificidade %s; OR = %s por km·h⁻¹) e em torno de %s km·h⁻¹ para o dia de vigor alto (AUC = %s; validação cruzada %s; sensibilidade %s/especificidade %s; OR = %s por km·h⁻¹) (Figura 20C; Tabela 14). Estratificando por tercis de aptidão, o terço mais apto (PV %s) mantém vigor semanal de %s contra %s do terço menos apto (PV %s), e fadiga de %s contra %s — gradientes amplos em ambas as direções.'%(
    c2('%.1f'%Tf['thr']),c2('%.2f'%Tf['auc']),c2('%.2f'%Tf['cv_auc']),c2('%.2f'%Tf['sens']),c2('%.2f'%Tf['spec']),c2('%.2f'%Tf['OR']),
    c2('%.1f'%Tv['thr']),c2('%.2f'%Tv['auc']),c2('%.2f'%Tv['cv_auc']),c2('%.2f'%Tv['sens']),c2('%.2f'%Tv['spec']),c2('%.2f'%Tv['OR']),
    c2('%.1f'%TE['Alta']['PV']),c2('%.1f'%TE['Alta']['Vigor']),c2('%.1f'%TE['Baixa']['Vigor']),c2('%.1f'%TE['Baixa']['PV']),
    c2('%.1f'%TE['Alta']['FadFis']),c2('%.1f'%TE['Baixa']['FadFis'])))
anchor=after(anchor,'Relação do PV com os dias de maior vigor. O grupo é mais vigoroso no Dia %d (vigor médio = %s), quando todos chegam descansados; nesse dia, porém, a aptidão quase não diferencia os atletas (ρ PV–vigor = %s, n.s.), porque o vigor está próximo do teto para todos. À medida que a semana avança e o vigor médio cai, o acoplamento PV–vigor se fortalece (ρ = %s no Dia 2, %s no Dia 5, %s no Dia 6; Figura 20D). A leitura é clara: a aptidão aeróbia não cria o pico de vigor do dia fresco — ela o PRESERVA sob carga. Os atletas com PV acima de ~%s km·h⁻¹ concentram os dias de vigor alto e resistem à queda; os de PV baixo perdem vigor mais cedo. Assim, o limiar do PV funciona como um marcador de quem sustenta o vigor ao longo do microciclo, não de quem parte mais vigoroso.'%(
    int(dmax),c2('%.1f'%vday[str(dmax)]),c2('%+.2f'%FULL['Vigor']['day'][str(dmax)]['rho']),
    c2('%+.2f'%FULL['Vigor']['day']['2']['rho']),c2('%+.2f'%FULL['Vigor']['day']['5']['rho']),c2('%+.2f'%FULL['Vigor']['day']['6']['rho']),
    c2('%.1f'%Tv['thr'])))
anchor=after(anchor,'Interpretação à luz da literatura. Os resultados alinham-se ao entendimento de que a capacidade aeróbia intermitente sustenta a recuperação entre esforços de alta intensidade e, portanto, a resistência à fadiga e a manutenção do estado de energia nas modalidades coletivas (BANGSBO; IAIA; KRUSTRUP, 2008; MICHALSIK; MADSEN; AAGAARD, 2013). O PV do teste de Carminatti, por integrar potência aeróbia, economia de corrida e capacidade de repetir esforços, é um índice ecológico dessa aptidão (CHAMARI et al., 2005). Transformar esse índice contínuo em regra de decisão por regressão logística e ponto de corte (limiar de Youden) é procedimento consolidado no rastreamento esportivo (SAW; MAIN; GASTIN, 2016).')
# Tabela 13 (per-day regression, fatigue & vigor)
rows13=[['Semana (acumulado)',c2('%+.2f'%Wf['rho']),c2('%.3f'%Wf['p']),c2('%+.2f'%Wv['rho']),c2('%.3f'%Wv['p'])]]
for dd in range(1,8):
    fa=FULL['FadFisica']['day'][str(dd)]; vi=FULL['Vigor']['day'][str(dd)]
    rows13.append(['Dia %d'%dd,c2('%+.2f'%fa['rho']),c2('%.3f'%fa['rho_p']),c2('%+.2f'%vi['rho']),c2('%.3f'%vi['rho_p'])])
anchor=table_after(anchor,'Tabela 13 – Correlação do PV (T-CAR 1) com a fadiga física e o vigor, por dia e no acumulado semanal (ρ de Spearman).',
    ['Recorte','ρ (PV × fadiga)','p','ρ (PV × vigor)','p'],rows13)
# Tabela 14 (thresholds)
anchor=table_after(anchor,'Tabela 14 – Ajuste do limiar logístico do PV (T-CAR 1) para o dia de fadiga alta e o dia de vigor alto.',
    ['Parâmetro','Dia de fadiga alta','Dia de vigor alto'],
    [['Limiar de Youden (km·h⁻¹)',c2('%.1f'%Tf['thr']),c2('%.1f'%Tv['thr'])],
     ['AUC [IC95%]','%s [%s–%s]'%(c2('%.2f'%Tf['auc']),c2('%.2f'%Tf['lo']),c2('%.2f'%Tf['hi'])),'%s [%s–%s]'%(c2('%.2f'%Tv['auc']),c2('%.2f'%Tv['lo']),c2('%.2f'%Tv['hi']))],
     ['AUC validada (LOOCV)',c2('%.2f'%Tf['cv_auc']),c2('%.2f'%Tv['cv_auc'])],
     ['Sensibilidade / Especificidade','%s / %s'%(c2('%.2f'%Tf['sens']),c2('%.2f'%Tf['spec'])),'%s / %s'%(c2('%.2f'%Tv['sens']),c2('%.2f'%Tv['spec']))],
     ['OR por km·h⁻¹',c2('%.2f'%Tf['OR']),c2('%.2f'%Tv['OR'])]])
anchor=fig_after(anchor,f'{FG}/x_limiar.png','Figura 20 – Ajuste do limiar do PV (T-CAR 1) e regressão aptidão–vigor–fadiga. (A) Quanto maior o PV, maior o vigor semanal. (B) A correlação PV–vigor é positiva em todos os dias e PV–fadiga negativa em todos os dias (* p<0,05). (C) Limiares logísticos: vigor alto acima de ~15,0 km·h⁻¹, fadiga alta abaixo de ~14,9 km·h⁻¹. (D) O grupo é mais vigoroso no Dia 1 (fresco), quando o PV pouco diferencia; o acoplamento PV–vigor cresce à medida que o vigor cai — a aptidão preserva o vigor sob carga.')

# ============ 6) SÍNTESE bullet on physiological param -> update to T-CAR 1 finding ============
try:
    pb=find('•  Parâmetro fisiológico: os índices do T-CAR 1')
    setpar(pb,'•  Parâmetro fisiológico: a aptidão de entrada (PV do T-CAR 1) diferencia o eixo energia–fadiga — associa-se a menos fadiga (ρ = −0,54) e mais vigor (ρ = +0,48) na semana, com limiares úteis (~14,9–15,0 km·h⁻¹) e discriminação que resiste à validação cruzada (AUC ≈ 0,62); a vantagem do mais apto não está no dia fresco, mas na preservação do vigor sob carga acumulada.')
except KeyError: pass

# ============ 7) REFERÊNCIAS ============
brandt=find('BRANDT, R.')
newp=d.add_paragraph('BANGSBO, J.; IAIA, F. M.; KRUSTRUP, P. The Yo-Yo intermittent recovery test: a useful tool for evaluation of physical performance in intermittent sports. Sports Medicine, v. 38, n. 1, p. 37–51, 2008. DOI: 10.2165/00007256-200838010-00004.')
brandt._p.addprevious(newp._p); _fmt(newp,size=12,just=False)
a=brandt._p
a=after(a,'CHAMARI, K. et al. Endurance training and testing with the ball in young elite soccer players. British Journal of Sports Medicine, v. 39, n. 1, p. 24–28, 2005. DOI: 10.1136/bjsm.2003.009985.',just=False)
a=after(a,'CLEVELAND, W. S. Robust locally weighted regression and smoothing scatterplots. Journal of the American Statistical Association, v. 74, n. 368, p. 829–836, 1979. DOI: 10.1080/01621459.1979.10481038.',just=False)
a=after(a,'JARIC, S. Muscle strength testing: use of normalisation for body size. Sports Medicine, v. 32, n. 10, p. 615–631, 2002. DOI: 10.2165/00007256-200232100-00002.',just=False)
morgan=find('MORGAN, W. P.'); b=morgan._p
b=after(b,'NEVILL, A. M.; RAMSBOTTOM, R.; WILLIAMS, C. Scaling physiological measurements for individuals of different body size. European Journal of Applied Physiology and Occupational Physiology, v. 65, n. 2, p. 110–117, 1992. DOI: 10.1007/BF00705066.',just=False)
b=after(b,'OBA, Y. et al. Allometric scaling of strength scores in NCAA Division I-A football athletes. Journal of Strength and Conditioning Research, v. 28, n. 12, p. 3330–3337, 2014. DOI: 10.1519/JSC.0000000000000548.',just=False)

d.save(OUT)
print('SAVED',OUT)
