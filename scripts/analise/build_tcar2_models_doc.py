# -*- coding: utf-8 -*-
import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
R=json.load(open('tcar2_models.json'))
def c2(s): return str(s).replace('.',',')
doc=Document()
sty=doc.styles['Normal']; sty.font.name='Times New Roman'; sty.font.size=Pt(12)
sty.element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')
pf=sty.paragraph_format; pf.line_spacing=1.5; pf.space_after=Pt(0)
s=doc.sections[0]; s.top_margin=Cm(3); s.left_margin=Cm(3); s.bottom_margin=Cm(2); s.right_margin=Cm(2)
_TN=[0]; _FN=[0]
def P(t='',just=True,size=12,after=6,before=0):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(size)
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(after); p.paragraph_format.space_before=Pt(before)
    if just: p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.first_line_indent=Cm(1.25)
    return p
def H(t,size=12):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(size)
    p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(6); p.paragraph_format.line_spacing=1.5
def _bd(cell):
    tcPr=cell._tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for edge in ['top','bottom','left','right']:
        e=OxmlElement(f'w:{edge}'); on=edge in ('top','bottom')
        e.set(qn('w:val'),'single' if on else 'nil'); e.set(qn('w:sz'),'6'); e.set(qn('w:color'),'000000'); b.append(e)
    tcPr.append(b)
def table(caption,header,rows,fonte='Fonte: os autores (2026).',widths=None,fs=10):
    _TN[0]+=1
    p=doc.add_paragraph(); r=p.add_run(f'Tabela {_TN[0]} – {caption}'); r.font.size=Pt(11)
    p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(3)
    t=doc.add_table(rows=1,cols=len(header)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,htxt in enumerate(header):
        cc=t.rows[0].cells[i]; cc.text=''; rr=cc.paragraphs[0].add_run(htxt); rr.bold=True; rr.font.size=Pt(fs); rr.font.name='Times New Roman'
        cc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; _bd(cc)
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=''; rr=cells[i].paragraphs[0].add_run(str(val)); rr.font.size=Pt(fs); rr.font.name='Times New Roman'
            cells[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if i>0 else WD_ALIGN_PARAGRAPH.LEFT; _bd(cells[i])
    if widths:
        for i,w in enumerate(widths):
            for rr in t.rows: rr.cells[i].width=Cm(w)
    pf=doc.add_paragraph(); rf=pf.add_run(fonte); rf.font.size=Pt(10); pf.paragraph_format.space_after=Pt(8)
def figure(path,caption,w=16.0):
    _FN[0]+=1
    pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; pp.add_run().add_picture(path,width=Cm(w))
    pc=doc.add_paragraph(); pc.alignment=WD_ALIGN_PARAGRAPH.CENTER; rc=pc.add_run(f'Figura {_FN[0]} – {caption}'); rc.font.size=Pt(11)
    pf=doc.add_paragraph(); pf.alignment=WD_ALIGN_PARAGRAPH.CENTER; rf=pf.add_run('Fonte: os autores (2026).'); rf.font.size=Pt(10); pf.paragraph_format.space_after=Pt(8)

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('AJUSTES, REGRESSÕES, LIMIARES E DERIVADAS COM O PICO DE VELOCIDADE DO T-CAR 2 COMO PARÂMETRO FISIOLÓGICO — COMPARAÇÃO COM O T-CAR 1')
r.bold=True; r.font.size=Pt(13); p.paragraph_format.space_after=Pt(12)
P('Reexecução do arcabouço de modelagem (regressão linear, ajuste e normalização alométrica, limiar logístico/curva ROC e limites/derivadas por tercil) usando o pico de velocidade do T-CAR 2 (reteste) como parâmetro fisiológico, com o T-CAR 1 (teste inicial) como comparador. Desfechos: eixo energia–fadiga do BRUMS ao longo do microciclo de 7 dias (27 atletas).',size=11,after=10)

H('Resumo')
P('Em TODAS as famílias de modelo, o PV do T-CAR 2 (aptidão pós-bloco) mostrou-se um parâmetro fisiológico MAIS FRACO e menos utilizável que o PV do T-CAR 1: a regressão da fadiga física semanal cai de ρ = %s (T-CAR 1, p = %s) para ρ = %s (T-CAR 2, n.s.); a curva ROC para o dia de fadiga alta cai de AUC = %s (limiar equilibrado ≈ %s km·h⁻¹) para AUC = %s (limiar ≈ %s km·h⁻¹, com especificidade insuficiente); o ajuste alométrico torna-se instável (expoente sem significado por associação quase nula); e a estratificação das derivadas por tercil deixa de ser monotônica. A conclusão é robusta e temporalmente coerente: o parâmetro fisiológico correto para modelar a fadiga e o humor DESTE microciclo é o T-CAR 1 (aptidão de entrada), não o T-CAR 2 (medido após o bloco de treino).'%(
    c2('%+.2f'%R['REG']['FadFisica']['PVini']['rho']),c2('%.3f'%R['REG']['FadFisica']['PVini']['p']),
    c2('%+.2f'%R['REG']['FadFisica']['PVfin']['rho']),
    c2('%.2f'%R['LIM']['PVini']['auc']),c2('%.1f'%R['LIM']['PVini']['youden_thr']),
    c2('%.2f'%R['LIM']['PVfin']['auc']),c2('%.1f'%R['LIM']['PVfin']['youden_thr'])),after=8)

H('1. Regressões (desfecho semanal ~ PV)')
CORE=[('FadFisica','Fadiga física'),('Vigor','Vigor'),('FadMental','Fadiga mental'),('TMD','PTH/TMD'),('Fadiga','Fadiga (BRUMS)')]
def cell(d): return '%s / %s / %s / %s'%(c2('%+.2f'%d['slope']),c2('%.2f'%d['r2']),c2('%+.2f'%d['rho']),c2('%.3f'%d['p']))
rows=[[lab,cell(R['REG'][v]['PVfin']),cell(R['REG'][v]['PVini'])] for v,lab in CORE]
table('Regressão do desfecho semanal sobre o PV (β / R² / ρ / p) — T-CAR 2 vs T-CAR 1 (n = 27)',
    ['Desfecho','T-CAR 2 (PV final)','T-CAR 1 (PV inicial)'],rows,widths=[3.6,5.6,5.6],fs=9.5)
P('Para todos os desfechos, o T-CAR 1 produz coeficientes maiores e significativos (fadiga física ρ = %s, p = %s; vigor ρ = %s, p = %s), enquanto o T-CAR 2 fica fraco e não significativo (fadiga física ρ = %s; vigor ρ = %s). O sinal permanece na direção esperada (mais aptidão → menos fadiga, mais vigor), mas a magnitude cai à metade ou menos com o T-CAR 2.'%(
    c2('%+.2f'%R['REG']['FadFisica']['PVini']['rho']),c2('%.3f'%R['REG']['FadFisica']['PVini']['p']),
    c2('%+.2f'%R['REG']['Vigor']['PVini']['rho']),c2('%.3f'%R['REG']['Vigor']['PVini']['p']),
    c2('%+.2f'%R['REG']['FadFisica']['PVfin']['rho']),c2('%+.2f'%R['REG']['Vigor']['PVfin']['rho'])),after=6)

H('2. Ajuste e normalização alométrica')
al=R['ALLO']['FadFisica']
P('No ajuste alométrico Y = a·PVᵇ da fadiga física sobre o PV do T-CAR 2, o expoente resultou b = %s com R² = %s (p = %s) — estatisticamente instável e sem significado prático, porque o PV final praticamente não se correlaciona com a fadiga. A normalização (fadiga/PVᵇ) reduz uma correlação que já é quase nula (ρ %s → %s), sem informação útil. Em contraste, no T-CAR 1 a normalização pelo PV removeu de fato o confundimento da aptidão (ρ da fadiga física de −0,52 para −0,19, conforme o manuscrito principal). Ou seja, o descritor que faz o escalonamento alométrico funcionar é o PV de entrada (T-CAR 1), não o PV pós-bloco.'%(
    c2('%+.2f'%al['b']),c2('%.2f'%al['r2']),c2('%.3f'%al['p']),c2('%+.2f'%al['rho_before']),c2('%+.2f'%al['rho_after'])),after=6)

H('3. Limiar logístico e curva ROC (dia de fadiga alta)')
L2=R['LIM']['PVfin']; L1=R['LIM']['PVini']
table('Limiar/ROC para classificar o dia de fadiga física alta — T-CAR 2 vs T-CAR 1',
    ['Métrica','T-CAR 2 (PV final)','T-CAR 1 (PV inicial)'],
    [['AUC [IC95%]','%s [%s–%s]'%(c2('%.2f'%L2['auc']),c2('%.2f'%L2['auc_lo']),c2('%.2f'%L2['auc_hi'])),
                    '%s [%s–%s]'%(c2('%.2f'%L1['auc']),c2('%.2f'%L1['auc_lo']),c2('%.2f'%L1['auc_hi']))],
     ['Limiar de Youden (km·h⁻¹)',c2('%.1f'%L2['youden_thr']),c2('%.1f'%L1['youden_thr'])],
     ['Sensibilidade / Especificidade','%s / %s'%(c2('%.2f'%L2['sens']),c2('%.2f'%L2['spec'])),'%s / %s'%(c2('%.2f'%L1['sens']),c2('%.2f'%L1['spec']))],
     ['OR por km·h⁻¹',c2('%.2f'%L2['OR']),c2('%.2f'%L1['OR'])],
     ['Kruskal por tercil (fad. física)','p = %s (não monotônico)'%c2('%.3f'%R['TERC_KW']['FadFisica']),'— (referência)']],
    widths=[5.0,3.9,3.9],fs=9.5)
P('O limiar do T-CAR 1 (≈ %s km·h⁻¹) é equilibrado e utilizável (sensibilidade %s, especificidade %s), reproduzindo o ponto de corte já descrito no manuscrito. O do T-CAR 2 (≈ %s km·h⁻¹) tem especificidade baixíssima (%s), classificando quase todos como "de risco" — inútil na prática. Além disso, os tercis de PV do T-CAR 2 não ordenam a fadiga de forma monotônica (o tercil médio é o mais fatigado), enquanto no T-CAR 1 o tercil menos apto é o mais fatigado.'%(
    c2('%.1f'%L1['youden_thr']),c2('%.2f'%L1['sens']),c2('%.2f'%L1['spec']),
    c2('%.1f'%L2['youden_thr']),c2('%.2f'%L2['spec'])),after=6)

H('4. Limites e derivadas da trajetória (por tercil de PV)')
D=R['DERIV']
rows=[]
for terc in ['Baixo','Médio','Alto']:
    d=D.get(terc,{}).get('FadFisica')
    if d: rows.append([terc,c2('%.1f'%d['dose']),c2('%+.2f'%d['vel_mean']),('%s'%c2('%.1f'%d['inflex'])) if d['inflex'] else '—',c2('%.1f'%d['rng'])])
table('Derivadas e limites da trajetória de fadiga física por tercil de PV do T-CAR 2',
    ['Tercil PV (T-CAR 2)','Dose (área)','Velocidade média','Ponto de inflexão (dia)','Amplitude'],
    rows,widths=[4.2,2.8,3.2,3.4,2.4],fs=9.5)
P('Estratificando a trajetória pela aptidão do T-CAR 2, a dose acumulada de fadiga (área sob a curva), a velocidade de mudança, o ponto de inflexão (≈ dia 4, meio da semana) e a amplitude ficam praticamente indistinguíveis entre tercis — e não monotônicos (o tercil médio acumula a maior dose). Isto é, o PV do T-CAR 2 não organiza a dinâmica da fadiga. Com o PV do T-CAR 1, o tercil menos apto mantém a trajetória mais alta ao longo da semana (Figura 1D), coerente com o acoplamento aptidão–fadiga já demonstrado.',after=6)

figure('/home/user/mdlucca/Artigos/figuras/x_tcar2_models.png',
    'Comparação T-CAR 1 vs T-CAR 2 como parâmetro fisiológico. (A) A regressão semana × PV é mais forte com o T-CAR 1 em todos os desfechos. (B) A curva ROC para o dia de fadiga alta tem AUC e limiar utilizáveis apenas com o T-CAR 1. (C) Os tercis de PV do T-CAR 2 não separam a trajetória de fadiga. (D) Os tercis de PV do T-CAR 1 separam (o tercil menos apto, em vermelho, permanece mais fatigado).')

H('5. Síntese e recomendação')
P('Reexecutados os ajustes, regressões, limiares e derivadas com o PV do T-CAR 2 como parâmetro fisiológico, o resultado é consistente entre todas as abordagens: o T-CAR 2 enfraquece a regressão (ρ da fadiga física de %s para %s), inviabiliza o ajuste alométrico (expoente instável), degrada o limiar/ROC (AUC de %s para %s e especificidade insuficiente) e não estratifica as derivadas. Recomenda-se, portanto, manter o PV do T-CAR 1 como o parâmetro fisiológico dos modelos do microciclo; o T-CAR 2 é apropriado como desfecho de adaptação (ganho de aptidão pós-bloco), não como preditor da fadiga e do humor desta semana. A explicação é temporal: a semana analisada é contemporânea à aptidão de entrada (T-CAR 1); o T-CAR 2 reflete um estado fisiológico posterior, medido após o bloco que também gerou o ganho de PV.'%(
    c2('%+.2f'%R['REG']['FadFisica']['PVini']['rho']),c2('%+.2f'%R['REG']['FadFisica']['PVfin']['rho']),
    c2('%.2f'%R['LIM']['PVini']['auc']),c2('%.2f'%R['LIM']['PVfin']['auc'])),after=6)

H('6. Limitações')
P('n = 27 (26 com PV final); o desfecho de fadiga tem efeito de piso; o limiar/ROC usa observações agrupadas por atleta (IC95% por bootstrap clusterizado). O ajuste alométrico exige desfecho estritamente positivo (aplicado à fadiga física; PTH/TMD, por assumir valores ≤ 0, não é escalável). Associações correlacionais.',after=6)

OUTP='/home/user/mdlucca/Artigos/TCAR2_Ajustes_Regressoes_Limiares_Derivadas.docx'
doc.save(OUTP)
print('SAVED',OUTP,'| Tabelas',_TN[0],'Figuras',_FN[0])
