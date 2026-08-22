# -*- coding: utf-8 -*-
# Artigo separado: capacidade aerobia maxima (pico de velocidade no T-CAR) x perfil de humor.
import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
FG='/home/user/mdlucca/Artigos/figuras'
J=json.load(open('/home/user/mdlucca/scripts/analise/aerobio_humor.json'))
doc=Document()
for s in doc.sections:
    s.top_margin=s.bottom_margin=Cm(2.5); s.left_margin=s.right_margin=Cm(3)
st=doc.styles['Normal']; st.font.name='Times New Roman'; st.font.size=Pt(12)
_FN=[0]; _TN=[0]
def P(t='',ind=True,after=0,it=False,bold=False):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(12); r.font.name='Times New Roman'; r.italic=it; r.bold=bold
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(after); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if ind: p.paragraph_format.first_line_indent=Cm(1.25)
    return p
def H1(t,n=None):
    p=doc.add_paragraph(); r=p.add_run((n+' ' if n else '')+t); r.bold=True; r.font.size=Pt(12); r.font.name='Times New Roman'
    p.paragraph_format.space_before=Pt(14); p.paragraph_format.space_after=Pt(6)
def H2(t,n=None):
    p=doc.add_paragraph(); r=p.add_run((n+' ' if n else '')+t); r.bold=True; r.font.size=Pt(12); r.font.name='Times New Roman'
    p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(4)
def figure(fn,cap,w=15.0):
    _FN[0]+=1
    pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; pp.add_run().add_picture(f'{FG}/{fn}',width=Cm(w)); pp.paragraph_format.space_before=Pt(6)
    pc=doc.add_paragraph(); pc.alignment=WD_ALIGN_PARAGRAPH.CENTER; rc=pc.add_run('Figura %d – %s'%(_FN[0],cap)); rc.font.size=Pt(10.5); rc.font.name='Times New Roman'
    ps=doc.add_paragraph(); ps.alignment=WD_ALIGN_PARAGRAPH.CENTER; rs=ps.add_run('Fonte: elaboração dos autores (2026).'); rs.font.size=Pt(9.5); rs.font.name='Times New Roman'; ps.paragraph_format.space_after=Pt(6)

def cm(x): return str(x).replace('.',',')
pv={r['dim']:r for r in J['pv_mood']}
ph=J['phys']

# ===== título =====
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('CAPACIDADE AERÓBIA MÁXIMA E PERFIL DE HUMOR EM ATLETAS DE HANDEBOL DE ELITE: '
            'O PICO DE VELOCIDADE NO T-CAR COMO MODERADOR DA FADIGA DE UM MICROCICLO'); r.bold=True; r.font.size=Pt(14)
p.paragraph_format.space_after=Pt(12)

# ===== resumo =====
H1('RESUMO')
P(f'Objetivo: examinar a relação entre a capacidade aeróbia máxima, indexada pelo pico de velocidade (PV) '
  f'no teste de Carminatti (T-CAR), e o perfil de humor de atletas de handebol de elite ao longo de um '
  f'microciclo de pré-temporada, e descrever a adaptação da aptidão ao período. Métodos: estudo observacional '
  f'com {ph["desc"]["n"]} atletas na avaliação física (capacidade aeróbia, potência de salto e capacidade de '
  f'sprints repetidos, pré e pós-microciclo) e {J["n_pvmood"]} atletas com humor diário (Escala de Humor de '
  f'Brunel, seis dimensões). A associação entre o PV e a média semanal de cada dimensão do humor foi estimada '
  f'por correlação de Spearman, com intervalos de confiança por reamostragem (bootstrap) e correção de '
  f'multiplicidade por taxa de falsas descobertas (BH-FDR). Resultados: a capacidade aeróbia aumentou ao longo '
  f'do microciclo (PV {cm(ph["TCARpv"]["pre"].split(" ")[0])} → {cm(ph["TCARpv"]["pos"].split(" ")[0])} km/h; '
  f'dz = {cm(ph["TCARpv"]["dz"])}), com efeito grande no grupo experimental (dz = {cm(ph["TCARpv"]["by_group"]["Experimental"]["dz"])}). '
  f'Atletas com maior PV apresentaram menor fadiga física (ρ = {cm(pv["FadFisica"]["r"]):.5}; p = {cm(round(pv["FadFisica"]["p"],3))}; '
  f'sobrevivente à correção FDR) e maior vigor (ρ = +{cm(round(pv["Vigor"]["r"],2))}; p = {cm(round(pv["Vigor"]["p"],3))}); '
  f'as dimensões afetivas negativas não se associaram à aptidão. Conclusão: a capacidade aeróbia máxima comporta-se '
  f'como um moderador da fadiga do microciclo, especificamente no eixo energia–fadiga; o mais apto tolera melhor a '
  f'carga, fadigando menos e preservando o vigor.', ind=False)
P('Palavras-chave: capacidade aeróbia; velocidade de pico; teste de Carminatti; estados de humor; BRUMS; handebol.', ind=False)

# ===== 1 introducao =====
H1('INTRODUÇÃO','1')
P('O handebol é uma modalidade coletiva intermitente de alta intensidade, na qual a capacidade aeróbia sustenta '
  'a recuperação entre esforços, a manutenção do ritmo de jogo e a tolerância às cargas concentradas da '
  'pré-temporada. Entre os testes de campo, o teste de Carminatti (T-CAR) fornece o pico de velocidade (PV) como '
  'índice de capacidade aeróbia máxima, com validade e reprodutibilidade estabelecidas frente a medidas '
  'laboratoriais e associação com o desempenho em jogo (DA SILVA et al., 2011; CARMINATTI et al., 2013; '
  'FERNANDES-DA-SILVA et al., 2016; FLORIANO et al., 2016).')
P('Paralelamente, o monitoramento dos estados de humor tornou-se uma ferramenta sensível e de baixo custo para '
  'acompanhar a resposta do atleta à carga, com a Escala de Humor de Brunel (BRUMS) amplamente empregada no '
  'esporte (ROHLFS et al., 2008; TERRY; LANE; FOGARTY, 2003; SELMI et al., 2023; DUIGNAN et al., 2020). O modelo '
  'aptidão–fadiga sugere que características estáveis de aptidão modulam o custo interno de um mesmo estímulo: '
  'atletas mais aptos tenderiam a experimentar menor perturbação afetiva sob cargas equivalentes. Contudo, a '
  'articulação direta entre um índice objetivo de capacidade aeróbia máxima e o perfil de humor ao longo de um '
  'microciclo real de pré-temporada ainda é pouco descrita no handebol de elite (REBELO et al., 2024; HENZE et '
  'al., 2025).')
P('Este estudo teve por objetivo examinar a relação entre a capacidade aeróbia máxima (PV no T-CAR) e o perfil '
  'de humor de atletas de handebol de elite ao longo de um microciclo de pré-temporada, além de descrever a '
  'adaptação da aptidão ao período. A hipótese foi a de que uma maior capacidade aeróbia se associaria a menor '
  'fadiga e maior vigor — isto é, ao eixo energia–fadiga do humor —, sem relação relevante com as dimensões '
  'afetivas negativas, tipicamente comprimidas pelo efeito de piso.')

# ===== 2 metodo =====
H1('MÉTODO','2')
H2('Amostra e desenho','2.1')
P(f'Trata-se de um estudo observacional e longitudinal. Participaram {ph["desc"]["n"]} atletas do sexo masculino '
  f'de uma equipe de handebol de elite na bateria de avaliação física (idade {cm(ph["desc"]["idade"])} anos; massa '
  f'{cm(ph["desc"]["massa"])} kg; estatura {cm(ph["desc"]["estatura"])} m), avaliados antes e depois do microciclo, '
  f'e {J["n_pvmood"]} atletas com registro diário de humor durante a semana. Todos os dados foram anonimizados '
  f'antes da análise. O microciclo concentrou as maiores cargas do período preparatório.')
H2('Capacidade aeróbia máxima (T-CAR)','2.2')
P('A capacidade aeróbia máxima foi avaliada pelo teste de Carminatti (T-CAR), um teste incremental intermitente '
  'de campo, do qual se extraiu o pico de velocidade (PV, km/h) como índice aeróbio máximo (DA SILVA et al., '
  '2011; CARMINATTI et al., 2013). Como medidas complementares de aptidão neuromuscular e anaeróbia foram '
  'registrados o salto com contramovimento (CMJ) e o teste de sprints repetidos (Baker), todos avaliados antes '
  'e depois do microciclo.')
H2('Perfil de humor (BRUMS)','2.3')
P('O humor foi avaliado pela Escala de Humor de Brunel (BRUMS), composta por 24 itens que compõem seis '
  'dimensões — tensão, depressão, raiva, vigor, fadiga e confusão —, com registro diário ao longo da semana '
  '(ROHLFS et al., 2008). Para cada atleta, calculou-se a média semanal de cada dimensão, além dos índices de '
  'fadiga física e mental e da perturbação total do humor (PTH).')
H2('Análise estatística','2.4')
P('A adaptação física pré–pós foi descrita por médias, e a magnitude da mudança, pelo tamanho de efeito para '
  'medidas pareadas (dz), com o teste t pareado apenas como referência. A associação entre o PV e a média '
  'semanal de cada dimensão do humor foi estimada pela correlação de Spearman (ρ), robusta a desvios de '
  'normalidade e ao efeito de piso. Os intervalos de confiança de 95% foram obtidos por reamostragem bootstrap '
  '(5.000 repetições) e a multiplicidade das nove correlações foi controlada pela taxa de falsas descobertas '
  '(Benjamini–Hochberg, FDR). Adotou-se α = 5%.')

# ===== 3 resultados =====
H1('RESULTADOS','3')
H2('Adaptação da aptidão ao microciclo','3.1')
P(f'A capacidade aeróbia máxima aumentou ao longo do microciclo: o PV passou de '
  f'{cm(ph["TCARpv"]["pre"].split(" ")[0])} para {cm(ph["TCARpv"]["pos"].split(" ")[0])} km/h '
  f'(dz = {cm(ph["TCARpv"]["dz"])}; Figura 1), com efeito grande no grupo experimental '
  f'(dz = {cm(ph["TCARpv"]["by_group"]["Experimental"]["dz"])}) e trivial no controle '
  f'(dz = {cm(ph["TCARpv"]["by_group"]["Controle"]["dz"])}). Também houve melhora do salto vertical '
  f'(CMJ, dz = {cm(ph["CMJ"]["dz"])}) e da capacidade de sprints repetidos (Baker soma, '
  f'dz = {cm(ph["BkSoma"]["dz"])}), o que confirma que o período induziu adaptação, sobretudo aeróbia.')
figure('aerobio_adaptacao.png','Adaptação da capacidade aeróbia máxima (pico de velocidade no T-CAR) do pré ao pós-microciclo, por grupo. Linhas finas: atletas; linhas grossas: média do grupo.',w=11.0)
H2('Capacidade aeróbia e perfil de humor','3.2')
P(f'A capacidade aeróbia máxima associou-se ao eixo energia–fadiga do humor (Tabela 1; Figura 2). Atletas com '
  f'maior PV apresentaram menor fadiga física (ρ = {cm(round(pv["FadFisica"]["r"],2))}; IC95% '
  f'[{cm(round(pv["FadFisica"]["lo"],2))}, {cm(round(pv["FadFisica"]["hi"],2))}]; p = {cm(round(pv["FadFisica"]["p"],3))}), '
  f'associação que permaneceu significativa após a correção por multiplicidade (FDR = {cm(round(pv["FadFisica"]["fdr"],3))}), '
  f'e maior vigor (ρ = +{cm(round(pv["Vigor"]["r"],2))}; IC95% [{cm(round(pv["Vigor"]["lo"],2))}, '
  f'{cm(round(pv["Vigor"]["hi"],2))}]; p = {cm(round(pv["Vigor"]["p"],3))}). A fadiga geral seguiu a mesma direção, '
  f'com magnitude moderada e significância limítrofe (ρ = {cm(round(pv["Fadiga"]["r"],2))}; p = {cm(round(pv["Fadiga"]["p"],3))}). '
  f'As dimensões afetivas negativas — tensão, depressão, raiva e confusão — não se associaram à aptidão '
  f'(|ρ| ≤ 0,12; p > 0,50), coerente com a sua concentração no piso da escala.')
# Tabela 1
_TN[0]+=1
cap=doc.add_paragraph(); rc=cap.add_run('Tabela %d – Correlação de Spearman entre o pico de velocidade no T-CAR e a média semanal de cada dimensão do humor.'%_TN[0]); rc.font.size=Pt(10.5); rc.font.name='Times New Roman'
tbl=doc.add_table(rows=1,cols=5); tbl.style='Light Grid Accent 1'
for j,h in enumerate(['Dimensão','ρ (Spearman)','IC95%','p','p (FDR)']): tbl.rows[0].cells[j].text=h
ORD=['Vigor','Fadiga','FadFisica','FadMental','TMD','Tensao','Depressao','Raiva','Confusao']
for k in ORD:
    r=pv[k]; c=tbl.add_row().cells
    c[0].text=r['lab']; c[1].text=('+' if r['r']>=0 else '')+cm(round(r['r'],2))
    c[2].text=f"[{cm(round(r['lo'],2))}, {cm(round(r['hi'],2))}]"; c[3].text=cm(round(r['p'],3)); c[4].text=cm(round(r['fdr'],3))
for row in tbl.rows:
    for cc in row.cells:
        for pp in cc.paragraphs:
            for rr in pp.runs: rr.font.size=Pt(10)
nt=doc.add_paragraph(); rn=nt.add_run(f'Nota: n = {J["n_pvmood"]} atletas. FDR: correção de Benjamini–Hochberg para as nove correlações. PTH: perturbação total do humor.'); rn.font.size=Pt(9.5); rn.italic=True
figure('aerobio_humor_assoc.png','Associação entre a capacidade aeróbia máxima (PV no T-CAR) e o humor. À esquerda, ρ de Spearman (IC95%) por dimensão (* FDR < 0,05; † p < 0,05); ao centro e à direita, dispersões da fadiga física e do vigor em função do PV.',w=15.5)

# ===== 4 discussao =====
H1('DISCUSSÃO','4')
P('Os resultados sustentam a hipótese de que a capacidade aeróbia máxima modula o custo afetivo de um microciclo '
  'intenso de pré-temporada, e o fazem de forma específica: a associação concentrou-se no eixo energia–fadiga — '
  'menor fadiga física e maior vigor nos atletas mais aptos — e esteve ausente nas dimensões afetivas negativas. '
  'Esse padrão é coerente com o modelo aptidão–fadiga: sob uma mesma carga externa, um sistema aeróbio mais '
  'desenvolvido impõe menor perturbação interna e recupera-se mais rápido entre as sessões, o que se traduz em '
  'um perfil de humor mais preservado no polo da energia.')
P('A especificidade do achado é informativa. A fadiga física foi a única dimensão a sobreviver à correção por '
  'multiplicidade, e o vigor apresentou associação positiva de magnitude moderada; já tensão, depressão, raiva e '
  'confusão — comprimidas pelo efeito de piso em atletas saudáveis — não se relacionaram à aptidão. Isso reforça '
  'a leitura, recorrente na literatura de monitoramento, de que o par vigor–fadiga carrega o sinal mais forte e '
  'mais interpretável do estado do atleta, ao passo que as dimensões negativas funcionam como sentinelas de '
  'estados raros, e não como termômetros graduais da carga (SAW; MAIN; GASTIN, 2016; SELMI et al., 2023).')
P('Do ponto de vista aplicado, o pico de velocidade no T-CAR — medida de campo simples e de baixo custo — '
  'identifica atletas mais vulneráveis à fadiga do microciclo, permitindo individualizar a prescrição e a '
  'recuperação. Como o estímulo intervalado costuma ser prescrito em relação à velocidade de pico individual, o '
  'custo interno é relativo à aptidão de cada atleta, o que ajuda a explicar por que a vulnerabilidade à fadiga '
  'da semana se associa a um traço aeróbio prévio, e não à intensidade absoluta corrida. As principais limitações '
  'são o tamanho amostral, próprio de um elenco de elite, e o desenho observacional, que impõe cautela causal; a '
  'convergência com a adaptação aeróbia objetiva do período, contudo, dá plausibilidade à interpretação.')

# ===== 5 conclusao =====
H1('CONCLUSÃO','5')
P('A capacidade aeróbia máxima, indexada pelo pico de velocidade no T-CAR, comporta-se como um moderador da '
  'fadiga de um microciclo de pré-temporada no handebol de elite: quanto maior a aptidão, menor a fadiga física '
  'e maior o vigor, sem relação com as dimensões afetivas negativas. O achado situa a aptidão aeróbia como um '
  'fator de proteção no eixo energia–fadiga do humor e recomenda o uso conjunto do T-CAR e do BRUMS no '
  'monitoramento individualizado da carga.')

# ===== referencias =====
H1('REFERÊNCIAS')
REFS=[
 'CARMINATTI, L. J.; POSSAMAI, C. A. P.; DE MORAES, M.; DA SILVA, J. F.; DE LUCAS, R. D.; DITTRICH, N.; GUGLIELMO, L. G. A. Intermittent versus continuous incremental field tests: are maximal variables interchangeable? Journal of Sports Science and Medicine, v. 12, n. 1, p. 165-170, 2013.',
 'DA SILVA, J. F.; GUGLIELMO, L. G. A.; CARMINATTI, L. J.; DE OLIVEIRA, F. R.; DITTRICH, N.; PATON, C. D. Validity and reliability of a new field test (Carminatti’s test) for soccer players compared with laboratory-based measures. Journal of Sports Sciences, v. 29, n. 15, p. 1621-1628, 2011.',
 'DUIGNAN, C.; DOHERTY, C.; CAULFIELD, B.; BLAKE, C. Single-item self-report measures of team-sport athlete wellbeing and their relationship with training load: a systematic review. Journal of Athletic Training, v. 55, n. 9, p. 944-953, 2020.',
 'FERNANDES-DA-SILVA, J.; CASTAGNA, C.; TEIXEIRA, A. S.; CARMINATTI, L. J.; GUGLIELMO, L. G. A. The peak velocity derived from the Carminatti Test is related to physical match performance in young soccer players. Journal of Sports Sciences, v. 34, n. 24, p. 2238-2245, 2016.',
 'FLORIANO, L. T.; DA SILVA, J. F.; TEIXEIRA, A. S.; SALVADOR, P. C. N.; DITTRICH, N.; CARMINATTI, L. J.; NASCIMENTO, L. L.; GUGLIELMO, L. G. A. Physiological responses during the time limit at 100% of the peak velocity in the Carminatti’s test in futsal players. Journal of Human Kinetics, v. 54, p. 91-101, 2016.',
 'HENZE, A. S.; KIRSTEN, J.; MATITS, L.; BIZJAK, D. A.; SCHULZ, S. V. W. Athlete monitoring in handball (ATHMON HB): a systematic review protocol. Systematic Reviews, v. 14, n. 1, art. 64, 2025.',
 'REBELO, A.; PEREIRA, J. R.; CUNHA, P.; COELHO-E-SILVA, M. J.; VALENTE-DOS-SANTOS, J. Training stress, neuromuscular fatigue and well-being in volleyball: a systematic review. BMC Sports Science, Medicine and Rehabilitation, v. 16, n. 1, art. 17, 2024.',
 'ROHLFS, I. C. P. M.; ROTTA, T. M.; LUFT, C. D. B.; ANDRADE, A.; KREBS, R. J.; CARVALHO, T. A Escala de Humor de Brunel (BRUMS): instrumento para detecção precoce da síndrome do excesso de treinamento. Revista Brasileira de Medicina do Esporte, v. 14, n. 3, p. 176-181, 2008.',
 'SAW, A. E.; MAIN, L. C.; GASTIN, P. B. Monitoring the athlete training response: subjective self-reported measures trump commonly used objective measures — a systematic review. British Journal of Sports Medicine, v. 50, n. 5, p. 281-291, 2016.',
 'SELMI, O.; OUERGUI, I.; MUSCELLA, A.; LEVITT, D. E.; SUZUKI, K.; BOUASSIDA, A. Monitoring mood state to improve performance in soccer players: a brief review. Frontiers in Psychology, v. 14, art. 1095238, 2023.',
 'TERRY, P. C.; LANE, A. M.; FOGARTY, G. J. Construct validity of the Profile of Mood States-Adolescents for use with adults. Psychology of Sport and Exercise, v. 4, n. 2, p. 125-139, 2003.',
]
for r in REFS:
    p=doc.add_paragraph(); rr=p.add_run(r); rr.font.size=Pt(11); rr.font.name='Times New Roman'
    p.paragraph_format.space_after=Pt(6); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

doc.save('/home/user/mdlucca/Artigos/Artigo_Aerobio_Humor.docx')
print('OK Artigo_Aerobio_Humor.docx  figs=%d tabs=%d refs=%d'%(_FN[0],_TN[0],len(REFS)))
