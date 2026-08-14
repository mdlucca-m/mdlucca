import warnings; warnings.filterwarnings('ignore')
import numpy as np, pandas as pd, json, os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
SC='/tmp/claude-0/-home-user-mdlucca/e1dba24c-b1d7-5908-9106-f2f4aaf3f56a/scratchpad'
FIG='/home/user/mdlucca/Artigos/figuras'
G=json.load(open(f'{SC}/pair_fatigue.json'))
grp,btw,per,PP,corr,eff,kmdc=G['grp'],G['btw'],G['per'],G['pp'],G['corr'],G['eff'],G['kmdc']
VARS=[('FadFisica','Fadiga física (0–10)'),('FadMental','Fadiga mental (0–10)')]

doc=Document(); st=doc.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(10.5)
def H(t,l=1): doc.add_heading(t,level=l)
def P(t,it=False,b=False):
    p=doc.add_paragraph(); r=p.add_run(t); r.italic=it; r.bold=b; return p
def fig(name,cap,w=6.4):
    path=f'{FIG}/{name}.png'
    if os.path.exists(path):
        doc.add_picture(path,width=Inches(w)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
        c=doc.add_paragraph(); r=c.add_run(cap); r.italic=True; r.font.size=Pt(8.5); c.alignment=WD_ALIGN_PARAGRAPH.CENTER
def table(headers,rows,cap=None):
    if cap: c=doc.add_paragraph(); r=c.add_run(cap); r.bold=True; r.font.size=Pt(9.5)
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Light Grid Accent 1'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        cell=t.rows[0].cells[i]; cell.paragraphs[0].add_run(str(h)).bold=True; cell.paragraphs[0].runs[0].font.size=Pt(8.5)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=str(v)
            if cells[i].paragraphs[0].runs: cells[i].paragraphs[0].runs[0].font.size=Pt(8.5)
    doc.add_paragraph()

# CAPA
doc.add_heading('Amplitude e Separação Ruído × Sinal — Fadiga física vs Fadiga mental',0)
P('Fadiga física (0–10) e Fadiga mental (0–10) ao longo do microciclo de choque de HIIT — 21 a 28 de abril de 2024',it=True)
P('Análise por grupo e por atleta · dentro da semana e entre dias · pré/mid/pós · remoção de ruído · 27 atletas · 456 observações · reanálise independente em Python. Atletas anonimizados (A01–A27).')
doc.add_paragraph()

# 1 método
H('1. Objetivo e método')
P('Aplica-se às duas autoavaliações de fadiga (física e mental, escalas 0–10) o mesmo protocolo: amplitude (range) no grupo e no atleta, decomposição da variância em sinal do microciclo / traço / ruído (modelo y ~ atleta + dia), limiares de decisão a partir do ruído (ETM, MDC95 individual e de grupo, SWC), curso intra-dia (pré→mid→pós) e o comportamento das análises quando o ruído é removido. A pergunta central deste par é a dissociação: a fadiga física e a mental respondem igual ao microciclo?')

# 2 grupo semana
H('2. Grupo — ao longo da semana')
rows=[[lab,f"{grp[v]['d1']:.2f}",f"{grp[v]['d7']:.2f}",f"{grp[v]['drift']:+.2f}",f"{grp[v]['sig']:.2f}",f"{grp[v]['etm']:.2f}",f"{grp[v]['snr']:.2f}"] for v,lab in VARS]
table(['Variável','Dia 1','Dia 7','Deriva D7−D1','Ampl. do sinal','ETM (ruído)','SNR'],rows,'Tabela 1 — Trajetória do grupo e amplitude do sinal')
P('A dissociação é imediata: a fadiga física quase dobra na semana (deriva +3,30; amplitude do sinal 3,30; SNR 2,02), enquanto a fadiga mental fica praticamente parada (deriva +0,30; amplitude 0,92; SNR 0,62 < 1 — a oscilação semanal é menor que o próprio ruído). A física é um marcador de carga; a mental, não.')
fig('pf1_trajetoria','Fig 1 — Fadiga física vs mental ao longo da semana (média ± DP; LOWESS). A física sobe com precipitação no D7; a mental permanece plana.')
fig('pf2_dissociacao','Fig 2 — Sobreposição das duas curvas suavizadas: a física acumula; a mental é estável — a assinatura da dissociação somática vs cognitiva.')
rows=[[lab,f"{grp[v]['pday']:.1f}%",f"{grp[v]['ptrait']:.1f}%",f"{grp[v]['pnoise']:.1f}%"] for v,lab in VARS]
table(['Variável','% Sinal (dia)','% Traço','% Ruído'],rows,'Tabela 2 — Decomposição da variância')
P('Na fadiga física, 12,0% da variância é sinal do microciclo — a maior fatia de sinal de todas as variáveis do estudo. Na fadiga mental, só 1,3% é sinal e 72,9% é traço: é um construto estável (quase um traço de personalidade dentro da janela), não um marcador reativo à carga.')

# 3 entre dias
H('3. Grupo — entre dias')
rows=[[lab]+[f"{btw[v]['sd_between'].get(str(d),'—')}" for d in range(1,8)]+[f"D{btw[v]['day_maxsd']}"] for v,lab in VARS]
table(['Variável','D1','D2','D3','D4','D5','D6','D7','Máx. disp.'],rows,'Tabela 3 — Dispersão ENTRE atletas por dia (DP)')
rows=[[lab]+[f"{btw[v]['dd'].get(str(d),'—')}" for d in range(2,8)]+[f"D{btw[v]['day_maxdd']}"] for v,lab in VARS]
table(['Variável','D1→2','D2→3','D3→4','D4→5','D5→6','D6→7','Maior passo'],rows,'Tabela 4 — |Mudança dia-a-dia| da média do grupo')
P('O maior passo da fadiga física ocorre no fim da semana (a precipitação do D7); a fadiga mental não tem um passo dominante — seus incrementos diários são todos pequenos e dentro do ruído. A dispersão entre atletas é alta e persistente nas duas.')

# 4 por atleta
H('4. Por atleta — amplitude e sinal-ruído individual')
rows=[[lab,f"{per[v]['amp_mean']:.2f} ± {per[v]['amp_sd']:.2f}",f"{per[v]['amp_min']:.0f}–{per[v]['amp_max']:.0f}",f"{grp[v]['sig']:.2f}",f"{per[v]['pnoise_mean']:.2f}",f"{per[v]['pct_snr_gt1']:.0f}%"] for v,lab in VARS]
table(['Variável','Ampl. atleta (média±DP)','Faixa','Ampl. sinal grupo','Ruído pessoal','% atletas SNR>1'],rows,'Tabela 5 — Amplitude por atleta vs sinal do grupo')
P('As duas variáveis têm amplitude individual grande (a física 6,0; a mental 5,0 em média) — mas na mental essa amplitude é quase toda ruído (a amplitude do sinal do grupo é só 0,92, e o atleta oscila 5 pontos em torno disso). Na física, a oscilação individual carrega um sinal maior. Ainda assim, em ambas a maioria dos atletas tem SNR pessoal > 1 sobre médias diárias.')
fig('pf3_amp_box','Fig 3 — Amplitude por atleta (caixa) vs amplitude do sinal do grupo (losango). A caixa fica muito acima do losango na mental — quase tudo é ruído individual.')

# 5 prepost
H('5. Pré → mid → pós por dia')
P('Curso intra-sessão: média ± DP de pré, mid (intra-sessão) e pós, Δ (pós−pré), dz pareado, Wilcoxon pré×pós e Friedman pré/mid/pós. Dia 1 (baseline) sem mid. * indica p<0,05.')
for vi,(v,lab) in enumerate(VARS):
    rows=[]
    for d in range(1,8):
        r=PP[v][str(d)]
        pstr='—' if r['p'] is None else (f"{r['p']:.3f}"+(' *' if r['p']<0.05 else ''))
        fstr='—' if r['pfried'] is None else (f"{r['pfried']:.3f}"+(' *' if r['pfried']<0.05 else ''))
        dzs='—' if r['dz'] is None else f"{r['dz']:+.2f}"
        mids='—' if r['mid_m'] is None else (f"{r['mid_m']:.2f} ± {r['mid_sd']:.2f}" if r['mid_sd'] is not None else f"{r['mid_m']:.2f}")
        rows.append([f"D{d}"+(" (HIIT)" if r['hiit'] else ""),f"{r['pre_m']:.2f} ± {r['pre_sd']:.2f}",mids,f"{r['pos_m']:.2f} ± {r['pos_sd']:.2f}",f"{r['delta']:+.2f}",dzs,pstr,fstr])
    table(['Dia','Pré (m ± DP)','Mid (m ± DP)','Pós (m ± DP)','Δ (pós−pré)','dz','p pré×pós','p Friedman'],rows,f'Tabela {6+vi} — {lab}: pré → mid → pós por dia')
P('A fadiga física responde agudamente pré→pós em praticamente todos os dias de treino (dz 0,5–1,0), com pico intra-sessão no mid. A fadiga mental responde de forma esparsa (apenas D1 e D7 atingem significância) e com efeitos pequenos — confirma que a sessão de HIIT taxa o corpo, não a percepção mental de esforço cognitivo.')
fig('pf4_prepos','Fig 4 — Pré → mid → pós por dia. A física (esquerda) sobe consistentemente pós-treino; a mental (direita) mal se move.')

# 6 veredito
H('6. Veredito em dois níveis — grupo vs indivíduo')
rows=[[lab,f"{grp[v]['sig']:.2f}",f"{grp[v]['mdc']:.2f}",('SIM' if grp[v]['detect_ind'] else 'NÃO'),f"{grp[v]['mdc_g']:.2f}",('SIM' if grp[v]['detect_grp'] else 'NÃO')] for v,lab in VARS]
table(['Variável','Ampl. do sinal','MDC95 individual','> MDC ind.?','MDC95 grupo (÷√n)','> MDC grupo?'],rows,'Tabela 8 — A oscilação semanal supera o ruído?')
P('A fadiga física é sinal robusto no grupo (amplitude 3,30 vs MDC do grupo 0,32 — supera 10×) mas fica abaixo do MDC95 individual de 1 coleta (4,54): decisão individual exige médias. A fadiga mental supera o MDC do grupo por pouco (amplitude 0,92 vs 0,29) mas seu sinal é tão pequeno que carece de relevância prática — e no indivíduo é indetectável.')

# 7 remover ruido
H('7. E se removermos o ruído? Como ficam as análises')
P('Remover o ruído (agregar, suavizar ou desatenuar) tem efeito diferente em cada nível — e aqui revela a natureza da relação entre as duas fadigas.')
rows=[['Fadiga física × mental — observada (bruta, Pearson)',f"{corr['r_obs']:+.2f}"],
      ['— Spearman (bruta)',f"{corr['rho_obs']:+.2f}"],
      ['— rmcorr (intra-atleta, momentânea)',f"{corr['rm']:+.2f}"],
      ['— desatenuada (erro de medida removido)',f"{corr['r_dis']:+.2f}"],
      ['— nível-grupo (médias diárias; ruído agregado)',f"{corr['r_den']:+.2f}"]]
table(['Estimativa da associação física × mental','r'],rows,'Tabela 9 — A correlação física × mental em vários níveis')
P('Duas verdades coexistem. (1) No nível momentâneo, física e mental sobem juntas quando removemos o erro de medida: a correlação bruta 0,50 passa a 0,78 desatenuada — há acoplamento real instante-a-instante (rmcorr 0,44). (2) No nível da trajetória semanal, elas dissociam: a correlação entre as médias diárias cai para 0,37, porque ao longo da semana a física acumula e a mental não. Ou seja: num dado momento, quem está fisicamente exausto tende a relatar mais fadiga mental; mas o microciclo, como processo, carrega apenas a física. Remover o ruído não apaga a dissociação — torna-a precisa.')
fig('pf5_scatter','Fig 5 — Física × mental: observações brutas (esquerda; r=0,50) e médias diárias do grupo (direita; r=0,37). A relação momentânea é maior que a relação de tendência semanal.')
rows=[[lab,f"{eff[v]['dz']:+.2f}",f"{eff[v]['dz_corr']:+.2f}",f"{grp[v]['rel']:.2f}"] for v,lab in VARS]
table(['Variável','dz observado','dz corrigido (ruído removido)','Confiab. (1 medida)'],rows,'Tabela 10 — Efeito D1→D7 corrigido pela atenuação')
P('Corrigido pelo ruído, o efeito da semana na fadiga física cresce de dz +1,74 para +2,36 (colossal); na mental permanece pequeno (+0,33 → +0,38). Remover o ruído amplia o que já é grande e não resgata o que é nulo.')
rows=[]
for v,lab in VARS:
    emg=next((k for k in [1,2,3,5,7] if grp[v]['sig']>kmdc[v][str(k)]),None)
    rows.append([lab,f"{grp[v]['sig']:.2f}",f"{kmdc[v]['1']:.2f}",f"{kmdc[v]['3']:.2f}",f"{kmdc[v]['5']:.2f}",f"{kmdc[v]['7']:.2f}",(f"k≥{emg}" if emg else "não emerge")])
table(['Variável','Ampl. sinal','MDC k=1','MDC k=3','MDC k=5','MDC k=7','Sinal supera o ruído em'],rows,'Tabela 11 — MDC95 individual cai com k coletas (ruído/√k)')
P('Para a fadiga física, a partir de ~3 coletas o sinal semanal supera o MDC individual — o monitoramento individual da fadiga física é viável com médias curtas. A fadiga mental não emerge nem com 7 coletas: por não haver sinal semanal, nenhuma quantidade de agregação a torna um marcador útil de carga.')
fig('pf6_mdc','Fig 6 — MDC95 individual em função do nº de coletas (k). A linha branca é a amplitude do sinal semanal: na física a curva do MDC cruza abaixo dela (~k=3); na mental, nunca.')

# 8 sintese
H('8. Síntese')
for t in [
 'Dissociação clara: a fadiga física quase dobra na semana (deriva +3,3; dz +1,74; 12% da variância é sinal); a fadiga mental fica plana (deriva +0,3; dz +0,33; 1,3% de sinal; SNR<1).',
 'Amplitude: as duas oscilam bastante no atleta, mas na mental a oscilação é quase toda ruído (sinal de grupo só 0,92); na física há sinal real (2,0× o ruído).',
 'Intra-sessão: a física sobe pós-treino em quase todos os dias (pico no mid); a mental só no D1 e D7.',
 'Removendo o ruído: física e mental acoplam-se no momento (0,50→0,78 desatenuada) mas dissociam na tendência semanal (0,37 nível-grupo); o efeito da física cresce (dz +1,74→+2,36), o da mental não.',
 'Uso prático: monitore a fadiga física (viável no indivíduo a partir de ~3 coletas); a fadiga mental não serve de marcador de carga aguda deste microciclo — o custo do HIIT é somático, não cognitivo.',
]:
    p=doc.add_paragraph(style='List Bullet'); p.add_run(t)
doc.add_paragraph()
P('Reprodutibilidade: scripts/analise/pair_fatigue.py, build_pair_docx.py. Consistente com Amplitude_Ruido_Sinal_Vigor_Fadiga_TMD.docx e Analise_Estatistica_Consolidada.md. Atletas anonimizados; nenhum dado bruto com nomes é distribuído.',it=True)

OUTP='/home/user/mdlucca/Artigos/Amplitude_Ruido_Sinal_Fadiga_Fisica_Mental.docx'
doc.save(OUTP); print('SAVED',OUTP)
d2=Document(OUTP); print('paras',len(d2.paragraphs),'tables',len(d2.tables),'imgs',len(d2.inline_shapes))
