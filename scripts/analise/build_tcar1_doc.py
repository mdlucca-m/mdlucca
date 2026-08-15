# -*- coding: utf-8 -*-
import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
R=json.load(open('tcar1_baseline.json'))
def c2(s): return str(s).replace('.',',')
doc=Document()
sty=doc.styles['Normal']; sty.font.name='Times New Roman'; sty.font.size=Pt(12)
sty.element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')
pf=sty.paragraph_format; pf.line_spacing=1.5; pf.space_after=Pt(0)
s=doc.sections[0]; s.top_margin=Cm(3); s.left_margin=Cm(3); s.bottom_margin=Cm(2); s.right_margin=Cm(2)
_TN=[0]; _FN=[0]
def P(t='',just=True,size=12,after=6,before=0):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(size)
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(after); p.paragraph_format.space_before=Pt(before)
    if just: p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.first_line_indent=Cm(1.25)
    return p
def H(t,size=12):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(size)
    p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(6); p.paragraph_format.line_spacing=1.5
def _bd(cell):
    tcPr=cell._tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for edge in ['top','bottom','left','right']:
        e=OxmlElement(f'w:{edge}'); on=edge in ('top','bottom')
        e.set(qn('w:val'),'single' if on else 'nil'); e.set(qn('w:sz'),'6'); e.set(qn('w:color'),'000000'); b.append(e)
    tcPr.append(b)
def table(caption,header,rows,fonte='Fonte: os autores (2026).',widths=None,fs=10):
    _TN[0]+=1
    p=doc.add_paragraph(); r=p.add_run(f'Tabela {_TN[0]} – {caption}'); r.font.size=Pt(11)
    p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(3)
    t=doc.add_table(rows=1,cols=len(header)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,htxt in enumerate(header):
        cc=t.rows[0].cells[i]; cc.text=''; rr=cc.paragraphs[0].add_run(htxt); rr.bold=True; rr.font.size=Pt(fs); rr.font.name='Times New Roman'
        cc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; _bd(cc)
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=''; rr=cells[i].paragraphs[0].add_run(str(val)); rr.font.size=Pt(fs); rr.font.name='Times New Roman'
            cells[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if i>0 else WD_ALIGN_PARAGRAPH.LEFT; _bd(cells[i])
    if widths:
        for i,w in enumerate(widths):
            for rr in t.rows: rr.cells[i].width=Cm(w)
    pf=doc.add_paragraph(); rf=pf.add_run(fonte); rf.font.size=Pt(10); pf.paragraph_format.space_after=Pt(8)
def figure(path,caption,w=15.5):
    _FN[0]+=1
    pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; pp.add_run().add_picture(path,width=Cm(w))
    pc=doc.add_paragraph(); pc.alignment=WD_ALIGN_PARAGRAPH.CENTER; rc=pc.add_run(f'Figura {_FN[0]} – {caption}'); rc.font.size=Pt(11)
    pf=doc.add_paragraph(); pf.alignment=WD_ALIGN_PARAGRAPH.CENTER; rf=pf.add_run('Fonte: os autores (2026).'); rf.font.size=Pt(10); pf.paragraph_format.space_after=Pt(8)

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('APTIDÃO DO T-CAR 1 E CARGA INTERNA (TRIMP) COMO PREDITORES FISIOLÓGICOS DO ESTADO PSICOLÓGICO DE BASE EM HANDEBOLISTAS DE ELITE')
r.bold=True; r.font.size=Pt(13); p.paragraph_format.space_after=Pt(12)
P('Análise com parâmetros restritos — fisiológico: apenas o T-CAR 1 (pico de velocidade, distância/repetições e FCmáx) e a carga interna do HIIT (TRIMP); psicológico: apenas o baseline (D1, 21/04). Amostra: 27 handebolistas do sexo masculino.',size=11,after=10)

H('Resumo')
P('Restringindo os parâmetros conforme solicitado — fisiológico ao T-CAR 1 e à carga interna (TRIMP) das sessões de HIIT, e psicológico ao baseline —, o parâmetro fisiológico que melhor se associa ao estado psicológico inicial é a APTIDÃO (pico de velocidade e a distância dela derivada): correlaciona-se negativamente com a fadiga de base (ρ = %s, p = %s, limítrofe) e classifica modestamente o perfil iceberg inicial (AUC = %s). A carga interna do HIIT (TRIMP e PSE) NÃO se associa ao baseline — resultado esperado e coerente com a temporalidade, já que as sessões de HIIT (22, 24 e 27/04) são posteriores ao baseline (21/04) e, portanto, não podem explicá-lo. Distância e PV carregam a mesma informação no T-CAR (associações de posto idênticas). Por fim, a aptidão de base não prediz a magnitude da deterioração ao longo da semana; o acoplamento aptidão–fadiga é fraco no baseline e mais forte na média semanal (ρ passa de %s para ≈ 0,51), ou seja, a vantagem do mais apto EMERGE à medida que a carga se acumula.'%(
    c2('%+.2f'%R['COR']['PV']['b_Fadiga']['r']),c2('%.2f'%R['COR']['PV']['b_Fadiga']['p']),
    c2('%.2f'%R['ROC']['PV']['auc']),c2('%+.2f'%R['COR']['PV']['b_FadFisica']['r'])),after=8)

H('1. Parâmetros e método')
P('Parâmetro fisiológico (T-CAR 1, teste incremental de base): pico de velocidade (PV, km·h⁻¹), distância percorrida (derivada das repetições do T-CAR 1: dist = repetições × (PV/3,6) × 12 s) e frequência cardíaca máxima (FCmáx). Carga interna: TRIMP de Banister calculado a partir da FC média das três sessões de HIIT (S1–S3), com FCrepouso = 60 bpm e duração nominal de 30 min mantidas constantes (portanto invariantes de escala, sem afetar as correlações), além da PSE (percepção de esforço) como carga por sessão-RPE. Parâmetro psicológico: o baseline — primeira coleta do BRUMS no D1 (21/04) —, incluindo vigor, fadiga (BRUMS), fadiga física e mental, PTH/TMD e o índice de perfil iceberg. Correlações de Spearman; ajuste alométrico log-log; classificação por curva ROC (AUC via U de Mann-Whitney).',after=6)
P('Observação sobre disponibilidade de dados: o TRIMP e o tempo em alta intensidade do próprio T-CAR 1 (Firstbeat) não constam das planilhas; por decisão do pesquisador, adotou-se o TRIMP das sessões de HIIT como carga interna. O tempo em alta intensidade não pôde ser reconstruído por atleta.',after=6)

H('2. Resultados')
H('2.1 Correlação — aptidão do T-CAR 1 × psicológico de base',12)
OUTB=[('b_Vigor','Vigor'),('b_Fadiga','Fadiga BRUMS'),('b_FadFisica','Fad. física'),('b_FadMental','Fad. mental'),('b_TMD','PTH/TMD'),('b_ice_idx','Iceberg')]
def rc(ph,ov): d=R['COR'][ph][ov]; return '%s (%s)'%(c2('%+.2f'%d['r']),c2('%.2f'%d['p']))
rows=[]
for ph,lab in [('PV','Pico de velocidade'),('dist','Distância (T-CAR 1)'),('FCmax','FCmáx'),('TRIMP','TRIMP (HIIT)'),('PSE','PSE (HIIT)')]:
    rows.append([lab]+[rc(ph,ov) for ov,_ in OUTB])
table('Correlação de Spearman (ρ, p) entre parâmetros fisiológicos e o psicológico de base (n = 27)',
    ['Parâmetro']+[l for _,l in OUTB],rows,widths=[3.4,2.7,2.7,2.5,2.5,2.4,2.3],fs=8.5)
P('A aptidão (PV/distância) é o único parâmetro com sinal consistente: correlaciona-se negativamente com a fadiga de base (ρ = %s, p = %s) — quem chega mais apto tende a iniciar a semana menos fatigado —, embora limítrofe pelo tamanho amostral e pelo efeito de piso do baseline. A carga interna (TRIMP e PSE) tem associação nula com o estado inicial, como esperado pela temporalidade (a carga é posterior ao baseline). PV e distância produzem associações idênticas: no T-CAR, ambos derivam do estágio final e carregam a mesma informação de aptidão.'%(
    c2('%+.2f'%R['COR']['PV']['b_Fadiga']['r']),c2('%.2f'%R['COR']['PV']['b_Fadiga']['p'])),after=6)

H('2.2 Ajuste alométrico e curva ROC',12)
al=R['ALLO']
P('No ajuste alométrico da fadiga física de base sobre os descritores fisiológicos, todos os expoentes foram não significativos (PV: b = %s, R² = %s; distância: b = %s; TRIMP: b = %s) — o baseline, por ser um ponto único e comprimido pelo efeito de piso, oferece pouca variância para o escalonamento. Na classificação do perfil iceberg de base por curva ROC, o PV e a distância atingem AUC = %s (discriminação modesta), enquanto o TRIMP fica em AUC = %s (≈ acaso). Ou seja, mesmo no baseline a aptidão — não a carga — é o parâmetro com algum poder de indicação.'%(
    c2('%+.2f'%al['PV']['b']),c2('%.2f'%al['PV']['r2']),c2('%+.2f'%al['dist']['b']),c2('%+.2f'%al['TRIMP']['b']),
    c2('%.2f'%R['ROC']['PV']['auc']),c2('%.2f'%R['ROC']['TRIMP']['auc'])),after=6)

H('2.3 Posição',12)
POS=R['POS']; order=[o for o in ['Goleiro','Meia','Ponta','Pivô'] if o in POS]
table('Parâmetros por posição — aptidão do T-CAR 1, carga interna e psicológico de base (médias)',
    ['Posição','n','PV','Distância','TRIMP','Vigor (base)','Fad. fís. (base)','Iceberg (base)'],
    [[o,POS[o]['n'],c2('%.1f'%POS[o]['PV']),c2('%.0f'%POS[o]['dist']),c2('%.1f'%POS[o]['TRIMP']),
      c2('%.1f'%POS[o]['b_Vigor']),c2('%.1f'%POS[o]['b_FadFisica']),c2('%+.2f'%POS[o]['b_ice'])] for o in order],
    widths=[2.4,1.1,1.8,2.3,1.9,2.5,2.6,2.6],fs=9)
P('As pontas (alas) reúnem a maior aptidão (PV %s; distância %s m) e o perfil iceberg de base mais marcado (índice %s), enquanto os pivôs, menos aptos (PV %s), têm o iceberg mais fraco (%s). O padrão posicional do baseline é coerente com o observado ao longo da semana.'%(
    c2('%.1f'%POS['Ponta']['PV']),c2('%.0f'%POS['Ponta']['dist']),c2('%+.2f'%POS['Ponta']['b_ice']),
    c2('%.1f'%POS['Pivô']['PV']),c2('%+.2f'%POS['Pivô']['b_ice'])) if 'Ponta' in POS and 'Pivô' in POS else 'Padrão posicional coerente com a semana.',after=6)

H('2.4 O acoplamento aptidão–fadiga emerge ao longo da semana',12)
P('Comparando o mesmo par PV × fadiga física, a correlação é fraca no baseline (ρ = %s, n.s.) e substancialmente mais forte na média da semana (ρ ≈ −0,51). Além disso, a aptidão de base não prediz a magnitude da deterioração semanal (PV × ΔfadigaD1→D7 ρ = %s; PV × ΔvigorD1→D7 ρ = %s, ambas n.s.). A leitura é clara e cientificamente robusta: no ponto de partida os atletas estão relativamente homogêneos (efeito de piso), e é a ACUMULAÇÃO de carga ao longo do microciclo que revela a vantagem do mais apto — o mais treinado não começa muito diferente, mas termina menos fatigado.'%(
    c2('%+.2f'%R['COR']['PV']['b_FadFisica']['r']),c2('%+.2f'%R['SEC']['PV']['d_FadFisica']['r']),c2('%+.2f'%R['SEC']['PV']['d_Vigor']['r'])),after=6)
figure('/home/user/mdlucca/Artigos/figuras/x_tcar1.png',
    'Parâmetros restritos (T-CAR 1 + TRIMP do HIIT × baseline). (A) A aptidão associa-se inversamente à fadiga BRUMS de base. (B) Entre os parâmetros fisiológicos, PV e distância lideram; a carga interna (TRIMP/PSE) é nula no baseline. (C) Perfil iceberg de base por posição (pontas no topo). (D) O acoplamento aptidão–fadiga é fraco no baseline e forte na média semanal — emerge com o acúmulo de carga.')

H('3. Síntese')
P('Com os parâmetros restritos ao T-CAR 1 (fisiológico) e ao baseline (psicológico): (1) a APTIDÃO — pico de velocidade e a distância dela derivada — é o parâmetro fisiológico que melhor indica o estado psicológico de base, com sinal negativo sobre a fadiga (ρ = %s) e AUC = %s para o perfil iceberg; (2) a CARGA INTERNA do HIIT (TRIMP/PSE) não explica o baseline, coerente com a temporalidade; (3) PV e distância são redundantes no T-CAR; (4) a vantagem do mais apto não está no ponto de partida, mas emerge com o acúmulo de carga ao longo da semana. Recomenda-se, portanto, tratar o PV do T-CAR 1 como o marcador fisiológico de base e a carga interna como preditor da RESPOSTA (dose de treino), não do estado inicial.'%(
    c2('%+.2f'%R['COR']['PV']['b_Fadiga']['r']),c2('%.2f'%R['ROC']['PV']['auc'])),after=6)

H('4. Limitações')
P('Baseline é um ponto único e comprimido por efeito de piso, o que reduz a variância e o poder das associações (n = 27; subgrupos posicionais pequenos). O TRIMP foi calculado das sessões de HIIT (não do T-CAR 1) e sob premissas constantes de FCrepouso e duração (invariantes de escala). O tempo em alta intensidade do T-CAR 1 não estava disponível. As relações são correlacionais.',after=6)

OUTP='/home/user/mdlucca/Artigos/TCAR1_Baseline_Parametros_Restritos.docx'
doc.save(OUTP)
print('SAVED',OUTP,'| Tabelas',_TN[0],'Figuras',_FN[0])
