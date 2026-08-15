import warnings; warnings.filterwarnings('ignore')
import numpy as np, pandas as pd, json, os
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
SC='/tmp/claude-0/-home-user-mdlucca/e1dba24c-b1d7-5908-9106-f2f4aaf3f56a/scratchpad'
FIG='/home/user/mdlucca/Artigos/figuras'
R=json.load(open(f'{SC}/semana.json')); APP=json.load(open(f'{SC}/app_data.json'))
REG=json.load(open(f'{SC}/reg.json')); BF=json.load(open(f'{SC}/baseline_final.json'))
desc=R['desc']; agu=R['agu']; cron=R['cron']; dec=R['dec']; rank=R['rank']
CORE=[('Vigor','Vigor'),('Fadiga','Fadiga (BRUMS)'),('FadFisica','Fadiga física'),('FadMental','Fadiga mental')]
def c2(x): return str(x).replace('.',',')

doc=Document(); s=doc.sections[0]; s.left_margin=Cm(3); s.top_margin=Cm(3); s.right_margin=Cm(2); s.bottom_margin=Cm(2)
n=doc.styles['Normal']; n.font.name='Times New Roman'; n.font.size=Pt(12); n.paragraph_format.line_spacing=1.5; n.paragraph_format.space_after=Pt(0)
FN='Times New Roman'
def _set(r,sz=12,b=False,it=False): r.font.name=FN; r.font.size=Pt(sz); r.bold=b; r.italic=it
def sec(num,txt):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(4); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    _set(p.add_run((f'{num} {txt}').upper() if '.' not in num else f'{num} {txt}'),12,b=True)
def sub(num,txt):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(2); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    _set(p.add_run(f'{num} {txt}'),12,b=True)
def P(t,indent=True,size=12,it=False):
    p=doc.add_paragraph(); _set(p.add_run(t),size,it=it); p.paragraph_format.line_spacing=1.5
    if indent: p.paragraph_format.first_line_indent=Cm(1.25)
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
def _img(name,label,num,titulo,w,fonte):
    cap=doc.add_paragraph(); cap.alignment=WD_ALIGN_PARAGRAPH.CENTER; cap.paragraph_format.space_before=Pt(6); cap.paragraph_format.line_spacing=1.0
    _set(cap.add_run(f'{label} {num} – {titulo}'),10)
    p=f'{FIG}/{name}.png'
    if os.path.exists(p): doc.add_picture(p,width=Inches(w)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    else: print('MISS',name)
    fp=doc.add_paragraph(); fp.alignment=WD_ALIGN_PARAGRAPH.CENTER; fp.paragraph_format.space_after=Pt(6); fp.paragraph_format.line_spacing=1.0
    _set(fp.add_run(f'Fonte: {fonte}'),10)
_GN=[0]; _FN=[0]
def graf(name,titulo,w=5.0,fonte='Elaborado pelos autores (2026).'): _GN[0]+=1; _img(name,'Gráfico',_GN[0],titulo,w,fonte)
def fig(name,titulo,w=5.0,fonte='Elaborado pelos autores (2026).'): _FN[0]+=1; _img(name,'Figura',_FN[0],titulo,w,fonte)
def _bd(t):
    tblPr=t._tbl.tblPr; b=OxmlElement('w:tblBorders')
    for e in ('top','bottom','insideH'):
        x=OxmlElement(f'w:{e}'); x.set(qn('w:val'),'single'); x.set(qn('w:sz'),'6'); x.set(qn('w:space'),'0'); x.set(qn('w:color'),'000000'); b.append(x)
    for e in ('left','right','insideV'):
        x=OxmlElement(f'w:{e}'); x.set(qn('w:val'),'none'); b.append(x)
    tblPr.append(b)
_TN=[0]
def table(titulo,headers,rows,fonte='Dados da pesquisa (2026).',fs=9.5):
    _TN[0]+=1
    cap=doc.add_paragraph(); cap.alignment=WD_ALIGN_PARAGRAPH.LEFT; cap.paragraph_format.space_before=Pt(6); cap.paragraph_format.line_spacing=1.0
    _set(cap.add_run(f'Tabela {_TN[0]} – {titulo}'),10)
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; _bd(t)
    for i,h in enumerate(headers):
        cl=t.rows[0].cells[i]; cl.paragraphs[0].line_spacing=1.0; _set(cl.paragraphs[0].add_run(str(h)),fs,b=True)
    for row in rows:
        cs=t.add_row().cells
        for i,v in enumerate(row):
            cs[i].paragraphs[0].line_spacing=1.0; _set(cs[i].paragraphs[0].add_run(str(v)),fs)
    fp=doc.add_paragraph(); fp.alignment=WD_ALIGN_PARAGRAPH.LEFT; fp.paragraph_format.space_after=Pt(6); fp.paragraph_format.line_spacing=1.0
    _set(fp.add_run(f'Fonte: {fonte}'),10)

# TÍTULO + RESUMO
tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
_set(tp.add_run('PERFIL DE HUMOR DE ATLETAS DE HANDEBOL NA ÚLTIMA SEMANA DA PRÉ-TEMPORADA COMPETITIVA'),14,b=True)
a2=doc.add_paragraph(); a2.alignment=WD_ALIGN_PARAGRAPH.CENTER; _set(a2.add_run('[Autoria e afiliação a preencher]'),11,it=True)
rp=doc.add_paragraph(); _set(rp.add_run('RESUMO'),12,b=True); rp.paragraph_format.line_spacing=1.0
P('Objetivo: avaliar e classificar o perfil de humor de atletas de handebol na última semana da pré-temporada competitiva. Método: 27 atletas responderam ao BRUMS-24 e a autoavaliações de fadiga física e mental duas vezes ao dia — a primeira coleta (pré) e a última (pós) — durante sete dias (456 observações), com análise descritiva, comparação do grupo entre a linha de base (Dia 1, pré) e o último dia (Dia 7, pós), regressão linear, derivadas, decomposição da variância e pontos de inflexão. Resultados: o humor deteriorou-se de forma concentrada no eixo energia–fadiga — da linha de base ao final, o vigor caiu 47% (dz = −1,02) e a fadiga física (+129%; dz = +2,12) e a fadiga do BRUMS (+116%; dz = +0,78) aumentaram, sendo a fadiga física a mais responsiva; a trajetória inflectiu-se em torno do Dia 4 e o perfil iceberg reduziu-se ao longo da semana. Conclusão: o custo do período é somático e acumulativo, sendo a fadiga física e o vigor os marcadores mais sensíveis para o monitoramento individualizado.',indent=False)
kw=doc.add_paragraph(); _set(kw.add_run('Palavras-chave: '),12,b=True); _set(kw.add_run('Estado de humor. BRUMS. Fadiga. Pré-temporada. Handebol.'),12); kw.paragraph_format.line_spacing=1.5

# 1 INTRODUÇÃO
sec('1','Introdução')
P('O monitoramento dos estados de humor é uma ferramenta prática, de baixo custo e não invasiva para acompanhar a resposta psicobiológica de atletas à carga de treinamento, sinalizando desadaptações antes que se manifestem em queda de desempenho (TERRY; LANE; FOGARTY, 2003; ROHLFS et al., 2008). A Brunel Mood Scale (BRUMS) operacionaliza o construto em seis dimensões — tensão, depressão, raiva, vigor, fadiga e confusão. Atletas saudáveis tendem a exibir o perfil iceberg, com vigor elevado destacando-se acima das dimensões negativas (MORGAN, 1985); o achatamento desse perfil indica fadiga acumulada.')
P('Este estudo tem por objetivo avaliar e classificar o perfil de humor de atletas de handebol durante a última semana da pré-temporada competitiva, descrevendo o comportamento geral do grupo por variável, a variação da linha de base (Dia 1, pré) ao último dia (Dia 7, pós), a resposta pré→pós e a sensibilidade relativa das variáveis, com aprofundamento por regressão, derivadas e pontos de inflexão.')

# 2 JUSTIFICATIVA
sec('2','Justificativa')
P('A última semana pré-competitiva é um período crítico, no qual a comissão técnica busca dissipar a fadiga acumulada preservando as adaptações, de modo a maximizar a prontidão para competir. O manejo inadequado da carga pode deteriorar o humor e sinalizar recuperação incompleta; ainda assim, o comportamento do humor nessa janela raramente é caracterizado de forma estruturada. Descrevê-lo quantitativamente — identificando quais variáveis são mais sensíveis e em que momento a trajetória se inflete — oferece subsídios objetivos para individualizar a recuperação.')

# 3 MÉTODO
sec('3','Método')
ag=APP['socio']['agg']; pc=APP['socio']['pos_counts']
P('Estudo descritivo de medidas repetidas, em condições ecológicas de treinamento, com 27 atletas de handebol do sexo masculino acompanhados por sete dias, cada um funcionando como sua própria referência (dados anonimizados, A01–A27). O humor foi avaliado pela BRUMS-24 (seis subescalas; a Perturbação Total do Humor, PTH = negativas − vigor) e a fadiga física e mental por escalas de 0 a 10, autoaplicadas por formulário eletrônico com carimbo de data/hora duas vezes ao dia (a primeira resposta tomada como “pré” e a última como “pós”), totalizando 456 respostas válidas. O processamento e a auditoria interna foram inteiramente roteirizados em Python (Figura 1).')
P('Computaram-se estatísticas descritivas com limites (amplitude); a variação da linha de base (Dia 1, pré) ao último dia (Dia 7, pós) e a resposta pré→pós foram avaliadas por Wilcoxon pareado com tamanho de efeito (dz); a tendência semanal por regressão linear (coeficiente por dia e R²). A forma da trajetória foi caracterizada por ajuste cúbico das médias diárias (velocidade, aceleração e ponto de inflexão) e por suavização LOWESS; a variância foi decomposta em sinal da semana, traço e ruído, derivando-se a relação sinal-ruído (SNR). Adotou-se α = 0,05.')
fig('fig_fluxograma','Fluxograma do processamento e da análise dos dados em Python, com auditoria interna por script',w=5.4)

# 4 RESULTADOS
sec('4','Resultados')
sub('4.1','Caracterização da amostra')
table('Caracterização sociodemográfica e antropométrica da amostra (n = 27)',['Variável','Média ± DP (amplitude)'],[
 ['Idade (anos)',f"{ag['idade'][0]:.0f} ± {ag['idade'][1]:.0f} ({ag['idade'][2]}–{ag['idade'][3]})"],
 ['Experiência (anos)',c2(f"{ag['exp'][0]:.1f} ± {ag['exp'][1]:.1f}")],
 ['Estatura (cm)',c2(f"{ag['estatura'][0]:.1f} ± {ag['estatura'][1]:.1f}")],
 ['Massa corporal (kg)',c2(f"{ag['massa'][0]:.1f} ± {ag['massa'][1]:.1f}")],
 ['Gordura corporal (%)',c2(f"{ag['pG'][0]:.1f} ± {ag['pG'][1]:.1f}")],
])
P(f"A amostra, homogênea, distribuiu-se por posição em {pc.get('Armador',0)} armadores, {pc.get('Ala',0)} alas, {pc.get('Pivô',0)} pivôs e {pc.get('Goleiro',0)} goleiros.")

sub('4.2','Estatística descritiva, distribuição e tendência')
rows=[[lab,c2(f"{desc[v]['mean']:.2f} ± {desc[v]['sd']:.2f}"),c2(f"{desc[v]['med']:.1f} [{desc[v]['iqr']:.1f}]"),f"{desc[v]['mn']:.0f}–{desc[v]['mx']:.0f}",c2(f"{desc[v]['amp']:.0f}")] for v,lab in [('Vigor','Vigor'),('Fadiga','Fadiga (BRUMS)'),('FadFisica','Fadiga física'),('FadMental','Fadiga mental'),('TMD','PTH/TMD')]]
table('Estatística descritiva e limites (amplitude) na semana',['Variável','M ± DP','Mediana [IIQ]','Mín.–Máx.','Amplitude'],rows)
rows=[[lab,c2(f"{REG[v]['slope']:+.3f}"),c2(f"{REG[v]['r2']:.3f}"),(c2(f"{REG[v]['p']:.4f}") if REG[v]['p']>=0.001 else '<0,001')] for v,lab in CORE]
table('Regressão linear do humor ao longo da semana (coeficiente por dia, R² e p)',['Variável','Coef./dia (b)','R²','p'],rows)
P('No conjunto da semana, o grupo preservou, em média, o perfil iceberg, com o vigor acima das subescalas negativas, mantidas próximas ao piso (Figura 2, radar); a fadiga foi a dimensão negativa mais expressiva, e a ampla amplitude da PTH revela heterogeneidade interindividual. Os histogramas (Gráfico 1) mostram o vigor e a fadiga física bem dispersos e as negativas concentradas no piso, e os diagramas de caixa por dia (Gráfico 2) evidenciam o deslocamento progressivo do vigor (para baixo) e das fadigas (para cima). A regressão linear (Tabela 3; Gráfico 3) confirma as tendências: a fadiga física apresenta a inclinação mais acentuada e o melhor ajuste (b = +0,39/dia; R² = 0,10; p < 0,001), seguida da fadiga do BRUMS e do vigor; a fadiga mental não exibe tendência (p = 0,60). Os R² baixos decorrem de a regressão ser computada no nível da observação, em que a variância é dominada por diferenças entre atletas e ruído — o que recomenda interpretar a tendência do grupo, e não a leitura isolada de um atleta.')
fig('sem_f_radar','Forma do perfil de humor (escores z) no Dia 1 e no Dia 7',w=4.6)
graf('sem_f_hist','Histogramas de distribuição das quatro variáveis na semana (linha tracejada = média)')
graf('sem_f_boxplot','Diagramas de caixa (boxplots) das quatro variáveis por dia')
graf('sem_f_regressao','Regressão linear de cada variável ao longo da semana (pontos = observações; reta = ajuste)')

sub('4.3','Comportamento e deterioração ao longo da semana')
P('As trajetórias (Gráfico 4) mostram que o sinal temporal se concentra no eixo energia–fadiga: o vigor declina e a fadiga — física e do BRUMS — eleva-se progressivamente rumo ao Dia 7, ao passo que a fadiga mental e as subescalas negativas permanecem estáveis. A decomposição da variância indica que o sinal da semana é a menor fração (fadiga física 13%, vigor 7%, fadiga do BRUMS 4%, fadiga mental 1%), com relação sinal-ruído favorável apenas à fadiga física (1,82) e ao vigor (1,52); a maior parte da variância é traço e ruído, o que sustenta o monitoramento por médias de coletas e por tendência. O Gráfico 5 sintetiza, em escores padronizados, a deterioração conjunta do eixo, e o mapa de calor por atleta (Figura 3) revela a heterogeneidade: alguns atletas oscilam pouco e outros amplificam em todas as variáveis.')
graf('sem_f_trajetoria','Comportamento semanal das quatro variáveis (M ± DP; suavização LOWESS pontilhada)')
graf('sem_f_deterioracao','Deterioração do eixo energia–fadiga em escores padronizados (z): vigor declina, fadigas ascendem',w=4.8)
fig('abnt_f_heatmap','Mapa de calor da amplitude semanal por atleta (escores z por variável)',w=4.8)

sub('4.4','Comparação geral do grupo: linha de base (Dia 1, pré) versus último dia (Dia 7, pós)')
rows=[[lab,c2(f"{BF[v]['base']:.2f} ± {BF[v]['base_sd']:.2f}"),c2(f"{BF[v]['fin']:.2f} ± {BF[v]['fin_sd']:.2f}"),c2(f"{BF[v]['delta']:+.2f}"),f"{BF[v]['pct']:+.0f}%",c2(f"{BF[v]['dz']:+.2f}"),(c2(f"{BF[v]['p']:.3f}") if BF[v]['p']>=0.001 else '<0,001')] for v,lab in CORE]
table('Humor e fadiga na linha de base (Dia 1, pré) e no último dia (Dia 7, pós); comparação pareada (Wilcoxon)',['Variável','Pré D1','Pós D7','Δ','% mud.','dz','p'],rows)
rows=[[r['lab'],c2(f"{r['dz_ag']:.2f}"),c2(f"{r['dz_cr']:.2f}"),c2(f"{r['snr']:.2f}"),f"{r['pday']:.0f}%"] for r in sorted(rank,key=lambda x:-x['dz_cr'])]
table('Sensibilidade/responsividade das variáveis (efeito agudo, crônico, SNR e % de sinal)',['Variável','|dz| agudo','|dz| crônico','SNR','% sinal'],rows)
P('A comparação geral do grupo entre a linha de base e o último dia confirma e amplia o quadro: o vigor caiu 47% (8,25 → 4,38; dz = −1,02) e a fadiga física quase triplicou em magnitude de efeito (3,69 → 8,44; +129%; dz = +2,12), com a fadiga do BRUMS (+116%; dz = +0,78) e a fadiga mental (+39%; dz = +0,82) também aumentando significativamente. A hierarquia de sensibilidade (Tabela 5) é clara — fadiga física > vigor > fadiga do BRUMS > fadiga mental —, com a fadiga física reunindo o maior efeito e a melhor SNR. O efeito crônico (linha de base ao final) supera o agudo pré→pós (fadiga física dz crônico = +2,12 versus agudo = +1,08), indicando acúmulo progressivo, e não apenas resposta pontual ao esforço diário.')

sub('4.5','Perfil iceberg e ponto de inflexão')
P('Os seis perfis de humor (Figura 4) descrevem as configurações possíveis do estado afetivo; ao longo da semana observou-se a migração do perfil saudável (iceberg) para o de fadiga, com queda do percentual em perfil iceberg e do índice-iceberg e precipitação no Dia 7 (Gráfico 6). A leitura das derivadas (Figura 5) descreve a forma da mudança: a aceleração troca de sinal uma única vez — o ponto de inflexão —, situado em torno do Dia 4 para o vigor, a fadiga física e a fadiga do BRUMS, marcando a transição da fase de acomodação para o aprofundamento da fadiga rumo ao Dia 7 e delimitando a janela oportuna para a intervenção na recuperação.')
graf('sem_f_iceberg','Perfil iceberg ao longo da semana: percentual em perfil iceberg (barras), índice-iceberg (linha) e regressão',w=5.4)
fig('fig_perfis_esquema','Esquema dos seis perfis de humor (escores z das subescalas; vigor em verde, negativas em vermelho)',w=5.2)
fig('sem_f_derivadas','Velocidade (linha cheia) e aceleração (tracejada) da trajetória de cada variável; linha vertical = ponto de inflexão')

sub('4.6','Comportamento individual entre atletas')
P('No nível individual, a heterogeneidade é marcante. O mapa de calor da amplitude por atleta (Figura 3) já mostrara que alguns atletas oscilam pouco e outros amplificam em todas as variáveis. As trajetórias individuais ao longo da semana (Figura 6) confirmam que, embora a direção do declínio do vigor e da elevação da fadiga física seja compartilhada pela maioria, a magnitude e o ponto de partida variam substancialmente entre atletas, de modo que escores idênticos podem representar estados distintos conforme a referência individual. A resposta aguda pré→pós por atleta (Figura 7) reproduz esse quadro no nível intradiário: a maioria reduz o vigor e aumenta a fadiga física, mas em amplitudes heterogêneas. Em conjunto, esses achados desaconselham o uso de pontos de corte normativos únicos e sustentam o referenciamento de cada atleta à sua própria linha de base.')
fig('sem_f_individual','Trajetórias individuais do vigor e da fadiga física ao longo da semana (linha em destaque = média do grupo)',w=5.4)
fig('sem_f_ind_prepos','Resposta individual pré→pós na semana (cada linha = um atleta; linha em destaque = média do grupo)',w=5.4)

# 5 CONSIDERAÇÕES
sec('5','Considerações Finais')
P('O perfil de humor deteriorou-se de forma coerente e concentrada no eixo energia–fadiga, tanto ao longo da semana quanto na comparação da linha de base ao último dia. A fadiga física foi a variável mais sensível e responsiva, seguida do vigor; a fadiga mental e as subescalas negativas foram pouco responsivas. O efeito crônico superou o agudo, indicando acúmulo, e as trajetórias inflectiram-se em torno do Dia 4. Recomenda-se acompanhar a fadiga física e o vigor por tendência individual e por médias de coletas, com atenção reforçada a partir do meio da semana. Como limitações, trata-se de estudo descritivo, sem grupo de comparação e restrito a um único período de uma equipe masculina.')

sec('REFERÊNCIAS','')
for r in [
 'MORGAN, W. P. Selected psychological factors limiting performance: a mental health model. In: CLARKE, D. H.; ECKERT, H. M. (Ed.). Limits of human performance. Champaign: Human Kinetics, 1985. p. 70-80.',
 'ROHLFS, I. C. P. M. et al. A Escala de Humor de Brunel (Brums): instrumento para detecção precoce da síndrome do excesso de treinamento. Revista Brasileira de Medicina do Esporte, v. 14, n. 3, p. 176-181, 2008.',
 'TERRY, P. C.; LANE, A. M.; FOGARTY, G. J. Construct validity of the Profile of Mood States — Adolescents for use with adults. Psychology of Sport and Exercise, v. 4, n. 2, p. 125-139, 2003.',
]:
    p=doc.add_paragraph(); _set(p.add_run(r),12); p.paragraph_format.line_spacing=1.0; p.paragraph_format.space_after=Pt(6); p.alignment=WD_ALIGN_PARAGRAPH.LEFT

OUTP='/home/user/mdlucca/Artigos/Perfil_Humor_Semana_PreTemporada_ABNT.docx'
doc.save(OUTP); print('SAVED',OUTP)
d2=Document(OUTP); print('tables',len(d2.tables),'imgs',len(d2.inline_shapes),'| Gráficos',_GN[0],'Figuras',_FN[0],'Tabelas',_TN[0])
