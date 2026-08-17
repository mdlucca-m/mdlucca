import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
OUT='/home/user/mdlucca/Artigos/submissao'
os.makedirs(OUT, exist_ok=True)

def letter(fname, journal, editor_line, body_paras):
    doc=Document()
    for s in doc.sections:
        s.top_margin=s.bottom_margin=Cm(2.5); s.left_margin=s.right_margin=Cm(2.5)
    st=doc.styles['Normal']; st.font.name='Times New Roman'; st.font.size=Pt(11)
    def P(t,before=0,after=8,align=WD_ALIGN_PARAGRAPH.JUSTIFY,bold=False):
        p=doc.add_paragraph(); r=p.add_run(t); r.bold=bold; r.font.size=Pt(11); r.font.name='Times New Roman'
        p.alignment=align; p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after); p.paragraph_format.line_spacing=1.15
        return p
    P('[Corresponding author name]\n[Affiliation]\n[Address]\n[E-mail]',after=2,align=WD_ALIGN_PARAGRAPH.LEFT)
    P('[Date]',after=10,align=WD_ALIGN_PARAGRAPH.LEFT)
    P(editor_line,after=2,align=WD_ALIGN_PARAGRAPH.LEFT,bold=True)
    P('Dear Editor,',after=8,align=WD_ALIGN_PARAGRAPH.LEFT)
    for para in body_paras: P(para)
    P('Sincerely,',before=6,after=2,align=WD_ALIGN_PARAGRAPH.LEFT)
    P('[Corresponding author name], on behalf of all authors',after=0,align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.save(os.path.join(OUT,fname)); print('SAVED',fname)

# ---------------- Paper 1 — Sports (MDPI) ----------------
letter('CoverLetter_Paper1_Sports.docx','Sports (MDPI)',
 'To the Editor-in-Chief, Sports (MDPI)',
 ['We are pleased to submit our manuscript entitled "Mood dynamics across a pre-competitive microcycle in '
  'elite handball: the energy-fatigue axis, profile migration, and trajectory modelling" for consideration as an '
  'Article in Sports.',
  'Self-report mood monitoring with the Brunel Mood Scale (BRUMS) is well established in performance sport, yet '
  'fine-grained descriptions of how mood behaves within a single microcycle remain scarce. In 27 elite male '
  'handball players monitored twice daily across a seven-day pre-season microcycle, we characterise the mood '
  'response and, moving beyond descriptive snapshots, model the daily trajectories: we quantify the rate and '
  'acceleration of change, locate the vigour-fatigue crossover, decompose each trajectory into signal and noise, '
  'and test the migration from the iceberg to the shark-fin profile using logistic trend analysis. The '
  'deterioration concentrated on the energy-fatigue axis and matched a functional-overreaching signature - an '
  'expected, reversible fatigue - with no emergence of at-risk mood profiles. All central findings were robust to '
  'a noise-filtering audit.',
  'We believe the manuscript fits the scope of Sports and complements the mood-profiling literature published in '
  'the journal, including the six-cluster taxonomy and its applications in Brazilian athletes.',
  'Companion-paper disclosure: this manuscript is one of two companion reports that draw on the same cohort and '
  'microcycle but address distinct primary outcomes. The companion paper - on the physiological and wellness '
  'correlates of the mood response (aerobic fitness, anaerobic capacity, sleep, and perceived stress) - is being '
  'submitted to Biology of Sport. The two papers cross-cite each other. We disclose this to ensure full '
  'transparency and to avoid any appearance of redundant publication.',
  'We confirm that this work is original, has not been published previously, and is not under consideration by '
  'any other journal. All authors have approved the submission and declare no conflicts of interest. Athlete data '
  'were fully anonymised, and the study was conducted in accordance with the Declaration of Helsinki with '
  'informed consent. Analysis scripts are available to support reproducibility.',
  'Thank you for considering our manuscript. We look forward to your response.'])

# ---------------- Paper 2 — Biology of Sport ----------------
letter('CoverLetter_Paper2_BiologyOfSport.docx','Biology of Sport',
 'To the Editor-in-Chief, Biology of Sport',
 ['We are pleased to submit our manuscript entitled "Physiological and wellness correlates of the mood response '
  'across a pre-competitive microcycle in elite handball: aerobic fitness, anaerobic capacity, sleep, and '
  'perceived stress" for consideration in Biology of Sport.',
  'In 27 elite male handball players monitored across a seven-day pre-season microcycle, we examined which '
  'physiological and wellness parameters track the magnitude of the mood response. Intermittent aerobic fitness '
  '(peak velocity in the Carminatti field test) emerged as the strongest physiological correlate of perceived '
  'fatigue and vigour, and we derived single and dual peak-velocity thresholds that flag athletes at higher risk '
  'of high-fatigue and low-vigour days. In contrast, repeated-sprint (anaerobic) ability did not track the mood '
  'response, and its apparent association reflected body mass, which we show by partial and allometric adjustment '
  '- a specificity finding. Daytime sleepiness followed the training load and tracked fatigue, whereas perceived '
  'stress remained stable, and the magnitude of the acute and chronic mood effects was largely independent of the '
  'individual profile.',
  'The study fits the scope of Biology of Sport, combining field-based physiological testing (including the '
  'Carminatti test) with athlete monitoring in an invasion team sport, and speaks directly to the multidomain '
  'monitoring framework.',
  'Companion-paper disclosure: this manuscript is one of two companion reports drawing on the same cohort and '
  'microcycle but with distinct primary outcomes. The companion paper - on the within-microcycle mood dynamics '
  'and trajectory modelling - is being submitted to Sports (MDPI). The two papers cross-cite each other. We '
  'disclose this for full transparency and to avoid any appearance of redundant publication.',
  'We confirm that this work is original, has not been published previously, and is not under consideration '
  'elsewhere. All authors approved the submission and declare no conflicts of interest. Athlete data were fully '
  'anonymised, and the study followed the Declaration of Helsinki with informed consent.',
  'Thank you for considering our manuscript.'])
print('OK')
