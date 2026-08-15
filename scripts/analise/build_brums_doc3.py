# -*- coding: utf-8 -*-
import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
R=json.load(open('brums_desc2.json')); RC=json.load(open('rci6.json'))
FG='/home/user/mdlucca/Artigos/figuras'
def c2(s): return str(s).replace('.',',')
doc=Document()
st=doc.styles['Normal']; st.font.name='Times New Roman'; st.font.size=Pt(12)
st.element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')
st.paragraph_format.line_spacing=1.5; st.paragraph_format.space_after=Pt(0)
s=doc.sections[0]; s.top_margin=Cm(3); s.left_margin=Cm(3); s.bottom_margin=Cm(2); s.right_margin=Cm(2)
_TN=[0]; _FN=[0]
def P(t='',just=True,size=12,after=6,bold=False,ind=True):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(size); r.bold=bold
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(after)
    if just: p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if just and ind: p.paragraph_format.first_line_indent=Cm(1.25)
    return p
def RUN(pairs,after=6,ind=True):  # mixed bold/normal runs in one justified paragraph
    p=doc.add_paragraph(); p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(after)
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if ind: p.paragraph_format.first_line_indent=Cm(1.25)
    for txt,bold in pairs:
        r=p.add_run(txt); r.bold=bold; r.font.size=Pt(12)
    return p
def H(t,size=12,before=10):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(size)
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.5
def _bd(c):
    tcPr=c._tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for e_ in ['top','bottom','left','right']:
        el=OxmlElement(f'w:{e_}'); on=e_ in ('top','bottom')
        el.set(qn('w:val'),'single' if on else 'nil'); el.set(qn('w:sz'),'6'); el.set(qn('w:color'),'000000'); b.append(el)
    tcPr.append(b)
def table(cap,header,rows,fonte='Fonte: dados da pesquisa (2026).',fs=9):
    _TN[0]+=1
    p=doc.add_paragraph(); r=p.add_run('Tabela %d – %s'%(_TN[0],cap)); r.font.size=Pt(11)
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
    t=doc.add_table(rows=1,cols=len(header)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,htx in enumerate(header):
        cc=t.rows[0].cells[i]; cc.text=''; rr=cc.paragraphs[0].add_run(htx); rr.bold=True; rr.font.size=Pt(fs); rr.font.name='Times New Roman'
        cc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; _bd(cc)
    for row in rows:
        cs=t.add_row().cells
        for i,val in enumerate(row):
            cs[i].text=''; rr=cs[i].paragraphs[0].add_run(str(val)); rr.font.size=Pt(fs); rr.font.name='Times New Roman'
            cs[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if i>0 else WD_ALIGN_PARAGRAPH.LEFT; _bd(cs[i])
    pf=doc.add_paragraph(); rf=pf.add_run(fonte); rf.font.size=Pt(9); pf.paragraph_format.space_after=Pt(6)
    return _TN[0]
def figure(path,cap,w=15.5):
    _FN[0]+=1
    pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; pp.add_run().add_picture(path,width=Cm(w)); pp.paragraph_format.space_before=Pt(6)
    pc=doc.add_paragraph(); pc.alignment=WD_ALIGN_PARAGRAPH.CENTER; rc=pc.add_run('Figura %d – %s'%(_FN[0],cap)); rc.font.size=Pt(11)
    pf=doc.add_paragraph(); pf.alignment=WD_ALIGN_PARAGRAPH.CENTER; rf=pf.add_run('Fonte: elaboração dos autores (2026).'); rf.font.size=Pt(9); pf.paragraph_format.space_after=Pt(6)
    return _FN[0]

pr=R['prepos']; d17=R['d1d7']; prof=R['profiles']; sm=R['sample']; foc=R['focus']; npk=R['neg_peak']; byd=R['byday']; desc=R['desc']
# logical order: energy-fatigue axis first, then PTH, then negatives
ORDER=[('Vigor','Vigor'),('Fadiga','Fadiga'),('TMD','Perturbação Total do Humor (PTH)'),('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Confusao','Confusão')]

# ===================== TÍTULO + RESUMO =====================
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('PERFIL DE HUMOR DE ATLETAS DE HANDEBOL DE ELITE NA ÚLTIMA SEMANA DE PRÉ-TEMPORADA: ANÁLISE DESCRITIVA, COMPARATIVA E DE CORRELAÇÃO DAS SEIS SUBESCALAS DO BRUMS')
r.bold=True; r.font.size=Pt(13); p.paragraph_format.space_after=Pt(10)
H('RESUMO',before=2)
RUN([('Objetivo: ',True),('avaliar o perfil de humor de atletas de handebol de elite na última semana de pré-temporada, '
 'caracterizando e comparando o comportamento das seis subescalas do BRUMS no nível do grupo e do sujeito. ',False),
 ('Método: ',True),('estudo descritivo-comparativo de medidas repetidas com %d atletas e %d observações do BRUMS-24 em '
 'sete dias (duas coletas/dia: pré e pós-treino). Aplicaram-se estatística descritiva exploratória, comparações pré→pós e '
 'Dia 1→Dia 7 (Wilcoxon; variação percentual, dz de Cohen com IC95%%, magnitude), correlação de Spearman e Índice de '
 'Mudança Confiável (RCI) individual. '%(sm['n'],sm['n_obs']),False),
 ('Resultados: ',True),('a deterioração concentrou-se no eixo energia–fadiga — vigor pré→pós −%s%% (dz = %s) e Dia 1→Dia 7 '
 'com efeito grande (dz = %s; p = %s); fadiga pré→pós +%s%% (dz = %s). No nível individual, %d%% dos atletas reduziram o '
 'vigor e %d%% aumentaram a fadiga de forma confiável. As subescalas negativas mantiveram-se próximas do piso, com '
 'acoplamento à fadiga fraco no dia de maior vigor (D1) e forte no de maior fadiga (D7: depressão ρ = %s; raiva ρ = %s). O '
 'perfil iceberg caiu de %d%% para %d%% e o de fadiga subiu de %d%% para %d%%. '%(
   c2('%.0f'%abs(pr['Vigor']['pct'])),c2('%.2f'%pr['Vigor']['dz']),c2('%.2f'%d17['Vigor']['dz']),c2('%.3f'%d17['Vigor']['p']),
   c2('%.0f'%pr['Fadiga']['pct']),c2('%.2f'%pr['Fadiga']['dz']),RC['Vigor']['pct_piora'],RC['Fadiga']['pct_piora'],
   c2('%+.2f'%foc['D7']['Depressao']['rho_fad']),c2('%+.2f'%foc['D7']['Raiva']['rho_fad']),
   int(prof['D1']['Iceberg']),int(prof['D7']['Iceberg']),int(prof['D1']['Barbatana tubarão']),int(prof['D7']['Barbatana tubarão'])),False),
 ('Conclusão: ',True),('o humor migra da prontidão para a fadiga funcional; a fadiga é a subescala mais sensível e o afeto '
 'negativo consolida-se com a fadiga apenas sob carga acumulada, o que fundamenta um monitoramento individualizado.',False)],after=6)
P('Palavras-chave: estados de humor; BRUMS; handebol; monitoramento; fadiga; atletas de elite.',size=11,after=8,ind=False)

# ===================== 1 INTRODUÇÃO =====================
H('1 INTRODUÇÃO')
P('O monitoramento dos estados psicológicos consolidou-se como componente essencial da gestão da carga de treino no '
 'esporte de rendimento. Escalas de autorrelato são práticas, econômicas e sensíveis: em revisão sistemática, as medidas '
 'subjetivas de bem-estar rastrearam as cargas aguda e crônica com sensibilidade e consistência superiores às de '
 'marcadores objetivos, respondendo à intensificação do treino e à redução de carga de modo coerente (SAW; MAIN; GASTIN, '
 '2016). Por isso, painéis de consenso recomendam o uso rotineiro de instrumentos subjetivos para monitorar a fadiga e '
 'orientar decisões de treino e recuperação em modalidades esportivas (THORPE et al., 2017; KELLMANN et al., 2018).')
P('Entre esses instrumentos, a Brunel Mood Scale (BRUMS) — versão abreviada e adaptada do Profile of Mood States — é uma '
 'das mais utilizadas na ciência do esporte. Composta por 24 adjetivos organizados em seis subescalas (tensão, depressão, '
 'raiva, vigor, fadiga e confusão), foi desenvolvida para aplicação rápida e repetida e possui propriedades psicométricas '
 'sólidas, com sucessivas validações transculturais que sustentam seu uso em diferentes populações e idiomas (TERRY et '
 'al., 1999; TERRY; LANE; FOGARTY, 2003; ROHLFS et al., 2008; LEW et al., 2023). Sua leitura pode ir além dos escores '
 'isolados: Morgan (1985) descreveu o “perfil iceberg” — vigor elevado sobre dimensões negativas baixas — como assinatura '
 'de prontidão e saúde mental positiva, e abordagens recentes formalizaram seis perfis prototípicos de humor (iceberg, '
 'Everest invertido, iceberg invertido, submerso, barbatana de tubarão e superfície), hoje amplamente empregados no '
 'rastreamento do estado psicológico e da prontidão de atletas de elite (PARSONS-SMITH; TERRY; MACHIN, 2017; TERRY et '
 'al., 2021; TAN et al., 2024). O valor aplicado do constructo é reforçado por metanálise que confirma a associação entre '
 'a perturbação do humor e o desempenho esportivo (LOCHBAUM et al., 2021).')
P('O handebol de quadra é uma modalidade coletiva de caráter marcadamente intermitente e de alta intensidade. Ao longo da '
 'partida, esforços máximos e explosivos — sprints curtos, saltos, arremessos, mudanças de direção e contatos físicos — '
 'alternam-se, de forma imprevisível, com períodos de recuperação incompleta, exigindo simultaneamente potência '
 'anaeróbia, capacidade aeróbia intermitente e elevada tolerância à fadiga (KARCHER; BUCHHEIT, 2014; WAGNER et al., '
 '2014). As demandas físicas variam com a posição tática e a fadiga já foi documentada pela queda do volume de corrida em '
 'alta intensidade ao longo do jogo (MICHALSIK; MADSEN; AAGAARD, 2013). Esse perfil de esforço eleva a carga interna nos '
 'microciclos de acúmulo e repercute diretamente no estado afetivo dos atletas: em handebolistas de elite, o humor e o '
 'estresse associam-se a indicadores de sobrecarga fisiológica e psicossomática (RATZ-SULYOK et al., 2026). O caráter '
 'intermitente e de contato do handebol, somado à variabilidade posicional, torna o acompanhamento subjetivo do humor '
 'particularmente informativo para a individualização da carga.')
P('A resposta afetiva à carga, contudo, é dinâmica: reduções de carga diminuem a fadiga percebida e melhoram o '
 'desempenho, ao passo que microciclos de acúmulo tendem a rebaixar o vigor e a elevar a fadiga. A última semana de '
 'pré-temporada concentra a carga que antecede a competição e constitui, portanto, uma janela crítica para descrever como '
 'cada uma das seis subescalas do BRUMS se comporta ao longo dos dias e dentro de cada dia de treino, como as dimensões '
 'se relacionam entre si — em particular, se e quando o afeto negativo (depressão, raiva, confusão) passa a acompanhar a '
 'fadiga — e em que medida a resposta média do grupo se traduz em mudanças confiáveis no nível individual. Apesar da '
 'ampla adoção da BRUMS, faltam descrições detalhadas, subescala a subescala e do grupo ao sujeito, do comportamento do '
 'humor em um microciclo pré-competitivo de handebol de elite. Endereçar essa lacuna, com estatística comparativa robusta '
 '(variação percentual, intervalos de confiança, tamanho de efeito e mudança confiável individual), fornece à comissão '
 'técnica informação diretamente acionável e contribui para a literatura de monitoramento em modalidades coletivas '
 'intermitentes.')

# ===================== 2 OBJETIVO =====================
H('2 OBJETIVO')
P('Objetivo geral: avaliar o perfil de humor de atletas de handebol de elite na última semana de pré-temporada, '
 'caracterizando e comparando o comportamento das seis subescalas do BRUMS e da Perturbação Total do Humor, do nível do '
 'grupo ao nível do sujeito.')
P('Objetivos específicos: (i) caracterizar exploratoriamente cada subescala (tendência central, dispersão, intervalo de '
 'confiança e coeficiente de variação); (ii) comparar, no grupo, a resposta aguda pré→pós e a variação do Dia 1 ao Dia 7 '
 'de cada subescala, com variação percentual, tamanho de efeito e magnitude; (iii) analisar a estrutura de correlação '
 'entre as subescalas, com ênfase nos dias de maior vigor e de maior fadiga; (iv) quantificar, no nível do sujeito, a '
 'proporção de atletas com mudança confiável (RCI) em cada subescala; e (v) classificar a evolução dos perfis de humor.')
P('Hipóteses: (H1) a deterioração concentra-se no eixo energia–fadiga (vigor ↓, fadiga ↑) com magnitude média a grande, '
 'enquanto as subescalas negativas de valência não fadiga permanecem estáveis por efeito de piso; (H2) o perfil de humor '
 'migra do iceberg para o perfil de fadiga (barbatana de tubarão); (H3) o afeto negativo acopla-se à fadiga somente sob '
 'carga acumulada (correlação maior no dia de maior fadiga do que no de maior vigor); e (H4) a resposta é heterogênea '
 'entre atletas, exigindo interpretação individualizada.')

# ===================== 3 MÉTODO =====================
H('3 MÉTODO')
H('3.1 Delineamento',12)
P('Estudo descritivo-comparativo, longitudinal e observacional, de medidas repetidas intraindividuais, conduzido em '
 'condições ecológicas de treinamento durante o microciclo pré-competitivo (21–27/04/2024).')
H('3.2 Amostra',12)
P('Participaram %d atletas de handebol do sexo masculino de nível competitivo, distribuídos nas posições da modalidade '
 '(%s). A caracterização sociodemográfica e antropométrica consta na Tabela 1.'%(
   sm['n'],', '.join('%s (%d)'%(k,v) for k,v in sm['pos'].items())))
def srow(lab,k):
    v=sm[k]; return [lab,c2('%.1f'%v['mean']),c2('%.1f'%v['sd']),'%s–%s'%(c2('%.1f'%v['lo']),c2('%.1f'%v['hi'])),'%s–%s'%(c2('%.1f'%v['mn']),c2('%.1f'%v['mx']))]
table('Caracterização sociodemográfica e antropométrica da amostra (n = %d).'%sm['n'],
    ['Variável','Média','DP','IC95%','Mín–Máx'],
    [srow('Idade (anos)','idade'),srow('Estatura (cm)','estatura'),srow('Massa corporal (kg)','massa'),srow('Percentual de gordura','pG'),srow('Experiência na modalidade (anos)','exp')])
H('3.3 Instrumento',12)
P('O humor foi avaliado pela BRUMS-24: 24 adjetivos em escala de 0 (“nada”) a 4 (“extremamente”), agrupados em seis '
 'subescalas de 0 a 16 pontos (tensão, depressão, raiva, vigor, fadiga e confusão). A Perturbação Total do Humor (PTH) '
 'resume o perfil: PTH = tensão + depressão + raiva + fadiga + confusão − vigor.')
H('3.4 Procedimentos e framework das coletas',12)
P('O BRUMS foi autoaplicado por formulário eletrônico com carimbo de data/hora, duas vezes por dia de treino — a primeira '
 'resposta tomada como pré e a última como pós —, ao longo de sete dias, totalizando %d observações válidas. A Figura 1 '
 'sintetiza o framework das coletas.'%sm['n_obs'])
figure(f'{FG}/xb2_framework.png','Framework das coletas: da amostra e do microciclo às observações do BRUMS, às subescalas e à classificação de perfis.',w=11.0)
H('3.5 Análise estatística',12)
P('A distribuição das subescalas foi inspecionada e, dado o predomínio de assimetria e efeito de piso nas dimensões '
 'negativas, adotaram-se procedimentos não paramétricos. Reportou-se estatística descritiva (média, mediana, '
 'desvio-padrão, IC95%, coeficiente de variação e amplitude), no geral, por dia e nos momentos pré e pós. As comparações '
 'pré→pós e Dia 1→Dia 7 empregaram o teste de Wilcoxon pareado, com variação absoluta e percentual, tamanho de efeito '
 'pareado (dz de Cohen) com IC95% e classificação de magnitude (trivial < 0,2; pequeno < 0,5; médio < 0,8; grande ≥ 0,8). '
 'A estrutura de associação entre subescalas foi examinada por correlação de Spearman (matriz entre atletas e correlações '
 'diárias), com ênfase nos dias de maior vigor e de maior fadiga. No nível do sujeito, calculou-se o Índice de Mudança '
 'Confiável (RCI = (D7 − D1)/EP_dif; EP_dif = DP_D1·√[2(1 − α)]; α = fidedignidade interna), considerando mudança '
 'confiável quando |RCI| ≥ 1,96 (JACOBSON; TRUAX, 1991). Cada observação foi classificada nos seis perfis de humor por '
 'proximidade aos protótipos em escores padronizados. Adotou-se α = 0,05.')

# ===================== 4 RESULTADOS =====================
H('4 RESULTADOS')
H('4.1 Análise descritiva exploratória',12)
P('A Tabela 2 apresenta a descritiva geral das seis subescalas e da PTH. As dimensões negativas de valência não fadiga '
 '(tensão, depressão, raiva e confusão) situam-se próximas do piso, com medianas baixas e coeficientes de variação '
 'elevados, enquanto o vigor e a fadiga concentram a média e a dispersão informativas — é no eixo energia–fadiga que '
 'reside a variabilidade útil do microciclo. A Figura 2 exibe as trajetórias diárias com bandas de confiança de 95%.')
def drow(k,lab):
    v=desc[k]; return [lab,c2('%.2f'%v['mean']),c2('%.1f'%v['md']),c2('%.2f'%v['sd']),
        '%s–%s'%(c2('%.2f'%v['lo']),c2('%.2f'%v['hi'])),c2('%.0f'%v['cv']),'%s–%s'%(c2('%.0f'%v['mn']),c2('%.0f'%v['mx']))]
table('Estatística descritiva geral das subescalas do BRUMS e da PTH (%d observações).'%sm['n_obs'],
    ['Subescala','Média','Mediana','DP','IC95%','CV (%)','Mín–Máx'],[drow(k,l) for k,l in ORDER])
figure(f'{FG}/xb2_traj.png','Trajetória das subescalas do BRUMS ao longo da semana (médias diárias; áreas sombreadas = IC95%).')

H('4.2 Comparações no nível do grupo por subescala',12)
P('As Tabelas 3 e 4 reúnem, para todas as subescalas, a resposta aguda pré→pós e a variação Dia 1 → Dia 7, com variação '
 'percentual, tamanho de efeito (dz) com IC95%, significância e magnitude. A seguir, cada subescala é interpretada em '
 'ordem lógica, iniciando pelo eixo energia–fadiga.')
def prow(k,lab):
    v=pr[k]; return [lab,c2('%.2f'%v['pre']),c2('%.2f'%v['pos']),c2('%+.2f'%v['delta']),c2('%+.0f'%v['pct'])+'%',
        '%s [%s, %s]'%(c2('%+.2f'%v['dz']),c2('%.2f'%v['lo']),c2('%.2f'%v['hi'])),c2('%.3f'%v['p']),v['mag']]
table('Resposta aguda pré → pós agregada na semana, por subescala.',
    ['Subescala','Pré','Pós','Δ','Δ%','dz [IC95%]','p','Magnitude'],[prow(k,l) for k,l in ORDER],fs=8.5)
def d17row(k,lab):
    v=d17[k]; return [lab,c2('%.2f'%v['d1']),c2('%.2f'%v['d7']),c2('%+.2f'%v['delta']),c2('%+.0f'%v['pct'])+'%',c2('%+.2f'%v['dz']),c2('%.3f'%v['p']),v['mag']]
table('Comparação Dia 1 → Dia 7, por subescala.',
    ['Subescala','Dia 1','Dia 7','Δ','Δ%','dz','p','Magnitude'],[d17row(k,l) for k,l in ORDER],fs=8.5)

def scaleblk(num,k,lab,text):
    H('4.2.%d %s'%(num,lab),12); P(text)
scaleblk(1,'Vigor','Vigor',
 'O vigor foi a subescala de maior média (%s ± %s) e apresentou a resposta mais expressiva: caiu %s%% do pré ao pós '
 '(Δ = %s; dz = %s; p = %s; magnitude média) e recuou do Dia 1 ao Dia 7 com efeito grande (Δ = %s; %s%%; dz = %s; p = %s). '
 'É a dimensão que melhor sinaliza a deterioração da prontidão ao longo do microciclo.'%(
   c2('%.2f'%desc['Vigor']['mean']),c2('%.2f'%desc['Vigor']['sd']),c2('%.0f'%abs(pr['Vigor']['pct'])),c2('%+.2f'%pr['Vigor']['delta']),c2('%.2f'%pr['Vigor']['dz']),c2('%.3f'%pr['Vigor']['p']),
   c2('%+.2f'%d17['Vigor']['delta']),c2('%.0f'%d17['Vigor']['pct']),c2('%.2f'%d17['Vigor']['dz']),c2('%.3f'%d17['Vigor']['p'])))
scaleblk(2,'Fadiga','Fadiga',
 'A fadiga elevou-se de forma consistente: +%s%% do pré ao pós (Δ = %s; dz = %s; p = %s) e +%s%% do Dia 1 ao Dia 7 '
 '(Δ = %s; dz = %s; p = %s), ambas de magnitude média. Junto ao vigor, compõe o eixo energia–fadiga onde se concentra a '
 'resposta ao acúmulo de carga.'%(
   c2('%.0f'%pr['Fadiga']['pct']),c2('%+.2f'%pr['Fadiga']['delta']),c2('%.2f'%pr['Fadiga']['dz']),c2('%.3f'%pr['Fadiga']['p']),
   c2('%.0f'%d17['Fadiga']['pct']),c2('%+.2f'%d17['Fadiga']['delta']),c2('%.2f'%d17['Fadiga']['dz']),c2('%.3f'%d17['Fadiga']['p'])))
scaleblk(3,'TMD','Perturbação Total do Humor (PTH)',
 'A PTH, síntese do perfil, aumentou %s%% do pré ao pós (Δ = %s; dz = %s; p = %s) e +%s%% do Dia 1 ao Dia 7 (dz = %s; '
 'p = %s), refletindo sobretudo o movimento do eixo energia–fadiga.'%(
   c2('%.0f'%pr['TMD']['pct']),c2('%+.2f'%pr['TMD']['delta']),c2('%.2f'%pr['TMD']['dz']),c2('%.3f'%pr['TMD']['p']),
   c2('%.0f'%d17['TMD']['pct']),c2('%.2f'%d17['TMD']['dz']),c2('%.3f'%d17['TMD']['p'])))
scaleblk(4,'Tensao','Tensão',
 'A tensão manteve-se baixa e sem resposta aguda relevante (pré→pós dz = %s; p = %s), mas decresceu do Dia 1 ao Dia 7 '
 '(Δ = %s; dz = %s; p = %s, magnitude média) — uma acomodação da apreensão inicial do microciclo, sendo mais alta no '
 'início da semana.'%(
   c2('%+.2f'%pr['Tensao']['dz']),c2('%.3f'%pr['Tensao']['p']),c2('%+.2f'%d17['Tensao']['delta']),c2('%.2f'%d17['Tensao']['dz']),c2('%.3f'%d17['Tensao']['p'])))
scaleblk(5,'Depressao','Depressão',
 'A depressão operou próxima do piso, com leve elevação aguda (pré→pós +%s%%; dz = %s; p = %s) e pico intermediário no '
 'Dia %d (%s), sugerindo um vale afetivo no meio da semana, sem instalar quadro clínico.'%(
   c2('%.0f'%pr['Depressao']['pct']),c2('%+.2f'%pr['Depressao']['dz']),c2('%.3f'%pr['Depressao']['p']),npk['Depressao']['peak_day'],c2('%.2f'%npk['Depressao']['peak_val'])))
scaleblk(6,'Raiva','Raiva',
 'A raiva foi a negativa mais reativa e atingiu o máximo no dia mais fadigado (pico no Dia %d = %s), acompanhando o '
 'acúmulo de carga, embora sem resposta aguda intra-dia significativa (pré→pós p = %s).'%(
   npk['Raiva']['peak_day'],c2('%.2f'%npk['Raiva']['peak_val']),c2('%.3f'%pr['Raiva']['p'])))
scaleblk(7,'Confusao','Confusão',
 'A confusão foi a subescala de menor expressão (média %s), mais alta no início e decrescente ao longo da semana, '
 'atingindo o mínimo (%s) — sem resposta aguda relevante (pré→pós dz = %s).'%(
   c2('%.2f'%desc['Confusao']['mean']),c2('%.2f'%npk['Confusao']['min_val']),c2('%+.2f'%pr['Confusao']['dz'])))

H('4.3 Estrutura de correlação e ênfase nos dias de maior vigor e de maior fadiga',12)
CT=R['corr_trait']; LB=R['corr_labels']
def gg(a,b): return c2('%+.2f'%CT[LB.index(a)][LB.index(b)])
P('A matriz de correlação entre atletas (Figura 3) revela um agrupamento coeso de afeto negativo: a depressão '
 'correlaciona-se com a raiva (ρ = %s), a confusão (ρ = %s) e a tensão (ρ = %s), e — de forma relevante — com a fadiga '
 '(ρ = %s); a raiva também acompanha a fadiga (ρ = %s). O eixo energia–fadiga aparece invertido (vigor × fadiga ρ = %s) e '
 'o vigor é praticamente independente das negativas.'%(
   gg('Depressão','Raiva'),gg('Depressão','Confusão'),gg('Depressão','Tensão'),gg('Depressão','Fadiga'),gg('Raiva','Fadiga'),gg('Vigor','Fadiga')))
figure(f'{FG}/xb2_heatmap.png','Matriz de correlação de Spearman entre as subescalas do BRUMS (nível entre atletas).',w=10.5)
P('O contraste entre os dias extremos é o achado mais expressivo (Figura 4; Tabela 5). No dia de maior vigor (Dia 1, '
 'grupo descansado), o afeto negativo é quase independente da fadiga (depressão ρ = %s; raiva ρ = %s; confusão ρ = %s). No '
 'dia de maior fadiga (Dia 7), o acoplamento torna-se forte (depressão ρ = %s; raiva ρ = %s; confusão ρ = %s), com '
 'associação negativa ao vigor. Ou seja, quando o grupo está fresco, quem está mais irritado ou abatido não é '
 'necessariamente o mais cansado; sob carga acumulada, as dimensões negativas consolidam-se com a fadiga e o perfil de '
 'humor “fecha”. O padrão cresce progressivamente ao longo dos dias, com recuo transitório no Dia 5 (menor carga).'%(
   c2('%+.2f'%foc['D1']['Depressao']['rho_fad']),c2('%+.2f'%foc['D1']['Raiva']['rho_fad']),c2('%+.2f'%foc['D1']['Confusao']['rho_fad']),
   c2('%+.2f'%foc['D7']['Depressao']['rho_fad']),c2('%+.2f'%foc['D7']['Raiva']['rho_fad']),c2('%+.2f'%foc['D7']['Confusao']['rho_fad'])))
figure(f'{FG}/xb2_corrfocus.png','Acoplamento das subescalas negativas com a fadiga: (A) dia de maior vigor (D1) vs. dia de maior fadiga (D7); (B) evolução dia a dia.')
def frow(neg,lab):
    return [lab,c2('%+.2f'%foc['D1'][neg]['rho_fad']),c2('%+.2f'%foc['D7'][neg]['rho_fad']),c2('%+.2f'%foc['D1'][neg]['rho_vig']),c2('%+.2f'%foc['D7'][neg]['rho_vig'])]
table('Correlação (ρ de Spearman) das subescalas negativas com a fadiga e o vigor no dia de maior vigor (D1) e no de maior fadiga (D7).',
    ['Subescala','ρ×Fadiga (D1)','ρ×Fadiga (D7)','ρ×Vigor (D1)','ρ×Vigor (D7)'],
    [frow('Depressao','Depressão'),frow('Raiva','Raiva'),frow('Confusao','Confusão'),frow('Tensao','Tensão')],fs=9)

H('4.4 Comportamento intra-dia (pré → pós por dia)',12)
P('O padrão intra-dia repete-se com regularidade: em todos os dias de treino o vigor cai e a fadiga e a PTH sobem do pré '
 'para o pós, indicando um choque agudo consistente a cada sessão, enquanto as subescalas negativas variam pouco dentro '
 'do dia (Figura 5).')
figure(f'{FG}/xb2_intraday.png','Escores pré e pós-treino por dia (vigor, fadiga, PTH) e variação intra-dia das subescalas negativas.')

H('4.5 Análise no nível do sujeito por subescala (mudança confiável)',12)
P('No nível individual, a mudança confiável (RCI, Dia 1 → Dia 7) distribui-se de modo desigual entre as subescalas '
 '(Tabela 6; Figura 6). A fadiga é a dimensão em que a maioria dos atletas deteriora de forma confiável (%d de %d; %d%%), '
 'seguida do vigor (%d%% de queda confiável). As subescalas negativas de valência não fadiga não pioram no nível '
 'individual — a tensão e a confusão, inclusive, exibem apenas melhoras confiáveis (acomodação), e a depressão e a raiva '
 'apresentam proporções equilibradas de piora e melhora. Confirma-se, no sujeito, que a deterioração real do microciclo '
 'se concentra no eixo energia–fadiga.'%(
   RC['Fadiga']['piora'],RC['Fadiga']['n'],RC['Fadiga']['pct_piora'],RC['Vigor']['pct_piora']))
def rcrow(k,lab):
    v=RC[k]; return [lab,c2('%.2f'%v['alpha']),c2('%.2f'%v['sediff']),v['n'],
        '%d (%d%%)'%(v['piora'],v['pct_piora']),'%d (%d%%)'%(v['estavel'],v['pct_est']),'%d (%d%%)'%(v['melhora'],v['pct_mel'])]
table('Mudança confiável individual (RCI, Dia 1 → Dia 7) por subescala: fidedignidade, erro-padrão da diferença e número de atletas (n = %d).'%RC['Vigor']['n'],
    ['Subescala','α','EP_dif','n','Piora confiável','Estável','Melhora confiável'],
    [rcrow(k,lab) for k,lab in [('Vigor','Vigor'),('Fadiga','Fadiga'),('Depressao','Depressão'),('Raiva','Raiva'),('Tensao','Tensão'),('Confusao','Confusão')]],fs=9)
figure(f'{FG}/xb2_rci6.png','Proporção de atletas com mudança confiável (RCI) por subescala, do Dia 1 ao Dia 7.',w=13.5)

H('4.6 Classificação nos perfis de humor',12)
P('Na linha de base (Dia 1) predomina o perfil iceberg (%d%%), marca de prontidão; no último dia (Dia 7) o perfil de '
 'fadiga (“barbatana de tubarão”) torna-se o mais frequente (%d%%) e o iceberg cai para %d%% (Figura 7; Tabela 7). Essa '
 'migração — do iceberg para a assinatura de fadiga funcional, sem instalação relevante de perfis de risco (submerso, '
 'iceberg invertido) — é a leitura integrada da semana.'%(
   int(prof['D1']['Iceberg']),int(prof['D7']['Barbatana tubarão']),int(prof['D7']['Iceberg'])))
figure(f'{FG}/xb2_profiles.png','Distribuição (%) dos seis perfis de humor por recorte: linha de base (D1), semana, último dia (D7), pré e pós.')
PROFR=[('Iceberg','Iceberg'),('Everest invertido','Everest invertido'),('Iceberg invertido','Iceberg invertido'),('Submerso','Submerso'),('Barbatana tubarão','Barbatana de tubarão'),('Superfície','Superfície')]
recs=[('D1','D1'),('Semana','overall'),('D7','D7'),('Pré','pre'),('Pós','pos')]
table('Distribuição (%) dos perfis de humor de Terry por recorte.',
    ['Perfil']+[l for l,_ in recs],
    [[lab]+[c2('%.0f'%prof[key][p]) for _,key in recs] for p,lab in PROFR],fs=9)

# ===================== 5 DISCUSSÃO =====================
H('5 DISCUSSÃO')
P('Os resultados confirmam as quatro hipóteses. A deterioração do humor no microciclo pré-competitivo concentrou-se no '
 'eixo energia–fadiga (H1): o vigor recuou com efeito grande (dz = %s) e a fadiga elevou-se com efeito médio, enquanto as '
 'subescalas negativas de valência não fadiga permaneceram próximas do piso — padrão coerente com a literatura de '
 'monitoramento, que identifica o vigor e a fadiga como as dimensões subjetivas mais sensíveis à carga (SAW; MAIN; '
 'GASTIN, 2016; THORPE et al., 2017). A migração do perfil iceberg para o de fadiga (H2) reproduz, em um microciclo de '
 'sete dias, o “derretimento do iceberg” descrito em fases de acúmulo, e é compatível com a associação entre perturbação '
 'do humor e desempenho (MORGAN, 1985; LOCHBAUM et al., 2021).'%c2('%.2f'%d17['Vigor']['dz']))
P('O achado de que o afeto negativo se acopla à fadiga apenas sob carga acumulada (H3) — depressão × fadiga passando de '
 'ρ = %s no dia de maior vigor para ρ = %s no de maior fadiga — sugere que, em atletas de elite, a irritabilidade e o '
 'abatimento não são ruído aleatório, mas sinais que se organizam com a exaustão, o que reforça o valor de interpretar as '
 'subescalas em conjunto e à luz do estado de fadiga. Por fim, a heterogeneidade individual (H4), com a fadiga sendo a '
 'única subescala em que a maioria dos atletas deteriora de forma confiável, sustenta o monitoramento individualizado por '
 'tendência e referenciado à linha de base de cada atleta, em linha com as recomendações de consenso (KELLMANN et al., '
 '2018). O caráter intermitente e de alta intensidade do handebol (KARCHER; BUCHHEIT, 2014) contextualiza a magnitude da '
 'resposta de fadiga observada.'%(c2('%+.2f'%foc['D1']['Depressao']['rho_fad']),c2('%+.2f'%foc['D7']['Depressao']['rho_fad'])))
P('Como limitações, destacam-se o tamanho amostral (n = %d) e o efeito de piso das subescalas negativas, que reduz a '
 'variância e a fidedignidade da tensão e da confusão, e o caráter observacional de fase única, que não permite inferência '
 'causal sobre a carga. Ainda assim, o desenho de medidas repetidas, o volume de observações e a estatística baseada em '
 'tamanho de efeito e mudança confiável conferem robustez descritiva aos achados.'%sm['n'])

# ===================== 6 CONSIDERAÇÕES FINAIS =====================
H('6 CONSIDERAÇÕES FINAIS')
P('Na última semana de pré-temporada, o perfil de humor de handebolistas de elite migrou da prontidão (iceberg) para a '
 'fadiga funcional (barbatana de tubarão), com a deterioração concentrada no eixo energia–fadiga — vigor em queda de '
 'efeito grande e fadiga em ascensão —, tanto no grupo quanto no sujeito, onde a fadiga foi a única subescala a piorar de '
 'forma confiável na maioria dos atletas. As dimensões negativas mantiveram-se estáveis e só se acoplaram à fadiga sob '
 'carga acumulada. Recomenda-se monitorar o eixo energia–fadiga por tendência e referenciado à linha de base individual, '
 'com atenção redobrada ao início do microciclo (choque de carga) e ao último dia (fadiga acumulada).')

# ===================== REFERÊNCIAS =====================
H('REFERÊNCIAS')
refs=[
 'BRANDT, R.; BEVILACQUA, G. G.; ANDRADE, A. Perceived sleep quality, mood states, and their relationship with performance among Brazilian elite athletes during a competitive period. Journal of Strength and Conditioning Research, v. 31, n. 4, p. 1033–1039, 2017.',
 'JACOBSON, N. S.; TRUAX, P. Clinical significance: a statistical approach to defining meaningful change in psychotherapy research. Journal of Consulting and Clinical Psychology, v. 59, n. 1, p. 12–19, 1991. DOI: 10.1037/0022-006X.59.1.12.',
 'KARCHER, C.; BUCHHEIT, M. On-court demands of elite handball, with special reference to playing positions. Sports Medicine, v. 44, n. 6, p. 797–814, 2014. DOI: 10.1007/s40279-014-0164-z.',
 'KELLMANN, M. et al. Recovery and performance in sport: consensus statement. International Journal of Sports Physiology and Performance, v. 13, n. 2, p. 240–245, 2018. DOI: 10.1123/ijspp.2017-0759.',
 'LEW, P. C. F. et al. Cross-cultural validation of the Malaysian Mood Scale and tests of between-group mood differences. International Journal of Environmental Research and Public Health, v. 20, n. 4, 3348, 2023. DOI: 10.3390/ijerph20043348.',
 'LOCHBAUM, M. et al. The Profile of Mood States and athletic performance: a meta-analysis of published studies. European Journal of Investigation in Health, Psychology and Education, v. 11, n. 1, p. 50–70, 2021. DOI: 10.3390/ejihpe11010005.',
 'MICHALSIK, L. B.; MADSEN, K.; AAGAARD, P. Match performance and physiological capacity of female elite team handball players. International Journal of Sports Medicine, v. 35, n. 7, p. 595–607, 2013. DOI: 10.1055/s-0033-1358713.',
 'MORGAN, W. P. Selected psychological factors limiting performance: a mental health model. In: CLARKE, D. H.; ECKERT, H. M. (Ed.). Limits of human performance. Champaign: Human Kinetics, 1985. p. 70–80.',
 'PARSONS-SMITH, R. L.; TERRY, P. C.; MACHIN, M. A. Identification and description of novel mood profile clusters. Frontiers in Psychology, v. 8, 1958, 2017. DOI: 10.3389/fpsyg.2017.01958.',
 'RATZ-SULYOK, F. Z. et al. Associations between endocrine status and stress, mood and psychosomatic status in elite handball players. Sports, v. 14, n. 7, 289, 2026. DOI: 10.3390/sports14070289.',
 'ROHLFS, I. C. P. M. et al. A Escala de Humor de Brunel (Brums): instrumento para detecção precoce da síndrome do excesso de treinamento. Revista Brasileira de Medicina do Esporte, v. 14, n. 3, p. 176–181, 2008.',
 'SAW, A. E.; MAIN, L. C.; GASTIN, P. B. Monitoring the athlete training response: subjective self-reported measures trump commonly used objective measures: a systematic review. British Journal of Sports Medicine, v. 50, n. 5, p. 281–291, 2016. DOI: 10.1136/bjsports-2015-094758.',
 'TAN, C. et al. The structural validity and latent profile characteristics of the Abbreviated Profile of Mood States among Chinese athletes. BMC Psychiatry, v. 24, n. 1, 636, 2024. DOI: 10.1186/s12888-024-06092-5.',
 'TERRY, P. C.; LANE, A. M.; LANE, H. J.; KEOHANE, L. Development and validation of a mood measure for adolescents. Journal of Sports Sciences, v. 17, n. 11, p. 861–872, 1999. DOI: 10.1080/026404199365425.',
 'TERRY, P. C.; LANE, A. M.; FOGARTY, G. J. Construct validity of the Profile of Mood States — Adolescents for use with adults. Psychology of Sport and Exercise, v. 4, n. 2, p. 125–139, 2003. DOI: 10.1016/S1469-0292(02)00035-8.',
 'TERRY, P. C. et al. Mood profiling for sustainable mental health among athletes. Sustainability, v. 13, n. 11, 6116, 2021. DOI: 10.3390/su13116116.',
 'THORPE, R. T. et al. Monitoring fatigue status in elite team-sport athletes: implications for practice. International Journal of Sports Physiology and Performance, v. 12, s2, p. S2-27–S2-34, 2017. DOI: 10.1123/ijspp.2016-0434.',
 'WAGNER, H. et al. Individual and team performance in team-handball: a review. Journal of Sports Science and Medicine, v. 13, n. 4, p. 808–816, 2014.']
for rf in refs:
    p=doc.add_paragraph(); r=p.add_run(rf); r.font.size=Pt(11); p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(6); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

OUTP='/home/user/mdlucca/Artigos/Perfil_Humor_BRUMS_Descritivo.docx'
doc.save(OUTP); print('SAVED',OUTP,'| Tabelas',_TN[0],'Figuras',_FN[0])
