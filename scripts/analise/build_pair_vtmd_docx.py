import warnings; warnings.filterwarnings('ignore')
import numpy as np, pandas as pd, json, os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
SC='/tmp/claude-0/-home-user-mdlucca/e1dba24c-b1d7-5908-9106-f2f4aaf3f56a/scratchpad'
FIG='/home/user/mdlucca/Artigos/figuras'
G=json.load(open(f'{SC}/pair_vigor_tmd.json'))
grp,btw,per,PP,corr,eff,kmdc=G['grp'],G['btw'],G['per'],G['pp'],G['corr'],G['eff'],G['kmdc']
VARS=[('Vigor','Vigor (0–20)'),('TMD','PTH/TMD')]

doc=Document(); st=doc.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(10.5)
def H(t,l=1): doc.add_heading(t,level=l)
def P(t,it=False,b=False):
    p=doc.add_paragraph(); r=p.add_run(t); r.italic=it; r.bold=b; return p
def fig(name,cap,w=6.4):
    path=f'{FIG}/{name}.png'
    if os.path.exists(path):
        doc.add_picture(path,width=Inches(w)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
        c=doc.add_paragraph(); r=c.add_run(cap); r.italic=True; r.font.size=Pt(8.5); c.alignment=WD_ALIGN_PARAGRAPH.CENTER
def table(headers,rows,cap=None):
    if cap: c=doc.add_paragraph(); r=c.add_run(cap); r.bold=True; r.font.size=Pt(9.5)
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Light Grid Accent 1'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        cell=t.rows[0].cells[i]; cell.paragraphs[0].add_run(str(h)).bold=True; cell.paragraphs[0].runs[0].font.size=Pt(8.5)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=str(v)
            if cells[i].paragraphs[0].runs: cells[i].paragraphs[0].runs[0].font.size=Pt(8.5)
    doc.add_paragraph()

# CAPA
doc.add_heading('Amplitude e Separação Ruído × Sinal — Vigor vs PTH/TMD',0)
P('Vigor (0–20) e Perturbação Total do Humor (PTH/TMD) ao longo do microciclo de choque de HIIT — 21 a 28 de abril de 2024',it=True)
P('Análise por grupo e por atleta · dentro da semana e entre dias · pré/mid/pós · remoção de ruído · 27 atletas · 456 observações · reanálise independente em Python. Atletas anonimizados (A01–A27).')
doc.add_paragraph()

# 1 método
H('1. Objetivo e método')
P('Aplica-se ao par Vigor (0–20) e Perturbação Total do Humor (PTH/TMD) o mesmo protocolo: amplitude (range) no grupo e no atleta, decomposição da variância em sinal do microciclo / traço / ruído (modelo y ~ atleta + dia), limiares de decisão a partir do ruído (ETM, MDC95 individual e de grupo, SWC), curso intra-dia (pré→mid→pós) e o comportamento das análises quando o ruído é removido. A pergunta central deste par é o acoplamento: como o Vigor (energia) e o PTH/TMD (perturbação global do humor, que já embute o vigor com sinal invertido) se movem juntos — e quanta informação o TMD acrescenta ao vigor.')

# 2 grupo semana
H('2. Grupo — ao longo da semana')
rows=[[lab,f"{grp[v]['d1']:.2f}",f"{grp[v]['d7']:.2f}",f"{grp[v]['drift']:+.2f}",f"{grp[v]['sig']:.2f}",f"{grp[v]['etm']:.2f}",f"{grp[v]['snr']:.2f}"] for v,lab in VARS]
table(['Variável','Dia 1','Dia 7','Deriva D7−D1','Ampl. do sinal','ETM (ruído)','SNR'],rows,'Tabela 1 — Trajetória do grupo e amplitude do sinal')
P('As duas variáveis movem-se em sentidos opostos, como esperado do mesmo eixo energia–humor: o Vigor cai (deriva −2,83) e o PTH/TMD sobe (deriva +5,95) ao longo da semana. O TMD tem a maior amplitude de sinal em unidades brutas (6,24) mas também o maior ETM (6,03), de modo que seu SNR (1,04) é menor que o do Vigor (1,36) — o TMD é uma medida mais ruidosa da mesma tendência.')
fig('pv1_trajetoria','Fig 1 — Vigor vs PTH/TMD ao longo da semana (média ± DP; LOWESS). O Vigor cai e o TMD sobe, ambos com precipitação no D7.')
fig('pv2_dissociacao','Fig 2 — Curvas padronizadas (z) sobrepostas: Vigor e PTH/TMD são espelhos — quando um sobe, o outro desce (acoplamento inverso).')
rows=[[lab,f"{grp[v]['pday']:.1f}%",f"{grp[v]['ptrait']:.1f}%",f"{grp[v]['pnoise']:.1f}%"] for v,lab in VARS]
table(['Variável','% Sinal (dia)','% Traço','% Ruído'],rows,'Tabela 2 — Decomposição da variância')
P('O Vigor tem 4,9% da variância como sinal do microciclo e o PTH/TMD 2,5%; ambos são dominados por traço (54–61%) e ruído (36–41%). O TMD, por agregar seis subescalas (várias no piso), dilui o sinal do eixo energia–fadiga em ruído das negativas — por isso rende menos sinal por unidade de ruído que o próprio Vigor isolado.')

# 3 entre dias
H('3. Grupo — entre dias')
rows=[[lab]+[f"{btw[v]['sd_between'].get(str(d),'—')}" for d in range(1,8)]+[f"D{btw[v]['day_maxsd']}"] for v,lab in VARS]
table(['Variável','D1','D2','D3','D4','D5','D6','D7','Máx. disp.'],rows,'Tabela 3 — Dispersão ENTRE atletas por dia (DP)')
rows=[[lab]+[f"{btw[v]['dd'].get(str(d),'—')}" for d in range(2,8)]+[f"D{btw[v]['day_maxdd']}"] for v,lab in VARS]
table(['Variável','D1→2','D2→3','D3→4','D4→5','D5→6','D6→7','Maior passo'],rows,'Tabela 4 — |Mudança dia-a-dia| da média do grupo')
P('O maior passo de ambos ocorre no fim da semana (a precipitação do D7): o Vigor despenca e o TMD dispara. A dispersão entre atletas é alta e persistente nas duas variáveis, sem um único dia dominante.')

# 4 por atleta
H('4. Por atleta — amplitude e sinal-ruído individual')
rows=[[lab,f"{per[v]['amp_mean']:.2f} ± {per[v]['amp_sd']:.2f}",f"{per[v]['amp_min']:.0f}–{per[v]['amp_max']:.0f}",f"{grp[v]['sig']:.2f}",f"{per[v]['pnoise_mean']:.2f}",f"{per[v]['pct_snr_gt1']:.0f}%"] for v,lab in VARS]
table(['Variável','Ampl. atleta (média±DP)','Faixa','Ampl. sinal grupo','Ruído pessoal','% atletas SNR>1'],rows,'Tabela 5 — Amplitude por atleta vs sinal do grupo')
P('A amplitude individual é grande nas duas (Vigor ~7,1; TMD ~20,1 em média) e muito maior que a amplitude do sinal do grupo (2,83 e 6,24) — o atleta oscila bem mais que a trajetória média, porque soma o próprio ruído. Ainda assim, a maioria dos atletas tem SNR pessoal > 1 sobre médias diárias.')
fig('pv3_amp_box','Fig 3 — Amplitude por atleta (caixa) vs amplitude do sinal do grupo (losango). Nos dois casos o atleta oscila bem mais que a trajetória média.')

# 5 prepost
H('5. Pré → mid → pós por dia')
P('Curso intra-sessão: média ± DP de pré, mid (intra-sessão) e pós, Δ (pós−pré), dz pareado, Wilcoxon pré×pós e Friedman pré/mid/pós. Dia 1 (baseline) sem mid. * indica p<0,05.')
for vi,(v,lab) in enumerate(VARS):
    rows=[]
    for d in range(1,8):
        r=PP[v][str(d)]
        pstr='—' if r['p'] is None else (f"{r['p']:.3f}"+(' *' if r['p']<0.05 else ''))
        fstr='—' if r['pfried'] is None else (f"{r['pfried']:.3f}"+(' *' if r['pfried']<0.05 else ''))
        dzs='—' if r['dz'] is None else f"{r['dz']:+.2f}"
        mids='—' if r['mid_m'] is None else (f"{r['mid_m']:.2f} ± {r['mid_sd']:.2f}" if r['mid_sd'] is not None else f"{r['mid_m']:.2f}")
        rows.append([f"D{d}"+(" (HIIT)" if r['hiit'] else ""),f"{r['pre_m']:.2f} ± {r['pre_sd']:.2f}",mids,f"{r['pos_m']:.2f} ± {r['pos_sd']:.2f}",f"{r['delta']:+.2f}",dzs,pstr,fstr])
    table(['Dia','Pré (m ± DP)','Mid (m ± DP)','Pós (m ± DP)','Δ (pós−pré)','dz','p pré×pós','p Friedman'],rows,f'Tabela {6+vi} — {lab}: pré → mid → pós por dia')
P('O curso intra-sessão é coerente e inverso: o Vigor cai de pré para pós e o PTH/TMD sobe, muitas vezes com pico no mid (a perturbação do humor é máxima no meio da sessão e recua um pouco ao final). A resposta aguda é mais nítida no D1 e no fim da semana.')
fig('pv4_prepos','Fig 4 — Pré → mid → pós por dia. O Vigor (esquerda) cai pós-treino; o PTH/TMD (direita) sobe, com pico no mid.')

# 6 veredito
H('6. Veredito em dois níveis — grupo vs indivíduo')
rows=[[lab,f"{grp[v]['sig']:.2f}",f"{grp[v]['mdc']:.2f}",('SIM' if grp[v]['detect_ind'] else 'NÃO'),f"{grp[v]['mdc_g']:.2f}",('SIM' if grp[v]['detect_grp'] else 'NÃO')] for v,lab in VARS]
table(['Variável','Ampl. do sinal','MDC95 individual','> MDC ind.?','MDC95 grupo (÷√n)','> MDC grupo?'],rows,'Tabela 8 — A oscilação semanal supera o ruído?')
P('Ambos são sinal robusto no grupo (a oscilação supera o MDC do grupo por 4–8×) mas ficam abaixo do MDC95 individual de 1 coleta (Vigor 5,77; TMD 16,71): a decisão individual exige médias de várias coletas — e o TMD, por ser mais ruidoso, precisa de mais coletas que o Vigor para a mesma confiança.')

# 7 remover ruido
H('7. E se removermos o ruído? Como ficam as análises')
P('Remover o ruído (agregar, suavizar ou desatenuar) revela o quão fortemente Vigor e PTH/TMD são espelhos um do outro.')
rows=[['Vigor × PTH/TMD — observada (bruta, Pearson)',f"{corr['r_obs']:+.2f}"],
      ['— Spearman (bruta)',f"{corr['rho_obs']:+.2f}"],
      ['— rmcorr (intra-atleta, momentânea)',f"{corr['rm']:+.2f}"],
      ['— desatenuada (erro de medida removido)',f"{corr['r_dis']:+.2f}"],
      ['— nível-grupo (médias diárias; ruído agregado)',f"{corr['r_den']:+.2f}"]]
table(['Estimativa da associação Vigor × PTH/TMD','r'],rows,'Tabela 9 — A correlação Vigor × PTH/TMD em vários níveis')
P('Vigor e PTH/TMD são fortemente acoplados de forma inversa, e removido o ruído tornam-se quase espelhos perfeitos: a correlação bruta −0,62 (Spearman −0,68; rmcorr intra-atleta −0,61) sobe em módulo para −0,78 no nível-grupo. A desatenuação leva a estimativa a −1,01 — um caso Heywood esperado, pois o próprio Vigor entra na fórmula do TMD com sinal invertido. Conclusão metodológica: o TMD não é um constructo independente do Vigor; para o eixo energia–fadiga, medir Vigor (e Fadiga) já captura quase toda a informação que o TMD carrega, com menos ruído.')
fig('pv5_scatter','Fig 5 — Vigor × PTH/TMD: observações brutas (esquerda; r=−0,62) e médias diárias do grupo (direita; r=−0,78). Removido o ruído, o acoplamento inverso fica ainda mais forte.')
rows=[[lab,f"{eff[v]['dz']:+.2f}",f"{eff[v]['dz_corr']:+.2f}",f"{grp[v]['rel']:.2f}"] for v,lab in VARS]
table(['Variável','dz observado','dz corrigido (ruído removido)','Confiab. (1 medida)'],rows,'Tabela 10 — Efeito D1→D7 corrigido pela atenuação')
P('Corrigido pela confiabilidade, o efeito da semana cresce nos dois: o Vigor de dz −0,95 para −1,24 e o PTH/TMD de +0,41 para +0,51. O efeito no Vigor é maior (e menos ruidoso) que no TMD — mais uma indicação de que o Vigor é o melhor portador do sinal do eixo.')
rows=[]
for v,lab in VARS:
    emg=next((k for k in [1,2,3,5,7] if grp[v]['sig']>kmdc[v][str(k)]),None)
    rows.append([lab,f"{grp[v]['sig']:.2f}",f"{kmdc[v]['1']:.2f}",f"{kmdc[v]['3']:.2f}",f"{kmdc[v]['5']:.2f}",f"{kmdc[v]['7']:.2f}",(f"k≥{emg}" if emg else "não emerge")])
table(['Variável','Ampl. sinal','MDC k=1','MDC k=3','MDC k=5','MDC k=7','Sinal supera o ruído em'],rows,'Tabela 11 — MDC95 individual cai com k coletas (ruído/√k)')
P('Para o Vigor, a partir de ~5 coletas o sinal semanal supera o MDC individual — o monitoramento individual é viável com médias curtas. O PTH/TMD, por ser muito ruidoso em unidades brutas, não emerge nem com 7 coletas: no indivíduo, prefira o Vigor (e a Fadiga física) ao TMD isolado.')
fig('pv6_mdc','Fig 6 — MDC95 individual em função do nº de coletas (k). A linha branca é a amplitude do sinal semanal: no Vigor a curva do MDC cruza abaixo dela (~k=5); no PTH/TMD, não cruza até k=7.')

# 8 sintese
H('8. Síntese')
for t in [
 'Acoplamento inverso: o Vigor cai (deriva −2,83; dz −0,95) e o PTH/TMD sobe (deriva +5,95; dz +0,41) — os dois lados do mesmo eixo energia–humor.',
 'Amplitude: o TMD tem a maior amplitude bruta (sinal 6,24) mas o pior SNR (1,04) — mais ruidoso que o Vigor (1,36) por embutir as negativas de piso.',
 'Intra-sessão: Vigor desce e TMD sobe pós-treino, com pico de perturbação no mid.',
 'Removendo o ruído: Vigor e TMD tornam-se quase espelhos (−0,62→−0,78 nível-grupo; −1,01 desatenuada = caso Heywood, pois o Vigor compõe o TMD); ambos os efeitos crescem (Vigor −0,95→−1,24).',
 'Uso prático: o TMD não acrescenta informação independente ao Vigor+Fadiga; para o eixo energia–fadiga, prefira o Vigor (menos ruidoso, viável no indivíduo a partir de ~5 coletas) ao TMD isolado.',
]:
    p=doc.add_paragraph(style='List Bullet'); p.add_run(t)
doc.add_paragraph()
P('Reprodutibilidade: scripts/analise/pair_vigor_tmd.py, build_pair_vtmd_docx.py. Consistente com Amplitude_Ruido_Sinal_Vigor_Fadiga_TMD.docx e Analise_Estatistica_Consolidada.md. Atletas anonimizados; nenhum dado bruto com nomes é distribuído.',it=True)

OUTP='/home/user/mdlucca/Artigos/Amplitude_Ruido_Sinal_Vigor_TMD.docx'
doc.save(OUTP); print('SAVED',OUTP)
d2=Document(OUTP); print('paras',len(d2.paragraphs),'tables',len(d2.tables),'imgs',len(d2.inline_shapes))
