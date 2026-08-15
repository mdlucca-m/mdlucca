import warnings; warnings.filterwarnings('ignore')
import numpy as np, pandas as pd, json, os
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
SC='/tmp/claude-0/-home-user-mdlucca/e1dba24c-b1d7-5908-9106-f2f4aaf3f56a/scratchpad'
FIG='/home/user/mdlucca/Artigos/figuras'
R=json.load(open(f'{SC}/semana2.json')); BF=json.load(open(f'{SC}/baseline_final.json')); REG=json.load(open(f'{SC}/reg.json'))
REL=json.load(open(f'{SC}/reliab.json')); dec=R['dec']; desc=R['desc']; rank=R['rank']; agu=R['agu']
ROC=json.load(open(f'{SC}/roc.json')); MV=json.load(open(f'{SC}/mv.json')); INF=R['ice']
CORE=[('Vigor','Vigor'),('Fadiga','Fadiga (BRUMS)'),('FadFisica','Fadiga física'),('FadMental','Fadiga mental'),('TMD','PTH/TMD')]
def c2(x): return str(x).replace('.',',')

doc=Document(); s=doc.sections[0]; s.left_margin=Cm(3); s.top_margin=Cm(3); s.right_margin=Cm(2); s.bottom_margin=Cm(2)
n=doc.styles['Normal']; n.font.name='Times New Roman'; n.font.size=Pt(12); n.paragraph_format.line_spacing=1.5; n.paragraph_format.space_after=Pt(0)
FN='Times New Roman'
def _set(r,sz=12,b=False,it=False): r.font.name=FN; r.font.size=Pt(sz); r.bold=b; r.italic=it
def sec(num,txt):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(4); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    _set(p.add_run((f'{num} {txt}').upper() if '.' not in num else f'{num} {txt}'),12,b=True)
def P(t,indent=True,size=12,it=False,b=False):
    p=doc.add_paragraph(); _set(p.add_run(t),size,it=it,b=b); p.paragraph_format.line_spacing=1.5
    if indent: p.paragraph_format.first_line_indent=Cm(1.25)
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; return p
_FN=[0]
def fig(name,titulo,w=5.4,fonte='Elaborado pelos autores (2026).'):
    _FN[0]+=1
    cap=doc.add_paragraph(); cap.alignment=WD_ALIGN_PARAGRAPH.CENTER; cap.paragraph_format.space_before=Pt(6); cap.paragraph_format.line_spacing=1.0
    _set(cap.add_run(f'Figura {_FN[0]} – {titulo}'),10)
    p=f'{FIG}/{name}.png'
    if os.path.exists(p): doc.add_picture(p,width=Inches(w)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    else: print('MISS',name)
    fp=doc.add_paragraph(); fp.alignment=WD_ALIGN_PARAGRAPH.CENTER; fp.paragraph_format.space_after=Pt(6); fp.paragraph_format.line_spacing=1.0
    _set(fp.add_run(f'Fonte: {fonte}'),10)
def _bd(t):
    tblPr=t._tbl.tblPr; b=OxmlElement('w:tblBorders')
    for e in ('top','bottom','insideH'):
        x=OxmlElement(f'w:{e}'); x.set(qn('w:val'),'single'); x.set(qn('w:sz'),'6'); x.set(qn('w:space'),'0'); x.set(qn('w:color'),'000000'); b.append(x)
    for e in ('left','right','insideV'):
        x=OxmlElement(f'w:{e}'); x.set(qn('w:val'),'none'); b.append(x)
    tblPr.append(b)
_TN=[0]
def table(titulo,headers,rows,fonte='Dados da pesquisa (2026).',fs=9):
    _TN[0]+=1
    cap=doc.add_paragraph(); cap.alignment=WD_ALIGN_PARAGRAPH.LEFT; cap.paragraph_format.space_before=Pt(6); cap.paragraph_format.line_spacing=1.0
    _set(cap.add_run(f'Tabela {_TN[0]} – {titulo}'),10)
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; _bd(t)
    for i,h in enumerate(headers):
        cl=t.rows[0].cells[i]; cl.paragraphs[0].line_spacing=1.0; _set(cl.paragraphs[0].add_run(str(h)),fs,b=True)
    for row in rows:
        cs=t.add_row().cells
        for i,v in enumerate(row):
            cs[i].paragraphs[0].line_spacing=1.0; _set(cs[i].paragraphs[0].add_run(str(v)),fs)
    fp=doc.add_paragraph(); fp.alignment=WD_ALIGN_PARAGRAPH.LEFT; fp.paragraph_format.space_after=Pt(6); fp.paragraph_format.line_spacing=1.0
    _set(fp.add_run(f'Fonte: {fonte}'),10)

# TÍTULO
tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
_set(tp.add_run('PERFIL DE HUMOR DE ATLETAS DE HANDEBOL NA ÚLTIMA SEMANA DA PRÉ-TEMPORADA COMPETITIVA: SINAL, RUÍDO E SENSIBILIDADE DOS MARCADORES'),14,b=True)
a2=doc.add_paragraph(); a2.alignment=WD_ALIGN_PARAGRAPH.CENTER; _set(a2.add_run('[Autoria e afiliação a preencher]'),11,it=True)

# RESUMO ESTRUTURADO
rp=doc.add_paragraph(); _set(rp.add_run('RESUMO'),12,b=True); rp.paragraph_format.line_spacing=1.0
def rlabel(lbl,txt):
    p=doc.add_paragraph(); p.paragraph_format.line_spacing=1.5; p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    _set(p.add_run(lbl+' '),12,b=True); _set(p.add_run(txt),12)
rlabel('Objetivo:','avaliar e classificar o perfil de humor de atletas de handebol na última semana da pré-temporada competitiva, identificando os marcadores mais sensíveis e separando o sinal do ruído de medida.')
rlabel('Métodos:','estudo descritivo de medidas repetidas com 27 atletas do sexo masculino, avaliados pelo BRUMS-24 e por escalas de fadiga física e mental duas vezes ao dia (pré e pós) durante sete dias (456 observações). Analisaram-se a comparação da linha de base (Dia 1, pré) ao último dia (Dia 7, pós), a decomposição da variância (sinal, traço e ruído), a discriminação por curva ROC com e sem filtragem do ruído e o comportamento dos perfis de humor.')
rlabel('Resultados:','da linha de base ao último dia, o vigor caiu 47% (dz = −1,02) e a fadiga física (+129%; dz = +2,12) e a do BRUMS (+116%; dz = +0,78) aumentaram, com deslocamento multivariado significativo do perfil (D de Mahalanobis = 2,02). A fadiga física foi o marcador mais sensível e o de melhor discriminação (AUC = 0,90 após filtragem). O sinal do microciclo representou apenas 1–13% da variância, e o perfil iceberg reduziu-se ao longo da semana, com inflexão da trajetória em torno do Dia 4.')
rlabel('Conclusão:','o custo do período é somático e acumulativo; a fadiga física e o vigor são os marcadores de eleição, e a decisão individual exige médias de coletas para superar o ruído de medida.')
kw=doc.add_paragraph(); _set(kw.add_run('Palavras-chave: '),12,b=True); _set(kw.add_run('Estado de humor. BRUMS. Fadiga. Relação sinal-ruído. Handebol.'),12); kw.paragraph_format.line_spacing=1.5

# 1 INTRODUÇÃO
sec('1','Introdução')
P('O monitoramento dos estados de humor é uma ferramenta prática e não invasiva para acompanhar a resposta psicobiológica de atletas à carga de treinamento, sinalizando desadaptações antes que se manifestem em queda de desempenho (TERRY; LANE; FOGARTY, 2003; ROHLFS et al., 2008). A Brunel Mood Scale (BRUMS) operacionaliza o construto em seis dimensões — tensão, depressão, raiva, vigor, fadiga e confusão — e resume o perfil na Perturbação Total do Humor (PTH). Atletas saudáveis tendem a exibir o perfil iceberg, com vigor destacando-se acima das dimensões negativas (MORGAN, 1985); seu achatamento indica fadiga acumulada.')
P('A última semana pré-competitiva é um período crítico, no qual a comissão técnica busca dissipar a fadiga preservando as adaptações. Contudo, boa parte da variação diária das medidas psicométricas não reflete a resposta ao treino, e sim erro de medida e flutuação aleatória — e interpretá-la como sinal gera decisões espúrias. Poucos estudos caracterizam de forma estruturada, e distinguindo sinal de ruído, o comportamento do humor nessa janela. O presente estudo teve por objetivo avaliar e classificar o perfil de humor de atletas de handebol na última semana da pré-temporada, identificar os marcadores mais sensíveis e quantificar quanto da oscilação é sinal do microciclo, traço estável e ruído.')

# 2 MÉTODOS
sec('2','Métodos')
ag=json.load(open(f'{SC}/app_data.json'))['socio']['agg']; pc=json.load(open(f'{SC}/app_data.json'))['socio']['pos_counts']
P(f"Participaram 27 atletas de handebol de elite do sexo masculino (idade {ag['idade'][0]:.0f} ± {ag['idade'][1]:.0f} anos; massa {ag['massa'][0]:.1f} ± {ag['massa'][1]:.1f} kg; {pc.get('Armador',0)} armadores, {pc.get('Ala',0)} alas, {pc.get('Pivô',0)} pivôs e {pc.get('Goleiro',0)} goleiros), monitorados durante os sete dias da última semana da pré-temporada, em delineamento descritivo de medidas repetidas, cada atleta como sua própria referência (dados anonimizados).")
P('O humor foi avaliado pelo BRUMS-24 (seis subescalas; PTH = negativas − vigor) e a fadiga física e mental por escalas de 0 a 10, autoaplicadas por formulário eletrônico com carimbo de data/hora duas vezes ao dia — a primeira resposta como “pré” e a última como “pós” —, totalizando 456 respostas válidas. A consistência interna foi estimada pelo α de Cronbach e pelo ω de McDonald.')
P('A variação da linha de base (Dia 1, pré) ao último dia (Dia 7, pós) e a resposta pré→pós foram avaliadas por Wilcoxon pareado com tamanho de efeito (dz); a magnitude multivariada, pela distância de Mahalanobis (Hotelling T²). A variância de cada variável foi decomposta (modelo atleta + dia) em sinal da semana, traço e ruído, derivando-se a relação sinal-ruído (SNR); a discriminação do último dia versus a linha de base foi quantificada pela área sob a curva ROC (AUC) com os dados brutos (com ruído) e com a média diária (sem ruído). A trajetória foi caracterizada por regressão, ajuste cúbico e ponto de inflexão. Adotou-se α = 0,05; software Python (pandas, SciPy, statsmodels). Análises complementares (decomposição de generalizabilidade, mudança confiável, componentes principais, correlação intra-atleta, análise fatorial confirmatória e ajuste alométrico) constam do material suplementar.')

# 3 RESULTADOS
sec('3','Resultados')
P('No conjunto da semana, o grupo preservou, em média, o perfil iceberg, com o vigor acima das subescalas negativas, mantidas próximas ao piso (Tabela 1). A consistência interna das subescalas foi adequada a boa (ω = 0,69–0,91), com exceção da tensão (α = 0,43; ω = 0,69), rebaixada pelo efeito de piso.')
rows=[[lab,c2(f"{desc[v]['mean']:.2f} ± {desc[v]['sd']:.2f}"),c2(f"{desc[v]['med']:.1f} [{desc[v]['iqr']:.1f}]"),f"{desc[v]['mn']:.0f}–{desc[v]['mx']:.0f}",(c2(f"{REL[k]['omega']:.2f}") if k else '—')] for v,lab,k in [('Vigor','Vigor','Vigor'),('Fadiga','Fadiga (BRUMS)','Fadiga'),('FadFisica','Fadiga física',None),('FadMental','Fadiga mental',None),('TMD','PTH/TMD',None),('Tensao','Tensão','Tensão'),('Depressao','Depressão','Depressão'),('Raiva','Raiva','Raiva'),('Confusao','Confusão','Confusão')]]
table('Estatística descritiva das variáveis de humor e fadiga na semana e confiabilidade (ω) das subescalas do BRUMS',['Variável','M ± DP','Mediana [IIQ]','Mín.–Máx.','ω'],rows)
P('Da linha de base ao último dia (Tabela 2), o vigor caiu de forma acentuada e a fadiga física, a fadiga do BRUMS, a fadiga mental e a PTH aumentaram; a fadiga física foi a variável mais responsiva (dz = +2,12), e o efeito crônico superou o agudo pré→pós (dz agudo da fadiga física = +1,08), indicando acúmulo. O deslocamento multivariado do perfil foi grande e significativo (Hotelling T²: F(4,17) = 18,17; p < 0,001; D de Mahalanobis = 2,02). As trajetórias (Figura 1) confirmam a deterioração concentrada no eixo energia–fadiga, com a fadiga mental e as subescalas negativas estáveis.')
rows=[[BF[v]['lab'],c2(f"{BF[v]['base']:.2f}"),c2(f"{BF[v]['fin']:.2f}"),f"{BF[v]['pct']:+.0f}%",c2(f"{BF[v]['dz']:+.2f}"),(c2(f"{BF[v]['p']:.3f}") if BF[v]['p']>=0.001 else '<0,001'),c2(f"{abs([r for r in rank if r['v']==v][0]['dz_ag']):.2f}")] for v,_ in CORE]
table('Comparação da linha de base (Dia 1, pré) ao último dia (Dia 7, pós): médias, variação, efeito crônico (dz) e efeito agudo pré→pós',['Variável','Pré D1','Pós D7','% mud.','dz crônico','p','|dz| agudo'],rows)
P('A classificação nos perfis mostrou a migração do iceberg para o de fadiga: o percentual em perfil iceberg e o índice-iceberg declinaram ao longo da semana, com precipitação no Dia 7 (Figura 2); a trajetória apresentou ponto de inflexão em torno do Dia 4, marcando a transição da acomodação para o aprofundamento da fadiga.')
fig('v_traj','Comportamento semanal das cinco variáveis (média ± desvio-padrão; suavização LOWESS pontilhada)',w=5.6)
fig('sem_f_iceberg','Perfil iceberg ao longo da semana: percentual em perfil iceberg (barras), índice-iceberg (linha) e regressão',w=5.4)
P('A decomposição da variância (Tabela 3) revelou que o sinal do microciclo é a menor fração (1–13%); a maior parte é traço estável e ruído, com SNR favorável apenas à fadiga física (1,82) e ao vigor (1,52). Na discriminação do último dia versus a linha de base, a fadiga física foi o melhor marcador (AUC = 0,85 com os dados brutos), e a filtragem do ruído (média das coletas do dia) elevou sua acurácia a 0,90, ao passo que variáveis sem sinal (fadiga mental, PTH) não se beneficiaram (Figura 3) — evidência direta de que a agregação de coletas aumenta o poder discriminativo apenas onde há sinal.')
rows=[[dec[v]['lab'],c2(f"{dec[v]['pday']:.0f}%"),c2(f"{dec[v]['ptrait']:.0f}%"),c2(f"{dec[v]['pnoise']:.0f}%"),c2(f"{dec[v]['snr']:.2f}"),c2(f"{ROC[v]['raw'][0]:.2f}"),c2(f"{ROC[v]['filt'][0]:.2f}")] for v,_ in CORE]
table('Decomposição da variância (sinal, traço e ruído), relação sinal-ruído (SNR) e discriminação por AUC (com e sem ruído)',['Variável','% Sinal','% Traço','% Ruído','SNR','AUC c/ ruído','AUC s/ ruído'],rows)
fig('x_roc','Curvas ROC (Dia 7 vs Dia 1) com ruído (cinza) e sem ruído (colorido, média diária filtrada)',w=5.4)

# 4 DISCUSSÃO
sec('4','Discussão')
P('Os achados mostram que, na última semana da pré-temporada, o humor se deteriora de forma coerente e concentrada no eixo energia–fadiga, com queda do vigor e elevação das fadigas, tanto na comparação da linha de base ao último dia quanto, em menor amplitude, dentro de cada dia. A fadiga física emergiu, em todas as análises, como o marcador mais precoce, sensível e discriminativo, seguida do vigor — enquanto a fadiga mental e as subescalas de afeto negativo, ancoradas no piso, mostraram-se pouco responsivas. Esse padrão indica que o custo do período é predominantemente somático, e não afetivo, coerente com a preservação do perfil sem colapso emocional.')
P('A principal contribuição metodológica é a separação explícita entre sinal e ruído. O sinal atribuível ao microciclo representou fração pequena da variância total (1–13%), sendo a maior parte diferenças estáveis entre atletas (traço) e erro de medida. Essa constatação tem consequência prática direta: a leitura de um único registro de um atleta é dominada pelo ruído, ao passo que a filtragem por médias de coletas recupera o sinal e eleva o poder discriminativo — como demonstrado pelo ganho de AUC da fadiga física (0,85 para 0,90) apenas nas variáveis que efetivamente carregam sinal. Recomenda-se, portanto, o monitoramento por tendência e por médias de coletas, com limiares individuais, em vez de valores absolutos isolados.')
P('A caracterização temporal acrescenta um marcador acionável: o ponto de inflexão em torno do Dia 4 delimita a transição da fase de acomodação para o aprofundamento da fadiga rumo ao Dia 7, indicando a janela oportuna para intervir na recuperação. A ampla heterogeneidade entre atletas — evidente na dispersão das trajetórias individuais — reforça a necessidade de referenciar cada indivíduo à sua própria linha de base. Como limitações, trata-se de estudo descritivo, sem grupo de comparação e restrito a um único microciclo de uma equipe masculina, o que circunscreve a generalização; a ausência de marcadores concomitantes de carga e desempenho impede inferências causais. Estudos futuros que articulem o monitoramento do humor a indicadores de carga e a desfechos competitivos poderão qualificar seu valor preditivo.')

# 5 CONCLUSÃO
sec('5','Conclusão')
P('Na última semana da pré-temporada, o perfil de humor deteriorou-se de forma concentrada no eixo energia–fadiga, com a fadiga física como marcador dominante e ponto de inflexão em torno do Dia 4. Como a maior parte da oscilação observada é traço e ruído, a decisão individual requer médias de coletas e limiares individualizados. Recomenda-se centrar o monitoramento na fadiga física e no vigor, com atenção reforçada a partir do meio da semana.')

sec('REFERÊNCIAS','')
for r in [
 'MORGAN, W. P. Selected psychological factors limiting performance: a mental health model. In: CLARKE, D. H.; ECKERT, H. M. (Ed.). Limits of human performance. Champaign: Human Kinetics, 1985. p. 70-80.',
 'ROHLFS, I. C. P. M. et al. A Escala de Humor de Brunel (Brums): instrumento para detecção precoce da síndrome do excesso de treinamento. Revista Brasileira de Medicina do Esporte, v. 14, n. 3, p. 176-181, 2008.',
 'TERRY, P. C.; LANE, A. M.; FOGARTY, G. J. Construct validity of the Profile of Mood States — Adolescents for use with adults. Psychology of Sport and Exercise, v. 4, n. 2, p. 125-139, 2003.',
]:
    p=doc.add_paragraph(); _set(p.add_run(r),12); p.paragraph_format.line_spacing=1.0; p.paragraph_format.space_after=Pt(6); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
doc.add_paragraph()
P('Material suplementar disponível: estatística descritiva completa e distribuições; regressão linear por variável; comportamento da PTH; limites, derivadas e ponto de inflexão; decomposição de generalizabilidade (traço/estado/ruído), mudança confiável (RCI) e componentes principais; modelo misto multivariado; correlação intra-atleta (rmcorr); análise fatorial confirmatória; e ajuste alométrico. Reprodutibilidade: scripts em Python; atletas anonimizados (A01–A27).',size=9,it=True,indent=False)

OUTP='/home/user/mdlucca/Artigos/Perfil_Humor_Periodico_Condensado.docx'
doc.save(OUTP); print('SAVED',OUTP)
d2=Document(OUTP); words=sum(len(p.text.split()) for p in d2.paragraphs)
print('Tabelas',_TN[0],'Figuras',_FN[0],'| palavras',words,'tables',len(d2.tables),'imgs',len(d2.inline_shapes))
