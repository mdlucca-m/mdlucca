# -*- coding: utf-8 -*-
import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
PV=json.load(open('pv_stats.json')); LIM=json.load(open('tcar_limiar.json')); ALLO=json.load(open('allo.json'))
DEN=json.load(open('denoise.json')); CRS=json.load(open('cross.json')); ROC=json.load(open('roc.json'))
def c2(s): return str(s).replace('.',',')
doc=Document()
stl=doc.styles['Normal']; stl.font.name='Times New Roman'; stl.font.size=Pt(12)
stl.element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')
stl.paragraph_format.line_spacing=1.5; stl.paragraph_format.space_after=Pt(0)
sec=doc.sections[0]; sec.top_margin=Cm(3); sec.left_margin=Cm(3); sec.bottom_margin=Cm(2); sec.right_margin=Cm(2)
def P(t='',just=True,size=12,after=6,bold=False,ind=True):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(size); r.bold=bold
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(after)
    if just: p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if just and ind: p.paragraph_format.first_line_indent=Cm(1.25)
    return p
def RUN(pairs,after=6,ind=True,just=True):
    p=doc.add_paragraph(); p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(after)
    if just: p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if ind: p.paragraph_format.first_line_indent=Cm(1.25)
    for txt,bold in pairs:
        r=p.add_run(txt); r.bold=bold; r.font.size=Pt(12)
    return p
def H(t,size=13,before=12,after=4):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(size)
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after); p.paragraph_format.line_spacing=1.5
def EQ(t):  # centered italic display equation
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(t); r.italic=True; r.font.size=Pt(12)
    p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(6)
def BUL(t,bold_lead=None):
    p=doc.add_paragraph(style='List Bullet');
    if bold_lead:
        r=p.add_run(bold_lead); r.bold=True; r.font.size=Pt(12)
    r=p.add_run(t); r.font.size=Pt(12); p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(3)
def _bd(c):
    tcPr=c._tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for e_ in ['top','bottom','left','right']:
        el=OxmlElement(f'w:{e_}'); on=e_ in ('top','bottom')
        el.set(qn('w:val'),'single' if on else 'nil'); el.set(qn('w:sz'),'6'); el.set(qn('w:color'),'000000'); b.append(el)
    tcPr.append(b)
def table(cap,header,rows,fs=9):
    p=doc.add_paragraph(); r=p.add_run(cap); r.bold=True; r.font.size=Pt(10.5)
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
    t=doc.add_table(rows=1,cols=len(header)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,htx in enumerate(header):
        cc=t.rows[0].cells[i]; cc.text=''; rr=cc.paragraphs[0].add_run(htx); rr.bold=True; rr.font.size=Pt(fs); rr.font.name='Times New Roman'
        cc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; _bd(cc)
    for row in rows:
        cs=t.add_row().cells
        for i,val in enumerate(row):
            cs[i].text=''; rr=cs[i].paragraphs[0].add_run(str(val)); rr.font.size=Pt(fs); rr.font.name='Times New Roman'
            cs[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT; _bd(cs[i])

# valores
pvm=c2('%.1f'%PV['desc']['pvini_m']); pvsd=c2('%.1f'%PV['desc']['pvini_sd'])
vig=PV['pv']['wk_Vigor']['TCAR1']; ff=PV['pv']['wk_FadFisica']['TCAR1']; fb=PV['pv']['wk_Fadiga']['TCAR1']
fit=ALLO['fitness']; trj=ALLO['FadFisica']; lp=LIM['LIM']['PVini']
relV=c2('%.2f'%DEN['rel']['Vigor']); relF=c2('%.2f'%DEN['rel']['Fadiga'])
etmV=c2('%.2f'%DEN['etm']['Vigor']); snrV=c2('%.2f'%DEN['snr_raw']['Vigor']); snrF=c2('%.2f'%DEN['snr_raw']['Fadiga'])
vcV=CRS['dec']['Vigor']['vc']; vcF=CRS['dec']['Fadiga']['vc']
mdc1=c2('%.1f'%DEN['kmdc']['Vigor']['1']); mdc2=c2('%.1f'%DEN['kmdc']['Vigor']['2'])
rocr=c2('%.2f'%ROC['FadFisica']['raw'][0]); rocf=c2('%.2f'%ROC['FadFisica']['filt'][0])

# ===== TÍTULO =====
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('NOTA METODOLÓGICA: RACIONAL CIENTÍFICO E ESTATÍSTICO DAS ESCOLHAS DE ANÁLISE'); r.bold=True; r.font.size=Pt(14)
p.paragraph_format.space_after=Pt(2)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Normalização pelo pico de velocidade, ajuste alométrico, derivadas e decomposição sinal–ruído'); r.italic=True; r.font.size=Pt(12)
p.paragraph_format.space_after=Pt(12)

P('Prezado(a) orientador(a), este documento explica, de forma didática e detalhada, por que cada uma das quatro escolhas '
 'analíticas foi adotada no estudo do perfil de humor dos handebolistas. Para cada uma, apresento: (i) o problema que ela '
 'resolve; (ii) o racional científico; (iii) a formalização estatística; (iv) o que encontramos nos nossos dados; e (v) a '
 'referência que a fundamenta. O fio condutor é único: sair da média do grupo e chegar a um sinal individual, confiável e '
 'acionável — separando o que é resposta real à carga do que é apenas diferença entre atletas, tamanho corporal ou ruído '
 'de medida.')

# ===== 0. O PROBLEMA CENTRAL =====
H('1. O problema central: três fontes de variação que se confundem')
P('Quando medimos o humor e a fadiga de 27 atletas ao longo de sete dias, a variação observada em qualquer escore mistura '
 'três coisas muito diferentes, que precisam ser separadas para uma leitura correta:')
BUL('a diferença estável entre atletas (uns são naturalmente mais vigorosos, outros mais fadigáveis).',bold_lead='Variação entre sujeitos — ')
BUL('a resposta real à carga do microciclo, comum ao grupo (o sinal que nos interessa).',bold_lead='Variação sistemática de dia — ')
BUL('a flutuação aleatória de uma medida isolada (humor do momento, erro de resposta).',bold_lead='Ruído de medida — ')
P('Se não separarmos essas fontes, corremos três riscos: (1) atribuir à “carga” o que é apenas um atleta menos apto; (2) '
 'atribuir à “aptidão” o que é apenas tamanho corporal; e (3) chamar de “mudança” o que é apenas ruído. As quatro '
 'escolhas metodológicas a seguir existem, cada uma, para neutralizar um desses riscos.')

# ===== 2. NORMALIZAÇÃO PELO PICO DE VELOCIDADE =====
H('2. Normalização pelo pico de velocidade (PV do T-CAR)')
RUN([('O problema. ',True),('A mesma sessão de treino não impõe o mesmo esforço relativo a todos. Um atleta com maior '
 'capacidade aeróbia intermitente recupera-se melhor entre os esforços de alta intensidade e, sob a mesma carga externa, '
 'sofre menor desgaste interno. Comparar a fadiga “crua” entre atletas, sem considerar a aptidão, confunde o efeito da '
 'carga com o efeito da capacidade física.',False)])
RUN([('O racional científico. ',True),('O handebol é uma modalidade intermitente de alta intensidade, na qual a aptidão '
 'aeróbia intermitente sustenta a reposição energética e a recuperação entre esforços. O pico de velocidade (PV) obtido '
 'no Teste de Carminatti (T-CAR) é um marcador de campo validado dessa capacidade, associado ao desempenho físico em '
 'partida (FERNANDES-DA-SILVA et al., 2016). Usá-lo como referência permite interpretar a resposta de humor–fadiga de '
 'cada atleta à luz da sua capacidade — ou seja, normalizar a leitura interindividual.',False)])
RUN([('A estatística. ',True),('Tratamos o PV como covariável fisiológica contínua e ajustamos uma reta de regressão '
 'entre o PV (T-CAR1, início da pré-temporada; %s ± %s km/h) e a resposta média semanal de cada dimensão. O coeficiente '
 'angular (β) diz quanto o desfecho muda por km/h de aptidão; o R² e o ρ de Spearman dizem quão forte é a relação. Um '
 'gráfico de dispersão com reta e banda de confiança de 95%% torna a relação visual e honesta quanto à incerteza.'%(pvm,pvsd),False)])
RUN([('O que encontramos. ',True),('Atletas mais aptos reportaram mais vigor (ρ = %s; p = %s) e menos fadiga física '
 '(ρ = %s; R² = %s; p = %s), com tendência semelhante para a fadiga do BRUMS (ρ = %s; p = %s). Estratificando em tercis '
 'de aptidão, a fadiga cresce dos mais aptos para os menos aptos (Kruskal-Wallis p = %s). Além disso, um limiar de '
 'PV ≈ %s km/h (índice de Youden) discriminou os dias de maior fadiga (AUC = %s; IC95%% %s–%s; sensibilidade %s; '
 'especificidade %s) — uma referência prática para sinalizar quem tende a acumular mais fadiga.'%(
   c2('%+.2f'%vig['rho']),c2('%.3f'%vig['rho_p']),c2('%+.2f'%ff['rho']),c2('%.2f'%ff['r2']),c2('%.3f'%ff['rho_p']),
   c2('%+.2f'%fb['rho']),c2('%.3f'%fb['rho_p']),c2('%.3f'%PV['terc_kruskal_fadfis']['p']),
   c2('%.1f'%lp['thr']),c2('%.2f'%lp['auc']),c2('%.2f'%lp['lo']),c2('%.2f'%lp['hi']),c2('%.2f'%lp['sens']),c2('%.2f'%lp['spec'])),False)])
RUN([('Em uma frase. ',True),('Normalizar pelo PV separa a “fadiga esperada, dado o nível de aptidão” da “fadiga '
 'excessiva”, que é o verdadeiro sinal de alerta para a comissão técnica.',False)])

# ===== 3. AJUSTE ALOMÉTRICO =====
H('3. Ajuste alométrico')
RUN([('O problema. ',True),('Uma tentação comum é “normalizar” dividindo pela massa corporal (por exemplo, PV por kg). '
 'Isso pressupõe que a relação entre as variáveis é uma reta que passa pela origem — o que quase nunca é verdade em '
 'variáveis biológicas. Essa razão simples introduz um viés: em vez de remover o efeito do tamanho, ela o inverte, '
 'penalizando os mais pesados. É a chamada “falácia da razão”.',False)])
RUN([('O racional científico. ',True),('Variáveis fisiológicas escalam com o tamanho corporal segundo uma lei de '
 'potência (Y = a · massa^b), não de forma proporcional. O ajuste alométrico estima o expoente b por regressão linear no '
 'espaço log–log (ln Y = ln a + b · ln massa) e usa esse b para tornar a variável independente do tamanho, de modo '
 'estatisticamente correto (NEVILL; LANE, 2007). Aplicamos a lógica alométrica em dois pontos.',False)])
RUN([('(a) Forma da trajetória no tempo. ',True),('Ajustamos a trajetória diária da fadiga como Y = a · dia^b. O expoente '
 'estimado foi b = %s (R² = %s) para a fadiga física — como b < 1, a curva é saturante: a fadiga sobe rápido no início do '
 'microciclo e desacelera ao longo da semana. Isso descreve o curso temporal de forma mais fiel do que uma reta e informa '
 'quando a acomodação fisiológica ocorre.'%(c2('%.2f'%trj['b']),c2('%.2f'%trj['r2'])),False)])
EQ('Y = a · dia^b    ⇔    ln(Y) = ln(a) + b · ln(dia)     (b < 1 ⇒ processo saturante)')
RUN([('(b) Remoção do tamanho corporal da relação aptidão–fadiga. ',True),('Estimamos PV = a · massa^b e obtivemos '
 'b = %s. Ao normalizar o PV pela massa com esse expoente, a correlação entre aptidão e fadiga física caiu de ρ = %s '
 '(p = %s) para ρ = %s (p = %s). A leitura honesta: parte da “proteção” contra a fadiga que atribuiríamos à aptidão é, na '
 'verdade, efeito do tamanho corporal. Sem o ajuste alométrico, superestimaríamos o papel da capacidade física.'%(
   c2('%.2f'%fit['b']),c2('%+.2f'%fit['rho_raw']),c2('%.3f'%fit['p_raw']),c2('%+.2f'%fit['rho_allo']),c2('%.3f'%fit['p_allo'])),False)])
RUN([('Nota importante. ',True),('O mesmo Nevill & Lane (2007) que fundamenta o ajuste alométrico das variáveis '
 'fisiológicas contínuas também alerta que dados Likert (como as subescalas do BRUMS) não devem ser log-transformados. '
 'Por isso fomos seletivos: aplicamos a lei de potência às variáveis físicas/contínuas e mantivemos as subescalas do '
 'humor em sua escala original, tratadas por estatística não paramétrica.',False)])

# ===== 4. DERIVADAS =====
H('4. Derivadas (velocidade e aceleração da mudança)')
RUN([('O problema. ',True),('A trajetória média diz onde os escores estão em cada dia, mas não diz quando a mudança é '
 'mais rápida. Para a comissão técnica, saber em qual dia o vigor despenca ou a fadiga dispara é mais acionável do que '
 'saber apenas o valor médio — é a diferença entre observar o nível e observar a velocidade.',False)])
RUN([('O racional científico. ',True),('Em monitoramento de carga, a taxa de variação (o quão rápido) sinaliza os '
 'momentos de maior estresse e os pontos ideais para intervir na recuperação. A primeira derivada da trajetória é a '
 'velocidade da mudança (pontos por dia); a segunda derivada é a aceleração (se a mudança está ganhando ou perdendo '
 'força).',False)])
RUN([('A estatística. ',True),('Ajustamos um polinômio às médias diárias e derivamos analiticamente. A velocidade indica '
 'o dia de maior taxa de variação; a aceleração indica inflexões (quando a subida da fadiga começa a saturar, ligando-se '
 'ao achado alométrico da seção anterior).',False)])
RUN([('O que encontramos. ',True),('A maior taxa de variação ocorre logo no Dia 1 (o choque de carga): o vigor cai a '
 '%s pontos/dia e a fadiga sobe a +%s pontos/dia. Ou seja, o momento mais crítico do microciclo, do ponto de vista da '
 'velocidade de deterioração, é a transição inicial — informação que uma simples média mascararia.'%(
   c2('%.2f'%CRS['dec']['Vigor']['pkvel_val']),c2('%.2f'%CRS['dec']['Fadiga']['pkvel_val'])),False)])

# ===== 5. RUÍDOS, FILTROS E SINAIS =====
H('5. Ruídos, filtros e sinais (decomposição sinal–ruído)')
RUN([('O problema. ',True),('Uma única resposta ao BRUMS não é o “humor verdadeiro” do atleta: é o humor verdadeiro mais '
 'um ruído (o estado do momento, um erro de resposta). Se não medirmos esse ruído, não sabemos se uma mudança observada '
 'é real ou apenas flutuação — e podemos agir sobre um falso alarme.',False)])
RUN([('O racional científico. ',True),('A ciência do esporte trata isso pelo arcabouço da confiabilidade e do erro típico '
 'de medida (HOPKINS, 2000): estima-se o ruído, e só se declara “mudança” o que o supera. É a mesma lógica que sustenta, '
 'no nível individual, o Índice de Mudança Confiável (RCI) que usamos.',False)])
RUN([('A estatística. ',True),('Decompusemos a variância de cada variável em três partes (atleta, dia e resíduo) por '
 'ANOVA. Dessa decomposição extraímos: a confiabilidade de uma medida (r = 1 − fração de ruído); o erro típico de medida '
 '(ETM); a amplitude do sinal (variação da média diária); a razão sinal–ruído (RSR = sinal/ETM); e a mínima mudança '
 'detectável, que diminui com o número de coletas (k):',False)])
EQ('MDC95 = 1,96 · √2 · ETM / √k')
RUN([('Além disso, ',False),('comparamos a discriminação entre o Dia 7 e o Dia 1 usando a curva ROC em dois cenários: '
 'com ruído (observações individuais) e sem ruído (médias diárias, que “filtram” o ruído ao agregá-lo). Se a área sob a '
 'curva (AUC) sobe ao filtrar, é porque havia sinal escondido sob o ruído.',False)])
RUN([('O que encontramos. ',True),('A confiabilidade de uma única medida foi moderada (r = %s no vigor; %s na fadiga) e a '
 'razão sinal–ruído favoreceu o eixo energia–fadiga (RSR = %s no vigor; %s na fadiga). A decomposição mostrou que a maior '
 'parte da variância é entre atletas (vigor: %s%% atleta, %s%% dia, %s%% resíduo), o que explica por que o monitoramento '
 'deve ser individualizado. Ao filtrar o ruído, a discriminação da fadiga física entre o início e o fim do microciclo '
 'subiu de AUC = %s para %s. E a mínima mudança detectável do vigor caiu de %s pontos (1 coleta) para %s pontos (2 '
 'coletas) — este é o motivo estatístico direto de termos coletado o BRUMS duas vezes ao dia: reduzir o ruído por um '
 'fator de √2 e tornar detectáveis mudanças menores.'%(
   relV,relF,snrV,snrF,c2('%.0f'%vcV['ath']),c2('%.0f'%vcV['day']),c2('%.0f'%vcV['res']),rocr,rocf,mdc1,mdc2),False)])
RUN([('Ressalva honesta. ',True),('Filtrar ruído não cria sinal onde não há. As subescalas negativas (tensão, depressão, '
 'raiva, confusão) permanecem próximas do piso mesmo após a filtragem — por isso concentramos a interpretação no eixo '
 'energia–fadiga, o único com sinal semanal aproveitável.',False)])

# ===== 6. SÍNTESE =====
H('6. Como as quatro escolhas se encaixam')
P('As quatro escolhas não são independentes: elas formam uma cadeia lógica que limpa progressivamente o sinal. Primeiro, '
 'a decomposição sinal–ruído garante que a mudança que observamos é real e não flutuação. Segundo, as derivadas dizem '
 'quando essa mudança real é mais intensa. Terceiro, a normalização pelo pico de velocidade explica por que alguns '
 'atletas respondem mais do que outros, à luz da aptidão. Quarto, o ajuste alométrico garante que essa explicação pela '
 'aptidão não seja, na verdade, um artefato do tamanho corporal. O resultado é uma leitura que sai da média do grupo e '
 'chega a um sinal individual, temporalmente localizado, fisiologicamente contextualizado e estatisticamente confiável.')

table('Quadro-resumo das escolhas metodológicas.',
    ['Escolha','Pergunta que responde','Método','Resultado no estudo'],
    [['Normalização pelo PV','A resposta depende da aptidão?','Regressão, tercis, limiar de Youden','Mais aptidão → mais vigor (ρ = %s) e menos fadiga (ρ = %s); limiar %s km/h'%(c2('%+.2f'%vig['rho']),c2('%+.2f'%ff['rho']),c2('%.1f'%lp['thr']))],
     ['Ajuste alométrico','O efeito é da aptidão ou do tamanho?','Lei de potência (log–log)','Trajetória saturante (b = %s); ρ aptidão–fadiga cai de %s para %s'%(c2('%.2f'%trj['b']),c2('%+.2f'%fit['rho_raw']),c2('%+.2f'%fit['rho_allo']))],
     ['Derivadas','Quando a mudança é mais rápida?','Derivada da trajetória','Maior taxa no Dia 1: vigor %s/dia, fadiga +%s/dia'%(c2('%.2f'%CRS['dec']['Vigor']['pkvel_val']),c2('%.2f'%CRS['dec']['Fadiga']['pkvel_val']))],
     ['Sinal–ruído','A mudança é real ou ruído?','Decomposição da variância, ETM, MDC, ROC','RSR favorece vigor/fadiga; AUC sobe %s→%s ao filtrar; 2 coletas/dia reduzem o MDC'%(rocr,rocf)]],fs=9)

# ===== REFERÊNCIAS =====
H('Referências')
for rf in [
 'FERNANDES-DA-SILVA, J. et al. The peak velocity derived from the Carminatti Test is related to physical match performance in young soccer players. Journal of Sports Sciences, v. 34, n. 24, p. 2238–2245, 2016.',
 'HOPKINS, W. G. Measures of reliability in sports medicine and science. Sports Medicine, v. 30, n. 1, p. 1–15, 2000.',
 'JACOBSON, N. S.; TRUAX, P. Clinical significance: a statistical approach to defining meaningful change in psychotherapy research. Journal of Consulting and Clinical Psychology, v. 59, n. 1, p. 12–19, 1991.',
 'NEVILL, A. M.; LANE, A. M. Why self-report “Likert” scale data should not be log-transformed. Journal of Sports Sciences, v. 25, n. 1, p. 1–2, 2007.']:
    p=doc.add_paragraph(); r=p.add_run(rf); r.font.size=Pt(11); p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(6); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

OUTP='/home/user/mdlucca/Artigos/Nota_Metodologica_Racional_Analises.docx'
doc.save(OUTP); print('SAVED',OUTP)
