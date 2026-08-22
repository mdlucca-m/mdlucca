# -*- coding: utf-8 -*-
# Tabela descritiva: dias de treinamento do microciclo + carga interna + classificacao dos perfis de humor.
import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

carga=json.load(open('/home/user/mdlucca/scripts/analise/hiit_carga_dia.json'))
perf=json.load(open('/home/user/mdlucca/scripts/analise/perfil_por_dia.json'))
PROF=['Iceberg','Superfície','Submerso','Barbatana de tubarão','Everest invertido','Iceberg invertido']
DATAS={1:'21/04',2:'22/04',3:'23/04',4:'24/04',5:'25/04',6:'26/04',7:'27/04'}
TIPO={1:'Técnico-tática',2:'HIIT (TIAI)',3:'Técnico-tática',4:'HIIT (TIAI)',5:'Técnico-tática',6:'Técnico-tática',7:'HIIT (TIAI)'}

doc=Document()
st=doc.styles['Normal']; st.font.name='Times New Roman'; st.font.size=Pt(11)
def cap(txt):
    p=doc.add_paragraph(); r=p.add_run(txt); r.bold=True; r.font.size=Pt(10.5); r.font.name='Times New Roman'
    p.paragraph_format.space_after=Pt(3); return p
def note(txt):
    p=doc.add_paragraph(); r=p.add_run(txt); r.italic=True; r.font.size=Pt(9); r.font.name='Times New Roman'
    p.paragraph_format.space_after=Pt(12); return p
def setcell(c,txt,bold=False,align='center',size=9.5):
    c.text=''; p=c.paragraphs[0]; p.alignment={'center':WD_ALIGN_PARAGRAPH.CENTER,'left':WD_ALIGN_PARAGRAPH.LEFT}[align]
    r=p.add_run(str(txt)); r.font.size=Pt(size); r.font.name='Times New Roman'; r.bold=bold

# ---------------- Tabela 1 ----------------
cap('Tabela 1. Caracterização do microciclo de pré-temporada: sessões de treinamento, carga interna e perfil de humor predominante ao longo dos sete dias.')
hdr=['Dia','Data','Sessão','%FCmáx','TRIMP (u.a.)','PSE (0–10)','sPSE (u.a.)','Perfil predominante','Iceberg (%)']
t=doc.add_table(rows=1,cols=len(hdr)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
for j,hh in enumerate(hdr): setcell(t.rows[0].cells[j],hh,bold=True)
for dia in range(1,8):
    c=t.add_row().cells
    cg=carga.get(str(dia))
    load=[f"{cg['pFC']:.1f}".replace('.',','),f"{cg['TRIMP']:.1f}".replace('.',','),
          f"{cg['PSE']:.1f}".replace('.',','),f"{cg['sPSE']:.1f}".replace('.',',')] if cg else ['—','—','—','—']
    pr=perf[str(dia)]
    vals=[f'D{dia}',DATAS[dia],TIPO[dia]]+load+[pr['pred'],str(pr['ice'])]
    for j,v in enumerate(vals): setcell(c[j],v,align='left' if j in(2,7) else 'center')
note('HIIT (TIAI): treinamento intervalado de alta intensidade. %FCmáx: percentual da '
     'frequência cardíaca máxima; TRIMP: impulso de treino (Banister); PSE: percepção subjetiva de esforço da sessão '
     '(0–10); sPSE: PSE da sessão × duração (u.a.). Sessões de HIIT em D2, D4 e D7. A carga interna foi quantificada apenas nas sessões de HIIT; nos '
     'dias técnico-táticos (—) não houve monitoramento fisiológico. Perfil predominante: perfil de humor mais '
     'frequente entre os atletas no dia, classificado nos seis perfis de Terry/Parsons-Smith a partir das seis '
     'dimensões do BRUMS. Iceberg (%): proporção de atletas com perfil iceberg no dia. Note a assinatura de fadiga '
     'acumulada nas sessões de HIIT: o %FCmáx e o TRIMP declinam (84,0 → 75,2) enquanto a PSE aumenta (8,3 → 8,9).')

# ---------------- Tabela 2 ----------------
cap('Tabela 2. Distribuição dos seis perfis de humor dos atletas ao longo dos sete dias do microciclo (número de atletas em cada perfil).')
hdr2=['Perfil']+[f'D{d}' for d in range(1,8)]
t2=doc.add_table(rows=1,cols=len(hdr2)); t2.style='Table Grid'; t2.alignment=WD_TABLE_ALIGNMENT.CENTER
for j,hh in enumerate(hdr2): setcell(t2.rows[0].cells[j],hh,bold=True)
for pf in PROF:
    c=t2.add_row().cells; setcell(c[0],pf,align='left')
    for d in range(1,8): setcell(c[d],perf[str(d)]['dist'][pf])
# linha n
c=t2.add_row().cells; setcell(c[0],'n (atletas)',bold=True,align='left')
for d in range(1,8): setcell(c[d],perf[str(d)]['n'],bold=True)
note('Perfis de humor classificados a partir das seis dimensões do BRUMS (tensão, depressão, raiva, vigor, fadiga e '
     'confusão) segundo o modelo de seis perfis de Terry/Parsons-Smith: iceberg (perfil de bem-estar, vigor elevado e '
     'dimensões negativas baixas); superfície (perfil achatado, todas as dimensões próximas da média); submerso; '
     'barbatana de tubarão (vigor e fadiga elevados); Everest invertido; e iceberg invertido. Observa-se a migração '
     'do perfil iceberg (41% no D1) para perfis achatados/de fadiga ao fim da semana (iceberg em apenas 10% no D7, '
     'com predomínio do perfil superfície).')

doc.save('/home/user/mdlucca/Artigos/Tabela_Descritiva_Microciclo.docx')
print('OK Tabela_Descritiva_Microciclo.docx (2 tabelas)')
