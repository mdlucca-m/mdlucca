# -*- coding: utf-8 -*-
import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
R=json.load(open('brums_desc2.json')); RC=json.load(open('rci6.json')); STAT=json.load(open('brums_stats3.json'))
S4=json.load(open('brums_stats4.json')); MS=json.load(open('model_stats.json')); MV=json.load(open('manova.json'))
PHJ=json.load(open('posthoc.json')); ROC=json.load(open('roc.json')); DEN=json.load(open('denoise.json'))
CRS=json.load(open('cross.json')); ALLO=json.load(open('allo.json'))
FG='/home/user/mdlucca/Artigos/figuras'
def c2(s): return str(s).replace('.',',')
doc=Document()
st=doc.styles['Normal']; st.font.name='Times New Roman'; st.font.size=Pt(12)
st.element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')
st.paragraph_format.line_spacing=1.5; st.paragraph_format.space_after=Pt(0)
s=doc.sections[0]; s.top_margin=Cm(3); s.left_margin=Cm(3); s.bottom_margin=Cm(2); s.right_margin=Cm(2)
_TN=[0]; _FN=[0]
def P(t='',just=True,size=12,after=6,bold=False,ind=True):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(size); r.bold=bold
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(after)
    if just: p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if just and ind: p.paragraph_format.first_line_indent=Cm(1.25)
    return p
def RUN(pairs,after=6,ind=True):
    p=doc.add_paragraph(); p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(after)
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if ind: p.paragraph_format.first_line_indent=Cm(1.25)
    for txt,bold in pairs:
        r=p.add_run(txt); r.bold=bold; r.font.size=Pt(12)
    return p
def H(t,size=12,before=10):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(size)
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.5
def _bd(c):
    tcPr=c._tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for e_ in ['top','bottom','left','right']:
        el=OxmlElement(f'w:{e_}'); on=e_ in ('top','bottom')
        el.set(qn('w:val'),'single' if on else 'nil'); el.set(qn('w:sz'),'6'); el.set(qn('w:color'),'000000'); b.append(el)
    tcPr.append(b)
def table(cap,header,rows,fonte='Fonte: dados da pesquisa (2026).',fs=9,note=None):
    _TN[0]+=1
    p=doc.add_paragraph(); r=p.add_run('Tabela %d – %s'%(_TN[0],cap)); r.font.size=Pt(11)
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
    t=doc.add_table(rows=1,cols=len(header)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,htx in enumerate(header):
        cc=t.rows[0].cells[i]; cc.text=''; rr=cc.paragraphs[0].add_run(htx); rr.bold=True; rr.font.size=Pt(fs); rr.font.name='Times New Roman'
        cc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; _bd(cc)
    for row in rows:
        cs=t.add_row().cells
        for i,val in enumerate(row):
            cs[i].text=''; rr=cs[i].paragraphs[0].add_run(str(val)); rr.font.size=Pt(fs); rr.font.name='Times New Roman'
            cs[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if i>0 else WD_ALIGN_PARAGRAPH.LEFT; _bd(cs[i])
    if note:
        pn=doc.add_paragraph(); rn=pn.add_run(note); rn.font.size=Pt(8.5); rn.italic=True; pn.paragraph_format.space_after=Pt(0)
    pf=doc.add_paragraph(); rf=pf.add_run(fonte); rf.font.size=Pt(9); pf.paragraph_format.space_after=Pt(6)
    return _TN[0]
def figure(path,cap,w=15.5):
    _FN[0]+=1
    pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; pp.add_run().add_picture(path,width=Cm(w)); pp.paragraph_format.space_before=Pt(6)
    pc=doc.add_paragraph(); pc.alignment=WD_ALIGN_PARAGRAPH.CENTER; rc=pc.add_run('Figura %d – %s'%(_FN[0],cap)); rc.font.size=Pt(11)
    pf=doc.add_paragraph(); pf.alignment=WD_ALIGN_PARAGRAPH.CENTER; rf=pf.add_run('Fonte: elaboração dos autores (2026).'); rf.font.size=Pt(9); pf.paragraph_format.space_after=Pt(6)
    return _FN[0]
def pstr(p): return '< 0,001' if p<0.001 else c2('%.3f'%p)
mvv=lambda tab,k,f: next(x[f] for x in MV[tab]['rows'] if x['k']==k)

pr=R['prepos']; d17=R['d1d7']; prof=R['profiles']; sm=R['sample']; foc=R['focus']; npk=R['neg_peak']; desc=R['desc']
CV=S4['cv']; FR=S4['friedman']; SENS=S4['sens']; PCT=S4['pct']; PT=S4['prof_trans']; PCNT=S4['prof_counts']
MSd=MS['desc']; MScorr=MS['corr']; MSord=MS['order']; PREV=MS['prev']
ORDER=[('Vigor','Vigor'),('Fadiga','Fadiga'),('TMD','Perturbação Total do Humor (PTH)'),('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Confusao','Confusão')]

# ===================== TÍTULO + RESUMO =====================
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('PERFIL DE HUMOR DE ATLETAS DE HANDEBOL DE ELITE NA ÚLTIMA SEMANA DE PRÉ-TEMPORADA: MONITORAMENTO DAS SEIS DIMENSÕES DO BRUMS DO GRUPO AO ATLETA')
r.bold=True; r.font.size=Pt(13); p.paragraph_format.space_after=Pt(10)
H('RESUMO',before=2)
RUN([('O monitoramento do humor é um indicador prático do bem-estar e da prontidão do atleta. ',False),
 ('O objetivo ',True),('deste estudo foi avaliar o perfil de humor de atletas de handebol de elite na última semana de '
 'pré-temporada, caracterizando e comparando as seis dimensões do BRUMS, do nível do grupo ao do atleta. Vinte e sete '
 'atletas do sexo masculino responderam ao BRUMS-24 duas vezes ao dia (pré e pós-treino) durante sete dias, totalizando '
 '%d observações. Os escores foram convertidos em escores T (M = 50; DP = 10) e as comparações intra-grupo (Dia 1 vs. Dia '
 '7; pré vs. pós) foram testadas por MANOVA de medidas repetidas, complementada por estatística não paramétrica '
 '(Shapiro-Wilk, Friedman, Wilcoxon, ρ de Spearman), pós-teste do modelo misto, coeficiente de correlação intraclasse '
 '(ICC), Índice de Mudança Confiável (RCI) individual, classificação nos seis perfis de humor e decomposição sinal–ruído. '%sm['n_obs'],False),
 ('Resultados: ',True),('a comparação Dia 1 → Dia 7 foi multivariadamente significativa (Wilks λ = %s; F(%d,%d) = %s; p = '
 '%s; η²ₚ = %s), com deterioração concentrada no eixo energia–fadiga — vigor (η²ₚ = %s; d = %s) e fadiga (η²ₚ = %s; d = '
 '%s). O vigor cai de forma mais acentuada no início do microciclo e a fadiga sobe mais ao final, cruzando-se no Dia 2 e '
 'invertendo-se no Dia 7. A prevalência de perfis mudou (χ² = %s; p = %s): o iceberg caiu e o perfil de fadiga subiu. No '
 'nível individual, a fadiga foi a única dimensão a piorar de forma confiável na maioria (%d%%). '%(
   c2('%.3f'%MV['d1d7']['wilks']),MV['d1d7']['df1'],MV['d1d7']['df2'],c2('%.2f'%MV['d1d7']['Fmv']),pstr(MV['d1d7']['p_mv']),c2('%.2f'%MV['d1d7']['eta_mv']),
   c2('%.2f'%mvv('d1d7','Vigor','eta')),c2('%+.2f'%mvv('d1d7','Vigor','d')),c2('%.2f'%mvv('d1d7','Fadiga','eta')),c2('%+.2f'%mvv('d1d7','Fadiga','d')),
   c2('%.2f'%PREV['chi']),pstr(PREV['p']),RC['Fadiga']['pct_piora']),False),
 ('Conclusão: ',True),('o perfil de humor migrou da prontidão (iceberg) para a fadiga funcional; o vigor e a fadiga são as '
 'dimensões mais sensíveis, o que fundamenta um monitoramento individualizado e referenciado à linha de base.',False)],after=6)
P('Palavras-chave: humor; BRUMS; handebol; monitoramento do atleta; fadiga; bem-estar psicológico.',size=11,after=8,ind=False)

# ===================== 1 INTRODUÇÃO =====================
H('1 INTRODUÇÃO')
H('1.1 Monitoramento do humor no esporte de elite',12,before=6)
P('As organizações esportivas de alto rendimento demonstram, cada vez mais, o compromisso de proteger o bem-estar '
 'psicológico dos atletas ao mesmo tempo em que buscam maximizar o desempenho, sendo comum o uso de algum indicador '
 'psicológico para rastrear o bem-estar e o risco de problemas de saúde mental (ROHLFS et al., 2023). Avaliações regulares '
 'do humor têm demonstrado, ao longo de décadas, utilidade preditiva tanto para o bem-estar quanto para o desempenho '
 'esportivo (SAW; MAIN; GASTIN, 2016; LOCHBAUM et al., 2021). O humor é definido como um conjunto de sentimentos, de '
 'natureza efêmera, que variam em intensidade e duração e geralmente envolvem mais de uma emoção; diferentemente das '
 'emoções, os estados de humor são menos intensos, mais duradouros e nem sempre possuem um gatilho específico, variando '
 'em valência e ativação (DE MIRANDA ROHLFS et al., 2025).')
P('O instrumento mais utilizado para essa finalidade é a Brunel Mood Scale (BRUMS), derivada abreviada do Profile of Mood '
 'States (POMS), que mensura seis dimensões — tensão, depressão, raiva, vigor, fadiga e confusão — em 24 itens de rápida '
 'aplicação e reaplicação, com sólidas propriedades psicométricas e sucessivas validações transculturais, incluindo a '
 'versão brasileira (ROHLFS et al., 2008; TERRY; LANE; FOGARTY, 2003; ROHLFS et al., 2023; LEW et al., 2023). O '
 'perfilamento do humor — em que os escores do indivíduo são comparados a dados normativos para gerar um perfil gráfico — '
 'é amplamente utilizado como triagem de bem-estar e prontidão (DE MIRANDA ROHLFS et al., 2025).')
H('1.2 Perfis de humor e a lacuna do monitoramento em microciclos',12,before=6)
P('Para além dos escores isolados, Morgan (1985) descreveu o “perfil iceberg” — vigor elevado sobre dimensões negativas '
 'baixas — como assinatura de prontidão. Abordagens recentes formalizaram seis perfis prototípicos de humor, replicados '
 'em grandes amostras internacionais: iceberg, Everest invertido, iceberg invertido, submerso, barbatana de tubarão e '
 'superfície (PARSONS-SMITH; TERRY; MACHIN, 2017; HAN; PARSONS-SMITH; TERRY, 2020; TERRY et al., 2021). O handebol de '
 'quadra é uma modalidade coletiva intermitente e de alta intensidade, cujo microciclo pré-competitivo concentra a carga '
 'que antecede a competição (KARCHER; BUCHHEIT, 2014; MICHALSIK; MADSEN; AAGAARD, 2013). Apesar da ampla adoção da BRUMS, '
 'faltam descrições detalhadas — dimensão a dimensão e do grupo ao sujeito — do comportamento do humor nesse período, '
 'inclusive de quando e onde, ao longo dos dias e dentro do dia, o vigor cai e a fadiga sobe.')
H('1.3 Objetivos e hipóteses',12,before=6)
P('O objetivo geral foi avaliar o perfil de humor de atletas de handebol de elite na última semana de pré-temporada, '
 'caracterizando e comparando o comportamento das seis dimensões do BRUMS e da Perturbação Total do Humor (PTH), do nível '
 'do grupo ao nível do sujeito. Como objetivos específicos, buscou-se: (i) caracterizar exploratoriamente cada dimensão; '
 '(ii) comparar, no grupo, a resposta Dia 1 → Dia 7 e pré → pós em escores T, identificando as dimensões mais sensíveis e '
 'os dias e momentos de maior variação; (iii) analisar a estrutura de correlação entre as dimensões; (iv) descrever a '
 'prevalência e a transição dos perfis de humor; e (v) quantificar, no sujeito, a mudança confiável (RCI) e a razão '
 'sinal–ruído das medidas.')
P('Formularam-se quatro hipóteses: (H1) a deterioração concentra-se no eixo energia–fadiga (vigor ↓, fadiga ↑), com '
 'magnitude média a grande, enquanto as dimensões negativas de valência não fadiga permanecem estáveis por efeito de '
 'piso; (H2) o perfil de humor migra do iceberg para o perfil de fadiga (barbatana de tubarão); (H3) o afeto negativo '
 'acopla-se à fadiga somente sob carga acumulada; e (H4) a resposta é heterogênea entre atletas, exigindo interpretação '
 'individualizada.')

# ===================== 2 MÉTODOS =====================
H('2 MATERIAIS E MÉTODOS')
H('2.1 Participantes',12,before=6)
P('Participaram %d atletas de handebol do sexo masculino de nível competitivo, avaliados em condições ecológicas de '
 'treinamento durante o microciclo pré-competitivo (21–27 de abril de 2024). A distribuição por posição consta na Tabela '
 '1 e a caracterização antropométrica na Tabela 2. Todos os procedimentos seguiram a Declaração de Helsinque, com '
 'consentimento informado.'%sm['n'])
t_dist=table('Distribuição demográfica e situacional da amostra (n = %d).'%sm['n'],
    ['Fonte','Grupo','n','%'],
    [['Sexo','Masculino',sm['n'],'100,0']]+
    [['Posição' if i==0 else '',k,v,c2('%.1f'%(100*v/sm['n']))] for i,(k,v) in enumerate(sm['pos'].items())]+
    [['Total','Todos',sm['n'],'100,0']],fs=9)
def srow(lab,k):
    v=sm[k]; return [lab,c2('%.1f'%v['mean']),c2('%.1f'%v['sd']),'%s–%s'%(c2('%.1f'%v['lo']),c2('%.1f'%v['hi'])),'%s–%s'%(c2('%.1f'%v['mn']),c2('%.1f'%v['mx']))]
t_anth=table('Caracterização antropométrica e de experiência da amostra (n = %d).'%sm['n'],
    ['Variável','Média','DP','IC95%','Mín–Máx'],
    [srow('Idade (anos)','idade'),srow('Estatura (cm)','estatura'),srow('Massa corporal (kg)','massa'),srow('Percentual de gordura','pG'),srow('Experiência na modalidade (anos)','exp')])
H('2.2 Medida do humor',12,before=6)
P('O humor foi avaliado pela BRUMS-24: 24 itens em escala Likert de 0 (“nada”) a 4 (“extremamente”), em seis subescalas '
 'de 0 a 16 pontos (tensão, depressão, raiva, vigor, fadiga e confusão); escores mais altos indicam maior nível da '
 'dimensão. A PTH resume o perfil: PTH = tensão + depressão + raiva + fadiga + confusão − vigor. A consistência interna '
 '(alfa de Cronbach) variou de %s a %s (Tabela 3). Para a interpretação normativa e a classificação de perfis, os escores '
 'foram convertidos em escores T (M = 50; DP = 10).'%(c2('%.2f'%RC['Tensao']['alpha']),c2('%.2f'%RC['Raiva']['alpha'])))
H('2.3 Procedimentos',12,before=6)
P('O BRUMS foi autoaplicado por formulário eletrônico com carimbo de data/hora, duas vezes por dia de treino — a primeira '
 'resposta como pré e a última como pós —, ao longo de sete dias, totalizando %d observações válidas. A Figura 1 sintetiza '
 'o framework das coletas.'%sm['n_obs'])
figure(f'{FG}/xb2_framework.png','Framework das coletas: da amostra e do microciclo às observações do BRUMS, às subescalas e à classificação de perfis.',w=11.0)
H('2.4 Análise de dados',12,before=6)
P('A normalidade foi verificada por Shapiro-Wilk e inspeção de histogramas, diagramas de caixa e dispersão. Como todas as '
 'dimensões violaram a normalidade (p < 0,001), com efeito de piso nas negativas, os dados não foram transformados '
 '(NEVILL; LANE, 2007) e adotaram-se procedimentos não paramétricos. Reportou-se descritiva (média, DP, amplitude, '
 'percentis, escore T), confiabilidade (alfa de Cronbach), coeficiente de variação entre atletas e intraindividual, e '
 'correlação de Spearman (ρ). As comparações intra-grupo (Dia 1 vs. Dia 7; pré vs. pós) usaram MANOVA de medidas '
 'repetidas em escores T (Wilks λ, F, η²ₚ), com testes univariados e Wilcoxon. A variação ao longo dos sete dias foi '
 'testada por Friedman com W de Kendall e por pós-teste das médias marginais de um modelo misto (Tukey/Holm); a '
 'consistência das medidas repetidas, por ICC(2,1) e ICC(2,k). A magnitude seguiu d de Cohen (0,20; 0,50; 0,80) e η²ₚ '
 '(0,01; 0,06; 0,14). No nível do sujeito, calculou-se o RCI (JACOBSON; TRUAX, 1991), com |RCI| ≥ 1,96. A confiabilidade '
 'de uma medida, o erro típico e a mínima mudança detectável foram estimados por decomposição da variância (HOPKINS, '
 '2000). Cada observação foi classificada nos seis perfis de humor por proximidade aos protótipos em escores T, com a '
 'mudança de prevalência testada por qui-quadrado. Adotou-se α = 0,05, com análises em Python (pandas, SciPy, statsmodels).')

# ===================== 3 RESULTADOS =====================
H('3 RESULTADOS')
H('3.1 Triagem de dados e análise descritiva exploratória',12,before=6)
P('O teste de Shapiro-Wilk rejeitou a normalidade em todas as dimensões (p < 0,001; Tabela 4), com as negativas '
 'concentradas no piso. A descritiva geral, a confiabilidade e as intercorrelações constam na Tabela 3; o vigor e a fadiga '
 'concentram a variabilidade informativa, enquanto as negativas têm medianas baixas e coeficientes de variação inflados '
 'pelo efeito de piso (Tabela 5).')
def cval(a,b):
    key='%s_%s'%(a,b); v=MScorr[key]; stx='*' if v['p']<0.001 else ''; return c2('%+.2f'%v['r'])+stx
def mrow(i,d):
    cells=[str(i+1)+' '+d['lab'],c2('%.2f'%d['M']),c2('%.2f'%d['SD']),'%d–%d'%(d['mn'],d['mx']),
        '%s–%s'%(c2('%.0f'%d['tmin']),c2('%.0f'%d['tmax'])),c2('%.2f'%d['alpha'])]
    for k in range(5):
        j=k+1; cells.append('—' if j<=i else cval(MSord[i],MSord[j]))
    return cells
t_desc=table('Descritivas, confiabilidade e intercorrelações das seis dimensões do BRUMS (%d observações).'%sm['n_obs'],
    ['Dimensão','M','DP','Amplitude','Escore T','α','2','3','4','5','6'],[mrow(i,d) for i,d in enumerate(MSd)],
    note='* p < 0,001 (Spearman). Amplitude = escore bruto 0–16; Escore T referenciado à amostra (M = 50; DP = 10).',fs=8)
SH=STAT['shapiro']; IC=STAT['icc']
def nrow(k,lab):
    sh=SH[k]; i=IC[k]; return [lab,c2('%.3f'%sh['W']),'< 0,001' if sh['p']<0.001 else c2('%.3f'%sh['p']),
        c2('%+.2f'%sh['skew']),c2('%.2f'%i['icc1']),c2('%.2f'%i['icck']),i['cls']]
t_norm=table('Normalidade (Shapiro-Wilk), assimetria e consistência (ICC) por dimensão.',
    ['Dimensão','W','p','Assimetria','ICC(2,1)','ICC(2,k)','Consistência'],
    [nrow(k,l) for k,l in [('Vigor','Vigor'),('Fadiga','Fadiga'),('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Confusao','Confusão')]],fs=9)
def cvrow(k,lab):
    v=CV[k]; return [lab,c2('%.0f'%v['inter']),c2('%.0f'%v['intra'])]
t_cv=table('Coeficiente de variação entre atletas e intraindividual, por dimensão.',
    ['Dimensão','CV entre atletas (%)','CV intraindividual (%)'],
    [cvrow(k,l) for k,l in [('Vigor','Vigor'),('Fadiga','Fadiga'),('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Confusao','Confusão')]],
    note='PTH excluída por assumir valores ≤ 0.',fs=9)
figure(f'{FG}/xb3_hist.png','Distribuição de frequências das seis dimensões do BRUMS (linha tracejada = mediana).')
figure(f'{FG}/xb3_box.png','Diagramas de caixa das seis dimensões por dia.')
figure(f'{FG}/xb4_splom.png','Matriz de dispersão entre as seis dimensões (médias semanais por atleta).',w=14.5)
figure(f'{FG}/xb2_traj.png','Trajetória das dimensões ao longo da semana (médias diárias; áreas sombreadas = IC95%).')

H('3.2 Perfil de humor e comparações intra-grupo (Dia 1 vs. Dia 7 e pré vs. pós)',12,before=6)
f_prof1=figure(f'{FG}/xb5_profile_d1d7.png','Perfil de humor em escores T (M = 50; DP = 10) no Dia 1 e no Dia 7.',w=14.0)
P('No Dia 1 o grupo apresenta assinatura próxima do iceberg — vigor acima da média populacional (T = %s) e fadiga abaixo '
 '(T = %s) —, invertendo-se no Dia 7 (Figura %d). A MANOVA de medidas repetidas confirmou diferença multivariada '
 'significativa (Wilks λ = %s; F(%d,%d) = %s; p = %s; η²ₚ = %s, grande); os testes univariados (Tabela %d) localizam o '
 'efeito no vigor (F = %s; p = %s; d = %s; η²ₚ = %s) e na fadiga (F = %s; p = %s; d = %s; η²ₚ = %s).'%(
   c2('%.1f'%mvv('d1d7','Vigor','m1')),c2('%.1f'%mvv('d1d7','Fadiga','m1')),f_prof1,
   c2('%.3f'%MV['d1d7']['wilks']),MV['d1d7']['df1'],MV['d1d7']['df2'],c2('%.2f'%MV['d1d7']['Fmv']),pstr(MV['d1d7']['p_mv']),c2('%.2f'%MV['d1d7']['eta_mv']),_TN[0]+1,
   c2('%.2f'%mvv('d1d7','Vigor','F')),pstr(mvv('d1d7','Vigor','p')),c2('%+.2f'%mvv('d1d7','Vigor','d')),c2('%.2f'%mvv('d1d7','Vigor','eta')),
   c2('%.2f'%mvv('d1d7','Fadiga','F')),pstr(mvv('d1d7','Fadiga','p')),c2('%+.2f'%mvv('d1d7','Fadiga','d')),c2('%.2f'%mvv('d1d7','Fadiga','eta'))))
def mvtable(tab,cap):
    rows=[]
    for x in MV[tab]['rows']:
        sig='*' if x['p']<0.05 else ''
        rows.append([x['lab'],c2('%.1f'%x['m1']),c2('%.1f'%x['s1']),c2('%.1f'%x['m2']),c2('%.1f'%x['s2']),c2('%.2f'%x['F']),pstr(x['p'])+sig,c2('%+.2f'%x['d']),c2('%.3f'%x['eta'])])
    mv=MV[tab]; note='Escores T. MANOVA: Wilks λ = %s; F(%d,%d) = %s; p = %s; η²ₚ = %s. * p < 0,05.'%(c2('%.3f'%mv['wilks']),mv['df1'],mv['df2'],c2('%.2f'%mv['Fmv']),pstr(mv['p_mv']),c2('%.2f'%mv['eta_mv']))
    hdr={'d1d7':['Dimensão','D1 M','D1 DP','D7 M','D7 DP','F','p','d','η²ₚ'],'prepos':['Dimensão','Pré M','Pré DP','Pós M','Pós DP','F','p','d','η²ₚ']}[tab]
    return table(cap,hdr,rows,note=note,fs=8.5)
t_mv1=mvtable('d1d7','Comparação Dia 1 → Dia 7 das seis dimensões em escores T (MANOVA de medidas repetidas; n = %d).'%MV['d1d7']['n'])
P('A resposta aguda pré → pós também foi multivariadamente significativa (Wilks λ = %s; F(%d,%d) = %s; p = %s; η²ₚ = %s), '
 'com o vigor caindo e a fadiga e a PTH subindo a cada sessão (Tabela %d; Figura %d).'%(
   c2('%.3f'%MV['prepos']['wilks']),MV['prepos']['df1'],MV['prepos']['df2'],c2('%.2f'%MV['prepos']['Fmv']),pstr(MV['prepos']['p_mv']),c2('%.2f'%MV['prepos']['eta_mv']),_TN[0]+1,_FN[0]+1))
t_mv2=mvtable('prepos','Resposta aguda pré → pós das seis dimensões em escores T (n = %d).'%MV['prepos']['n'])
figure(f'{FG}/xb5_profile_prepos.png','Perfil de humor em escores T nos momentos pré e pós-treino.',w=14.0)
P('A leitura conjunta padronizada (Figura %d) e o contraste entre os dias extremos (Figura %d) confirmam o cruzamento '
 'vigor↓/fadiga↑. A ordenação de sensibilidade (Tabela %d) situa o vigor (|dz| = %s) e a fadiga (|dz| = %s) como as mais '
 'sensíveis; o Friedman acusou diferença entre os sete dias para vigor (p = %s), fadiga (p = %s), tensão (p = %s) e '
 'confusão (p = %s).'%(_FN[0]+1,_FN[0]+2,_TN[0]+1,c2('%.2f'%SENS['Vigor']['absdz']),c2('%.2f'%SENS['Fadiga']['absdz']),
   pstr(FR['Vigor']['p']),pstr(FR['Fadiga']['p']),pstr(FR['Tensao']['p']),pstr(FR['Confusao']['p'])))
figure(f'{FG}/xb4_allz.png','Comportamento conjunto das sete variáveis ao longo da semana (escores padronizados z).')
figure(f'{FG}/xb4_dumbbell.png','Comparação Dia 1 versus Dia 7 dos escores de todas as variáveis.',w=14.5)
snames={'Vigor':'Vigor','Fadiga':'Fadiga','TMD':'PTH','Tensao':'Tensão','Depressao':'Depressão','Raiva':'Raiva','Confusao':'Confusão'}
sord=sorted(snames,key=lambda k:-SENS[k]['absdz'])
def frow2(k,rank):
    fr=FR[k]; return [rank,snames[k],c2('%+.2f'%SENS[k]['dz']),c2('%.1f'%fr['chi']),pstr(fr['p']),c2('%.2f'%fr['W']),'sim' if fr['p']<0.05 else 'não']
t_sens=table('Sensibilidade à variação semanal e teste de Friedman (7 dias), ordenadas por |dz| (Dia 1 → Dia 7).',
    ['Ordem','Variável','dz (D1→D7)','Friedman χ²','p','W de Kendall','Difere entre dias'],
    [frow2(k,i+1) for i,k in enumerate(sord)],note='Casos completos (n = %d).'%FR['Vigor']['n'],fs=8.5)
P('Cada dimensão é apresentada individualmente, com banda de confiança de 95%%, diagramas de caixa por dia e o efeito do '
 'microciclo (Figuras %d a %d).'%(_FN[0]+1,_FN[0]+6))
for k,fn,lab in [('Vigor','xb4_v_Vigor.png','Vigor'),('Fadiga','xb4_v_Fadiga.png','Fadiga'),('Tensao','xb4_v_Tensao.png','Tensão'),('Depressao','xb4_v_Depressao.png','Depressão'),('Raiva','xb4_v_Raiva.png','Raiva'),('Confusao','xb4_v_Confusao.png','Confusão')]:
    figure(f'{FG}/{fn}','%s ao longo da semana: médias diárias (banda = IC95%%), diagramas de caixa por dia e efeito Dia 1 → Dia 7.'%lab,w=12.5)

H('3.3 Estrutura de correlação (dias de maior vigor e de maior fadiga)',12,before=6)
PA=STAT['pairs']; FP=STAT['focusp']
P('No dia de maior vigor (Dia 1), o afeto negativo é praticamente independente da fadiga (depressão ρ = %s, p = %s; raiva '
 'ρ = %s, p = %s); no dia de maior fadiga (Dia 7), o acoplamento torna-se forte (depressão ρ = %s, p = %s; raiva ρ = %s, '
 'p = %s), confirmando a H3 (Tabela %d).'%(
   c2('%+.2f'%FP['D1']['Depressao']['rho_fad']),pstr(FP['D1']['Depressao']['p_fad']),c2('%+.2f'%FP['D1']['Raiva']['rho_fad']),pstr(FP['D1']['Raiva']['p_fad']),
   c2('%+.2f'%FP['D7']['Depressao']['rho_fad']),pstr(FP['D7']['Depressao']['p_fad']),c2('%+.2f'%FP['D7']['Raiva']['rho_fad']),pstr(FP['D7']['Raiva']['p_fad']),_TN[0]+1))
def frow(neg,lab):
    d1=FP['D1'][neg]; d7=FP['D7'][neg]
    return [lab,'%s (%s)'%(c2('%+.2f'%d1['rho_fad']),pstr(d1['p_fad'])),'%s (%s)'%(c2('%+.2f'%d7['rho_fad']),pstr(d7['p_fad'])),
            '%s (%s)'%(c2('%+.2f'%d1['rho_vig']),pstr(d1['p_vig'])),'%s (%s)'%(c2('%+.2f'%d7['rho_vig']),pstr(d7['p_vig']))]
t_focus=table('Correlação (ρ; p) das dimensões negativas com a fadiga e o vigor no dia de maior vigor (D1) e no de maior fadiga (D7).',
    ['Dimensão','ρ×Fadiga D1 (p)','ρ×Fadiga D7 (p)','ρ×Vigor D1 (p)','ρ×Vigor D7 (p)'],
    [frow('Depressao','Depressão'),frow('Raiva','Raiva'),frow('Confusao','Confusão'),frow('Tensao','Tensão')],fs=8.5)

H('3.4 Prevalência e transição dos perfis de humor',12,before=6)
P('A distribuição dos perfis mudou entre os dias extremos (χ² = %s; gl = %d; p = %s; Figura %d; Tabela %d): o iceberg caiu '
 'de %s%% para %s%% e a barbatana de tubarão subiu para %s%%, sem instalação relevante de perfis de risco extremo.'%(
   c2('%.2f'%PREV['chi']),PREV['dof'],pstr(PREV['p']),_FN[0]+1,_TN[0]+1,
   c2('%.0f'%(100*PREV['D1']['Iceberg']/PREV['n_d1'])),c2('%.0f'%(100*PREV['D7']['Iceberg']/PREV['n_d7'])),c2('%.0f'%(100*PREV['D7']['Barbatana tubarão']/PREV['n_d7']))))
figure(f'{FG}/xb5_prev.png','Prevalência (%) dos seis perfis de humor no Dia 1 e no Dia 7.',w=14.0)
PROFR=[('Iceberg','Iceberg'),('Everest invertido','Everest invertido'),('Iceberg invertido','Iceberg invertido'),('Submerso','Submerso'),('Barbatana tubarão','Barbatana de tubarão'),('Superfície','Superfície')]
def prow(p,lab):
    d1=PREV['D1'][p]; d7=PREV['D7'][p]; ov=PREV['overall'][p]; nov=sum(PREV['overall'].values())
    return [lab,'%d (%s)'%(d1,c2('%.1f'%(100*d1/PREV['n_d1']))),'%d (%s)'%(ov,c2('%.1f'%(100*ov/nov))),'%d (%s)'%(d7,c2('%.1f'%(100*d7/PREV['n_d7'])))]
t_prev=table('Prevalência dos perfis de humor por recorte: n (%%) no Dia 1, na semana e no Dia 7.',
    ['Perfil','Dia 1 n (%)','Semana n (%)','Dia 7 n (%)'],[prow(p,lab) for p,lab in PROFR],
    note='Qui-quadrado Dia 1 × Dia 7: χ² = %s; gl = %d; p = %s.'%(c2('%.2f'%PREV['chi']),PREV['dof'],pstr(PREV['p'])),fs=9)

H('3.5 Análise no nível do atleta: trajetórias, mudança confiável e transição de perfis',12,before=6)
P('As trajetórias individuais de vigor e fadiga (Figura %d) mostram ampla dispersão entre atletas. A mudança confiável '
 '(RCI, Dia 1 → Dia 7) distribui-se de modo desigual (Tabela %d; Figura %d): a fadiga é a dimensão em que a maioria '
 'deteriora de forma confiável (%d de %d; %d%%), seguida do vigor (%d%%); as negativas não pioram. A classificação '
 'individual (Tabela %d) mostra que o iceberg era o perfil de %d dos %d atletas na linha de base e a barbatana de tubarão '
 'passou a ser o de %d no último dia.'%(
   _FN[0]+1,_TN[0]+1,_FN[0]+2,RC['Fadiga']['piora'],RC['Fadiga']['n'],RC['Fadiga']['pct_piora'],RC['Vigor']['pct_piora'],
   _TN[0]+2,PCNT['iceberg_D1'],PCNT['n'],PCNT['shark_D7']))
figure(f'{FG}/xb4_spaghetti.png','Trajetórias individuais de vigor e fadiga (linhas cinzas = atletas; linha colorida = média do grupo).')
def rcrow(k,lab):
    v=RC[k]; return [lab,c2('%.2f'%v['alpha']),c2('%.2f'%v['sediff']),v['n'],'%d (%d%%)'%(v['piora'],v['pct_piora']),'%d (%d%%)'%(v['estavel'],v['pct_est']),'%d (%d%%)'%(v['melhora'],v['pct_mel'])]
t_rci=table('Mudança confiável individual (RCI, Dia 1 → Dia 7) por dimensão (n = %d).'%RC['Vigor']['n'],
    ['Dimensão','α','EP_dif','n','Piora confiável','Estável','Melhora confiável'],
    [rcrow(k,lab) for k,lab in [('Vigor','Vigor'),('Fadiga','Fadiga'),('Depressao','Depressão'),('Raiva','Raiva'),('Tensao','Tensão'),('Confusao','Confusão')]],fs=9)
figure(f'{FG}/xb2_rci6.png','Proporção de atletas com mudança confiável (RCI) por dimensão, do Dia 1 ao Dia 7.',w=13.5)
ABBR={'Iceberg':'Iceberg','Everest invertido':'Everest inv.','Iceberg invertido':'Iceberg inv.','Submerso':'Submerso','Barbatana tubarão':'Barbatana','Superfície':'Superfície'}
def tag(d1,d7):
    if d7=='Iceberg' and d1!='Iceberg': return 'melhora'
    if d7=='Barbatana tubarão' and d1!='Barbatana tubarão': return 'fadiga'
    if d1==d7: return 'estável'
    return 'transição'
rows_tr=[[aid,ABBR.get(PT[aid]['D1'],PT[aid]['D1']),ABBR.get(PT[aid]['D7'],PT[aid]['D7']),tag(PT[aid]['D1'],PT[aid]['D7'])] for aid in sorted(PT)]
t_trans=table('Perfil de humor de cada atleta no Dia 1 e no Dia 7 e classificação da transição (n = %d).'%PCNT['n'],
    ['Atleta','Perfil no Dia 1','Perfil no Dia 7','Transição'],rows_tr,fs=8.5)

H('3.6 Pós-teste, cruzamento vigor–fadiga e decomposição sinal–ruído',12,before=6)
KD=PHJ['keydays']
def emmc(v,d): return c2('%.2f'%PHJ[v]['emm'][str(d)])
def sig1(v,d): return '' if d==1 or PHJ[v]['pairs']['1_%d'%d]['ptukey']>=0.05 else '*'
P('O pós-teste do modelo misto localiza os dias críticos (Tabela %d; Figura %d): a queda do vigor é mais acentuada no Dia '
 '%d → Dia %d (Δ = %s), refletindo o choque de carga, enquanto a fadiga sobe mais no Dia %d → Dia %d (Δ = %s). Em cada '
 'dia, o vigor cai e a fadiga sobe da coleta pré para a pós (padrão em dente de serra); as curvas cruzam-se pela primeira '
 'vez no Dia 2 (entre pré e pós) e a fadiga supera definitivamente o vigor no Dia 7, no qual o vigor atinge o mínimo e a '
 'fadiga o máximo (Figura %d).'%(
   _TN[0]+1,_FN[0]+1,KD['vigor_maxdrop_day']-1,KD['vigor_maxdrop_day'],c2('%+.2f'%KD['vigor_maxdrop']),
   KD['fadbrums_maxrise_day']-1,KD['fadbrums_maxrise_day'],c2('%+.2f'%KD['fadbrums_maxrise']),_FN[0]+2))
def phrow(d): return ['Dia %d'%d,emmc('Vigor',d)+sig1('Vigor',d),emmc('Fadiga',d)+sig1('Fadiga',d)]
t_ph=table('Médias marginais estimadas (modelo misto) de vigor e fadiga por dia e pós-teste vs. Dia 1.',
    ['Dia','Vigor','Fadiga'],[phrow(d) for d in range(1,8)],note='* diferença significativa vs. Dia 1 (Tukey, p < 0,05).',fs=9)
figure(f'{FG}/ph_emm.png','Trajetória diária (médias marginais do modelo misto) com pós-teste vs. Dia 1 e destaque dos dias de maior variação.')
figure(f'{FG}/x_cross.png','Cruzamento vigor × fadiga por dia e momento (pré → pós dentro de cada dia); linhas pontilhadas = cruzamentos.',w=15.0)
dV=CRS['dec']['Vigor']; dF=CRS['dec']['Fadiga']
P('A decomposição da variância mostra que o humor é dominado pela diferença entre atletas (%s%%–%s%%), com componente de '
 'dia (sinal do microciclo) pequeno mas presente no vigor (%s%%) e na fadiga (%s%%) e quase nulo nas negativas (Tabela '
 '%d; Figura %d). Consistentemente, a razão sinal–ruído favorece o eixo energia–fadiga (RSR = %s no vigor; %s na fadiga) e '
 'a discriminação entre o Dia 7 e o Dia 1 melhora ao remover o ruído (médias diárias). A mínima mudança detectável reduz-'
 'se com o número de coletas, o que fundamenta a agregação de múltiplas aferições; a remoção de ruído, contudo, não cria '
 'sinal nas subescalas negativas de piso.'%(
   c2('%.0f'%min(CRS['dec'][k]['vc']['ath'] for k in ['Vigor','Fadiga','TMD','Tensao','Depressao'])),
   c2('%.0f'%max(CRS['dec'][k]['vc']['ath'] for k in ['Vigor','Fadiga','TMD','Tensao','Depressao'])),
   c2('%.0f'%dV['vc']['day']),c2('%.0f'%dF['vc']['day']),_TN[0]+1,_FN[0]+1,c2('%.2f'%DEN['snr_raw']['Vigor']),c2('%.2f'%DEN['snr_raw']['Fadiga'])))
figure(f'{FG}/x_vardecomp.png','Decomposição da variância de cada variável do humor em componentes entre atletas, dia e resíduo.',w=13.5)
def snrow(k,lab):
    return [lab,c2('%.2f'%DEN['rel'][k]),c2('%.2f'%DEN['etm'][k]),c2('%.2f'%DEN['snr_raw'][k]),c2('%.2f'%ROC[k]['raw'][0]),c2('%.2f'%ROC[k]['filt'][0]),c2('%.1f'%DEN['kmdc'][k]['1']),c2('%.1f'%DEN['kmdc'][k]['2'])]
t_snr=table('Decomposição sinal–ruído: confiabilidade, erro típico, razão sinal–ruído, discriminação (ROC) e mudança detectável.',
    ['Variável','r (1 medida)','ETM','RSR','AUC c/ ruído','AUC s/ ruído','MDC95 (1)','MDC95 (2)'],
    [snrow('Vigor','Vigor'),snrow('Fadiga','Fadiga (BRUMS)'),snrow('TMD','PTH'),snrow('FadMental','Fadiga mental')],
    note='ETM = erro típico de medida; RSR = razão sinal–ruído; AUC = área sob a curva ROC (Dia 7 vs. Dia 1); MDC95 (k) = mínima mudança detectável com k coletas.',fs=8.5)

# ===================== 4 DISCUSSÃO =====================
H('4 DISCUSSÃO')
P('A deterioração do humor no microciclo pré-competitivo concentrou-se no eixo energia–fadiga (H1): a MANOVA Dia 1 → Dia 7 '
 'foi multivariadamente significativa (η²ₚ = %s), com os maiores efeitos no vigor (η²ₚ = %s) e na fadiga (η²ₚ = %s), '
 'enquanto as negativas permaneceram próximas do piso — coerente com a literatura que identifica o vigor e a fadiga como '
 'as dimensões mais sensíveis à carga (SAW; MAIN; GASTIN, 2016; THORPE et al., 2017). A migração do perfil iceberg para o '
 'de fadiga (H2; χ² = %s; p = %s) reproduz, em sete dias, o “derretimento do iceberg” e é compatível com a associação '
 'entre perturbação do humor e desfechos esportivos (MORGAN, 1985; LOCHBAUM et al., 2021).'%(
   c2('%.2f'%MV['d1d7']['eta_mv']),c2('%.2f'%mvv('d1d7','Vigor','eta')),c2('%.2f'%mvv('d1d7','Fadiga','eta')),c2('%.2f'%PREV['chi']),pstr(PREV['p'])))
P('O afeto negativo acoplou-se à fadiga apenas sob carga acumulada (H3): a correlação depressão × fadiga passou de ρ = %s '
 'no dia de maior vigor para ρ = %s no de maior fadiga, sugerindo que a irritabilidade e o abatimento se organizam com a '
 'exaustão. A resposta foi heterogênea entre atletas (H4), com a fadiga sendo a única dimensão em que a maioria deteriora '
 'de forma confiável, sustentando o monitoramento individualizado (KELLMANN et al., 2018). O pós-teste e a análise de '
 'cruzamento acrescentam o “quando”: o vigor cai mais no choque inicial de carga e a fadiga acumula-se rumo ao Dia 7, com '
 'as curvas cruzando-se já no Dia 2 — informação diretamente acionável para escalonar a recuperação. A decomposição '
 'sinal–ruído demonstra que a agregação de múltiplas coletas melhora a confiabilidade e a detecção da mudança (HOPKINS, '
 '2000), sendo o eixo energia–fadiga o único a carregar sinal semanal aproveitável.'%(
   c2('%+.2f'%foc['D1']['Depressao']['rho_fad']),c2('%+.2f'%foc['D7']['Depressao']['rho_fad'])))
P('Como limitações, destacam-se o tamanho amostral (n = %d), o efeito de piso das dimensões negativas — que reduz a '
 'variância e a confiabilidade da tensão e da confusão — e o caráter observacional de fase única, que não permite '
 'inferência causal. A conversão em escores T foi referenciada à própria amostra. Estudos futuros devem ampliar a amostra '
 'com múltiplos microciclos e integrar o humor a marcadores de carga e de aptidão física.'%sm['n'])

# ===================== 5 CONCLUSÕES =====================
H('5 CONCLUSÕES')
P('Na última semana de pré-temporada, o perfil de humor de handebolistas de elite migrou da prontidão (iceberg) para a '
 'fadiga funcional (barbatana de tubarão), com a deterioração concentrada no eixo energia–fadiga — vigor em queda de '
 'efeito grande e fadiga em ascensão, as dimensões mais sensíveis —, tanto no grupo quanto no sujeito, onde a fadiga foi '
 'a única a piorar de forma confiável na maioria. O vigor cai mais no início do microciclo e a fadiga acumula-se rumo ao '
 'último dia, com as curvas cruzando-se no Dia 2. Recomenda-se monitorar o eixo energia–fadiga por tendência e '
 'referenciado à linha de base individual, agregando múltiplas coletas para melhorar a razão sinal–ruído.')

# ===================== REFERÊNCIAS =====================
H('REFERÊNCIAS')
refs=[
 'DE MIRANDA ROHLFS, I. C. P. et al. Mood states, injury status, and countermovement jump performance in Brazilian high-level sports. Sports, v. 13, n. 9, 303, 2025. DOI: 10.3390/sports13090303.',
 'HAN, C.; PARSONS-SMITH, R. L.; TERRY, P. C. Mood profiling in Singapore: cross-cultural validation and potential applications of mood profile clusters. Frontiers in Psychology, v. 11, 665, 2020. DOI: 10.3389/fpsyg.2020.00665.',
 'HOPKINS, W. G. Measures of reliability in sports medicine and science. Sports Medicine, v. 30, n. 1, p. 1–15, 2000. DOI: 10.2165/00007256-200030010-00001.',
 'JACOBSON, N. S.; TRUAX, P. Clinical significance: a statistical approach to defining meaningful change in psychotherapy research. Journal of Consulting and Clinical Psychology, v. 59, n. 1, p. 12–19, 1991. DOI: 10.1037/0022-006X.59.1.12.',
 'KARCHER, C.; BUCHHEIT, M. On-court demands of elite handball, with special reference to playing positions. Sports Medicine, v. 44, n. 6, p. 797–814, 2014. DOI: 10.1007/s40279-014-0164-z.',
 'KELLMANN, M. et al. Recovery and performance in sport: consensus statement. International Journal of Sports Physiology and Performance, v. 13, n. 2, p. 240–245, 2018. DOI: 10.1123/ijspp.2017-0759.',
 'LEW, P. C. F. et al. Cross-cultural validation of the Malaysian Mood Scale and tests of between-group mood differences. International Journal of Environmental Research and Public Health, v. 20, n. 4, 3348, 2023. DOI: 10.3390/ijerph20043348.',
 'LOCHBAUM, M. et al. The Profile of Mood States and athletic performance: a meta-analysis of published studies. European Journal of Investigation in Health, Psychology and Education, v. 11, n. 1, p. 50–70, 2021. DOI: 10.3390/ejihpe11010005.',
 'MICHALSIK, L. B.; MADSEN, K.; AAGAARD, P. Match performance and physiological capacity of female elite team handball players. International Journal of Sports Medicine, v. 35, n. 7, p. 595–607, 2013. DOI: 10.1055/s-0033-1358713.',
 'MORGAN, W. P. Selected psychological factors limiting performance: a mental health model. In: CLARKE, D. H.; ECKERT, H. M. (Ed.). Limits of human performance. Champaign: Human Kinetics, 1985. p. 70–80.',
 'NEVILL, A. M.; LANE, A. M. Why self-report “Likert” scale data should not be log-transformed. Journal of Sports Sciences, v. 25, n. 1, p. 1–2, 2007. DOI: 10.1080/02640410601111183.',
 'PARSONS-SMITH, R. L.; TERRY, P. C.; MACHIN, M. A. Identification and description of novel mood profile clusters. Frontiers in Psychology, v. 8, 1958, 2017. DOI: 10.3389/fpsyg.2017.01958.',
 'ROHLFS, I. C. P. M. et al. A Escala de Humor de Brunel (Brums): instrumento para detecção precoce da síndrome do excesso de treinamento. Revista Brasileira de Medicina do Esporte, v. 14, n. 3, p. 176–181, 2008.',
 'ROHLFS, I. C. P. M. et al. Psychometric characteristics of the Brazil Mood Scale among youth and elite athletes using two response time frames. Sports, v. 11, n. 12, 244, 2023. DOI: 10.3390/sports11120244.',
 'SAW, A. E.; MAIN, L. C.; GASTIN, P. B. Monitoring the athlete training response: subjective self-reported measures trump commonly used objective measures: a systematic review. British Journal of Sports Medicine, v. 50, n. 5, p. 281–291, 2016. DOI: 10.1136/bjsports-2015-094758.',
 'TERRY, P. C.; LANE, A. M.; FOGARTY, G. J. Construct validity of the Profile of Mood States — Adolescents for use with adults. Psychology of Sport and Exercise, v. 4, n. 2, p. 125–139, 2003. DOI: 10.1016/S1469-0292(02)00035-8.',
 'TERRY, P. C. et al. Mood profiling for sustainable mental health among athletes. Sustainability, v. 13, n. 11, 6116, 2021. DOI: 10.3390/su13116116.',
 'THORPE, R. T. et al. Monitoring fatigue status in elite team-sport athletes: implications for practice. International Journal of Sports Physiology and Performance, v. 12, s2, p. S2-27–S2-34, 2017. DOI: 10.1123/ijspp.2016-0434.']
for rf in refs:
    p=doc.add_paragraph(); r=p.add_run(rf); r.font.size=Pt(11); p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(6); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

OUTP='/home/user/mdlucca/Artigos/Artigo1_Perfil_Humor_BRUMS.docx'
doc.save(OUTP); print('SAVED',OUTP,'| Tabelas',_TN[0],'Figuras',_FN[0])
