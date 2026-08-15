import warnings; warnings.filterwarnings('ignore')
import numpy as np, pandas as pd, json, os
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
SC='/tmp/claude-0/-home-user-mdlucca/e1dba24c-b1d7-5908-9106-f2f4aaf3f56a/scratchpad'; FIG='/home/user/mdlucca/Artigos/figuras'
R=json.load(open(f'{SC}/semana2.json')); REG=json.load(open(f'{SC}/reg.json')); REL=json.load(open(f'{SC}/reliab.json'))
MV=json.load(open(f'{SC}/mv.json')); DX=json.load(open(f'{SC}/decomp_extra.json')); CFA=json.load(open(f'{SC}/cfa.json'))
DE=json.load(open(f'{SC}/deriv_exp.json')); ALLO=json.load(open(f'{SC}/allo.json')); ANORM=json.load(open(f'{SC}/allonorm.json'))
desc=R['desc']; dec=R['dec']; der=R['der']; fc=DX['facet']; sl=DX['slope']; ac=DX['acf']; rc=DX['rci']
CORE=[('Vigor','Vigor'),('Fadiga','Fadiga (BRUMS)'),('FadFisica','Fadiga física'),('FadMental','Fadiga mental'),('TMD','PTH/TMD')]
def c2(x): return str(x).replace('.',',')

doc=Document(); s=doc.sections[0]; s.left_margin=Cm(3); s.top_margin=Cm(3); s.right_margin=Cm(2); s.bottom_margin=Cm(2)
n=doc.styles['Normal']; n.font.name='Times New Roman'; n.font.size=Pt(12); n.paragraph_format.line_spacing=1.5; n.paragraph_format.space_after=Pt(0)
FN='Times New Roman'
def _set(r,sz=12,b=False,it=False): r.font.name=FN; r.font.size=Pt(sz); r.bold=b; r.italic=it
def sec(t):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(3); _set(p.add_run(t),12,b=True)
def P(t,indent=True,size=11,it=False):
    p=doc.add_paragraph(); _set(p.add_run(t),size,it=it); p.paragraph_format.line_spacing=1.4
    if indent: p.paragraph_format.first_line_indent=Cm(1.25)
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
_FN=[0]; _TN=[0]
def fig(name,titulo,w=5.6):
    _FN[0]+=1; cap=doc.add_paragraph(); cap.alignment=WD_ALIGN_PARAGRAPH.CENTER; cap.paragraph_format.space_before=Pt(4); cap.paragraph_format.line_spacing=1.0
    _set(cap.add_run(f'Figura S{_FN[0]} – {titulo}'),9.5)
    p=f'{FIG}/{name}.png'
    if os.path.exists(p): doc.add_picture(p,width=Inches(w)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    else: print('MISS',name)
    fp=doc.add_paragraph(); fp.alignment=WD_ALIGN_PARAGRAPH.CENTER; fp.paragraph_format.space_after=Pt(6); fp.paragraph_format.line_spacing=1.0; _set(fp.add_run('Fonte: elaborado pelos autores (2026).'),9)
def _bd(t):
    tblPr=t._tbl.tblPr; b=OxmlElement('w:tblBorders')
    for e in ('top','bottom','insideH'):
        x=OxmlElement(f'w:{e}'); x.set(qn('w:val'),'single'); x.set(qn('w:sz'),'6'); x.set(qn('w:space'),'0'); x.set(qn('w:color'),'000000'); b.append(x)
    for e in ('left','right','insideV'):
        x=OxmlElement(f'w:{e}'); x.set(qn('w:val'),'none'); b.append(x)
    tblPr.append(b)
def table(titulo,headers,rows,fs=8.5):
    _TN[0]+=1; cap=doc.add_paragraph(); cap.alignment=WD_ALIGN_PARAGRAPH.LEFT; cap.paragraph_format.space_before=Pt(4); cap.paragraph_format.line_spacing=1.0
    _set(cap.add_run(f'Tabela S{_TN[0]} – {titulo}'),9.5)
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; _bd(t)
    for i,h in enumerate(headers):
        cl=t.rows[0].cells[i]; cl.paragraphs[0].line_spacing=1.0; _set(cl.paragraphs[0].add_run(str(h)),fs,b=True)
    for row in rows:
        cs=t.add_row().cells
        for i,v in enumerate(row): cs[i].paragraphs[0].line_spacing=1.0; _set(cs[i].paragraphs[0].add_run(str(v)),fs)
    fp=doc.add_paragraph(); fp.paragraph_format.space_after=Pt(6); fp.paragraph_format.line_spacing=1.0; _set(fp.add_run('Fonte: dados da pesquisa (2026).'),9)

# CABEÇALHO
tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER; _set(tp.add_run('MATERIAL SUPLEMENTAR'),14,b=True)
st=doc.add_paragraph(); st.alignment=WD_ALIGN_PARAGRAPH.CENTER; _set(st.add_run('Perfil de humor de atletas de handebol na última semana da pré-temporada competitiva: sinal, ruído e sensibilidade dos marcadores'),11,it=True)
P('Este material reúne as análises complementares referidas no manuscrito principal, mantendo o mesmo conjunto de dados (27 atletas, 456 observações, coletas pré/pós). Atletas anonimizados (A01–A27).',indent=False,size=10)

sec('S1. Estatística descritiva completa e distribuições')
rows=[[desc[v]['lab'],c2(f"{desc[v]['mean']:.2f} ± {desc[v]['sd']:.2f}"),c2(f"{desc[v]['med']:.1f} [{desc[v]['iqr']:.1f}]"),f"{desc[v]['mn']:.0f}–{desc[v]['mx']:.0f}",c2(f"{desc[v]['amp']:.0f}"),c2(f"{desc[v]['skew']:+.2f}"),c2(f"{desc[v]['floor']:.0f}%")] for v in ['Tensao','Depressao','Raiva','Vigor','Fadiga','Confusao','TMD','FadFisica','FadMental']]
table('Estatística descritiva completa do humor e da fadiga na semana',['Variável','M ± DP','Mediana [IIQ]','Mín.–Máx.','Ampl.','Assim.','% piso'],rows)
fig('v_hist','Histogramas de distribuição das cinco variáveis centrais'); fig('v_box','Diagramas de caixa por dia')

sec('S2. Tendência semanal (regressão linear)')
rows=[[lab,c2(f"{REG[v]['slope']:+.3f}"),c2(f"{REG[v]['r2']:.3f}"),(c2(f"{REG[v]['p']:.4f}") if REG[v]['p']>=0.001 else '<0,001')] for v,lab in CORE]
table('Regressão linear ao longo da semana (coef./dia, R², p)',['Variável','Coef./dia','R²','p'],rows)
fig('v_reg','Regressão linear de cada variável ao longo da semana')

sec('S3. Comportamento da Perturbação Total do Humor (PTH/TMD)')
fig('v_tmd','Trajetória, resposta pré/pós por dia e distribuição da PTH/TMD')

sec('S4. Limites, derivadas e ponto de inflexão')
rows=[[lab,f"{desc[v]['mn']:.0f}–{desc[v]['mx']:.0f}",c2(f"{DE[v]['amp_sig']:.2f}"),c2(f"{DE[v]['pct_scale']:.0f}%"),c2(f"{DE[v]['auc']:.1f}"),(c2(f"dia {der[v]['infl']:.1f}") if der[v]['infl'] else 'n/d')] for v,lab in CORE]
table('Limites, amplitude do sinal, fração da escala, dose acumulada e ponto de inflexão',['Variável','Mín.–Máx.','Ampl. sinal','% escala','Dose (AUC)','Inflexão'],rows)
fig('x_deriv_exp','Velocidade e aceleração por dia, para cada variável')

sec('S5. Decomposição de generalizabilidade, mudança confiável e dimensionalidade')
rows=[[fc[v]['lab'],c2(f"{fc[v]['ptrait']:.0f}%"),c2(f"{fc[v]['pstate']:.0f}%"),c2(f"{fc[v]['pnoise']:.0f}%"),c2(f"{fc[v]['icc1']:.2f}"),c2(f"{fc[v]['phi']:.2f}"),(c2(f"{sl[v]['sd']:.2f}") if sl[v]['sd'] else '—'),c2(f"{ac[v]:+.2f}"),f"{rc[v]['pct']:.0f}%"] for v,_ in CORE]
table('Facetas (traço/estado/ruído), fidedignidade (ICC, Φ), heterogeneidade do declínio (DP das inclinações), autocorrelação e mudança confiável (RCI)',['Variável','%Traço','%Estado','%Ruído','ICC','Φ','DP incl.','ACF(1)','%RCI'],rows,fs=8)
fig('x_facet','Decomposição em três facetas (estado, traço, ruído)'); fig('x_pca','Variância explicada pelos componentes principais (BRUMS)',w=5.0)

sec('S6. Modelo misto multivariado e correlações')
rows=[['Efeito conjunto do dia (LRT)',f"χ²({MV['dfd']}) = {MV['lr']:.1f}".replace('.',','),'p < 0,001'],
 ['Hotelling T² (D1 vs D7)',f"F({MV['dfn']},{MV['dfden']}) = {MV['F']:.2f}".replace('.',','),'p < 0,001'],
 ['D de Mahalanobis (D1→D7)',c2(f"{MV['mahal']:.2f}"),'—']]
table('Testes multivariados do efeito do microciclo (eixo energia–fadiga)',['Teste','Estatística','p'],rows)
fig('x_trait_corr','Correlações entre-atletas (traço)',w=4.6); fig('x_rmcorr','Correlação intra-atleta (rmcorr)',w=4.6)

sec('S7. Análise fatorial confirmatória (6 fatores do BRUMS)')
rows=[['χ² (df)',f"{CFA['chi2']:.0f} ({CFA['df']})"],['CFI',c2(f"{CFA['cfi']:.3f}")],['TLI',c2(f"{CFA['tli']:.3f}")],['RMSEA',c2(f"{CFA['rmsea']:.3f}")],['SRMR',c2(f"{CFA['srmr']:.3f}")]]
table('Índices de ajuste da AFC (estimação por máxima verossimilhança; 24 itens; observações aninhadas)',['Índice','Valor'],rows)

sec('S8. Ajuste alométrico e normalização por massa e por pico de velocidade')
rows=[[ALLO[v]['lab'],c2(f"{ALLO[v]['a']:.2f}"),c2(f"{ALLO[v]['b']:.2f}"),c2(f"{ALLO[v]['r2']:.2f}")] for v in ['FadFisica','Fadiga','TMD']]
table('Ajuste alométrico da trajetória de fadiga (Y = a·diaᵇ)',['Variável','a','b','R²'],rows)
def star(p): return '*' if p<0.05 else ''
rows=[[ANORM[v]['lab'],c2(f"{ANORM[v]['b_mass']:+.2f}"),c2(f"{ANORM[v]['rho_mass']:+.2f}")+star(ANORM[v]['p_rho_mass'])+' → '+c2(f"{ANORM[v]['rho_after_mass']:+.2f}"),c2(f"{ANORM[v]['b_pv']:+.2f}"),c2(f"{ANORM[v]['rho_pv']:+.2f}")+star(ANORM[v]['p_rho_pv'])+' → '+c2(f"{ANORM[v]['rho_after_pv']:+.2f}")] for v in ['FadFisica','Fadiga','TMD','Vigor']]
table('Normalização alométrica: expoentes (b) e correlação (ρ) com a covariável, bruta → após normalização, por massa e por pico de velocidade (* p < 0,05; a queda de ρ para ≈ 0 confirma a remoção da dependência)',['Variável','b massa','ρ massa (bruto→norm.)','b PV','ρ PV (bruto→norm.)'],rows)
P('A normalização alométrica ajusta cada variável pela covariável elevada a um expoente estimado (Y/massaᵇ e Y/PVᵇ). Pela massa corporal, os expoentes não atingiram significância e as correlações brutas foram fracas — as variáveis de humor e fadiga são essencialmente independentes do porte corporal, tornando a normalização por massa dispensável. Pela aptidão aeróbia (pico de velocidade do T-CAR), evidenciou-se um gradiente significativo: atletas mais aptos apresentaram menor fadiga física (ρ = −0,52; p < 0,05) e maior vigor (ρ = +0,48; p < 0,05), com expoentes negativos para as fadigas e positivo para o vigor. Portanto, o pico de velocidade — e não a massa — é a covariável relevante, e a normalização por PVᵇ produz um índice de fadiga ajustado à aptidão.')
fig('x_allo','Ajuste alométrico da trajetória de fadiga (linear e log-log)',w=5.2)
fig('x_allonorm','Relações alométricas (log-log) com a massa corporal (esq.) e com o pico de velocidade (dir.)',w=5.6)

P('Reprodutibilidade: todas as análises são reproduzíveis pelos scripts em Python do repositório; atletas anonimizados (A01–A27); nenhum dado bruto com identificação é distribuído.',indent=False,size=9,it=True)

OUTP='/home/user/mdlucca/Artigos/Material_Suplementar_Periodico.docx'
doc.save(OUTP); print('SAVED',OUTP)
d2=Document(OUTP); print('Tabelas S',_TN[0],'Figuras S',_FN[0],'| tables',len(d2.tables),'imgs',len(d2.inline_shapes))
