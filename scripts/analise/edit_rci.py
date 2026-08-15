# -*- coding: utf-8 -*-
import json, numpy as np
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
DOC='/home/user/mdlucca/Artigos/PERFIL_HUMOR_COMPLETO.docx'
R=json.load(open('rci.json'))
def c2(s): return str(s).replace('.',',')
d=Document(DOC)
def _fmt(p,size=12,just=True,bold=False,center=False):
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(6)
    if center: p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    elif just: p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    for r in p.runs: r.font.name='Times New Roman'; r.font.size=Pt(size); r.bold=bold
def after(el,text='',**kw):
    p=d.add_paragraph(text); el.addnext(p._p); _fmt(p,**kw); return p._p
def head(el,text): return after(el,text,bold=True,just=False)
def _bd(cell):
    tcPr=cell._tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for edge in ['top','bottom','left','right']:
        e=OxmlElement(f'w:{edge}'); on=edge in ('top','bottom')
        e.set(qn('w:val'),'single' if on else 'nil'); e.set(qn('w:sz'),'6'); e.set(qn('w:color'),'000000'); b.append(e)
    tcPr.append(b)
def table_after(el,caption,header,rows,fonte='Fonte: dados da pesquisa (2026).',fs=9):
    cap=d.add_paragraph(caption); el.addnext(cap._p); _fmt(cap,size=11,just=False)
    t=d.add_table(rows=1,cols=len(header)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; cap._p.addnext(t._tbl)
    for i,htx in enumerate(header):
        cc=t.rows[0].cells[i]; cc.text=''; rr=cc.paragraphs[0].add_run(htx); rr.bold=True; rr.font.size=Pt(fs); rr.font.name='Times New Roman'
        cc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; _bd(cc)
    for row in rows:
        cs=t.add_row().cells
        for i,val in enumerate(row):
            cs[i].text=''; rr=cs[i].paragraphs[0].add_run(str(val)); rr.font.size=Pt(fs); rr.font.name='Times New Roman'
            cs[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if i>0 else WD_ALIGN_PARAGRAPH.LEFT; _bd(cs[i])
    fon=d.add_paragraph(fonte); t._tbl.addnext(fon._p); _fmt(fon,size=10,just=False); return fon._p
def fig_after(el,path,caption,w=16.0):
    pimg=d.add_paragraph(); el.addnext(pimg._p); pimg.alignment=WD_ALIGN_PARAGRAPH.CENTER; pimg.add_run().add_picture(path,width=Cm(w))
    cap=d.add_paragraph(caption); pimg._p.addnext(cap._p); _fmt(cap,size=11,just=False,center=True)
    fon=d.add_paragraph('Fonte: elaboração dos autores (2026).'); cap._p.addnext(fon._p); _fmt(fon,size=10,just=False,center=True); return fon._p
def find(prefix):
    for p in d.paragraphs:
        if p.text.strip().startswith(prefix): return p
    raise KeyError(prefix)

A=R['alpha']; SE=R['sediff']; S=R['summary']; RCI=R['rci']
p6=find('6 SÍNTESE DOS PRINCIPAIS ACHADOS'); anchor=p6._p.getprevious()
anchor=head(anchor,'5.14 Mudança confiável individual (RCI) do Dia 1 ao Dia 7')
anchor=after(anchor,
 'O Índice de Mudança Confiável (RCI) avalia, para cada atleta, se a mudança do Dia 1 ao Dia 7 excede o erro de medida '
 'esperado, distinguindo deterioração real de flutuação de medida (JACOBSON; TRUAX, 1991). Calculou-se RCI = (D7 − D1) / '
 'EP_dif, com EP_dif = DP_D1 · √[2(1 − α)] e α = fidedignidade interna de cada escala (Vigor α = %s; Fadiga α = %s; '
 'PTH/TMD α = %s), o que resulta em erros-padrão da diferença de %s, %s e %s pontos, respectivamente. Uma mudança é '
 'confiável quando |RCI| ≥ 1,96 (nível de 95%%).'%(
   c2('%.2f'%A['Vigor']),c2('%.2f'%A['Fadiga']),c2('%.2f'%A['TMD']),
   c2('%.2f'%SE['Vigor']),c2('%.2f'%SE['Fadiga']),c2('%.2f'%SE['TMD'])))
anchor=after(anchor,
 'Dos %d atletas com medida no Dia 1 e no Dia 7 (Tabela 18; Figura 22), a fadiga aumentou de forma confiável em %d (%d%%), '
 'a PTH/TMD em %d (%d%%) e o vigor caiu de forma confiável em %d (%d%%); um atleta (A05) apresentou melhora confiável em '
 'fadiga e PTH. A maioria, portanto, encerra a semana com deterioração real — acima do ruído de medida — no eixo '
 'energia–fadiga, e a fadiga é o marcador que mais frequentemente ultrapassa o limiar individual, coerente com o efeito '
 'agudo (§5.13) e a trajetória de grupo. A síntese por variável está na Tabela 19.'%(
   S['Vigor']['n'],S['Fadiga']['piora'],int(S['Fadiga']['pct_piora']),S['TMD']['piora'],int(S['TMD']['pct_piora']),
   S['Vigor']['piora'],int(S['Vigor']['pct_piora'])))

# Tabela 18 – per-athlete RCI
ids=sorted(RCI['Vigor'].keys())
def cell(v,aid):
    r=RCI[v].get(aid)
    if r is None: return '—'
    base=c2('%+.1f'%r)
    det=(v=='Vigor' and r<=-1.96) or (v!='Vigor' and r>=1.96)
    imp=(v=='Vigor' and r>=1.96) or (v!='Vigor' and r<=-1.96)
    if det: return base+(' ↓*' if v=='Vigor' else ' ↑*')
    if imp: return base+' (m)'
    return base
rows18=[[aid,cell('Vigor',aid),cell('Fadiga',aid),cell('TMD',aid)] for aid in ids if RCI['Vigor'].get(aid) is not None]
anchor=table_after(anchor,'Tabela 18 – Índice de Mudança Confiável (RCI) do Dia 1 ao Dia 7 por atleta. * = mudança confiável de deterioração (|RCI| ≥ 1,96); ↓ queda de vigor, ↑ aumento de fadiga/PTH; (m) = melhora confiável.',
    ['Atleta','RCI Vigor','RCI Fadiga','RCI PTH/TMD'],rows18,fs=9)
# Tabela 19 – resumo
def pct(v,k): return '%d (%d%%)'%(S[v][k],int(S[v]['pct_'+k]))
anchor=table_after(anchor,'Tabela 19 – Síntese do RCI: fidedignidade, erro-padrão da diferença e número de atletas com mudança confiável.',
    ['Escala','α','EP_dif','n','Piora confiável','Estável','Melhora confiável'],
    [['Vigor',c2('%.2f'%A['Vigor']),c2('%.2f'%SE['Vigor']),S['Vigor']['n'],pct('Vigor','piora'),pct('Vigor','estavel'),pct('Vigor','melhora')],
     ['Fadiga',c2('%.2f'%A['Fadiga']),c2('%.2f'%SE['Fadiga']),S['Fadiga']['n'],pct('Fadiga','piora'),pct('Fadiga','estavel'),pct('Fadiga','melhora')],
     ['PTH/TMD',c2('%.2f'%A['TMD']),c2('%.2f'%SE['TMD']),S['TMD']['n'],pct('TMD','piora'),pct('TMD','estavel'),pct('TMD','melhora')]],fs=9.5)
anchor=fig_after(anchor,'/home/user/mdlucca/Artigos/figuras/x_rci.png',
    'Figura 22 – RCI individual do Dia 1 ao Dia 7 por atleta (ordenado). Vermelho = deterioração confiável (queda de vigor / aumento de fadiga ou PTH); verde = melhora confiável; cinza = mudança dentro do erro de medida. Linhas tracejadas: ±1,96.')

# Referência JACOBSON & TRUAX (alfabético, antes de JARIC)
try:
    jaric=find('JARIC, S.')
    np_=d.add_paragraph('JACOBSON, N. S.; TRUAX, P. Clinical significance: a statistical approach to defining meaningful change in psychotherapy research. Journal of Consulting and Clinical Psychology, v. 59, n. 1, p. 12–19, 1991. DOI: 10.1037/0022-006X.59.1.12.')
    jaric._p.addprevious(np_._p); _fmt(np_,size=12,just=False)
except KeyError: pass

d.save(DOC); print('SAVED',DOC)
