# -*- coding: utf-8 -*-
import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
R=json.load(open('brums_desc.json'))
FG='/home/user/mdlucca/Artigos/figuras'
def c2(s): return str(s).replace('.',',')
doc=Document()
st=doc.styles['Normal']; st.font.name='Times New Roman'; st.font.size=Pt(12)
st.element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')
st.paragraph_format.line_spacing=1.5; st.paragraph_format.space_after=Pt(0)
s=doc.sections[0]; s.top_margin=Cm(3); s.left_margin=Cm(3); s.bottom_margin=Cm(2); s.right_margin=Cm(2)
_TN=[0]; _FN=[0]
def P(t='',just=True,size=12,after=6,bold=False,ind=True):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(size); r.bold=bold
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(after)
    if just: p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if just and ind: p.paragraph_format.first_line_indent=Cm(1.25)
    return p
def H(t,size=12):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(size)
    p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.5
def _bd(c):
    tcPr=c._tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for e_ in ['top','bottom','left','right']:
        el=OxmlElement(f'w:{e_}'); on=e_ in ('top','bottom')
        el.set(qn('w:val'),'single' if on else 'nil'); el.set(qn('w:sz'),'6'); el.set(qn('w:color'),'000000'); b.append(el)
    tcPr.append(b)
def table(cap,header,rows,fonte='Fonte: dados da pesquisa (2026).',fs=9):
    _TN[0]+=1
    p=doc.add_paragraph(); r=p.add_run('Tabela %d – %s'%(_TN[0],cap)); r.font.size=Pt(11)
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
    t=doc.add_table(rows=1,cols=len(header)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,htx in enumerate(header):
        cc=t.rows[0].cells[i]; cc.text=''; rr=cc.paragraphs[0].add_run(htx); rr.bold=True; rr.font.size=Pt(fs); rr.font.name='Times New Roman'
        cc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; _bd(cc)
    for row in rows:
        cs=t.add_row().cells
        for i,val in enumerate(row):
            cs[i].text=''; rr=cs[i].paragraphs[0].add_run(str(val)); rr.font.size=Pt(fs); rr.font.name='Times New Roman'
            cs[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if i>0 else WD_ALIGN_PARAGRAPH.LEFT; _bd(cs[i])
    pf=doc.add_paragraph(); rf=pf.add_run(fonte); rf.font.size=Pt(9); pf.paragraph_format.space_after=Pt(6)
def figure(path,cap,w=15.0):
    _FN[0]+=1
    pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; pp.add_run().add_picture(path,width=Cm(w)); pp.paragraph_format.space_before=Pt(6)
    pc=doc.add_paragraph(); pc.alignment=WD_ALIGN_PARAGRAPH.CENTER; rc=pc.add_run('Figura %d – %s'%(_FN[0],cap)); rc.font.size=Pt(11)
    pf=doc.add_paragraph(); pf.alignment=WD_ALIGN_PARAGRAPH.CENTER; rf=pf.add_run('Fonte: elaboração dos autores (2026).'); rf.font.size=Pt(9); pf.paragraph_format.space_after=Pt(6)

SUBS=[('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Vigor','Vigor'),('Fadiga','Fadiga'),('Confusao','Confusão'),('TMD','PTH/TMD')]

# ===== TÍTULO =====
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('PERFIL DE HUMOR DE ATLETAS DE HANDEBOL DE ELITE NA ÚLTIMA SEMANA DE PRÉ-TEMPORADA: ANÁLISE DESCRITIVA E COMPARATIVA DAS VARIÁVEIS DO BRUMS')
r.bold=True; r.font.size=Pt(13); p.paragraph_format.space_after=Pt(10)

H('RESUMO')
pr=R['prepos']; d17=R['d1d7']; prof=R['profiles']
P('Objetivo: avaliar o perfil de humor de atletas de handebol de elite na última semana de pré-temporada. Método: estudo '
 'descritivo de medidas repetidas com %d atletas, %d observações do BRUMS-24 em sete dias, duas coletas/dia (pré e '
 'pós-treino). Analisou-se o comportamento geral do grupo, a resposta pré→pós, o padrão diário e intra-dia, e a '
 'classificação nos seis perfis de humor, com percentuais, intervalos de confiança e tamanhos de efeito. Resultados: a '
 'deterioração concentra-se no eixo energia–fadiga — do pré ao pós, o vigor cai %s%% (dz = %s) e a fadiga sobe %s%% '
 '(dz = %s); do Dia 1 ao Dia 7, o vigor recua com efeito grande (dz = %s). O perfil iceberg cai de %d%% para %d%% e o '
 'perfil de fadiga (barbatana de tubarão) sobe de %d%% para %d%%. Conclusão: o humor migra da prontidão para a fadiga '
 'funcional, sem instalação de perfis de risco, com o eixo energia–fadiga como marcador sensível do microciclo.'%(
   R['sample']['n'],R['sample']['n_obs'],c2('%.0f'%abs(pr['Vigor']['pct'])),c2('%.2f'%pr['Vigor']['dz']),
   c2('%.0f'%pr['Fadiga']['pct']),c2('%.2f'%pr['Fadiga']['dz']),c2('%.2f'%d17['Vigor']['dz']),
   int(prof['D1']['Iceberg']),int(prof['D7']['Iceberg']),int(prof['D1']['Barbatana tubarão']),int(prof['D7']['Barbatana tubarão'])),after=8)

H('1 INTRODUÇÃO')
P('O monitoramento dos estados de humor é uma ferramenta prática e não invasiva para acompanhar a resposta de atletas à '
 'carga de treino e sinalizar precocemente a fadiga acumulada (SAW; MAIN; GASTIN, 2016). A Brunel Mood Scale (BRUMS), '
 'derivada do Profile of Mood States, mede seis dimensões — tensão, depressão, raiva, vigor, fadiga e confusão — e é um '
 'dos instrumentos de humor mais usados no esporte, com validações transculturais recentes que reforçam sua robustez '
 '(TERRY; LANE; FOGARTY, 2003; ROHLFS et al., 2008; LEW et al., 2023). Mais do que escores isolados, o humor pode ser '
 'lido como um perfil integrado: Morgan (1985) descreveu o “perfil iceberg”, e trabalhos recentes formalizaram seis '
 'perfis prototípicos, úteis ao rastreamento da prontidão e da saúde mental (PARSONS-SMITH; TERRY; MACHIN, 2017; TERRY '
 'et al., 2021; TAN et al., 2024).')
P('Em modalidades coletivas e intermitentes de alta demanda, como o handebol, a resposta afetiva oscila conforme a fase '
 'de treino, e o perfil de humor associa-se ao sono, ao estresse e ao desempenho (BRANDT; BEVILACQUA; ANDRADE, 2017; '
 'MICHALSIK; MADSEN; AAGAARD, 2013; RATZ-SULYOK et al., 2026). Metanálises confirmam a relação entre a perturbação do '
 'humor e o desempenho esportivo (LOCHBAUM et al., 2021). A última semana de pré-temporada concentra a carga que antecede '
 'a competição e é uma janela crítica para observar, de forma descritiva e comparativa, como cada dimensão do humor se '
 'comporta ao longo dos dias e dentro de cada dia.')

H('2 OBJETIVO')
P('Avaliar o perfil de humor de atletas de handebol de elite na última semana de pré-temporada, descrevendo e comparando '
 'o comportamento das variáveis do BRUMS no grupo (geral, pré→pós, por dia e intra-dia) e classificando os perfis de '
 'humor, com estimativas de variação percentual, intervalos de confiança e tamanho de efeito, do nível do grupo ao '
 'nível do sujeito.')

H('3 MÉTODO')
H('3.1 Amostra e desenho',12)
sm=R['sample']
P('Participaram %d atletas de handebol do sexo masculino de nível competitivo (Tabela 1). O monitoramento ocorreu ao '
 'longo de sete dias consecutivos (21–27/04/2024), com duas aplicações do BRUMS por dia de treino — a primeira tomada '
 'como pré e a última como pós —, totalizando %d observações válidas. A Figura 1 sintetiza o framework das coletas. '
 'Distribuição por posição: %s.'%(sm['n'],sm['n_obs'],
   ', '.join('%s (%d)'%(k,v) for k,v in sm['pos'].items())))
def srow(lab,k,unit=''):
    v=sm[k]; return [lab,c2('%.1f'%v['mean']),c2('%.1f'%v['sd']),'%s–%s'%(c2('%.1f'%v['lo']),c2('%.1f'%v['hi'])),'%s–%s'%(c2('%.1f'%v['mn']),c2('%.1f'%v['mx']))]
table('Caracterização sociodemográfica e antropométrica da amostra (n = %d).'%sm['n'],
    ['Variável','Média','DP','IC95%','Mín–Máx'],
    [srow('Idade (anos)','idade'),srow('Estatura (cm)','estatura'),srow('Massa (kg)','massa'),
     srow('% de gordura','pG'),srow('Experiência (anos)','exp')])
figure(f'{FG}/xb_framework.png','Framework das coletas: dos atletas às observações do BRUMS e à classificação de perfis.',w=10.5)
H('3.2 Instrumento e procedimentos',12)
P('O humor foi avaliado pela BRUMS-24 (24 adjetivos, escala 0–4; seis subescalas de 0–16 e a Perturbação Total do Humor, '
 'PTH = tensão + depressão + raiva + fadiga + confusão − vigor). Os questionários foram autoaplicados por formulário '
 'eletrônico com carimbo de data/hora.')
H('3.3 Análise estatística',12)
P('Estatística descritiva (média, mediana, desvio-padrão, IC95%, coeficiente de variação e amplitude) de cada subescala '
 'e da PTH, no geral, por dia e nos momentos pré e pós. Comparações pré→pós e Dia 1→Dia 7 por teste de Wilcoxon pareado, '
 'com variação percentual, tamanho de efeito pareado (dz de Cohen) e classificação de magnitude (trivial < 0,2; pequeno '
 '< 0,5; médio < 0,8; grande ≥ 0,8). Classificação de cada observação nos seis perfis de humor de Terry por proximidade '
 'aos protótipos em escores padronizados. Análises no nível do grupo e no nível do sujeito.')

H('4 RESULTADOS')
H('4.1 Análise descritiva geral do grupo',12)
P('As subescalas negativas situam-se próximas do piso (medianas baixas, coeficientes de variação elevados), de modo que '
 'a variabilidade útil concentra-se no eixo energia–fadiga (Tabela 2): o vigor e a fadiga são as dimensões de maior '
 'média e dispersão informativa.')
def drow(k,lab):
    v=R['desc'][k]; return [lab,c2('%.2f'%v['mean']),c2('%.1f'%v['md']),c2('%.2f'%v['sd']),
        '%s–%s'%(c2('%.2f'%v['lo']),c2('%.2f'%v['hi'])),c2('%.0f'%v['cv']),'%s–%s'%(c2('%.0f'%v['mn']),c2('%.0f'%v['mx']))]
table('Estatística descritiva geral do grupo por subescala do BRUMS (%d observações).'%R['sample']['n_obs'],
    ['Subescala','Média','Mediana','DP','IC95%','CV (%)','Mín–Máx'],[drow(k,l) for k,l in SUBS])

H('4.2 Comportamento ao longo da semana e do Dia 1 ao Dia 7',12)
P('Ao longo da semana (Figura 2), o vigor declina e a fadiga e a PTH ascendem, ao passo que as subescalas negativas de '
 'valência não fadiga oscilam próximas do piso. Na comparação Dia 1 → Dia 7 (Tabela 3), o vigor recua com efeito grande '
 '(Δ = %s; %s%%; dz = %s; p = %s) e a fadiga sobe com efeito médio (Δ = %s; %s%%; dz = %s; p = %s); a tensão e a confusão, '
 'ao contrário, decrescem — acomodação do estado inicial de ativação.'%(
   c2('%+.2f'%d17['Vigor']['delta']),c2('%.0f'%d17['Vigor']['pct']),c2('%.2f'%d17['Vigor']['dz']),c2('%.3f'%d17['Vigor']['p']),
   c2('%+.2f'%d17['Fadiga']['delta']),c2('%.0f'%d17['Fadiga']['pct']),c2('%.2f'%d17['Fadiga']['dz']),c2('%.3f'%d17['Fadiga']['p'])))
figure(f'{FG}/xb_traj.png','Trajetória das subescalas do BRUMS ao longo da semana (médias diárias; barras = IC95% no eixo energia–fadiga).')
def d17row(k,lab):
    v=d17[k]; return [lab,c2('%.2f'%v['d1']),c2('%.2f'%v['d7']),c2('%+.2f'%v['delta']),c2('%+.0f'%v['pct'])+'%',c2('%+.2f'%v['dz']),c2('%.3f'%v['p']),v['mag']]
table('Comparação Dia 1 → Dia 7 por subescala: médias, variação, tamanho de efeito e magnitude.',
    ['Subescala','Dia 1','Dia 7','Δ','Δ%','dz','p','Magnitude'],[d17row(k,l) for k,l in SUBS],fs=8.5)

H('4.3 Resposta aguda pré→pós da semana',12)
P('Agregando toda a semana, a resposta aguda ao treino (pré → pós) é significativa e de magnitude média no eixo '
 'energia–fadiga (Tabela 4): o vigor cai %s%% (dz = %s; p = %s), a fadiga sobe %s%% (dz = %s; p = %s) e a PTH aumenta %s%% '
 '(dz = %s). As demais subescalas apresentam resposta trivial a pequena, sem significância — a deterioração aguda poupa o '
 'afeto negativo geral.'%(
   c2('%.0f'%abs(pr['Vigor']['pct'])),c2('%.2f'%pr['Vigor']['dz']),c2('%.3f'%pr['Vigor']['p']),
   c2('%.0f'%pr['Fadiga']['pct']),c2('%.2f'%pr['Fadiga']['dz']),c2('%.3f'%pr['Fadiga']['p']),
   c2('%.0f'%pr['TMD']['pct']),c2('%.2f'%pr['TMD']['dz'])))
def prow(k,lab):
    v=pr[k]; return [lab,c2('%.2f'%v['pre']),c2('%.2f'%v['pos']),c2('%+.2f'%v['delta']),c2('%+.0f'%v['pct'])+'%',
        '%s [%s, %s]'%(c2('%+.2f'%v['dz']),c2('%.2f'%v['lo']),c2('%.2f'%v['hi'])),c2('%.3f'%v['p']),v['mag']]
table('Resposta pré → pós agregada na semana: médias, variação percentual, dz (IC95%), p e magnitude.',
    ['Subescala','Pré','Pós','Δ','Δ%','dz [IC95%]','p','Magnitude'],[prow(k,l) for k,l in SUBS],fs=8.5)

H('4.4 Comportamento intra-dia (pré → pós por dia)',12)
P('O padrão intra-dia repete-se com regularidade ao longo da semana (Figura 3; Tabela 5): em todos os dias de treino o '
 'vigor cai e a fadiga e a PTH sobem do pré para o pós, sinal de um choque agudo consistente a cada sessão.')
figure(f'{FG}/xb_intraday.png','Escores pré e pós-treino por dia para vigor, fadiga e PTH/TMD (médias do grupo).')
ID=R['intraday']; hdr=['Variável']+['Dia %d'%d for d in range(1,8)]
def idrow(k,lab):
    row=[lab]
    for d in range(1,8):
        c=ID[k].get(str(d));
        if c and c.get('dz')==c.get('dz'):
            star='*' if (c.get('p') is not None and c['p']<0.05) else ''
            row.append('%s (%s)%s'%(c2('%+.1f'%c['delta']),c2('%+.1f'%c['dz']),star))
        else: row.append('—')
    return row
table('Variação intra-dia (Δ pós−pré e dz) por dia. * p < 0,05.',hdr,
    [idrow('Vigor','Vigor'),idrow('Fadiga','Fadiga'),idrow('TMD','PTH/TMD')],fs=8.5)

H('4.5 Perfis de humor e classificação do grupo',12)
P('Na linha de base (Dia 1) predomina o perfil iceberg (%d%%), marca de prontidão; no último dia (Dia 7), o perfil de '
 'fadiga (“barbatana de tubarão”) torna-se o mais frequente (%d%%) e o iceberg cai para %d%% (Figura 4; Tabela 6). Essa '
 'migração — do iceberg para a assinatura de fadiga funcional, sem instalação relevante de perfis de risco (submerso, '
 'iceberg invertido) — é a leitura central da semana.'%(
   int(prof['D1']['Iceberg']),int(prof['D7']['Barbatana tubarão']),int(prof['D7']['Iceberg'])))
figure(f'{FG}/xb_profiles.png','Distribuição (%) dos seis perfis de humor por recorte: linha de base (D1), semana, último dia (D7), pré e pós.')
PROF=[('Iceberg','Iceberg'),('Everest invertido','Everest invertido'),('Iceberg invertido','Iceberg invertido'),('Submerso','Submerso'),('Barbatana tubarão','Barbatana de tubarão'),('Superfície','Superfície')]
recs=[('D1','D1'),('Semana','overall'),('D7','D7'),('Pré','pre'),('Pós','pos')]
table('Distribuição (%) dos perfis de humor de Terry por recorte.',
    ['Perfil']+[l for l,_ in recs],
    [[lab]+[c2('%.0f'%prof[key][p]) for _,key in recs] for p,lab in PROF],fs=9)

H('4.6 Análise no nível do sujeito',12)
sub=R['subject']
P('No nível individual, a resposta é heterogênea, mas majoritariamente de deterioração no eixo energia–fadiga: do Dia 1 '
 'ao Dia 7, %d%% dos atletas reduzem o vigor e %d%% aumentam a fadiga. As médias semanais individuais variam amplamente '
 '(vigor de %s a %s; fadiga de %s a %s), o que reforça a recomendação de interpretar o humor referenciado à linha de base '
 'de cada atleta, e não apenas pela média do grupo.'%(
   int(sub['Vigor']['pct_worse']),int(sub['Fadiga']['pct_worse']),
   c2('%.1f'%sub['Vigor']['wk_min']),c2('%.1f'%sub['Vigor']['wk_max']),
   c2('%.1f'%sub['Fadiga']['wk_min']),c2('%.1f'%sub['Fadiga']['wk_max'])))

H('5 SÍNTESE DOS PRINCIPAIS ACHADOS')
for b in [
 'A deterioração do humor concentra-se no eixo energia–fadiga (vigor ↓, fadiga ↑, PTH ↑); as subescalas negativas de valência não fadiga permanecem próximas do piso.',
 'Resposta aguda pré→pós de magnitude média e significativa: vigor −%s%% (dz = %s), fadiga +%s%% (dz = %s), PTH +%s%% (dz = %s).'%(c2('%.0f'%abs(pr['Vigor']['pct'])),c2('%.2f'%pr['Vigor']['dz']),c2('%.0f'%pr['Fadiga']['pct']),c2('%.2f'%pr['Fadiga']['dz']),c2('%.0f'%pr['TMD']['pct']),c2('%.2f'%pr['TMD']['dz'])),
 'Do Dia 1 ao Dia 7, o vigor recua com efeito grande (dz = %s) e a fadiga sobe com efeito médio (dz = %s); tensão e confusão se acomodam.'%(c2('%.2f'%d17['Vigor']['dz']),c2('%.2f'%d17['Fadiga']['dz'])),
 'O perfil de humor migra do iceberg (%d%% no D1) para a fadiga funcional (barbatana de tubarão, %d%% no D7), sem instalação de perfis de risco.'%(int(prof['D1']['Iceberg']),int(prof['D7']['Barbatana tubarão'])),
 'No nível do sujeito, a maioria deteriora de forma coerente com o grupo, com ampla variabilidade individual — monitorar por tendência e por linha de base individual.']:
    p=doc.add_paragraph(); r=p.add_run('•  '+b); r.font.size=Pt(12); p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(4); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

H('REFERÊNCIAS')
refs=[
 'BRANDT, R.; BEVILACQUA, G. G.; ANDRADE, A. Perceived sleep quality, mood states, and their relationship with performance among Brazilian elite athletes during a competitive period. Journal of Strength and Conditioning Research, v. 31, n. 4, p. 1033–1039, 2017.',
 'LEW, P. C. F. et al. Cross-cultural validation of the Malaysian Mood Scale and tests of between-group mood differences. International Journal of Environmental Research and Public Health, v. 20, n. 4, 3348, 2023. DOI: 10.3390/ijerph20043348.',
 'LOCHBAUM, M. et al. The Profile of Mood States and athletic performance: a meta-analysis of published studies. European Journal of Investigation in Health, Psychology and Education, v. 11, n. 1, p. 50–70, 2021. DOI: 10.3390/ejihpe11010005.',
 'MICHALSIK, L. B.; MADSEN, K.; AAGAARD, P. Match performance and physiological capacity of female elite team handball players. International Journal of Sports Medicine, v. 35, n. 7, p. 595–607, 2013. DOI: 10.1055/s-0033-1358713.',
 'MORGAN, W. P. Selected psychological factors limiting performance: a mental health model. In: CLARKE, D. H.; ECKERT, H. M. (Ed.). Limits of human performance. Champaign: Human Kinetics, 1985. p. 70–80.',
 'PARSONS-SMITH, R. L.; TERRY, P. C.; MACHIN, M. A. Identification and description of novel mood profile clusters. Frontiers in Psychology, v. 8, 1958, 2017. DOI: 10.3389/fpsyg.2017.01958.',
 'RATZ-SULYOK, F. Z. et al. Associations between endocrine status and stress, mood and psychosomatic status in elite handball players. Sports, v. 14, n. 7, 289, 2026. DOI: 10.3390/sports14070289.',
 'ROHLFS, I. C. P. M. et al. A Escala de Humor de Brunel (Brums): instrumento para detecção precoce da síndrome do excesso de treinamento. Revista Brasileira de Medicina do Esporte, v. 14, n. 3, p. 176–181, 2008.',
 'SAW, A. E.; MAIN, L. C.; GASTIN, P. B. Monitoring the athlete training response: subjective self-reported measures trump commonly used objective measures: a systematic review. British Journal of Sports Medicine, v. 50, n. 5, p. 281–291, 2016. DOI: 10.1136/bjsports-2015-094758.',
 'TAN, C. et al. The structural validity and latent profile characteristics of the Abbreviated Profile of Mood States among Chinese athletes. BMC Psychiatry, v. 24, n. 1, 636, 2024. DOI: 10.1186/s12888-024-06092-5.',
 'TERRY, P. C.; LANE, A. M.; FOGARTY, G. J. Construct validity of the Profile of Mood States — Adolescents for use with adults. Psychology of Sport and Exercise, v. 4, n. 2, p. 125–139, 2003. DOI: 10.1016/S1469-0292(02)00035-8.',
 'TERRY, P. C. et al. Mood profiling for sustainable mental health among athletes. Sustainability, v. 13, n. 11, 6116, 2021. DOI: 10.3390/su13116116.']
for rf in refs:
    p=doc.add_paragraph(); r=p.add_run(rf); r.font.size=Pt(11); p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(6); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

OUTP='/home/user/mdlucca/Artigos/Perfil_Humor_BRUMS_Descritivo.docx'
doc.save(OUTP); print('SAVED',OUTP,'| Tabelas',_TN[0],'Figuras',_FN[0])
