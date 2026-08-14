import warnings; warnings.filterwarnings('ignore')
import numpy as np, pandas as pd, json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
SC='/tmp/claude-0/-home-user-mdlucca/e1dba24c-b1d7-5908-9106-f2f4aaf3f56a/scratchpad'
FIG='/home/user/mdlucca/Artigos/figuras'
hum=pd.read_csv(f'{SC}/humor.csv'); hp=pd.read_csv(f'{SC}/hum_prof.csv')
norm=json.load(open(f'{SC}/norm_data.json')); app=json.load(open(f'{SC}/app_data.json'))

doc=Document()
# base styles
st=doc.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(10.5)
def H(t,l=1): doc.add_heading(t,level=l)
def P(t,it=False,b=False):
    p=doc.add_paragraph(); r=p.add_run(t); r.italic=it; r.bold=b; return p
def fig(name,cap,w=6.3):
    import os
    path=f'{FIG}/{name}.png'
    if os.path.exists(path):
        doc.add_picture(path,width=Inches(w))
        doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
        c=doc.add_paragraph(); r=c.add_run(cap); r.italic=True; r.font.size=Pt(8.5); c.alignment=WD_ALIGN_PARAGRAPH.CENTER
def table(headers,rows,cap=None):
    if cap: c=doc.add_paragraph(); r=c.add_run(cap); r.bold=True; r.font.size=Pt(9.5)
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Light Grid Accent 1'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        cell=t.rows[0].cells[i]; cell.paragraphs[0].add_run(str(h)).bold=True
        cell.paragraphs[0].runs[0].font.size=Pt(8.5)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=str(v); cells[i].paragraphs[0].runs[0].font.size=Pt(8.5) if cells[i].paragraphs[0].runs else None
    doc.add_paragraph()

# ============ CAPA ============
ti=doc.add_heading('Análise Estatística do Monitoramento do Humor e da Carga',0)
P('Microciclo de choque de HIIT em atletas de handebol de elite — 21 a 28 de abril de 2024',it=True)
P('Relatório analítico completo · 27 atletas · 456 observações · reanálise independente (Python) · via não paramétrica. Atletas anonimizados (A01–A27).')
doc.add_paragraph()

# ============ 1. RESUMO ============
H('1. Resumo executivo')
P('O humor deteriorou-se de forma não linear (cúbica) ao longo do microciclo, precipitando-se no Dia 7, com a mudança concentrada no eixo energia–fadiga (vigor↓, fadiga↑). A fadiga física foi o marcador mais precoce, sensível e confiável. O efeito do HIIT é de acúmulo (não choque agudo isolado) e confunde-se com o volume. As variáveis de humor/bem-estar são não normais, exigindo métodos não paramétricos. O perfil "iceberg" caiu de 40% para 17% e a "barbatana de tubarão" subiu de 2% para 28% (fadiga somática sem colapso afetivo). A carga externa é fixa (prescrição relativa), mas a interna revela fadiga (FC de pico↓, PSE↑; session-RPE > TRIMP). A variância é dominada por traço, exigindo médias de ≥3–7 coletas para decisão individual.')

# ============ 2. MÉTODOS ============
H('2. Abordagem estatística')
P('Delineamento observacional longitudinal, medidas repetidas (observações ⊂ dias ⊂ atletas). Unidade amostral para contrastes: o atleta (agregação por atleta), corrigindo a pseudorreplicação. α=0,05; Python (numpy, scipy, statsmodels). A normalidade foi testada por Shapiro–Wilk (com assimetria/curtose) em dois níveis (distribuição bruta e escores de mudança por atleta). Como as variáveis de humor/bem-estar são não normais, adotou-se a via NÃO PARAMÉTRICA: mediana (IQR), Spearman/rmcorr, Wilcoxon/Mann-Whitney, Kruskal-Wallis, PERMANOVA (permutação restrita ao atleta), bootstrap e fatores de Bayes/ROPE. Variáveis contínuas normais (aptidão/antropometria) admitem via paramétrica.')

# ============ 3. AMOSTRA ============
H('3. Caracterização da amostra')
ag=app['socio']['agg']; pc=app['socio']['pos_counts']; n=sum(pc.values())
table(['Variável','Valor'],[
 ['Idade (anos)',f"{ag['idade'][0]} ± {ag['idade'][1]} ({ag['idade'][2]}–{ag['idade'][3]})"],
 ['Experiência (anos)',f"{ag['exp'][0]} ± {ag['exp'][1]}"],
 ['Estatura (cm)',f"{ag['estatura'][0]} ± {ag['estatura'][1]}"],
 ['Massa (kg)',f"{ag['massa'][0]} ± {ag['massa'][1]}"],
 ['Gordura (%)',f"{ag['pG'][0]} ± {ag['pG'][1]}"],
 ['T-CAR PV (km/h)',f"{ag['PV'][0]} ± {ag['PV'][1]}"],
 ['Posição',' · '.join(f"{k} {round(100*v/n)}%" for k,v in pc.items())],
],'Tabela 1. Características da amostra (n=27).')

# ============ 4. NORMALIDADE ============
H('4. Teste de normalidade (Shapiro–Wilk)')
rows=[[r['lab'],r['n'],r['W'],(f"{r['p']:.1e}"),r['skew'],('sim' if r['normal'] else 'NÃO'),('paramétrica' if r['normal'] else 'não paramétrica')] for r in norm['dist']]
table(['Variável','n','W','Shapiro p','Assim.','Normal?','Família'],rows,'Tabela 2. Normalidade da distribuição (14/18 não normais).')
P('Decisão: humor e bem-estar não normais → via NÃO PARAMÉTRICA. Nos escores de mudança (D7−D1), o eixo energia–fadiga é aproximadamente normal, mas as subescalas negativas não (Wilcoxon).',it=True)

# ============ 5. DESCRITIVAS ============
H('5. Estatística descritiva (mediana [IQR]; % piso)')
def med(v):
    s=hum[v].dropna(); return [round(s.median(),1),round(s.quantile(.75)-s.quantile(.25),1),round(s.skew(),2),round(100*(s==s.min()).mean(),1)]
rows=[]
for v,l in [('Vigor','Vigor'),('Fadiga','Fadiga BRUMS'),('FadFisica','Fadiga física'),('FadMental','Fadiga mental'),('TMD','PTH/TMD'),('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Confusao','Confusão')]:
    m=med(v); rows.append([l,f"{m[0]} [{m[1]}]",m[2],f"{m[3]}%"])
table(['Variável','Mediana [IQR]','Assim.','% piso'],rows,'Tabela 3. Descritivas robustas (456 observações).')
fig('assoc_scatter','Figura 1. Associação Fadiga física × Vigor (oposição energia–fadiga).',5.6)

# ============ 6. PSICOMETRIA ============
H('6. Confiabilidade e estrutura (psicometria)')
table(['Índice','Resultado'],[
 ['Consistência interna (α)','Fadiga 0,80 · Depressão 0,85 · Raiva 0,87 · Tensão 0,43 (piso)'],
 ['Adequação fatorial','KMO=0,835; Bartlett χ²(276)=5.228; p<0,001; 6 fatores'],
 ['Decomposição de variância','Traço (ICC atleta) 0,26–0,67 — humor dominantemente traço'],
 ['Precisão da decisão','1 coleta Φ=0,57; ≥3 coletas Φ≥0,80; ICC(1,7)=0,76–0,94'],
 ['Lei do piso','|dz| = 0,588 − 0,0063·(%piso); ρ=−0,92; R²=0,85'],
],'Tabela 4. Síntese psicométrica.')

# ============ 7. EFEITO DO DIA ============
H('7. Efeito do dia e trajetória temporal')
table(['Variável','F (aprox.)','p (FDR)','ICC atleta'],[
 ['Fadiga física','18,9','<0,001','0,46'],['Vigor','8,4','<0,001','0,57'],
 ['Fadiga (BRUMS)','8,2','<0,001','0,59'],['PTH/TMD','5,1','<0,001','0,59'],
 ['Depressão','1,6','0,150 (n.s.)','0,68'],
],'Tabela 5. Efeito do dia (modelo misto; FDR).')
P('Forma da trajetória: componentes ortogonais LINEAR + CÚBICA significativas (quadrática ausente); crescimento cúbico vence por AIC/LRT (p<0,001). Três fases: subida → alívio no meio → precipitação no Dia 7.')
fig('fig1_trajetoria','Figura 2. Trajetória semanal (faixas = dias de HIIT).')
fig('fig7_cubico','Figura 3. Ajuste polinomial cúbico com IC95% (subida–alívio–precipitação).')
fig('fig2_derivadas','Figura 4. Limites e derivadas: velocidade e aceleração da mudança.')

# ============ 8. RESPOSTA AGUDA ============
H('8. Resposta aguda pré→pós (agregada por atleta; Wilcoxon/dz; FDR)')
table(['Variável','Δ','dz [IC95%]','p (FDR)'],[
 ['Fadiga física','+1,81','+0,97 [0,69; 1,37]','<0,001 ✔'],
 ['Fadiga (BRUMS)','+1,76','+0,59 [0,36; 0,88]','0,019 ✔'],
 ['PTH/TMD','+3,14','+0,57 [0,29; 0,89]','0,019 ✔'],
 ['Vigor','−0,80','−0,45 [−0,80; −0,12]','0,061'],
 ['Subescalas negativas','—','|dz| ≤ 0,40','n.s.'],
],'Tabela 6. Resposta aguda; só o eixo energia–fadiga sobrevive à correção.')
fig('fig9_efeitos','Figura 5. Tamanhos de efeito D1→D7 (dz, IC95% bootstrap).',5.6)

# ============ 9. MULTIVARIADA ============
H('9. Confirmação multivariada e associações')
table(['Teste','Estatística','p'],[
 ['PERMANOVA — pré vs pós','pseudo-F=1,75','0,003'],
 ['PERMANOVA — HIIT vs não-HIIT','pseudo-F=2,71','0,008'],
 ['Hotelling T² — eixo energia–fadiga','F(2,25)=5,59','0,010'],
 ['Misto multivariado — LRT do dia','χ²(4)=118,9','≈9×10⁻²⁵'],
],'Tabela 7. Multivariada (permutação restrita ao atleta).')
P('Estrutura entre-atletas: vigor↔fadiga = −0,65; fadiga↔TMD = +0,86 → eixo bipolar energia↔fadiga como dimensão de traço.')
fig('assoc_corr','Figura 6. Matriz de correlação (ρ de Spearman) — associações entre as variáveis.')
fig('fig6_correlacao','Figura 7. Correlação entre-atletas (modelo misto multivariado).',5.6)

# ============ 10. PERFIS ============
H('10. Perfis de humor (Parsons-Smith) — análise transversal')
tab=pd.crosstab(hp['dia'],hp['perfil'],normalize='index').mul(100).round(0)
order=['Iceberg','Barbatana tubarão','Superfície','Iceberg invertido','Everest invertido','Submerso']
rows=[[f"D{d}"]+[f"{tab.loc[d,c]:.0f}%" if c in tab.columns else '0%' for c in order] for d in [1,4,7]]
table(['Dia']+[c[:10] for c in order],rows,'Tabela 8. Distribuição dos perfis (%) — início/meio/fim.')
P('O grupo começa saudável (iceberg 40%, índice-iceberg z=+0,52) e termina fatigado (iceberg 17%, barbatana 28%, índice=−0,47). Mann-Whitney D1×D7 p=0,0015. Pré→pós: humor piora dentro do dia (Wilcoxon p=0,0005). Maior variação no Dia 6 (Fligner p=0,26, variância homogênea).')
fig('perfil_1_area','Figura 8. Distribuição dos 6 perfis ao longo da semana (iceberg↓, barbatana↑).')
fig('perfil_6_radar','Figura 9. Forma do perfil: Dia 1 (iceberg) vs Dia 7 (achatado).',5.6)

# ============ 11. FADIGA / RELAÇÕES ============
H('11. Relações: fadiga física × mental × BRUMS')
table(['Par','Spearman ρ','rmcorr'],[
 ['Fadiga física × Fadiga BRUMS','+0,68','+0,64'],
 ['Fadiga física × Fadiga mental','+0,50','+0,44'],
 ['Fadiga mental × Fadiga BRUMS','+0,46','+0,45'],
],'Tabela 9. Associações da tríade de fadiga.')
P('Dissociação: fadiga física e do BRUMS sobem juntas e quase dobram (4→7,5); a fadiga mental permanece estável (4,8→5,1). O custo do microciclo é somático, não afetivo — coerente com o perfil "barbatana".')
fig('perfil_5_fadiga','Figura 10. Fadiga física × BRUMS × mental (LOWESS): física e BRUMS sobem juntas; mental dissocia.')

# ============ 12. CARGA ============
H('12. Carga interna e externa')
table(['Domínio','Resultado'],[
 ['Externa (4×4 min @104% PV)','Velocidade 16,5 ± 1,2 km/h; distância 2.929 ± 207 m/sessão; total ≈8,8 km (fixa entre sessões)'],
 ['Interna (S1→S3)','FC pico 184→181 bpm (Friedman p<0,001); TRIMP 50,6→46,4 (cai); PSE/sRPE 196→208 (sobe)'],
 ['Dissociação','TRIMP cardíaco subestima sob fadiga; session-RPE é o marcador interno fiel'],
 ['Acoplamento (rmcorr)','FC pico × vigor +0,57; × TMD −0,48; × fadiga física −0,51'],
],'Tabela 10. Carga do grupo.')
fig('fig3_carga_interna','Figura 11. Dissociação da carga interna (FC/TRIMP↓ vs PSE/sRPE↑).')

# ============ 13. SEGMENTAÇÃO ============
H('13. Segmentação por aptidão (não paramétrica)')
table(['Grupo (tercil T-CAR)','n','D1','D7','dz D1→D7'],[
 ['Aptidão baixa','10','4,9','8,71','+1,62'],['Aptidão média','7','3,71','7,02','+1,50'],['Aptidão alta','8','3,94','7,25','+1,85'],
],'Tabela 11. Fadiga física por tercil de aptidão.')
P('Kruskal-Wallis no D7: H=4,46; p=0,107 (não significativo com n≈8–10/grupo), embora o tamanho de efeito entre extremos seja grande (g=−1,18). Aptidão × fadiga da semana ρ=−0,54 (p=0,005); robusto ao escalonamento alométrico (ρ=−0,52).')

# ============ 14. MODELOS ============
H('14. Modelos avançados (ROC, bayesiano, polinomial)')
table(['Tarefa (ROC)','Melhor variável','AUC [IC95%]'],[
 ['Discriminar sessão de HIIT','Fadiga física','0,54 [0,48; 0,60]'],
 ['Discriminar acúmulo (D7 vs D1)','Fadiga física','0,86 [0,77; 0,92]'],
 ['Discriminar dia de fadiga alta','Fadiga (BRUMS)','0,84 [0,80; 0,88]'],
],'Tabela 12. Discriminância (ROC): o sinal está no acúmulo, não na sessão.')
P('Bayesiano: fadiga física BF₁₀≈2444 (extrema); PTH 22,4; Fadiga 11,3; Vigor 5,9 (efeito). Confusão BF₁₀≈0,23 e 93% na ROPE → evidência para equivalência (ausência de efeito, não só p>0,05). Polinomial: cúbico vence (AIC; LRT p<0,001).')
fig('fig8_roc','Figura 12. Curvas ROC: acúmulo (AUC=0,86) vs sessão (AUC≈0,5).',5.6)

# ============ 15. CONCLUSÕES ============
H('15. Conclusões')
for t in [
 'Pressupostos: humor/bem-estar não normais → via não paramétrica (a paramétrica não é reportada onde a variável é não normal).',
 'Fenômeno: deterioração não linear (cúbica), precipitação no Dia 7, concentrada no eixo energia–fadiga; fadiga física como marcador-sentinela.',
 'HIIT: efeito de acúmulo (só o D7 difere; interação Condição×Momento n.s.), confundido com volume — não é dose-resposta.',
 'Perfis: iceberg→barbatana de tubarão (fadiga somática sem colapso afetivo); fadiga mental dissocia.',
 'Carga: externa fixa; interna revela fadiga; session-RPE > TRIMP.',
 'Medida: variância dominada por traço ⇒ decisão individual por médias de ≥3–7 coletas e MDC; a não-resposta das negativas é efeito piso (Bayes favorável ao nulo).',
 'Robustez: convergência entre modelos mistos, permutação, bayesiano, bootstrap e ROC.',
]:
    doc.add_paragraph(t,style='List Bullet')

# ============ 16. REFERÊNCIAS ============
H('16. Referências')
for r in [
 'Beedie, C., Terry, P. C., & Lane, A. M. (2000). The profile of mood states and athletic performance: Two meta-analyses. Journal of Applied Sport Psychology, 12(1), 49–68.',
 'Bonfim, B. M. A. do, Bonuzzi, G. M. G., & Monteiro, C. B. de M. (2023). Human Movement, 24(4), 72–79.',
 'Horta, T. A. G., Lima, P., & Matta, G. (2020). Revista Brasileira de Medicina do Esporte, 26(2), 158–161.',
 'Yang, S., Yin, Y., & Qiu, Z. (2024). Frontiers in Neuroscience, 18, 1341972.',
 'Liu, H., Yang, W., & Liu, H. (2023). BMC Sports Science, Medicine and Rehabilitation, 15(1).',
 'Parsons-Smith, R. L., Terry, P. C., & Machin, M. A. (2017). Frontiers in Psychology, 8, 1958.',
]:
    p=doc.add_paragraph(r); p.paragraph_format.space_after=Pt(3); p.runs[0].font.size=Pt(9)

doc.save('/home/user/mdlucca/Artigos/Relatorio_Estatistico_Completo.docx')
print('DOCX salvo.')
