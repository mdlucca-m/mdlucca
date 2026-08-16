# -*- coding: utf-8 -*-
import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
R=json.load(open('brums_desc2.json')); RC=json.load(open('rci6.json')); STAT=json.load(open('brums_stats3.json'))
S4=json.load(open('brums_stats4.json')); MS=json.load(open('model_stats.json')); MV=json.load(open('manova.json'))
PV=json.load(open('pv_stats.json')); LIM=json.load(open('tcar_limiar.json')); EX=json.load(open('extra.json')); TCD=json.load(open('tcar_desc.json'))
FG='/home/user/mdlucca/Artigos/figuras'
def c2(s): return str(s).replace('.',',')
doc=Document()
stl=doc.styles['Normal']; stl.font.name='Times New Roman'; stl.font.size=Pt(12)
stl.element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')
stl.paragraph_format.line_spacing=1.5; stl.paragraph_format.space_after=Pt(0)
sec=doc.sections[0]; sec.top_margin=Cm(3); sec.left_margin=Cm(3); sec.bottom_margin=Cm(2); sec.right_margin=Cm(2)
_TN=[0]; _FN=[0]
def P(t='',just=True,size=12,after=6,bold=False,ind=True):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(size); r.bold=bold
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(after)
    if just: p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if just and ind: p.paragraph_format.first_line_indent=Cm(1.25)
    return p
def RUN(pairs,after=6,ind=True):
    p=doc.add_paragraph(); p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(after)
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if ind: p.paragraph_format.first_line_indent=Cm(1.25)
    for txt,bold in pairs:
        r=p.add_run(txt); r.bold=bold; r.font.size=Pt(12)
    return p
def H(t,size=12,before=10):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(size)
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.5
def _bd(c):
    tcPr=c._tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for e_ in ['top','bottom','left','right']:
        el=OxmlElement(f'w:{e_}'); on=e_ in ('top','bottom')
        el.set(qn('w:val'),'single' if on else 'nil'); el.set(qn('w:sz'),'6'); el.set(qn('w:color'),'000000'); b.append(el)
    tcPr.append(b)
def table(cap,header,rows,fonte='Fonte: dados da pesquisa (2026).',fs=9,note=None):
    _TN[0]+=1
    p=doc.add_paragraph(); r=p.add_run('Tabela %d – %s'%(_TN[0],cap)); r.font.size=Pt(11)
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
    t=doc.add_table(rows=1,cols=len(header)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,htx in enumerate(header):
        cc=t.rows[0].cells[i]; cc.text=''; rr=cc.paragraphs[0].add_run(htx); rr.bold=True; rr.font.size=Pt(fs); rr.font.name='Times New Roman'
        cc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; _bd(cc)
    for row in rows:
        cs=t.add_row().cells
        for i,val in enumerate(row):
            cs[i].text=''; rr=cs[i].paragraphs[0].add_run(str(val)); rr.font.size=Pt(fs); rr.font.name='Times New Roman'
            cs[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if i>0 else WD_ALIGN_PARAGRAPH.LEFT; _bd(cs[i])
    if note:
        pn=doc.add_paragraph(); rn=pn.add_run(note); rn.font.size=Pt(8.5); rn.italic=True; pn.paragraph_format.space_after=Pt(0)
    pf=doc.add_paragraph(); rf=pf.add_run(fonte); rf.font.size=Pt(9); pf.paragraph_format.space_after=Pt(6)
    return _TN[0]
def figure(path,cap,w=15.5):
    _FN[0]+=1
    pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; pp.add_run().add_picture(path,width=Cm(w)); pp.paragraph_format.space_before=Pt(6)
    pc=doc.add_paragraph(); pc.alignment=WD_ALIGN_PARAGRAPH.CENTER; rc=pc.add_run('Figura %d – %s'%(_FN[0],cap)); rc.font.size=Pt(11)
    pf=doc.add_paragraph(); pf.alignment=WD_ALIGN_PARAGRAPH.CENTER; rf=pf.add_run('Fonte: elaboração dos autores (2026).'); rf.font.size=Pt(9); pf.paragraph_format.space_after=Pt(6)
    return _FN[0]
def pstr(p): return '< 0,001' if p<0.001 else c2('%.3f'%p)
mvv=lambda tab,k,f: next(x[f] for x in MV[tab]['rows'] if x['k']==k)
def S(n): return 'Material Suplementar (Tabela S%d)'%n if False else n

pr=R['prepos']; d17=R['d1d7']; prof=R['profiles']; sm=R['sample']; foc=R['focus']; desc=R['desc']
SENS=S4['sens']; FR=S4['friedman']; PCT=S4['pct']; PCNT=S4['prof_counts']
MSd=MS['desc']; MScorr=MS['corr']; MSord=MS['order']; PREV=MS['prev']
LP=LIM['LIM']['PVini']; TC=LIM['TERC']; ep=EX['epw']; ps=EX['pss']; sc=EX['sleepy_cmp']; hi=EX['hiit']

# ===================== TÍTULO + RESUMO =====================
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('PERFIL DE HUMOR DE ATLETAS DE HANDEBOL DE ELITE NA ÚLTIMA SEMANA DE PRÉ-TEMPORADA: MONITORAMENTO DAS SEIS DIMENSÕES DO BRUMS E A APTIDÃO INTERMITENTE COMO COVARIÁVEL')
r.bold=True; r.font.size=Pt(13); p.paragraph_format.space_after=Pt(10)
H('RESUMO',before=2)
RUN([('Objetivo: ',True),('avaliar o perfil de humor de handebolistas de elite na última semana de pré-temporada, do grupo '
 'ao atleta, tendo a aptidão aeróbia intermitente (pico de velocidade do T-CAR) como covariável fisiológica. ',False),
 ('Método: ',True),('%d atletas do sexo masculino responderam ao BRUMS-24 duas vezes ao dia por sete dias (%d '
 'observações). Os escores foram convertidos em escores T e as comparações intra-grupo (Dia 1 vs. Dia 7; pré vs. pós) '
 'testadas por MANOVA de medidas repetidas, com correlação de Spearman, classificação nos seis perfis de humor, Índice '
 'de Mudança Confiável (RCI) individual e associação com o pico de velocidade, a sonolência (Epworth) e a carga de HIIT. '%(sm['n'],sm['n_obs']),False),
 ('Resultados: ',True),('a deterioração concentrou-se no eixo energia–fadiga (Wilks λ = %s; p = %s; η²ₚ = %s), com vigor '
 '(d = %s) e fadiga (d = %s); o perfil iceberg caiu e o de fadiga subiu (χ² = %s; p = %s). No nível individual, a fadiga '
 'foi a única dimensão a piorar de forma confiável na maioria (%d%%). Maior aptidão intermitente associou-se a mais vigor '
 '(ρ = %s) e menos fadiga (ρ = %s); a sonolência acompanhou a fadiga (ρ = %s) e os dias de HIIT concentraram pior humor. '%(
   c2('%.3f'%MV['d1d7']['wilks']),pstr(MV['d1d7']['p_mv']),c2('%.2f'%MV['d1d7']['eta_mv']),c2('%+.2f'%mvv('d1d7','Vigor','d')),c2('%+.2f'%mvv('d1d7','Fadiga','d')),
   c2('%.2f'%PREV['chi']),pstr(PREV['p']),RC['Fadiga']['pct_piora'],
   c2('%+.2f'%PV['pv']['wk_Vigor']['TCAR1']['rho']),c2('%+.2f'%PV['pv']['wk_FadFisica']['TCAR1']['rho']),c2('%+.2f'%ep['corr']['Fadiga']['rho'])),False),
 ('Conclusão: ',True),('o humor migra da prontidão para a fadiga funcional; o vigor e a fadiga são as dimensões mais '
 'sensíveis, e a aptidão intermitente e o sono modulam a resposta, fundamentando um monitoramento individualizado.',False)],after=6)
P('Palavras-chave: humor; BRUMS; handebol; monitoramento do atleta; aptidão intermitente; fadiga.',size=11,after=8,ind=False)
P('Nota: análises complementares (pós-teste, ajuste alométrico, derivadas, decomposição sinal–ruído e figuras detalhadas) '
 'constam no Material Suplementar.',size=10,after=8,ind=False)

# ===================== 1 INTRODUÇÃO =====================
H('1 INTRODUÇÃO')
H('1.1 Monitoramento do humor e handebol',12,before=6)
P('As organizações esportivas de elite utilizam rotineiramente indicadores psicológicos para rastrear o bem-estar e a '
 'prontidão, e avaliações regulares do humor têm utilidade preditiva para o bem-estar e o desempenho (SAW; MAIN; GASTIN, '
 '2016; ROHLFS et al., 2023). A Brunel Mood Scale (BRUMS), derivada do POMS, mensura seis dimensões (tensão, depressão, '
 'raiva, vigor, fadiga e confusão) e é amplamente validada, inclusive no Brasil (TERRY; LANE; FOGARTY, 2003; ROHLFS et '
 'al., 2008, 2023). O handebol é uma modalidade coletiva intermitente e de alta intensidade, cujo microciclo '
 'pré-competitivo concentra a carga que antecede a competição (KARCHER; BUCHHEIT, 2014; MICHALSIK; MADSEN; AAGAARD, '
 '2013).')
H('1.2 Perfis de humor, aptidão intermitente e lacuna',12,before=6)
P('Além dos escores isolados, seis perfis prototípicos de humor (iceberg, Everest invertido, iceberg invertido, submerso, '
 'barbatana de tubarão e superfície) são amplamente empregados no rastreamento da prontidão (MORGAN, 1985; PARSONS-SMITH; '
 'TERRY; MACHIN, 2017; HAN; PARSONS-SMITH; TERRY, 2020). A tolerância à carga intermitente depende da aptidão aeróbia '
 'intermitente, cujo marcador de campo — o pico de velocidade do Teste de Carminatti (T-CAR) — associa-se ao desempenho '
 'físico em modalidades intermitentes (FERNANDES-DA-SILVA et al., 2016). Faltam, contudo, descrições detalhadas do '
 'comportamento do humor em um microciclo de handebol de elite, do grupo ao atleta, e de como a aptidão física modula '
 'essa resposta.')
H('1.3 Objetivos e hipóteses',12,before=6)
P('O objetivo foi avaliar o perfil de humor de handebolistas de elite na última semana de pré-temporada, do grupo ao '
 'atleta, tendo a aptidão aeróbia intermitente como covariável. Hipóteses: (H1) a deterioração concentra-se no eixo '
 'energia–fadiga; (H2) o perfil migra do iceberg para o de fadiga; (H3) o afeto negativo acopla-se à fadiga sob carga '
 'acumulada; e (H4) a resposta é heterogênea e modulada pela aptidão e pelo sono.')

# ===================== 2 MÉTODOS =====================
H('2 MATERIAIS E MÉTODOS')
H('2.1 Participantes',12,before=6)
P('Participaram %d atletas de handebol do sexo masculino de nível competitivo (idade %s ± %s anos; estatura %s ± %s cm; '
 'massa %s ± %s kg), avaliados em condições ecológicas durante o microciclo pré-competitivo (21–27 de abril de 2024), '
 'conforme a Declaração de Helsinque, com consentimento informado. A distribuição por posição consta na Tabela 1 e a '
 'antropometria na Tabela 2.'%(sm['n'],c2('%.1f'%sm['idade']['mean']),c2('%.1f'%sm['idade']['sd']),
   c2('%.1f'%sm['estatura']['mean']),c2('%.1f'%sm['estatura']['sd']),c2('%.1f'%sm['massa']['mean']),c2('%.1f'%sm['massa']['sd'])))
t_dist=table('Distribuição demográfica da amostra (n = %d).'%sm['n'],['Fonte','Grupo','n','%'],
    [['Sexo','Masculino',sm['n'],'100,0']]+[['Posição' if i==0 else '',k,v,c2('%.1f'%(100*v/sm['n']))] for i,(k,v) in enumerate(sm['pos'].items())]+[['Total','Todos',sm['n'],'100,0']],fs=9)
def srow(lab,k):
    v=sm[k]; return [lab,c2('%.1f'%v['mean']),c2('%.1f'%v['sd']),'%s–%s'%(c2('%.1f'%v['lo']),c2('%.1f'%v['hi'])),'%s–%s'%(c2('%.1f'%v['mn']),c2('%.1f'%v['mx']))]
t_anth=table('Caracterização antropométrica e de experiência (n = %d).'%sm['n'],['Variável','Média','DP','IC95%','Mín–Máx'],
    [srow('Idade (anos)','idade'),srow('Estatura (cm)','estatura'),srow('Massa corporal (kg)','massa'),srow('Percentual de gordura','pG'),srow('Experiência (anos)','exp')])
H('2.2 Medida do humor',12,before=6)
P('O humor foi avaliado pela BRUMS-24 (seis subescalas de 0 a 16), autoaplicada duas vezes ao dia (pré e pós-treino) por '
 'sete dias. A Perturbação Total do Humor (PTH) resume o perfil (soma das negativas − vigor). Para a classificação de '
 'perfis, os escores foram convertidos em escores T (M = 50; DP = 10).')
H('2.3 Aptidão intermitente (T-CAR) — covariável fisiológica',12,before=6)
P('A aptidão aeróbia intermitente foi avaliada pelo Teste de Carminatti (T-CAR), teste de campo progressivo e '
 'intermitente (12 s de corrida em vaivém por 6 s de recuperação, com incrementos até a exaustão), tendo o pico de '
 'velocidade (PV) como desfecho (FERNANDES-DA-SILVA et al., 2016). O T-CAR foi aplicado em 15 de abril de 2024, quatro '
 'dias de treino antes do início do microciclo monitorado (21–27 de abril de 2024), de modo a funcionar como parâmetro '
 'fisiológico (covariável) de linha de base. A caracterização completa do desempenho no teste consta na Tabela 3; o PV '
 'inicial (T-CAR1) foi a referência das análises.')
def tcrow(o): return [o['lab'],o['n'],c2('%.1f'%o['m']),c2('%.1f'%o['sd']),'%s–%s'%(c2('%.1f'%o['mn']),c2('%.1f'%o['mx']))]
t_tcar=table('Estatística descritiva do desempenho no Teste de Carminatti (T-CAR), covariável fisiológica.',
    ['Parâmetro','n','Média','DP','Mín–Máx'],[tcrow(o) for o in TCD],
    note='PV = pico de velocidade; FC = frequência cardíaca; TRIMP = training impulse. T-CAR1 (linha de base) e T-CAR2 (final).',fs=8.5)
H('2.4 Procedimentos e análise de dados',12,before=6)
P('Registraram-se ainda a sonolência (Epworth; JOHNS, 1991), o estresse percebido (PSS; COHEN; KAMARCK; MERMELSTEIN, '
 '1983) e as sessões de HIIT (Dias 2, 4 e 7). A normalidade foi verificada por Shapiro-Wilk; como todas as dimensões '
 'violaram a normalidade (p < 0,001), com efeito de piso nas negativas, os dados não foram transformados (NEVILL; LANE, '
 '2007) e adotaram-se procedimentos não paramétricos. Reportou-se descritiva, confiabilidade (alfa de Cronbach), '
 'coeficiente de variação e correlação de Spearman. As comparações intra-grupo (Dia 1 vs. Dia 7; pré vs. pós) usaram '
 'MANOVA de medidas repetidas em escores T (Wilks λ, F, η²ₚ) com testes univariados e magnitude por d de Cohen (0,20; '
 '0,50; 0,80) e η²ₚ (0,01; 0,06; 0,14). No sujeito, calculou-se o RCI (JACOBSON; TRUAX, 1991). A relação do pico de '
 'velocidade com o humor foi analisada por regressão (β, R², ρ), tercis (Kruskal-Wallis) e limiar de Youden (AUC). '
 'Análises complementares (pós-teste do modelo misto, ajuste alométrico, derivadas e decomposição sinal–ruído) constam '
 'no Material Suplementar. Adotou-se α = 0,05 (Python: pandas, SciPy, statsmodels).')

# ===================== 3 RESULTADOS =====================
H('3 RESULTADOS')
H('3.1 Análise descritiva exploratória',12,before=6)
P('Todas as dimensões violaram a normalidade (Shapiro-Wilk p < 0,001), com as negativas no piso; a consistência das '
 'medidas repetidas (ICC) foi moderada a boa (ICC(2,k) ≥ 0,81). A descritiva, a confiabilidade e as intercorrelações '
 'constam na Tabela 4: o vigor e a fadiga concentram a variabilidade informativa, enquanto as negativas têm medianas '
 'baixas; a fadiga associa-se à depressão e à raiva, e o vigor mantém-se relativamente independente do polo negativo '
 '(distribuições e diagramas no Material Suplementar).')
def cval(a,b):
    key='%s_%s'%(a,b); v=MScorr[key]; stx='*' if v['p']<0.001 else ''; return c2('%+.2f'%v['r'])+stx
def mrow(i,d):
    cells=[str(i+1)+' '+d['lab'],c2('%.2f'%d['M']),c2('%.2f'%d['SD']),'%d–%d'%(d['mn'],d['mx']),'%s–%s'%(c2('%.0f'%d['tmin']),c2('%.0f'%d['tmax'])),c2('%.2f'%d['alpha'])]
    for k in range(5):
        j=k+1; cells.append('—' if j<=i else cval(MSord[i],MSord[j]))
    return cells
t_desc=table('Descritivas, confiabilidade (α) e intercorrelações (ρ) das seis dimensões do BRUMS (%d observações).'%sm['n_obs'],
    ['Dimensão','M','DP','Amplitude','Escore T','α','2','3','4','5','6'],[mrow(i,d) for i,d in enumerate(MSd)],
    note='* p < 0,001 (Spearman). Escore T referenciado à amostra (M = 50; DP = 10).',fs=8)
H('3.2 Comparações intra-grupo e sensibilidade',12,before=6)
f_prof=figure(f'{FG}/xb5_profile_d1d7.png','Perfil de humor em escores T (M = 50; DP = 10) no Dia 1 e no Dia 7.',w=13.5)
P('No Dia 1 o grupo apresenta assinatura iceberg (vigor T = %s; fadiga T = %s), invertendo-se no Dia 7 (Figura %d). A '
 'MANOVA foi multivariadamente significativa (Wilks λ = %s; F(%d,%d) = %s; p = %s; η²ₚ = %s), com efeito no vigor '
 '(η²ₚ = %s; d = %s) e na fadiga (η²ₚ = %s; d = %s) (Tabela %d). O vigor e a fadiga foram as dimensões mais sensíveis à '
 'variação semanal (|dz| = %s e %s), com diferença entre os sete dias (Friedman p < 0,05); a resposta aguda pré → pós '
 'reproduziu o padrão (Wilks λ = %s; p = %s; detalhes no Material Suplementar). A trajetória semanal está na Figura %d.'%(
   c2('%.1f'%mvv('d1d7','Vigor','m1')),c2('%.1f'%mvv('d1d7','Fadiga','m1')),f_prof,
   c2('%.3f'%MV['d1d7']['wilks']),MV['d1d7']['df1'],MV['d1d7']['df2'],c2('%.2f'%MV['d1d7']['Fmv']),pstr(MV['d1d7']['p_mv']),c2('%.2f'%MV['d1d7']['eta_mv']),
   c2('%.2f'%mvv('d1d7','Vigor','eta')),c2('%+.2f'%mvv('d1d7','Vigor','d')),c2('%.2f'%mvv('d1d7','Fadiga','eta')),c2('%+.2f'%mvv('d1d7','Fadiga','d')),_TN[0]+1,
   c2('%.2f'%SENS['Vigor']['absdz']),c2('%.2f'%SENS['Fadiga']['absdz']),c2('%.3f'%MV['prepos']['wilks']),pstr(MV['prepos']['p_mv']),_FN[0]+1))
def mvrow(x):
    sig='*' if x['p']<0.05 else ''
    return [x['lab'],c2('%.1f'%x['m1']),c2('%.1f'%x['s1']),c2('%.1f'%x['m2']),c2('%.1f'%x['s2']),c2('%.2f'%x['F']),pstr(x['p'])+sig,c2('%+.2f'%x['d']),c2('%.3f'%x['eta'])]
mv=MV['d1d7']
t_mv=table('Comparação Dia 1 → Dia 7 das seis dimensões em escores T (MANOVA de medidas repetidas; n = %d).'%mv['n'],
    ['Dimensão','D1 M','D1 DP','D7 M','D7 DP','F','p','d','η²ₚ'],[mvrow(x) for x in mv['rows']],
    note='Wilks λ = %s; F(%d,%d) = %s; p = %s; η²ₚ = %s. * p < 0,05.'%(c2('%.3f'%mv['wilks']),mv['df1'],mv['df2'],c2('%.2f'%mv['Fmv']),pstr(mv['p_mv']),c2('%.2f'%mv['eta_mv']),),fs=8.5)
figure(f'{FG}/xb4_allz.png','Comportamento conjunto das sete variáveis ao longo da semana (escores padronizados z).')
H('3.3 Correlação nos dias de maior vigor e de maior fadiga',12,before=6)
FP=STAT['focusp']
P('O acoplamento do afeto negativo à fadiga depende da carga (Tabela %d): no dia de maior vigor (Dia 1) a depressão é '
 'independente da fadiga (ρ = %s; p = %s), mas no dia de maior fadiga (Dia 7) torna-se forte (ρ = %s; p = %s), o mesmo '
 'ocorrendo com a raiva — confirmando a H3.'%(_TN[0]+1,
   c2('%+.2f'%FP['D1']['Depressao']['rho_fad']),pstr(FP['D1']['Depressao']['p_fad']),c2('%+.2f'%FP['D7']['Depressao']['rho_fad']),pstr(FP['D7']['Depressao']['p_fad'])))
def frow(neg,lab):
    d1=FP['D1'][neg]; d7=FP['D7'][neg]
    return [lab,'%s (%s)'%(c2('%+.2f'%d1['rho_fad']),pstr(d1['p_fad'])),'%s (%s)'%(c2('%+.2f'%d7['rho_fad']),pstr(d7['p_fad']))]
t_focus=table('Correlação (ρ; p) das dimensões negativas com a fadiga no dia de maior vigor (D1) e no de maior fadiga (D7).',
    ['Dimensão','ρ × Fadiga D1 (p)','ρ × Fadiga D7 (p)'],
    [frow('Depressao','Depressão'),frow('Raiva','Raiva'),frow('Confusao','Confusão'),frow('Tensao','Tensão')],fs=9)
H('3.4 Prevalência dos perfis de humor',12,before=6)
P('A distribuição dos perfis mudou entre os dias extremos (χ² = %s; gl = %d; p = %s; Figura %d; Tabela %d): o iceberg caiu '
 'de %s%% para %s%% e a barbatana de tubarão subiu para %s%%, sem instalação relevante de perfis de risco extremo.'%(
   c2('%.2f'%PREV['chi']),PREV['dof'],pstr(PREV['p']),_FN[0]+1,_TN[0]+1,
   c2('%.0f'%(100*PREV['D1']['Iceberg']/PREV['n_d1'])),c2('%.0f'%(100*PREV['D7']['Iceberg']/PREV['n_d7'])),c2('%.0f'%(100*PREV['D7']['Barbatana tubarão']/PREV['n_d7']))))
figure(f'{FG}/xb5_prev.png','Prevalência (%) dos seis perfis de humor no Dia 1 e no Dia 7.',w=13.5)
PROFR=[('Iceberg','Iceberg'),('Everest invertido','Everest invertido'),('Iceberg invertido','Iceberg invertido'),('Submerso','Submerso'),('Barbatana tubarão','Barbatana de tubarão'),('Superfície','Superfície')]
def prow(p,lab):
    d1=PREV['D1'][p]; d7=PREV['D7'][p]; ov=PREV['overall'][p]; nov=sum(PREV['overall'].values())
    return [lab,'%d (%s)'%(d1,c2('%.1f'%(100*d1/PREV['n_d1']))),'%d (%s)'%(ov,c2('%.1f'%(100*ov/nov))),'%d (%s)'%(d7,c2('%.1f'%(100*d7/PREV['n_d7'])))]
t_prev=table('Prevalência dos perfis de humor por recorte: n (%%) no Dia 1, na semana e no Dia 7.',
    ['Perfil','Dia 1 n (%)','Semana n (%)','Dia 7 n (%)'],[prow(p,lab) for p,lab in PROFR],
    note='Qui-quadrado Dia 1 × Dia 7: χ² = %s; gl = %d; p = %s.'%(c2('%.2f'%PREV['chi']),PREV['dof'],pstr(PREV['p'])),fs=9)
H('3.5 Análise no nível do atleta (mudança confiável)',12,before=6)
P('A mudança confiável (RCI, Dia 1 → Dia 7) foi desigual entre as dimensões (Tabela %d; Figura %d): a fadiga é a única em '
 'que a maioria deteriora de forma confiável (%d de %d; %d%%), seguida do vigor (%d%%); as negativas não pioram. O iceberg '
 'era o perfil de %d dos %d atletas na linha de base e a barbatana de tubarão passou a ser o de %d no último dia '
 '(trajetórias e transições individuais no Material Suplementar).'%(_TN[0]+1,_FN[0]+1,
   RC['Fadiga']['piora'],RC['Fadiga']['n'],RC['Fadiga']['pct_piora'],RC['Vigor']['pct_piora'],PCNT['iceberg_D1'],PCNT['n'],PCNT['shark_D7']))
def rcrow(k,lab):
    v=RC[k]; return [lab,c2('%.2f'%v['alpha']),v['n'],'%d (%d%%)'%(v['piora'],v['pct_piora']),'%d (%d%%)'%(v['estavel'],v['pct_est']),'%d (%d%%)'%(v['melhora'],v['pct_mel'])]
t_rci=table('Mudança confiável individual (RCI, Dia 1 → Dia 7) por dimensão (n = %d).'%RC['Vigor']['n'],
    ['Dimensão','α','n','Piora confiável','Estável','Melhora confiável'],
    [rcrow(k,lab) for k,lab in [('Vigor','Vigor'),('Fadiga','Fadiga'),('Depressao','Depressão'),('Raiva','Raiva'),('Tensao','Tensão'),('Confusao','Confusão')]],fs=9)
figure(f'{FG}/xb2_rci6.png','Proporção de atletas com mudança confiável (RCI) por dimensão, do Dia 1 ao Dia 7.',w=13.0)
H('3.6 Aptidão intermitente, sono e carga de HIIT',12,before=6)
def preg(col): return PV['pv']['wk_'+col]['TCAR1']
P('Atletas com maior pico de velocidade (T-CAR1) reportaram mais vigor (ρ = %s; p = %s) e menos fadiga física (ρ = %s; '
 'R² = %s; p = %s) ao longo da semana (Tabela %d; Figura %d); a fadiga diferiu entre tercis de aptidão (Kruskal-Wallis '
 'p = %s) e um limiar de PV ≈ %s km/h discriminou os dias de maior fadiga (AUC = %s). A sonolência (%d de %d atletas com '
 'Epworth > 10) associou-se à fadiga (ρ = %s; p = %s), enquanto o estresse percebido não (p > 0,05). Os dias de HIIT '
 'apresentaram pior humor (vigor dz = %s; fadiga dz = %s; p < 0,05), embora a resposta aguda pré → pós não fosse maior '
 'nesses dias (Material Suplementar).'%(
   c2('%+.2f'%preg('Vigor')['rho']),pstr(preg('Vigor')['rho_p']),c2('%+.2f'%preg('FadFisica')['rho']),c2('%.2f'%preg('FadFisica')['r2']),pstr(preg('FadFisica')['rho_p']),
   _TN[0]+1,_FN[0]+1,pstr(PV['terc_kruskal_fadfis']['p']),c2('%.1f'%LP['thr']),c2('%.2f'%LP['auc']),
   ep['hi'],ep['n'],c2('%+.2f'%ep['corr']['Fadiga']['rho']),pstr(ep['corr']['Fadiga']['p']),
   c2('%+.2f'%hi['Vigor']['dz']),c2('%+.2f'%hi['Fadiga']['dz'])))
def pregrow(col,lab):
    v=preg(col); return [lab,c2('%+.2f'%v['slope']),c2('%.2f'%v['r2']),c2('%+.2f'%v['rho']),pstr(v['rho_p'])]
t_pv=table('Regressão do pico de velocidade do T-CAR1 sobre as médias semanais de humor e fadiga (n = %d atletas).'%preg('Vigor')['n'],
    ['Desfecho semanal','β','R²','ρ','p'],
    [pregrow('Vigor','Vigor (BRUMS)'),pregrow('Fadiga','Fadiga (BRUMS)'),pregrow('TMD','PTH'),pregrow('FadFisica','Fadiga física')],fs=9)
figure(f'{FG}/pv7_scatter.png','Dispersão e reta de regressão entre o pico de velocidade do T-CAR1 e as médias semanais de vigor, fadiga (BRUMS), PTH e fadiga física (faixa = IC95%).',w=15.0)

# ===================== 4 DISCUSSÃO =====================
H('4 DISCUSSÃO')
P('A deterioração do humor concentrou-se no eixo energia–fadiga (H1), com efeito grande no vigor e médio na fadiga, '
 'enquanto as negativas permaneceram próximas do piso — coerente com o vigor e a fadiga como dimensões mais sensíveis à '
 'carga (SAW; MAIN; GASTIN, 2016; THORPE et al., 2017). A migração do iceberg para o perfil de fadiga (H2) reproduz o '
 '“derretimento do iceberg” e alinha-se aos perfis prototípicos validados (MORGAN, 1985; HAN; PARSONS-SMITH; TERRY, '
 '2020). O afeto negativo acoplou-se à fadiga apenas sob carga acumulada (H3), e a resposta foi heterogênea entre atletas '
 '(H4), com a fadiga sendo a única dimensão a piorar de forma confiável na maioria — o que sustenta o monitoramento '
 'individualizado (KELLMANN et al., 2018).')
P('A aptidão aeróbia intermitente modulou a resposta: atletas mais aptos reportaram mais vigor e menos fadiga, coerente '
 'com o papel da capacidade intermitente na tolerância ao esforço do handebol (KARCHER; BUCHHEIT, 2014; FERNANDES-DA-'
 'SILVA et al., 2016); o limiar de pico de velocidade oferece uma referência prática. A sonolência acompanhou a fadiga, '
 'ao passo que o estresse percebido não, e os dias de HIIT concentraram pior humor por acúmulo no microciclo. Como '
 'limitações, destacam-se o tamanho amostral (n = %d), o efeito de piso das negativas e o desenho observacional de fase '
 'única. Análises complementares no Material Suplementar (pós-teste, alométrico, sinal–ruído) reforçam a robustez dos '
 'achados.'%sm['n'])

# ===================== 5 CONCLUSÕES =====================
H('5 CONCLUSÕES')
P('Na última semana de pré-temporada, o perfil de humor de handebolistas de elite migrou da prontidão (iceberg) para a '
 'fadiga funcional, com a deterioração concentrada no eixo energia–fadiga — as dimensões mais sensíveis —, tanto no grupo '
 'quanto no sujeito. A aptidão aeróbia intermitente e o sono modularam a resposta. Recomenda-se monitorar o eixo '
 'energia–fadiga por tendência, referenciado à linha de base individual e à aptidão física de cada atleta.')

# ===================== REFERÊNCIAS =====================
H('REFERÊNCIAS')
refs=[
 'COHEN, S.; KAMARCK, T.; MERMELSTEIN, R. A global measure of perceived stress. Journal of Health and Social Behavior, v. 24, n. 4, p. 385–396, 1983.',
 'FERNANDES-DA-SILVA, J. et al. The peak velocity derived from the Carminatti Test is related to physical match performance in young soccer players. Journal of Sports Sciences, v. 34, n. 24, p. 2238–2245, 2016.',
 'HAN, C.; PARSONS-SMITH, R. L.; TERRY, P. C. Mood profiling in Singapore: cross-cultural validation and potential applications of mood profile clusters. Frontiers in Psychology, v. 11, 665, 2020.',
 'JACOBSON, N. S.; TRUAX, P. Clinical significance: a statistical approach to defining meaningful change. Journal of Consulting and Clinical Psychology, v. 59, n. 1, p. 12–19, 1991.',
 'JOHNS, M. W. A new method for measuring daytime sleepiness: the Epworth Sleepiness Scale. Sleep, v. 14, n. 6, p. 540–545, 1991.',
 'KARCHER, C.; BUCHHEIT, M. On-court demands of elite handball, with special reference to playing positions. Sports Medicine, v. 44, n. 6, p. 797–814, 2014.',
 'KELLMANN, M. et al. Recovery and performance in sport: consensus statement. International Journal of Sports Physiology and Performance, v. 13, n. 2, p. 240–245, 2018.',
 'MICHALSIK, L. B.; MADSEN, K.; AAGAARD, P. Match performance and physiological capacity of female elite team handball players. International Journal of Sports Medicine, v. 35, n. 7, p. 595–607, 2013.',
 'MORGAN, W. P. Selected psychological factors limiting performance: a mental health model. In: CLARKE, D. H.; ECKERT, H. M. (Ed.). Limits of human performance. Champaign: Human Kinetics, 1985. p. 70–80.',
 'NEVILL, A. M.; LANE, A. M. Why self-report “Likert” scale data should not be log-transformed. Journal of Sports Sciences, v. 25, n. 1, p. 1–2, 2007.',
 'PARSONS-SMITH, R. L.; TERRY, P. C.; MACHIN, M. A. Identification and description of novel mood profile clusters. Frontiers in Psychology, v. 8, 1958, 2017.',
 'ROHLFS, I. C. P. M. et al. A Escala de Humor de Brunel (Brums). Revista Brasileira de Medicina do Esporte, v. 14, n. 3, p. 176–181, 2008.',
 'ROHLFS, I. C. P. M. et al. Psychometric characteristics of the Brazil Mood Scale among youth and elite athletes. Sports, v. 11, n. 12, 244, 2023.',
 'SAW, A. E.; MAIN, L. C.; GASTIN, P. B. Monitoring the athlete training response: subjective self-reported measures trump commonly used objective measures. British Journal of Sports Medicine, v. 50, n. 5, p. 281–291, 2016.',
 'TERRY, P. C.; LANE, A. M.; FOGARTY, G. J. Construct validity of the Profile of Mood States — Adolescents for use with adults. Psychology of Sport and Exercise, v. 4, n. 2, p. 125–139, 2003.',
 'THORPE, R. T. et al. Monitoring fatigue status in elite team-sport athletes. International Journal of Sports Physiology and Performance, v. 12, s2, p. S2-27–S2-34, 2017.']
for rf in refs:
    p=doc.add_paragraph(); r=p.add_run(rf); r.font.size=Pt(11); p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(6); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

OUTP='/home/user/mdlucca/Artigos/Artigo_Principal_Perfil_Humor.docx'
doc.save(OUTP); print('SAVED',OUTP,'| Tabelas',_TN[0],'Figuras',_FN[0])
