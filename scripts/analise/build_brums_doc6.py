# -*- coding: utf-8 -*-
import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
R=json.load(open('brums_desc2.json')); RC=json.load(open('rci6.json')); STAT=json.load(open('brums_stats3.json'))
S4=json.load(open('brums_stats4.json')); MS=json.load(open('model_stats.json')); MV=json.load(open('manova.json'))
PV=json.load(open('pv_stats.json')); LIM=json.load(open('tcar_limiar.json'))
PHJ=json.load(open('posthoc.json')); ALLO=json.load(open('allo.json')); ROC=json.load(open('roc.json'))
DEN=json.load(open('denoise.json')); DERIV=json.load(open('deriv_exp.json'))
CRS=json.load(open('cross.json')); EX=json.load(open('extra.json'))
FG='/home/user/mdlucca/Artigos/figuras'
def c2(s): return str(s).replace('.',',')
doc=Document()
st=doc.styles['Normal']; st.font.name='Times New Roman'; st.font.size=Pt(12)
st.element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')
st.paragraph_format.line_spacing=1.5; st.paragraph_format.space_after=Pt(0)
s=doc.sections[0]; s.top_margin=Cm(3); s.left_margin=Cm(3); s.bottom_margin=Cm(2); s.right_margin=Cm(2)
_TN=[0]; _FN=[0]
def P(t='',just=True,size=12,after=6,bold=False,ind=True):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(size); r.bold=bold
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(after)
    if just: p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if just and ind: p.paragraph_format.first_line_indent=Cm(1.25)
    return p
def RUN(pairs,after=6,ind=True):
    p=doc.add_paragraph(); p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(after)
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if ind: p.paragraph_format.first_line_indent=Cm(1.25)
    for txt,bold in pairs:
        r=p.add_run(txt); r.bold=bold; r.font.size=Pt(12)
    return p
def H(t,size=12,before=10):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(size)
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.5
def _bd(c,span=False):
    tcPr=c._tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for e_ in ['top','bottom','left','right']:
        el=OxmlElement(f'w:{e_}'); on=e_ in ('top','bottom')
        el.set(qn('w:val'),'single' if on else 'nil'); el.set(qn('w:sz'),'6'); el.set(qn('w:color'),'000000'); b.append(el)
    tcPr.append(b)
def table(cap,header,rows,fonte='Fonte: dados da pesquisa (2026).',fs=9,note=None):
    _TN[0]+=1
    p=doc.add_paragraph(); r=p.add_run('Tabela %d – %s'%(_TN[0],cap)); r.font.size=Pt(11)
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
    t=doc.add_table(rows=1,cols=len(header)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,htx in enumerate(header):
        cc=t.rows[0].cells[i]; cc.text=''; rr=cc.paragraphs[0].add_run(htx); rr.bold=True; rr.font.size=Pt(fs); rr.font.name='Times New Roman'
        cc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; _bd(cc)
    for row in rows:
        cs=t.add_row().cells
        for i,val in enumerate(row):
            cs[i].text=''; rr=cs[i].paragraphs[0].add_run(str(val)); rr.font.size=Pt(fs); rr.font.name='Times New Roman'
            cs[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if i>0 else WD_ALIGN_PARAGRAPH.LEFT; _bd(cs[i])
    if note:
        pn=doc.add_paragraph(); rn=pn.add_run(note); rn.font.size=Pt(8.5); rn.italic=True; pn.paragraph_format.space_after=Pt(0)
    pf=doc.add_paragraph(); rf=pf.add_run(fonte); rf.font.size=Pt(9); pf.paragraph_format.space_after=Pt(6)
    return _TN[0]
def figure(path,cap,w=15.5):
    _FN[0]+=1
    pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; pp.add_run().add_picture(path,width=Cm(w)); pp.paragraph_format.space_before=Pt(6)
    pc=doc.add_paragraph(); pc.alignment=WD_ALIGN_PARAGRAPH.CENTER; rc=pc.add_run('Figura %d – %s'%(_FN[0],cap)); rc.font.size=Pt(11)
    pf=doc.add_paragraph(); pf.alignment=WD_ALIGN_PARAGRAPH.CENTER; rf=pf.add_run('Fonte: elaboração dos autores (2026).'); rf.font.size=Pt(9); pf.paragraph_format.space_after=Pt(6)
    return _FN[0]
def pstr(p): return '< 0,001' if p<0.001 else c2('%.3f'%p)

pr=R['prepos']; d17=R['d1d7']; prof=R['profiles']; sm=R['sample']; foc=R['focus']; npk=R['neg_peak']; desc=R['desc']
CV=S4['cv']; FR=S4['friedman']; SENS=S4['sens']; PCT=S4['pct']; PT=S4['prof_trans']; PCNT=S4['prof_counts']
MSd=MS['desc']; MScorr=MS['corr']; MSord=MS['order']; PREV=MS['prev']
ORDER=[('Vigor','Vigor'),('Fadiga','Fadiga'),('TMD','Perturbação Total do Humor (PTH)'),('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Confusao','Confusão')]

# ===================== TÍTULO + RESUMO =====================
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('PERFIL DE HUMOR DE ATLETAS DE HANDEBOL DE ELITE NA ÚLTIMA SEMANA DE PRÉ-TEMPORADA: MONITORAMENTO DAS SEIS DIMENSÕES DO BRUMS DO GRUPO AO ATLETA')
r.bold=True; r.font.size=Pt(13); p.paragraph_format.space_after=Pt(10)
H('RESUMO',before=2)
RUN([('O monitoramento do humor é um indicador prático do bem-estar e da prontidão do atleta. ',False),
 ('O objetivo ',True),('deste estudo foi avaliar o perfil de humor de atletas de handebol de elite na última semana de '
 'pré-temporada, caracterizando e comparando as seis dimensões do BRUMS, do nível do grupo ao do atleta. Vinte e sete '
 'atletas do sexo masculino responderam ao BRUMS-24 duas vezes ao dia (pré e pós-treino) durante sete dias, totalizando '
 '%d observações. Os escores foram convertidos em escores T (M = 50; DP = 10) e as comparações intra-grupo (Dia 1 vs. Dia '
 '7; pré vs. pós) foram testadas por MANOVA de medidas repetidas, complementada por estatística não paramétrica (Shapiro-'
 'Wilk, Friedman, Wilcoxon, ρ de Spearman), coeficiente de correlação intraclasse (ICC), Índice de Mudança Confiável (RCI) '
 'individual e classificação nos seis perfis de humor. '%sm['n_obs'],False),
 ('Resultados: ',True),('a comparação Dia 1 → Dia 7 foi multivariadamente significativa (Wilks λ = %s; F(%d,%d) = %s; p = '
 '%s; η²ₚ = %s), com deterioração concentrada no eixo energia–fadiga — vigor (F = %s; η²ₚ = %s; d = %s) e fadiga (F = %s; '
 'η²ₚ = %s; d = %s). A resposta aguda pré → pós também foi significativa (Wilks λ = %s; p = %s). O afeto negativo '
 'acoplou-se à fadiga apenas sob carga acumulada (depressão × fadiga: ρ = %s no D1 → ρ = %s no D7). A prevalência de '
 'perfis mudou do Dia 1 ao Dia 7 (χ² = %s; p = %s): o perfil iceberg caiu e o de fadiga (“barbatana de tubarão”) subiu. No '
 'nível individual, a fadiga foi a única dimensão a piorar de forma confiável na maioria dos atletas (%d%%). '%(
   c2('%.3f'%MV['d1d7']['wilks']),MV['d1d7']['df1'],MV['d1d7']['df2'],c2('%.2f'%MV['d1d7']['Fmv']),pstr(MV['d1d7']['p_mv']),c2('%.2f'%MV['d1d7']['eta_mv']),
   c2('%.2f'%next(x['F'] for x in MV['d1d7']['rows'] if x['k']=='Vigor')),c2('%.2f'%next(x['eta'] for x in MV['d1d7']['rows'] if x['k']=='Vigor')),c2('%+.2f'%next(x['d'] for x in MV['d1d7']['rows'] if x['k']=='Vigor')),
   c2('%.2f'%next(x['F'] for x in MV['d1d7']['rows'] if x['k']=='Fadiga')),c2('%.2f'%next(x['eta'] for x in MV['d1d7']['rows'] if x['k']=='Fadiga')),c2('%+.2f'%next(x['d'] for x in MV['d1d7']['rows'] if x['k']=='Fadiga')),
   c2('%.3f'%MV['prepos']['wilks']),pstr(MV['prepos']['p_mv']),
   c2('%+.2f'%foc['D1']['Depressao']['rho_fad']),c2('%+.2f'%foc['D7']['Depressao']['rho_fad']),
   c2('%.2f'%PREV['chi']),pstr(PREV['p']),RC['Fadiga']['pct_piora']),False),
 ('Conclusão: ',True),('na última semana de pré-temporada, o perfil de humor migrou da prontidão (iceberg) para a fadiga '
 'funcional; o vigor e a fadiga são as dimensões mais sensíveis, o que fundamenta um monitoramento individualizado e '
 'referenciado à linha de base.',False)],after=6)
P('Palavras-chave: humor; BRUMS; handebol; monitoramento do atleta; fadiga; bem-estar psicológico.',size=11,after=8,ind=False)

# ===================== 1 INTRODUÇÃO =====================
H('1 INTRODUÇÃO')
H('1.1 Monitoramento do humor no esporte de elite',12,before=6)
P('As organizações esportivas de alto rendimento demonstram, cada vez mais, o compromisso de proteger o bem-estar '
 'psicológico dos atletas ao mesmo tempo em que buscam maximizar o desempenho, sendo comum o uso de algum indicador '
 'psicológico para rastrear o bem-estar e o risco de problemas de saúde mental (ROHLFS et al., 2023). Avaliações regulares '
 'do humor têm demonstrado, ao longo de décadas, utilidade preditiva tanto para o bem-estar quanto para o desempenho '
 'esportivo (SAW; MAIN; GASTIN, 2016; LOCHBAUM et al., 2021). O humor é definido como um conjunto de sentimentos, de '
 'natureza efêmera, que variam em intensidade e duração e geralmente envolvem mais de uma emoção; diferentemente das '
 'emoções, os estados de humor são menos intensos, mais duradouros e nem sempre possuem um gatilho específico, variando '
 'em valência (positiva–negativa) e ativação (DE MIRANDA ROHLFS et al., 2025).')
P('O perfilamento do humor — processo em que os escores do indivíduo são comparados a dados normativos para gerar um '
 'perfil gráfico — é amplamente utilizado para identificar padrões de resposta e desvios em relação a um humor típico, '
 'funcionando como triagem de bem-estar e de risco psicológico (DE MIRANDA ROHLFS et al., 2025). O instrumento mais '
 'utilizado para essa finalidade é a Brunel Mood Scale (BRUMS), derivada abreviada do Profile of Mood States (POMS), que '
 'mensura seis dimensões — tensão, depressão, raiva, vigor, fadiga e confusão — em 24 itens de rápida aplicação e '
 'reaplicação, com sólidas propriedades psicométricas e sucessivas validações transculturais, incluindo a versão '
 'brasileira (ROHLFS et al., 2008; TERRY; LANE; FOGARTY, 2003; ROHLFS et al., 2023; LEW et al., 2023).')
H('1.2 Handebol: modalidade coletiva intermitente e de alta intensidade',12,before=6)
P('O handebol de quadra é uma modalidade coletiva de caráter marcadamente intermitente e de alta intensidade. Ao longo da '
 'partida, esforços máximos e explosivos — sprints curtos, saltos, arremessos, mudanças de direção e contatos físicos — '
 'alternam-se, de forma imprevisível, com períodos de recuperação incompleta, exigindo simultaneamente potência '
 'anaeróbia, capacidade aeróbia intermitente e elevada tolerância à fadiga (KARCHER; BUCHHEIT, 2014; WAGNER et al., 2014). '
 'As demandas variam com a posição tática e a fadiga já foi documentada pela queda do volume de corrida em alta '
 'intensidade ao longo do jogo (MICHALSIK; MADSEN; AAGAARD, 2013). Esse perfil de esforço eleva a carga interna nos '
 'microciclos de acúmulo e repercute no estado afetivo: em handebolistas de elite, o humor e o estresse associam-se a '
 'indicadores de sobrecarga fisiológica e psicossomática (RATZ-SULYOK et al., 2026). O caráter intermitente e de contato '
 'do handebol torna o acompanhamento subjetivo do humor particularmente informativo para a individualização da carga.')
P('A tolerância a esse padrão de esforço depende diretamente das capacidades físicas do atleta, em especial da aptidão '
 'aeróbia intermitente, que sustenta a reposição energética e a recuperação entre os esforços de alta intensidade e, '
 'assim, atenua a instalação da fadiga ao longo do jogo e do microciclo (KARCHER; BUCHHEIT, 2014; MICHALSIK; AAGAARD, '
 '2015). Um marcador de campo válido dessa capacidade é o pico de velocidade obtido no Teste de Carminatti (T-CAR), um '
 'teste intermitente progressivo cujo pico de velocidade se associa ao desempenho físico em partida e discrimina níveis '
 'de aptidão em modalidades intermitentes (FERNANDES-DA-SILVA et al., 2016). Integrar esse parâmetro fisiológico ao '
 'monitoramento do humor permite verificar se a aptidão intermitente modula (e, portanto, ajuda a normalizar '
 'interindividualmente) a resposta afetiva de vigor e fadiga à carga do microciclo — hipótese ainda pouco explorada em '
 'handebolistas de elite.')
H('1.3 Perfis de humor e a lacuna do monitoramento em microciclos',12,before=6)
P('Para além dos escores isolados, Morgan (1985) descreveu o “perfil iceberg” — vigor elevado sobre dimensões negativas '
 'baixas — como assinatura de prontidão e saúde mental positiva. Abordagens recentes formalizaram seis perfis '
 'prototípicos de humor, replicados em grandes amostras internacionais: iceberg, Everest invertido, iceberg invertido, '
 'submerso, barbatana de tubarão e superfície (PARSONS-SMITH; TERRY; MACHIN, 2017; HAN; PARSONS-SMITH; TERRY, 2020; TERRY '
 'et al., 2021). Esses perfis vêm sendo aplicados ao rastreamento da prontidão e do risco em atletas, inclusive na '
 'associação entre perfis negativos e a ocorrência de lesão (DE MIRANDA ROHLFS et al., 2025). Apesar da ampla adoção da '
 'BRUMS, faltam descrições detalhadas — dimensão a dimensão e do grupo ao sujeito — do comportamento do humor em um '
 'microciclo pré-competitivo de handebol de elite, período em que a carga que antecede a competição se concentra e no '
 'qual o afeto negativo pode ou não passar a acompanhar a fadiga.')
H('1.4 Objetivos e hipóteses',12,before=6)
P('O objetivo geral foi avaliar o perfil de humor de atletas de handebol de elite na última semana de pré-temporada, '
 'caracterizando e comparando o comportamento das seis dimensões do BRUMS e da Perturbação Total do Humor (PTH), do nível '
 'do grupo ao nível do sujeito. Como objetivos específicos, buscou-se: (i) caracterizar exploratoriamente cada dimensão '
 '(tendência central, dispersão, percentis, confiabilidade e coeficiente de variação); (ii) comparar, no grupo, a '
 'resposta Dia 1 → Dia 7 e pré → pós em escores T (MANOVA), identificando as dimensões mais sensíveis; (iii) analisar a '
 'estrutura de correlação entre as dimensões, com ênfase nos dias de maior vigor e de maior fadiga; (iv) descrever a '
 'prevalência e a transição dos perfis de humor; e (v) quantificar, no sujeito, a mudança confiável (RCI).')
P('Formularam-se quatro hipóteses: (H1) a deterioração concentra-se no eixo energia–fadiga (vigor ↓, fadiga ↑), com '
 'magnitude média a grande, enquanto as dimensões negativas de valência não fadiga permanecem estáveis por efeito de '
 'piso; (H2) o perfil de humor migra do iceberg para o perfil de fadiga (barbatana de tubarão); (H3) o afeto negativo '
 'acopla-se à fadiga somente sob carga acumulada (correlação maior no dia de maior fadiga do que no de maior vigor); e '
 '(H4) a resposta é heterogênea entre atletas, exigindo interpretação individualizada.')

# ===================== 2 MATERIAIS E MÉTODOS =====================
H('2 MATERIAIS E MÉTODOS')
H('2.1 Participantes',12,before=6)
P('Participaram %d atletas de handebol do sexo masculino de nível competitivo, avaliados em condições ecológicas de '
 'treinamento durante o microciclo pré-competitivo (21–27 de abril de 2024). A distribuição da amostra por posição de '
 'jogo é apresentada na Tabela 1 e a caracterização antropométrica na Tabela 2. Todos os procedimentos seguiram os '
 'princípios éticos da Declaração de Helsinque, com consentimento informado dos participantes.'%sm['n'])
t_dist=table('Distribuição demográfica e situacional da amostra (n = %d).'%sm['n'],
    ['Fonte','Grupo','n','%'],
    [['Sexo','Masculino',sm['n'],'100,0']]+
    [['Posição' if i==0 else '',k,v,c2('%.1f'%(100*v/sm['n']))] for i,(k,v) in enumerate(sm['pos'].items())]+
    [['Total','Todos',sm['n'],'100,0']],fs=9)
def srow(lab,k):
    v=sm[k]; return [lab,c2('%.1f'%v['mean']),c2('%.1f'%v['sd']),'%s–%s'%(c2('%.1f'%v['lo']),c2('%.1f'%v['hi'])),'%s–%s'%(c2('%.1f'%v['mn']),c2('%.1f'%v['mx']))]
t_anth=table('Caracterização antropométrica e de experiência da amostra (n = %d).'%sm['n'],
    ['Variável','Média','DP','IC95%','Mín–Máx'],
    [srow('Idade (anos)','idade'),srow('Estatura (cm)','estatura'),srow('Massa corporal (kg)','massa'),srow('Percentual de gordura','pG'),srow('Experiência na modalidade (anos)','exp')])
H('2.2 Medida do humor',12,before=6)
P('O humor foi avaliado pela BRUMS-24, composta por 24 itens respondidos em escala Likert de 0 (“nada”) a 4 '
 '(“extremamente”), agrupados em seis subescalas de 0 a 16 pontos (tensão, depressão, raiva, vigor, fadiga e confusão); '
 'escores mais altos indicam maior nível da dimensão. A Perturbação Total do Humor resume o perfil: PTH = tensão + '
 'depressão + raiva + fadiga + confusão − vigor. A consistência interna, estimada pelo alfa de Cronbach, variou de %s '
 '(tensão) a %s (raiva) nesta amostra (Tabela 4). Para a interpretação normativa e a classificação de perfis, os escores '
 'brutos foram convertidos em escores T (M = 50; DP = 10).'%(
   c2('%.2f'%RC['Tensao']['alpha']),c2('%.2f'%RC['Raiva']['alpha'])))
H('2.3 Medida da aptidão intermitente (T-CAR)',12,before=6)
P('A aptidão aeróbia intermitente foi avaliada pelo Teste de Carminatti (T-CAR), teste de campo progressivo e '
 'intermitente composto por repetições de 12 s de corrida em vaivém (ida e volta) intercaladas por 6 s de recuperação, '
 'com incrementos sucessivos de velocidade guiados por sinal sonoro até a exaustão voluntária. O desfecho analisado foi o '
 'pico de velocidade (PV), tomado como marcador da capacidade intermitente e do limiar de esforço (FERNANDES-DA-SILVA et '
 'al., 2016). O teste foi aplicado no início da pré-temporada (T-CAR1; PV = %s ± %s km/h) e reaplicado ao final '
 '(T-CAR2; PV = %s ± %s km/h), permitindo caracterizar a aptidão inicial de cada atleta e sua variação (ΔPV = %s ± %s '
 'km/h). O PV inicial (T-CAR1) foi utilizado como parâmetro fisiológico de referência para normalizar e interpretar '
 'interindividualmente a resposta de humor–fadiga ao microciclo.'%(
   c2('%.1f'%PV['desc']['pvini_m']),c2('%.1f'%PV['desc']['pvini_sd']),c2('%.1f'%PV['desc']['pv_m']),c2('%.1f'%PV['desc']['pv_sd']),
   c2('%.1f'%PV['desc']['dpv_m']),c2('%.1f'%PV['desc']['dpv_sd'])))
H('2.4 Procedimentos',12,before=6)
P('O BRUMS foi autoaplicado por formulário eletrônico com carimbo de data/hora, duas vezes por dia de treino — a primeira '
 'resposta tomada como pré e a última como pós —, ao longo de sete dias, totalizando %d observações válidas. Em conjunto, '
 'foram registradas a sonolência diurna (Escala de Sonolência de Epworth; JOHNS, 1991), o estresse percebido (Escala de '
 'Estresse Percebido — PSS; COHEN; KAMARCK; MERMELSTEIN, 1983) e a realização de sessão de treino intervalado de alta '
 'intensidade (HIIT) em cada dia (Dias 2, 4 e 7). A Figura 1 sintetiza o framework das coletas.'%sm['n_obs'])
figure(f'{FG}/xb2_framework.png','Framework das coletas: da amostra e do microciclo às observações do BRUMS, às subescalas e à classificação de perfis.',w=11.0)
H('2.5 Análise de dados',12,before=6)
P('A normalidade foi verificada pelo teste de Shapiro-Wilk e pela inspeção de histogramas, diagramas de caixa e gráficos '
 'de dispersão. Como todas as dimensões violaram a normalidade (Shapiro-Wilk p < 0,001), com forte assimetria positiva e '
 'efeito de piso nas dimensões negativas, os dados não foram transformados (NEVILL; LANE, 2007) e adotaram-se '
 'procedimentos não paramétricos como referência. Reportou-se estatística descritiva (média, desvio-padrão, amplitude, '
 'percentis e escore T), confiabilidade (alfa de Cronbach) e coeficiente de variação entre atletas e intraindividual. As '
 'relações entre dimensões foram quantificadas por correlação de Spearman (ρ). Para as comparações intra-grupo, os '
 'escores foram convertidos em escores T e submetidos à análise multivariada de variância (MANOVA) de medidas repetidas '
 '(Dia 1 vs. Dia 7; pré vs. pós), reportada por lambda de Wilks, F, p e eta-quadrado parcial (η²ₚ), com testes '
 'univariados de acompanhamento e teste de Wilcoxon como verificação não paramétrica. A variação ao longo dos sete dias '
 'foi testada por Friedman com W de Kendall; a consistência das medidas repetidas, pelo coeficiente de correlação '
 'intraclasse (ICC(2,1) e ICC(2,k)). A magnitude foi interpretada por d de Cohen (0,20 = pequeno; 0,50 = médio; 0,80 = '
 'grande) e por η²ₚ (0,01 = pequeno; 0,06 = médio; 0,14 = grande). No nível do sujeito, calculou-se o Índice de Mudança '
 'Confiável (RCI = (D7 − D1)/EP_dif; EP_dif = DP_D1·√[2(1 − α)]), com mudança confiável quando |RCI| ≥ 1,96 (JACOBSON; '
 'TRUAX, 1991). Cada observação foi classificada nos seis perfis de humor por proximidade aos protótipos em escores T, e '
 'a mudança de prevalência entre dias foi testada pelo qui-quadrado. A relação entre o pico de velocidade do T-CAR e as '
 'respostas de humor–fadiga (médias semanais de vigor, fadiga, PTH e fadiga física) foi examinada por regressão linear '
 '(coeficiente β, R² e ρ de Spearman) com reta de ajuste, banda de confiança de 95%% e gráfico de dispersão; os atletas '
 'foram ainda estratificados em tercis de aptidão (T-CAR1), comparados por Kruskal-Wallis, e determinou-se um limiar de '
 'PV que melhor discrimina os dias de maior fadiga (índice de Youden, com área sob a curva ROC e intervalo de confiança '
 'por reamostragem bootstrap). Adotou-se α = 0,05, com as análises conduzidas em Python (bibliotecas pandas, SciPy e '
 'statsmodels).')

# ===================== 3 RESULTADOS =====================
H('3 RESULTADOS')
H('3.1 Triagem de dados e análise descritiva exploratória',12,before=6)
P('O teste de Shapiro-Wilk rejeitou a normalidade em todas as dimensões (p < 0,001; Tabela 4), com as negativas '
 'concentradas no piso — padrão típico do humor, em que os escores negativos exibem muitos valores baixos e poucos '
 'elevados. Nenhuma transformação foi aplicada. A descritiva geral, a confiabilidade e as intercorrelações entre as seis '
 'dimensões constam na Tabela 3. O vigor (escore T de %s a %s) e a fadiga concentram a variabilidade informativa, '
 'enquanto as dimensões negativas apresentam medianas baixas e coeficientes de variação inflados pelo efeito de piso '
 '(Tabela 5). As dimensões negativas intercorrelacionam-se entre si e a fadiga associa-se à depressão e à raiva, ao passo '
 'que o vigor mantém-se relativamente independente do polo negativo.'%(
   c2('%.0f'%next(d['tmin'] for d in MSd if d['k']=='Vigor')),c2('%.0f'%next(d['tmax'] for d in MSd if d['k']=='Vigor'))))
# Rohlfs-style descriptives + reliability + intercorrelation
def cval(a,b):
    key='%s_%s'%(a,b); v=MScorr[key]; st='*' if v['p']<0.001 else ''; return c2('%+.2f'%v['r'])+st
def mrow(i,d):
    cells=[str(i+1)+' '+d['lab'],c2('%.2f'%d['M']),c2('%.2f'%d['SD']),'%d–%d'%(d['mn'],d['mx']),
        '%s–%s'%(c2('%.0f'%d['tmin']),c2('%.0f'%d['tmax'])),c2('%.2f'%d['alpha'])]
    for k in range(5):  # columns labelled 2..6 => subscale index j=k+1
        j=k+1
        cells.append('—' if j<=i else cval(MSord[i],MSord[j]))
    return cells
t_desc=table('Descritivas, confiabilidade e intercorrelações das seis dimensões do BRUMS (%d observações).'%sm['n_obs'],
    ['Dimensão','M','DP','Amplitude','Escore T','α','2','3','4','5','6'],
    [mrow(i,d) for i,d in enumerate(MSd)],
    note='* p < 0,001 (correlação de Spearman). Amplitude = escore bruto 0–16; Escore T referenciado à amostra (M = 50; DP = 10).',fs=8)
SH=STAT['shapiro']; IC=STAT['icc']
def nrow(k,lab):
    sh=SH[k]; i=IC[k]; return [lab,c2('%.3f'%sh['W']),'< 0,001' if sh['p']<0.001 else c2('%.3f'%sh['p']),
        c2('%+.2f'%sh['skew']),c2('%.2f'%i['icc1']),c2('%.2f'%i['icck']),i['cls']]
t_norm=table('Normalidade (Shapiro-Wilk), assimetria e consistência das medidas repetidas (ICC) por dimensão.',
    ['Dimensão','W','p','Assimetria','ICC(2,1)','ICC(2,k)','Consistência'],
    [nrow(k,l) for k,l in [('Vigor','Vigor'),('Fadiga','Fadiga'),('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Confusao','Confusão')]],fs=9)
def cvrow(k,lab):
    v=CV[k]; return [lab,c2('%.0f'%v['inter']),c2('%.0f'%v['intra'])]
t_cv=table('Coeficiente de variação entre atletas e intraindividual, por dimensão.',
    ['Dimensão','CV entre atletas (%)','CV intraindividual (%)'],
    [cvrow(k,l) for k,l in [('Vigor','Vigor'),('Fadiga','Fadiga'),('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Confusao','Confusão')]],
    note='PTH excluída por assumir valores ≤ 0, que tornam o CV ininterpretável.',fs=9)
f_hist=figure(f'{FG}/xb3_hist.png','Distribuição de frequências das seis dimensões do BRUMS (linha tracejada = mediana).')
f_box=figure(f'{FG}/xb3_box.png','Diagramas de caixa das seis dimensões por dia (mediana, quartis e valores atípicos).')
f_splom=figure(f'{FG}/xb4_splom.png','Matriz de dispersão entre as seis dimensões (médias semanais por atleta).',w=14.5)

H('3.2 Perfil de humor e comparações intra-grupo (Dia 1 vs. Dia 7 e pré vs. pós)',12,before=6)
mvv=lambda tab,k,f: next(x[f] for x in MV[tab]['rows'] if x['k']==k)
P('O perfil de humor em escores T ilustra a transição do microciclo (Figura %d): no Dia 1 o grupo apresenta assinatura '
 'próxima do iceberg — vigor acima da média populacional (T = %s) e fadiga abaixo (T = %s) —, ao passo que no Dia 7 o '
 'padrão inverte-se (vigor T = %s; fadiga T = %s). A MANOVA de medidas repetidas confirmou diferença multivariada '
 'significativa entre o Dia 1 e o Dia 7 (Wilks λ = %s; F(%d,%d) = %s; p = %s; η²ₚ = %s). Os testes univariados (Tabela %d) '
 'localizam o efeito no eixo energia–fadiga: vigor (F = %s; p = %s; d = %s; η²ₚ = %s, grande) e fadiga (F = %s; p = %s; '
 'd = %s; η²ₚ = %s, grande), com contribuição adicional da tensão e da confusão e ausência de efeito em depressão e '
 'raiva.'%(
   0 if False else 0,  # placeholder replaced below
   c2('%.1f'%mvv('d1d7','Vigor','m1')),c2('%.1f'%mvv('d1d7','Fadiga','m1')),c2('%.1f'%mvv('d1d7','Vigor','m2')),c2('%.1f'%mvv('d1d7','Fadiga','m2')),
   c2('%.3f'%MV['d1d7']['wilks']),MV['d1d7']['df1'],MV['d1d7']['df2'],c2('%.2f'%MV['d1d7']['Fmv']),pstr(MV['d1d7']['p_mv']),c2('%.2f'%MV['d1d7']['eta_mv']),
   0,  # placeholder Tabela num
   c2('%.2f'%mvv('d1d7','Vigor','F')),pstr(mvv('d1d7','Vigor','p')),c2('%+.2f'%mvv('d1d7','Vigor','d')),c2('%.2f'%mvv('d1d7','Vigor','eta')),
   c2('%.2f'%mvv('d1d7','Fadiga','F')),pstr(mvv('d1d7','Fadiga','p')),c2('%+.2f'%mvv('d1d7','Fadiga','d')),c2('%.2f'%mvv('d1d7','Fadiga','eta'))))
# The above P has placeholders; rebuild cleanly instead:
# (remove the placeholder paragraph just added)
doc.paragraphs[-1]._p.getparent().remove(doc.paragraphs[-1]._p)
f_prof1=figure(f'{FG}/xb5_profile_d1d7.png','Perfil de humor em escores T (M = 50; DP = 10) no Dia 1 e no Dia 7.',w=14.0)
def mvtable(tab,cap):
    rows=[]
    for x in MV[tab]['rows']:
        sig='*' if x['p']<0.05 else ''
        rows.append([x['lab'],c2('%.1f'%x['m1']),c2('%.1f'%x['s1']),c2('%.1f'%x['m2']),c2('%.1f'%x['s2']),
            c2('%.2f'%x['F']),pstr(x['p'])+sig,c2('%+.2f'%x['d']),c2('%.3f'%x['eta'])])
    mv=MV[tab]
    note='Escores T. MANOVA: Wilks λ = %s; F(%d,%d) = %s; p = %s; η²ₚ = %s. * p < 0,05.'%(
        c2('%.3f'%mv['wilks']),mv['df1'],mv['df2'],c2('%.2f'%mv['Fmv']),pstr(mv['p_mv']),c2('%.2f'%mv['eta_mv']))
    hdr={'d1d7':['Dimensão','D1 M','D1 DP','D7 M','D7 DP','F','p','d','η²ₚ'],
         'prepos':['Dimensão','Pré M','Pré DP','Pós M','Pós DP','F','p','d','η²ₚ']}[tab]
    return table(cap,hdr,rows,note=note,fs=8.5)
t_mv1=table  # placeholder to keep name
p1=P('O perfil de humor em escores T ilustra a transição do microciclo (Figura %d): no Dia 1 o grupo apresenta assinatura '
 'próxima do iceberg — vigor acima da média populacional (T = %s) e fadiga abaixo (T = %s) —, ao passo que no Dia 7 o '
 'padrão inverte-se (vigor T = %s; fadiga T = %s). A MANOVA de medidas repetidas confirmou diferença multivariada '
 'significativa entre o Dia 1 e o Dia 7 (Wilks λ = %s; F(%d,%d) = %s; p = %s; η²ₚ = %s, grande). Os testes univariados '
 '(Tabela %d) localizam o efeito no eixo energia–fadiga: vigor (F = %s; p = %s; d = %s; η²ₚ = %s) e fadiga (F = %s; '
 'p = %s; d = %s; η²ₚ = %s), com contribuição da tensão e da confusão e ausência de efeito em depressão e raiva.'%(
   f_prof1,c2('%.1f'%mvv('d1d7','Vigor','m1')),c2('%.1f'%mvv('d1d7','Fadiga','m1')),c2('%.1f'%mvv('d1d7','Vigor','m2')),c2('%.1f'%mvv('d1d7','Fadiga','m2')),
   c2('%.3f'%MV['d1d7']['wilks']),MV['d1d7']['df1'],MV['d1d7']['df2'],c2('%.2f'%MV['d1d7']['Fmv']),pstr(MV['d1d7']['p_mv']),c2('%.2f'%MV['d1d7']['eta_mv']),
   _TN[0]+1,
   c2('%.2f'%mvv('d1d7','Vigor','F')),pstr(mvv('d1d7','Vigor','p')),c2('%+.2f'%mvv('d1d7','Vigor','d')),c2('%.2f'%mvv('d1d7','Vigor','eta')),
   c2('%.2f'%mvv('d1d7','Fadiga','F')),pstr(mvv('d1d7','Fadiga','p')),c2('%+.2f'%mvv('d1d7','Fadiga','d')),c2('%.2f'%mvv('d1d7','Fadiga','eta'))))
# move the just-added paragraph above the figure? keep order: text then figure already placed. Reorder: we added figure before text. Move figure after text:
figp=None
t_mv_d1d7=mvtable('d1d7','Comparação Dia 1 → Dia 7 das seis dimensões em escores T (MANOVA de medidas repetidas; n = %d).'%MV['d1d7']['n'])
P('A resposta aguda pré → pós, agregada na semana, também foi multivariadamente significativa (Wilks λ = %s; F(%d,%d) = '
 '%s; p = %s; η²ₚ = %s), novamente conduzida pelo eixo energia–fadiga: o vigor cai e a fadiga e a PTH sobem do pré para o '
 'pós, indicando um choque agudo consistente a cada sessão (Tabela %d; Figura %d).'%(
   c2('%.3f'%MV['prepos']['wilks']),MV['prepos']['df1'],MV['prepos']['df2'],c2('%.2f'%MV['prepos']['Fmv']),pstr(MV['prepos']['p_mv']),c2('%.2f'%MV['prepos']['eta_mv']),
   _TN[0]+1,_FN[0]+1))
t_mv_prepos=mvtable('prepos','Resposta aguda pré → pós das seis dimensões em escores T (MANOVA de medidas repetidas; n = %d).'%MV['prepos']['n'])
f_prof2=figure(f'{FG}/xb5_profile_prepos.png','Perfil de humor em escores T (M = 50; DP = 10) nos momentos pré e pós-treino.',w=14.0)
P('A leitura conjunta das sete variáveis padronizadas por dia (Figura %d) e o contraste direto entre os dias extremos '
 '(Figura %d) confirmam o cruzamento vigor↓/fadiga↑. A ordenação de sensibilidade (Tabela %d) situa o vigor (|dz| = %s) e '
 'a fadiga (|dz| = %s) como as dimensões mais sensíveis; o teste de Friedman acusou diferença entre os sete dias para '
 'vigor (p = %s), fadiga (p = %s), tensão (p = %s) e confusão (p = %s).'%(
   _FN[0]+1,_FN[0]+2,_TN[0]+1,c2('%.2f'%SENS['Vigor']['absdz']),c2('%.2f'%SENS['Fadiga']['absdz']),
   pstr(FR['Vigor']['p']),pstr(FR['Fadiga']['p']),pstr(FR['Tensao']['p']),pstr(FR['Confusao']['p'])))
f_allz=figure(f'{FG}/xb4_allz.png','Comportamento conjunto das sete variáveis ao longo da semana (escores padronizados z).')
f_dumb=figure(f'{FG}/xb4_dumbbell.png','Comparação Dia 1 versus Dia 7 dos escores de todas as variáveis.',w=14.5)
snames={'Vigor':'Vigor','Fadiga':'Fadiga','TMD':'PTH','Tensao':'Tensão','Depressao':'Depressão','Raiva':'Raiva','Confusao':'Confusão'}
sord=sorted(snames,key=lambda k:-SENS[k]['absdz'])
def eta_cls(e): return 'pequeno' if e<0.06 else ('médio' if e<0.14 else 'grande')
def frow2(k,rank):
    fr=FR[k]; sg='sim' if fr['p']<0.05 else 'não'
    return [rank,snames[k],c2('%+.2f'%SENS[k]['dz']),c2('%.1f'%fr['chi']),pstr(fr['p']),c2('%.2f'%fr['W']),sg]
t_sens=table('Sensibilidade à variação semanal e teste de Friedman (7 dias) por variável, ordenadas por |dz| (Dia 1 → Dia 7).',
    ['Ordem','Variável','dz (D1→D7)','Friedman χ²','p','W de Kendall','Difere entre dias'],
    [frow2(k,i+1) for i,k in enumerate(sord)],
    note='Friedman e W de Kendall sobre casos completos (n = %d).'%FR['Vigor']['n'],fs=8.5)
P('Cada dimensão é ainda apresentada individualmente, com banda de confiança de 95%%, diagramas de caixa por dia e o '
 'efeito do microciclo (Figuras %d a %d), iniciando pelo eixo energia–fadiga.'%(_FN[0]+1,_FN[0]+6))
FVAR=[('Vigor','xb4_v_Vigor.png','Vigor'),('Fadiga','xb4_v_Fadiga.png','Fadiga'),('Tensao','xb4_v_Tensao.png','Tensão'),
      ('Depressao','xb4_v_Depressao.png','Depressão'),('Raiva','xb4_v_Raiva.png','Raiva'),('Confusao','xb4_v_Confusao.png','Confusão')]
for k,fn,lab in FVAR:
    figure(f'{FG}/{fn}','%s ao longo da semana: médias diárias (banda = IC95%%), diagramas de caixa por dia e efeito Dia 1 → Dia 7.'%lab,w=12.5)

H('3.3 Estrutura de correlação e ênfase nos dias de maior vigor e de maior fadiga',12,before=6)
PA=STAT['pairs']; FP=STAT['focusp']
P('O contraste entre os dias extremos é o achado correlacional mais expressivo (Tabela %d). No dia de maior vigor (Dia 1, '
 'grupo descansado), o afeto negativo é praticamente independente da fadiga (depressão ρ = %s, p = %s; raiva ρ = %s, p = '
 '%s). No dia de maior fadiga (Dia 7), o acoplamento torna-se forte e significativo (depressão ρ = %s, p = %s; raiva ρ = '
 '%s, p = %s). Ou seja, quando o grupo está fresco, quem está mais irritado ou abatido não é necessariamente o mais '
 'cansado; sob carga acumulada, as dimensões negativas consolidam-se com a fadiga e o perfil de humor “fecha” (Tabela '
 '%d).'%(_TN[0]+1,
   c2('%+.2f'%FP['D1']['Depressao']['rho_fad']),pstr(FP['D1']['Depressao']['p_fad']),c2('%+.2f'%FP['D1']['Raiva']['rho_fad']),pstr(FP['D1']['Raiva']['p_fad']),
   c2('%+.2f'%FP['D7']['Depressao']['rho_fad']),pstr(FP['D7']['Depressao']['p_fad']),c2('%+.2f'%FP['D7']['Raiva']['rho_fad']),pstr(FP['D7']['Raiva']['p_fad']),_TN[0]+2))
def frow(neg,lab):
    d1=FP['D1'][neg]; d7=FP['D7'][neg]
    return [lab,'%s (%s)'%(c2('%+.2f'%d1['rho_fad']),pstr(d1['p_fad'])),'%s (%s)'%(c2('%+.2f'%d7['rho_fad']),pstr(d7['p_fad'])),
            '%s (%s)'%(c2('%+.2f'%d1['rho_vig']),pstr(d1['p_vig'])),'%s (%s)'%(c2('%+.2f'%d7['rho_vig']),pstr(d7['p_vig']))]
t_focus=table('Correlação (ρ; p) das dimensões negativas com a fadiga e o vigor no dia de maior vigor (D1) e no de maior fadiga (D7).',
    ['Dimensão','ρ×Fadiga D1 (p)','ρ×Fadiga D7 (p)','ρ×Vigor D1 (p)','ρ×Vigor D7 (p)'],
    [frow('Depressao','Depressão'),frow('Raiva','Raiva'),frow('Confusao','Confusão'),frow('Tensao','Tensão')],fs=8.5)
t_pairs=table('Correlação de Spearman entre as dimensões do BRUMS (nível entre atletas, n = %d): coeficiente (ρ) e valor de p.'%sm['n'],
    ['Par de dimensões','ρ','p'],
    [['%s × %s'%(pr_['a'],pr_['b']),c2('%+.2f'%pr_['rho']),pstr(pr_['p'])] for pr_ in PA],fs=9)

H('3.4 Prevalência e transição dos perfis de humor',12,before=6)
P('A distribuição dos seis perfis de humor mudou significativamente entre os dias extremos (χ² = %s; gl = %d; p = %s; '
 'Figura %d; Tabela %d). Na linha de base (Dia 1) predomina o perfil iceberg (%s%%), marca de prontidão; no último dia '
 '(Dia 7) o perfil de fadiga (“barbatana de tubarão”) torna-se o mais frequente (%s%%) e o iceberg cai para %s%%, sem '
 'instalação relevante de perfis de risco extremo (submerso, iceberg invertido).'%(
   c2('%.2f'%PREV['chi']),PREV['dof'],pstr(PREV['p']),_FN[0]+1,_TN[0]+1,
   c2('%.0f'%(100*PREV['D1']['Iceberg']/PREV['n_d1'])),c2('%.0f'%(100*PREV['D7']['Barbatana tubarão']/PREV['n_d7'])),
   c2('%.0f'%(100*PREV['D7']['Iceberg']/PREV['n_d7']))))
f_prev=figure(f'{FG}/xb5_prev.png','Prevalência (%) dos seis perfis de humor no Dia 1 e no Dia 7.',w=14.0)
PROFR=[('Iceberg','Iceberg'),('Everest invertido','Everest invertido'),('Iceberg invertido','Iceberg invertido'),
       ('Submerso','Submerso'),('Barbatana tubarão','Barbatana de tubarão'),('Superfície','Superfície')]
def prow(p,lab):
    d1=PREV['D1'][p]; d7=PREV['D7'][p]; ov=PREV['overall'][p]; nov=sum(PREV['overall'].values())
    return [lab,'%d (%s)'%(d1,c2('%.1f'%(100*d1/PREV['n_d1']))),'%d (%s)'%(ov,c2('%.1f'%(100*ov/nov))),'%d (%s)'%(d7,c2('%.1f'%(100*d7/PREV['n_d7'])))]
t_prev=table('Prevalência dos perfis de humor por recorte: n (%%) no Dia 1, na semana e no Dia 7.',
    ['Perfil','Dia 1 n (%)','Semana n (%)','Dia 7 n (%)'],
    [prow(p,lab) for p,lab in PROFR],
    note='Qui-quadrado de associação Dia 1 × Dia 7: χ² = %s; gl = %d; p = %s.'%(c2('%.2f'%PREV['chi']),PREV['dof'],pstr(PREV['p'])),fs=9)

H('3.5 Análise no nível do atleta: trajetórias, mudança confiável e transição de perfis',12,before=6)
P('As trajetórias individuais do vigor e da fadiga (Figura %d) revelam ampla dispersão entre atletas em torno da média do '
 'grupo. A mudança confiável (RCI, Dia 1 → Dia 7) distribui-se de modo desigual (Tabela %d; Figura %d): a fadiga é a '
 'dimensão em que a maioria dos atletas deteriora de forma confiável (%d de %d; %d%%), seguida do vigor (%d%%). As '
 'dimensões negativas de valência não fadiga não pioram no nível individual — a tensão e a confusão exibem apenas melhoras '
 'confiáveis. A classificação individual (Tabela %d) mostra que, dos %d atletas com casos completos, o iceberg era o '
 'perfil de %d na linha de base e a barbatana de tubarão passou a ser o de %d no último dia, com transições '
 'heterogêneas.'%(
   _FN[0]+1,_TN[0]+1,_FN[0]+2,RC['Fadiga']['piora'],RC['Fadiga']['n'],RC['Fadiga']['pct_piora'],RC['Vigor']['pct_piora'],
   _TN[0]+2,PCNT['n'],PCNT['iceberg_D1'],PCNT['shark_D7']))
f_spa=figure(f'{FG}/xb4_spaghetti.png','Trajetórias individuais de vigor e fadiga ao longo da semana (linhas cinzas = atletas; linha colorida = média do grupo).')
def rcrow(k,lab):
    v=RC[k]; return [lab,c2('%.2f'%v['alpha']),c2('%.2f'%v['sediff']),v['n'],
        '%d (%d%%)'%(v['piora'],v['pct_piora']),'%d (%d%%)'%(v['estavel'],v['pct_est']),'%d (%d%%)'%(v['melhora'],v['pct_mel'])]
t_rci=table('Mudança confiável individual (RCI, Dia 1 → Dia 7) por dimensão: confiabilidade, erro-padrão da diferença e número de atletas (n = %d).'%RC['Vigor']['n'],
    ['Dimensão','α','EP_dif','n','Piora confiável','Estável','Melhora confiável'],
    [rcrow(k,lab) for k,lab in [('Vigor','Vigor'),('Fadiga','Fadiga'),('Depressao','Depressão'),('Raiva','Raiva'),('Tensao','Tensão'),('Confusao','Confusão')]],fs=9)
f_rci=figure(f'{FG}/xb2_rci6.png','Proporção de atletas com mudança confiável (RCI) por dimensão, do Dia 1 ao Dia 7.',w=13.5)
ABBR={'Iceberg':'Iceberg','Everest invertido':'Everest inv.','Iceberg invertido':'Iceberg inv.','Submerso':'Submerso','Barbatana tubarão':'Barbatana','Superfície':'Superfície'}
def tag(d1,d7):
    if d7=='Iceberg' and d1!='Iceberg': return 'melhora'
    if d7=='Barbatana tubarão' and d1!='Barbatana tubarão': return 'fadiga'
    if d1==d7: return 'estável'
    return 'transição'
rows_tr=[[aid,ABBR.get(PT[aid]['D1'],PT[aid]['D1']),ABBR.get(PT[aid]['D7'],PT[aid]['D7']),tag(PT[aid]['D1'],PT[aid]['D7'])] for aid in sorted(PT)]
t_trans=table('Perfil de humor de cada atleta no Dia 1 e no Dia 7 e classificação da transição (n = %d).'%PCNT['n'],
    ['Atleta','Perfil no Dia 1','Perfil no Dia 7','Transição'],rows_tr,fs=8.5)

H('3.6 Aptidão intermitente (T-CAR) e resposta de humor–fadiga',12,before=6)
LP=LIM['LIM']['PVini']; TC=LIM['TERC']
def preg(col,tab='TCAR1'):
    v=PV['pv']['wk_'+col][tab]; return v
P('O pico de velocidade inicial (T-CAR1) relacionou-se de forma coerente com a resposta de humor–fadiga da semana '
 '(Figura %d; Tabela %d): atletas com maior aptidão intermitente reportaram mais vigor (β = %s; R² = %s; ρ = %s; p = %s) '
 'e menos fadiga física (β = %s; R² = %s; ρ = %s; p = %s), com tendência semelhante para a fadiga do BRUMS (ρ = %s; p = '
 '%s). A dispersão com reta de regressão evidencia esse gradiente, no qual o eixo energia–fadiga se organiza ao longo da '
 'capacidade física do atleta.'%(
   _FN[0]+1,_TN[0]+1,
   c2('%+.2f'%preg('Vigor')['slope']),c2('%.2f'%preg('Vigor')['r2']),c2('%+.2f'%preg('Vigor')['rho']),pstr(preg('Vigor')['rho_p']),
   c2('%+.2f'%preg('FadFisica')['slope']),c2('%.2f'%preg('FadFisica')['r2']),c2('%+.2f'%preg('FadFisica')['rho']),pstr(preg('FadFisica')['rho_p']),
   c2('%+.2f'%preg('Fadiga')['rho']),pstr(preg('Fadiga')['rho_p'])))
f_pvsc=figure(f'{FG}/pv7_scatter.png','Dispersão e reta de regressão entre o pico de velocidade do T-CAR1 e as médias semanais de vigor, fadiga (BRUMS), PTH e fadiga física (linha tracejada = limiar de %s km/h; faixa sombreada = IC95%%).'%c2('%.1f'%LP['thr']),w=15.0)
def pregrow(col,lab):
    v=preg(col); return [lab,c2('%+.2f'%v['slope']),c2('%.2f'%v['r2']),c2('%+.2f'%v['rho']),pstr(v['rho_p'])]
t_pv=table('Regressão do pico de velocidade do T-CAR1 sobre as médias semanais de humor e fadiga (n = %d atletas).'%preg('Vigor')['n'],
    ['Desfecho semanal','β','R²','ρ','p'],
    [pregrow('Vigor','Vigor (BRUMS)'),pregrow('Fadiga','Fadiga (BRUMS)'),pregrow('TMD','PTH'),pregrow('FadFisica','Fadiga física')],
    note='β = inclinação por km/h; ρ = correlação de Spearman.',fs=9)
P('A estratificação em tercis de aptidão confirma o padrão (Figura %d): a fadiga física decresce do tercil de menor para '
 'o de maior pico de velocidade (%s → %s; Kruskal-Wallis p = %s). Por fim, um limiar de PV ≈ %s km/h discriminou os dias '
 'de maior fadiga com desempenho aceitável (AUC = %s [IC95%% %s–%s]; sensibilidade = %s; especificidade = %s), indicando '
 'que abaixo desse pico de velocidade cresce a probabilidade de dias de fadiga elevada — uma referência prática para '
 'individualizar a carga (Figura %d).'%(
   _FN[0]+1,c2('%.1f'%TC['Baixa']['FadFis']),c2('%.1f'%TC['Alta']['FadFis']),pstr(PV['terc_kruskal_fadfis']['p']),
   c2('%.1f'%LP['thr']),c2('%.2f'%LP['auc']),c2('%.2f'%LP['lo']),c2('%.2f'%LP['hi']),c2('%.2f'%LP['sens']),c2('%.2f'%LP['spec']),_FN[0]+2))
f_pvterc=figure(f'{FG}/pv8_tercis.png','Fadiga semanal (BRUMS e física) por tercil de aptidão intermitente (T-CAR1).',w=12.5)
f_pvlim=figure(f'{FG}/pv9_limiar.png','Pico de velocidade do T-CAR1 e fadiga física semanal, com reta de regressão, IC95%% e limiar discriminante.',w=13.5)

H('3.7 Análises complementares: pós-teste, ajustes e decomposição sinal–ruído',12,before=6)
KD=PHJ['keydays']
def emmc(v,d): return c2('%.2f'%PHJ[v]['emm'][str(d)])
def sig1(v,d):
    if d==1: return ''
    return '*' if PHJ[v]['pairs']['1_%d'%d]['ptukey']<0.05 else ''
P('O pós-teste das médias marginais estimadas por modelo misto localiza no tempo a deterioração e identifica os dias '
 'críticos (Tabela %d; Figura %d). A queda do vigor é mais acentuada logo no início do microciclo — Dia %d → Dia %d '
 '(Δ = %s) — refletindo o choque de carga, ao passo que a fadiga (física e do BRUMS) sobe de forma mais intensa no final '
 '— Dia %d → Dia %d (Δ = %s na fadiga física; Δ = %s na fadiga do BRUMS). Ambas culminam no Dia 7, onde o vigor atinge o '
 'mínimo e a fadiga o máximo, com todas as comparações Dia 1 vs. Dia 2–7 significativas para o vigor e para a fadiga '
 'física (Tukey p < 0,05).'%(
   _TN[0]+1,_FN[0]+1,KD['vigor_maxdrop_day']-1,KD['vigor_maxdrop_day'],c2('%+.2f'%KD['vigor_maxdrop']),
   KD['fadfis_maxrise_day']-1,KD['fadfis_maxrise_day'],c2('%+.2f'%KD['fadfis_maxrise']),c2('%+.2f'%KD['fadbrums_maxrise'])))
def phrow(d):
    return ['Dia %d'%d,emmc('Vigor',d)+sig1('Vigor',d),emmc('Fadiga',d)+sig1('Fadiga',d),emmc('FadFisica',d)+sig1('FadFisica',d)]
t_ph=table('Médias marginais estimadas (modelo misto) por dia e pós-teste em relação ao Dia 1.',
    ['Dia','Vigor','Fadiga (BRUMS)','Fadiga física'],[phrow(d) for d in range(1,8)],
    note='* diferença significativa vs. Dia 1 (Tukey, p < 0,05).',fs=9)
f_ph=figure(f'{FG}/ph_emm.png','Trajetória diária (médias marginais do modelo misto) de vigor e fadiga, com pós-teste vs. Dia 1 e destaque dos dias de maior variação.')
P('O ajuste alométrico da trajetória (Y = a · dia^b) descreve o curso temporal da fadiga como um processo saturante '
 '(b < 1): a fadiga física ajusta-se com b = %s (R² = %s) e a fadiga do BRUMS com b = %s (R² = %s), indicando subida '
 'rápida no início e desaceleração ao longo da semana (Tabela %d; Figura %d). No plano interindividual, o escalonamento '
 'alométrico da aptidão (PV = a · massa^%s) mostrou que parte da associação bruta entre pico de velocidade e fadiga '
 'reflete o tamanho corporal: a correlação aptidão × fadiga física caiu de ρ = %s (p = %s) para ρ = %s (p = %s) após '
 'normalizar o PV pela massa, sugerindo que a proteção contra a fadiga é, em parte, um efeito de dimensão corporal.'%(
   c2('%.2f'%ALLO['FadFisica']['b']),c2('%.2f'%ALLO['FadFisica']['r2']),c2('%.2f'%ALLO['Fadiga']['b']),c2('%.2f'%ALLO['Fadiga']['r2']),
   _TN[0]+1,_FN[0]+1,c2('%.2f'%ALLO['fitness']['b']),
   c2('%+.2f'%ALLO['fitness']['rho_raw']),pstr(ALLO['fitness']['p_raw']),c2('%+.2f'%ALLO['fitness']['rho_allo']),pstr(ALLO['fitness']['p_allo'])))
def allorow(k,lab):
    v=ALLO[k]; return [lab,c2('%.2f'%v['a']),c2('%.2f'%v['b']),c2('%.2f'%v['r2']),pstr(v['p'])]
t_allo=table('Ajuste alométrico da trajetória diária (Y = a · dia^b) das variáveis de fadiga.',
    ['Variável','a','b','R²','p'],
    [allorow('FadFisica','Fadiga física'),allorow('Fadiga','Fadiga (BRUMS)'),allorow('TMD','PTH')],
    note='b < 1 = processo saturante (subida rápida seguida de desaceleração).',fs=9)
f_allo=figure(f'{FG}/x_allo.png','Ajuste alométrico (lei de potência) da trajetória de fadiga em escala linear e log-log.',w=15.0)
P('Por fim, a decomposição sinal–ruído quantifica a confiabilidade de uma medida isolada e a detectabilidade da mudança '
 '(Tabela %d; Figuras %d e %d). A confiabilidade de uma única aferição foi moderada (r = %s para o vigor; %s para a '
 'fadiga do BRUMS), e a razão sinal–ruído favoreceu o eixo energia–fadiga (RSR = %s no vigor; %s na fadiga). '
 'Consequentemente, a discriminação entre o Dia 7 e o Dia 1 melhora ao remover o ruído (médias diárias): a área sob a '
 'curva ROC da fadiga física passa de %s para %s, sendo a variável mais discriminante do microciclo. No nível '
 'individual, a mínima mudança detectável de uma medida (MDC95 = %s pontos de vigor) reduz-se com o número de coletas, o '
 'que fundamenta a agregação de múltiplas aferições por dia adotada neste estudo. Ressalva importante: a remoção de ruído '
 'não cria sinal onde não há — as subescalas negativas, próximas do piso, permanecem sem amplitude relevante mesmo após '
 'a filtragem.'%(
   _TN[0]+1,_FN[0]+1,_FN[0]+2,c2('%.2f'%DEN['rel']['Vigor']),c2('%.2f'%DEN['rel']['Fadiga']),
   c2('%.2f'%DEN['snr_raw']['Vigor']),c2('%.2f'%DEN['snr_raw']['Fadiga']),
   c2('%.2f'%ROC['FadFisica']['raw'][0]),c2('%.2f'%ROC['FadFisica']['filt'][0]),c2('%.1f'%DEN['kmdc']['Vigor']['1'])))
def snrow(k,lab):
    return [lab,c2('%.2f'%DEN['rel'][k]),c2('%.2f'%DEN['etm'][k]),c2('%.2f'%DEN['snr_raw'][k]),
        c2('%.2f'%ROC[k]['raw'][0]),c2('%.2f'%ROC[k]['filt'][0]),c2('%.1f'%DEN['kmdc'][k]['1']),c2('%.1f'%DEN['kmdc'][k]['2'])]
t_snr=table('Decomposição sinal–ruído: confiabilidade, erro típico, razão sinal–ruído, discriminação (ROC) e mudança detectável.',
    ['Variável','r (1 medida)','ETM','RSR','AUC c/ ruído','AUC s/ ruído','MDC95 (1)','MDC95 (2)'],
    [snrow('Vigor','Vigor'),snrow('Fadiga','Fadiga (BRUMS)'),snrow('TMD','PTH'),snrow('FadMental','Fadiga mental')],
    note='ETM = erro típico de medida; RSR = razão sinal–ruído; AUC = área sob a curva ROC (Dia 7 vs. Dia 1); MDC95 (k) = mínima mudança detectável com k coletas.',fs=8.5)
f_roc=figure(f'{FG}/x_roc.png','Curvas ROC (Dia 7 vs. Dia 1) com e sem ruído (medidas brutas vs. médias diárias).',w=14.0)
f_deriv=figure(f'{FG}/x_deriv_exp.png','Velocidade (barras) e aceleração (linha) diárias da trajetória de cada variável, evidenciando os dias de maior taxa de variação.',w=15.5)

H('3.8 Limites, derivadas e cruzamento vigor–fadiga (dia e momento)',12,before=6)
dV=CRS['dec']['Vigor']; dF=CRS['dec']['Fadiga']; ac=CRS['acute']; xc=CRS['cross_pp'][0]
P('A leitura conjunta de limites, derivadas e decomposição precisa o “onde” e o “quando” da resposta (Tabela %d; Figuras '
 '%d e %d). Em cada dia de treino, o vigor cai e a fadiga sobe da coleta pré para a pós (variação aguda média de %s no '
 'vigor e %s na fadiga), configurando um padrão em dente de serra. Ao longo da semana, a maior taxa de variação (derivada) '
 'ocorre logo no Dia %d — o choque de carga —, quando o vigor decresce a %s pontos/dia e a fadiga cresce a %s '
 'pontos/dia. O vigor supera a fadiga no Dia 1; as duas curvas cruzam-se pela primeira vez já no Dia 2 (entre as coletas '
 'pré e pós), oscilam próximas ao longo da semana e a fadiga passa a superar definitivamente o vigor no Dia 7, no qual o '
 'vigor atinge o mínimo (pós-treino) e a fadiga o máximo.'%(
   _TN[0]+1,_FN[0]+1,_FN[0]+2,c2('%+.2f'%ac['Vigor']['delta']),c2('%+.2f'%ac['Fadiga']['delta']),
   dV['pkvel_day'],c2('%.2f'%dV['pkvel_val']),c2('%+.2f'%dF['pkvel_val'])))
f_cross=figure(f'{FG}/x_cross.png','Cruzamento vigor × fadiga por dia e momento (pré → pós dentro de cada dia); linhas pontilhadas = cruzamentos.',w=15.0)
P('A decomposição da variância mostra que o humor é dominado pela diferença entre atletas (%s%%–%s%% da variância total), '
 'com componente sistemática de dia (o sinal do microciclo) pequena mas presente no vigor (%s%%) e na fadiga (%s%%) e '
 'praticamente nula nas subescalas negativas de piso; o restante é ruído de medida (Figura %d). Isso reforça, do ponto de '
 'vista da variância, por que o monitoramento deve ser individualizado e por que o eixo energia–fadiga é o único a '
 'carregar sinal semanal aproveitável.'%(
   c2('%.0f'%min(CRS['dec'][k]['vc']['ath'] for k in ['Vigor','Fadiga','TMD','Tensao','Depressao'])),
   c2('%.0f'%max(CRS['dec'][k]['vc']['ath'] for k in ['Vigor','Fadiga','TMD','Tensao','Depressao'])),
   c2('%.0f'%dV['vc']['day']),c2('%.0f'%dF['vc']['day']),_FN[0]+1))
f_vd=figure(f'{FG}/x_vardecomp.png','Decomposição da variância de cada variável do humor em componentes entre atletas, dia (microciclo) e resíduo.',w=13.5)
def ldrow(k,lab):
    d=CRS['dec'][k]; return [lab,c2('%.2f'%d['amp_sig']),c2('%.2f'%d['amp_intra']),'D%d'%d['pkvel_day'],c2('%+.2f'%d['pkvel_val']),
        c2('%.0f'%d['vc']['ath']),c2('%.0f'%d['vc']['day']),c2('%.0f'%d['vc']['res'])]
t_ld=table('Limites (amplitude), derivada máxima e decomposição da variância por variável do humor.',
    ['Variável','Amplitude do sinal','Amplitude intra-atleta','Dia de maior taxa','Taxa (pontos/dia)','Var. entre atletas (%)','Var. dia (%)','Var. resíduo (%)'],
    [ldrow(k,l) for k,l in [('Vigor','Vigor'),('Fadiga','Fadiga'),('TMD','PTH'),('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Confusao','Confusão')]],
    note='Amplitude do sinal = variação da média diária; amplitude intra-atleta = variação média intraindividual na semana; derivada por ajuste cúbico.',fs=8)

H('3.9 Sono, estresse percebido e humor',12,before=6)
ep=EX['epw']; ps=EX['pss']; sc=EX['sleepy_cmp']
def cx(c): return '%s (p = %s)'%(c2('%+.2f'%c['rho']),pstr(c['p']))
P('A sonolência diurna (Epworth = %s ± %s; %d dos %d atletas com sonolência excessiva, escore > 10) associou-se ao pior '
 'humor: correlacionou-se positivamente com a fadiga (ρ = %s; p = %s) e com a PTH (ρ = %s; p = %s) e negativamente, em '
 'tendência, com o vigor (ρ = %s; p = %s) no nível entre atletas (Tabela %d; Figura %d). Consistentemente, os atletas '
 'mais sonolentos apresentaram fadiga semanal marcadamente maior do que os menos sonolentos (%s vs. %s; p = %s). O '
 'estresse percebido (PSS = %s ± %s), por sua vez, não se relacionou de forma significativa com nenhuma dimensão do humor '
 '(todos p > 0,05), indicando que, nesta amostra e janela, a fadiga do microciclo é mais acompanhada pela sonolência do '
 'que pelo estresse psicológico geral.'%(
   c2('%.1f'%ep['m']),c2('%.1f'%ep['sd']),ep['hi'],ep['n'],
   c2('%+.2f'%ep['corr']['Fadiga']['rho']),pstr(ep['corr']['Fadiga']['p']),c2('%+.2f'%ep['corr']['TMD']['rho']),pstr(ep['corr']['TMD']['p']),
   c2('%+.2f'%ep['corr']['Vigor']['rho']),pstr(ep['corr']['Vigor']['p']),_TN[0]+1,_FN[0]+1,
   c2('%.2f'%sc['Fadiga']['hi']),c2('%.2f'%sc['Fadiga']['lo']),pstr(sc['Fadiga']['p']),c2('%.1f'%ps['m']),c2('%.1f'%ps['sd'])))
def sprow(k,lab):
    e=ep['corr'][k]; p=ps['corr'][k]; return [lab,cx(e),cx(p)]
t_sono=table('Correlação de Spearman da sonolência (Epworth) e do estresse percebido (PSS) com o humor (nível entre atletas, n = %d).'%ep['n'],
    ['Desfecho semanal','Epworth ρ (p)','PSS ρ (p)'],
    [sprow('Vigor','Vigor'),sprow('Fadiga','Fadiga'),sprow('TMD','PTH'),sprow('FadFisica','Fadiga física')],fs=9)
f_sono=figure(f'{FG}/x_sono.png','Sonolência (Epworth) e fadiga: dispersão com reta de regressão (linha tracejada = ponto de corte 10) e comparação do humor entre atletas sonolentos e não sonolentos.',w=15.0)

H('3.10 Resposta ao estímulo: dias com vs. sem HIIT',12,before=6)
hi=EX['hiit']; ha=EX['hiit_acute']
P('Os dias com sessão de treino intervalado de alta intensidade (HIIT: Dias 2, 4 e 7) apresentaram pior humor do que os '
 'dias sem HIIT (Dias 1, 3, 5 e 6): menor vigor (%s vs. %s; p = %s; dz = %s), maior fadiga (%s vs. %s; p = %s; dz = %s), '
 'maior PTH (p = %s) e maior fadiga física (p = %s) (Tabela %d; Figura %d). Contudo, a resposta aguda pré → pós não foi '
 'maior nos dias de HIIT (variação do vigor %s vs. %s; da fadiga %s vs. %s; ambos p > 0,05), sugerindo que a diferença '
 'reflete sobretudo o posicionamento das sessões no microciclo — inclusive o Dia 7, de fadiga acumulada — e não uma '
 'resposta intra-sessão específica ao HIIT.'%(
   c2('%.2f'%hi['Vigor']['hiit']),c2('%.2f'%hi['Vigor']['nohiit']),pstr(hi['Vigor']['p']),c2('%+.2f'%hi['Vigor']['dz']),
   c2('%.2f'%hi['Fadiga']['hiit']),c2('%.2f'%hi['Fadiga']['nohiit']),pstr(hi['Fadiga']['p']),c2('%+.2f'%hi['Fadiga']['dz']),
   pstr(hi['TMD']['p']),pstr(hi['FadFisica']['p']),_TN[0]+1,_FN[0]+1,
   c2('%+.2f'%ha['Vigor']['hiit']),c2('%+.2f'%ha['Vigor']['nohiit']),c2('%+.2f'%ha['Fadiga']['hiit']),c2('%+.2f'%ha['Fadiga']['nohiit'])))
def hrow(k,lab):
    v=hi[k]; return [lab,c2('%.2f'%v['hiit']),c2('%.2f'%v['nohiit']),c2('%+.2f'%(v['hiit']-v['nohiit'])),pstr(v['p']),c2('%+.2f'%v['dz'])]
t_hiit=table('Humor e fadiga nos dias com vs. sem HIIT (média por atleta; Wilcoxon pareado, n = %d).'%hi['Vigor']['n'],
    ['Variável','Dias com HIIT','Dias sem HIIT','Δ','p','dz'],
    [hrow('Vigor','Vigor'),hrow('Fadiga','Fadiga'),hrow('TMD','PTH'),hrow('FadFisica','Fadiga física')],
    note='Dias com HIIT = 2, 4 e 7; dias sem HIIT = 1, 3, 5 e 6.',fs=9)
f_hiit=figure(f'{FG}/x_hiit.png','Humor e fadiga nos dias com vs. sem HIIT (* p < 0,05, Wilcoxon pareado).',w=12.5)

# ===================== 4 DISCUSSÃO =====================
H('4 DISCUSSÃO')
H('4.1 O eixo energia–fadiga como marcador central da carga (H1)',12,before=6)
P('A deterioração do humor no microciclo pré-competitivo concentrou-se no eixo energia–fadiga, confirmando a H1. A MANOVA '
 'Dia 1 → Dia 7 foi multivariadamente significativa (η²ₚ = %s), com os maiores efeitos univariados no vigor (η²ₚ = %s) e '
 'na fadiga (η²ₚ = %s), enquanto as dimensões negativas de valência não fadiga permaneceram próximas do piso. O padrão é '
 'coerente com a literatura de monitoramento, que identifica o vigor e a fadiga como as dimensões subjetivas mais '
 'sensíveis à carga (SAW; MAIN; GASTIN, 2016; THORPE et al., 2017), e reproduz, em sete dias, o “derretimento do iceberg” '
 'descrito em fases de acúmulo.'%(
   c2('%.2f'%MV['d1d7']['eta_mv']),c2('%.2f'%mvv('d1d7','Vigor','eta')),c2('%.2f'%mvv('d1d7','Fadiga','eta'))))
H('4.2 Migração do perfil iceberg para a fadiga funcional (H2)',12,before=6)
P('A prevalência de perfis mudou significativamente entre os dias extremos (χ² = %s; p = %s), com queda do iceberg e '
 'ascensão da barbatana de tubarão, confirmando a H2. Essa migração — da prontidão para a assinatura de fadiga funcional, '
 'sem instalação relevante de perfis de risco extremo — é compatível com o arcabouço dos seis perfis prototípicos '
 'validados em grandes amostras (PARSONS-SMITH; TERRY; MACHIN, 2017; HAN; PARSONS-SMITH; TERRY, 2020) e com a associação '
 'entre perturbação do humor e desfechos esportivos (LOCHBAUM et al., 2021; DE MIRANDA ROHLFS et al., 2025).'%(
   c2('%.2f'%PREV['chi']),pstr(PREV['p'])))
H('4.3 Acoplamento do afeto negativo à fadiga sob carga acumulada (H3)',12,before=6)
P('O afeto negativo acoplou-se à fadiga apenas sob carga acumulada (H3): a correlação depressão × fadiga passou de ρ = %s '
 'no dia de maior vigor para ρ = %s no de maior fadiga, com padrão análogo para a raiva. Isso sugere que, em atletas de '
 'elite, a irritabilidade e o abatimento não são ruído aleatório, mas sinais que se organizam com a exaustão, reforçando '
 'o valor de interpretar as dimensões em conjunto e à luz do estado de fadiga — leitura consistente com o uso do perfil '
 'de humor como triagem de bem-estar (ROHLFS et al., 2023).'%(
   c2('%+.2f'%foc['D1']['Depressao']['rho_fad']),c2('%+.2f'%foc['D7']['Depressao']['rho_fad'])))
H('4.4 Heterogeneidade individual e implicações para o monitoramento (H4)',12,before=6)
P('A resposta foi heterogênea entre atletas (H4): a fadiga foi a única dimensão em que a maioria deteriorou de forma '
 'confiável (%d%%), com trajetórias e transições de perfil diversas. Esse achado sustenta um monitoramento individualizado, '
 'por tendência e referenciado à linha de base de cada atleta, em linha com as recomendações de consenso (KELLMANN et al., '
 '2018). Do ponto de vista aplicado, recomenda-se acompanhar o eixo energia–fadiga com atenção redobrada ao início do '
 'microciclo (choque de carga) e ao último dia (fadiga acumulada), sinalizando para avaliação complementar os atletas que '
 'reportem os perfis mais negativos.'%RC['Fadiga']['pct_piora'])
H('4.5 Aptidão intermitente como moderador da resposta de humor–fadiga',12,before=6)
P('A inclusão do pico de velocidade do T-CAR como parâmetro fisiológico agregou uma leitura interindividual à resposta '
 'afetiva: a maior aptidão aeróbia intermitente associou-se a mais vigor (ρ = %s) e a menos fadiga física (ρ = %s) ao '
 'longo da semana, com gradiente confirmado entre tercis de aptidão (p = %s). Esse achado é coerente com o papel da '
 'capacidade intermitente na tolerância ao esforço repetido do handebol (KARCHER; BUCHHEIT, 2014; MICHALSIK; AAGAARD, '
 '2015) e com a validade do pico de velocidade do T-CAR como marcador de desempenho físico em modalidades intermitentes '
 '(FERNANDES-DA-SILVA et al., 2016). Na prática, normalizar a resposta de humor pela aptidão física ajuda a distinguir a '
 'fadiga esperada — de atletas menos aptos sob a mesma carga — daquela que sinaliza sobrecarga, refinando a '
 'individualização do monitoramento; o limiar de pico de velocidade identificado oferece uma referência objetiva para '
 'sinalizar atletas sob maior risco de dias de fadiga elevada.'%(
   c2('%+.2f'%PV['pv']['wk_Vigor']['TCAR1']['rho']),c2('%+.2f'%PV['pv']['wk_FadFisica']['TCAR1']['rho']),pstr(PV['terc_kruskal_fadfis']['p'])))
H('4.6 Robustez: pós-teste, ajustes e razão sinal–ruído',12,before=6)
P('O conjunto de análises complementares reforça e refina os achados centrais. O pós-teste do modelo misto confirma que '
 'a deterioração do vigor se concentra no choque inicial de carga (Dia 1 → Dia 2), enquanto a fadiga se acumula '
 'progressivamente e atinge o máximo no Dia 7 — informação diretamente acionável para escalonar a recuperação nesses '
 'pontos do microciclo. O ajuste alométrico da trajetória (b < 1) descreve a fadiga como um processo saturante, coerente '
 'com a acomodação fisiológica ao longo da semana, e o escalonamento alométrico da aptidão evidencia que parte da '
 'proteção contra a fadiga atribuída ao pico de velocidade decorre do tamanho corporal — um cuidado interpretativo '
 'relevante ao normalizar respostas por marcadores de aptidão (NEVILL; LANE, 2007). Por fim, a decomposição sinal–ruído '
 'demonstra que a agregação de múltiplas coletas melhora a confiabilidade e a detecção da mudança (HOPKINS, 2000): a '
 'discriminação entre o início e o fim do microciclo é máxima para a fadiga física e aumenta ao filtrar o ruído, ao passo '
 'que as subescalas negativas de piso permanecem sem sinal — o que delimita, de forma honesta, onde o monitoramento do '
 'humor é informativo.')
H('4.7 Sono, estresse e o estímulo de HIIT',12,before=6)
P('A associação da sonolência (Epworth) com a fadiga e a PTH, e a ausência de relação do estresse percebido (PSS) com o '
 'humor, reforçam que, neste microciclo, a deterioração afetiva acompanha marcadores de recuperação insuficiente (sono) '
 'mais do que a carga psicológica geral — achado convergente com a literatura que liga qualidade de sono e humor em '
 'atletas de elite (ANDRADE et al., 2016). Que quase metade dos atletas tenha reportado sonolência excessiva sublinha o '
 'valor de monitorar o sono em conjunto com o humor. Os dias de HIIT concentraram o pior humor, mas a resposta aguda '
 'pré → pós não diferiu dos demais dias, indicando que o efeito é de acúmulo ao longo do microciclo (com o Dia 7 somando '
 'HIIT e fadiga terminal) e não de um choque intra-sessão exclusivo do HIIT — leitura coerente com o caráter intermitente '
 'e progressivo da carga.')
H('4.8 Limitações e direções futuras',12,before=6)
P('Como limitações, destacam-se o tamanho amostral (n = %d) e o efeito de piso das dimensões negativas, que reduz a '
 'variância e a confiabilidade da tensão e da confusão e infla seus coeficientes de variação, e o caráter observacional '
 'de fase única, que não permite inferência causal sobre a carga. A conversão em escores T foi referenciada à própria '
 'amostra, e não a tabelas normativas nacionais, o que recomenda cautela na comparação absoluta com outras populações. '
 'Estudos futuros devem integrar o humor a marcadores de carga externa e interna, testar o valor preditivo dos perfis '
 'negativos para lesão e sobrecarga (DE MIRANDA ROHLFS et al., 2025) e ampliar a amostra com múltiplos microciclos.'%sm['n'])

# ===================== 5 CONCLUSÕES =====================
H('5 CONCLUSÕES')
P('Na última semana de pré-temporada, o perfil de humor de handebolistas de elite migrou da prontidão (iceberg) para a '
 'fadiga funcional (barbatana de tubarão), com a deterioração concentrada no eixo energia–fadiga — vigor em queda de '
 'efeito grande e fadiga em ascensão, as duas dimensões mais sensíveis —, tanto no grupo (MANOVA multivariadamente '
 'significativa) quanto no sujeito, onde a fadiga foi a única dimensão a piorar de forma confiável na maioria dos atletas. '
 'As dimensões negativas mantiveram-se estáveis e só se acoplaram à fadiga sob carga acumulada. Reforça-se o valor do '
 'perfilamento do humor como componente do monitoramento do atleta, monitorando-se o eixo energia–fadiga por tendência e '
 'referenciado à linha de base individual.')

# ===================== REFERÊNCIAS =====================
H('REFERÊNCIAS')
refs=[
 'ANDRADE, A. et al. Sleep quality, mood and performance: a study of elite Brazilian volleyball athletes. Journal of Sports Science and Medicine, v. 15, n. 4, p. 601–605, 2016.',
 'ANDRADE, A. et al. Effect of practice exergames on the mood states and self-esteem of elementary school boys and girls during physical education classes: a cluster-randomized controlled trial. PLoS ONE, v. 15, n. 6, e0232392, 2020. DOI: 10.1371/journal.pone.0232392.',
 'BRANDT, R.; BEVILACQUA, G. G.; ANDRADE, A. Perceived sleep quality, mood states, and their relationship with performance among Brazilian elite athletes during a competitive period. Journal of Strength and Conditioning Research, v. 31, n. 4, p. 1033–1039, 2017.',
 'COHEN, S.; KAMARCK, T.; MERMELSTEIN, R. A global measure of perceived stress. Journal of Health and Social Behavior, v. 24, n. 4, p. 385–396, 1983. DOI: 10.2307/2136404.',
 'DE MIRANDA ROHLFS, I. C. P. et al. Mood states, injury status, and countermovement jump performance in Brazilian high-level sports. Sports, v. 13, n. 9, 303, 2025. DOI: 10.3390/sports13090303.',
 'FERNANDES-DA-SILVA, J. et al. The peak velocity derived from the Carminatti Test is related to physical match performance in young soccer players. Journal of Sports Sciences, v. 34, n. 24, p. 2238–2245, 2016. DOI: 10.1080/02640414.2015.1093646.',
 'HAN, C.; PARSONS-SMITH, R. L.; TERRY, P. C. Mood profiling in Singapore: cross-cultural validation and potential applications of mood profile clusters. Frontiers in Psychology, v. 11, 665, 2020. DOI: 10.3389/fpsyg.2020.00665.',
 'HOPKINS, W. G. Measures of reliability in sports medicine and science. Sports Medicine, v. 30, n. 1, p. 1–15, 2000. DOI: 10.2165/00007256-200030010-00001.',
 'JACOBSON, N. S.; TRUAX, P. Clinical significance: a statistical approach to defining meaningful change in psychotherapy research. Journal of Consulting and Clinical Psychology, v. 59, n. 1, p. 12–19, 1991. DOI: 10.1037/0022-006X.59.1.12.',
 'JOHNS, M. W. A new method for measuring daytime sleepiness: the Epworth Sleepiness Scale. Sleep, v. 14, n. 6, p. 540–545, 1991. DOI: 10.1093/sleep/14.6.540.',
 'KARCHER, C.; BUCHHEIT, M. On-court demands of elite handball, with special reference to playing positions. Sports Medicine, v. 44, n. 6, p. 797–814, 2014. DOI: 10.1007/s40279-014-0164-z.',
 'KELLMANN, M. et al. Recovery and performance in sport: consensus statement. International Journal of Sports Physiology and Performance, v. 13, n. 2, p. 240–245, 2018. DOI: 10.1123/ijspp.2017-0759.',
 'LEW, P. C. F. et al. Cross-cultural validation of the Malaysian Mood Scale and tests of between-group mood differences. International Journal of Environmental Research and Public Health, v. 20, n. 4, 3348, 2023. DOI: 10.3390/ijerph20043348.',
 'LOCHBAUM, M. et al. The Profile of Mood States and athletic performance: a meta-analysis of published studies. European Journal of Investigation in Health, Psychology and Education, v. 11, n. 1, p. 50–70, 2021. DOI: 10.3390/ejihpe11010005.',
 'MICHALSIK, L. B.; AAGAARD, P. Physical demands in elite team handball: comparisons between male and female players. Journal of Sports Medicine and Physical Fitness, v. 55, n. 9, p. 878–891, 2015.',
 'MICHALSIK, L. B.; MADSEN, K.; AAGAARD, P. Match performance and physiological capacity of female elite team handball players. International Journal of Sports Medicine, v. 35, n. 7, p. 595–607, 2013. DOI: 10.1055/s-0033-1358713.',
 'MORGAN, W. P. Selected psychological factors limiting performance: a mental health model. In: CLARKE, D. H.; ECKERT, H. M. (Ed.). Limits of human performance. Champaign: Human Kinetics, 1985. p. 70–80.',
 'NEVILL, A. M.; LANE, A. M. Why self-report “Likert” scale data should not be log-transformed. Journal of Sports Sciences, v. 25, n. 1, p. 1–2, 2007. DOI: 10.1080/02640410601111183.',
 'PARSONS-SMITH, R. L.; TERRY, P. C.; MACHIN, M. A. Identification and description of novel mood profile clusters. Frontiers in Psychology, v. 8, 1958, 2017. DOI: 10.3389/fpsyg.2017.01958.',
 'RATZ-SULYOK, F. Z. et al. Associations between endocrine status and stress, mood and psychosomatic status in elite handball players. Sports, v. 14, n. 7, 289, 2026. DOI: 10.3390/sports14070289.',
 'ROHLFS, I. C. P. M. et al. A Escala de Humor de Brunel (Brums): instrumento para detecção precoce da síndrome do excesso de treinamento. Revista Brasileira de Medicina do Esporte, v. 14, n. 3, p. 176–181, 2008.',
 'ROHLFS, I. C. P. M. et al. Psychometric characteristics of the Brazil Mood Scale among youth and elite athletes using two response time frames. Sports, v. 11, n. 12, 244, 2023. DOI: 10.3390/sports11120244.',
 'SAW, A. E.; MAIN, L. C.; GASTIN, P. B. Monitoring the athlete training response: subjective self-reported measures trump commonly used objective measures: a systematic review. British Journal of Sports Medicine, v. 50, n. 5, p. 281–291, 2016. DOI: 10.1136/bjsports-2015-094758.',
 'TERRY, P. C.; LANE, A. M.; FOGARTY, G. J. Construct validity of the Profile of Mood States — Adolescents for use with adults. Psychology of Sport and Exercise, v. 4, n. 2, p. 125–139, 2003. DOI: 10.1016/S1469-0292(02)00035-8.',
 'TERRY, P. C. et al. Mood profiling for sustainable mental health among athletes. Sustainability, v. 13, n. 11, 6116, 2021. DOI: 10.3390/su13116116.',
 'THORPE, R. T. et al. Monitoring fatigue status in elite team-sport athletes: implications for practice. International Journal of Sports Physiology and Performance, v. 12, s2, p. S2-27–S2-34, 2017. DOI: 10.1123/ijspp.2016-0434.',
 'WAGNER, H. et al. Individual and team performance in team-handball: a review. Journal of Sports Science and Medicine, v. 13, n. 4, p. 808–816, 2014.']
for rf in refs:
    p=doc.add_paragraph(); r=p.add_run(rf); r.font.size=Pt(11); p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(6); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

OUTP='/home/user/mdlucca/Artigos/Perfil_Humor_BRUMS_Descritivo.docx'
doc.save(OUTP); print('SAVED',OUTP,'| Tabelas',_TN[0],'Figuras',_FN[0])
