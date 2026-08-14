import warnings; warnings.filterwarnings('ignore')
import numpy as np, pandas as pd, json, os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
SC='/tmp/claude-0/-home-user-mdlucca/e1dba24c-b1d7-5908-9106-f2f4aaf3f56a/scratchpad'
FIG='/home/user/mdlucca/Artigos/figuras'
amp=json.load(open(f'{SC}/amp_docx.json')); grp=amp['grp']; per=amp['per']
DN=json.load(open(f'{SC}/denoise.json')); PF=json.load(open(f'{SC}/pair_fatigue.json'))
FF=PF['grp']['FadFisica']

doc=Document()
sec=doc.sections[0]
sec.top_margin=Inches(0.8); sec.bottom_margin=Inches(0.8); sec.left_margin=Inches(0.9); sec.right_margin=Inches(0.9)
st=doc.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(10);
from docx.enum.text import WD_LINE_SPACING
st.paragraph_format.line_spacing=1.06; st.paragraph_format.space_after=Pt(4)
def H(t,l=1): doc.add_heading(t,level=l)
def P(t,it=False,b=False,size=None,just=True):
    p=doc.add_paragraph(); r=p.add_run(t); r.italic=it; r.bold=b
    if size: r.font.size=Pt(size)
    if just: p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    return p
def fig(name,cap,w=5.5):
    path=f'{FIG}/{name}.png'
    if os.path.exists(path):
        doc.add_picture(path,width=Inches(w)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
        c=doc.add_paragraph(); r=c.add_run(cap); r.italic=True; r.font.size=Pt(8); c.alignment=WD_ALIGN_PARAGRAPH.CENTER
    else: print('MISS',name)
def table(headers,rows,cap=None,fs=8):
    if cap: c=doc.add_paragraph(); r=c.add_run(cap); r.bold=True; r.font.size=Pt(8.5)
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Light Grid Accent 1'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        cell=t.rows[0].cells[i]; cell.paragraphs[0].add_run(str(h)).bold=True; cell.paragraphs[0].runs[0].font.size=Pt(fs)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=str(v)
            if cells[i].paragraphs[0].runs: cells[i].paragraphs[0].runs[0].font.size=Pt(fs)

# ===== TÍTULO / APRESENTAÇÃO =====
ti=doc.add_heading('',0); ti.add_run('Do relatório descritivo à análise evolutiva do humor: aprofundamento estatístico, processamento em Python e auditoria interna por script'); ti.alignment=WD_ALIGN_PARAGRAPH.CENTER
P('Monitoramento psicométrico (BRUMS-24) em microciclo de choque de HIIT · handebol de elite · 27 atletas · 456 observações · 21–28/04/2024 · atletas anonimizados (A01–A27)',it=True,size=9,just=False).alignment=WD_ALIGN_PARAGRAPH.CENTER

H('1. Apresentação')
P('Este documento tem dois propósitos, articulados para a apreciação do orientador. Primeiro, dar transparência ao processo analítico: descrever, de forma verificável, como os arquivos brutos foram processados e analisados em Python e como cada resultado foi submetido a uma auditoria interna por script — isto é, recomputado de maneira independente do relatório original, para confirmar sua fidedignidade. Segundo, evoluir o relatório descritivo inicial (perfil de humor ao longo da semana) para uma análise evolutiva, que caracteriza não apenas o "quanto" o humor mudou, mas a forma da mudança no tempo, a fração dessa mudança que é sinal real versus ruído de medida, e as condições sob as quais a variação é interpretável no grupo e no atleta. A argumentação segue do dado bruto à conclusão, mostrando por que cada passo foi tomado.')

# ===== 2. PROCESSAMENTO E AUDITORIA =====
H('2. Processamento em Python e auditoria interna por script')
P('Os dados chegaram em planilhas de trabalho (Excel) contendo as respostas do BRUMS-24, as autoavaliações de fadiga e os registros de sessão. O tratamento seguiu um pipeline inteiramente roteirizado em Python (bibliotecas pandas, numpy, SciPy e statsmodels), o que garante que qualquer número deste relatório possa ser regenerado a partir da base, sem edição manual:')
for t in [
 '(i) Extração e leitura das planilhas (openpyxl/pandas), com conferência de tipos, faixas válidas e valores ausentes;',
 '(ii) Anonimização irreversível na fronteira do processamento: cada atleta recebeu um código A01–A27 e nenhum nome real integra as bases derivadas, as figuras ou os relatórios — uma varredura automática (leak-scan) roda antes de cada versionamento para impedir vazamento de identificação;',
 '(iii) Construção de bases derivadas por finalidade (humor por observação, itens do BRUMS, momentos pré/mid/pós, sessões de treino), preservando a estrutura aninhada observação ⊂ dia ⊂ atleta;',
 '(iv) Auditoria interna: cada estatística de destaque do relatório descritivo foi recomputada por um script independente e confrontada com o valor original; divergências seriam sinalizadas automaticamente.',
]:
    p=doc.add_paragraph(style='List Bullet'); p.add_run(t).font.size=Pt(9.5)
P('O resultado da auditoria é de concordância. Os achados longitudinais centrais — a queda do vigor e a elevação da fadiga do Dia 1 ao Dia 7, com seus tamanhos de efeito e níveis de significância — foram reproduzidos de forma idêntica; pequenas diferenças restringiram-se a métricas cuja definição admite variação operacional (a proporção de "perfil iceberg", sensível ao critério adotado, e a resposta aguda pré→pós, sensível a quais dias e a quais respostas do dia são tomadas como "antes" e "depois"). A Tabela 1 sintetiza a auditoria.')
audit=[
 ['Vigor — Dia 1 → Dia 7','7,6 → 4,5; dz −0,95','7,6 → 4,5; dz −0,95; p = 0,001','idêntico'],
 ['Fadiga — Dia 1 → Dia 7','4,3 → 7,5; dz +0,72','4,3 → 7,5; dz +0,72; p = 0,004','idêntico'],
 ['PTH/TMD — Dia 1 → Dia 7','dz +0,41; p = 0,104','dz +0,41; p = 0,104','idêntico'],
 ['Tensão — Dia 1 → Dia 7','dz −0,59','dz −0,59; p = 0,011','idêntico'],
 ['Perfil iceberg — D1 → D7','74% → 38%','71% → 33% (critério clássico)','concordante (Δ ≤ 5 pp; definição)'],
 ['Resposta aguda pré→pós','vigor↓, fadiga↑, PTH↑ (peq.–mod.)','mesma direção e ordem de magnitude','concordante (operacional)'],
 ['Amostra / observações','27 atletas / 456','27 atletas / 456','idêntico'],
]
table(['Métrica','Relatório descritivo (reportado)','Recomputado por script (auditoria)','Concordância'],audit,
      'Tabela 1 — Auditoria interna: reprodução independente das estatísticas de destaque',fs=8)

# ===== 3. ANÁLISE EVOLUTIVA =====
H('3. Análise evolutiva do humor')
P('Confirmada a base descritiva, a análise foi aprofundada em cinco camadas que respondem, de forma encadeada, à pergunta central do orientador: a evolução do humor no microciclo é real, e o que ela significa para a decisão prática?')

H('3.1. A forma da evolução ao longo da semana',2)
P('A trajetória do eixo energia–fadiga não é uma reta. Ajustando polinômios ortogonais no modelo misto, as componentes linear e cúbica são significativas (a quadrática, não), e o crescimento cúbico vence por critérios de informação (AIC/LRT, p < 0,001). Isso descreve três fases: uma subida inicial da fadiga, um alívio parcial no meio da semana e uma precipitação no Dia 7 — o maior passo diário de toda a série ocorre justamente de D6 para D7. A leitura das derivadas torna essa forma precisa: para a PTH/TMD, a velocidade diária da mudança percorre valores da ordem de +4,8 (D1) → +1,3 → −0,7 → −1,2 → −0,3 → +2,2 → +6,1 (D7), cruzando zero por volta dos dias 3 e 5,5 (o "alívio") e atingindo o módulo máximo no Dia 7. Ou seja, a velocidade não tende a zero (não haveria platô/saturação): ela apenas cruza zero no ponto de virada para logo reacelerar rumo ao pico do último dia, enquanto a aceleração muda de sinal uma única vez, no ponto de inflexão (~Dia 4). A evolução é, portanto, de acúmulo com desfecho abrupto — precipitação, não saturação —, e não de deriva linear uniforme.')
fig('adx1_grupo_semana','Figura 1 — Evolução semanal do grupo (média ± DP; suavização LOWESS pontilhada; linhas vermelhas = dias de HIIT). Vigor cai, Fadiga e PTH/TMD sobem com precipitação no Dia 7; a Fadiga mental permanece plana.')

H('3.2. A evolução dentro da sessão (pré → intra-sessão → pós)',2)
P('A mesma lógica de acúmulo reaparece em miniatura dentro de cada dia. Acrescentando o momento intra-sessão (mid) entre o pré e o pós, observa-se que o curso agudo não é monotônico: a fadiga frequentemente atinge o pico no meio da sessão e recua ligeiramente ao final, de modo que uma leitura apenas pré/pós subestima o ponto de maior carga percebida. A resposta aguda é mais nítida no Dia 1 (choque inicial) e no fim da semana, e sua direção (vigor↓, fadiga↑, PTH↑) é a mesma da evolução semanal — a sessão é um análogo, em menor amplitude, do processo crônico.')
fig('adx6_prepos_dia','Figura 2 — Evolução intra-dia (pré → mid → pós) das quatro variáveis do eixo energia–fadiga. O asterisco marca dias com resposta pareada significativa (Wilcoxon, p < 0,05).')

H('3.3. Sinal, traço e ruído: quanto da evolução é real',2)
P('Toda observação psicométrica combina três fontes: o sinal do microciclo (a evolução sistemática que interessa), o traço (as diferenças estáveis entre atletas) e o ruído (erro de medida e flutuação aleatória). Decompondo a variância de cada variável (modelo atleta + dia), o sinal do microciclo é, em todas elas, a menor fração (1–12%): a maior parte é traço (32–73%) e ruído (26–64%). A fadiga física concentra o maior sinal (12%) e a melhor relação sinal-ruído (SNR 2,0), confirmando-a como o marcador mais limpo; a fadiga mental é quase pura estrutura de traço (73%) e não responde ao microciclo (SNR < 1). Este é o argumento que sustenta por que a fadiga física e o par vigor–fadiga — e não as subescalas negativas nem a PTH agregada — carregam o sinal útil.')
rows=[['Fadiga física',f"{FF['sig']:.2f}",f"{FF['pday']:.0f}%",f"{FF['ptrait']:.0f}%",f"{FF['pnoise']:.0f}%",f"{FF['snr']:.2f}"]]
for v,lab in [('Vigor','Vigor'),('Fadiga','Fadiga BRUMS'),('TMD','PTH/TMD'),('FadMental','Fadiga mental')]:
    g=grp[v]; rows.append([lab,f"{g['sig']:.2f}",f"{g['pday']:.0f}%",f"{g['ptrait']:.0f}%",f"{g['pnoise']:.0f}%",f"{g['snr']:.2f}"])
table(['Variável','Ampl. do sinal','% Sinal','% Traço','% Ruído','SNR'],rows,'Tabela 2 — Decomposição da evolução em sinal, traço e ruído',fs=8)
fig('amp_a_variancia','Figura 3 — Origem da variância por variável: sinal do microciclo (vermelho) vs traço estável (azul) vs ruído (cinza).',w=5.3)
P('Há ainda uma razão psicométrica para as subescalas negativas não evoluírem: o efeito de piso. Cerca de 50% a 80% das respostas de tensão, depressão, raiva e confusão são zero, e a magnitude da resposta aguda de cada variável é quase perfeitamente predita pela sua proporção de piso (|dz| ≈ 0,588 − 0,0063·%piso; ρ = −0,92; R² = 0,85). Uma variável ancorada no piso não tem para onde descer e pouca margem para subir de forma detectável — sua "não evolução" é uma consequência métrica, não a ausência do fenômeno. Isso reforça, por um caminho independente, que o sinal deve ser buscado no par vigor–fadiga.')

H('3.4. Detectabilidade: a evolução é legível no grupo e no atleta?',2)
P('A distinção entre sinal e ruído tem consequência prática direta. Derivando do ruído a mudança mínima detectável (MDC95), a evolução semanal do grupo supera o limiar de ruído da média por 4 a 10 vezes — é, portanto, inequívoca no coletivo. No indivíduo, porém, a mesma oscilação fica abaixo do MDC95 de uma única coleta: não por ser irreal, mas porque o ruído da média de 27 atletas encolhe por √n, enquanto o de um atleta isolado, não. A recomendação do relatório descritivo — referenciar cada atleta à sua própria linha de base — ganha aqui uma quantificação: a decisão individual confiável requer médias de aproximadamente 3 a 5 coletas, que trazem o ruído para baixo da amplitude do sinal.')
rows=[['Fadiga física',f"{FF['sig']:.2f}",f"{FF['mdc']:.2f}",'não',f"{FF['mdc_g']:.2f}",'SIM']]
for v,lab in [('Vigor','Vigor'),('Fadiga','Fadiga BRUMS'),('TMD','PTH/TMD')]:
    g=grp[v]; rows.append([lab,f"{g['sig']:.2f}",f"{g['mdc']:.2f}",('SIM' if g['detect_ind'] else 'não'),f"{g['mdc_g']:.2f}",('SIM' if g['detect_grp'] else 'não')])
table(['Variável','Ampl. sinal','MDC95 individual','Detecta no indivíduo?','MDC95 grupo','Detecta no grupo?'],rows,
      'Tabela 3 — Detectabilidade em dois níveis (a oscilação semanal supera o ruído?)',fs=8)

H('3.5. Removendo o ruído: as conclusões ficam mais nítidas',2)
P('Por fim, perguntou-se: e se removêssemos o ruído? Removê-lo (por agregação, suavização ou correção pela confiabilidade) não altera a direção das conclusões — apenas as torna mais precisas. As associações do eixo energia–fadiga, uma vez descontado o erro de medida, aproximam-se de uma dimensão bipolar quase perfeita (vigor × fadiga passa de −0,44 para −0,93 no nível-grupo), e os tamanhos de efeito crescem (a fadiga física de dz +1,74 para +2,36; o vigor de −0,95 para −1,25): a leitura bruta subestimava o custo do microciclo. Crucialmente, remover o ruído não fabrica sinal onde não há — as subescalas negativas, ancoradas no piso, permanecem planas. A não resposta delas é um fenômeno real, não um artefato de ruído.')
Robs=pd.DataFrame(DN['Robs']); Rden=pd.DataFrame(DN['Rden'])
rows=[['Vigor × Fadiga',f"{Robs.loc['Vigor','Fadiga']:+.2f}",f"{Rden.loc['Vigor','Fadiga']:+.2f}"],
      ['Fadiga × PTH/TMD',f"{Robs.loc['Fadiga','TMD']:+.2f}",f"{Rden.loc['Fadiga','TMD']:+.2f}"],
      ['Fadiga física × Fadiga mental','+0,50','+0,37 (tendência) / +0,78 desatenuada']]
table(['Par','Correlação observada (com ruído)','Correlação sem ruído'],rows,'Tabela 4 — O efeito de remover o ruído sobre as associações',fs=8)

H('3.6. Contrastes entre variáveis: dissociação e redundância',2)
P('A análise evolutiva por pares esclarece o que cada variável acrescenta. A fadiga física e a fadiga mental dissociam ao longo da semana: a física quase dobra (dz +1,74) enquanto a mental permanece plana (dz +0,33). No instante, porém, acoplam-se quando removido o erro de medida (correlação 0,50 → 0,78 desatenuada) — quem está fisicamente exausto tende a relatar mais fadiga mental naquele momento —, mas a trajetória semanal só carrega a física (correlação de tendência 0,37). O custo do microciclo é somático, não cognitivo. Já o Vigor e a PTH/TMD são espelhos inversos do mesmo eixo (correlação −0,62 → −0,78 sem ruído); como o próprio vigor entra na fórmula da PTH com sinal invertido, a PTH não é independente e, tendo o pior sinal-ruído (por diluir o eixo no ruído das negativas de piso), pode ser substituída, sem perda de informação, pelo par vigor + fadiga. A análise evolutiva, portanto, não só descreve a mudança: hierarquiza os instrumentos por quanto sinal útil cada um entrega.')
fig('pf2_dissociacao','Figura 4 — Dissociação da fadiga física (acumula) e da fadiga mental (estável) ao longo da semana (curvas suavizadas).',w=5.3)

# ===== 4. ARGUMENTAÇÃO =====
H('4. Por que as conclusões são robustas')
P('A confiança nos achados não repousa em um único teste, mas na convergência de métodos independentes sobre a mesma conclusão. O efeito do dia é sustentado por modelos mistos, por análise de tendência polinomial, por permutação (PERMANOVA), por abordagem bayesiana e por curvas ROC — e, sobretudo, pela auditoria interna que reproduziu os números do relatório descritivo a partir do zero. A decomposição sinal-traço-ruído acrescenta a essa convergência uma leitura de magnitude e de confiabilidade: mostra que a evolução observada é pequena em fração de variância, porém real e direcional, e explicita as condições (agregação de coletas) sob as quais é acionável no indivíduo. O quadro é, assim, coerente em todas as escalas — semanal e intradiária, coletiva e individual, bruta e sem ruído.')

# ===== 5. SÍNTESE =====
H('5. Síntese e recomendações')
for t in [
 'A evolução do humor é real, não linear (cúbica) e concentrada no eixo energia–fadiga, com precipitação no Dia 7; a fadiga física é o marcador mais precoce, sensível e confiável.',
 'A maior parte da oscilação diária de um atleta é traço e ruído (sinal do microciclo = 1–12% da variância); interpretar oscilações isoladas como resposta ao treino é arriscado.',
 'A leitura do grupo é robusta com uma coleta/dia; a decisão individual exige médias de ≥3–5 coletas e limiares baseados no MDC, não valores absolutos pontuais.',
 'Removido o ruído, as conclusões ficam mais fortes, não diferentes; as subescalas negativas e a fadiga mental não são marcadores úteis de carga aguda deste microciclo — a PTH agregada pode ser substituída, sem perda, pelo par vigor + fadiga.',
]:
    p=doc.add_paragraph(style='List Bullet'); p.add_run(t).font.size=Pt(9.5)

# ===== 6. REPRODUTIBILIDADE =====
H('6. Reprodutibilidade e governança de dados')
P('Todo o processamento e as análises são reproduzíveis pelos scripts em Python depositados em scripts/analise/ (extração, normalidade, amplitude, decomposição sinal-ruído, pré/mid/pós, remoção de ruído e auditoria), executados sobre bases derivadas da janela de 21–28/04/2024. Os atletas são identificados exclusivamente por códigos A01–A27; nenhum dado bruto com identificação pessoal é distribuído, e uma varredura automática impede vazamento de nomes a cada versionamento. As figuras estão disponíveis em alta resolução (4K) e os relatórios detalhados de cada camada acompanham este documento.',it=False)
P('Documento gerado a partir do relatório descritivo enviado, unificado à reanálise evolutiva. Processamento e auditoria em Python (pandas, numpy, SciPy, statsmodels). Atletas anonimizados (A01–A27).',it=True,size=8.5)

OUTP='/home/user/mdlucca/Artigos/Analise_Evolutiva_Humor_Orientador.docx'
doc.save(OUTP); print('SAVED',OUTP)
d2=Document(OUTP); print('paras',len(d2.paragraphs),'tables',len(d2.tables),'imgs',len(d2.inline_shapes))
