# -*- coding: utf-8 -*-
import json, numpy as np, pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
FG='/home/user/mdlucca/Artigos/figuras/paper1'
T=json.load(open('temporal.json')); AR=json.load(open('all_results.json')); h=pd.read_csv('hum_prof.csv')
def c2(s): return str(s).replace('.',',')

doc=Document()
stl=doc.styles['Normal']; stl.font.name='Times New Roman'; stl.font.size=Pt(12)
stl.element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')
stl.paragraph_format.line_spacing=1.5; stl.paragraph_format.space_after=Pt(0)
sec=doc.sections[0]; sec.top_margin=Cm(2.5); sec.left_margin=Cm(3.0); sec.bottom_margin=Cm(2); sec.right_margin=Cm(2)
_TN=[0]; _FN=[0]
def P(t='',ind=True,after=0):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(12); r.font.name='Times New Roman'
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(after); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if ind: p.paragraph_format.first_line_indent=Cm(1.25)
def LI(t):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(12); r.font.name='Times New Roman'
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.left_indent=Cm(1.25); p.paragraph_format.space_after=Pt(0)
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
def H1(t,n=None):
    p=doc.add_paragraph(); r=p.add_run((n+' ' if n else '')+t); r.bold=True; r.font.size=Pt(12); r.font.name='Times New Roman'
    p.paragraph_format.space_before=Pt(14); p.paragraph_format.space_after=Pt(6)
def H2(t,n=None):
    p=doc.add_paragraph(); r=p.add_run((n+' ' if n else '')+t); r.bold=True; r.font.size=Pt(12); r.font.name='Times New Roman'
    p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(4)
def _bd(c):
    tcPr=c._tc.get_or_add_tcPr(); bd=OxmlElement('w:tcBorders')
    for e in ['top','bottom']:
        el=OxmlElement('w:'+e); el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'6'); el.set(qn('w:color'),'000000'); bd.append(el)
    tcPr.append(bd)
def table(cap,header,rows,fs=9,note=None,src=True):
    _TN[0]+=1
    p=doc.add_paragraph(); r=p.add_run('Tabela %d – %s'%(_TN[0],cap)); r.font.size=Pt(11); r.font.name='Times New Roman'
    p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(2); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    t=doc.add_table(rows=1,cols=len(header)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=True
    for i,ht in enumerate(header):
        cc=t.rows[0].cells[i]; cc.text=''; rr=cc.paragraphs[0].add_run(ht); rr.bold=True; rr.font.size=Pt(fs); rr.font.name='Times New Roman'
        cc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if i>0 else WD_ALIGN_PARAGRAPH.LEFT; _bd(cc)
    for row in rows:
        cs=t.add_row().cells
        for i,v in enumerate(row):
            cs[i].text=''; rr=cs[i].paragraphs[0].add_run(str(v)); rr.font.size=Pt(fs); rr.font.name='Times New Roman'
            cs[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if i>0 else WD_ALIGN_PARAGRAPH.LEFT; _bd(cs[i])
    if note:
        pn=doc.add_paragraph(); rn=pn.add_run(note); rn.font.size=Pt(9); rn.font.name='Times New Roman'; pn.paragraph_format.space_after=Pt(0)
    if src:
        ps=doc.add_paragraph(); rs=ps.add_run('Fonte: dados da pesquisa (2026).'); rs.font.size=Pt(9); rs.font.name='Times New Roman'; ps.paragraph_format.space_after=Pt(6)
def figure_bos(fn,cap,w=13.5):
    # figura no formato Biology of Sport: legenda "Figure N." em ingles, sem linha de fonte
    _FN[0]+=1
    pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; pp.add_run().add_picture(f'{FG}/{fn}',width=Cm(w)); pp.paragraph_format.space_before=Pt(6)
    pc=doc.add_paragraph(); pc.alignment=WD_ALIGN_PARAGRAPH.CENTER; rc=pc.add_run('Figure %d. %s'%(_FN[0],cap)); rc.font.size=Pt(10.5); rc.font.name='Times New Roman'
    pc.paragraph_format.space_after=Pt(6)
def figure(fn,cap,w=15.5):
    _FN[0]+=1
    pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; pp.add_run().add_picture(f'{FG}/{fn}',width=Cm(w)); pp.paragraph_format.space_before=Pt(6)
    pc=doc.add_paragraph(); pc.alignment=WD_ALIGN_PARAGRAPH.CENTER; rc=pc.add_run('Figura %d – %s'%(_FN[0],cap)); rc.font.size=Pt(10.5); rc.font.name='Times New Roman'
    pc.paragraph_format.space_after=Pt(0)
    ps=doc.add_paragraph(); ps.alignment=WD_ALIGN_PARAGRAPH.CENTER; rs=ps.add_run('Fonte: elaboração dos autores (2026).'); rs.font.size=Pt(9.5); rs.font.name='Times New Roman'; ps.paragraph_format.space_after=Pt(6)

# ===== TÍTULO =====
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('DINÂMICA DO HUMOR EM UM MICROCICLO PRÉ-COMPETITIVO DE ATLETAS DE HANDEBOL DE ELITE: '
            'PERFIS E TRAJETÓRIAS DAS SEIS DIMENSÕES DO BRUMS'); r.bold=True; r.font.size=Pt(14)
p.paragraph_format.space_after=Pt(12)

# ===== RESUMO =====
H1('RESUMO')
P('Objetivo: descrever e caracterizar a dinâmica do humor de atletas de handebol de elite ao longo de um '
  'microciclo pré-competitivo, com a análise das seis dimensões do BRUMS e da perturbação total do humor, do '
  'comportamento dos seis perfis de humor, da resposta aguda pré e pós-treino e da forma temporal das trajetórias, '
  'com destaque para o eixo energia-fadiga como achado dominante. Método: 27 atletas do sexo masculino responderam '
  'ao BRUMS-24 durante sete dias, com uma coleta de linha de base e duas coletas diárias (pré e pós-treino) nos '
  'seis dias de treino, e um total de 286 observações. Além da estatística descritiva, da consistência interna e '
  'da classificação dos perfis, as comparações entre dias e entre pré e pós-treino incluíram o tamanho e a '
  'magnitude do efeito, e as trajetórias das dimensões foram analisadas por suavização, segundas derivadas, ajuste '
  'polinomial, localização dos cruzamentos entre as dimensões, separação entre sinal e ruído e comparação '
  'sequencial do efeito agudo (pré para pós) e de recuperação (pós para pré do dia seguinte) por teste de Wilcoxon '
  'pareado. Resultados: entre as seis dimensões, a deterioração concentrou-se no eixo energia-fadiga: o vigor caiu '
  'e a fadiga subiu do primeiro para o último dia, com efeito grande no vigor e médio na fadiga (d = -0,95 e d = 0,72) e confirmação '
  'multivariada (Wilks lambda = 0,181; p < 0,001), ao passo que as dimensões negativas se mantiveram junto ao '
  'piso. As trajetórias suavizadas apresentaram um ponto de inflexão na metade da semana, e a fadiga ultrapassou '
  'o vigor de forma definitiva no fim do microciclo, com a PTH a superar ambas as dimensões. O vigor e a fadiga '
  'concentraram a maior razão sinal/ruído, ao passo que as dimensões negativas mostraram forte efeito de piso. A '
  'análise sequencial revelou um padrão em dente de serra, com piora aguda no treino, sobretudo no Dia 6 (PTH com '
  'dz = 0,75), e recuperação apenas parcial entre as sessões. A prevalência dos perfis deslocou-se do iceberg '
  '(48% no primeiro dia) para a barbatana de tubarão (22% no último dia), com aumento da chance desse perfil a '
  'cada dia (OR = 1,37). Conclusão: entre as seis dimensões, o humor migrou da prontidão para a fadiga funcional, '
  'em um padrão compatível com sobre-esforço funcional, o que recomenda centrar o monitoramento no par '
  'vigor-fadiga e no cruzamento entre as suas curvas.')
pk=doc.add_paragraph(); rr=pk.add_run('Palavras-chave: '); rr.bold=True; rr.font.size=Pt(11)
pk.add_run('humor; BRUMS; handebol; perfis de humor; fadiga; monitoramento do atleta.').font.size=Pt(11)
pk.paragraph_format.space_after=Pt(4)

# ===== 1 INTRODUÇÃO =====
H1('INTRODUÇÃO','1')
P('O monitoramento do estado do atleta consolidou-se como parte da gestão do treinamento no esporte de '
  'rendimento, e integra medidas objetivas de carga externa e interna a instrumentos subjetivos de autorrelato '
  '(HELWIG et al., 2023; KELLMANN et al., 2018). Entre esses instrumentos, os questionários de humor destacam-se '
  'pela praticidade, pelo baixo custo e pela sensibilidade às variações da carga, e mostram utilidade para o '
  'acompanhamento do bem-estar e do desempenho (SAW; MAIN; GASTIN, 2016; LOCHBAUM et al., 2021). Revisões '
  'sistemáticas indicam, inclusive, que as medidas subjetivas costumam responder às alterações de carga com '
  'sensibilidade igual ou superior à de muitos marcadores objetivos, o que sustenta o seu uso rotineiro (SAW; '
  'MAIN; GASTIN, 2016). Por esse motivo, documentos de consenso recomendam a vigilância regular da fadiga e da '
  'recuperação como base para as decisões de treino (KELLMANN et al., 2018).')
P('A Escala de Humor de Brunel (BRUMS) operacionaliza essa avaliação por meio de seis dimensões, a saber, tensão, '
  'depressão, raiva, vigor, fadiga e confusão, com propriedades psicométricas replicadas em diferentes idiomas e '
  'culturas (TERRY; LANE; FOGARTY, 2003; TERRY et al., 2022). A versão em português dispõe de validação para o '
  'contexto brasileiro e foi concebida, desde a origem, como ferramenta de detecção precoce de sinais associados '
  'ao excesso de treinamento (ROHLFS et al., 2008). A partir das seis dimensões calcula-se a perturbação total do '
  'humor (PTH), um índice-resumo do desequilíbrio afetivo que sintetiza o estado do atleta em um único valor e '
  'cuja associação com o desempenho tem sido documentada em meta-análise (LOCHBAUM et al., 2021).')
P('Para além dos escores isolados, o modelo dos perfis de humor organiza as seis dimensões em padrões '
  'reconhecíveis. O clássico perfil iceberg, com o vigor acima da média e as dimensões negativas rebaixadas, foi '
  'descrito no modelo de saúde mental como marca do atleta em prontidão (MORGAN, 1985). Estudos posteriores, com '
  'grandes amostras e análise de agrupamento, identificaram e replicaram seis perfis distintos, entre os quais o '
  'iceberg, a barbatana de tubarão, que sinaliza fadiga com vigor ainda preservado, e perfis de maior risco, como '
  'o iceberg invertido e o submerso (PARSONS-SMITH; TERRY; MACHIN, 2017; LUOJUMÄKI et al., 2026). A leitura por '
  'perfis aproxima o dado psicométrico da linguagem do treinador, facilita a comunicação com a comissão técnica e '
  'tem sido proposta como recurso de rastreio da saúde mental no esporte (TERRY et al., 2021; HAN; PARSONS-SMITH; '
  'TERRY, 2020; LEW et al., 2023).')
P('Entre as seis dimensões, o vigor e a fadiga formam o eixo mais responsivo à carga e concentram boa parte do '
  'valor prático do monitoramento, embora as demais dimensões componham o quadro completo do estado afetivo. Sob '
  'intensificação do treino, o vigor tende a cair e a fadiga a subir, um padrão observado em atletas de diferentes '
  'modalidades e sensível também à privação de sono (FERREIRA et al., 2026; PIERCE, 2002). Esse eixo responde '
  'ainda de forma aguda a cada sessão ou competição, com oscilações mensuráveis entre os momentos pré e '
  'pós-esforço, o que recomenda coletas repetidas dentro do dia, e não apenas entre dias (DO NASCIMENTO et al., '
  '2026).')
P('A leitura conjunta dessas variações remete ao continuum do sobre-esforço. A distinção entre o sobre-esforço '
  'funcional, o sobre-esforço não funcional e a síndrome de overtraining é hoje descrita como um processo gradual, '
  'de difícil diagnóstico por um único marcador, no qual a fadiga central e a piora do humor figuram entre os '
  'sinais mais precoces (ROETE et al., 2021; LA TORRE et al., 2023; MANESCU et al., 2026). Nesse quadro, o '
  'acompanhamento do humor oferece um marcador sensível, de baixo custo e não invasivo, complementar aos '
  'marcadores fisiológicos, para vigiar a janela em que o sobre-esforço funcional é buscado de forma planejada '
  '(THORPE et al., 2017).')
P('O handebol oferece um cenário exigente para essa vigilância. A modalidade é coletiva, intermitente e de alta '
  'intensidade, com sprints curtos, mudanças de direção, saltos, arremessos e contato físico, o que impõe elevada '
  'demanda neuromuscular e psicofisiológica, variável por posição de jogo (KARCHER; BUCHHEIT, 2014; GARCÍA-SÁNCHEZ '
  'et al., 2023; CARTON-LLORENTE et al., 2023). O monitoramento longitudinal da carga interna ao longo da '
  'temporada e a atenção à fadiga, inclusive a mental, já foram descritos como relevantes para o rendimento na '
  'modalidade (STRUZIK; NADOBNIK; STEPIEN-SLODKOWSKA, 2026; STAIANO et al., 2025). A acumulação de carga na '
  'semana tende, assim, a corroer o vigor e a elevar a fadiga, um padrão que, quando controlado, caracteriza o '
  'sobre-esforço funcional e antecede a recuperação planejada.')
# --- Justificativa incorporada ---
P('O monitoramento do humor reúne três atributos que o tornam atraente para a rotina do esporte de rendimento: é '
  'sensível às variações da carga, tem baixo custo e não é invasivo, o que favorece a coleta frequente sem '
  'acréscimo relevante de logística (SAW; MAIN; GASTIN, 2016; HELWIG et al., 2023). Em um calendário competitivo '
  'denso, no qual as medidas fisiológicas e endócrinas nem sempre são viáveis no dia a dia, um marcador subjetivo '
  'bem escolhido oferece à comissão técnica uma leitura rápida do estado do atleta e um apoio concreto às decisões '
  'de treino e recuperação (KELLMANN et al., 2018).')
P('A literatura, contudo, ainda descreve de forma incompleta a dinâmica do humor dentro de um único microciclo de '
  'handebol de elite. A maior parte dos estudos compara momentos isolados ou temporadas inteiras, sem capturar, ao '
  'mesmo tempo, a variação entre dias e a variação aguda entre pré e pós-treino, e sem descrever a forma temporal '
  'das trajetórias das seis dimensões (DE MIRANDA ROHLFS et al., 2024; DO NASCIMENTO et al., 2026). Essa lacuna é '
  'relevante porque é justamente no interior do microciclo que se instala o sobre-esforço funcional buscado de '
  'forma planejada, e cuja passagem para o sobre-esforço não funcional deve ser vigiada (ROETE et al., 2021; '
  'MANESCU et al., 2026).')
P('Este estudo justifica-se, portanto, por reunir, em um mesmo delineamento, a descrição completa das seis '
  'dimensões e dos perfis de humor, a quantificação do tamanho do efeito e uma análise da forma temporal das '
  'trajetórias que vai além da comparação de médias. O uso de suavização, de segundas derivadas, de ajuste '
  'polinomial e da localização dos cruzamentos exatos entre as dimensões, com destaque para o par vigor-fadiga e '
  'a PTH, converte a ideia qualitativa de inversão do eixo energia-fadiga em eventos datados, o que oferece à '
  'prática do handebol de elite um conjunto de marcadores objetivos e visuais, complementares ao monitoramento de '
  'carga e de fadiga já descrito na modalidade (STRUZIK; NADOBNIK; STEPIEN-SLODKOWSKA, 2026; STAIANO et al., '
  '2025). Poucos estudos descrevem, dentro de um único microciclo de handebol de elite, a migração dos perfis de '
  'humor e a forma temporal exata das trajetórias das dimensões, com coletas pré e pós-treino que capturam também '
  'a variação dentro do dia (DE MIRANDA ROHLFS et al., 2024; BIRD et al., 2025; RATZ-SULYOK et al., 2026). A '
  'Figura 1 resume o quadro conceitual que orienta este trabalho.')
figure('fig01.png','Quadro conceitual do monitoramento do humor: a carga do microciclo altera as seis dimensões do BRUMS, com queda do vigor e elevação da fadiga e da PTH, o que desloca os perfis de humor em direção à barbatana de tubarão, dentro da janela do sobre-esforço funcional, e retroalimenta as decisões de treino e recuperação.',w=14.5)

# ===== 2 OBJETIVOS =====
H1('OBJETIVOS','2')
H2('Objetivo geral','2.1')
P('Caracterizar as alterações do perfil de humor de atletas de handebol de elite ao longo de um microciclo '
  'pré-competitivo, com a descrição das seis dimensões do BRUMS.')
H2('Objetivos específicos','2.2')
P('São objetivos específicos deste estudo:')
LI('a) descrever a estatística descritiva, a consistência interna e os limiares de mudança das seis dimensões do BRUMS e da perturbação total do humor;')
LI('b) caracterizar os seis perfis de humor representados na amostra e quantificar a sua migração entre o primeiro e o último dia do microciclo;')
LI('c) estimar a magnitude do efeito da mudança de cada dimensão entre o primeiro e o último dia, entre os sete dias e entre os momentos pré e pós-treino;')
LI('d) modelar a forma temporal das trajetórias das dimensões por meio de suavização e da análise das segundas derivadas, com a localização dos pontos de inflexão;')
LI('e) ajustar modelos polinomiais às trajetórias e localizar os cruzamentos exatos entre as dimensões, com atenção ao vigor, à fadiga e à perturbação total do humor;')
LI('f) examinar a estrutura dimensional e as correlações entre as dimensões do humor e discutir o conjunto como um quadro de sobre-esforço funcional útil ao monitoramento aplicado.')

# ===== 3 MÉTODO =====
H1('MÉTODO','3')
H2('Caracterização da pesquisa','3.1')
P('Trata-se de uma pesquisa de abordagem quantitativa, de natureza descritiva e correlacional, com delineamento '
  'observacional e longitudinal e medidas repetidas no mesmo grupo de atletas. A rotina de treino não foi '
  'manipulada pelos pesquisadores, e os mesmos atletas foram avaliados de forma repetida ao longo de um microciclo '
  'pré-competitivo, sem grupo controle, de modo que cada atleta serve de referência para si mesmo e as comparações '
  'são feitas entre os momentos da semana.')
H2('Aspectos éticos','3.2')
P('A pesquisa foi conduzida de acordo com os princípios da Declaração de Helsinque e com as normas nacionais de '
  'pesquisa com seres humanos, e foi aprovada pelo Comitê de Ética em Pesquisa (CEP) da instituição responsável '
  '(CAAE: [inserir número de aprovação]). Antes do início das coletas, a equipe recebeu a explicação detalhada dos '
  'procedimentos e dos objetivos do estudo, e todos os atletas concordaram em participar de forma voluntária e '
  'assinaram o termo de consentimento livre e esclarecido. Os dados foram tratados de forma confidencial e '
  'anônima, e a participação não interferiu na rotina de treino da equipe.')
H2('População e amostra','3.3')
P('Participaram 27 atletas de handebol do sexo masculino, de nível de elite, com idade de 22,0 ± 3,8 anos e 10,5 '
  '± 3,8 anos de prática, integrantes do elenco principal durante a pré-temporada. A amostra reúne as quatro '
  'posições táticas da modalidade (armador, ala, pivô e goleiro), o que permite descrever também as diferenças por '
  'função de jogo. A seleção dos participantes foi por conveniência, dada a natureza de elite do grupo, e '
  'incluíram-se os atletas que responderam ao instrumento nos momentos previstos ao longo do microciclo.')
H2('Procedimentos da pesquisa','3.4')
P('O delineamento cobriu sete dias de um microciclo pré-competitivo. No primeiro dia foi realizada uma coleta de '
  'linha de base, em condição de repouso, e nos seis dias subsequentes foram feitas duas coletas diárias, uma '
  'antes e outra depois do treino, o que totalizou 286 observações. A estrutura de duas medidas diárias permite '
  'separar o efeito imediato de cada sessão, que é a variação da manhã para o fim do dia, do efeito acumulado ao '
  'longo dos dias. A Figura 2 sintetiza o delineamento e o plano de análise.')
figure('fig02.png','Organograma do delineamento do estudo: da amostra e do microciclo às coletas de linha de base e pré/pós-treino, às 286 observações do BRUMS-24 e aos blocos do plano de análise.',w=15.0)
H2('Instrumentos de coleta','3.5')
P('O humor foi avaliado pela Escala de Humor de Brunel na versão de 24 itens (BRUMS-24), que reúne seis dimensões '
  '(tensão, depressão, raiva, vigor, fadiga e confusão), com escores de 0 a 16 em cada dimensão, a partir das '
  'quais se calcula também a perturbação total do humor (PTH), obtida pela soma das dimensões negativas menos o '
  'vigor. O instrumento foi respondido de forma virtual pelos atletas, o que reduz o tempo de resposta e permite o '
  'registro imediato dos dados. A BRUMS dispõe de validação para o contexto brasileiro e apresenta propriedades '
  'psicométricas replicadas em diferentes idiomas (ROHLFS et al., 2008; TERRY; LANE; FOGARTY, 2003; ROHLFS et '
  'al., 2023).')
H2('Análise estatística','3.6')
P('Na análise, empregaram-se estatística descritiva das seis dimensões e da PTH e avaliação da consistência '
  'interna por alfa de Cronbach e por ômega de McDonald, a estrutura dimensional por análise fatorial confirmatória '
  'de seis fatores correlacionados (com a exclusão do item invariante tensao_1) e o funcionamento dos itens por um '
  'modelo de resposta gradual de Samejima. Como o teste de Shapiro-Wilk apontou desvios da normalidade em parte das '
  'dimensões, adotaram-se testes não paramétricos: o teste de Wilcoxon para a comparação entre pré e pós-treino, o '
  'teste de Friedman com o W de Kendall para a comparação entre os sete dias e a correlação de Spearman para as '
  'relações entre as dimensões, sempre com o cálculo do tamanho do efeito, classificado por magnitude (trivial, '
  'pequeno, médio ou grande no d de Cohen e no dz; trivial, pequeno, moderado ou grande no W de Kendall). A '
  'confirmação multivariada da diferença entre o primeiro e o último dia recorreu à MANOVA em escores T. Cada '
  'observação foi classificada em um dos seis perfis de humor a partir dos escores padronizados, e a distribuição '
  'dos perfis entre o primeiro e o último dia foi avaliada pelo teste do qui-quadrado, com a tendência de migração '
  'resumida por uma regressão logística ao longo dos dias. A consistência das medidas repetidas foi estimada pelo '
  'coeficiente de correlação intraclasse (ICC), do qual se derivaram os limiares de mudança.')
P('A forma temporal das trajetórias foi analisada por três abordagens complementares, aplicadas às seis dimensões '
  'e à PTH. Primeiro, as séries foram suavizadas por spline, e a segunda derivada de cada curva localizou o ponto '
  'de inflexão, definido como a raiz da segunda derivada. Segundo, cada trajetória foi modelada por um ajuste '
  'polinomial de grau três, em duas resoluções, sobre as sete médias diárias e sobre os quatorze pontos pré e '
  'pós-treino, do qual se extraíram a equação, o coeficiente de determinação (R²) e a inflexão analítica '
  '(x = -b/3a). Terceiro, sobre as curvas ajustadas foram localizados os cruzamentos exatos entre as dimensões, ou '
  'seja, os dias em que uma curva iguala e ultrapassa a outra. Para separar o sinal do ruído, calculou-se a razão '
  'sinal/ruído de cada dimensão, definida como a razão entre a amplitude do sinal filtrado e o desvio-padrão do '
  'resíduo. As rotinas computacionais foram executadas em Python 3.11. A manipulação e a tabulação dos dados usaram '
  'o pandas e o NumPy; os testes não paramétricos (Wilcoxon, Friedman) e as demais estatísticas, o SciPy; os modelos '
  'lineares mistos e a regressão logística de tendência dos perfis, o statsmodels; e a normalidade multivariada '
  '(Henze-Zirkler), o pingouin. A PERMANOVA e o teste de homogeneidade das dispersões multivariadas (PERMDISP) foram '
  'implementados especificamente para este estudo, com distância euclidiana sobre escores padronizados e permutação '
  'restrita ao estrato do atleta, de modo a respeitar o delineamento de medidas repetidas — restrição que os '
  'procedimentos de uma via dos pacotes genéricos não contemplam. As figuras foram geradas com Plotly e Matplotlib. '
  'Adotou-se o nível de significância de 5%. O fluxo de decisão que orientou a escolha de cada método, segundo as '
  'propriedades dos dados, está resumido na Figura 30.')
P('Como modelagem complementar, a trajetória de cada dimensão foi ajustada por modelos lineares mistos de '
  'crescimento, com intercepto e inclinação aleatórios por atleta, e a capacidade preditiva do perfil de humor foi '
  'avaliada por regressão logística — distinguindo a fase tardia da inicial do microciclo — com validação cruzada '
  'agrupada por atleta (para evitar vazamento entre treino e teste) e curva de aprendizado, usando a biblioteca '
  'scikit-learn.')
P('A adequação da via multivariada paramétrica foi verificada antes de sua interpretação. A normalidade '
  'multivariada das seis dimensões foi rejeitada em todos os subgrupos comparados (teste de Henze-Zirkler, '
  'p < 0,001; assimetria e curtose multivariadas de Mardia, p < 0,001), e a homogeneidade das matrizes de '
  'covariância não se sustentou (teste M de Box: dia 1 vs. dia 7, χ²(21) = 71,8, p < 0,001; pré vs. pós, '
  'χ²(21) = 38,3, p = 0,012). Como esses são justamente '
  'os dois pressupostos da MANOVA e do Hotelling T², e ambos foram violados — em coerência, ademais, com o efeito '
  'de piso e com a rota não paramétrica já adotada nas análises univariadas —, o deslocamento multivariado do '
  'perfil de humor foi testado por análise de variância multivariada permutacional (PERMANOVA; Anderson, 2001), '
  'sobre uma matriz de distâncias euclidianas calculada a partir de escores padronizados. Para respeitar o '
  'delineamento de medidas repetidas, a significância foi obtida por 9 999 permutações restritas ao atleta — '
  'esquema que constrói a hipótese nula correta para fatores intraindividuais e evita a inflação do erro por '
  'pseudorreplicação —, e o tamanho de efeito foi expresso pelo R² (fração da variância multivariada explicada '
  'pelo fator). Como a PERMANOVA é sensível a diferenças de dispersão, aplicou-se o teste de homogeneidade das '
  'dispersões multivariadas (PERMDISP; Anderson, 2006), também com permutação restrita ao atleta, para distinguir '
  'deslocamento de centroide de heterogeneidade de dispersão. A MANOVA em escores T e o Hotelling T² foram '
  'mantidos como verificações paramétricas de sensibilidade.')

# ===== 4 RESULTADOS =====
H1('RESULTADOS','4')
P('Os resultados são apresentados em três níveis de profundidade, em ordem lógica. O primeiro nível, descritivo, '
  'resume as seis dimensões e a PTH, avalia a consistência interna e os limiares de mudança, e caracteriza os '
  'perfis de humor e a sua migração. O segundo nível, exploratório, compara os dias e os momentos pré e pós-treino '
  'com o cálculo do tamanho e da magnitude do efeito, e examina as relações entre as dimensões e a estrutura de '
  'componentes principais. O terceiro nível, avançado, modela a forma temporal das trajetórias de todas as '
  'dimensões por suavização, segundas derivadas, ajuste polinomial e localização dos cruzamentos, separa o sinal '
  'do ruído em cada dimensão e decompõe a resposta em efeito agudo e de recuperação ao longo das treze coletas. '
  'Cada análise é descrita e, em seguida, interpretada quanto ao seu significado para o monitoramento do atleta.')

H2('Descrição e consistência das dimensões','4.1')
P('A Tabela 1 resume as seis dimensões do BRUMS e a PTH no conjunto das observações. O vigor apresentou a maior '
  'média entre as dimensões, ao passo que as dimensões negativas concentraram-se em valores baixos, com médias '
  'próximas do limite inferior da escala. Esse padrão indica um grupo, em geral, bem ajustado, no qual a fadiga e '
  'o vigor ocupam a maior parte da faixa de resposta e sustentam a leitura do estado do atleta.')
table('Estatística descritiva das dimensões do BRUMS e da perturbação total do humor (286 observações).',
  ['Dimensão','Média','DP','Mín.-Máx.'],
  [['Tensão','1,4','1,8','0-8'],['Depressão','1,0','2,4','0-16'],['Raiva','1,6','2,8','0-14'],
   ['Vigor','5,7','3,3','0-15'],['Fadiga','5,5','4,1','0-16'],['Confusão','0,5','1,3','0-8'],['PTH','4,4','10,1','-12-52']],
  note='DP: desvio-padrão. PTH: perturbação total do humor. Escores das dimensões variam de 0 a 16.')
P('A Figura 3 apresenta os diagramas de caixa das seis dimensões, em ordem canônica e em painéis separados, com a '
  'mediana e a média de cada uma. Como as dimensões diferem muito em amplitude, cada painel usa a escala da '
  'própria variável, o que torna a distribuição de cada uma mais legível. O vigor e a fadiga percorrem a maior '
  'parte da escala e apresentam média e mediana próximas, sinal de distribuições razoavelmente simétricas, ao '
  'passo que as dimensões negativas concentram a caixa junto ao zero e exibem a média acima da mediana, o que '
  'revela assimetria positiva e um efeito de piso.')
figure('fig03.png','Diagramas de caixa das seis dimensões do BRUMS, em painéis separados e ordem canônica; cada painel usa a escala da própria variável, com a mediana (linha sólida) e a média (linha tracejada e losango).')
P('A consistência interna foi adequada nas duas dimensões do eixo energia-fadiga (alfa de 0,68 para o vigor e '
  '0,80 para a fadiga), o que reforça a confiança na sua medida. A PTH, por reunir as seis dimensões em um único '
  'índice, associou-se de forma forte ao vigor e à fadiga (correlações de 0,67 e 0,77, respectivamente), e essas '
  'duas dimensões explicaram 70% da sua variância, o que confirma o eixo energia-fadiga como o núcleo do sinal '
  'dentro do conjunto das seis dimensões.')
P('A estrutura de seis fatores correlacionados do BRUMS, ajustada por análise fatorial confirmatória sobre os 23 '
  'itens — após a exclusão do item tensao_1, com 99,6% das respostas no piso e variância praticamente nula —, '
  'apresentou ajuste aceitável (CFI = 0,87; TLI = 0,85; RMSEA = 0,08; SRMR = 0,09), coerente com a literatura de '
  'validação da escala. A consistência interna, avaliada também pelo ômega de McDonald, foi de boa a alta na raiva, '
  'na depressão e na fadiga (ômega de 0,83 a 0,87) e moderada no vigor e na confusão (0,77 e 0,68), reduzindo-se na '
  'tensão (0,54), limitada a três itens após a exclusão. A estrutura e a fidedignidade das medidas são, portanto, '
  'sólidas, e o eixo energia-fadiga reúne, ao lado de boa consistência, a maior parte do sinal interpretável.')
table('Consistência das medidas repetidas ao longo da semana (coeficiente de correlação intraclasse, ICC).',
  ['Dimensão','ICC(2,1)','ICC(2,k)','Consistência'],
  [['Vigor','0,53','0,89','moderada'],['Fadiga','0,61','0,92','moderada'],['Tensão','0,64','0,93','moderada'],
   ['Depressão','0,68','0,94','moderada'],['Raiva','0,39','0,82','pobre'],['Confusão','0,36','0,80','pobre']],
  note='ICC(2,1) = medida isolada; ICC(2,k) = média das medidas. Consistência: < 0,50 pobre; 0,50-0,75 moderada; > 0,75 boa.')
P('A consistência das medidas repetidas ao longo da semana (Tabela 2) foi moderada a boa, com os valores mais '
  'baixos na raiva e na confusão, dimensões mais reativas de um dia para o outro, o que também recomenda cautela '
  'na leitura isolada dessas dimensões.')
P('Os limiares de mudança de todas as dimensões (Tabela 3), derivados do ICC, permitem separar a variação real do '
  'ruído de medida. No vigor, o erro-padrão de medida foi de 2,3 ponto e a mudança mínima detectável de 5,3 (90%) '
  'a 6,3 (95%) pontos; na fadiga, de 2,5 e de 5,9 a 7,0 pontos. Esses valores superam a menor mudança relevante '
  '(SWC), o que indica que, no plano individual, apenas oscilações dessa ordem podem ser lidas como mudança real, '
  'ao passo que variações menores se confundem com o ruído. As dimensões negativas, de menor amplitude, '
  'apresentaram limiares proporcionalmente menores. Por esse motivo, a interpretação deste estudo apoia-se nas '
  'tendências de grupo e não em leituras isoladas por atleta.')
# limiares todas as dimensoes
DIMS=[('Vigor','Vigor'),('Fadiga','Fadiga'),('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),
      ('Confusao','Confusão'),('TMD','PTH')]
DP={'Vigor':3.3,'Fadiga':4.1,'Tensao':1.8,'Depressao':2.4,'Raiva':2.8,'Confusao':1.3,'TMD':10.1}
rows=[]
for k,lab in DIMS:
    l=T['limiares'].get(k,{}); icc=l.get('icc',0); dp=DP[k]
    sem=dp*np.sqrt(max(0,1-icc)); mdc90=1.65*np.sqrt(2)*sem; mdc95=1.96*np.sqrt(2)*sem; swc=0.2*dp
    rows.append([lab,c2('%.1f'%sem),c2('%.1f'%mdc90),c2('%.1f'%mdc95),c2('%.1f'%swc)])
table('Erro de medida e limiares de mudança de todas as dimensões e da PTH (escala 0-16): erro-padrão de medida (SEM), mudança mínima detectável (MDC) e menor mudança relevante (SWC).',
  ['Dimensão','SEM','MDC90','MDC95','SWC'],rows,
  note='SEM = DP × raiz(1 - ICC); MDC = 1,65 (90%) ou 1,96 (95%) × raiz(2) × SEM; SWC = 0,2 × DP. Uma mudança individual só é tomada como real quando excede a MDC.')

H2('Perfis de humor e sua migração','4.2')
P('Os seis perfis de humor descritos na literatura estiveram representados na amostra. A Figura 4 apresenta cada '
  'perfil identificado no estudo em escores T, com a respectiva prevalência, o que permite reconhecer a sua forma '
  'característica. O perfil iceberg exprime prontidão, com o vigor acima da média e as dimensões negativas abaixo, '
  'ao passo que a barbatana de tubarão sinaliza pico isolado de fadiga, o submerso reúne todas as dimensões abaixo '
  'da média e o Everest invertido eleva todas as dimensões negativas (PARSONS-SMITH; TERRY; MACHIN, 2017).')
figure('fig04.png','Perfis de humor identificados na amostra, em escores T (M = 50; DP = 10) nas seis dimensões; prevalência de cada perfil entre parênteses.')
P('A comparação entre o primeiro e o último dia do microciclo revelou a reconfiguração do perfil médio do grupo. '
  'A Figura 5 mostra que, no primeiro dia, o perfil assumiu o formato iceberg, com o vigor no topo e as dimensões '
  'negativas abaixo da média populacional. No último dia, o perfil inverteu-se para a forma de barbatana de '
  'tubarão, com a fadiga no topo e o vigor rebaixado, o que traduz a acumulação da carga ao longo da semana.')
figure('fig05.png','Perfil de humor em escores T no primeiro e no último dia do microciclo.',w=13.5)
P('Essa mudança de forma correspondeu a um deslocamento da prevalência dos perfis (Tabela 4). O perfil iceberg '
  'caiu de 48% no primeiro dia para 22% no último, enquanto a barbatana de tubarão subiu de 4% para 22% e o perfil '
  'submerso passou de 4% para 14%. A reorganização categórica não alcançou significância no teste do qui-quadrado '
  '(qui-quadrado = 8,96; p = 0,111), resultado esperado pela baixa contagem por célula quando poucas observações '
  'se distribuem por seis perfis. A regressão logística de tendência, porém, confirmou o aumento da chance do '
  'perfil de barbatana de tubarão a cada dia (OR = 1,37), sem que os perfis de maior risco à saúde mental se '
  'instalassem.')
table('Distribuição dos seis perfis de humor no primeiro e no último dia do microciclo: n (%).',
  ['Perfil','Dia 1, n (%)','Dia 7, n (%)'],
  [['Iceberg','13 (48,1%)','8 (21,6%)'],['Everest invertido','4 (14,8%)','4 (10,8%)'],
   ['Iceberg invertido','2 (7,4%)','2 (5,4%)'],['Submerso','1 (3,7%)','5 (13,5%)'],
   ['Barbatana de tubarão','1 (3,7%)','8 (21,6%)'],['Superfície','6 (22,2%)','10 (27,0%)']],
  note='qui-quadrado = 8,96; p = 0,111. A migração é descritiva e converge com a queda do vigor e a elevação da fadiga.')

H2('Diferença entre o primeiro e o último dia (com tamanho de efeito)','4.3')
P('A comparação entre o primeiro e o último dia do microciclo quantificou a magnitude da mudança em cada dimensão '
  '(Tabela 5). O efeito foi grande no vigor, que caiu 41% (dz = -0,95), e médio na fadiga, que subiu 75% (dz = 0,72), '
  'ao passo que as demais dimensões apresentaram efeitos de menor magnitude. A análise multivariada confirmou a '
  'diferença global entre os dois dias (Wilks lambda = 0,181; p < 0,001), e, sob a correção de Bonferroni, apenas '
  'o vigor e a fadiga permaneceram significativos, o que concentra o efeito no eixo energia-fadiga sem prescindir '
  'da leitura das demais dimensões.')
table('Diferença das dimensões do BRUMS e da PTH entre o primeiro e o último dia do microciclo (com tamanho e magnitude do efeito).',
  ['Dimensão','Dia 1 (M)','Dia 7 (M)','Variação (%)','p','dz','Magnitude'],
  [['Vigor','7,64','4,49','-41%','< 0,001','-0,95','grande'],['Fadiga','4,26','7,46','+75%','0,004','+0,72','médio'],
   ['Tensão','2,14','0,94','-56%','0,011','-0,59','médio'],['Depressão','1,02','1,27','+24%','0,752','+0,10','trivial'],
   ['Raiva','2,48','2,59','+5%','0,802','+0,02','trivial'],['Confusão','1,10','0,51','-53%','0,068','-0,46','pequeno'],
   ['PTH','3,36','8,28','+147%','0,104','+0,41','pequeno']],fs=8.6,
  note='dz = tamanho de efeito intraindividual; magnitude: trivial (< 0,2); pequeno (0,2-0,5); médio (0,5-0,8); grande (> 0,8). PTH: perturbação total do humor.')
P('A análise multivariada permutacional, livre do pressuposto de normalidade, corroborou o deslocamento do perfil '
  'de humor (Tabela 6). O perfil das seis dimensões diferiu entre o pré e o pós-treino (PERMANOVA: pseudo-F = 2,52; '
  'R² = 0,010; p = 0,002). O tamanho de efeito, embora pequeno em fração de variância (1,0%), é direcional e '
  'coerente com a leitura de sinal/ruído: a mudança é real, porém ocupa parcela reduzida de uma variância total '
  'dominada pelo efeito de piso das dimensões negativas. O teste PERMDISP indicou que as dispersões multivariadas '
  'são homogêneas (F = 0,80; p = 0,342), de modo que a diferença reflete deslocamento do centroide, e não '
  'heterogeneidade de variância. A conclusão mostrou-se robusta à inclusão das coletas intermediárias (fator momento com três '
  'níveis: pseudo-F = 1,75; p = 0,002) e convergiu com o eixo energia-fadiga do Hotelling T² (D de Mahalanobis = '
  '0,66; p = 0,010); o Hotelling nas seis dimensões foi limítrofe (p = 0,054), o que decorre do menor poder do '
  'teste paramétrico sob não normalidade, e não de ausência de efeito.')
table('Análise multivariada permutacional (PERMANOVA) do perfil de humor e teste de homogeneidade das dispersões (PERMDISP).',
  ['Fator','pseudo-F','R²','p','PERMDISP (p)','Interpretação'],
  [['Momento (pré vs. pós)','2,52','0,010','0,002','0,342','deslocamento de centroide (dispersões homogêneas)'],
   ['Momento (pré/intermediário/pós)†','1,75','0,008','0,002','0,024','robustez à inclusão das coletas intermediárias']],fs=8.6,
  note='Distância euclidiana sobre escores padronizados; 9 999 permutações restritas ao atleta. R² = fração da variância multivariada explicada pelo fator. †Análise de robustez.')

H2('Comparação das dimensões entre os dias (Friedman)','4.4')
P('A comparação das dimensões entre os sete dias pelo teste de Friedman (Tabela 6) apontou variação significativa '
  'no vigor (p < 0,001; W = 0,21) e na fadiga (p = 0,007; W = 0,16), com magnitude pequena a moderada, além de '
  'diferenças na tensão e na confusão. O vigor foi máximo no Dia 1 e mínimo no Dia 7, enquanto a fadiga percorreu '
  'o caminho inverso, com pico no Dia 7, um padrão coerente com o acúmulo de carga dentro da faixa funcional.')
table('Médias diárias das dimensões do BRUMS e teste de Friedman com W de Kendall e magnitude do efeito (comparação entre os sete dias).',
  ['Dim.','D1','D2','D3','D4','D5','D6','D7','χ²','p','W','Magn.'],
  [['Tensão','2,2','1,6','1,1','1,4','1,0','1,5','0,9','13,2','0,039','0,12','pequeno'],
   ['Depressão','1,0','1,2','0,7','1,1','0,7','1,1','1,3','1,8','0,936','0,02','trivial'],
   ['Raiva','2,0','1,7','1,4','1,4','0,6','1,7','2,6','12,3','0,056','0,11','pequeno'],
   ['Vigor','7,6','5,7','5,7','5,3','5,6','5,7','4,5','14,7','0,022','0,13','pequeno'],
   ['Fadiga','4,0','5,2','5,0','5,8','5,3','5,8','7,5','13,2','0,040','0,12','pequeno'],
   ['Confusão','1,0','0,5','0,3','0,4','0,2','0,6','0,5','25,8','< 0,001','0,23','pequeno'],
   ['PTH','2,5','4,6','2,9','4,8','2,2','4,8','8,3','7,6','0,269','0,07','trivial']],fs=8.0,
  note='W de Kendall = tamanho de efeito do teste de Friedman; magnitude: trivial (< 0,1); pequeno (0,1-0,3); moderado (0,3-0,5); grande (> 0,5).')
P('A Figura 6 ilustra a trajetória do vigor, da fadiga e da PTH ao longo da semana, com a queda progressiva do '
  'vigor e a elevação da fadiga e da PTH em direção ao fim do microciclo.')
figure('fig06.png','Trajetória de vigor, fadiga e perturbação total do humor ao longo dos sete dias (curvas suavizadas; ponto de inflexão e extremos sinalizados).')
P('A Figura 7 detalha, por meio de diagramas de caixa por dia, a distribuição de cada dimensão ao longo da '
  'semana. O deslocamento das caixas do vigor para baixo e das caixas da fadiga para cima, sobretudo na segunda '
  'metade do microciclo, torna visível a mesma deterioração do eixo energia-fadiga, ao passo que as dimensões '
  'negativas de menor expressão mantêm caixas comprimidas junto ao zero.')
figure('fig07.png','Diagramas de caixa das seis dimensões do BRUMS por dia do microciclo (áreas sombreadas: início e acúmulo da semana).')
P('A comparação direta da PTH entre três momentos do microciclo (Figura 8) ilustra a mesma tendência: a '
  'perturbação total do humor aumentou do primeiro para o último dia, e a diferença foi significativa entre o Dia '
  '1 e o Dia 7 (p = 0,008), ao passo que as comparações entre dias adjacentes não alcançaram significância.')
figure('fig08.png','Perturbação total do humor (PTH) no Dia 1, no Dia 4 e no Dia 7, com as comparações par a par (teste de Mann-Whitney; * p < 0,05; ** p < 0,01; ns = não significativo).',w=13.5)
P('A distribuição completa da PTH em cada dia (Figura 9) e a sua distribuição acumulada (Figura 10) confirmam o '
  'deslocamento ao longo da semana. Os diagramas de violino mostram a densidade da PTH a subir do Dia 1 ao Dia 7, '
  'e as curvas de distribuição acumulada revelam um deslocamento sistemático para a direita, com o Dia 7 situado '
  'acima do Dia 1 em praticamente toda a faixa de escores.')
figure('fig09.png','Distribuição da PTH por dia (diagramas de violino com caixa e média; a escala foi ajustada à faixa central da distribuição para facilitar a leitura).')
figure('fig10.png','Distribuição acumulada (ECDF) da PTH no Dia 1, no Dia 4 e no Dia 7; o deslocamento das curvas para a direita indica o aumento da perturbação ao longo do microciclo.')

H2('Suavização, derivadas e cruzamento das trajetórias das dimensões','4.5')
P('Para separar o sinal do ruído, as trajetórias das dimensões foram suavizadas sobre os pontos pré e pós-treino '
  'por meio de uma spline, e a segunda derivada de cada curva localizou o seu ponto de inflexão, ou seja, o '
  'momento em que a concavidade muda e a taxa de variação atinge o seu limite (Figura 11; Tabela 7). Depois de '
  'removido o ruído, o vigor, a fadiga e a PTH exibiram um único ponto de inflexão, situado na metade da semana '
  '(em torno do Dia 3,8), o que marca a transição entre a fase inicial de prontidão e a fase de acúmulo de carga. '
  'A curva suavizada do vigor subiu até um máximo em torno do Dia 5,0 e caiu ao seu mínimo no fim da semana, '
  'enquanto a fadiga e a PTH percorreram o caminho inverso e atingiram os seus máximos no último dia.')
figure('fig11.png','Trajetórias suavizadas do vigor, da fadiga e da PTH sobre os doze pontos pré e pós-treino. Marcadores translúcidos: sinal bruto; linha grossa: sinal suavizado; linha pontilhada: ponto de inflexão (segunda derivada nula); triângulos: máximo e mínimo.')
table('Ponto de inflexão (segunda derivada nula) e extremos das trajetórias suavizadas do eixo energia-fadiga.',
  ['Dimensão','Inflexão (dia)','Máximo: dia (escore)','Mínimo: dia (escore)'],
  [['Vigor','Dia 3,8','Dia 5,0 (5,9)','Dia 7,2 (4,2)'],['Fadiga','Dia 3,8','Dia 7,2 (7,7)','Dia 1,8 (4,5)'],
   ['PTH','Dia 3,5','Dia 7,2 (9,9)','Dia 4,7 (3,2)']],
  note='Curvas suavizadas por spline sobre os 12 pontos pré/pós; inflexão = raiz da segunda derivada.')
P('A mesma inflexão foi examinada por dois caminhos complementares. O painel da segunda derivada localiza o '
  'cruzamento por zero de cada curva e realça os extremos com o respectivo valor (Figura 12). Em paralelo, um '
  'ajuste polinomial de grau três resume cada trajetória por uma equação e reproduz a inflexão de forma analítica, '
  'na raiz da segunda derivada (x = -b/3a), em duas resoluções (Figura 13). Sobre as sete médias diárias, o ajuste '
  'foi forte, com R² de 0,96 para o vigor, 0,95 para a fadiga e 0,80 para a PTH, e inflexões nos dias 4,2, 4,0 e '
  '3,9. Sobre os quatorze pontos pré e pós-treino, incluída a linha de base do primeiro dia, o R² foi de 0,80, '
  '0,75 e 0,60, com inflexões praticamente idênticas. As duas resoluções convergem para a transição na metade da '
  'semana e reforçam, por via paramétrica, o achado da spline.')
figure('fig12.png','Análise por derivadas do eixo energia-fadiga. Em cada variável, o painel superior traz a curva suavizada com o máximo e o mínimo realçados, e o painel inferior traz a segunda derivada (concavidade) com o cruzamento por zero destacado, que marca o ponto de inflexão da trajetória.')
figure('fig13.png','Ajuste polinomial de grau três das trajetórias do vigor, da fadiga e da PTH em duas resoluções: sobre as sete médias diárias (marcadores quadrados e linha cheia) e sobre os quatorze pontos pré e pós-treino (pontos translúcidos e linha tracejada). Cada painel traz a equação do ajuste diário, o R² de cada resolução e a inflexão analítica.')
P('Para atender à leitura conjunta de todas as dimensões, a Figura 14 apresenta as trajetórias suavizadas das seis '
  'dimensões em uma mesma escala. A leitura evidencia que apenas o par vigor-fadiga ocupa a faixa central da '
  'escala e se cruza ao longo da semana, ao passo que a tensão, a depressão, a raiva e a confusão permanecem junto '
  'ao piso, com pequenas oscilações. Esse contraste mostra, de forma direta, por que o sinal útil do monitoramento '
  'se concentra no eixo energia-fadiga, sem que as demais dimensões deixem de ser descritas.')
figure('cruz_todas.png','Trajetórias suavizadas das seis dimensões do BRUMS ao longo do microciclo, na mesma escala; os losangos assinalam os cruzamentos entre o vigor e a fadiga (dias 2,9, 3,6 e 5,7).',w=16.0)
P('Como o vigor e a fadiga compartilham a mesma escala, as suas trajetórias podem ser comparadas de forma direta '
  'e os seus cruzamentos exatos podem ser localizados (Figura 15). No primeiro dia, o vigor superou a fadiga por '
  'ampla margem (média de 7,6 contra 4,0 pontos). O vigor caiu e a fadiga subiu até que as duas curvas se '
  'igualaram pela primeira vez em torno do dia 2,5 (escore 5,3). A partir daí, o vigor e a fadiga percorreram uma '
  'faixa estreita e próxima e voltaram a cruzar-se nos dias 4,0 e 6,0, sinal de um equilíbrio instável entre '
  'energia e fadiga durante a maior parte da semana. Após o último cruzamento, no dia 6,0 (escore 5,8), a fadiga '
  'afastou-se de forma definitiva acima do vigor. No mesmo trecho final, a perturbação total do humor, que vinha '
  'abaixo das duas dimensões, subiu e ultrapassou primeiro o vigor (dia 6,4; escore 5,5) e depois a fadiga (dia '
  '6,6; escore 6,6), o que marca a deterioração conjunta do estado de humor no fim do microciclo.')
figure('fig14.png','Cruzamentos exatos das trajetórias de vigor, fadiga e PTH ao longo do microciclo (curvas do ajuste polinomial sobre as médias diárias; losangos: pontos de cruzamento, com o dia e o escore).')
P('Para examinar o cruzamento decisivo do eixo energia-fadiga com maior detalhe, a Figura 16 amplia a região do '
  'último cruzamento e representa a diferença suavizada entre o vigor e a fadiga com a área sombreada de acordo com '
  'o sinal: a área é positiva enquanto o vigor domina e negativa quando a fadiga passa a dominar. Sobre as treze '
  'coletas, o vigor e a fadiga cruzaram-se três vezes (dias 2,9, 3,6 e 5,7), o que confirma um longo trecho de '
  'quase equilíbrio no miolo da semana. O cruzamento do dia 5,8 é o definitivo: a partir dele a diferença '
  'torna-se cada vez mais negativa e alcança cerca de três pontos de vantagem da fadiga ao fim do microciclo, sem '
  'retorno.')
figure('fig15.png','Cruzamento vigor por fadiga. Painel superior: curvas suavizadas com a área sombreada entre elas e os três cruzamentos assinalados. Painel inferior: recorte ampliado da diferença suavizada (vigor - fadiga), com a área sombreada segundo o sinal e o cruzamento definitivo destacado.')
P('A análise por suavização e derivadas foi então estendida às seis dimensões e à PTH (Figura 17; Tabela 8). A '
  'separação entre sinal e ruído variou muito entre as dimensões. O vigor e a fadiga apresentaram a maior razão '
  'sinal/ruído (S/R de 8,4 e 6,7), com trajetórias limpas e um único ponto de inflexão na metade da semana. Já as '
  'dimensões negativas exibiram sinal fraco e forte ruído (S/R entre 2,7 e 4,4), com trajetórias irregulares e de '
  'pequena amplitude, marcadas por um forte efeito de piso.')
figure('fig16.png','Suavização e derivadas das seis dimensões do BRUMS e da PTH, sobre as treze coletas. Marcadores: sinal bruto (ruído); linha grossa: sinal suavizado; linha pontilhada: inflexão (segunda derivada nula); triângulos: máximo e mínimo. S/R: razão sinal/ruído (amplitude); piso: percentual de escores nulos.')
table('Separação entre sinal e ruído, efeito de piso e ponto de inflexão de cada dimensão (curvas suavizadas sobre as 13 coletas).',
  ['Dimensão','Amplitude do sinal','Ruído (DP)','S/R','Piso (% de zeros)','Inflexão (1º dia)'],
  [['Vigor','4,1','0,49','8,4','9%','Dia 3,8'],['Fadiga','4,5','0,67','6,7','10%','Dia 4,0'],
   ['PTH','9,2','1,96','4,7','43%','Dia 3,5'],['Tensão','1,1','0,28','3,7','46%','Dia 4,8'],
   ['Depressão','0,9','0,33','2,7','67%','Dia 3,2'],['Raiva','1,8','0,54','3,4','62%','Dia 3,0'],
   ['Confusão','0,8','0,18','4,4','79%','Dia 5,4']],fs=8.6,
  note='Amplitude do sinal = máximo - mínimo da curva suavizada; ruído = desvio-padrão dos resíduos; S/R = razão sinal/ruído; piso = percentual de observações com escore zero.')
P('A leitura das trajetórias como um sistema em movimento é sintetizada a seguir. Sobre a curva suavizada — que '
  'atua como um filtro que retém a tendência lenta (o sinal) e descarta a oscilação de curto prazo (o ruído) —, a '
  'primeira derivada mede a velocidade de desgaste e a segunda, a sua aceleração; para a fadiga e a PTH, o ponto de '
  'inflexão em que a segunda derivada se anula, em torno do dia 4, marca a passagem da desaceleração para a '
  'aceleração da fadiga (Figura da leitura cinemática). Estendida a todas as dimensões, a análise só se sustenta '
  'onde a razão sinal/ruído é suficiente: no eixo energia-fadiga o ajuste é confiável (R² ajustado de 0,88 no vigor '
  'e na fadiga e de 0,60 na PTH), e é apenas moderado na confusão (0,70); já nas dimensões negativas comprimidas '
  'pelo piso o polinômio ajusta sobretudo ruído — no caso da depressão, com R² ajustado negativo —, o que '
  'desautoriza a leitura de derivadas nessas dimensões e delimita, de forma honesta, o alcance do método ao par '
  'energia-fadiga.')
figure('deriv_cinematica.png','Leitura cinemática do eixo energia-fadiga. Para a fadiga e a PTH, à esquerda a trajetória das médias diárias com o ajuste cúbico e o ponto de inflexão; à direita, a velocidade (primeira derivada) e a aceleração (segunda derivada), com o cruzamento por zero da aceleração (segunda derivada nula) na inflexão de meia-semana.',w=15.0)
figure('deriv_dimensoes.png','Trajetória e derivadas das seis dimensões do BRUMS e da PTH ao longo dos sete dias, com o R² ajustado do ajuste cúbico de cada dimensão. A inflexão é assinalada apenas quando o ajuste é confiável (R² ajustado maior ou igual a 0,5); as dimensões negativas travadas no efeito de piso não apresentam trajetória modelável e são apresentadas em cinza.',w=16.0)

H2('Dias de pico e comparação de cada dia ao primeiro','4.6')
P('A localização dos picos sintetiza a dinâmica semanal (Tabela 9). O vigor foi máximo no Dia 1 e mínimo no Dia '
  '7; a fadiga foi máxima no Dia 7; a tensão, no Dia 1; a raiva, no Dia 7; e a depressão, no Dia 7. Em síntese, o '
  'início da semana reúne maior prontidão, com vigor e tensão mais altos, ao passo que o final concentra a fadiga '
  'e a raiva.')
table('Dia de maior e de menor expressão de cada dimensão do humor (médias diárias).',
  ['Dimensão','Dia de maior valor','Valor','Dia de menor valor','Valor'],
  [['Tensão','Dia 1','2,15','Dia 7','1,05'],['Depressão','Dia 7','1,35','Dia 1','0,59'],
   ['Raiva','Dia 7','2,54','Dia 5','0,57'],['Vigor','Dia 1','8,30','Dia 7','4,43'],
   ['Fadiga','Dia 7','7,22','Dia 1','3,37'],['Confusão','Dia 1','1,04','Dia 5','0,21']])
P('No pós-teste que compara todos os dias entre si (Tabela 10; Figura 18), o vigor diferiu de forma significativa '
  'em 9 dos 21 pares de dias e a fadiga em 6, sempre no sentido de piora em relação aos primeiros dias, o que '
  'confirma a deterioração progressiva do eixo energia-fadiga ao longo do microciclo.')
table('Médias marginais estimadas por dia do vigor e da fadiga, com comparação de cada dia ao Dia 1 (pós-teste de Tukey).',
  ['Dia','Vigor','Fadiga'],
  [['Dia 1','8,30','3,37'],['Dia 2','5,45*','5,18'],['Dia 3','5,65*','4,84'],['Dia 4','5,31*','5,70*'],
   ['Dia 5','5,80*','5,01'],['Dia 6','5,63*','5,72*'],['Dia 7','4,10*','6,93*']],
  note='* diferença significativa em relação ao Dia 1 (Tukey, p < 0,05).')
figure('fig17.png','Médias marginais diárias do vigor e da fadiga, com comparação de todos os dias ao Dia 1 (* p < 0,05).')

H2('Variação entre pré e pós-treino e efeito agudo/recuperação','4.7')
P('A comparação entre pré e pós-treino pelo teste de Wilcoxon (Tabela 11) evidenciou uma resposta aguda coerente '
  'com o esforço: o vigor caiu (d = -0,47) e a fadiga e a PTH subiram do momento pré para o pós, com efeito de '
  'magnitude pequena a moderada. Essa oscilação dentro do dia somou-se à tendência semanal e ajudou a compor o '
  'quadro de fadiga funcional observado no fim do microciclo. A Figura 19 apresenta esses escores por dia e mostra '
  'que a diferença entre pré e pós-treino se mantém ao longo da semana.')
table('Comparação entre pré e pós-treino das dimensões do BRUMS e da PTH (teste de Wilcoxon e d de Cohen).',
  ['Dimensão','Pré (M)','Pós (M)','Variação (%)','p','d','Magnitude'],
  [['Tensão','1,30','1,60','+23%','0,057','+0,43','pequeno'],['Depressão','1,03','1,21','+17%','0,687','+0,24','pequeno'],
   ['Raiva','1,31','1,51','+15%','0,745','+0,13','trivial'],['Vigor','5,76','4,94','-14%','0,025','-0,47','pequeno'],
   ['Fadiga','4,61','6,28','+36%','0,005','+0,57','médio'],['Confusão','0,41','0,41','-1%','0,706','-0,01','trivial'],
   ['PTH','2,91','6,06','+108%','0,009','+0,58','médio']],fs=8.6,
  note='p do teste de Wilcoxon; d = tamanho de efeito de Cohen. PTH: perturbação total do humor.')
figure('fig18.png','Escores de vigor, fadiga e perturbação total do humor no pré e no pós-treino, por dia.')
P('Para descrever a resposta aguda com maior rigor, cada uma das treze coletas foi comparada à coleta '
  'imediatamente anterior por meio do teste de Wilcoxon pareado, o que separa o efeito agudo do treino (transição '
  'pré para pós, dentro do dia) do efeito de recuperação entre sessões (transição pós para pré do dia seguinte). A '
  'Figura 20 representa esse percurso em forma de dente de serra e a Tabela 12 reúne o tamanho de efeito de cada '
  'transição no eixo energia-fadiga.')
figure('fig19.png','Efeito agudo (pré para pós, segmentos laranja) e de recuperação (pós para pré do dia seguinte, segmentos azuis) da PTH ao longo das treze coletas, com o vigor e a fadiga como contexto. Os rótulos indicam o tamanho de efeito das transições significativas da PTH (* p < 0,05, Wilcoxon pareado).')
table('Transições sequenciais entre coletas: efeito agudo (pré para pós) e de recuperação (pós para pré) no eixo energia-fadiga (tamanho de efeito dz; teste de Wilcoxon pareado).',
  ['Transição','Tipo','Vigor (dz)','Fadiga (dz)','PTH (dz)'],
  [['D1 base → D2 pré','Recuperação','-0,69*','+0,20','+0,49*'],['D2 pré → D2 pós','Agudo','-0,39','+0,64*','+0,31'],
   ['D2 pós → D3 pré','Recuperação','+0,49*','-0,51*','-0,56*'],['D3 pré → D3 pós','Agudo','-0,22','+0,28','+0,16'],
   ['D3 pós → D4 pré','Recuperação','+0,19','-0,03','-0,20'],['D4 pré → D4 pós','Agudo','-0,26','+0,23','+0,50'],
   ['D4 pós → D5 pré','Recuperação','+0,43','-0,55','-0,63*'],['D5 pré → D5 pós','Agudo','-0,12','+0,38','+0,17'],
   ['D5 pós → D6 pré','Recuperação','+0,36','-0,22','-0,23'],['D6 pré → D6 pós','Agudo','-0,64*','+0,36','+0,75*'],
   ['D6 pós → D7 pré','Recuperação','-0,03','-0,07','-0,17'],['D7 pré → D7 pós','Agudo','-0,20','+0,48','+0,53']],fs=8.4,
  note='dz = tamanho de efeito intraindividual; * transição significativa (p < 0,05). Agudo = pré para pós no mesmo dia; Recuperação = pós para pré do dia seguinte. PTH: perturbação total do humor.')
P('O padrão foi consistente com o de um sistema que oscila sob carga e se recupera de forma parcial. As transições '
  'agudas tenderam a elevar a fadiga e a PTH e a reduzir o vigor, com destaque para a sessão do Dia 6, na qual a '
  'PTH subiu com efeito grande (dz = 0,75) e o vigor caiu (dz = -0,64). As transições de recuperação, entre o pós '
  'de um dia e o pré do seguinte, moveram-se no sentido inverso e devolveram parte do vigor e reduziram a fadiga e '
  'a PTH, com o episódio mais nítido entre o Dia 2 e o Dia 3 (PTH com dz = -0,56). Ainda assim, a recuperação '
  'noturna não anulou por completo o efeito agudo acumulado, o que explica a deriva descendente do vigor e a '
  'subida líquida da fadiga ao longo da semana.')

H2('Relações entre as dimensões (Spearman)','4.8')
P('As correlações de Spearman entre as dimensões (Tabela 13) mostraram que as dimensões negativas se associam '
  'entre si, com destaque para os pares depressão-raiva e tensão-confusão, e que a fadiga se relaciona com a '
  'depressão e a raiva. O vigor manteve-se relativamente independente das demais dimensões, o que reforça a sua '
  'leitura como um polo próprio do eixo energia-fadiga.')
table('Correlações de Spearman significativas entre as dimensões do BRUMS (n = 27 atletas).',
  ['Par de dimensões','rho','p'],
  [['Vigor × Fadiga','-0,40','0,037'],['Fadiga × Depressão','+0,42','0,028'],['Fadiga × Raiva','+0,49','0,010'],
   ['Tensão × Confusão','+0,45','0,017'],['Depressão × Raiva','+0,53','0,005'],['Depressão × Confusão','+0,43','0,027']],
  note='rho = coeficiente de correlação de Spearman. Apresentam-se apenas os pares com p < 0,05.')
P('A relação entre o vigor e a fadiga, além de negativa, tornou-se mais estreita à medida que a carga se acumulou '
  '(Figura 21): a regressão por fase do microciclo mostrou um acoplamento fraco na fase inicial (r = -0,35) e mais '
  'forte na fase de acúmulo (r = -0,51), o que sugere que os dois polos do eixo energia-fadiga passam a variar de '
  'forma mais solidária sob fadiga acumulada.')
figure('fig20.png','Relação entre vigor e fadiga por fase do microciclo, com a reta de regressão de cada grupo (início: dias 1 a 4; acúmulo: dias 5 a 7).',w=13.5)

H2('Estrutura dimensional (análise de componentes principais)','4.9')
P('Uma análise exploratória de componentes principais resumiu a estrutura das seis dimensões. Os dois primeiros '
  'componentes explicaram 64% da variância (41% no primeiro e 23% no segundo). O círculo de correlação (Figura 22) '
  'mostra que as dimensões negativas (depressão, raiva, confusão e fadiga) se projetam juntas no lado positivo do '
  'primeiro componente, que funciona como um eixo geral de perturbação, ao passo que o vigor aponta em sentido '
  'oposto. Essa oposição reforça, por via independente, a centralidade do eixo energia-fadiga dentro da estrutura '
  'das seis dimensões.')
figure('fig21.png','Círculo de correlação da análise de componentes principais das seis dimensões do BRUMS. A espessura de cada seta é proporcional à contribuição da variável ao plano dos dois primeiros componentes.',w=13.0)

H2('Funcionamento dos itens (teoria de resposta ao item)','4.10')
P('O modelo de resposta gradual de Samejima detalhou o funcionamento de cada item ao longo do contínuo do traço '
  '(Tabela 15; Figura 25). Os itens do eixo energia-fadiga foram os mais discriminativos (parâmetro de discriminação '
  'a entre 1,3 e 2,2 na fadiga e no vigor), ao passo que alguns itens travados no piso — os relativos a sonolência, '
  'a estado de alerta e à tensão extrema — exibiram discriminação próxima de zero e limiares fora da faixa '
  'observável, sinal de degenerescência amostral sob o efeito de piso, e não de falha do item em populações sem '
  'esse teto de resposta. A função de informação do teste confirma essa leitura: a escala de vigor informa numa '
  'faixa ampla do traço, com pico próximo à média, ao passo que a de fadiga informa melhor em níveis acima da '
  'média — ou seja, discrimina com mais precisão justamente quando a fadiga se eleva, que é a condição de interesse '
  'no monitoramento da carga.')
table('Discriminação (parâmetro a) dos itens do eixo energia-fadiga no modelo de resposta gradual.',
  ['Item','Dimensão','Discriminação (a)','Leitura'],
  [['Animado','Vigor','0,94','adequado'],['Com disposição','Vigor','2,15','alto'],['Com energia','Vigor','1,74','alto'],
   ['Alerta','Vigor','0,11','degenerado (piso)'],['Esgotado','Fadiga','2,18','alto'],['Exausto','Fadiga','1,85','alto'],
   ['Cansado','Fadiga','1,35','adequado'],['Sonolento','Fadiga','0,25','degenerado (piso)']],fs=8.6,
  note='a = parâmetro de discriminação (métrica logística); itens degenerados apresentam poucas categorias endossadas e limiares fora da faixa observável, efeito da concentração de respostas no piso.')
figure('tri_informacao.png','Função de informação do teste (modelo de resposta gradual de Samejima) para as escalas de vigor e de fadiga, calculada sobre os itens estáveis. O eixo horizontal é o traço latente em desvios-padrão; a linha tracejada assinala o ponto de máxima informação.',w=15.0)

H2('Modelos de crescimento e capacidade preditiva','4.11')
P('A dinâmica descrita pelas médias e derivadas foi formalizada por modelos mistos de crescimento, com intercepto e '
  'inclinação aleatórios por atleta (Figura 26). O vigor decresceu de forma significativa ao longo da semana (efeito '
  'do dia, p = 0,003), a fadiga acumulou-se de modo aproximadamente linear (cerca de 0,3 ponto por dia) e a '
  'perturbação total do humor descreveu uma curva em U — coerentes, no nível populacional, com as trajetórias '
  'suavizadas, agora estimadas respeitando a variação entre atletas.')
P('A capacidade de o perfil de humor prever a fase do microciclo foi avaliada por regressão logística (fase tardia, '
  'dias 5 a 7, versus inicial, dias 1 a 3), com validação cruzada agrupada por atleta para evitar vazamento entre '
  'treino e teste. A discriminação foi apenas modesta (AUC = 0,58; Figura 27), com a fadiga como principal preditor. '
  'A curva de aprendizado (Figura 28) manteve-se plana e próxima do acaso, com pequena distância entre treino e '
  'validação: o desempenho limitado não decorre de sobreajuste nem de falta de dados, e sim de um sinal individual '
  'intrinsecamente baixo, pois o efeito de piso comprime a informação disponível em uma coleta isolada. Esse '
  'resultado é, ele próprio, um achado: a leitura do humor é robusta como tendência de grupo, mas não sustenta, '
  'isoladamente, a predição no nível do atleta — o que reforça a recomendação de interpretar o indivíduo apenas '
  'acima da mudança mínima detectável.')
P('Como verificação de robustez, a mesma tarefa e o mesmo esquema de validação agrupado por atleta foram '
  'reaplicados a dois algoritmos flexíveis de gradient boosting — XGBoost e LightGBM —, capazes de captar '
  'interações e não linearidades. Ambos ficaram no nível do acaso (AUC = 0,50 nos dois casos) e, portanto, abaixo '
  'do teto obtido pela regressão logística (AUC = 0,58; Figura 29). O fato de modelos mais expressivos não '
  'superarem — e sequer igualarem — um modelo linear simples indica que o limite de predição não decorre da forma '
  'funcional escolhida, mas da ausência de estrutura individual explorável no perfil de humor, em convergência com '
  'a curva de aprendizado plana e com o efeito de piso.')
figure('curvas_crescimento.png','Curvas de crescimento do humor (modelos mistos): trajetórias individuais (linhas finas) e curva populacional (linha grossa) para o vigor, a fadiga e a perturbação total do humor.',w=16.0)
figure('roc_preditivo.png','Curva ROC do modelo preditivo (regressão logística) para a fase do microciclo (tardia versus inicial) a partir do perfil de humor, com validação cruzada agrupada por atleta.',w=10.5)
figure('curva_aprendizado.png','Curva de aprendizado do modelo preditivo: AUC de treino e de validação cruzada em função do tamanho do conjunto de treino.',w=13.0)
figure('comparacao_modelos_auc.png','Comparação da capacidade preditiva (AUC, validação cruzada por atleta) entre a regressão logística e os modelos de gradient boosting (XGBoost e LightGBM) na classificação da fase do microciclo; a linha tracejada indica o nível do acaso.',w=11.5)
H2('Síntese dos principais achados','4.12')
P('Tomados em conjunto, os resultados desenham uma história coerente do microciclo pré-competitivo. No plano das '
  'medidas, o BRUMS mostrou estrutura de seis fatores com ajuste aceitável e boa confiabilidade no eixo '
  'energia-fadiga, onde também se concentra a maior informação psicométrica; as dimensões negativas, comprimidas '
  'pelo efeito de piso, funcionam como sentinelas de estados raros, e não como termômetros graduais. No plano do '
  'comportamento, a semana produziu uma deterioração ordenada do humor: o vigor caiu e a fadiga subiu com efeito '
  'grande entre o primeiro e o último dia, a maioria dos atletas migrou do perfil iceberg para perfis de menor '
  'energia, e a análise multivariada por permutação (PERMANOVA), livre do pressuposto de normalidade e respeitando '
  'as medidas repetidas, confirmou que o perfil de humor se deslocou de fato ao longo da semana (p = 0,002). No '
  'plano da forma temporal, a suavização das curvas — um filtro que separa a tendência lenta do ruído — e as '
  'derivadas dataram esse deslocamento: a fadiga e a PTH desaceleram até a metade da semana e voltam a acelerar a '
  'partir de um ponto de inflexão em torno do dia 4, marcador objetivo do momento em que o desgaste passa a se '
  'acumular. E a razão sinal/ruído, ao qualificar onde a informação é forte o bastante para sustentar inferência, '
  'fecha a triangulação: as vias multivariada, dinâmica e psicométrica apontam para o mesmo núcleo — o par '
  'vigor-fadiga como o eixo dominante do estado do atleta —, o que confere ao achado um alicerce que não depende de '
  'nenhum método isolado. A recuperação apenas parcial entre sessões, revelada pela decomposição do efeito agudo e '
  'da recuperação, dá a esse quadro o seu conteúdo prático: uma fadiga funcional, induzida de forma planejada, que '
  'deve ser monitorada justamente na relação entre a carga aguda e a recuperação entre as sessões.')
figure_bos('arvore_decisoes_bos.png','Analytical decision tree of the study: the choice of each statistical method according to the properties of the data (distribution, floor effect, type of comparison and multivariate assumptions), with the parallel measurement-quality track (ICC, α/ω, CFA and IRT) and the cross-cutting rules applied throughout the analysis.',w=13.5)

# ===== 5 DISCUSSÃO =====
H1('DISCUSSÃO','5')
P('O conjunto dos resultados sustenta uma tese central: em um microciclo pré-competitivo de handebol de elite, o '
  'estado de humor, descrito nas suas seis dimensões, comporta-se como um sistema dinâmico com um eixo dominante, '
  'o eixo energia-fadiga, que se deteriora de modo ordenado sob a carga e cujo curso pode ser lido, datado e '
  'quantificado. Cada camada de análise, da descrição simples à modelagem das trajetórias, converge para esse '
  'mesmo núcleo por caminhos independentes, o que confere robustez à interpretação e afasta a leitura de achados '
  'isolados. A discussão a seguir organiza essa convergência em quatro eixos: o significado do padrão observado, a '
  'sua tradução por perfis de humor, a contribuição metodológica das análises de trajetória e a sua aplicação '
  'prática, com atenção reflexiva aos limites psicométricos do instrumento e à natureza dos dados.')
H2('Migração do humor e o eixo energia-fadiga','5.1')
P('Os resultados descrevem, em um microciclo pré-competitivo de handebol de elite, a migração do humor da '
  'prontidão para a fadiga funcional. A queda do vigor e a elevação da fadiga entre o primeiro e o último dia '
  'alcançaram efeito grande e foram confirmadas pela análise multivariada, o que situa o eixo energia-fadiga como '
  'o principal responsável pela mudança do estado do atleta, entre as seis dimensões. Esse padrão reproduz, dentro '
  'de um microciclo de handebol, o comportamento já documentado em outras modalidades sob acúmulo de carga e sob '
  'treino intensificado (PIERCE, 2002; FERREIRA et al., 2026; THORPE et al., 2017).')
P('À tendência semanal somou-se uma resposta aguda coerente com o esforço. Entre o momento pré e o pós-treino, o '
  'vigor caiu e a fadiga e a PTH subiram, com magnitude de pequena a moderada, o que acrescenta uma camada de '
  'variação dentro do dia à deterioração observada entre dias. Esse comportamento converge com o de outros '
  'esportes coletivos, nos quais o humor responde de forma imediata ao evento esportivo, e reforça o valor de '
  'coletas repetidas dentro do dia, e não apenas entre dias (DO NASCIMENTO et al., 2026). Estudos experimentais '
  'com atletas de esportes coletivos confirmam esse curso: a fadiga percebida, medida pelo POMS, eleva-se logo '
  'após o esforço e retorna à linha de base apenas 24 a 48 horas depois, o que dá base fisiológica à resposta '
  'aguda e à recuperação apenas parcial entre sessões próximas (CROSS et al., 2022).')
P('A perturbação total do humor comportou-se como um integrador desse eixo. O vigor e a fadiga associaram-se de '
  'forma forte à PTH e explicaram a maior parte da sua variância, ao passo que as dimensões negativas, '
  'concentradas em valores baixos, contribuíram pouco para a variação. A consistência interna adequada do vigor e '
  'da fadiga e a sua correlação intraclasse de moderada a boa convergem para a mesma conclusão, a saber, que o '
  'núcleo do sinal do monitoramento reside no par energia-fadiga (LOCHBAUM et al., 2021; ROHLFS et al., 2023).')
P('Do ponto de vista psicométrico, a concentração das dimensões negativas junto ao limite inferior da escala é '
  'informativa, e não um mero artefato. O forte efeito de piso, com percentuais de escores nulos que chegaram a '
  '79% na confusão e a 67% na depressão, reduz a variância disponível dessas dimensões, comprime a sua '
  'distribuição e rebaixa a sua razão sinal/ruído, o que limita a capacidade de detectar mudança e ajuda a '
  'explicar a sua menor confiabilidade entre medidas repetidas. Em atletas saudáveis e bem ajustados, esse piso é '
  'esperado e, longe de invalidar o instrumento, delimita o seu uso: as dimensões negativas funcionam como '
  'sentinelas de estados clínicos raros, ao passo que o vigor e a fadiga, livres do piso e com sinal limpo, '
  'carregam a informação útil ao monitoramento da carga (TERRY et al., 2022; ROHLFS et al., 2023). Esse raciocínio '
  'recomenda cautela ao interpretar médias e testes das dimensões de piso e favorece, para elas, a leitura por '
  'prevalência e por perfis em vez da leitura por escore contínuo.')
H2('Perfis de humor e sobre-esforço funcional','5.2')
P('A leitura por perfis acrescentou clareza a essa descrição. O deslocamento do perfil iceberg para a barbatana '
  'de tubarão traduz, em uma única imagem, a acumulação da carga sem os sinais de comprometimento da saúde mental, '
  'uma vez que os perfis de maior risco, como o iceberg invertido e o submerso, não se instalaram. Quando '
  'comparados à distribuição de referência de uma grande amostra brasileira e aos padrões descritos em análises de '
  'agrupamento, os nossos dados partem de um predomínio semelhante de iceberg e elevam a barbatana de tubarão '
  'apenas ao fim da semana (PARSONS-SMITH; TERRY; MACHIN, 2017; LUOJUMÄKI et al., 2026; DE MIRANDA ROHLFS et al., '
  '2024).')
P('Essa reconfiguração tem valor comunicativo direto. O perfil condensa as seis dimensões em uma imagem única, o '
  'que aproxima o dado psicométrico da linguagem da comissão técnica e favorece a decisão de treino e de '
  'recuperação. A mesma leitura por perfis tem sido proposta como recurso de rastreio da saúde mental no esporte, '
  'o que amplia o alcance do monitoramento para além do desempenho (TERRY et al., 2021; HAN; PARSONS-SMITH; TERRY, '
  '2020; LEW et al., 2023).')
H2('Contribuição metodológica: derivadas, ajuste polinomial e cruzamentos','5.3')
P('Além de comparar médias, este estudo mapeou a forma temporal exata das trajetórias das dimensões. A suavização '
  'das séries e o cálculo da segunda derivada localizaram, em cada dimensão do eixo energia-fadiga, um único ponto '
  'de inflexão situado na metade da semana, no qual a concavidade muda e a taxa de variação atinge o seu limite. '
  'Esse marcador objetivo demarca a transição entre a fase inicial de prontidão e a fase de acúmulo de carga e '
  'oferece uma leitura que a simples comparação entre o primeiro e o último dia não revela.')
P('Lidas como um sistema em movimento, essas trajetórias admitem uma interpretação cinemática direta: a primeira '
  'derivada expressa a velocidade de desgaste — a taxa instantânea com que o vigor cai e a fadiga sobe —, e a '
  'segunda derivada expressa a sua aceleração, isto é, se o desgaste está ganhando ou perdendo ritmo. O ponto de '
  'inflexão, no qual a segunda derivada se anula em torno da metade da semana (dias 4,0 a 4,2 no eixo energia-fadiga), '
  'marca a passagem de uma fase de desaceleração para uma fase de aceleração da fadiga e funciona como um sinal de '
  'alerta precoce que a comparação entre o primeiro e o último dia não capta. Essa leitura articula-se com o '
  'resultado multivariado: o deslocamento do perfil de humor foi significativo ao longo da semana '
  '(PERMANOVA, p = 0,002), enquanto a resposta aguda isolada teve magnitude apenas pequena a '
  'moderada — coerente com um estímulo que não atua como um degrau estático, e sim alterando a velocidade da '
  'trajetória de fadiga nos dias subsequentes. Assim, o achado dinâmico e o multivariado convergem: o efeito '
  'relevante é crônico e acumulado, e a sua datação objetiva reside na inflexão de meia-semana.')
P('Essa convergência de métodos é o que confere robustez ao achado temporal. A PERMANOVA estabelece, sem pressupor '
  'normalidade e respeitando o delineamento de medidas repetidas, que o perfil multivariado de humor de fato se '
  'desloca ao longo da semana (p = 0,002); a suavização por spline atua como um filtro que separa a tendência lenta '
  'do ruído de alta frequência, e as derivadas datam objetivamente esse deslocamento na inflexão de meia-semana; e '
  'a razão sinal/ruído funciona como o critério de referência para distinguir sinal de ruído, qualificando em quais '
  'dimensões a informação é forte o bastante para sustentar inferência — o eixo energia-fadiga, com razão de 8,4 e '
  '6,7, ante as dimensões negativas que mal ultrapassam o ruído. Assim, a via multivariada (PERMANOVA), a via '
  'dinâmica (limites e derivadas sobre as curvas filtradas) e a via de qualidade de sinal (sinal/ruído) apontam '
  'para o mesmo núcleo por caminhos independentes, o que afasta a dependência de qualquer método isolado e sustenta '
  'a conclusão sobre um alicerce triangulado.')
P('O ajuste polinomial de grau três acrescentou uma síntese paramétrica a essa análise. O modelo reproduziu a '
  'inflexão de forma analítica, na raiz da segunda derivada, e resumiu cada trajetória por uma equação, com ajuste '
  'forte sobre as médias diárias (R² de 0,96 a 0,95 no eixo energia-fadiga). As duas resoluções, sobre sete médias '
  'diárias e sobre quatorze pontos pré e pós-treino, apontaram inflexões quase idênticas, o que indica que o '
  'achado da transição na metade da semana não depende do método nem da granularidade da medida. A convergência '
  'entre a spline não paramétrica e o polinômio paramétrico reforça a robustez do resultado. A opção pelo grau '
  'três, e não por um polinômio de grau superior, apoiou-se no coeficiente de determinação ajustado: o grau quatro '
  'elevava apenas marginalmente o R² bruto, mas reduzia o R² ajustado (de 0,88 para 0,86 na fadiga e de 0,60 para '
  '0,43 na PTH) e introduzia inflexões espúrias — sinal de sobreajuste com apenas sete médias diárias —, de modo '
  'que a cúbica, com um único ponto de inflexão, é a descrição mais parcimoniosa e estável da forma temporal.')
P('A localização dos cruzamentos exatos representa a contribuição mais original desta abordagem. Quando todas as '
  'dimensões são postas na mesma escala, apenas o par vigor-fadiga ocupa a faixa central e se cruza, ao passo que '
  'as demais permanecem junto ao piso, o que dá base empírica à escolha desse par como foco do monitoramento. O '
  'vigor superou a fadiga por ampla margem no início e foi alcançado por ela ainda na primeira metade da semana, '
  'mas apenas ao fim do microciclo a fadiga afastou-se de forma definitiva acima do vigor, seguida pela PTH, que '
  'ultrapassou primeiro o vigor e depois a fadiga. A datação desses cruzamentos converte a ideia qualitativa de '
  'inversão do eixo energia-fadiga em um evento com dia e escore definidos, o que fornece um candidato a marcador '
  'temporal do início do sobre-esforço funcional, coerente com o continuum descrito entre sobre-esforço e '
  'overtraining (ROETE et al., 2021; MANESCU et al., 2026; LA TORRE et al., 2023).')
P('O rigor deste conjunto de análises apoia-se na convergência de métodos independentes, e não em um único '
  'procedimento. A escolha de testes não paramétricos, adotada após a verificação da normalidade, protege as '
  'comparações contra os desvios de distribuição observados nas dimensões negativas, e o relato sistemático do '
  'tamanho e da magnitude do efeito, ao lado dos valores de p, evita a leitura de significância como se fosse '
  'relevância. Na análise de trajetória, a spline não paramétrica e o ajuste polinomial paramétrico localizaram a '
  'mesma inflexão, e as duas resoluções do polinômio apontaram valores quase idênticos, o que caracteriza uma '
  'triangulação metodológica. Os limiares de mudança derivados do coeficiente de correlação intraclasse delimitam '
  'ainda o que pode ser lido no plano individual e o que só se sustenta como tendência de grupo, o que mantém a '
  'inferência dentro dos limites do delineamento.')
P('A decomposição entre sinal e ruído acrescenta uma leitura reflexiva sobre a qualidade da informação de cada '
  'dimensão. Ao separar a tendência lenta (o sinal) das flutuações de alta frequência (o ruído de medida e a '
  'variabilidade biológica do dia a dia), a suavização mostrou que o vigor e a fadiga concentram a maior razão '
  'sinal/ruído (cerca de 8,4 e 6,7), enquanto as dimensões negativas, comprimidas pelo piso, mal ultrapassam o '
  'ruído. Esse resultado converge com a análise psicométrica e com a estrutura de componentes principais, e '
  'fundamenta, por um terceiro caminho, a decisão de centrar o monitoramento no eixo energia-fadiga, no qual o '
  'sinal é forte o bastante para sustentar inferências sobre a carga (SAW; MAIN; GASTIN, 2016; HELWIG et al., '
  '2023).')
P('A análise sequencial das treze coletas revelou, por fim, a microestrutura da fadiga funcional. A decomposição '
  'de cada dia em um efeito agudo (pré para pós) e um efeito de recuperação (pós para pré do dia seguinte) mostrou '
  'um sistema que oscila em dente de serra: a sessão eleva a perturbação e reduz o vigor, e o intervalo entre '
  'sessões devolve parte do estado, com destaque para a carga elevada do Dia 6 (PTH com dz = 0,75) e para a '
  'recuperação entre o Dia 2 e o Dia 3 (PTH com dz = -0,56). O ponto central é que a recuperação foi apenas '
  'parcial: como o retorno noturno não anulou o efeito agudo, o saldo acumulou-se e produziu a deriva descendente '
  'do vigor ao longo da semana. Essa leitura dá conteúdo operacional ao conceito de sobre-esforço funcional, no '
  'qual a fadiga é induzida de forma planejada e monitorada justamente pela relação entre a carga aguda e a '
  'recuperação entre sessões, e sugere que o desequilíbrio persistente entre esses dois termos seria o sinal de '
  'alerta para a transição ao sobre-esforço não funcional (KELLMANN et al., 2018; ROETE et al., 2021; MANESCU et '
  'al., 2026). Análises recentes de séries temporais em uma temporada completa de futebol profissional, com '
  'aprendizagem de máquina e modelos de mediação, mostram que a recuperação subjetiva é determinada sobretudo pela '
  'fadiga e que a carga do dia anterior atua por meio dela, o que confere plausibilidade empírica ao acoplamento '
  'entre carga aguda e recuperação aqui descrito (SIMONELLI; FORMENTI; ROSSI, 2025).')
H2('Por que este modelo supera a abordagem clássica e onde está a inovação','5.4')
P('O monitoramento do humor no esporte consolidou-se, ao longo das últimas décadas, em torno de uma rotina de '
  'análise de variância de medidas repetidas e testes t conduzida em pacotes de uso corrente, como o SPSS. Essa '
  'tradição responde bem a uma pergunta — houve diferença entre os dias? —, mas o faz sob pressupostos que os '
  'dados de humor de atletas com frequência violam: a normalidade e a esfericidade são exigidas justamente onde a '
  'compressão pelo efeito de piso produz assimetria e inflação de zeros, e o tempo é tratado como fator '
  'categórico, o que reduz uma trajetória contínua a uma sequência de médias e descarta a forma da mudança. Ao '
  'assumir os atletas como intercambiáveis, essa rotina também ignora a variação individual de trajetória. O '
  'produto é uma leitura que informa se o humor mudou, sem informar quando, com que velocidade, para quem e com '
  'que confiabilidade — precisamente as perguntas que orientam a decisão de ajuste de carga.')
P('O modelo analítico proposto neste estudo troca essa pergunta única por um encadeamento de métodos escolhidos '
  'em função das propriedades dos dados e resumido em uma árvore de decisões (Figura 30). Nenhuma de suas peças é, '
  'isoladamente, inédita: a suavização e a razão sinal/ruído provêm do processamento de sinais; a PERMANOVA, da '
  'ecologia de comunidades; a teoria de resposta ao item, da psicometria; os modelos mistos, da estatística '
  'longitudinal. A contribuição está em integrá-las em um único fluxo aplicado ao monitoramento diário do humor no '
  'esporte coletivo — um domínio que, historicamente, se acomodou à ANOVA e ao teste t sobre amostras pequenas e '
  'distribuições dominadas pelo piso, ou seja, nas condições em que esses testes clássicos são menos confiáveis. '
  'Importar o cálculo (derivadas, limites e ponto de inflexão) e a teoria de sinais (filtragem e razão '
  'sinal/ruído) para a leitura de estados psicológicos, e amarrar a seleção de cada método a uma regra explícita '
  'sobre a natureza dos dados, é o que distingue esta abordagem da prática consolidada e o que, até onde se pôde '
  'verificar, ainda não havia sido reunido dessa forma para a dinâmica do humor em uma pré-temporada de handebol '
  'de elite.')
P('O ganho é, ao mesmo tempo, estatístico e prático. Ao dispensar pressupostos que os dados não cumprem, ao '
  'descrever a forma da mudança e não apenas a sua presença, ao separar o sinal do ruído e ao respeitar a '
  'dependência das medidas repetidas, o modelo produz inferências mais defensáveis sobre os mesmos dados. E, ao '
  'delimitar honestamente o teto de previsibilidade individual — a regressão logística alcançou AUC de apenas 0,58 '
  'e os modelos flexíveis de gradient boosting (XGBoost e LightGBM) não a superaram —, evita a promessa exagerada '
  'de predição individual que a leitura clássica, silenciosa quanto à confiabilidade, não teria como sustentar. O '
  'resultado é um retrato acionável para a comissão técnica: onde, quando, quanto e com que segurança o humor se '
  'altera ao longo do microciclo.')
P('A robustez desses achados foi examinada por duas análises de sensibilidade. Na sensibilidade das variáveis, '
  'três lentes independentes — a contribuição multivariada (PERMANOVA univariada), a importância preditiva (por '
  'permutação) e a responsividade (tamanho de efeito e razão sinal/ruído) — convergiram sobre a fadiga como a '
  'dimensão de que os resultados mais dependem, seguida do vigor, o mais responsivo, porém de menor poder '
  'discriminativo isolado; as dimensões afetivas negativas contribuíram pouco (Figura 31). Na sensibilidade '
  'analítica, o efeito da semana permaneceu estável ao remover cada atleta (leave-one-athlete-out), à troca da '
  'janela de dias e à variação do suavizador, com a inflexão da fadiga fixada em torno do dia 4 em todos os '
  'filtros (Figura 32). Assim, as conclusões centrais não dependem de nenhuma variável isolada nem de uma '
  'escolha analítica particular.')
figure('sensibilidade_variaveis.png','Sensibilidade das variáveis: contribuição multivariada (pseudo-F univariado), importância preditiva (queda de AUC por permutação) e responsividade (|dz| do D1 ao D7) de cada dimensão do BRUMS.',w=16.0)
figure('sensibilidade_analitica.png','Sensibilidade analítica: (A) leave-one-athlete-out do efeito D1→D7; (B) janela de dias (com e sem baseline); (C) estabilidade do ponto de inflexão sob diferentes suavizadores.',w=16.0)
H2('Aplicação prática e perspectivas','5.5')
P('As características do handebol ajudam a explicar esse comportamento e orientam a aplicação. A modalidade impõe '
  'esforço intermitente de alta intensidade, com sprints curtos, mudanças de direção, saltos e contato, o que gera '
  'elevada demanda neuromuscular e psicofisiológica ao longo da semana e sustenta tanto a tendência de acúmulo '
  'quanto a resposta aguda entre o pré e o pós-treino (KARCHER; BUCHHEIT, 2014; GARCÍA-SÁNCHEZ et al., 2023; '
  'CARTON-LLORENTE et al., 2023). Na prática, entre as seis dimensões, o par vigor-fadiga concentra o sinal útil, '
  'e o cruzamento entre as suas curvas oferece um alerta simples e visual, capaz de complementar o monitoramento '
  'de carga interna e de fadiga já descrito na modalidade (STRUZIK; NADOBNIK; STEPIEN-SLODKOWSKA, 2026; STAIANO et '
  'al., 2025).')
P('O acompanhamento subjetivo do humor deve ser lido como parte de um monitoramento multidomínio, e não como '
  'medida isolada. A sua sensibilidade, o baixo custo e o caráter não invasivo tornam-no complementar aos '
  'marcadores fisiológicos, endócrinos e de bem-estar praticados em esportes coletivos (RATZ-SULYOK et al., 2026; '
  'BIRD et al., 2025; HELWIG et al., 2023). Evidências recentes que combinam o POMS a sensores vestíveis em '
  'atletas de elite reforçam que a integração de medidas subjetivas e objetivas, analisada de forma '
  'individualizada, supera as médias de grupo na leitura do equilíbrio entre carga e recuperação, o que aponta o '
  'caminho para o refinamento futuro do monitoramento do humor (SPETZ et al., 2025). Como agenda futura, '
  'recomenda-se integrar medidas objetivas de carga a este delineamento, ampliar a amostra e testar se a datação '
  'do cruzamento do eixo energia-fadiga antecipa desfechos de fadiga e de desempenho, de modo a validar o marcador '
  'temporal aqui proposto (KELLMANN et al., 2018; SAW; MAIN; GASTIN, 2016).')
H2('Limitações','5.6')
P('O estudo tem limitações, entre as quais a amostra de um único clube, o recorte de um microciclo e a ausência '
  'de medidas objetivas de carga neste recorte, o que restringe a generalização e impede inferências de causa. O '
  'pequeno número de observações distribuído por seis perfis reduz a potência do teste categórico, e o erro de '
  'medida de uma leitura isolada recomenda cautela na interpretação individual, de modo que a inferência se apoia '
  'nas tendências de grupo. Ainda assim, a descrição oferece um retrato direto da dinâmica do humor no handebol de '
  'elite e sustenta a recomendação prática de acompanhar as seis dimensões, com foco no eixo energia-fadiga, ao '
  'longo da semana, integrado a um monitoramento multidomínio do estado do atleta (OSTAPIUK-KAROLCZUK et al., '
  '2025).')

# ===== 6 CONCLUSÃO =====
H1('CONCLUSÃO','6')
P('Em um microciclo pré-competitivo de handebol de elite, o estado de humor, descrito nas suas seis dimensões, '
  'deteriorou-se de forma ordenada ao longo da semana, e essa deterioração concentrou-se no eixo energia-fadiga. O '
  'vigor caiu e a fadiga subiu com efeito grande entre o primeiro e o último dia, resultado confirmado pela '
  'análise multivariada, ao passo que as dimensões negativas permaneceram junto ao piso e com baixa razão '
  'sinal-ruído, o que delimita o valor informativo de cada dimensão.')
P('A leitura por perfis traduziu esse percurso na migração do iceberg para a barbatana de tubarão, sem a '
  'instalação dos perfis de risco à saúde mental, um quadro compatível com o sobre-esforço funcional buscado de '
  'forma planejada na pré-temporada. A modelagem das trajetórias localizou um ponto de inflexão na metade da '
  'semana e datou os cruzamentos entre o vigor, a fadiga e a perturbação total do humor, o que converte a inversão '
  'qualitativa do eixo energia-fadiga em eventos com dia e escore definidos e oferece um marcador temporal '
  'objetivo e visual. A decomposição da resposta em efeito agudo e de recuperação revelou um padrão em dente de '
  'serra, com recuperação apenas parcial entre as sessões, mecanismo que explica a deriva descendente do vigor ao '
  'longo do microciclo.')
P('Como desfecho aplicado, os resultados recomendam centrar o monitoramento no par vigor-fadiga e no cruzamento '
  'entre as suas curvas, um alerta simples e de baixo custo, integrado a um monitoramento multidomínio do estado '
  'do atleta. Como desfecho científico, o estudo entrega um conjunto de marcadores temporais reprodutíveis, '
  'obtidos por triangulação entre a suavização não paramétrica, o ajuste polinomial e a datação dos cruzamentos, '
  'que ampliam a leitura tradicional por comparação de médias. As perspectivas incluem integrar medidas objetivas '
  'de carga ao delineamento, ampliar a amostra e verificar se a datação do cruzamento do eixo energia-fadiga '
  'antecipa desfechos de fadiga e de desempenho, o que permitiria validar o marcador temporal proposto e '
  'incorporá-lo à rotina da comissão técnica.')

# ===== REFERÊNCIAS =====
H1('REFERÊNCIAS')
refs=[
 'ANDERSON, M. J. A new method for non-parametric multivariate analysis of variance. Austral Ecology, v. 26, n. 1, p. 32-46, 2001. DOI: 10.1111/j.1442-9993.2001.01070.pp.x.',
 'ANDERSON, M. J. Distance-based tests for homogeneity of multivariate dispersions. Biometrics, v. 62, n. 1, p. 245-253, 2006. DOI: 10.1111/j.1541-0420.2005.00440.x.',
 'BIRD, S. P. et al. Wellness, mood, sleep, and performance in a women’s national basketball team during international competition. Journal of Human Kinetics, v. 96, p. 163-175, 2025. DOI: 10.5114/jhk/200117.',
 'CARTON-LLORENTE, A. et al. Worst-case scenario analysis of physical demands in elite men handball players by playing position through big data analytics. Biology of Sport, v. 40, n. 4, p. 1219-1227, 2023. DOI: 10.5114/biolsport.2023.126665.',
 'CROSS, R. et al. Acute neuromuscular response to team sports-specific running, resistance, and concurrent training: a crossover study. Medicine and Science in Sports and Exercise, v. 54, n. 3, p. 456-465, 2022. DOI: 10.1249/MSS.0000000000002804.',
 'DE MIRANDA ROHLFS, I. C. P. et al. Prevalence of specific mood profile clusters among elite and youth athletes at a Brazilian sports club. Sports, v. 12, n. 7, 195, 2024. DOI: 10.3390/sports12070195.',
 'DO NASCIMENTO, M. H. et al. Acute psychological responses to official match outcomes in male youth volleyball: an observational repeated-measures study within a single national-level team. Frontiers in Psychology, v. 17, 1826372, 2026. DOI: 10.3389/fpsyg.2026.1826372.',
 'FERREIRA, A. B. M. et al. Impact of sleep restriction and intensified training on mucosal immunity and psychological responses in young soccer players. Journal of Strength and Conditioning Research, v. 40, n. 7, p. e703-e713, 2026. DOI: 10.1519/JSC.0000000000005416.',
 'GARCÍA-SÁNCHEZ, C. et al. Physical demands during official competitions in elite handball: a systematic review. International Journal of Environmental Research and Public Health, v. 20, n. 4, 3353, 2023. DOI: 10.3390/ijerph20043353.',
 'HAN, C.; PARSONS-SMITH, R. L.; TERRY, P. C. Mood profiling in Singapore: cross-cultural validation and potential applications of mood profile clusters. Frontiers in Psychology, v. 11, 665, 2020. DOI: 10.3389/fpsyg.2020.00665.',
 'HELWIG, J. et al. Relationships between external, wearable sensor-based, and internal parameters: a systematic review. Sensors, v. 23, n. 2, 827, 2023. DOI: 10.3390/s23020827.',
 'KARCHER, C.; BUCHHEIT, M. On-court demands of elite handball, with special reference to playing positions. Sports Medicine, v. 44, n. 6, p. 797-814, 2014. DOI: 10.1007/s40279-014-0164-z.',
 'KELLMANN, M. et al. Recovery and performance in sport: consensus statement. International Journal of Sports Physiology and Performance, v. 13, n. 2, p. 240-245, 2018. DOI: 10.1123/ijspp.2017-0759.',
 'LA TORRE, M. E. et al. The potential role of nutrition in overtraining syndrome: a narrative review. Nutrients, v. 15, n. 23, 4916, 2023. DOI: 10.3390/nu15234916.',
 'LEW, P. C. F. et al. Cross-cultural validation of the Malaysian Mood Scale and tests of between-group mood differences. International Journal of Environmental Research and Public Health, v. 20, n. 4, 3348, 2023. DOI: 10.3390/ijerph20043348.',
 'LOCHBAUM, M. et al. The Profile of Mood States and athletic performance: a meta-analysis of published studies. European Journal of Investigation in Health, Psychology and Education, v. 11, n. 1, p. 50-70, 2021. DOI: 10.3390/ejihpe11010005.',
 'LUOJUMÄKI, R. J. et al. Exploring mood profile clusters across physical activity level, gender and age in a Finnish population. European Journal of Sport Science, v. 26, n. 2, e70131, 2026. DOI: 10.1002/ejsc.70131.',
 'MANESCU, D. C. et al. Molecular biomarkers of training responses: a systems framework for exercise adaptation and athlete monitoring. International Journal of Molecular Sciences, v. 27, n. 8, 3601, 2026. DOI: 10.3390/ijms27083601.',
 'MORGAN, W. P. Selected psychological factors limiting performance: a mental health model. In: CLARKE, D. H.; ECKERT, H. M. (Ed.). Limits of human performance. Champaign: Human Kinetics, 1985. p. 70-80.',
 'OSTAPIUK-KAROLCZUK, J. et al. Biochemical and psychological markers of fatigue and recovery in mixed martial arts athletes during strength and conditioning training. Scientific Reports, v. 15, n. 1, 24234, 2025. DOI: 10.1038/s41598-025-09719-z.',
 'PARSONS-SMITH, R. L.; TERRY, P. C.; MACHIN, M. A. Identification and description of novel mood profile clusters. Frontiers in Psychology, v. 8, 1958, 2017. DOI: 10.3389/fpsyg.2017.01958.',
 'PIERCE, E. F. Relationship between training volume and mood states in competitive swimmers during a 24-week season. Perceptual and Motor Skills, v. 94, n. 3, p. 1009-1012, 2002. DOI: 10.2466/pms.2002.94.3.1009.',
 'RATZ-SULYOK, F. Z. et al. Associations between endocrine status and stress, mood and psychosomatic status in elite handball players. Sports, v. 14, n. 7, 289, 2026. DOI: 10.3390/sports14070289.',
 'ROETE, A. J. et al. A systematic review on markers of functional overreaching in endurance athletes. International Journal of Sports Physiology and Performance, v. 16, n. 8, p. 1065-1073, 2021. DOI: 10.1123/ijspp.2021-0024.',
 'ROHLFS, I. C. P. M. et al. A Escala de Humor de Brunel (Brums): instrumento para detecção precoce da síndrome do excesso de treinamento. Revista Brasileira de Medicina do Esporte, v. 14, n. 3, p. 176-181, 2008.',
 'ROHLFS, I. C. P. M. et al. Psychometric characteristics of the Brazil Mood Scale among youth and elite athletes using two response time frames. Sports, v. 11, n. 12, 244, 2023. DOI: 10.3390/sports11120244.',
 'SAW, A. E.; MAIN, L. C.; GASTIN, P. B. Monitoring the athlete training response: subjective self-reported measures trump commonly used objective measures: a systematic review. British Journal of Sports Medicine, v. 50, n. 5, p. 281-291, 2016. DOI: 10.1136/bjsports-2015-094758.',
 'SIMONELLI, C.; FORMENTI, D.; ROSSI, A. Subjective recovery in professional soccer players: a machine learning and mediation approach. Journal of Sports Sciences, v. 43, n. 5, p. 448-455, 2025. DOI: 10.1080/02640414.2025.2461932.',
 'SPETZ, L. et al. Validating subjective ratings with wearable data for a nuanced understanding of load-recovery status in elite endurance athletes. Sports Medicine - Open, v. 11, n. 1, 154, 2025. DOI: 10.1186/s40798-025-00958-y.',
 'STAIANO, W. et al. Overcoming mental fatigue through mindfulness: improving physical and cognitive performance in elite handball players. Journal of Science and Medicine in Sport, v. 29, n. 1, p. 91-99, 2025. DOI: 10.1016/j.jsams.2025.08.004.',
 'STRUZIK, A.; NADOBNIK, J.; STEPIEN-SLODKOWSKA, M. TRIMP and session-RPE monitoring in elite women’s handball: a full-season descriptive analysis. Scientific Reports, v. 16, n. 1, 2026. DOI: 10.1038/s41598-026-53134-x.',
 'TERRY, P. C.; LANE, A. M.; FOGARTY, G. J. Construct validity of the Profile of Mood States, Adolescents for use with adults. Psychology of Sport and Exercise, v. 4, n. 2, p. 125-139, 2003. DOI: 10.1016/S1469-0292(02)00035-8.',
 'TERRY, P. C. et al. Mood profiling for sustainable mental health among athletes. Sustainability, v. 13, n. 11, 6116, 2021. DOI: 10.3390/su13116116.',
 'TERRY, P. C. et al. Validation of a Lithuanian-language version of the Brunel Mood Scale: the BRUMS-LTU. International Journal of Environmental Research and Public Health, v. 19, n. 8, 4867, 2022. DOI: 10.3390/ijerph19084867.',
 'THORPE, R. T. et al. Monitoring fatigue status in elite team-sport athletes: implications for practice. International Journal of Sports Physiology and Performance, v. 12, n. S2, p. S227-S234, 2017. DOI: 10.1123/ijspp.2016-0434.']
for rf in refs:
    p=doc.add_paragraph(); r=p.add_run(rf); r.font.size=Pt(11); r.font.name='Times New Roman'
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(6); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

doc.save('/home/user/mdlucca/Artigos/Paper1_Humor_revisado.docx')
print('OK Paper1_Humor_revisado.docx  figs=%d tabs=%d refs=%d'%(_FN[0],_TN[0],len(refs)))
