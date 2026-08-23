# -*- coding: utf-8 -*-
import json, numpy as np
import matplotlib; matplotlib.use('Agg'); import matplotlib.pyplot as plt
res=json.load(open('perfil_variacao.json'))
COL={'Iceberg':'#2f9e44','Superfície':'#1971c2','Submerso':'#7048e8','Barbatana de tubarão':'#e8590c',
     'Everest invertido':'#e0525b','Iceberg invertido':'#f59f00'}
days=np.arange(1,8)
plt.rcParams.update({'font.family':'DejaVu Sans','font.size':11,'axes.spines.top':False,'axes.spines.right':False})
fig,ax=plt.subplots(figsize=(11,6),dpi=200)
for d in [2,4,7]: ax.axvspan(d-.18,d+.18,color='#e8590c',alpha=.08,zorder=0)
big={'Iceberg','Superfície'}
for r in res:
    p=r['perfil']; y=r['pct']; lw=3.4 if p in big else 1.8; a=1 if p in big else .8
    ax.plot(days,y,'-o',color=COL[p],lw=lw,ms=7 if p in big else 5,alpha=a,zorder=3 if p in big else 2,
            label=f"{p}  (Δ {r['delta']:+.0f}"+(" pp*" if r['qp']<0.05 else " pp")+")")
    ax.annotate(f"{y[-1]:.0f}%",(7,y[-1]),textcoords='offset points',xytext=(8,0),fontsize=9,color=COL[p],fontweight='bold',va='center')
ax.annotate('1ª sessão de HIIT\n(maior salto: superfície +31 pp)',(2,42),textcoords='offset points',xytext=(12,18),
            fontsize=8.5,color='#495057',arrowprops=dict(arrowstyle='->',color='#adb5bd'))
ax.set_xticks(days); ax.set_xlabel('Dia do microciclo'); ax.set_ylabel('Prevalência do perfil (% de atletas)')
ax.set_ylim(0,50); ax.set_xlim(.7,7.7)
ax.set_title('Variação da prevalência de cada perfil de humor ao longo da semana\nfaixas = dias de HIIT · * variação significativa entre dias (Cochran Q, p<0,05)',
             fontweight='bold',fontsize=12,loc='left')
ax.legend(frameon=False,fontsize=9,loc='upper center',ncol=3,bbox_to_anchor=(0.5,-0.12))
fig.tight_layout(); fig.savefig('/home/user/mdlucca/Artigos/figuras/perfil_variacao.png',bbox_inches='tight',facecolor='white')
print('figura ok')

# ---- tabela docx ----
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
def cm(x): return str(x).replace('.',',')
doc=Document(); st=doc.styles['Normal']; st.font.name='Times New Roman'; st.font.size=Pt(11)
p=doc.add_paragraph(); r=p.add_run('Tabela. Variação da prevalência dos seis perfis de humor ao longo dos sete dias do microciclo.'); r.bold=True; r.font.size=Pt(10.5)
hdr=['Perfil','D1','D2','D3','D4','D5','D6','D7','Δ(D7−D1)','Amplitude','Maior salto','Cochran Q (p)']
t=doc.add_table(rows=1,cols=len(hdr)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
def setc(c,txt,b=False,sz=9,al='center'):
    c.text=''; pp=c.paragraphs[0]; pp.alignment={'center':WD_ALIGN_PARAGRAPH.CENTER,'left':WD_ALIGN_PARAGRAPH.LEFT}[al]
    rr=pp.add_run(str(txt)); rr.font.size=Pt(sz); rr.font.name='Times New Roman'; rr.bold=b
for j,hh in enumerate(hdr): setc(t.rows[0].cells[j],hh,b=True,sz=8.5)
for r in res:
    c=t.add_row().cells
    qp=cm(r['qp'])+('*' if r['qp']<0.05 else '')
    vals=[r['perfil']]+[f"{x:.0f}" for x in r['pct']]+[f"{r['delta']:+.0f}",cm(r['amp']),cm(r['maxstep']),qp]
    for j,v in enumerate(vals): setc(c[j],v,al='left' if j==0 else 'center',sz=8.5)
nt=doc.add_paragraph(); rn=nt.add_run('Valores em % de atletas classificados em cada perfil por dia. Δ(D7−D1): mudança líquida do primeiro ao último dia (pontos percentuais); '
 'Amplitude: máximo − mínimo ao longo da semana; Maior salto: maior variação entre dois dias consecutivos; Cochran Q: teste de homogeneidade da prevalência '
 'entre os sete dias (subconjunto balanceado de 19 atletas com registro em todos os dias); * p < 0,05. Os dois maiores movimentos são espelhados — o perfil '
 'iceberg cai 31 pp (41% → 10%) e o superfície sobe 27 pp (11% → 38%) —, com o maior salto isolado no D1→D2 (entrada do HIIT).')
rn.italic=True; rn.font.size=Pt(9)
doc.save('/home/user/mdlucca/Artigos/Tabela_Variacao_Perfis.docx')
print('tabela ok')
