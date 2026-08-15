# -*- coding: utf-8 -*-
import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
DOC='/home/user/mdlucca/Artigos/PERFIL_HUMOR_COMPLETO.docx'
R=json.load(open('robust_extra.json'))
def c2(s): return str(s).replace('.',',')
d=Document(DOC)
def _fmt(p,size=12,just=True,bold=False,center=False):
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(6)
    if center: p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    elif just: p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    for r in p.runs: r.font.name='Times New Roman'; r.font.size=Pt(size); r.bold=bold
def after(el,text='',**kw):
    p=d.add_paragraph(text); el.addnext(p._p); _fmt(p,**kw); return p._p
def head(el,text): return after(el,text,bold=True,just=False)
def _bd(cell):
    tcPr=cell._tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for edge in ['top','bottom','left','right']:
        e=OxmlElement(f'w:{edge}'); on=edge in ('top','bottom')
        e.set(qn('w:val'),'single' if on else 'nil'); e.set(qn('w:sz'),'6'); e.set(qn('w:color'),'000000'); b.append(e)
    tcPr.append(b)
def table_after(el,caption,header,rows,fonte='Fonte: dados da pesquisa (2026).',fs=9.5):
    cap=d.add_paragraph(caption); el.addnext(cap._p); _fmt(cap,size=11,just=False)
    t=d.add_table(rows=1,cols=len(header)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; cap._p.addnext(t._tbl)
    for i,htx in enumerate(header):
        cc=t.rows[0].cells[i]; cc.text=''; rr=cc.paragraphs[0].add_run(htx); rr.bold=True; rr.font.size=Pt(fs); rr.font.name='Times New Roman'
        cc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; _bd(cc)
    for row in rows:
        cs=t.add_row().cells
        for i,val in enumerate(row):
            cs[i].text=''; rr=cs[i].paragraphs[0].add_run(str(val)); rr.font.size=Pt(fs); rr.font.name='Times New Roman'
            cs[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if i>0 else WD_ALIGN_PARAGRAPH.LEFT; _bd(cs[i])
    fon=d.add_paragraph(fonte); t._tbl.addnext(fon._p); _fmt(fon,size=10,just=False); return fon._p
def fig_after(el,path,caption,w=16.0):
    pimg=d.add_paragraph(); el.addnext(pimg._p); pimg.alignment=WD_ALIGN_PARAGRAPH.CENTER; pimg.add_run().add_picture(path,width=Cm(w))
    cap=d.add_paragraph(caption); pimg._p.addnext(cap._p); _fmt(cap,size=11,just=False,center=True)
    fon=d.add_paragraph('Fonte: elaboração dos autores (2026).'); cap._p.addnext(fon._p); _fmt(fon,size=10,just=False,center=True); return fon._p
def find(prefix):
    for p in d.paragraphs:
        if p.text.strip().startswith(prefix): return p
    raise KeyError(prefix)

AC=R['acute']; MM=R['mm']
p6=find('6 SÍNTESE DOS PRINCIPAIS ACHADOS'); anchor=p6._p.getprevious()
anchor=head(anchor,'5.13 Efeito agudo pré→pós: a resposta mais robusta do microciclo e sua hierarquia')
anchor=after(anchor,
 'A resposta aguda dentro do dia (do pré ao pós-treino), agregada por atleta-dia ao longo de toda a semana (%d pares), é '
 'o achado estatisticamente mais robusto do estudo. A fadiga física é o marcador mais reativo, com efeito grande '
 '(dz = %s; IC95%% %s–%s; p = 5×10⁻¹³), seguida da fadiga do BRUMS (dz = %s), da perturbação total do humor (dz = %s) e '
 'do vigor (dz = %s) — todos com p < 0,001 (Tabela 17; Figura 21). A fadiga mental responde de forma pequena, porém '
 'significativa (dz = %s; p = 0,007), enquanto as subescalas negativas de valência não fadiga — depressão, tensão, raiva '
 'e confusão — não apresentam resposta aguda (intervalos de confiança cruzando o zero). Confirma-se, com máxima robustez, '
 'que a deterioração aguda do treino se concentra no eixo energia–fadiga e poupa o afeto negativo geral.'%(
   AC['FadFisica']['n'],c2('%+.2f'%AC['FadFisica']['dz']),c2('%.2f'%AC['FadFisica']['lo']),c2('%.2f'%AC['FadFisica']['hi']),
   c2('%+.2f'%AC['Fadiga']['dz']),c2('%+.2f'%AC['TMD']['dz']),c2('%+.2f'%AC['Vigor']['dz']),c2('%+.2f'%AC['FadMental']['dz'])))
anchor=after(anchor,
 'A magnitude do efeito agudo é estável ao longo da semana (interação momento×dia não significativa: fadiga física '
 'p = 0,12; vigor p = 0,08), ou seja, cada sessão produz um choque de fadiga de tamanho semelhante, sem sensibilização '
 'progressiva estatisticamente detectável.')
anchor=after(anchor,
 'Um modelo multinível (efeitos mistos; %d observações; intercepto e inclinação aleatórios por atleta) confirma, com mais '
 'rigor do que as médias por atleta, tanto a trajetória de grupo — queda do vigor (β = %s por dia; p = %s) e subida da '
 'fadiga física (β = %s por dia; p < 10⁻⁷) — quanto o efeito de nível da aptidão: cada km·h⁻¹ de PV soma %s ao vigor e '
 'subtrai %s da fadiga física (p = %s e p = %s). A interação aptidão×dia não é significativa: a vantagem do mais apto está '
 'no patamar do eixo energia–fadiga, não numa desaceleração distinguível da trajetória — coerente com a leitura de §5.12 '
 '(a aptidão preserva o vigor sob carga).'%(
   R['n_obs'],c2('%+.2f'%MM['Vigor']['dia_c']['b']),c2('%.3f'%MM['Vigor']['dia_c']['p']),
   c2('%+.2f'%MM['FadFisica']['dia_c']['b']),c2('%+.2f'%MM['Vigor']['PV_c']['b']),c2('%.2f'%abs(MM['FadFisica']['PV_c']['b'])),
   c2('%.3f'%MM['Vigor']['PV_c']['p']),c2('%.3f'%MM['FadFisica']['PV_c']['p'])))
# Tabela 17
order=['FadFisica','Fadiga','TMD','Vigor','FadMental','Depressao','Tensao','Raiva','Confusao']
def prow(v):
    a=AC[v]; p='<0,001' if a['p']<0.001 else c2('%.3f'%a['p'])
    return [a['lab'],c2('%+.2f'%a['delta']),c2('%+.2f'%a['dz']),'%s–%s'%(c2('%.2f'%a['lo']),c2('%.2f'%a['hi'])),p]
anchor=table_after(anchor,'Tabela 17 – Efeito agudo pré→pós por subescala (agregado por atleta-dia): variação média, tamanho de efeito (dz) com IC95% e significância.',
    ['Subescala','Δ (pós−pré)','dz','IC95% (dz)','p'],[prow(v) for v in order])
anchor=fig_after(anchor,'/home/user/mdlucca/Artigos/figuras/x_agudo.png',
    'Figura 21 – Hierarquia do efeito agudo pré→pós (dz com IC95%). A fadiga física tem efeito grande; o eixo energia–fadiga (fadiga BRUMS, PTH, vigor, fadiga mental) responde de forma significativa; as subescalas negativas de valência não fadiga (cinza) não respondem. Linhas pontilhadas: limiares de efeito pequeno/médio/grande.')

# SÍNTESE: novo bullet
try:
    ph=find('•  Dinâmica:')
    nb=after(ph._p,'•  Resposta aguda (achado mais robusto): do pré ao pós-treino, a fadiga física sobe com efeito grande (dz = +0,76; p = 5×10⁻¹³) e o vigor cai (dz = −0,39); o eixo energia–fadiga concentra toda a resposta aguda, enquanto o afeto negativo (depressão, tensão, raiva, confusão) permanece inerte. O choque é estável a cada sessão, sem sensibilização progressiva.')
except KeyError: pass

d.save(DOC); print('SAVED',DOC)
