import warnings; warnings.filterwarnings('ignore')
import numpy as np, pandas as pd, json, os
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
SC='/tmp/claude-0/-home-user-mdlucca/e1dba24c-b1d7-5908-9106-f2f4aaf3f56a/scratchpad'
FIG='/home/user/mdlucca/Artigos/figuras'
R=json.load(open(f'{SC}/semana2.json')); REG=json.load(open(f'{SC}/reg.json')); BF=json.load(open(f'{SC}/baseline_final.json'))
DX=json.load(open(f'{SC}/decomp_extra.json')); REL=json.load(open(f'{SC}/reliab.json')); MV=json.load(open(f'{SC}/mv.json'))
RMC=json.load(open(f'{SC}/rmcorr.json')); CFA=json.load(open(f'{SC}/cfa.json'))
DE=json.load(open(f'{SC}/deriv_exp.json')); ROC=json.load(open(f'{SC}/roc.json')); ALLO=json.load(open(f'{SC}/allo.json'))
socio=R['socio']; desc=R['desc']; cron=R['cron']; agu=R['agu']; dec=R['dec']; der=R['der']; rank=R['rank']
CORE=[('Vigor','Vigor'),('Fadiga','Fadiga (BRUMS)'),('FadFisica','Fadiga física'),('FadMental','Fadiga mental'),('TMD','PTH/TMD')]
def c2(x): return str(x).replace('.',',')

doc=Document(); s=doc.sections[0]; s.left_margin=Cm(3); s.top_margin=Cm(3); s.right_margin=Cm(2); s.bottom_margin=Cm(2)
n=doc.styles['Normal']; n.font.name='Times New Roman'; n.font.size=Pt(12); n.paragraph_format.line_spacing=1.5; n.paragraph_format.space_after=Pt(0)
FN='Times New Roman'
def _set(r,sz=12,b=False,it=False): r.font.name=FN; r.font.size=Pt(sz); r.bold=b; r.italic=it
def sec(num,txt):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(4); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    _set(p.add_run((f'{num} {txt}').upper() if '.' not in num else f'{num} {txt}'),12,b=True)
def sub(num,txt):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(2); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    _set(p.add_run(f'{num} {txt}'),12,b=True)
def P(t,indent=True,size=12,it=False):
    p=doc.add_paragraph(); _set(p.add_run(t),size,it=it); p.paragraph_format.line_spacing=1.5
    if indent: p.paragraph_format.first_line_indent=Cm(1.25)
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
def _img(name,label,num,titulo,w,fonte):
    cap=doc.add_paragraph(); cap.alignment=WD_ALIGN_PARAGRAPH.CENTER; cap.paragraph_format.space_before=Pt(6); cap.paragraph_format.line_spacing=1.0
    _set(cap.add_run(f'{label} {num} – {titulo}'),10)
    p=f'{FIG}/{name}.png'
    if os.path.exists(p): doc.add_picture(p,width=Inches(w)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    else: print('MISS',name)
    fp=doc.add_paragraph(); fp.alignment=WD_ALIGN_PARAGRAPH.CENTER; fp.paragraph_format.space_after=Pt(6); fp.paragraph_format.line_spacing=1.0
    _set(fp.add_run(f'Fonte: {fonte}'),10)
_GN=[0]; _FN=[0]; _TN=[0]
def graf(name,titulo,w=5.6,fonte='Elaborado pelos autores (2026).'): _GN[0]+=1; _img(name,'Gráfico',_GN[0],titulo,w,fonte)
def fig(name,titulo,w=5.6,fonte='Elaborado pelos autores (2026).'): _FN[0]+=1; _img(name,'Figura',_FN[0],titulo,w,fonte)
def _bd(t):
    tblPr=t._tbl.tblPr; b=OxmlElement('w:tblBorders')
    for e in ('top','bottom','insideH'):
        x=OxmlElement(f'w:{e}'); x.set(qn('w:val'),'single'); x.set(qn('w:sz'),'6'); x.set(qn('w:space'),'0'); x.set(qn('w:color'),'000000'); b.append(x)
    for e in ('left','right','insideV'):
        x=OxmlElement(f'w:{e}'); x.set(qn('w:val'),'none'); b.append(x)
    tblPr.append(b)
def table(titulo,headers,rows,fonte='Dados da pesquisa (2026).',fs=8.5):
    _TN[0]+=1
    cap=doc.add_paragraph(); cap.alignment=WD_ALIGN_PARAGRAPH.LEFT; cap.paragraph_format.space_before=Pt(6); cap.paragraph_format.line_spacing=1.0
    _set(cap.add_run(f'Tabela {_TN[0]} – {titulo}'),10)
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; _bd(t)
    for i,h in enumerate(headers):
        cl=t.rows[0].cells[i]; cl.paragraphs[0].line_spacing=1.0; _set(cl.paragraphs[0].add_run(str(h)),fs,b=True)
    for row in rows:
        cs=t.add_row().cells
        for i,v in enumerate(row):
            cs[i].paragraphs[0].line_spacing=1.0; _set(cs[i].paragraphs[0].add_run(str(v)),fs)
    fp=doc.add_paragraph(); fp.alignment=WD_ALIGN_PARAGRAPH.LEFT; fp.paragraph_format.space_after=Pt(6); fp.paragraph_format.line_spacing=1.0
    _set(fp.add_run(f'Fonte: {fonte}'),10)

# TÍTULO + RESUMO
tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
_set(tp.add_run('PERFIL DE HUMOR DE ATLETAS DE HANDEBOL NA ÚLTIMA SEMANA DA PRÉ-TEMPORADA COMPETITIVA'),14,b=True)
a2=doc.add_paragraph(); a2.alignment=WD_ALIGN_PARAGRAPH.CENTER; _set(a2.add_run('[Autoria e afiliação a preencher]'),11,it=True)
rp=doc.add_paragraph(); _set(rp.add_run('RESUMO'),12,b=True); rp.paragraph_format.line_spacing=1.0
P('Objetivo: avaliar e classificar o perfil de humor de atletas de handebol na última semana da pré-temporada competitiva. Método: 27 atletas responderam ao BRUMS-24 (seis subescalas e a Perturbação Total do Humor, PTH) e a autoavaliações de fadiga física e mental duas vezes ao dia — a primeira coleta (pré) e a última (pós) — durante sete dias (456 observações). Analisaram-se estatísticas descritivas completas, a comparação do grupo entre a linha de base (Dia 1, pré) e o último dia (Dia 7, pós), a resposta pré→pós, a regressão linear, a decomposição da variância (sinal, traço e ruído), as derivadas e os pontos de inflexão. Resultados: o humor deteriorou-se de forma concentrada no eixo energia–fadiga — da linha de base ao final, o vigor caiu 47% (dz = −1,02) e a fadiga física (+129%; dz = +2,12), a fadiga do BRUMS (+116%; dz = +0,78), a fadiga mental (+39%; dz = +0,82) e a PTH aumentaram; a fadiga física foi a mais responsiva. A trajetória inflectiu-se em torno do Dia 4 e o perfil iceberg reduziu-se ao longo da semana. Conclusão: o custo do período é somático e acumulativo, sendo a fadiga física e o vigor os marcadores mais sensíveis para o monitoramento individualizado.',indent=False)
kw=doc.add_paragraph(); _set(kw.add_run('Palavras-chave: '),12,b=True); _set(kw.add_run('Estado de humor. BRUMS. Fadiga. Pré-temporada. Handebol.'),12); kw.paragraph_format.line_spacing=1.5

# 1 INTRODUÇÃO
sec('1','Introdução')
P('O monitoramento dos estados de humor é uma ferramenta prática, de baixo custo e não invasiva para acompanhar a resposta psicobiológica de atletas à carga de treinamento, sinalizando desadaptações antes que se manifestem em queda de desempenho (TERRY; LANE; FOGARTY, 2003; ROHLFS et al., 2008). A Brunel Mood Scale (BRUMS) operacionaliza o construto em seis dimensões — tensão, depressão, raiva, vigor, fadiga e confusão — e resume o perfil na Perturbação Total do Humor (PTH), índice que integra as dimensões negativas subtraídas do vigor. Atletas saudáveis tendem a exibir o perfil iceberg, com vigor elevado destacando-se acima das dimensões negativas (MORGAN, 1985); o achatamento desse perfil, com a convergência entre vigor e fadiga, indica fadiga acumulada e risco de excesso de treinamento.')
P('A caracterização do comportamento do humor em diferentes escalas temporais — no conjunto da semana, na variação da linha de base ao último dia e na resposta aguda dentro de cada dia — e, simultaneamente, nos níveis coletivo e individual, fornece um panorama interpretável para calibrar a carga e individualizar a recuperação. Este estudo tem por objetivo avaliar e classificar o perfil de humor de atletas de handebol durante a última semana da pré-temporada competitiva, descrevendo o comportamento geral do grupo por variável (incluindo a PTH), a variação da linha de base (Dia 1, pré) ao último dia (Dia 7, pós), a resposta pré→pós e a sensibilidade relativa das variáveis, com aprofundamento por regressão, decomposição da variância, derivadas e pontos de inflexão.')

# 2 JUSTIFICATIVA
sec('2','Justificativa')
P('A última semana pré-competitiva é um período crítico do planejamento, no qual a comissão técnica busca dissipar a fadiga acumulada preservando as adaptações, de modo a maximizar a prontidão para competir. O manejo inadequado da carga nesse intervalo pode deteriorar o humor e sinalizar recuperação incompleta, com repercussões diretas sobre o desempenho; ainda assim, o comportamento do humor nessa janela raramente é caracterizado de forma estruturada nas equipes. Descrevê-lo quantitativamente — com estatística descritiva completa, identificando quais variáveis são mais sensíveis, como se comporta o índice global (PTH) e em que momento a trajetória se inflete — oferece à comissão técnica subsídios objetivos para individualizar a recuperação e intervir no momento oportuno.')

# 3 MÉTODO
sec('3','Método')
ag=socio; pc=json.load(open(f'{SC}/app_data.json'))['socio']['pos_counts']
P('Estudo descritivo de medidas repetidas, em condições ecológicas de treinamento, com 27 atletas de handebol do sexo masculino acompanhados por sete dias, cada um funcionando como sua própria referência (dados anonimizados, A01–A27). O humor foi avaliado pela BRUMS-24 (seis subescalas de quatro itens, escore de 0 a 16; PTH = tensão + depressão + raiva + fadiga + confusão − vigor) e a fadiga física e mental por escalas de autoavaliação de 0 a 10, autoaplicadas por formulário eletrônico com carimbo de data/hora duas vezes ao dia (a primeira resposta tomada como “pré” e a última como “pós”), totalizando 456 respostas válidas. O processamento e a auditoria interna foram inteiramente roteirizados em Python (Figura 1).')
P('Computaram-se estatísticas descritivas completas (média, desvio-padrão, coeficiente de variação, mediana, intervalo interquartílico, mínimo, máximo, amplitude e assimetria). A variação da linha de base (Dia 1, pré) ao último dia (Dia 7, pós) e a resposta pré→pós foram avaliadas por Wilcoxon pareado com tamanho de efeito (dz); a tendência semanal por regressão linear (coeficiente por dia e R²). A variância de cada variável foi decomposta (modelo atleta + dia) em sinal da semana, traço e ruído, derivando-se o erro típico de medida (ETM), a mudança mínima detectável (MDC95) e a relação sinal-ruído (SNR). A forma da trajetória foi caracterizada por ajuste cúbico das médias diárias (velocidade, aceleração e ponto de inflexão) e por suavização LOWESS. Adotou-se α = 0,05; software Python (pandas, numpy, SciPy, statsmodels).')
fig('fig_fluxograma','Fluxograma do processamento e da análise dos dados em Python, com auditoria interna por script',w=5.4)

# 4 RESULTADOS
sec('4','Resultados')
sub('4.1','Caracterização sociodemográfica e antropométrica')
rows=[]
for k in ['idade','exp','estatura','massa','pG','PV','CMJ']:
    if k in socio:
        d=socio[k]; rows.append([d['lab'],c2(f"{d['mean']:.1f} ± {d['sd']:.1f}"),(c2(f"{d['cv']:.1f}") if d['cv'] else '—'),c2(f"{d['med']:.1f}"),c2(f"{d['mn']:.1f}–{d['mx']:.1f}")])
table('Estatística descritiva sociodemográfica e antropométrica da amostra (n = 27)',['Variável','M ± DP','CV (%)','Mediana','Mín.–Máx.'],rows)
P(f"A amostra, homogênea, distribuiu-se por posição em {pc.get('Armador',0)} armadores, {pc.get('Ala',0)} alas, {pc.get('Pivô',0)} pivôs e {pc.get('Goleiro',0)} goleiros. Trata-se de atletas adultos, de estatura e massa elevadas e boa aptidão aeróbia (pico de velocidade no T-CAR ≈ {socio['PV']['mean']:.1f} km/h), com coeficientes de variação moderados, exceto na experiência e na gordura corporal, mais dispersas.")

sub('4.2','Estatística descritiva das variáveis de humor e fadiga')
rows=[]
for v,lab in [('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Vigor','Vigor'),('Fadiga','Fadiga (BRUMS)'),('Confusao','Confusão'),('TMD','PTH/TMD'),('FadFisica','Fadiga física'),('FadMental','Fadiga mental')]:
    d=desc[v]; rows.append([lab,c2(f"{d['mean']:.2f} ± {d['sd']:.2f}"),c2(f"{d['med']:.1f} [{d['iqr']:.1f}]"),f"{d['mn']:.0f}–{d['mx']:.0f}",c2(f"{d['amp']:.0f}"),c2(f"{d['skew']:+.2f}"),c2(f"{d['floor']:.0f}%")])
table('Estatística descritiva completa do humor e da fadiga na semana',['Variável','M ± DP','Mediana [IIQ]','Mín.–Máx.','Ampl.','Assim.','% piso'],rows)
P('No conjunto da semana, o grupo preservou, em média, o perfil iceberg, com o vigor acima das subescalas negativas (Figura 2, radar). As subescalas negativas concentram-se no piso (tensão, depressão, raiva e confusão com 50–80% dos escores no valor mínimo e assimetria positiva acentuada), enquanto o vigor e a fadiga física são aproximadamente simétricos e bem dispersos (Gráfico 1, histogramas). A fadiga foi a dimensão negativa mais expressiva, e a ampla amplitude da PTH (−12 a 52) revela marcada heterogeneidade interindividual. Os diagramas de caixa por dia (Gráfico 2) evidenciam o deslocamento progressivo do vigor (para baixo) e das fadigas e da PTH (para cima).')
fig('sem_f_radar','Forma do perfil de humor (escores z) no Dia 1 e no Dia 7',w=4.4)
graf('v_hist','Histogramas de distribuição das cinco variáveis centrais na semana (linha tracejada = média)')
graf('v_box','Diagramas de caixa (boxplots) das cinco variáveis centrais por dia')

sub('4.3','Tendência e comportamento ao longo da semana')
rows=[[lab,c2(f"{REG[v]['slope']:+.3f}"),c2(f"{REG[v]['r2']:.3f}"),(c2(f"{REG[v]['p']:.4f}") if REG[v]['p']>=0.001 else '<0,001')] for v,lab in CORE]
table('Regressão linear do humor ao longo da semana (coeficiente por dia, R² e p)',['Variável','Coef./dia (b)','R²','p'],rows)
P('A regressão linear (Tabela 3; Gráfico 3) confirma a direção e a significância das tendências: a fadiga física apresenta a inclinação mais acentuada e o melhor ajuste (b = +0,39/dia; R² = 0,10; p < 0,001), seguida da fadiga do BRUMS (b = +0,41/dia; p < 0,001), do vigor (b = −0,27/dia; p = 0,005) e da PTH; a fadiga mental não exibe tendência (p = 0,60). As trajetórias (Gráfico 4) mostram o vigor declinando e as fadigas e a PTH ascendendo rumo ao Dia 7, com as bandas de desvio-padrão amplas confirmando a heterogeneidade. O Gráfico 5 sintetiza, em escores padronizados, a deterioração conjunta do eixo. Os R² baixos decorrem de a regressão ser computada no nível da observação, em que a variância é dominada por diferenças entre atletas e ruído — o que recomenda interpretar a tendência do grupo, e não a leitura isolada de um atleta.')
graf('v_reg','Regressão linear de cada variável ao longo da semana (pontos = observações; reta = ajuste)')
graf('v_traj','Comportamento semanal das cinco variáveis (M ± DP; suavização LOWESS pontilhada)')
graf('v_deter','Deterioração do eixo energia–fadiga em escores padronizados (z): vigor declina, fadigas e PTH ascendem',w=5.0)

sub('4.4','Comportamento da Perturbação Total do Humor (PTH/TMD)')
P('A PTH, índice global do perfil, comporta-se como esperado do eixo energia–humor: eleva-se ao longo da semana (b = +0,41/dia; p < 0,001; Tabela 3) e responde agudamente pré→pós (dz = +0,66), com a maior amplitude bruta entre as variáveis (Gráfico 6). Contudo, por integrar seis subescalas — várias ancoradas no piso —, a PTH tem a pior relação sinal-ruído do eixo (SNR = 0,96) e apenas 3% da variância atribuível ao sinal da semana, o que a torna um indicador global útil para a leitura de conjunto, porém menos sensível e mais ruidoso do que o vigor e a fadiga física isolados. A precipitação da PTH no Dia 7 é o ponto máximo da série, coerente com o achatamento do perfil iceberg.')
graf('v_tmd','Comportamento da PTH/TMD na semana: trajetória (M ± DP; LOWESS), resposta pré vs pós por dia e distribuição por dia')

sub('4.5','Linha de base (Dia 1, pré) versus último dia (Dia 7, pós) e sensibilidade')
rows=[[lab,c2(f"{BF[v]['base']:.2f} ± {BF[v]['base_sd']:.2f}"),c2(f"{BF[v]['fin']:.2f} ± {BF[v]['fin_sd']:.2f}"),c2(f"{BF[v]['delta']:+.2f}"),f"{BF[v]['pct']:+.0f}%",c2(f"{BF[v]['dz']:+.2f}"),(c2(f"{BF[v]['p']:.3f}") if BF[v]['p']>=0.001 else '<0,001')] for v,lab in CORE]
table('Humor e fadiga na linha de base (Dia 1, pré) e no último dia (Dia 7, pós); comparação pareada (Wilcoxon)',['Variável','Pré D1','Pós D7','Δ','% mud.','dz','p'],rows)
rows=[[r['lab'],c2(f"{r['dz_ag']:.2f}"),c2(f"{r['dz_cr']:.2f}"),c2(f"{r['snr']:.2f}"),f"{r['pday']:.0f}%"] for r in sorted(rank,key=lambda x:-x['dz_cr'])]
table('Sensibilidade/responsividade das variáveis (efeito agudo, crônico, SNR e % de sinal)',['Variável','|dz| agudo','|dz| crônico','SNR','% sinal'],rows)
P('A comparação geral do grupo entre a linha de base e o último dia confirma e amplia o quadro: o vigor caiu 47% (8,25 → 4,38; dz = −1,02) e a fadiga física quase triplicou em magnitude de efeito (3,69 → 8,44; +129%; dz = +2,12), com a fadiga do BRUMS (+116%), a fadiga mental (+39%) e a PTH também aumentando. A hierarquia de sensibilidade (Tabela 5) é clara — fadiga física > vigor > fadiga do BRUMS > PTH ≈ fadiga mental —, com a fadiga física reunindo o maior efeito e a melhor SNR. O efeito crônico supera o agudo pré→pós (fadiga física dz crônico = +2,12 versus agudo = +1,08), indicando acúmulo progressivo, e não apenas resposta pontual ao esforço diário.')

sub('4.6','Limites, decomposição da variância, derivadas e ponto de inflexão')
rows=[[lab,f"{desc[v]['mn']:.0f}–{desc[v]['mx']:.0f}",c2(f"{desc[v]['amp']:.0f}"),c2(f"{dec[v]['pday']:.0f}%"),c2(f"{dec[v]['ptrait']:.0f}%"),c2(f"{dec[v]['pnoise']:.0f}%"),c2(f"{dec[v]['snr']:.2f}")] for v,lab in CORE]
table('Limites (amplitude) e decomposição da variância (sinal, traço, ruído) e SNR',['Variável','Mín.–Máx.','Ampl.','% Sinal','% Traço','% Ruído','SNR'],rows)
rows=[[lab,c2(f"{der[v]['vel_d1']:+.2f}"),c2(f"{der[v]['vel_d7']:+.2f}"),(c2(f"dia {der[v]['infl']:.1f}") if der[v]['infl'] else 'n/d'),(', '.join(c2(f"{x:.1f}") for x in der[v]['turn']) if der[v]['turn'] else '—')] for v,lab in CORE]
table('Derivadas (velocidade no D1 e no D7) e ponto de inflexão por variável',['Variável','Vel. D1','Vel. D7','Inflexão','Viradas (vel = 0)'],rows)
P('A decomposição mostra que o sinal da semana é a menor fração da variância (fadiga física 13%, vigor 7%, fadiga do BRUMS 4%, PTH 3%, fadiga mental 1%), com SNR favorável apenas à fadiga física (1,82) e ao vigor (1,52); a maior parte é traço e ruído, o que sustenta o monitoramento por médias de coletas e por tendência.')
rows=[[lab,c2(f"{DE[v]['amp_tot']:.0f}"),c2(f"{DE[v]['amp_intra']:.1f}"),c2(f"{DE[v]['amp_sig']:.2f}"),c2(f"{DE[v]['pct_scale']:.0f}%"),c2(f"{DE[v]['auc']:.1f}")] for v,lab in CORE]
table('Amplitude em três níveis (total, intra-atleta e do sinal), fração da escala utilizada e dose acumulada (área sob a trajetória)',['Variável','Ampl. total','Ampl. intra-atleta','Ampl. do sinal','% da escala','Dose (AUC)'],rows)
P('Os limites revelam uma hierarquia constante — amplitude do sinal < amplitude intra-atleta < amplitude total —: o atleta oscila 2 a 3 vezes mais que a trajetória média do grupo, porque sua oscilação carrega o próprio ruído, e o sinal do microciclo ocupa apenas 5–33% da escala. A dose acumulada (área sob a trajetória) sintetiza a carga psicométrica total do período. As derivadas (Figuras 3 e 4) descrevem a forma da mudança: a velocidade da fadiga física e da fadiga do BRUMS é máxima nas extremidades (início e Dia 7), com um mínimo no meio da semana; a aceleração troca de sinal uma única vez — o ponto de inflexão —, em torno do Dia 4 para o vigor, a fadiga física e a fadiga do BRUMS (e ≈ Dia 3,5 para a PTH). O sinal da terceira derivada (jerk) é positivo para as fadigas, indicando que a aceleração da deterioração cresce rumo ao fim da semana — a assinatura da precipitação do Dia 7, e não de uma saturação.')
fig('v_deriv','Velocidade (linha cheia) e aceleração (tracejada) da trajetória de cada variável; linha vertical = ponto de inflexão')
fig('x_deriv_exp','Velocidade (barras) e aceleração (linha pontilhada) por dia, para cada variável — perfil completo das derivadas')

sub('4.7','Perfil iceberg ao longo da semana')
tab=R['perfil_tab']
rows=[[f"Dia {dd}",f"{tab['Iceberg'][str(dd)]:.0f}",f"{tab['Barbatana tubarão'][str(dd)]:.0f}",f"{tab['Superfície'][str(dd)]:.0f}",f"{tab['Iceberg invertido'][str(dd)]:.0f}",f"{tab['Everest invertido'][str(dd)]:.0f}"] for dd in [1,4,7]]
table('Distribuição dos seis perfis de humor (%) em dias-chave',['Dia','Iceberg','Barbatana','Superfície','Iceberg inv.','Everest inv.'],rows)
P('A classificação nos seis perfis (Figura 4) mostra a migração do perfil saudável para o de fadiga: o iceberg cai do Dia 1 ao Dia 7 e a barbatana de tubarão — fadiga em pico com vigor ainda preservado — cresce expressivamente. O percentual em perfil iceberg e o índice-iceberg declinam ao longo da semana, com precipitação no Dia 7 (Gráfico 7). Os perfis de colapso afetivo permanecem raros, reforçando que o custo do período é somático, e não afetivo.')
graf('sem_f_iceberg','Perfil iceberg ao longo da semana: percentual em perfil iceberg (barras), índice-iceberg (linha) e regressão',w=5.4)
fig('fig_perfis_esquema','Esquema dos seis perfis de humor (escores z das subescalas; vigor em verde, negativas em vermelho)',w=5.2)

sub('4.8','Comportamento individual entre atletas')
P('No nível individual, a heterogeneidade é marcante. A variação da linha de base ao último dia difere amplamente entre atletas (Gráfico 8): quase todos reduzem o vigor e elevam a fadiga física, mas em magnitudes distintas, havendo atletas com variação modesta e outros com grande deterioração. As trajetórias individuais e a resposta pré→pós (Figura 5) confirmam a direção compartilhada com dispersão substancial: escores idênticos podem representar estados de recuperação distintos conforme a referência de cada atleta. Esses achados desaconselham pontos de corte normativos únicos e sustentam o referenciamento de cada atleta à sua própria linha de base, com decisões apoiadas em médias de coletas.')
graf('v_delta','Variação individual da linha de base ao último dia (Δ D7 − D1) por atleta, para o vigor e a fadiga física')
fig('v_individual','Comportamento individual: trajetórias semanais (superior) e resposta pré→pós (inferior) do vigor e da fadiga física (linha em destaque = média do grupo)',w=5.8)

sub('4.9','Decomposição avançada: generalizabilidade, heterogeneidade e dimensionalidade')
fc=DX['facet']
rows=[[fc[v]['lab'],c2(f"{fc[v]['ptrait']:.0f}%"),c2(f"{fc[v]['pstate']:.0f}%"),c2(f"{fc[v]['pnoise']:.0f}%"),c2(f"{fc[v]['icc1']:.2f}"),c2(f"{fc[v]['phi']:.2f}")] for v,_ in CORE]
table('Decomposição em três facetas (traço, estado e ruído), fidedignidade de uma medida (ICC) e dependabilidade (Φ, k = 14 coletas)',['Variável','% Traço (atleta)','% Estado (dia)','% Ruído (intra-dia)','ICC₁','Φ'],rows)
P('A teoria da generalizabilidade permite separar o que o modelo de dois fatores agregava: além do traço (diferenças estáveis entre atletas) e do ruído (erro intra-dia), isola-se o estado — a variância atribuível ao dia, isto é, ao sinal do microciclo. Essa fatia de estado é pequena em todas as variáveis (3–9%), confirmando que a oscilação semanal, embora real, é modesta frente ao traço e ao ruído (Gráfico 9). A fidedignidade de uma única medida (ICC₁) é moderada (0,33–0,67), mas a dependabilidade da média das coletas da semana é excelente (Φ = 0,86–0,96), o que fundamenta, de forma quantitativa, a recomendação de decidir por médias e não por registros isolados.')
graf('x_facet','Decomposição da variância em três facetas: estado (sinal do microciclo), traço e ruído',w=5.4)
sl=DX['slope']; ac=DX['acf']
rows=[[fc[v]['lab'],c2(f"{sl[v]['mean']:+.3f}") if sl[v]['mean'] is not None else '—',c2(f"{sl[v]['sd']:.3f}") if sl[v]['sd'] is not None else '—',c2(f"{ac[v]:+.2f}"),('branco' if abs(ac[v])<0.2 else 'estruturado')] for v,_ in CORE]
table('Heterogeneidade do declínio (inclinação individual) e natureza do ruído (autocorrelação de lag-1 dos resíduos)',['Variável','Inclinação média/dia','DP entre atletas','ACF(1)','Ruído'],rows)
P('O modelo de inclinação aleatória revela quão uniformemente os atletas se deterioram: a fadiga física apresenta a menor dispersão de inclinações (DP = 0,09/dia) — todos pioram praticamente no mesmo ritmo, o que a torna previsível —, ao passo que a PTH tem inclinações muito heterogêneas (DP = 1,29/dia), com atletas melhorando e outros piorando (Figura 6). A autocorrelação de lag-1 dos resíduos é próxima de zero na maioria das variáveis, indicando ruído essencialmente branco (erro de medida puro), o que valida a decomposição; a fadiga física é a exceção (ACF = −0,27), sugerindo leve reversão à média — parte de seu resíduo é flutuação lenta, não apenas erro.')
fig('x_slopes','Distribuição das inclinações individuais (unidade por dia) por variável — heterogeneidade do declínio entre atletas',w=5.4)
rc=DX['rci']
rows=[[fc[v]['lab'],c2(f"{rc[v]['etm']:.2f}"),c2(f"{rc[v]['mdc']:.2f}"),f"{rc[v]['pct']:.0f}%"] for v,_ in CORE]
table('Mudança confiável por atleta (RCI): erro típico, mudança mínima detectável e proporção de atletas com mudança individual confiável (D1→D7)',['Variável','ETM','MDC95','% atletas com mudança confiável'],rows)
P('No nível individual, a proporção de atletas cuja mudança da linha de base ao último dia supera o limiar de ruído (|RCI| > 1,96) é baixa — fadiga do BRUMS 19%, vigor e PTH 14%, fadiga física 10% e fadiga mental 0% —, o que confirma, por um índice per capita, que uma única coleta é insuficiente para a decisão individual confiável, reforçando o uso de médias de coletas. Por fim, a análise de componentes principais das seis subescalas (Gráfico 10) mostra que o humor é essencialmente bidimensional: o primeiro componente explica 40% e o segundo 23% da variância (63% acumulados), correspondendo ao eixo energia–fadiga e a um eixo de afeto negativo — ou seja, o "sinal" do humor concentra-se em poucas dimensões, e o monitoramento pode focar-se nelas.')
graf('x_pca','Variância explicada pelos componentes principais das seis subescalas do BRUMS (barras) e variância acumulada (linha)',w=5.0)

sub('4.10','Confiabilidade interna das subescalas (α de Cronbach e ω de McDonald)')
rows=[[sub,str(REL[sub]['k']),c2(f"{REL[sub]['alpha']:.2f}"),c2(f"{REL[sub]['omega']:.2f}"),c2(f"{REL[sub]['mic']:.2f}")] for sub in ['Tensão','Depressão','Raiva','Vigor','Fadiga','Confusão']]
table('Confiabilidade (consistência interna) das seis subescalas do BRUMS',['Subescala','Nº itens','α de Cronbach','ω de McDonald','r̄ inter-itens'],rows)
P('A consistência interna das subescalas foi, em geral, adequada a boa, com o ω de McDonald — estimador menos enviesado que o α para poucos itens — variando de 0,69 a 0,91: excelente para depressão (0,91), raiva (0,91) e fadiga (0,87), e boa para vigor (0,83), confusão (0,79) e mesmo tensão (0,69). O α de Cronbach acompanha esse padrão, exceto na tensão (α = 0,43), cujo valor é rebaixado não por falha do construto, mas pelo forte efeito de piso (baixa variância e baixa correlação média inter-itens, r̄ = 0,16), como se vê ao contrastá-lo com o ω. Em conjunto, os escores das subescalas são suficientemente fidedignos para o monitoramento, com a ressalva de que a tensão deve ser lida com cautela neste contexto.')

sub('4.11','Modelo misto multivariado do eixo energia–fadiga')
P('O eixo energia–fadiga foi modelado conjuntamente (vigor, fadiga do BRUMS, fadiga física e fadiga mental; a PTH foi excluída por ser derivada dessas dimensões). A estrutura entre-atletas (Figura 7) confirma o eixo como dimensão de traço: o vigor correlaciona-se negativamente com as três fadigas (r = −0,27 a −0,38), enquanto a fadiga do BRUMS e a fadiga física correlacionam-se fortemente entre si (r = +0,74) — atletas mais fatigados numa dimensão tendem a sê-lo nas demais.')
fig('x_trait_corr','Correlações entre-atletas (nível de traço) das quatro variáveis do eixo energia–fadiga',w=4.8)
rows=[
 ['Modelo misto multivariado — efeito conjunto do dia (LRT)',f"χ²({MV['dfd']}) = {MV['lr']:.1f}".replace('.',','),('p < 0,001' if MV['pjoint']<0.001 else c2(f"p = {MV['pjoint']:.3f}"))],
 ['Hotelling T² — Dia 1 vs Dia 7 (multivariado pareado)',f"F({MV['dfn']},{MV['dfden']}) = {MV['F']:.2f}".replace('.',','),('p < 0,001' if MV['pF']<0.001 else c2(f"p = {MV['pF']:.3f}"))],
 ['Tamanho de efeito multivariado — D de Mahalanobis (D1→D7)',c2(f"D = {MV['mahal']:.2f}"),'—'],
]
table('Testes multivariados do efeito do microciclo sobre o eixo energia–fadiga (quatro variáveis conjuntas)',['Teste','Estatística','p'],rows)
P('O modelo misto multivariado — que ajusta as quatro variáveis conjuntamente, com efeitos aleatórios de atleta — indica um efeito do dia altamente significativo sobre o perfil como um todo (razão de verossimilhança χ²(4) = 109,1; p ≈ 10⁻²²), confirmando que a deterioração é um deslocamento coordenado do vetor de humor, e não fruto de uma variável isolada. A comparação multivariada da linha de base ao último dia (Hotelling T²) é igualmente significativa (F(4,17) = 18,17; p < 0,001), com tamanho de efeito multivariado muito grande (D de Mahalanobis = 2,02) — magnitude que integra, num único índice, o afastamento conjunto do vigor e das fadigas entre o início e o fim da semana.')

sub('4.12','Acoplamento intra-atleta das variáveis (correlação de medidas repetidas, rmcorr)')
P('Enquanto a correlação entre-atletas (Figura 7) descreve como os atletas se ordenam entre si, a correlação de medidas repetidas (rmcorr) descreve como as variáveis co-variam dentro de cada atleta ao longo da semana, removendo as diferenças de nível entre indivíduos. A matriz rmcorr (Figura 8) confirma o eixo energia–fadiga também no nível intraindividual: quando um atleta piora, o vigor cai e as fadigas sobem de forma acoplada — o vigor correlaciona-se negativamente com a fadiga do BRUMS (rmcorr = −0,48), a fadiga física (−0,50) e a PTH (−0,59), enquanto a fadiga do BRUMS e a fadiga física sobem juntas (rmcorr = +0,65) e ambas acompanham a PTH (+0,72 e +0,51). Esse acoplamento intraindividual reforça que o monitoramento de qualquer marcador do eixo captura, em boa parte, a mesma informação de estado.')
fig('x_rmcorr','Matriz de correlação intra-atleta (rmcorr) entre as variáveis do eixo energia–fadiga',w=4.8)

sub('4.13','Análise fatorial confirmatória da estrutura de seis fatores do BRUMS')
rows=[
 ['χ² (df)',f"{CFA['chi2']:.0f} ({CFA['df']})",'—'],
 ['CFI',c2(f"{CFA['cfi']:.3f}"),'≥ 0,90 (aceitável); ≥ 0,95 (bom)'],
 ['TLI',c2(f"{CFA['tli']:.3f}"),'≥ 0,90'],
 ['RMSEA',c2(f"{CFA['rmsea']:.3f}"),'≤ 0,08 (aceitável); ≤ 0,06 (bom)'],
 ['SRMR',c2(f"{CFA['srmr']:.3f}"),'≤ 0,08'],
]
table('Índices de ajuste da análise fatorial confirmatória do modelo de seis fatores do BRUMS (24 itens; estimação por máxima verossimilhança)',['Índice','Valor','Referência'],rows)
P('A análise fatorial confirmatória testou a estrutura hipotetizada de seis fatores correlacionados (cada subescala com quatro itens). O ajuste foi aceitável, ainda que limítrofe: RMSEA = 0,080 e SRMR = 0,084 nos limiares de aceitação, com CFI = 0,86 e TLI = 0,84 pouco abaixo do valor de referência de 0,90. As cargas padronizadas foram, em geral, elevadas (depressão 0,71–0,87; raiva 0,71–0,85; fadiga até 0,91), com exceção de itens isolados ancorados no piso (na tensão, no vigor e na confusão), que apresentaram cargas baixas e explicam a parte do desajuste — coerente com a menor confiabilidade da tensão (Tabela 12). As correlações entre fatores reproduzem a estrutura afetiva esperada (vigor × fadiga = −0,55; depressão × confusão = +0,72). Cabe a ressalva metodológica de que as 457 observações são medidas repetidas aninhadas em 27 atletas — não independentes —, de modo que os índices devem ser lidos como aproximação confirmatória; ainda assim, sustentam a validade de construto do instrumento neste contexto, com a tensão como a dimensão mais frágil.')

sub('4.14','Discriminação (curva ROC) com e sem ruído: o efeito da filtragem')
def aucstr(t): return c2('%.2f [%.2f; %.2f]'%(t[0],t[1],t[2]))
rows=[[ROC[v]['lab'],aucstr(ROC[v]['raw']),aucstr(ROC[v]['filt'])] for v,_ in CORE]
table('Área sob a curva ROC (AUC) na discriminação do último dia (Dia 7) versus a linha de base (Dia 1), com ruído (observações) e sem ruído (média diária filtrada), com IC95% bootstrap',['Variável','AUC com ruído [IC95%]','AUC sem ruído [IC95%]'],rows)
P('A curva ROC quantifica o quanto cada variável distingue uma observação do último dia de uma observação da linha de base. Com os dados brutos (observações individuais, com ruído), a fadiga física já é o melhor discriminador (AUC = 0,85), seguida do vigor (0,74) e da fadiga do BRUMS (0,73); a fadiga mental fica no acaso (0,52). Ao filtrar o ruído — tomando a média das coletas do dia —, a discriminação melhora nas variáveis responsivas: a fadiga física sobe para AUC = 0,90 e o vigor para 0,78 (Figura 9), enquanto as variáveis sem sinal (fadiga mental, PTH) não se beneficiam, como esperado. Isso demonstra, de forma direta, que a filtragem do ruído aumenta o poder discriminativo apenas onde há sinal a recuperar, e reforça o uso de médias de coletas para a decisão.')
fig('x_roc','Curvas ROC (Dia 7 vs Dia 1) com ruído (cinza, observações) e sem ruído (colorido, média diária filtrada)',w=5.6)

sub('4.15','Ajuste alométrico: forma da acumulação e escalonamento pela aptidão')
rows=[[ALLO[v]['lab'],c2(f"{ALLO[v]['a']:.2f}"),c2(f"{ALLO[v]['b']:.2f}"),c2(f"{ALLO[v]['r2']:.2f}"),(c2(f"{ALLO[v]['aic_pow']:.1f}")+' / '+c2(f"{ALLO[v]['aic_lin']:.1f}"))] for v in ['FadFisica','Fadiga','TMD']]
table('Ajuste alométrico (potência) da trajetória de fadiga, Y = a · diaᵇ, e comparação de AIC (potência / linear)',['Variável','a','b (expoente)','R²','AIC pot./lin.'],rows)
P('A acumulação da fadiga ao longo da semana foi modelada por uma função de potência (alométrica) Y = a·diaᵇ, cujo expoente b resume a forma do crescimento (b > 1 acelera; b < 1 desacelera/satura). Para a fadiga física, o ajuste é bom (a = 4,34; b = 0,24; R² = 0,88) e supera o linear por AIC, indicando um crescimento sublinear em média — rápido no início e mais lento no miolo da semana (Figura 10, escala log-log) —, sobre o qual se sobrepõe a precipitação do Dia 7 que apenas o modelo cúbico captura integralmente. Complementarmente, a aptidão aeróbia foi escalonada alometricamente pela massa corporal (pico de velocidade do T-CAR = a·massa^(−0,19)); a aptidão bruta associa-se negativamente à fadiga média da semana (ρ = −0,49; p = 0,014) — atletas mais aptos fatigam menos —, mas essa associação enfraquece após o escalonamento alométrico (ρ = −0,31; p = 0,135), sugerindo que parte do efeito protetor da aptidão é explicada pelo porte corporal.')
fig('x_allo','Ajuste alométrico (potência) da trajetória de fadiga em escala linear e log-log',w=5.4)

# 5 CONSIDERAÇÕES
sec('5','Considerações Finais')
P('O perfil de humor deteriorou-se de forma coerente e concentrada no eixo energia–fadiga, tanto ao longo da semana quanto na comparação da linha de base ao último dia. A fadiga física foi a variável mais sensível e responsiva, com maior fração de sinal e melhor relação sinal-ruído, seguida do vigor; a fadiga do BRUMS ocupou posição intermediária, e a PTH — embora útil como índice global — mostrou-se mais ruidosa, enquanto a fadiga mental e as subescalas negativas foram pouco responsivas. O efeito crônico superou o agudo, indicando acúmulo, e as trajetórias inflectiram-se em torno do Dia 4. Recomenda-se acompanhar a fadiga física e o vigor por tendência individual e por médias de coletas, com atenção reforçada a partir do meio da semana. Como limitações, trata-se de estudo descritivo, sem grupo de comparação e restrito a um único período de uma equipe masculina.')

sec('REFERÊNCIAS','')
for r in [
 'MORGAN, W. P. Selected psychological factors limiting performance: a mental health model. In: CLARKE, D. H.; ECKERT, H. M. (Ed.). Limits of human performance. Champaign: Human Kinetics, 1985. p. 70-80.',
 'ROHLFS, I. C. P. M. et al. A Escala de Humor de Brunel (Brums): instrumento para detecção precoce da síndrome do excesso de treinamento. Revista Brasileira de Medicina do Esporte, v. 14, n. 3, p. 176-181, 2008.',
 'TERRY, P. C.; LANE, A. M.; FOGARTY, G. J. Construct validity of the Profile of Mood States — Adolescents for use with adults. Psychology of Sport and Exercise, v. 4, n. 2, p. 125-139, 2003.',
]:
    p=doc.add_paragraph(); _set(p.add_run(r),12); p.paragraph_format.line_spacing=1.0; p.paragraph_format.space_after=Pt(6); p.alignment=WD_ALIGN_PARAGRAPH.LEFT

OUTP='/home/user/mdlucca/Artigos/Perfil_Humor_Semana_PreTemporada_ABNT.docx'
doc.save(OUTP); print('SAVED',OUTP)
d2=Document(OUTP); print('Tabelas',_TN[0],'Gráficos',_GN[0],'Figuras',_FN[0],'| imgs',len(d2.inline_shapes),'tables',len(d2.tables))
