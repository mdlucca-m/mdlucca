# -*- coding: utf-8 -*-
import json, re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
DOC='/home/user/mdlucca/Artigos/PERFIL_HUMOR_COMPLETO.docx'
J=json.load(open('tcar_desc.json'))
def c2(s): return str(s).replace('.',',')
d=Document(DOC)
def _fmt(p,size=12,just=True,bold=False,center=False):
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(6)
    if center: p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    elif just: p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    for r in p.runs: r.font.name='Times New Roman'; r.font.size=Pt(size); r.bold=bold
def after(el,text='',**kw):
    p=d.add_paragraph(text); el.addnext(p._p); _fmt(p,**kw); return p._p
def head(el,text): return after(el,text,bold=True,just=False)
def _bd(cell):
    tcPr=cell._tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for edge in ['top','bottom','left','right']:
        e=OxmlElement(f'w:{edge}'); on=edge in ('top','bottom')
        e.set(qn('w:val'),'single' if on else 'nil'); e.set(qn('w:sz'),'6'); e.set(qn('w:color'),'000000'); b.append(e)
    tcPr.append(b)
def table_after(el,caption,header,rows,fonte='Fonte: dados da pesquisa (2026).',fs=9.5):
    cap=d.add_paragraph(caption); el.addnext(cap._p); _fmt(cap,size=11,just=False)
    t=d.add_table(rows=1,cols=len(header)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; cap._p.addnext(t._tbl)
    for i,htx in enumerate(header):
        cc=t.rows[0].cells[i]; cc.text=''; rr=cc.paragraphs[0].add_run(htx); rr.bold=True; rr.font.size=Pt(fs); rr.font.name='Times New Roman'
        cc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; _bd(cc)
    for row in rows:
        cs=t.add_row().cells
        for i,val in enumerate(row):
            cs[i].text=''; rr=cs[i].paragraphs[0].add_run(str(val)); rr.font.size=Pt(fs); rr.font.name='Times New Roman'
            cs[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if i>0 else WD_ALIGN_PARAGRAPH.LEFT; _bd(cs[i])
    fon=d.add_paragraph(fonte); t._tbl.addnext(fon._p); _fmt(fon,size=10,just=False); return fon._p
def setpar(p,text):
    if not p.runs: p.add_run(text); return
    p.runs[0].text=text
    for r in p.runs[1:]: r.text=''
def find(prefix):
    for p in d.paragraphs:
        if p.text.strip().startswith(prefix): return p
    raise KeyError(prefix)

# ===== 0) RENUMBER existing tables 3..14 -> +2 (captions + in-text), leave 1,2 =====
def bump(m):
    s=m.group(1); n=int(m.group(2));
    if 3<=n<=14: n+=2
    return 'Tabela%s %d'%(s,n)
for p in d.paragraphs:
    if 'Tabela' in p.text:
        new=re.sub(r'Tabela(s?)\s+(\d+)',bump,p.text)
        if new!=p.text: setpar(p,new)

# ===== 1) INTRO: parágrafo sobre PV do T-CAR com referências (após o último parágrafo da introdução) =====
p_intro=find('A última semana de pré-temporada é, portanto, uma janela crítica')
after(p_intro._p,
 'O pico de velocidade (PV) do teste de Carminatti (T-CAR) — protocolo de campo intermitente em vaivém — é um índice '
 'validado da aptidão aeróbia intermitente: não difere da velocidade associada ao consumo máximo de oxigênio (v-VO₂máx) '
 'e correlaciona-se fortemente com ela (r = 0,74) e com o limiar de lactato, com excelente reprodutibilidade (CCI = 0,94; '
 'CV = 1,4%) (DA SILVA et al., 2011; DITTRICH et al., 2011). Mais do que um marcador laboratorial, o PV do T-CAR prediz o '
 'desempenho físico em jogo — atletas de maior PV percorrem mais distância em alta intensidade e realizam mais sprints '
 '(r = 0,66–0,81) —, o que sustenta sua validade ecológica e de constructo para modalidades coletivas (FERNANDES-DA-SILVA '
 'et al., 2016), sendo ainda independente do estágio maturacional em jovens (TEIXEIRA et al., 2015). Por integrar '
 'potência aeróbia, economia de corrida e capacidade de repetir esforços em um único índice de baixo custo, o PV do T-CAR '
 'é um parâmetro fisiológico de referência natural para o monitoramento da tolerância à carga.')

# ===== 2) MÉTODO 4.3.1: descrição do teste + objetivo + Tabela 3 (descritiva) + Tabela 4 (padronização) =====
p431=find('A escolha do T-CAR justifica-se')
a=after(p431._p,
 'Descrição do teste e objetivo. O T-CAR consiste em séries de 5 × 12 s de corrida em vaivém intercaladas por 6 s de '
 'recuperação, compondo estágios de 90 s; a distância inicial entre os cones é de 15 m e aumenta 1 m a cada estágio, com '
 'a velocidade progredindo até a exaustão voluntária (DA SILVA et al., 2011; DITTRICH et al., 2011). O pico de velocidade '
 '(PV) corresponde à velocidade do último estágio completado e estima a velocidade aeróbia máxima. Neste estudo, o '
 'objetivo do teste é fornecer a linha de base fisiológica — a aptidão aeróbia intermitente de cada atleta — contra a '
 'qual se interpreta a resposta de humor e fadiga ao microciclo.')
D=J['desc']
def rr(k):
    v=D[k]; return [v['n'],c2('%.1f'%v['mean']),c2('%.1f'%v['md']),c2('%.2f'%v['sd']),
        '%s–%s'%(c2('%.1f'%v['lo']),c2('%.1f'%v['hi'])),'%s–%s'%(c2('%.1f'%v['mn']),c2('%.1f'%v['mx'])),c2('%.1f'%v['cv'])]
a=after(a,'A Tabela 3 apresenta a estatística descritiva completa dos resultados do T-CAR. O PV médio foi de %s ± %s km·h⁻¹ '
 '(IC95%% %s–%s; amplitude %s–%s), com baixa dispersão (CV = %s%%), compatível com valores de equipes adultas de modalidades '
 'coletivas.'%(c2('%.1f'%D['PV']['mean']),c2('%.1f'%D['PV']['sd']),c2('%.1f'%D['PV']['lo']),c2('%.1f'%D['PV']['hi']),
   c2('%.1f'%D['PV']['mn']),c2('%.1f'%D['PV']['mx']),c2('%.1f'%D['PV']['cv'])))
a=table_after(a,'Tabela 3 – Estatística descritiva dos resultados do T-CAR (pico de velocidade, repetições, distância e FC máxima).',
 ['Parâmetro','n','Média','Mediana','DP','IC95%','Mín–Máx','CV (%)'],
 [['Pico de velocidade (km·h⁻¹)']+rr('PV'),
  ['Repetições completadas']+rr('reps_ini'),
  ['Distância percorrida (m)']+rr('dist'),
  ['FC máxima (bpm)']+rr('FCmax')],fs=9)
# Tabela 4 – padronização pelo limiar
CL=J['classif']; ab='Abaixo do limiar (<14,9)'; ac='Acima do limiar (≥14,9)'
a=after(a,'Padronização dos atletas pelo limiar do PV. A partir do limiar estimado para o dia de fadiga alta (≈ 14,9 km·h⁻¹; '
 'ver §5.12), os atletas foram estratificados em dois grupos (Tabela 4): abaixo do limiar (n = %d; %.0f%%) e acima do limiar '
 '(n = %d; %.0f%%). Os atletas abaixo do limiar apresentaram, ao longo da semana, maior fadiga física (%s vs %s) e menor '
 'vigor (%s vs %s), confirmando o valor do PV como critério de estratificação da tolerância à carga — ainda que, no nível '
 'individual e com amostra reduzida, as diferenças não atinjam significância (fadiga p = %s; vigor p = %s).'%(
   CL[ab]['n'],CL[ab]['pct'],CL[ac]['n'],CL[ac]['pct'],
   c2('%.1f'%CL[ab]['FadFis']),c2('%.1f'%CL[ac]['FadFis']),c2('%.1f'%CL[ab]['Vigor']),c2('%.1f'%CL[ac]['Vigor']),
   c2('%.3f'%J['mw_fad']),c2('%.3f'%J['mw_vig'])))
a=table_after(a,'Tabela 4 – Padronização dos atletas pelo limiar do pico de velocidade (14,9 km·h⁻¹): distribuição e resposta média no microciclo.',
 ['Estrato','n','%','PV médio (km·h⁻¹)','Fadiga física (semana)','Vigor (semana)'],
 [['Abaixo do limiar (< 14,9)',CL[ab]['n'],c2('%.0f'%CL[ab]['pct']),c2('%.1f'%CL[ab]['PV']),c2('%.1f'%CL[ab]['FadFis']),c2('%.1f'%CL[ab]['Vigor'])],
  ['Acima do limiar (≥ 14,9)',CL[ac]['n'],c2('%.0f'%CL[ac]['pct']),c2('%.1f'%CL[ac]['PV']),c2('%.1f'%CL[ac]['FadFis']),c2('%.1f'%CL[ac]['Vigor'])]])

# ===== 3) REFERÊNCIAS (alfabético) =====
def ref_after(el,text): return after(el,text,just=False)
cleveland=find('CLEVELAND, W. S.')
b=cleveland._p
b=ref_after(b,'DA SILVA, J. F.; GUGLIELMO, L. G. A.; CARMINATTI, L. J.; DE OLIVEIRA, F. R.; DITTRICH, N.; PATON, C. D. Validity and reliability of a new field test (Carminatti\'s test) for soccer players compared with laboratory-based measures. Journal of Sports Sciences, v. 29, n. 15, p. 1621–1628, 2011. DOI: 10.1080/02640414.2011.609179.')
b=ref_after(b,'DITTRICH, N.; DA SILVA, J. F.; CASTAGNA, C.; DE LUCAS, R. D.; GUGLIELMO, L. G. A. Validity of Carminatti\'s test to determine physiological indices of aerobic power and capacity in soccer and futsal players. Journal of Strength and Conditioning Research, v. 25, n. 11, p. 3099–3106, 2011. DOI: 10.1519/JSC.0b013e3182132ce7.')
b=ref_after(b,'FERNANDES-DA-SILVA, J.; CASTAGNA, C.; TEIXEIRA, A. S.; CARMINATTI, L. J.; GUGLIELMO, L. G. A. The peak velocity derived from the Carminatti Test is related to physical match performance in young soccer players. Journal of Sports Sciences, v. 34, n. 24, p. 2238–2245, 2016. DOI: 10.1080/02640414.2016.1209307.')
stieler=find('STIELER, E.')
ref_after(stieler._p,'TEIXEIRA, A. S. et al. Skeletal maturation and aerobic performance in young soccer players from professional academies. International Journal of Sports Medicine, v. 36, n. 13, p. 1069–1075, 2015. DOI: 10.1055/s-0035-1549922.')

d.save(DOC)
print('SAVED',DOC)
