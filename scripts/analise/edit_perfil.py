# -*- coding: utf-8 -*-
import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SRC='/root/.claude/uploads/e1dba24c-b1d7-5908-9106-f2f4aaf3f56a/557fba4a-PERFIL_HUMOR_COMPLETO.docx'
OUT='/home/user/mdlucca/Artigos/PERFIL_HUMOR_COMPLETO.docx'
FIG='/home/user/mdlucca/Artigos/figuras/x_normalizacao.png'
N=json.load(open('normalizacao.json'))
M=json.load(open('tcar2_models.json'))
def c2(s): return str(s).replace('.',',')
d=Document(SRC)

# ---------- helpers: insert paragraphs AFTER a given xml element ----------
def _fmt(p,size=12,just=True,bold=False,center=False):
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(6)
    if center: p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    elif just: p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    for r in p.runs: r.font.name='Times New Roman'; r.font.size=Pt(size); r.bold=bold
def new_after(anchor_el,text='',**kw):
    p=d.add_paragraph(text); anchor_el.addnext(p._p); _fmt(p,**kw); return p._p
def head_after(anchor_el,text):
    return new_after(anchor_el,text,bold=True,just=False)
def fig_after(anchor_el,path,caption,fonte='Fonte: elaboração dos autores (2026).',w=16.0):
    pimg=d.add_paragraph(); anchor_el.addnext(pimg._p); pimg.alignment=WD_ALIGN_PARAGRAPH.CENTER
    pimg.add_run().add_picture(path,width=Cm(w)); pimg.paragraph_format.space_before=Pt(6)
    cap=d.add_paragraph(caption); pimg._p.addnext(cap._p); _fmt(cap,size=11,just=False,center=True)
    fon=d.add_paragraph(fonte); cap._p.addnext(fon._p); _fmt(fon,size=10,just=False,center=True)
    return fon._p

paras=d.paragraphs
def find(prefix):
    for p in paras:
        if p.text.strip().startswith(prefix): return p
    raise KeyError(prefix)

# ================= 1) CORRIGIR: completar a lista truncada em 4.5 (termina em "(7)") =================
p45=find('(1) Estatística descritiva exploratória')
add=(' (7) Padronização (escore z) das subescalas para a classificação de perfis e ajuste/normalização '
     'alométrica dos parâmetros do T-CAR 2 pela massa corporal (ver 4.5.1). (8) Suavização das médias diárias '
     'por regressão local ponderada (LOWESS) para separar sinal de ruído. (9) Modelagem preditiva de alta fadiga '
     'pelos parâmetros fisiológicos com validação cruzada (leave-one-out) e curvas ROC aparente e validada.')
# append text to the last run so o parágrafo deixa de terminar em "(7)"
run=p45.runs[-1]; run.text=run.text.rstrip()
if run.text.endswith('(7)'): run.text=run.text[:-3].rstrip()  # remove o "(7)" truncado
run.text=run.text+add

# ================= 2) INCLUIR seção de método 4.5.1 (após 4.5) =================
anchor=p45._p
anchor=head_after(anchor,'4.5.1 Normalização, ajuste alométrico e suavização — o que são e por que aplicá-las')
anchor=new_after(anchor,'Três operações preparam os dados antes das inferências, cada uma com uma razão metodológica específica e respaldo na literatura.')
anchor=new_after(anchor,
 '(a) Padronização (escore z) das subescalas. A classificação nos seis perfis de Terry baseia-se na menor distância '
 'euclidiana de cada observação aos protótipos. Como as subescalas têm variâncias muito diferentes — a fadiga varia '
 'cerca de três vezes mais do que a confusão (Figura 19A) —, usar os escores brutos faria as subescalas de maior '
 'amplitude dominarem a distância e distorcerem o perfil atribuído. A padronização (média 0, desvio 1) coloca todas as '
 'dimensões em pé de igualdade, condição necessária para uma classificação de perfil válida e comparável entre atletas '
 '(PARSONS-SMITH; TERRY; MACHIN, 2017).')
anchor=new_after(anchor,
 '(b) Normalização alométrica do parâmetro fisiológico. Medidas fisiológicas dependem do tamanho corporal; compará-las '
 'entre atletas de massas diferentes por razão simples (por quilograma) introduz viés sistemático, pois a relação entre '
 'a variável e a massa raramente é de proporção direta (NEVILL; RAMSBOTTOM; WILLIAMS, 1992; JARIC, 2002). O ajuste '
 'alométrico Y = a·massaᵇ estima o expoente b que torna a medida independente do tamanho, de modo que a medida '
 'normalizada (Y/massaᵇ) deixa de se correlacionar com a massa. No futebol e em outras modalidades coletivas, esse '
 'procedimento é padrão para comparar a aptidão entre posições e indivíduos sem o confundimento do porte físico '
 '(CHAMARI et al., 2005; OBA et al., 2014). Aqui ele permite testar a aptidão “limpa” do efeito de massa como possível '
 'diferenciador de quem fadiga mais.')
anchor=new_after(anchor,
 '(c) Suavização das curvas (LOWESS). O sinal do microciclo é de baixa frequência e pequeno em relação ao ruído de '
 'medida do dia a dia — apenas uma fração da variância diária é sinal de estado, o restante é oscilação. A regressão '
 'local ponderada (LOWESS) estima a tendência sistemática sem impor forma paramétrica, separando o sinal (a deriva '
 'semanal de vigor e fadiga) do ruído (a oscilação aleatória em torno do piso das subescalas negativas) (CLEVELAND, '
 '1979). Essa filtragem é o que sustenta afirmar, com segurança, que a migração do perfil iceberg para o perfil de '
 'fadiga é tendência real e não artefato de medida — em linha com a recomendação de monitorar o atleta por tendência '
 'referenciada à sua linha de base individual (SAW; MAIN; GASTIN, 2016).')

# ================= 3) INCLUIR seção de resultados 5.11 (após a Fonte da Figura 18, antes de "6 SÍNTESE") =================
# âncora: a linha "Fonte..." logo após a Figura 18 (par. imediatamente antes de "6 SÍNTESE DOS PRINCIPAIS ACHADOS")
p6=find('6 SÍNTESE DOS PRINCIPAIS ACHADOS')
# inserir antes de p6 => usar o elemento anterior como âncora
prev=p6._p.getprevious()
anchor=head_after(prev,'5.11 Normalização, ajuste alométrico e suavização das curvas')
anchor=new_after(anchor,
 'A Figura 19 consolida as três operações e suas consequências práticas. (A) As subescalas do BRUMS têm desvios-padrão '
 'muito desiguais (fadiga = %s; confusão = %s), o que justifica a padronização antes da classificação de perfis: sem '
 'ela, fadiga e vigor dominariam a distância aos protótipos e enviesariam o perfil atribuído.'%(
   c2('%.1f'%N['sds']['Fadiga']),c2('%.1f'%N['sds']['Confusão'])))
anchor=new_after(anchor,
 '(B) A normalização alométrica pela massa remove o confundimento do tamanho na fadiga física: a correlação massa–fadiga '
 'cai de |ρ| = %s para %s após o ajuste, isolando o efeito da aptidão do efeito do porte. Aplicada ao PV do T-CAR 2, a '
 'normalização quase não altera a associação (|ρ| %s → %s), porque o PV pós-bloco já se correlaciona fracamente com a '
 'fadiga — resultado coerente com o expoente alométrico instável (b = %s) e com o nulo robusto da modelagem (§5.10).'%(
   c2('%.2f'%N['mass_raw']),c2('%.2f'%N['mass_norm']),c2('%.2f'%N['pv_raw']),c2('%.2f'%N['pv_norm']),
   c2('%+.2f'%M['ALLO']['FadFisica']['b'])))
anchor=new_after(anchor,
 '(C) A suavização por LOWESS extrai, da nuvem de observações, a tendência de queda do vigor e de subida da fadiga ao '
 'longo da semana — ou seja, o movimento do eixo energia–fadiga é sinal de baixa frequência, não ruído. (D) A tendência '
 'suavizada do índice iceberg cai de forma sistemática, cruza o zero por volta do Dia 3 e despenca no Dia 7 — a leitura '
 'contínua da mesma migração iceberg→fadiga descrita pelas frequências de perfil (§5.2). Em conjunto, as três operações '
 'não são passos meramente técnicos: a padronização torna a classificação de perfil válida, a normalização alométrica '
 'isola a aptidão do tamanho corporal, e a suavização confirma que a deterioração do humor ao longo do microciclo é '
 'sinal e não ruído.')
anchor=fig_after(anchor,FIG,
 'Figura 19 – Normalização e suavização. (A) Desvios-padrão desiguais das subescalas justificam a padronização (escore z). '
 '(B) A normalização alométrica remove o confundimento do tamanho (massa: |ρ| 0,33→0,06; PV do T-CAR 2: 0,14→0,08). '
 '(C) A suavização (LOWESS) separa o sinal semanal (vigor ↓, fadiga ↑) do ruído. (D) A tendência suavizada do índice '
 'iceberg cai ao longo da semana — a mesma migração iceberg→fadiga em leitura contínua.')

# adicionar ponteiro em 5.10 para 5.11
p510=find('Testou-se se os parâmetros do T-CAR 2')
r=p510.add_run(' O detalhamento da normalização, do ajuste alométrico e da suavização consta na seção 5.11.')
r.font.name='Times New Roman'; r.font.size=Pt(12)

# ================= 4) REFERÊNCIAS (inserir em ordem alfabética) =================
refBRANDT=find('BRANDT, R.')
refMORGAN=find('MORGAN, W. P.')
# após BRANDT: CHAMARI, CLEVELAND, JARIC
a=refBRANDT._p
a=new_after(a,'CHAMARI, K. et al. Endurance training and testing with the ball in young elite soccer players. British Journal of Sports Medicine, v. 39, n. 1, p. 24–28, 2005. DOI: 10.1136/bjsm.2003.009985.',just=False)
a=new_after(a,'CLEVELAND, W. S. Robust locally weighted regression and smoothing scatterplots. Journal of the American Statistical Association, v. 74, n. 368, p. 829–836, 1979. DOI: 10.1080/01621459.1979.10481038.',just=False)
a=new_after(a,'JARIC, S. Muscle strength testing: use of normalisation for body size. Sports Medicine, v. 32, n. 10, p. 615–631, 2002. DOI: 10.2165/00007256-200232100-00002.',just=False)
# após MORGAN: NEVILL, OBA
b=refMORGAN._p
b=new_after(b,'NEVILL, A. M.; RAMSBOTTOM, R.; WILLIAMS, C. Scaling physiological measurements for individuals of different body size. European Journal of Applied Physiology and Occupational Physiology, v. 65, n. 2, p. 110–117, 1992. DOI: 10.1007/BF00705066.',just=False)
b=new_after(b,'OBA, Y. et al. Allometric scaling of strength scores in NCAA Division I-A football athletes. Journal of Strength and Conditioning Research, v. 28, n. 12, p. 3330–3337, 2014. DOI: 10.1519/JSC.0000000000000548.',just=False)

d.save(OUT)
print('SAVED',OUT)
