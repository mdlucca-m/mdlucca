import warnings; warnings.filterwarnings('ignore')
import numpy as np, pandas as pd, json, os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
SC='/tmp/claude-0/-home-user-mdlucca/e1dba24c-b1d7-5908-9106-f2f4aaf3f56a/scratchpad'
FIG='/home/user/mdlucca/Artigos/figuras'
amp=json.load(open(f'{SC}/amp_docx.json'))          # 4-var: grp/btw/per
PP=json.load(open(f'{SC}/prepost.json'))
DN=json.load(open(f'{SC}/denoise.json'))
PF=json.load(open(f'{SC}/pair_fatigue.json'))       # física×mental
PV=json.load(open(f'{SC}/pair_vigor_tmd.json'))     # vigor×TMD
NORM=json.load(open(f'{SC}/norm_data.json'))
APP=json.load(open(f'{SC}/app_data.json'))
grp=amp['grp']; btw=amp['btw']; per=amp['per']

doc=Document()
# ---- styles ----
st=doc.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(10.5)
for i in range(1,4):
    try: doc.styles[f'Heading {i}'].font.color.rgb=RGBColor(0x1a,0x1a,0x1a)
    except Exception: pass
def H(t,l=1): doc.add_heading(t,level=l)
def P(t,it=False,b=False,center=False,size=None):
    p=doc.add_paragraph(); r=p.add_run(t); r.italic=it; r.bold=b
    if size: r.font.size=Pt(size)
    if center: p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY if not center else p.alignment
    return p
def fig(name,cap,w=6.3):
    path=f'{FIG}/{name}.png'
    if os.path.exists(path):
        doc.add_picture(path,width=Inches(w)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
        c=doc.add_paragraph(); r=c.add_run(cap); r.italic=True; r.font.size=Pt(8.5); c.alignment=WD_ALIGN_PARAGRAPH.CENTER
    else: print('MISSING FIG',name)
def table(headers,rows,cap=None,fs=8.0):
    if cap: c=doc.add_paragraph(); r=c.add_run(cap); r.bold=True; r.font.size=Pt(9)
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Light Grid Accent 1'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        cell=t.rows[0].cells[i]; cell.paragraphs[0].add_run(str(h)).bold=True; cell.paragraphs[0].runs[0].font.size=Pt(fs)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=str(v)
            if cells[i].paragraphs[0].runs: cells[i].paragraphs[0].runs[0].font.size=Pt(fs)
    doc.add_paragraph()

V4=[('Vigor','Vigor'),('Fadiga','Fadiga BRUMS'),('TMD','PTH/TMD'),('FadMental','Fadiga mental')]
FF=PF['grp']['FadFisica']  # fadiga física stats

# ================= TÍTULO =================
ti=doc.add_heading('',0)
ti.add_run('Sinal, ruído e amplitude no monitoramento psicométrico da fadiga e do humor durante um microciclo de choque de HIIT em atletas de handebol de elite')
ti.alignment=WD_ALIGN_PARAGRAPH.CENTER
P('[Autoria e afiliação a preencher]',center=True,it=True,size=10)
P('Estudo observacional longitudinal · 27 atletas · 456 observações · janela de 21 a 28 de abril de 2024 · reanálise estatística independente em Python. Atletas anonimizados (A01–A27).',center=True,it=True,size=9)
doc.add_paragraph()

# ================= RESUMO =================
H('Resumo')
P('Contexto e objetivo. O monitoramento psicométrico de atletas confunde rotineiramente flutuação real (sinal) com erro de medida (ruído). Este estudo quantifica a amplitude das respostas de humor e fadiga a um microciclo de choque de HIIT e separa, para cada variável, quanto da oscilação é sinal do microciclo, quanto é traço estável e quanto é ruído — determinando quando a variação observada é interpretável no grupo e no atleta individual.', b=False)
P('Métodos. Vinte e sete atletas de handebol de elite responderam ao BRUMS e a autoavaliações de fadiga física e mental (0–10) em três momentos diários (pré, intra-sessão [mid] e pós) ao longo de sete dias com três sessões de HIIT 4×4 min a 104% do pico de velocidade do T-CAR. A variância de cada variável foi decomposta (modelo atleta + dia) em sinal do microciclo, traço e ruído; derivaram-se o erro típico de medida (ETM), a mudança mínima detectável (MDC95, individual e de grupo) e a menor mudança relevante (SWC). O efeito da remoção do ruído foi avaliado por agregação, suavização e desatenuação pela confiabilidade.')
P('Resultados. As variáveis de humor são não normais (14/18; Shapiro-Wilk p < 0,01), impondo métodos não paramétricos. O sinal do microciclo é a menor fração da variância (1–12%), dominada por traço (32–73%) e ruído (26–64%); a fadiga física concentra o maior sinal (12%; SNR 2,02; dz D1→D7 +1,74). No grupo, a oscilação semanal supera o MDC de grupo por 4–10×, mas no indivíduo fica abaixo do MDC95 de uma coleta — a decisão individual exige médias de ≥3–5 coletas. Removido o ruído, as associações do eixo energia–fadiga fortalecem-se (Vigor×Fadiga −0,44→−0,93 no nível-grupo) e os efeitos crescem (Vigor dz −0,95→−1,25), sem criar sinal onde há efeito de piso. Contrastes entre pares: fadiga física e mental acoplam-se no momento (r 0,50→0,78 desatenuada) mas dissociam na trajetória semanal (0,37); Vigor e PTH/TMD são espelhos inversos (−0,62→−0,78), com o TMD adicionando pouca informação além do Vigor.')
P('Conclusão. A deterioração do humor no microciclo é real, não linear e concentrada no eixo energia–fadiga, com a fadiga física como marcador mais precoce, sensível e confiável. Removido o ruído, as conclusões ficam mais nítidas — não diferentes. A decisão individual requer agregação de coletas; a fadiga mental e as subescalas de piso não servem de marcadores de carga aguda deste microciclo.')
kp=doc.add_paragraph(); kp.add_run('Palavras-chave: ').bold=True; kp.add_run('monitoramento da carga; BRUMS; erro típico de medida; mudança mínima detectável; relação sinal-ruído; HIIT; handebol.')
doc.add_paragraph()

# ================= 1. INTRODUÇÃO =================
H('1. Introdução')
P('O monitoramento psicométrico do estado de humor e da fadiga é amplamente empregado para individualizar a carga e prevenir o excesso de treinamento. Instrumentos como o Brunel Mood Scale (BRUMS) e escalas de autoavaliação de fadiga são rápidos e sensíveis, mas partilham um problema raramente enfrentado de forma explícita: uma parte substancial da variação dia a dia de um atleta não reflete o efeito do treino, e sim erro de medida e flutuação aleatória. Interpretar essa flutuação como resposta ao treino gera decisões espúrias.')
P('Três quantidades distintas coexistem em cada medida: (i) o sinal do microciclo — a mudança sistemática que os dias de treino impõem; (ii) o traço — as diferenças estáveis entre atletas; e (iii) o ruído — erro de medida e variação intraindividual não sistemática. Separá-las é pré-requisito para saber qual variação merece ação. As ferramentas psicométricas para tal — erro típico de medida (ETM), mudança mínima detectável (MDC95) e menor mudança relevante (SWC) — são conhecidas, mas raramente aplicadas de forma integrada a um microciclo real, e menos ainda para responder à pergunta prática: "removido o ruído, como ficam as análises?".')
P('Este estudo aplica essa decomposição a um microciclo de choque de HIIT em atletas de handebol de elite. Os objetivos são: (1) descrever a amplitude (range) das respostas de humor e fadiga no grupo e no atleta; (2) decompor a variância de cada variável em sinal do microciclo, traço e ruído; (3) estabelecer limiares de detectabilidade (MDC individual vs de grupo) e determinar quando a variação é interpretável; (4) demonstrar o efeito da remoção do ruído sobre associações e tamanhos de efeito; e (5) contrastar três pares de variáveis do eixo afetivo-energético — o eixo energia–fadiga completo, fadiga física × mental e Vigor × PTH/TMD.')

# ================= 2. MÉTODOS =================
H('2. Materiais e Métodos')
H('2.1. Amostra e delineamento',2)
ag=APP['socio']['agg']; pc=APP['socio']['pos_counts']
P(f"Vinte e sete atletas de handebol de elite do sexo masculino (idade {ag['idade'][0]:.0f} ± {ag['idade'][1]:.0f} anos, faixa {ag['idade'][2]}–{ag['idade'][3]}; experiência {ag['exp'][0]:.1f} ± {ag['exp'][1]:.1f} anos; estatura {ag['estatura'][0]:.1f} ± {ag['estatura'][1]:.1f} cm; massa {ag['massa'][0]:.1f} ± {ag['massa'][1]:.1f} kg; gordura {ag['pG'][0]:.1f} ± {ag['pG'][1]:.1f}%; pico de velocidade no T-CAR {ag['PV'][0]:.2f} ± {ag['PV'][1]:.2f} km/h) foram monitorados durante um microciclo de choque de sete dias (21–28/04/2024), totalizando 456 observações. As posições distribuíram-se em armadores (11), alas (9), pivôs (4) e goleiros (3). O delineamento é observacional longitudinal, com medidas repetidas aninhadas (observações ⊂ dias ⊂ atletas).")
H('2.2. Instrumentos e carga',2)
el=APP['extload']
P(f"O humor foi avaliado pelo BRUMS (subescalas Tensão, Depressão, Raiva, Vigor, Fadiga e Confusão) e sua Perturbação Total do Humor (PTH/TMD); a fadiga física e a fadiga mental por escalas de autoavaliação de 0–10. Coletas em três momentos: pré-sessão, intra-sessão (mid) e pós-sessão. As sessões de HIIT (dias 2, 4 e 7) consistiram em 4×4 min a 104% do pico de velocidade do T-CAR. A carga externa, derivada da prescrição relativa, foi fixa entre sessões (velocidade média {el['vel_mean']:.1f} ± {el['vel_sd']:.1f} km/h; distância {el['dist_mean']:.0f} ± {el['dist_sd']:.0f} m/sessão; total ≈ {el['dist_total']/1000:.1f} km).")
H('2.3. Decomposição sinal–traço–ruído e limiares',2)
P('A variância de cada variável foi decomposta por um modelo de dois fatores (y ~ atleta + dia, soma de quadrados tipo I) em três parcelas somando 100%: sinal do microciclo (efeito do dia), traço (entre atletas) e ruído (resíduo). Do ruído derivaram-se: o erro típico de medida ETM = √(QM residual); a mudança mínima detectável individual MDC95 = 1,96·√2·ETM; a MDC95 de grupo = MDC95/√n (n ≈ 27/dia); e a menor mudança relevante SWC = 0,2·DP-entre-atletas. A relação sinal-ruído (SNR) foi definida como a amplitude do sinal semanal dividida pelo ETM. A confiabilidade de uma medida foi estimada por r_xx = 1 − fração de ruído.')
H('2.4. Remoção do ruído',2)
P('O efeito de "remover o ruído" foi avaliado por três operações equivalentes: (a) agregação (a média de n coletas reduz o ruído por √n); (b) suavização da trajetória (LOWESS); e (c) desatenuação das correlações e tamanhos de efeito pela confiabilidade (r_verdadeiro = r_observado/√(r_xx·r_yy); dz_corrigido = dz/√r_xx).')
H('2.5. Análise estatística',2)
P('A normalidade foi testada por Shapiro–Wilk (com assimetria e curtose). Dada a não normalidade das variáveis de humor, adotou-se a via não paramétrica: descritivas por mediana [IQR]; associações por Spearman (ρ) e correlação de medidas repetidas (rmcorr); contrastes intra-sujeito por Wilcoxon e Friedman; entre dias, por passos da média e dispersão; tamanhos de efeito por dz pareado com IC bootstrap. A unidade amostral para contrastes é o atleta (agregação), corrigindo a pseudorreplicação. α = 0,05. Software: Python 3 (numpy, scipy, statsmodels). Toda a análise é reproduzível pelos scripts em scripts/analise/.')

# ================= 3. RESULTADOS =================
H('3. Resultados')

H('3.1. Normalidade e descritivas',2)
dist=[d for d in NORM['dist'] if d['level']=='obs']
rows=[]
for d in dist[:9]:
    rows.append([d['lab'],d['n'],f"{d['W']:.3f}",(f"{d['p']:.1e}"),f"{d['skew']:+.2f}",('não' if not d['normal'] else 'sim')])
table(['Variável','n','Shapiro W','p','Assimetria','Normal?'],rows,'Tabela 1 — Teste de normalidade (Shapiro–Wilk) das distribuições brutas')
P('Quatorze das dezoito variáveis são não normais (todas as de humor/bem-estar; p de 3×10⁻⁴ a 3×10⁻³⁵); apenas velocidade, distância, T-CAR PV e CMJ são normais. As subescalas negativas concentram 50–80% das respostas no piso (zero), o que fundamenta a via não paramétrica e antecipa sua baixa responsividade.')

H('3.2. Amplitude',2)
rows=[]
for v,lab in V4:
    g=grp[v]; p=per[v]
    rows.append([lab,f"{g['sig']:.2f}",f"{p['amp_mean']:.2f} ± {p['amp_sd']:.2f}"])
rows.insert(0,['Fadiga física',f"{FF['sig']:.2f}",f"{PF['per']['FadFisica']['amp_mean']:.2f} ± {PF['per']['FadFisica']['amp_sd']:.2f}"])
table(['Variável','Amplitude do sinal (grupo)','Amplitude intra-atleta (média ± DP)'],rows,'Tabela 2 — Amplitude do sinal do microciclo vs amplitude individual')
P('Em todas as variáveis a amplitude do sinal (a trajetória média do grupo) é muito menor que a amplitude intra-atleta: o atleta oscila 2–3× mais que o microciclo médio, porque sua oscilação carrega o próprio ruído. A hierarquia amplitude-do-sinal < amplitude-intra-atleta < amplitude-total é universal.')
fig('adx1_grupo_semana','Figura 1 — Trajetória do grupo ao longo da semana (média ± DP; LOWESS pontilhado; linhas vermelhas = dias de HIIT). Vigor cai; Fadiga e PTH/TMD sobem com precipitação no Dia 7; Fadiga mental permanece plana.')

H('3.3. Decomposição sinal–traço–ruído',2)
rows=[['Fadiga física',f"{FF['pday']:.1f}%",f"{FF['ptrait']:.1f}%",f"{FF['pnoise']:.1f}%",f"{FF['etm']:.2f}",f"{FF['snr']:.2f}"]]
for v,lab in V4:
    g=grp[v]; rows.append([lab,f"{g['pday']:.1f}%",f"{g['ptrait']:.1f}%",f"{g['pnoise']:.1f}%",f"{g['etm']:.2f}",f"{g['snr']:.2f}"])
table(['Variável','% Sinal (dia)','% Traço','% Ruído','ETM','SNR'],rows,'Tabela 3 — Decomposição da variância e relação sinal-ruído')
P('O sinal do microciclo é sempre a menor fatia da variância (1–12%); o grosso é traço estável (32–73%) e ruído (26–64%). A fadiga física lidera em sinal (12%) e SNR (2,02); a fadiga mental é quase pura estrutura de traço (73%; SNR 0,62 < 1). As subescalas negativas são dominadas por ruído.')
fig('amp_a_variancia','Figura 2 — Partição da variância (%) por variável: sinal do microciclo (vermelho) vs traço (azul) vs ruído (cinza).')

H('3.4. Trajetória semanal e variação entre dias',2)
P('O maior passo da trajetória concentra-se no fim da semana (D6→D7) para a Fadiga e o PTH/TMD — a precipitação do Dia 7. A dispersão entre atletas mantém-se elevada em todos os dias (heterogeneidade intrínseca, coerente com ICC de traço 0,3–0,7), sem um único dia dominante. A forma da trajetória é não linear: componentes linear e cúbica significativas, com crescimento cúbico vencendo por AIC/LRT (subida → alívio no meio da semana → precipitação no D7).')

H('3.5. Curso intra-sessão (pré → mid → pós)',2)
# representative: FadFisica and Fadiga
for v,lab in [('FadFisica','Fadiga física'),('Fadiga','Fadiga BRUMS')]:
    src=PF['pp'][v] if v=='FadFisica' else PP[v]
    rows=[]
    for d in range(1,8):
        r=src[str(d)]
        pstr='—' if r['p'] is None else (('<0.001*' if r['p']<0.001 else f"{r['p']:.3f}"+('*' if r['p']<0.05 else '')))
        mids='—' if r['mid_m'] is None else f"{r['mid_m']:.2f}"
        rows.append([f"D{d}"+(" H" if r['hiit'] else ""),f"{r['pre_m']:.2f}",mids,f"{r['pos_m']:.2f}",f"{r['delta']:+.2f}",pstr])
    table(['Dia','Pré','Mid','Pós','Δ(pós−pré)','p'],rows,f'Tabela {"4" if v=="FadFisica" else "5"} — {lab}: pré → mid → pós por dia (H = HIIT; * p<0,05)',fs=8)
P('A resposta aguda pré→pós é maior no Dia 1 (choque inicial) e reaparece no fim da semana. O momento intermediário (mid) revela que o curso intra-sessão não é monotônico: a fadiga frequentemente pica no meio da sessão e recua um pouco ao final. A maioria dos dias isolados não atinge significância pareada — o sinal robusto está no acúmulo semanal, não na sessão isolada.')
fig('adx6_prepos_dia','Figura 3 — Pré → mid → pós por dia para as quatro variáveis do eixo energia–fadiga (* Wilcoxon pareado pré×pós p<0,05).')

H('3.6. Variabilidade individual',2)
fig('adx3_heatmap_atleta','Figura 4 — Amplitude semanal por atleta (z por variável). Alguns atletas são consistentemente estáveis; outros amplificam em todas as variáveis — a heterogeneidade é estrutural.')

H('3.7. Detectabilidade em dois níveis (grupo vs indivíduo)',2)
rows=[['Fadiga física',f"{FF['sig']:.2f}",f"{FF['mdc']:.2f}",('SIM' if FF['sig']>FF['mdc'] else 'NÃO'),f"{FF['mdc_g']:.2f}",'SIM']]
for v,lab in V4:
    g=grp[v]; rows.append([lab,f"{g['sig']:.2f}",f"{g['mdc']:.2f}",('SIM' if g['detect_ind'] else 'NÃO'),f"{g['mdc_g']:.2f}",('SIM' if g['detect_grp'] else 'NÃO')])
table(['Variável','Ampl. sinal','MDC95 individual','> ind.?','MDC95 grupo','> grupo?'],rows,'Tabela 6 — A oscilação semanal supera o ruído? Veredito em dois níveis')
P('No grupo, a oscilação semanal supera o MDC de grupo por 4–10× — o efeito do microciclo é sinal real e robusto na média. No indivíduo, a mesma oscilação fica abaixo do MDC95 de uma coleta: não é contradição, é o ganho √n da média de 27 atletas. A decisão individual exige médias de ≥3–5 coletas (Φ ≥ 0,80).')

H('3.8. Efeito da remoção do ruído',2)
Robs=pd.DataFrame(DN['Robs']); Rden=pd.DataFrame(DN['Rden'])
rows=[]
for x,y,l in [('Vigor','Fadiga','Vigor × Fadiga'),('Fadiga','TMD','Fadiga × TMD'),('Vigor','TMD','Vigor × TMD'),('FadMental','Fadiga','Fadiga mental × Fadiga')]:
    rows.append([l,f"{Robs.loc[x,y]:+.2f}",f"{Rden.loc[x,y]:+.2f}"])
table(['Par','r observado (com ruído)','r denoised (nível-grupo)'],rows,'Tabela 7 — Correlações antes e depois de remover o ruído')
P('Removido o ruído, as associações do eixo energia–fadiga fortalecem-se: Vigor × Fadiga passa de −0,44 para −0,93 no nível-grupo. As correlações brutas eram atenuadas pelo ruído de medida — não porque as relações fossem fracas, mas porque cada medida é imprecisa.')
fig('den1_corr','Figura 5 — Matriz de correlações do eixo energia–fadiga: observada (com ruído) vs denoised (nível-grupo).')
dzo={'Vigor':-0.95,'Fadiga':0.72,'TMD':0.41,'FadMental':0.33,'FadFisica':1.74}; dzc={'Vigor':-1.25,'Fadiga':0.91,'TMD':0.51,'FadMental':0.38,'FadFisica':2.36}
rows=[['Fadiga física',f"{dzo['FadFisica']:+.2f}",f"{dzc['FadFisica']:+.2f}"]]+[[lab,f"{dzo[v]:+.2f}",f"{dzc[v]:+.2f}"] for v,lab in V4]
table(['Variável','dz observado','dz corrigido (ruído removido)'],rows,'Tabela 8 — Efeito D1→D7 corrigido pela atenuação do ruído')
P('Corrigido pela confiabilidade, o efeito acumulado cresce: a fadiga física de +1,74 para +2,36 e o Vigor de −0,95 para −1,25. O ruído estava encolhendo os efeitos observados; a magnitude verdadeira do custo do microciclo é maior que a leitura bruta. Contudo, remover o ruído não cria sinal onde há efeito de piso: as subescalas negativas permanecem planas.')
fig('den3_mdc_k','Figura 6 — MDC95 individual em função do número de coletas (k); a linha branca é a amplitude do sinal semanal. A partir de ~5 coletas o Vigor e a Fadiga tornam-se individualmente detectáveis.')

H('3.9. Contrastes entre pares de variáveis',2)
P('Fadiga física × fadiga mental. As duas dissociam ao longo da semana: a física quase dobra (dz +1,74) enquanto a mental fica plana (dz +0,33). No momento, porém, acoplam-se quando removido o erro de medida (r 0,50 → 0,78 desatenuada; rmcorr 0,44); na trajetória semanal, a correlação cai para 0,37. Num dado instante, quem está fisicamente exausto tende a relatar mais fadiga mental, mas o microciclo, como processo, carrega apenas a física — o custo é somático, não cognitivo.')
fig('pf2_dissociacao','Figura 7 — Fadiga física vs mental (LOWESS): a física acumula; a mental é estável — dissociação somática vs cognitiva.')
P('Vigor × PTH/TMD. São espelhos inversos do mesmo eixo: r −0,62 (bruto) → −0,78 (nível-grupo); a desatenuação atinge −1,01 (caso Heywood, pois o Vigor compõe o TMD com sinal invertido). O TMD tem a maior amplitude bruta mas o pior SNR (1,04 vs 1,36 do Vigor), por diluir o eixo energia–fadiga no ruído das negativas de piso. Conclusão metodológica: o TMD não é independente do Vigor e acrescenta pouca informação além de Vigor + Fadiga, com mais ruído.')
fig('pv2_dissociacao','Figura 8 — Vigor vs PTH/TMD (z): curvas espelhadas — acoplamento inverso do eixo energia–humor.')

# ================= 4. DISCUSSÃO =================
H('4. Discussão')
P('Os resultados sustentam três mensagens centrais. Primeiro, a deterioração do humor no microciclo é real e não linear, mas ocupa uma fração pequena da variância total (1–12%): a maior parte do que se observa no dia a dia de um atleta é traço estável e ruído. Ignorar essa proporção leva a superinterpretar oscilações que são, em grande medida, aleatórias.')
P('Segundo, a detectabilidade depende do nível de análise. A média do grupo remove o ruído por √n e torna o sinal semanal inequívoco; a leitura de um único atleta com uma única coleta, ao contrário, fica abaixo do MDC95 e é, portanto, não confiável. A solução prática não é abandonar o monitoramento individual, e sim agregá-lo: médias de três a cinco coletas trazem o ruído para baixo da amplitude do sinal e restauram a validade da decisão individual — o que reconcilia a utilidade clínica do instrumento com sua confiabilidade psicométrica.')
P('Terceiro, remover o ruído esclarece o fenômeno sem alterá-lo. As associações do eixo energia–fadiga, uma vez desatenuadas, aproximam-se de uma dimensão bipolar quase perfeita, e os tamanhos de efeito crescem — a leitura bruta subestimava o custo do microciclo. Mas o denoising não resgata as subescalas de piso, cuja não resposta é um fenômeno real (efeito piso), não um artefato de ruído. A fadiga física emerge, em todas as análises, como o marcador mais precoce, sensível e confiável, ao passo que a fadiga mental e o PTH/TMD isolado são, respectivamente, não responsivo e redundante para este eixo.')
P('Do ponto de vista aplicado, recomenda-se monitorar a fadiga física (e o Vigor) com médias curtas de coletas, usar limiares individuais baseados no MDC, ler a resposta do grupo livremente e evitar sobreinterpretar variações individuais isoladas ou variáveis de piso. O TMD pode ser substituído, sem perda de informação, pelo par Vigor + Fadiga, com menor ruído.')

# ================= 5. LIMITAÇÕES =================
H('5. Limitações')
P('O estudo é observacional e restrito a uma janela de sete dias, sem grupo-controle; a carga externa fixa (prescrição relativa) impede a análise dose-resposta e confunde o efeito do HIIT com o volume acumulado. Os seis perfis de humor dependem de padronização intra-amostral (sem normas de escore-T). As estimativas de confiabilidade e de ruído assumem o modelo aditivo atleta + dia; estruturas de erro mais complexas (autocorrelação intraindividual) não foram modeladas. A amostra, embora de elite, é de tamanho moderado (n = 27), o que limita a potência das análises de moderação por subgrupos.')

# ================= 6. CONCLUSÃO =================
H('6. Conclusão')
P('Num microciclo de choque de HIIT, a resposta do humor é real, não linear e concentrada no eixo energia–fadiga, com a fadiga física como marcador dominante. A maior parte da oscilação observada, porém, é traço e ruído; separá-los mostra que o sinal do microciclo é robusto no grupo mas exige agregação de coletas para a decisão individual. Removido o ruído, as conclusões do estudo tornam-se mais nítidas — associações mais fortes, efeitos maiores — sem mudar de direção, e sem manufaturar sinal onde há efeito de piso. Recomenda-se centrar o monitoramento na fadiga física e no Vigor, com limiares de MDC e médias de ≥3–5 coletas para o nível individual.')

# ================= REFERÊNCIAS =================
H('Referências')
refs=[
 'Terry, P. C., Lane, A. M., & Fogarty, G. J. (2003). Construct validity of the POMS-A for use with adults. Psychology of Sport and Exercise, 4(2), 125–139.',
 'Parsons-Smith, R. L., Terry, P. C., & Machin, M. A. (2017). Identification and description of novel mood profile clusters. Frontiers in Psychology, 8, 1958.',
 'Hopkins, W. G. (2000). Measures of reliability in sports medicine and science. Sports Medicine, 30(1), 1–15.',
 'Weakley, J., et al. (2022). A comparison of methods for quantifying the smallest worthwhile change. Sports Medicine – Open, 8, 20.',
 'Saw, A. E., Main, L. C., & Gastin, P. B. (2016). Monitoring the athlete training response: subjective self-reported measures. British Journal of Sports Medicine, 50(5), 281–291.',
]
for i,r in enumerate(refs,1):
    p=doc.add_paragraph(); p.add_run(f'{i}. {r}').font.size=Pt(9)

doc.add_paragraph()
P('Reprodutibilidade e materiais suplementares. Todas as análises são reproduzíveis pelos scripts em scripts/analise/ (amplitude, amp_docx_compute, prepost_compute, denoise, pair_fatigue, pair_vigor_tmd) sobre a base derivada da janela. Relatórios-companheiros e apps interativos em Artigos/. Atletas anonimizados (A01–A27); nenhum dado bruto com identificação é distribuído.',it=True,size=9)

OUTP='/home/user/mdlucca/Artigos/Artigo_Cientifico_Sinal_Ruido_Amplitude.docx'
doc.save(OUTP); print('SAVED',OUTP)
d2=Document(OUTP); print('paras',len(d2.paragraphs),'tables',len(d2.tables),'imgs',len(d2.inline_shapes))
