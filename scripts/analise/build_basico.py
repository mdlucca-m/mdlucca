# -*- coding: utf-8 -*-
import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
R=json.load(open('brums_desc2.json')); STAT=json.load(open('brums_stats3.json'))
PV=json.load(open('pv_stats.json')); LIM=json.load(open('tcar_limiar.json')); TCD=json.load(open('tcar_desc.json'))
MS=json.load(open('model_stats.json'))
FG='/home/user/mdlucca/Artigos/figuras'
def c2(s): return str(s).replace('.',',')
doc=Document()
stl=doc.styles['Normal']; stl.font.name='Times New Roman'; stl.font.size=Pt(12)
stl.element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')
stl.paragraph_format.line_spacing=1.5; stl.paragraph_format.space_after=Pt(0)
sec=doc.sections[0]; sec.top_margin=Cm(3); sec.left_margin=Cm(3); sec.bottom_margin=Cm(2); sec.right_margin=Cm(2)
_TN=[0]; _FN=[0]
def P(t='',just=True,size=12,after=6,bold=False,ind=True):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(size); r.bold=bold
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(after)
    if just: p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if just and ind: p.paragraph_format.first_line_indent=Cm(1.25)
    return p
def H(t,size=12,before=10):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(size)
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.5
def _bd(c):
    tcPr=c._tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for e_ in ['top','bottom','left','right']:
        el=OxmlElement(f'w:{e_}'); on=e_ in ('top','bottom')
        el.set(qn('w:val'),'single' if on else 'nil'); el.set(qn('w:sz'),'6'); el.set(qn('w:color'),'000000'); b.append(el)
    tcPr.append(b)
def table(cap,header,rows,fonte='Fonte: dados da pesquisa (2026).',fs=9,note=None):
    _TN[0]+=1
    p=doc.add_paragraph(); r=p.add_run('Tabela %d – %s'%(_TN[0],cap)); r.font.size=Pt(11)
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
    t=doc.add_table(rows=1,cols=len(header)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,htx in enumerate(header):
        cc=t.rows[0].cells[i]; cc.text=''; rr=cc.paragraphs[0].add_run(htx); rr.bold=True; rr.font.size=Pt(fs); rr.font.name='Times New Roman'
        cc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; _bd(cc)
    for row in rows:
        cs=t.add_row().cells
        for i,val in enumerate(row):
            cs[i].text=''; rr=cs[i].paragraphs[0].add_run(str(val)); rr.font.size=Pt(fs); rr.font.name='Times New Roman'
            cs[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if i>0 else WD_ALIGN_PARAGRAPH.LEFT; _bd(cs[i])
    if note:
        pn=doc.add_paragraph(); rn=pn.add_run(note); rn.font.size=Pt(8.5); rn.italic=True; pn.paragraph_format.space_after=Pt(0)
    pf=doc.add_paragraph(); rf=pf.add_run(fonte); rf.font.size=Pt(9); pf.paragraph_format.space_after=Pt(6)
    return _TN[0]
def figure(path,cap,w=15.0):
    _FN[0]+=1
    pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; pp.add_run().add_picture(path,width=Cm(w)); pp.paragraph_format.space_before=Pt(6)
    pc=doc.add_paragraph(); pc.alignment=WD_ALIGN_PARAGRAPH.CENTER; rc=pc.add_run('Figura %d – %s'%(_FN[0],cap)); rc.font.size=Pt(11)
    pf=doc.add_paragraph(); pf.alignment=WD_ALIGN_PARAGRAPH.CENTER; rf=pf.add_run('Fonte: elaboração dos autores (2026).'); rf.font.size=Pt(9); pf.paragraph_format.space_after=Pt(6)
    return _FN[0]
def pstr(p): return '< 0,001' if p<0.001 else c2('%.3f'%p)

pr=R['prepos']; prof=R['profiles']; sm=R['sample']; desc=R['desc']; PREV=MS['prev']
SH=STAT['shapiro']; LP=LIM['LIM']['PVini']
ORD=[('Vigor','Vigor'),('Fadiga','Fadiga'),('TMD','PTH'),('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Confusao','Confusão')]

# ===== TÍTULO =====
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('PERFIL DE HUMOR DE ATLETAS DE HANDEBOL DE ELITE NA ÚLTIMA SEMANA DE PRÉ-TEMPORADA')
r.bold=True; r.font.size=Pt(13); p.paragraph_format.space_after=Pt(10)

# ===== 1 INTRODUÇÃO (robusta) =====
H('1 INTRODUÇÃO')
P('O acompanhamento do estado psicológico dos atletas consolidou-se como parte essencial da gestão do treinamento no '
 'esporte de rendimento. Ao longo das últimas décadas, evidências têm demonstrado que instrumentos de autorrelato do '
 'humor são práticos, econômicos e sensíveis às variações da carga de treino, apresentando utilidade preditiva tanto '
 'para o bem-estar quanto para o desempenho esportivo (SAW; MAIN; GASTIN, 2016; LOCHBAUM et al., 2021). Por essa razão, '
 'documentos de consenso recomendam o uso rotineiro de medidas subjetivas para monitorar a fadiga e orientar decisões de '
 'treino e recuperação (KELLMANN et al., 2018).')
P('Entre esses instrumentos, a Escala de Humor de Brunel (BRUMS) — versão abreviada do Profile of Mood States (POMS) — é '
 'uma das mais utilizadas na ciência do esporte. Composta por 24 itens organizados em seis subescalas (tensão, depressão, '
 'raiva, vigor, fadiga e confusão), foi concebida para aplicação rápida e repetida e possui propriedades psicométricas '
 'sólidas, com sucessivas validações transculturais, incluindo a versão brasileira (TERRY; LANE; FOGARTY, 2003; ROHLFS '
 'et al., 2008; ROHLFS et al., 2023). Além da leitura isolada de cada subescala, a literatura recente formalizou seis '
 'perfis prototípicos de humor — iceberg, Everest invertido, iceberg invertido, submerso, barbatana de tubarão e '
 'superfície —, replicados em grandes amostras internacionais e utilizados no rastreamento da prontidão e do estado '
 'psicológico de atletas de elite (PARSONS-SMITH; TERRY; MACHIN, 2017; HAN; PARSONS-SMITH; TERRY, 2020; TERRY et al., '
 '2021).')
P('O handebol de quadra é uma modalidade coletiva de caráter marcadamente intermitente e de alta intensidade, na qual '
 'esforços máximos e explosivos — sprints, saltos, arremessos, mudanças de direção e contatos — alternam-se, de forma '
 'imprevisível, com períodos de recuperação incompleta, exigindo elevada tolerância à fadiga (KARCHER; BUCHHEIT, 2014; '
 'WAGNER et al., 2014). Esse padrão de esforço eleva a carga interna nos microciclos de acúmulo e repercute no estado '
 'afetivo dos atletas, sobretudo na última semana de pré-temporada, quando a carga que antecede a competição se '
 'concentra. A tolerância a essa carga depende, em boa medida, da aptidão aeróbia intermitente, que sustenta a '
 'recuperação entre os esforços de alta intensidade; um marcador de campo válido dessa capacidade é o pico de velocidade '
 'obtido no Teste de Carminatti (T-CAR), associado ao desempenho físico em modalidades intermitentes (FERNANDES-DA-SILVA '
 'et al., 2016). Descrever como o humor se comporta ao longo desse microciclo, e relacioná-lo a esse parâmetro '
 'fisiológico, fornece à comissão técnica informação diretamente aplicável ao monitoramento individualizado.')

# ===== 2 OBJETIVO (robusto) =====
H('2 OBJETIVO')
P('Objetivo geral: caracterizar e analisar o perfil de humor de atletas de handebol de elite ao longo da última semana '
 'de pré-temporada, descrevendo o comportamento das seis dimensões do BRUMS e da Perturbação Total do Humor (PTH), do '
 'início ao fim da semana e dentro do dia de treino (pré e pós).')
P('Objetivos específicos: (i) caracterizar a amostra quanto aos aspectos sociodemográficos e antropométricos; (ii) '
 'verificar a normalidade das distribuições das variáveis do humor; (iii) descrever a resposta aguda (pré → pós) de cada '
 'dimensão ao longo da semana; (iv) descrever o comportamento das variáveis do dia 1 ao dia 7; (v) classificar e '
 'acompanhar a evolução dos perfis de humor; e (vi) analisar a aptidão aeróbia intermitente (pico de velocidade do '
 'T-CAR) como parâmetro fisiológico, estabelecendo um limiar de pico de velocidade associado aos dias de maior fadiga.')

# ===== 3 MÉTODO =====
H('3 MÉTODO')
H('3.1 Delineamento e amostra',12,before=6)
P('Estudo descritivo e observacional, de medidas repetidas, conduzido em condições ecológicas de treinamento durante o '
 'microciclo pré-competitivo (21 a 27 de abril de 2024). Participaram %d atletas de handebol do sexo masculino de nível '
 'competitivo, distribuídos nas posições da modalidade. Todos os procedimentos respeitaram os princípios éticos da '
 'Declaração de Helsinque, com consentimento informado.'%sm['n'])
H('3.2 Instrumentos',12,before=6)
P('O humor foi avaliado pela BRUMS-24, com 24 itens respondidos em escala de 0 (“nada”) a 4 (“extremamente”), agrupados '
 'em seis subescalas de 0 a 16 pontos (tensão, depressão, raiva, vigor, fadiga e confusão). A PTH resume o perfil '
 '(soma das dimensões negativas menos o vigor). A aptidão aeróbia intermitente foi avaliada pelo Teste de Carminatti '
 '(T-CAR), teste de campo progressivo e intermitente (repetições de 12 s de corrida em vaivém por 6 s de recuperação, '
 'até a exaustão), tendo o pico de velocidade (PV) como desfecho (FERNANDES-DA-SILVA et al., 2016). O T-CAR foi aplicado '
 'em 15 de abril de 2024, quatro dias de treino antes do início do microciclo monitorado, de modo a servir como '
 'parâmetro fisiológico de linha de base.')
H('3.3 Procedimentos',12,before=6)
P('O BRUMS foi autoaplicado por formulário eletrônico com registro de data e hora, duas vezes por dia de treino — a '
 'primeira resposta tomada como pré e a última como pós —, ao longo de sete dias, totalizando %d observações válidas.'%sm['n_obs'])
H('3.4 Análise estatística',12,before=6)
P('Os dados foram analisados por estatística descritiva (média, desvio-padrão, mediana e amplitude). A normalidade das '
 'distribuições foi verificada pelo teste de Shapiro-Wilk. As diferenças entre os momentos pré e pós-treino foram '
 'testadas pelo teste de Wilcoxon para amostras pareadas. A evolução dos perfis de humor foi descrita por distribuição '
 'de frequências. A relação entre o pico de velocidade do T-CAR e a fadiga foi analisada por regressão, e um limiar de '
 'pico de velocidade foi estabelecido para discriminar os dias de maior fadiga. Adotou-se nível de significância de '
 '5% (p < 0,05).')

# ===== 4 RESULTADOS =====
H('4 RESULTADOS')
H('4.1 Caracterização da amostra',12,before=6)
P('A Tabela 1 apresenta a distribuição da amostra por posição e a Tabela 2, a caracterização sociodemográfica e '
 'antropométrica dos %d atletas.'%sm['n'])
table('Distribuição da amostra por posição de jogo (n = %d).'%sm['n'],['Posição','n','%'],
    [[k,v,c2('%.1f'%(100*v/sm['n']))] for k,v in sm['pos'].items()]+[['Total',sm['n'],'100,0']],fs=9)
def srow(lab,k):
    v=sm[k]; return [lab,c2('%.1f'%v['mean']),c2('%.1f'%v['sd']),'%s–%s'%(c2('%.1f'%v['mn']),c2('%.1f'%v['mx']))]
table('Caracterização sociodemográfica e antropométrica da amostra (n = %d).'%sm['n'],
    ['Variável','Média','DP','Mín–Máx'],
    [srow('Idade (anos)','idade'),srow('Estatura (cm)','estatura'),srow('Massa corporal (kg)','massa'),srow('Percentual de gordura (%)','pG'),srow('Experiência na modalidade (anos)','exp')])
H('4.2 Normalidade das distribuições',12,before=6)
P('O teste de Shapiro-Wilk indicou que as seis dimensões do humor não seguem distribuição normal (p < 0,001; Tabela 3), '
 'com as dimensões negativas concentradas em escores baixos. Por esse motivo, as comparações foram conduzidas por testes '
 'não paramétricos.')
table('Teste de normalidade (Shapiro-Wilk) das dimensões do BRUMS.',['Dimensão','Estatística (W)','p'],
    [[lab,c2('%.3f'%SH[k]['W']),'< 0,001' if SH[k]['p']<0.001 else c2('%.3f'%SH[k]['p'])] for k,lab in [('Vigor','Vigor'),('Fadiga','Fadiga'),('Tensao','Tensão'),('Depressao','Depressão'),('Raiva','Raiva'),('Confusao','Confusão')]],fs=9)
H('4.3 Estatística descritiva das dimensões do humor',12,before=6)
P('A Tabela 4 apresenta a descritiva geral das seis dimensões e da PTH. O vigor e a fadiga concentram as maiores médias '
 'e a maior variabilidade, enquanto as dimensões negativas (tensão, depressão, raiva e confusão) mantêm-se em níveis '
 'baixos ao longo da semana.')
def drow(k,lab):
    v=desc[k]; return [lab,c2('%.2f'%v['mean']),c2('%.1f'%v['md']),c2('%.2f'%v['sd']),'%s–%s'%(c2('%.0f'%v['mn']),c2('%.0f'%v['mx']))]
table('Estatística descritiva das dimensões do BRUMS e da PTH (%d observações).'%sm['n_obs'],
    ['Dimensão','Média','Mediana','DP','Mín–Máx'],[drow(k,l) for k,l in ORD],fs=9)
H('4.4 Diferenças entre pré e pós-treino ao longo da semana',12,before=6)
P('Considerando o conjunto da semana, observou-se resposta aguda consistente do humor ao treino (Tabela 5): o vigor '
 'reduziu-se do pré para o pós (−%s%%; p = %s), enquanto a fadiga (+%s%%; p = %s) e a PTH (+%s%%; p = %s) aumentaram. As '
 'dimensões negativas de valência não fadiga variaram pouco dentro do dia.'%(
   c2('%.0f'%abs(pr['Vigor']['pct'])),pstr(pr['Vigor']['p']),c2('%.0f'%pr['Fadiga']['pct']),pstr(pr['Fadiga']['p']),
   c2('%.0f'%pr['TMD']['pct']),pstr(pr['TMD']['p'])))
def ppr(k,lab):
    v=pr[k]; return [lab,c2('%.2f'%v['pre']),c2('%.2f'%v['pos']),c2('%+.2f'%v['delta']),c2('%+.0f'%v['pct'])+'%',pstr(v['p'])]
table('Diferenças entre os momentos pré e pós-treino (média agregada da semana; teste de Wilcoxon).',
    ['Dimensão','Pré (M)','Pós (M)','Δ','Variação (%)','p'],[ppr(k,l) for k,l in ORD],fs=9)
H('4.5 Comportamento das variáveis ao longo da semana',12,before=6)
P('Do dia 1 ao dia 7, o perfil de humor deslocou-se de forma clara: o vigor tendeu a diminuir e a fadiga a aumentar '
 'progressivamente, atingindo, respectivamente, o menor e o maior valor no último dia (Figura 1). As dimensões negativas '
 'mantiveram-se relativamente estáveis e em níveis baixos ao longo de toda a semana.')
figure(f'{FG}/xb2_traj.png','Comportamento das dimensões do humor ao longo da semana (médias diárias; áreas sombreadas = intervalo de confiança de 95%).')
H('4.6 Classificação dos perfis de humor',12,before=6)
P('A classificação nos seis perfis de humor evidenciou a migração ao longo da semana (Figura 2; Tabela 6). No primeiro '
 'dia predominou o perfil iceberg — vigor alto sobre dimensões negativas baixas, marca de prontidão —, presente em %s%% '
 'das avaliações; no último dia, o perfil de fadiga (“barbatana de tubarão”) tornou-se o mais frequente (%s%%) e o '
 'iceberg reduziu-se para %s%%, sem instalação relevante de perfis de risco.'%(
   c2('%.0f'%(100*PREV['D1']['Iceberg']/PREV['n_d1'])),c2('%.0f'%(100*PREV['D7']['Barbatana tubarão']/PREV['n_d7'])),c2('%.0f'%(100*PREV['D7']['Iceberg']/PREV['n_d7']))))
figure(f'{FG}/xb5_prev.png','Distribuição (%) dos seis perfis de humor no primeiro (Dia 1) e no último dia (Dia 7) da semana.',w=13.5)
PROFR=[('Iceberg','Iceberg'),('Everest invertido','Everest invertido'),('Iceberg invertido','Iceberg invertido'),('Submerso','Submerso'),('Barbatana tubarão','Barbatana de tubarão'),('Superfície','Superfície')]
def prow(p,lab):
    d1=PREV['D1'][p]; d7=PREV['D7'][p]
    return [lab,'%d (%s)'%(d1,c2('%.1f'%(100*d1/PREV['n_d1']))),'%d (%s)'%(d7,c2('%.1f'%(100*d7/PREV['n_d7'])))]
table('Distribuição dos perfis de humor no primeiro e no último dia da semana: n (%).',
    ['Perfil','Dia 1 n (%)','Dia 7 n (%)'],[prow(p,lab) for p,lab in PROFR],fs=9)

# ===== 4.7 T-CAR + LIMIAR =====
H('4.7 Aptidão intermitente (T-CAR) e limiar de pico de velocidade',12,before=6)
P('O desempenho no T-CAR está descrito na Tabela 7. O pico de velocidade inicial (T-CAR1) foi de %s ± %s km/h, servindo '
 'de parâmetro fisiológico de referência.'%(c2('%.1f'%PV['desc']['pvini_m']),c2('%.1f'%PV['desc']['pvini_sd'])))
def tcrow(o): return [o['lab'],o['n'],c2('%.1f'%o['m']),c2('%.1f'%o['sd']),'%s–%s'%(c2('%.1f'%o['mn']),c2('%.1f'%o['mx']))]
table('Estatística descritiva do desempenho no Teste de Carminatti (T-CAR).',
    ['Parâmetro','n','Média','DP','Mín–Máx'],[tcrow(o) for o in TCD],
    note='PV = pico de velocidade; FC = frequência cardíaca; TRIMP = training impulse.',fs=8.5)
P('Para verificar se a aptidão física ajuda a explicar a fadiga observada, relacionou-se o pico de velocidade à fadiga '
 'da semana. A análise mostrou associação inversa e significativa: atletas com maior pico de velocidade tenderam a '
 'relatar menos fadiga (ρ = %s; p = %s). Em seguida, ajustou-se um modelo de regressão para estimar a probabilidade de um '
 'atleta apresentar um “dia de fadiga elevada” em função do seu pico de velocidade e, a partir dele, identificou-se o '
 'ponto de corte que melhor separa os atletas — o limiar de pico de velocidade (Figura 3).'%(
   c2('%+.2f'%PV['pv']['wk_FadFisica']['TCAR1']['rho']),pstr(PV['pv']['wk_FadFisica']['TCAR1']['rho_p'])))
figure(f'{FG}/pv9_limiar.png','Relação entre o pico de velocidade do T-CAR e a fadiga da semana, com reta de regressão e o limiar de pico de velocidade (linha tracejada).',w=13.5)
P('O limiar identificado foi de aproximadamente %s km/h, com bom poder de discriminação (área sob a curva = %s; '
 'sensibilidade = %s; especificidade = %s). Na prática, isso significa que atletas com pico de velocidade abaixo desse '
 'valor apresentaram maior probabilidade de vivenciar dias de fadiga elevada ao longo do microciclo. O porquê dessa '
 'análise é direto: ao dispor de um limiar objetivo de aptidão física, a comissão técnica pode individualizar a carga — '
 'reforçando a recuperação dos atletas menos aptos, que tendem a fadigar mais sob a mesma carga de treino, e '
 'distinguindo a fadiga esperada daquela que possa sinalizar sobrecarga.'%(
   c2('%.1f'%LP['thr']),c2('%.2f'%LP['auc']),c2('%.2f'%LP['sens']),c2('%.2f'%LP['spec'])))

# ===== 5 DISCUSSÃO =====
H('5 DISCUSSÃO')
P('Os resultados descrevem, de forma clara, o comportamento do humor no microciclo pré-competitivo. A resposta '
 'concentrou-se no eixo energia–fadiga: o vigor caiu e a fadiga aumentou, tanto de forma aguda (do pré para o pós de '
 'cada treino) quanto ao longo da semana (do dia 1 ao dia 7), enquanto as dimensões negativas de valência não fadiga '
 'permaneceram estáveis e em níveis baixos. Esse padrão é coerente com a literatura de monitoramento, que aponta o vigor '
 'e a fadiga como as dimensões mais sensíveis à carga de treino (SAW; MAIN; GASTIN, 2016; KELLMANN et al., 2018).')
P('A migração do perfil iceberg, típico de prontidão, para o perfil de fadiga ao final da semana reflete o acúmulo de '
 'carga da pré-temporada e reproduz o “derretimento do iceberg” descrito na literatura (MORGAN, 1985; HAN; PARSONS-SMITH; '
 'TERRY, 2020). Por fim, a relação entre a aptidão aeróbia intermitente e a fadiga — sintetizada no limiar de pico de '
 'velocidade — reforça a utilidade de integrar um parâmetro fisiológico simples ao monitoramento do humor, permitindo '
 'individualizar a interpretação da resposta afetiva à carga.')

# ===== 6 CONSIDERAÇÕES FINAIS =====
H('6 CONSIDERAÇÕES FINAIS')
P('Na última semana de pré-temporada, o humor dos handebolistas migrou da prontidão para a fadiga, com o vigor em queda '
 'e a fadiga em ascensão, do início ao fim da semana e dentro de cada treino. A aptidão aeróbia intermitente mostrou-se '
 'associada à fadiga, e o limiar de pico de velocidade oferece uma referência prática para individualizar a carga. '
 'Recomenda-se o acompanhamento contínuo do humor, com atenção especial ao vigor e à fadiga.')

# ===== REFERÊNCIAS =====
H('REFERÊNCIAS')
refs=[
 'FERNANDES-DA-SILVA, J. et al. The peak velocity derived from the Carminatti Test is related to physical match performance in young soccer players. Journal of Sports Sciences, v. 34, n. 24, p. 2238–2245, 2016.',
 'HAN, C.; PARSONS-SMITH, R. L.; TERRY, P. C. Mood profiling in Singapore: cross-cultural validation and potential applications of mood profile clusters. Frontiers in Psychology, v. 11, 665, 2020.',
 'KARCHER, C.; BUCHHEIT, M. On-court demands of elite handball, with special reference to playing positions. Sports Medicine, v. 44, n. 6, p. 797–814, 2014.',
 'KELLMANN, M. et al. Recovery and performance in sport: consensus statement. International Journal of Sports Physiology and Performance, v. 13, n. 2, p. 240–245, 2018.',
 'LOCHBAUM, M. et al. The Profile of Mood States and athletic performance: a meta-analysis of published studies. European Journal of Investigation in Health, Psychology and Education, v. 11, n. 1, p. 50–70, 2021.',
 'MORGAN, W. P. Selected psychological factors limiting performance: a mental health model. In: CLARKE, D. H.; ECKERT, H. M. (Ed.). Limits of human performance. Champaign: Human Kinetics, 1985. p. 70–80.',
 'PARSONS-SMITH, R. L.; TERRY, P. C.; MACHIN, M. A. Identification and description of novel mood profile clusters. Frontiers in Psychology, v. 8, 1958, 2017.',
 'ROHLFS, I. C. P. M. et al. A Escala de Humor de Brunel (Brums): instrumento para detecção precoce da síndrome do excesso de treinamento. Revista Brasileira de Medicina do Esporte, v. 14, n. 3, p. 176–181, 2008.',
 'ROHLFS, I. C. P. M. et al. Psychometric characteristics of the Brazil Mood Scale among youth and elite athletes. Sports, v. 11, n. 12, 244, 2023.',
 'SAW, A. E.; MAIN, L. C.; GASTIN, P. B. Monitoring the athlete training response: subjective self-reported measures trump commonly used objective measures. British Journal of Sports Medicine, v. 50, n. 5, p. 281–291, 2016.',
 'TERRY, P. C.; LANE, A. M.; FOGARTY, G. J. Construct validity of the Profile of Mood States — Adolescents for use with adults. Psychology of Sport and Exercise, v. 4, n. 2, p. 125–139, 2003.',
 'TERRY, P. C. et al. Mood profiling for sustainable mental health among athletes. Sustainability, v. 13, n. 11, 6116, 2021.',
 'WAGNER, H. et al. Individual and team performance in team-handball: a review. Journal of Sports Science and Medicine, v. 13, n. 4, p. 808–816, 2014.']
for rf in refs:
    p=doc.add_paragraph(); r=p.add_run(rf); r.font.size=Pt(11); p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(6); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

OUTP='/home/user/mdlucca/Artigos/Relatorio_Basico_Perfil_Humor.docx'
doc.save(OUTP); print('SAVED',OUTP,'| Tabelas',_TN[0],'Figuras',_FN[0])
