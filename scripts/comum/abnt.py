#!/usr/bin/env python3
"""Renderizador ABNT compartilhado pelos documentos do projeto.

Formato conforme NBR 14724: A4 retrato, margens de 3 cm à esquerda e no topo e
2 cm à direita e embaixo, Times New Roman 12, recuo de 1,25 cm na primeira
linha. Tabelas com título acima e fonte abaixo, conforme as normas de
apresentação tabular; figuras com legenda acima e fonte abaixo.

O documento é descrito por uma lista de blocos, cada um uma tupla:

    ("h1", texto)                    seção de primeiro nível
    ("h2", texto)                    seção de segundo nível
    ("h3", texto)                    seção de terceiro nível
    ("p", texto)                     parágrafo com recuo
    ("p0", texto)                    parágrafo sem recuo
    ("nota", texto)                  parágrafo menor, sem recuo
    ("lista", [item, ...])           itens marcados
    ("tab", dicionário)              tabela de apresentação tabular
    ("fig", caminho, largura, legenda)
    ("quebra",)                      quebra de página
"""
from __future__ import annotations

import re
import shutil
import zipfile
from pathlib import Path

import docx
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

FONTE = "Times New Roman"
CORPO, MENOR, MIUDO = 12, 10, 8
PRETO = RGBColor(0, 0, 0)


def preparar(doc, entrelinha: float = 1.0):
    s = doc.sections[0]
    s.orientation = WD_ORIENT.PORTRAIT
    s.page_width, s.page_height = Cm(21.0), Cm(29.7)
    s.left_margin, s.top_margin = Cm(3.0), Cm(3.0)
    s.right_margin, s.bottom_margin = Cm(2.0), Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = FONTE
    normal.font.size = Pt(CORPO)
    normal.font.color.rgb = PRETO
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONTE)

    for nivel in (1, 2, 3):
        st = doc.styles[f"Heading {nivel}"]
        st.font.name = FONTE
        st.font.size = Pt(CORPO if nivel < 3 else MENOR + 1)
        st.font.bold = True
        st.font.italic = False
        st.font.all_caps = (nivel == 1)
        st.font.color.rgb = PRETO
        st.element.rPr.rFonts.set(qn("w:eastAsia"), FONTE)
        pf = st.paragraph_format
        pf.space_before = Pt({1: 12, 2: 9, 3: 7}[nivel])
        pf.space_after = Pt(4)
        pf.line_spacing = entrelinha
        pf.keep_with_next = True


def paragrafo(doc, texto, *, recuo=True, tamanho=CORPO, espaco=1.0,
              centro=False, negrito=False, italico=False, justificar=True,
              depois=0, marcador=False):
    p = doc.add_paragraph(style="List Bullet" if marcador else None)
    pf = p.paragraph_format
    pf.line_spacing = espaco
    pf.space_after = Pt(depois)
    if not marcador:
        pf.first_line_indent = Cm(1.25) if recuo else Cm(0)
    else:
        pf.left_indent, pf.first_line_indent = Cm(1.0), Cm(-0.4)
    p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if centro else
                   WD_ALIGN_PARAGRAPH.JUSTIFY if justificar else
                   WD_ALIGN_PARAGRAPH.LEFT)
    for trecho, forte in _partes(texto):
        r = p.add_run(trecho)
        r.font.name, r.font.size, r.font.color.rgb = FONTE, Pt(tamanho), PRETO
        r.bold, r.italic = (negrito or forte), italico
    return p


_MARCA = re.compile(r"\*\*(.+?)\*\*")


def _partes(texto: str):
    """Divide o texto nos trechos entre asteriscos duplos, que saem em negrito."""
    pos = 0
    for m in _MARCA.finditer(texto):
        if m.start() > pos:
            yield texto[pos:m.start()], False
        yield m.group(1), True
        pos = m.end()
    if pos < len(texto):
        yield texto[pos:], False


# O esquema OOXML exige esta ordem dentro de <w:tcBorders>.
ORDEM_BORDAS = ("top", "left", "bottom", "right", "insideH", "insideV")


def definir_bordas(celula, visiveis: set[str], cor="000000", sz=6):
    tcPr = celula._tc.get_or_add_tcPr()
    antigas = tcPr.find(qn("w:tcBorders"))
    if antigas is not None:
        tcPr.remove(antigas)
    marcas = tcPr.makeelement(qn("w:tcBorders"), {})
    for lado in ORDEM_BORDAS:
        if lado in visiveis:
            atributos = {qn("w:val"): "single", qn("w:sz"): str(sz),
                         qn("w:space"): "0", qn("w:color"): cor}
        else:
            atributos = {qn("w:val"): "none", qn("w:sz"): "0",
                         qn("w:space"): "0", qn("w:color"): "auto"}
        marcas.append(marcas.makeelement(qn(f"w:{lado}"), atributos))
    tcPr.append(marcas)


def inserir_tabela(doc, dados, fonte_tabela: str):
    """Apresentação tabular: sem bordas laterais, com traço no topo, sob o
    cabeçalho e no rodapé. Uma linha cujo primeiro item começa por asterisco
    duplo vira subtítulo em negrito, sem bordas."""
    paragrafo(doc, f"Tabela {dados['numero']} - {dados['titulo']}",
              recuo=False, justificar=False, tamanho=MENOR, espaco=1.0,
              depois=3)

    n = len(dados["cabecalho"])
    corpo = 8 if n <= 5 else 7.5 if n <= 6 else 6.5
    tab = doc.add_table(rows=1, cols=n)
    tab.alignment = WD_TABLE_ALIGNMENT.CENTER
    tab.autofit = True
    tab.style = "Table Grid"

    for i, texto in enumerate(dados["cabecalho"]):
        cel = tab.rows[0].cells[i]
        cel.text = ""
        p = cel.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(texto)
        r.bold = True
        r.font.name, r.font.size, r.font.color.rgb = FONTE, Pt(corpo), PRETO

    for linha in dados["linhas"]:
        celulas = tab.add_row().cells
        subtitulo = str(linha[0]).startswith("**")
        for i, texto in enumerate(linha):
            celulas[i].text = ""
            p = celulas[i].paragraphs[0]
            p.alignment = (WD_ALIGN_PARAGRAPH.LEFT if i == 0
                           else WD_ALIGN_PARAGRAPH.CENTER)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(str(texto).replace("**", ""))
            r.bold = subtitulo
            r.font.name, r.font.size, r.font.color.rgb = FONTE, Pt(corpo), PRETO

    ultima = len(tab.rows) - 1
    for i, linha in enumerate(tab.rows):
        if i == 0:
            visiveis = {"top", "bottom"}
        elif i == ultima:
            visiveis = {"bottom"}
        else:
            visiveis = set()
        for cel in linha.cells:
            definir_bordas(cel, visiveis)

    paragrafo(doc, fonte_tabela, recuo=False, justificar=False, tamanho=MIUDO,
              espaco=1.0, depois=1)
    if dados.get("nota"):
        paragrafo(doc, dados["nota"], recuo=False, tamanho=MIUDO, espaco=1.0,
                  depois=6)


def inserir_figura(doc, caminho: Path, legenda: str, largura_cm: float,
                   fonte_figura: str):
    paragrafo(doc, legenda, recuo=False, justificar=False, tamanho=MENOR,
              espaco=1.0, depois=3)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(caminho), width=Cm(largura_cm))
    paragrafo(doc, fonte_figura, recuo=False, justificar=False, tamanho=MIUDO,
              espaco=1.0, depois=6)


def montar(blocos, saida: Path, *, titulo: str, subtitulo: str = "",
           abertura: list[tuple[str, str]] = (), entrelinha: float = 1.0,
           fonte_tabela="Fonte: dados da pesquisa (2026).",
           fonte_figura="Fonte: elaborada pelos autores (2026).",
           dir_figuras: Path | None = None) -> dict:
    doc = docx.Document()
    preparar(doc, entrelinha)

    paragrafo(doc, titulo, recuo=False, centro=True, negrito=True, espaco=1.0,
              depois=3)
    if subtitulo:
        paragrafo(doc, subtitulo, recuo=False, centro=True, italico=True,
                  tamanho=MENOR, espaco=1.0, depois=8)
    for rotulo, texto in abertura:
        paragrafo(doc, rotulo, recuo=False, justificar=False, negrito=True,
                  tamanho=MENOR, espaco=1.0, depois=2)
        paragrafo(doc, texto, recuo=False, tamanho=MENOR, espaco=1.0, depois=6)

    conta = {"parágrafos": 0, "tabelas": 0, "figuras": 0, "seções": 0}
    for bloco in blocos:
        tipo = bloco[0]
        if tipo in ("h1", "h2", "h3"):
            doc.add_heading(bloco[1], level=int(tipo[1]))
            conta["seções"] += tipo == "h1"
        elif tipo == "p":
            paragrafo(doc, bloco[1], depois=4)
            conta["parágrafos"] += 1
        elif tipo == "p0":
            paragrafo(doc, bloco[1], recuo=False, depois=4)
            conta["parágrafos"] += 1
        elif tipo == "nota":
            paragrafo(doc, bloco[1], recuo=False, tamanho=MENOR, depois=5)
        elif tipo == "lista":
            for item in bloco[1]:
                paragrafo(doc, item, marcador=True, depois=2)
        elif tipo == "tab":
            inserir_tabela(doc, bloco[1], fonte_tabela)
            conta["tabelas"] += 1
        elif tipo == "fig":
            base = dir_figuras or Path(".")
            inserir_figura(doc, base / bloco[1], bloco[3], bloco[2],
                           fonte_figura)
            conta["figuras"] += 1
        elif tipo == "quebra":
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        else:
            raise ValueError(f"bloco desconhecido: {tipo}")

    saida.parent.mkdir(parents=True, exist_ok=True)
    doc.save(saida)
    limpar_settings(saida)
    return conta


def limpar_settings(caminho: Path):
    """Remove o <w:zoom> sem atributo que o template do python-docx grava e
    que viola o esquema OOXML."""
    temp = caminho.with_suffix(".tmp.docx")
    with zipfile.ZipFile(caminho) as orig, \
            zipfile.ZipFile(temp, "w", zipfile.ZIP_DEFLATED) as novo:
        for item in orig.infolist():
            dados = orig.read(item.filename)
            if item.filename == "word/settings.xml":
                txt = dados.decode("utf-8")
                if "w:percent" not in txt:
                    txt = re.sub(r"<w:zoom\b[^>]*/>", "", txt)
                    txt = re.sub(r"<w:zoom\b[^>]*>.*?</w:zoom>", "", txt,
                                 flags=re.S)
                dados = txt.encode("utf-8")
            novo.writestr(item, dados)
    shutil.move(temp, caminho)
