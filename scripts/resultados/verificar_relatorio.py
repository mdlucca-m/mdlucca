#!/usr/bin/env python3
"""Verifica o relatório de resultados: formato ABNT, presença das tabelas e
figuras, e conformidade de estilo.

    python3 scripts/resultados/verificar_relatorio.py data/RESULTADOS.docx
"""
from __future__ import annotations

import argparse
import re
import sys
import zipfile
from pathlib import Path

import docx

sys.path.insert(0, str(Path(__file__).resolve().parent))
import conteudo  # noqa: E402
from verificar_estilo import achar_gerundios  # noqa: E402

falhas: list[str] = []
aprovados = 0


def conferir(cond: bool, desc: str, detalhe: str = "") -> None:
    global aprovados
    if cond:
        aprovados += 1
        print(f"  ✓ {desc}")
    else:
        print(f"  ✗ {desc}" + (f" → {detalhe}" if detalhe else ""))
        falhas.append(desc)


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("arquivo", type=Path)
    a = ap.parse_args()
    d = docx.Document(a.arquivo)
    texto = "\n".join(p.text for p in d.paragraphs)
    for t in d.tables:
        for linha in t.rows:
            texto += "\n" + " ".join(c.text for c in linha.cells)

    print("── formato ABNT (NBR 14724) ──")
    s = d.sections[0]
    conferir(abs(s.page_width.cm - 21.0) < 0.1 and abs(s.page_height.cm - 29.7) < 0.1,
             "A4 retrato", f"{s.page_width.cm:.1f} x {s.page_height.cm:.1f} cm")
    conferir(all(abs(m.cm - v) < 0.1 for m, v in
                 ((s.left_margin, 3), (s.top_margin, 3),
                  (s.right_margin, 2), (s.bottom_margin, 2))),
             "margens 3/3/2/2 cm")
    normal = d.styles["Normal"].font
    conferir(normal.name == "Times New Roman" and normal.size.pt == 12,
             "Times New Roman 12 no corpo")
    # Parágrafos de corpo são os do módulo de conteúdo; notas de tabela e
    # legendas seguem outro formato (espaço simples, sem recuo), de propósito.
    do_corpo = {t for s_ in conteudo.SECOES
                for t in s_.get("paragrafos", []) + s_.get("pos_tabela", [])}
    corpo = [p for p in d.paragraphs if p.text in do_corpo]
    conferir(len(corpo) == len(do_corpo),
             f"{len(do_corpo)} parágrafos de corpo localizados", str(len(corpo)))
    conferir(corpo and all(abs((p.paragraph_format.line_spacing or 0) - 1.5) < 0.01
                           for p in corpo),
             "entrelinha 1,5 no corpo")
    conferir(corpo and all(
        p.paragraph_format.first_line_indent is not None
        and abs(p.paragraph_format.first_line_indent.cm - 1.25) < 0.05
        for p in corpo), "recuo de 1,25 cm na primeira linha")
    notas = [p for p in d.paragraphs
             if p.text.startswith("Nota:") or p.text.startswith("Fonte:")]
    conferir(notas and all(abs((p.paragraph_format.line_spacing or 0) - 1.0) < 0.01
                           for p in notas),
             "espaço simples nas notas e fontes", f"{len(notas)} itens")

    print("\n── estilo pedido ──")
    conferir("—" not in texto, "sem travessão")
    conferir(not re.findall(r"(?<!\d)–(?!\d)", texto), "sem traço de meia risca")
    g = achar_gerundios(texto)
    conferir(not g, "sem gerúndio", f"{len(g)}: {g[:5]}")

    print("\n── conteúdo ──")
    conferir(len(d.tables) == len(conteudo.TABELAS),
             f"{len(conteudo.TABELAS)} tabelas presentes", str(len(d.tables)))
    for chave, t in conteudo.TABELAS.items():
        marca = f"Tabela {t['numero']} - {t['titulo']}"
        conferir(marca in texto, f"Tabela {t['numero']} com título e numeração")
    esperadas = sum(1 for s_ in conteudo.SECOES if "figura" in s_)
    z = zipfile.ZipFile(a.arquivo)
    imgs = sum(1 for n in z.namelist() if n.startswith("word/media/"))
    conferir(imgs == esperadas, f"{esperadas} figuras embutidas", str(imgs))
    for s_ in conteudo.SECOES:
        if "legenda_figura" in s_:
            conferir(s_["legenda_figura"] in texto,
                     f"legenda presente: {s_['legenda_figura'][:34]}...")
    conferir(all(s_["titulo"] in texto for s_ in conteudo.SECOES),
             f"{len(conteudo.SECOES)} seções presentes")
    conferir(texto.count(conteudo.FONTE_TABELA) == len(conteudo.TABELAS),
             "fonte declarada sob cada tabela")
    conferir(texto.count(conteudo.FONTE_FIGURA) == esperadas,
             "fonte declarada sob cada figura")

    print("\n── análise estatística explicada ──")
    for termo in ("modelo linear misto", "pseudorreplicação", "bootstrap",
                  "Benjamini-Hochberg", "PERMANOVA", "fatores de Bayes",
                  "efeito de desenho", "d de Cohen", "efeito piso",
                  "diferença em diferenças", "E-value"):
        conferir(termo.lower() in texto.lower(), f"explica: {termo}")

    if falhas:
        print(f"\nFALHOU: {len(falhas)} de {len(falhas) + aprovados}")
        for f_ in falhas:
            print(f"   · {f_}")
        return 1
    print(f"\nOK: {aprovados} verificações passaram")
    return 0


if __name__ == "__main__":
    sys.exit(main())
