#!/usr/bin/env python3
"""Gera o manuscrito corrigido em ABNT a partir da biblioteca e do original.

    python3 scripts/gerar_manuscrito.py \
        --original ESTUDO_RS_HANDEBOL_ABNT.docx \
        --db data/BIBLIOTECA_HANDEBOL.sqlite \
        --saida data/ESTUDO_RS_HANDEBOL_CORRIGIDO.docx

Formato conforme NBR 14724: A4 retrato, margens 3/3/2/2 cm, Times New Roman 12,
entrelinha 1,5 no corpo e simples nas referências, recuo de 1,25 cm na primeira
linha, títulos com estilos de título para permitir sumário automático.
Referências conforme NBR 6023, com o título do periódico em itálico.
"""
from __future__ import annotations

import argparse
import sqlite3
import sys
from pathlib import Path

import docx
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

sys.path.insert(0, str(Path(__file__).resolve().parent))

from curadoria import extracao, referencias, tabelas
from curadoria.elegibilidade import triar
from manuscrito import correcoes, secoes

FONTE = "Times New Roman"
CORPO_PT = 12
NOTA_PT = 10


# ── Formatação base ─────────────────────────────────────────────────────────
def preparar(doc: docx.Document) -> None:
    s = doc.sections[0]
    s.orientation = WD_ORIENT.PORTRAIT
    s.page_width, s.page_height = Cm(21.0), Cm(29.7)      # A4 retrato (A1)
    s.left_margin, s.top_margin = Cm(3.0), Cm(3.0)        # NBR 14724
    s.right_margin, s.bottom_margin = Cm(2.0), Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = FONTE
    normal.font.size = Pt(CORPO_PT)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONTE)

    for nivel, tamanho, negrito, caixa in ((1, CORPO_PT, True, True),
                                           (2, CORPO_PT, True, False),
                                           (3, CORPO_PT, False, False)):
        st = doc.styles[f"Heading {nivel}"]
        st.font.name = FONTE
        st.font.size = Pt(tamanho)
        st.font.bold = negrito
        st.font.all_caps = caixa
        st.font.color.rgb = None
        st.element.rPr.rFonts.set(qn("w:eastAsia"), FONTE)
        pf = st.paragraph_format
        pf.space_before, pf.space_after = Pt(18), Pt(12)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.keep_with_next = True


def corpo(doc, texto: str, *, recuo=True, justificar=True, tamanho=CORPO_PT,
          espaco=1.5, italico=False, negrito=False, centro=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = espaco
    pf.space_after = Pt(0 if espaco == 1.5 else 6)
    pf.first_line_indent = Cm(1.25) if recuo else Cm(0)
    p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if centro else
                   WD_ALIGN_PARAGRAPH.JUSTIFY if justificar else
                   WD_ALIGN_PARAGRAPH.LEFT)
    r = p.add_run(texto)
    r.font.name, r.font.size = FONTE, Pt(tamanho)
    r.italic, r.bold = italico, negrito
    return p


def marcador(doc, texto: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(texto)
    r.font.name, r.font.size = FONTE, Pt(CORPO_PT)
    return p


def titulo(doc, texto: str, nivel: int):
    return doc.add_heading(texto, level=nivel)


def inserir_tabela(doc, t: tabelas.Tabela) -> None:
    corpo(doc, f"Tabela {t.numero} – {t.titulo}", recuo=False, justificar=False,
          tamanho=NOTA_PT, espaco=1.0)
    tab = doc.add_table(rows=1, cols=len(t.cabecalho))
    tab.style = "Table Grid"
    tab.autofit = True
    for i, h in enumerate(t.cabecalho):
        cel = tab.rows[0].cells[i]
        cel.text = ""
        r = cel.paragraphs[0].add_run(h)
        r.bold = True
        r.font.name, r.font.size = FONTE, Pt(NOTA_PT)
    for linha in t.linhas:
        cels = tab.add_row().cells
        for i, valor in enumerate(linha):
            texto = str(valor)
            forte = texto.startswith("**")
            cels[i].text = ""
            r = cels[i].paragraphs[0].add_run(texto.replace("**", ""))
            r.bold = forte
            r.font.name, r.font.size = FONTE, Pt(NOTA_PT)
    corpo(doc, f"Fonte: {t.fonte}", recuo=False, justificar=False,
          tamanho=NOTA_PT, espaco=1.0)
    if t.nota:
        corpo(doc, f"Nota: {t.nota}", recuo=False, tamanho=NOTA_PT, espaco=1.0)
    corpo(doc, "", recuo=False, espaco=1.0)


def inserir_referencia(doc, ref: referencias.Referencia) -> None:
    """NBR 6023: alinhada à esquerda, espaço simples, periódico em itálico (A3)."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.0
    pf.space_after = Pt(12)
    pf.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for texto, italico in ((ref.antes, False), (ref.periodico, True),
                           (ref.depois, False)):
        r = p.add_run(texto)
        r.font.name, r.font.size = FONTE, Pt(CORPO_PT)
        r.italic = italico


# ── Leitura do original ─────────────────────────────────────────────────────
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _texto(el) -> str:
    return " ".join("".join(t.text or "" for t in el.iter(f"{W}t")).split())


def blocos_do_original(caminho: Path) -> list[tuple[str, object]]:
    """Parágrafos e tabelas do original, na ordem em que aparecem.

    Ler só os parágrafos descartaria os Quadros 1 a 6 e as Tabelas 1 e 2, que
    são conteúdo do método e permanecem válidos.
    """
    import zipfile
    from xml.etree import ElementTree as ET
    z = zipfile.ZipFile(caminho)
    raiz = ET.fromstring(z.read("word/document.xml"))
    saida: list[tuple[str, object]] = []
    for el in raiz.find(f"{W}body"):
        if el.tag == f"{W}p":
            saida.append(("p", _texto(el)))
        elif el.tag == f"{W}tbl":
            linhas = [[_texto(tc) for tc in tr.findall(f"{W}tc")]
                      for tr in el.findall(f"{W}tr")]
            saida.append(("tbl", linhas))
    return saida


def paragrafos_do_original(caminho: Path) -> list[str]:
    return [v for k, v in blocos_do_original(caminho) if k == "p"]


def inserir_tabela_bruta(doc, linhas: list[list[str]]) -> None:
    """Reproduz um Quadro ou Tabela do original, com a mesma formatação das
    tabelas geradas."""
    if not linhas:
        return
    ncols = max(len(l) for l in linhas)
    tab = doc.add_table(rows=0, cols=ncols)
    tab.style = "Table Grid"
    for i, linha in enumerate(linhas):
        cels = tab.add_row().cells
        for j in range(ncols):
            cels[j].text = ""
            r = cels[j].paragraphs[0].add_run(linha[j] if j < len(linha) else "")
            r.bold = (i == 0)
            r.font.name, r.font.size = FONTE, Pt(NOTA_PT)


def aplicar_correcoes(paras: list[str]) -> list[str]:
    saida = []
    for p in paras:
        if any(p.startswith(pref) for pref in correcoes.REMOVER):
            continue
        troca = next((v for k, v in correcoes.SUBSTITUIR.items()
                      if p.startswith(k)), None)
        saida.append(troca if troca is not None else p)
    return saida


# ── Montagem ────────────────────────────────────────────────────────────────
SECOES_SUBSTITUIDAS = ("4 RESULTADOS", "5 DISCUSSÃO", "6 LIMITAÇÕES",
                       "7 ESQUELETO", "REFERÊNCIAS")


def montar(original: Path, db: Path, saida: Path) -> dict:
    con = sqlite3.connect(db)
    decisoes = triar(con)
    linhas_extracao = extracao.montar(con, decisoes)
    tabs = {t.numero: t for t in tabelas.todas(con, decisoes)}

    diag_alertas = extracao.diagnostico(linhas_extracao)["com_alerta"]
    doc = docx.Document()
    preparar(doc)
    paras = aplicar_correcoes(paragrafos_do_original(original))

    # ── Folha de rosto e pré-textuais (A6, A8) ──
    for i, linha in enumerate(paras[:4]):
        corpo(doc, linha, recuo=False, justificar=False, centro=True,
              negrito=(i == 0), espaco=1.5)
    doc.add_page_break()

    titulo(doc, "RESUMO", 1)
    corpo(doc, "[A PREENCHER: Resumo] Redigir resumo estruturado de até 250 "
               "palavras — objetivo, método, resultados principais e conclusão —, "
               "conforme a NBR 6028, em parágrafo único e espaço simples. "
               "Depende de: conclusão da triagem.", espaco=1.0)
    corpo(doc, "Palavras-chave: handebol; psicologia do esporte; ansiedade "
               "competitiva; motivação; revisão de escopo.", recuo=False, espaco=1.0)
    corpo(doc, "", recuo=False)
    titulo(doc, "ABSTRACT", 1)
    corpo(doc, "[A PREENCHER: Abstract] English version of the structured "
               "abstract. Depends on: completion of screening.", espaco=1.0)
    corpo(doc, "Keywords: handball; sport psychology; competitive anxiety; "
               "motivation; scoping review.", recuo=False, espaco=1.0)
    doc.add_page_break()

    titulo(doc, "SUMÁRIO", 1)
    corpo(doc, "[Inserir sumário automático no Word: Referências › Sumário. Os "
               "títulos deste documento usam os estilos de título nativos, de "
               "modo que o sumário se gera e se atualiza sozinho.]",
          recuo=False, espaco=1.0, italico=True)
    doc.add_page_break()

    # ── Corpo: introdução até 3.12, do original corrigido ──
    inseridas = set()
    blocos_orig = blocos_do_original(original)
    # Aplica as correções mantendo as tabelas em posição.
    corrigidos: list[tuple[str, object]] = []
    for tipo, valor in blocos_orig:
        if tipo == "tbl":
            corrigidos.append((tipo, valor))
            continue
        if any(valor.startswith(pref) for pref in correcoes.REMOVER):
            continue
        if valor in correcoes.TITULOS:
            corrigidos.append(("p", correcoes.TITULOS[valor]))
            continue
        troca = next((v for k, v in correcoes.SUBSTITUIR.items()
                      if valor.startswith(k)), None)
        corrigidos.append(("p", troca if troca is not None else valor))

    vistos_p = 0
    i = 0
    while i < len(corrigidos):
        tipo, valor = corrigidos[i]
        if tipo == "tbl":
            # A Tabela 1 do original divergia da procedência real dos registros
            # (achado G1); entra a versão gerada dos dados.
            if valor and valor[0][:1] == ["Base"]:
                inserir_tabela(doc, tabs[1])
                inseridas.add(1)
            else:
                inserir_tabela_bruta(doc, valor)
            i += 1
            continue
        p = valor
        vistos_p += 1
        if vistos_p <= 4:      # folha de rosto, já emitida
            i += 1
            continue
        if any(p.startswith(s) for s in SECOES_SUBSTITUIDAS):
            break
        if not p:
            i += 1
            continue
        if p.startswith("Quadro ") or p.startswith("Tabela ") or p.startswith("Figura "):
            corpo(doc, p, recuo=False, justificar=False, tamanho=NOTA_PT, espaco=1.0)
        elif p.startswith("Fonte:") or p.startswith("Nota:"):
            corpo(doc, p, recuo=False, tamanho=NOTA_PT, espaco=1.0)
        elif p.startswith("•"):
            marcador(doc, p.lstrip("• "))
        elif p[0].isdigit() and " " in p and p.split()[0].replace(".", "").isdigit():
            nivel = 1 if "." not in p.split()[0] else 2
            titulo(doc, p, nivel)
        else:
            corpo(doc, p)
        i += 1

    # ── Seções geradas dos dados ──
    blocos = (secoes.resultados(con, decisoes, linhas_extracao)
              + secoes.discussao(con, decisoes)
              + secoes.limitacoes(decisoes, con)      # 6
              + secoes.conclusao(decisoes))           # 7
    for estilo, texto in blocos:
        if estilo == "Título 1":
            titulo(doc, texto, 1)
        elif estilo == "Título 2":
            titulo(doc, texto, 2)
        elif estilo == "Marcador":
            marcador(doc, texto)
        elif estilo == "Tabela":
            inserir_tabela(doc, tabs[int(texto)])
            inseridas.add(int(texto))
        else:
            corpo(doc, texto)

    # ── Divisão de tarefas ──
    titulo(doc, "8 DIVISÃO DE TAREFAS", 1)
    corpo(doc, "A triagem independente é requisito metodológico: a concordância "
               "entre revisores só tem significado se as decisões forem tomadas "
               "sem conhecimento mútuo. O quadro abaixo lista as tarefas "
               "atribuídas ao bolsista, na ordem de execução.")
    quadro_tarefas = next((v for k, v in blocos_orig if k == "tbl"
                           and v and v[0][:1] == ["Etapa"] and len(v[0]) == 3), None)
    if quadro_tarefas:
        corpo(doc, "Quadro 5 – Tarefas do bolsista de iniciação científica",
              recuo=False, justificar=False, tamanho=NOTA_PT, espaco=1.0)
        inserir_tabela_bruta(doc, quadro_tarefas)
        corpo(doc, "Fonte: Elaborada pelos autores.", recuo=False,
              justificar=False, tamanho=NOTA_PT, espaco=1.0)
    corpo(doc, f"Acrescenta-se uma tarefa que a versão anterior não previa: "
               f"conferir contra o texto completo as {diag_alertas} linhas do "
               "Apêndice B assinaladas com alerta. São os campos que a extração "
               "automatizada não conseguiu estabelecer com segurança, e a "
               "conferência deles antecede qualquer tabulação de amostra, idade "
               "ou delineamento.")

    # ── Referências (A2, A3, A4) ──
    doc.add_page_break()
    titulo(doc, "REFERÊNCIAS", 1)
    citadas = citadas_no_texto(paras)
    refs = referencias_do_original(original, citadas)
    for ref in refs:
        inserir_referencia(doc, ref)

    # ── Apêndices ──
    doc.add_page_break()
    titulo(doc, "APÊNDICE A – ESTRATÉGIAS DE BUSCA NA ÍNTEGRA", 1)
    corpo(doc, "As consultas abaixo são geradas dos blocos de termos da Tabela 8 "
               "e são exatamente as submetidas a cada base, de modo que apêndice "
               "e busca não podem divergir. Reproduza-as com "
               "scripts/atualizar_buscas.py --imprimir-consultas.")
    from busca import estrategia
    for nome, fn in estrategia.CONSULTAS.items():
        q = fn()
        corpo(doc, f"{nome.upper()} ({len(q)} caracteres)", recuo=False,
              negrito=True, tamanho=NOTA_PT, espaco=1.0)
        corpo(doc, q, recuo=False, justificar=False, tamanho=NOTA_PT, espaco=1.0)
    inserir_tabela(doc, tabs[8])

    doc.add_page_break()
    titulo(doc, "APÊNDICE B – TABELA DE EXTRAÇÃO", 1)
    diag = extracao.diagnostico(linhas_extracao)
    corpo(doc, f"A tabela abaixo cobre os {diag['linhas']} registros elegíveis e "
               "é gerada da biblioteca, não digitada. Campos ausentes aparecem "
               "como n.d.; campos internamente incoerentes aparecem como “a "
               "conferir”, com o motivo na coluna de alertas. "
               f"{diag['com_alerta']} linhas trazem ao menos um alerta, e é essa "
               "a lista de conferência do bolsista.")
    t9 = tabelas.Tabela(9, f"Extração dos {diag['linhas']} registros elegíveis",
                        extracao.CABECALHO,
                        [l.como_lista() for l in linhas_extracao],
                        nota="Gerada de data/BIBLIOTECA_HANDEBOL.sqlite por "
                             "scripts/gerar_manuscrito.py.")
    inserir_tabela(doc, t9)

    saida.parent.mkdir(parents=True, exist_ok=True)
    doc.save(saida)
    _corrigir_settings(saida)
    con.close()
    return {"paragrafos": len(paras), "tabelas": sorted(inseridas | {8, 9}),
            "referencias": len(refs), "linhas_extracao": diag["linhas"]}


# Sobrenomes em versalete carregam diacríticos fora da faixa A-Ý — Ś de Świdwa,
# Ø de Jørgensen, Ł de Łukasz. A classe é definida por propriedade Unicode.
PALAVRA = __import__("re").compile(r"[^\W\d_][\w'’\-]{2,}", __import__("re").UNICODE)


def _sobrenomes(texto: str) -> set[str]:
    return {m.group(0) for m in PALAVRA.finditer(texto)
            if m.group(0) == m.group(0).upper()}


def _corrigir_settings(caminho: Path) -> None:
    """O template do python-docx grava <w:zoom> sem o atributo w:percent, que o
    esquema OOXML exige. Word tolera; validadores não. Remove-se o elemento.
    """
    import re
    import shutil
    import zipfile
    temp = caminho.with_suffix(".tmp.docx")
    with zipfile.ZipFile(caminho) as orig, \
            zipfile.ZipFile(temp, "w", zipfile.ZIP_DEFLATED) as novo:
        for item in orig.infolist():
            dados = orig.read(item.filename)
            if item.filename == "word/settings.xml":
                texto = dados.decode("utf-8")
                if "w:percent" not in texto:
                    texto = re.sub(r"<w:zoom\b[^>]*/>", "", texto)
                    texto = re.sub(r"<w:zoom\b[^>]*>.*?</w:zoom>", "", texto,
                                   flags=re.S)
                dados = texto.encode("utf-8")
            novo.writestr(item, dados)
    shutil.move(temp, caminho)


def citadas_no_texto(paras: list[str]) -> set[str]:
    corte = next((i for i, p in enumerate(paras) if p.startswith("REFERÊNCIAS")),
                 len(paras))
    return _sobrenomes("\n".join(paras[:corte]))


def referencias_do_original(caminho: Path, citadas: set[str]) -> list:
    """Reformata as referências do original pela NBR 6023 e descarta as não
    citadas (achado A4)."""
    import re
    paras = paragrafos_do_original(caminho)
    inicio = next(i for i, p in enumerate(paras) if p.startswith("REFERÊNCIAS"))
    fim = next(i for i, p in enumerate(paras) if p.startswith("APÊNDICE A"))
    saida = []
    for linha in paras[inicio:fim]:
        if " DOI: " not in linha:
            continue
        sobrenome = PALAVRA.match(linha)
        if not sobrenome or sobrenome.group(0) not in citadas:
            continue
        saida.append(_referencia_de_linha(linha))
    return sorted(saida, key=lambda r: r.antes)


def _referencia_de_linha(linha: str) -> referencias.Referencia:
    """Converte a referência já redigida do original, corrigindo o ponto duplo
    (A2) e isolando o periódico para o itálico (A3)."""
    import re
    linha = re.sub(r"\.\.(\s)", r".\1", linha.strip())
    m = re.match(r"^(.*?\.\s)([^.]+?)(,\s*(?:v\.|n\.|p\.|\d{4}).*)$", linha)
    if not m:
        return referencias.Referencia(linha, "", "")
    return referencias.Referencia(m.group(1), m.group(2).strip(), m.group(3))


def main() -> int:
    p = argparse.ArgumentParser(description=__doc__,
                                formatter_class=argparse.RawDescriptionHelpFormatter)
    p.add_argument("--original", type=Path, required=True)
    p.add_argument("--db", type=Path, default=Path("data/BIBLIOTECA_HANDEBOL.sqlite"))
    p.add_argument("--saida", type=Path,
                   default=Path("data/ESTUDO_RS_HANDEBOL_CORRIGIDO.docx"))
    a = p.parse_args()
    info = montar(a.original, a.db, a.saida)
    print(f"gerado: {a.saida}")
    for k, v in info.items():
        print(f"  {k}: {v}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
